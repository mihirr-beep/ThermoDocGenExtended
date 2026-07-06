"""One-time build tool: fix image/caption layout in the .docx templates.

The source datasheets (and the first generated templates) place the caption
ABOVE the image and carry "< Plot Size ... >" size hints. The lab wants the
image ON TOP with its caption/name directly BELOW and NO spacing around them.

This rewrites every template in word_templates/ in place:
  * for each caption paragraph ("Photo N:" / "Figure N:") immediately followed
    by an image placeholder paragraph ("{{ img_... }}"), the caption is moved to
    sit AFTER the image;
  * "< ... >" size-hint paragraphs are removed;
  * the image + caption paragraphs get zero space-before/after so they read as
    one tight block.

Idempotent: running twice is a no-op (a caption already below its image stays).

    python -m datasheet_gen.fix_template_layout          # fix all templates
    python -m datasheet_gen.fix_template_layout CE ESD   # fix specific codes
"""
import os
import re
import sys

import docx
from docx.shared import Pt

TPL_DIR = os.path.join(os.path.dirname(__file__), "word_templates")

CAPTION_RE = re.compile(r"^\s*(Photo|Figure)\s*\d*\s*:", re.I)
HINT_RE = re.compile(r"^\s*<.*>\s*$")
# a paragraph whose whole content is a single {{ var }} referencing an image
IMG_PLACEHOLDER_RE = re.compile(
    r"^\s*\{\{\s*(img_\w+|[\w.]*plot_\w+|photo_\w+|signature|\w*figure\w*|\w*diagram\w*)\s*\}\}\s*$",
    re.I,
)


def _text(p):
    return "".join(r.text or "" for r in p.runs) or p.text or ""


def _zero_spacing(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _mark_loop_table_headers(doc):
    """Any template table that contains a `{%tr for ... %}` loop row produces a
    dynamic number of data rows: mark its first row as a repeating header
    (w:tblHeader) so multi-page tables keep their header. Idempotent."""
    from docx.oxml.ns import qn
    n = 0
    for tbl in doc.tables:
        texts = " ".join(c.text for r in tbl.rows for c in r.cells)
        if "{%tr for" not in texts:
            continue
        tr = tbl.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        if not trPr.findall(qn("w:tblHeader")):
            trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))
            n += 1
    return n


def fix_doc(path):
    doc = docx.Document(path)
    paras = doc.paragraphs
    changed = 0
    changed += _mark_loop_table_headers(doc)

    # 1) drop "< size hint >" paragraphs
    for p in list(paras):
        if HINT_RE.match(_text(p)):
            p._p.getparent().remove(p._p)
            changed += 1

    # 2) re-order caption/image pairs so the image sits on top
    paras = doc.paragraphs  # refresh after removals
    for i, p in enumerate(paras[:-1]):
        if not CAPTION_RE.match(_text(p)):
            continue
        nxt = paras[i + 1]
        if IMG_PLACEHOLDER_RE.match(_text(nxt)):
            # move the caption to directly AFTER the image placeholder
            nxt._p.addnext(p._p)
            _zero_spacing(nxt)   # image
            _zero_spacing(p)     # caption (now below)
            changed += 1

    doc.save(path)
    return changed


def _iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def fix_checkbox_placeholders():
    """Convert `{{ key }}` -> `{{r key }}` for every checkbox-rendered field so
    docxtpl inserts the human-ticked checkbox runs instead of escaped text.

    Covers the CE template's two classification fields plus every schema field
    that declares a "checkbox" option list. Idempotent.
    """
    import json

    targets = {"IEC-FRM-504_CE.docx": ["classification_group", "classification_class"]}
    schema_dir = os.path.join(os.path.dirname(__file__), "schemas")
    for fn in sorted(os.listdir(schema_dir)):
        if not fn.endswith(".json"):
            continue
        schema = json.load(open(os.path.join(schema_dir, fn), encoding="utf-8"))
        keys = []
        for sec in schema.get("sections", []):
            for it in sec.get("items", []):
                fields = it.get("fields", []) if it.get("type") == "fields" else (
                    [it] if it.get("type") in ("field", "textarea") else [])
                keys += [f["key"] for f in fields if f.get("checkbox")]
        if keys:
            targets[fn.replace(".json", ".docx")] = keys

    for tpl_name, keys in targets.items():
        path = os.path.join(TPL_DIR, tpl_name)
        if not os.path.exists(path):
            print(f"SKIP (missing): {tpl_name}")
            continue
        doc = docx.Document(path)
        n = 0
        for p in _iter_all_paragraphs(doc):
            for run in p.runs:
                for key in keys:
                    plain = "{{ " + key + " }}"
                    if plain in (run.text or ""):
                        run.text = run.text.replace(plain, "{{r " + key + " }}")
                        n += 1
        if n:
            doc.save(path)
        print(f"OK {tpl_name:24s} - {n} checkbox placeholder(s) -> {{{{r }}}}")


def remove_stray_deviation_na():
    """CE template: drop the literal 'NA' paragraph the source doc left behind
    after DEVIATION FROM THE STANDARD (the {{ deviation }} placeholder already
    carries the value). Idempotent."""
    path = os.path.join(TPL_DIR, "IEC-FRM-504_CE.docx")
    doc = docx.Document(path)
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if _text(p).strip() == "{{ deviation }}":
            j = i + 1
            while j < len(paras) and not _text(paras[j]).strip():
                j += 1
            if j < len(paras) and _text(paras[j]).strip().upper() == "NA":
                paras[j]._p.getparent().remove(paras[j]._p)
                doc.save(path)
                print("OK IEC-FRM-504_CE.docx      - removed stray 'NA' after {{ deviation }}")
                return
    print("OK IEC-FRM-504_CE.docx      - no stray 'NA' (already clean)")


def main(codes=None):
    files = ([os.path.join(TPL_DIR, f"{c}.docx") for c in codes]
             if codes else
             [os.path.join(TPL_DIR, f) for f in sorted(os.listdir(TPL_DIR)) if f.endswith(".docx")])
    for path in files:
        if not os.path.exists(path):
            print(f"SKIP (missing): {os.path.basename(path)}")
            continue
        n = fix_doc(path)
        print(f"OK {os.path.basename(path):24s} - {n} caption/hint fix(es)")


if __name__ == "__main__":
    main(sys.argv[1:] or None)
