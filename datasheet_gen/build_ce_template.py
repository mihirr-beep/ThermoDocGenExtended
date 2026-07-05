"""Build-time tool: templatize the official IEC-FRM-504 CE datasheet into a
docxtpl template covering EVERY section of the document, with image anchors.

Source : D:/THERMO/DocGenerator/OneDrive_1_22-6-2026/IEC-FRM-504  CE Test Data sheet.docx
Output : datasheet_gen/word_templates/IEC-FRM-504_CE.docx

Run from project root:  python datasheet_gen/build_ce_template.py
"""
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DEFAULT_SRC = os.path.join(
    "D:/THERMO/DocGenerator/OneDrive_1_22-6-2026",
    "IEC-FRM-504  CE Test Data sheet.docx",
)
OUT_DIR = os.path.join(os.path.dirname(__file__), "word_templates")
OUT = os.path.join(OUT_DIR, "IEC-FRM-504_CE.docx")


def set_cell(cell, value):
    cell.text = ""
    cell.paragraphs[0].add_run(value)


def find_para(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def remove_para(p):
    p._p.getparent().remove(p._p)


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    para = Paragraph(new_p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def clear_rows_after(table, keep_index):
    for row in list(table.rows[keep_index + 1:]):
        row._tr.getparent().remove(row._tr)


def add_tag_row(table, tag_text, before=None, after=None):
    row = table.add_row()
    set_cell(row.cells[0], tag_text)
    tr = row._tr
    tr.getparent().remove(tr)
    if before is not None:
        before._tr.addprevious(tr)
    else:
        after._tr.addnext(tr)


def wrap_loop(table, data_index, loop_var):
    data = table.rows[data_index]
    add_tag_row(table, "{%tr for " + loop_var + " %}", before=data)
    add_tag_row(table, "{%tr endfor %}", after=data)


def insert_measure_table_after(doc, paragraph, loop_var, item):
    header = ["Frequency (MHz)", "Q-peak (dBµV)", "Limit (dBµV)", "Margin (dB)",
              "Frequency (MHz)", "Average (dBµV)", "Limit (dBµV)", "Margin (dB)"]
    cols = ["qp_freq", "qp", "qp_limit", "qp_margin", "avg_freq", "avg", "avg_limit", "avg_margin"]
    tbl = doc.add_table(rows=4, cols=len(header))
    tbl.style = "Table Grid"
    for j, h in enumerate(header):
        set_cell(tbl.rows[0].cells[j], h)
    set_cell(tbl.rows[1].cells[0], "{%tr for " + loop_var + " %}")
    for j, c in enumerate(cols):
        set_cell(tbl.rows[2].cells[j], "{{ " + item + "." + c + " }}")
    set_cell(tbl.rows[3].cells[0], "{%tr endfor %}")
    paragraph._p.addnext(tbl._tbl)
    return tbl


SPEC_MAP = {
    "Product Standard": ("{{ product_standard }}", ""),
    "Basic Standard": ("{{ basic_standard }}", ""),
    # {{r }}: rendered as RichText (human-ticked checkboxes) by the generator
    "Classification": ("{{r classification_group }}", "{{r classification_class }}"),
    "Test Port": ("{{ test_port }}", ""),
    "Coupling Method": ("{{ coupling_method }}", ""),
    "Frequency Range": ("{{ frequency_range }}", ""),
    "Resolution Bandwidth": ("{{ resolution_bandwidth }}", ""),
    "Step size": ("{{ step_size }}", ""),
    "Detector": ("{{ detector }}", ""),
    "Measurement time": ("{{ measurement_time }}", ""),
    "Test Mode": ("{{ test_mode }}", ""),
    "EUT Modification state": ("{{ eut_modification_state }}", ""),
    # Checkboxes (Tabletop | Floor standing), one per cell, ticked via RichText
    "EUT Configuration": ("{{r eut_config_tabletop }}", "{{r eut_config_floor }}"),
    "EUT Input Voltage": ("{{ eut_voltage_frequency }}", ""),
    "Ambient Temperature": ("{{ ambient_temperature }}", ""),
    "Relative Humidity": ("{{ relative_humidity }}", ""),
    "Test Date": ("{{ test_date }}", ""),
    "Tested by": ("{{ tested_by }}", ""),
}


def main(src):
    doc = Document(src)
    t = doc.tables

    # Table 0 - EUT details
    eut = {"Job Number": "{{ job_number }}", "EUT Name": "{{ eut_name }}",
           "EUT Model/SKU Number": "{{ eut_model }}", "EUT Serial Number": "{{ eut_serial }}"}
    for row in t[0].rows:
        if row.cells[0].text.strip() in eut:
            set_cell(row.cells[1], eut[row.cells[0].text.strip()])

    # Table 1 - Modification record (loop)
    mod = t[1]
    for j, c in enumerate(["state", "description", "fitted_by", "date"]):
        set_cell(mod.rows[1].cells[j], "{{ m." + c + " }}")
    clear_rows_after(mod, 1)
    wrap_loop(mod, 1, "m in modifications")

    # Table 2 - Measurement uncertainty
    set_cell(t[2].rows[1].cells[1], "{{ measurement_uncertainty }}")

    # Functional check paragraph
    fc = find_para(doc, "Functional check is conducted")
    if fc:
        fc.text = "Functional check is conducted as per SOP reference number: {{ sop_reference }}."

    # Table 3 - Test specification.
    # Most value rows are a single cell spanning columns 1+2 (a horizontal merge);
    # python-docx returns that one merged cell as BOTH cells[1] and cells[2]. We must
    # only write the 2nd value column when it is a DISTINCT cell, otherwise writing c2
    # overwrites (wipes) the c1 placeholder we just set in the merged cell.
    for row in t[3].rows:
        label = row.cells[0].text.strip()
        match = next((k for k in SPEC_MAP if label.startswith(k)), None)
        if match is None:
            continue
        c1, c2 = SPEC_MAP[match]
        set_cell(row.cells[1], c1)
        if len(row.cells) >= 3 and row.cells[2]._tc is not row.cells[1]._tc:
            set_cell(row.cells[2], c2)

    # Deviation paragraph ("NA")
    dev_head = find_para(doc, "DEVIATION FROM THE STANDARD")
    if dev_head:
        na = insert_paragraph_after(dev_head, "{{ deviation }}")
        # remove the original "NA" paragraph that followed (it may be separated
        # from the placeholder by empty spacer paragraphs — skip those)
        nxt = na._p.getnext()
        while nxt is not None and nxt.tag.endswith('}p'):
            np = Paragraph(nxt, na._parent)
            txt = np.text.strip()
            if not txt:
                nxt = nxt.getnext()
                continue
            if txt.upper() == "NA":
                remove_para(np)
            break

    # Table 4 - Test limits (3 bands per the DS504 sheet: 0.15-0.50, 0.50-5, 5-30).
    # The source doc only ships two rows (0.15-0.50 and 0.50-30); we relabel the
    # second to 0.50-5 and clone it to add the 5-30 band so regeneration stays in
    # sync with the form/service (which prefill all three bands from the class).
    import copy
    from docx.table import _Row
    limits_tbl = t[4]
    row_050 = None
    for row in limits_tbl.rows:
        lbl = row.cells[0].text.strip()
        if lbl.startswith("0.15"):
            set_cell(row.cells[1], "{{ limit_qp_015_050 }}")
            set_cell(row.cells[2], "{{ limit_avg_015_050 }}")
        elif lbl.startswith("0.50"):
            set_cell(row.cells[0], "0.50 to 5")
            set_cell(row.cells[1], "{{ limit_qp_050_5 }}")
            set_cell(row.cells[2], "{{ limit_avg_050_5 }}")
            row_050 = row
    if row_050 is not None:
        new_tr = copy.deepcopy(row_050._tr)
        row_050._tr.addnext(new_tr)
        nr = _Row(new_tr, limits_tbl)
        set_cell(nr.cells[0], "5 to 30")
        set_cell(nr.cells[1], "{{ limit_qp_5_30 }}")
        set_cell(nr.cells[2], "{{ limit_avg_5_30 }}")

    # Test procedure: make the first procedure line editable, drop the boilerplate
    proc = find_para(doc, "The test procedure was in accordance")
    if proc:
        proc.text = "{{r test_procedure }}"   # RichText: bold "LISN (Voltage Method):" header
        for prefix in ("The EUT was placed", "LISN (Voltage Method)",
                       "The conducted emission was measured"):
            p = find_para(doc, prefix)
            if p:
                remove_para(p)

    # Measurement data: images + Line/Neutral tables after captions
    fig1, tab1 = find_para(doc, "Figure 1:"), find_para(doc, "Table 1:")
    fig2, tab2 = find_para(doc, "Figure 2:"), find_para(doc, "Table 2:")
    photo1 = find_para(doc, "Photo 1:")
    if fig1:
        insert_paragraph_after(fig1, "{{ plot_line }}")
    if tab1:
        insert_measure_table_after(doc, tab1, "r in line_rows", "r")
    if fig2:
        insert_paragraph_after(fig2, "{{ plot_neutral }}")
    if tab2:
        insert_measure_table_after(doc, tab2, "r in neutral_rows", "r")
    if photo1:
        insert_paragraph_after(photo1, "{{ photo_setup }}")

    # Table 5 - Equipment (loop)
    eq = t[5]
    for j, c in enumerate(["name", "make", "model", "serial", "cal_due"]):
        set_cell(eq.rows[1].cells[j], "{{ e." + c + " }}")
    clear_rows_after(eq, 1)
    wrap_loop(eq, 1, "e in equipment")

    # Table 6 - Software
    set_cell(t[6].rows[1].cells[0], "{{ software_used }}")
    set_cell(t[6].rows[1].cells[1], "{{ software_version }}")

    # Result paragraph + sign-off
    result = find_para(doc, "Conducted Emissions from the EUT as per Class")
    if result:
        result.text = ("Conducted Emissions from the EUT as per "
                       "{{r result_class_label }} limit: {{r result_checkbox }}")
    signoff = {"Name": "{{ tested_by_name }}", "Signature": "{{ signature }}", "Date": "{{ tested_by_date }}"}
    for row in t[7].rows:
        if row.cells[0].text.strip() in signoff:
            set_cell(row.cells[1], signoff[row.cells[0].text.strip()])

    # Remove the source form's MANUAL page breaks. They were placed assuming fixed
    # content; with our dynamic rows + fitted images the content reaches the page
    # boundary and the forced break then creates a stray blank page. Letting Word
    # paginate naturally avoids the blank page (no manual backspace needed).
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    removed = 0
    for br in list(doc.element.body.iter(W + "br")):
        if br.get(W + "type") == "page":
            br.getparent().remove(br)
            removed += 1

    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT)
    print(f"Saved template -> {OUT}  (removed {removed} manual page break(s))")

    # Put each measurement image ON TOP with its caption directly BELOW (no
    # spacing). build_ce_template inserts the image after the caption, so reuse
    # the shared layout fixer to reorder + tighten. Keeps rebuilds consistent.
    from fix_template_layout import fix_doc
    fix_doc(OUT)
    print("Applied image-above-caption layout fix.")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_SRC
    if not os.path.exists(src):
        raise SystemExit(f"Source not found: {src}")
    main(src)
