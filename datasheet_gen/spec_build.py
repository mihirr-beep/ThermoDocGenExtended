"""Generic templatizer: convert any IEC-FRM-50x datasheet .docx into a docxtpl
template + a JSON schema describing its form (sections, fields, tables, images).

Run from project root:  python datasheet_gen/spec_build.py
Produces, for every non-CE test in the registry:
    datasheet_gen/word_templates/<CODE>.docx
    datasheet_gen/schemas/<CODE>.json
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from registry import REGISTRY, GENERIC_CODES, source_path

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
HERE = os.path.dirname(__file__)
TPL_DIR = os.path.join(HERE, "word_templates")
SCHEMA_DIR = os.path.join(HERE, "schemas")


def _t(cell):
    return cell.text.strip()


def slugify(label, used, prefix=""):
    s = re.sub(r"\(.*?\)", "", label or "")
    s = re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_").lower()
    s = (prefix + s)[:46]
    if not s or s[0].isdigit():   # Jinja identifiers can't start with a digit
        s = "f_" + s
    s = s.strip("_") or "field"
    base, i = s, 2
    while s in used:
        s = f"{base}_{i}"
        i += 1
    used.add(s)
    return s


def set_cell(cell, value):
    cell.text = ""
    cell.paragraphs[0].add_run(value)


def is_header_row(r):
    text = _t(r.cells[0]).lower()
    return any(x in text for x in [
        "test level", "name of the signal line", "s. no.", "s.no."
    ])


def is_observation_grid(table):
    if len(table.rows) < 3:
        return False
    # Check if Row 1 has column headers like +2, -2, +4
    row1_text = "".join([cell.text for cell in table.rows[1].cells]).lower()
    return any(x in row1_text for x in ["+2", "-2", "+4", "-4", "+8", "-8"])


def get_column_header(table, r_idx, c_idx):
    headers = []
    seen_cells = set()
    for r in table.rows[:r_idx]:
        if not is_header_row(r):
            continue
        cell = r.cells[c_idx]
        if cell._tc is r.cells[0]._tc:
            continue
        tc_id = id(cell._tc)
        if tc_id in seen_cells:
            continue
        seen_cells.add(tc_id)
        text = cell.text.strip()
        if text:
            headers.append(re.sub(r'\s+', ' ', text))
    return " ".join(headers)


def _classify(table):
    if is_observation_grid(table):
        return "kv"
    rows = table.rows
    if len(rows) < 2:
        return "kv"
    body0 = [_t(r.cells[0]) for r in rows[1:]]
    empty = sum(1 for x in body0 if x == "")
    numeric = sum(1 for x in body0 if x.replace(".", "").replace("-", "").isdigit())
    return "loop" if (empty + numeric) >= max(1, len(body0)) * 0.5 else "kv"


def _process_table(table, used, prefix):
    ncols = len(table.columns)
    rows = table.rows
    if _classify(table) == "kv":
        fields = []
        for r_idx, r in enumerate(rows):
            # Skip rows that act as table headers
            if is_header_row(r):
                continue
            
            # Identify unique cells in the row to handle merged cells correctly
            unique_cells = []
            seen_tc = set()
            for cell in r.cells:
                tc_id = id(cell._tc)
                if tc_id not in seen_tc:
                    seen_tc.add(tc_id)
                    unique_cells.append(cell)
            
            n_unique = len(unique_cells)
            if n_unique < 2:
                continue
            
            # Determine label and value start column
            first_cell_text = _t(r.cells[0])
            first_cell_clean = first_cell_text.strip(" .")
            
            if first_cell_clean.isdigit():
                # Serial-numbered observation row!
                second_cell_text = _t(r.cells[1])
                if second_cell_text:
                    label = second_cell_text
                else:
                    label = f"{prefix} Point {first_cell_clean}"
                val_start_col = 2
            elif ncols >= 3 and r.cells[0]._tc is r.cells[1]._tc:
                label = _t(r.cells[0])
                val_start_col = 2
            else:
                if ncols >= 4:
                    l0 = _t(r.cells[0])
                    l1 = _t(r.cells[1])
                    if not l1:
                        # Col 1 is empty in original, so it is a value column!
                        label = l0
                        val_start_col = 1
                    else:
                        # Col 1 contains sub-label
                        label = f"{l0} - {l1}"
                        val_start_col = 2
                else:
                    label = _t(r.cells[0])
                    val_start_col = 1
            
            if not label:
                continue
            
            # Find unique cells in value columns area
            val_cells = []
            seen_val_tc = set()
            for j in range(val_start_col, ncols):
                cell = r.cells[j]
                # Skip if cell is merged with label cells
                if val_start_col == 2 and (cell._tc is r.cells[0]._tc or cell._tc is r.cells[1]._tc):
                    continue
                if val_start_col == 1 and cell._tc is r.cells[0]._tc:
                    continue
                tc_id = id(cell._tc)
                if tc_id not in seen_val_tc:
                    seen_val_tc.add(tc_id)
                    val_cells.append((j, cell))
            
            base_key = slugify(label, used)
            is_sig = "signature" in label.lower()
            
            if len(val_cells) == 1:
                col_idx, cell = val_cells[0]
                key = base_key
                default_val = cell.text.strip()
                set_cell(cell, "{{ " + key + " }}")
                # Clear other cells in same row pointing to different cell in value columns area
                for j in range(val_start_col, ncols):
                    if r.cells[j]._tc is not cell._tc:
                        if val_start_col == 2 and (r.cells[j]._tc is r.cells[0]._tc or r.cells[j]._tc is r.cells[1]._tc):
                            continue
                        if val_start_col == 1 and r.cells[j]._tc is r.cells[0]._tc:
                            continue
                        set_cell(r.cells[j], "")
                fields.append({
                    "key": key,
                    "label": label,
                    "input": "image" if is_sig else "text",
                    "default": default_val
                })
            elif len(val_cells) > 1:
                # Multiple value cells (e.g. Power Line and Signal Line)
                for col_idx, cell in val_cells:
                    suffix = get_column_header(table, r_idx, col_idx)
                    if suffix:
                        key = slugify(label + " " + suffix, used)
                        display_label = f"{label} - {suffix}"
                    else:
                        key = slugify(label + f" col_{col_idx}", used)
                        display_label = f"{label} (Col {col_idx + 1})"
                    
                    default_val = cell.text.strip()
                    set_cell(cell, "{{ " + key + " }}")
                    fields.append({
                        "key": key,
                        "label": display_label,
                        "input": "image" if is_sig else "text",
                        "default": default_val
                    })
        return {"type": "fields", "fields": fields}

    # Tables with merged header/data cells can't be safely turned into docxtpl
    # row-loops (the {%tr%} tags don't relocate cleanly) -> leave them static
    # (empty, filled manually in Word) so the template never breaks.
    def _has_hmerge(row):
        ids = [id(c._tc) for c in row.cells]
        return len(set(ids)) < len(ids)

    if len(rows) >= 2 and (_has_hmerge(rows[0]) or _has_hmerge(rows[1])):
        return {"type": "static_table"}

    # loop table: row0 = header; row1 = data row -> placeholders; drop the rest
    hdr = rows[0]
    columns = [{"key": f"c{j}", "label": (_t(hdr.cells[j]) or f"Column {j + 1}")} for j in range(ncols)]
    tkey = slugify("rows", used, prefix + "_")
    drow = rows[1]
    for j in range(ncols):
        set_cell(drow.cells[j], "{{ r.c" + str(j) + " }}")
    for r in list(rows[2:]):
        r._tr.getparent().remove(r._tr)

    def add_tag(text, before=None, after=None):
        nr = table.add_row()
        set_cell(nr.cells[0], text)
        tr = nr._tr
        tr.getparent().remove(tr)
        (before._tr.addprevious if before else after._tr.addnext)(tr)

    add_tag("{%tr for r in " + tkey + " %}", before=drow)
    add_tag("{%tr endfor %}", after=drow)
    return {"type": "table", "key": tkey, "columns": columns}


def build_one(code):
    form_no, name, _ = REGISTRY[code]
    doc = Document(source_path(code))
    used = set()
    schema = {"code": code, "name": name, "form": form_no, "sections": []}
    section = {"title": name, "items": []}
    schema["sections"].append(section)
    ti = 0
    img_n = 0
    mode = None  # special-paragraph mode set by certain headings

    for c in list(doc.element.body.iterchildren()):
        tag = c.tag.split("}")[-1]
        if tag == "tbl":
            t = doc.tables[ti]
            ti += 1
            prefix = slugify(section["title"], set(), "")[:20] if section else "rows"
            section["items"].append(_process_table(t, used, prefix))
            continue
        if tag != "p":
            continue
        p = Paragraph(c, doc)
        text = p.text.strip()
        style = p.style.name or ""

        if style.startswith("Heading"):
            section = {"title": text, "items": []}
            schema["sections"].append(section)
            up = text.upper()
            if "DEVIATION" in up:
                mode = "deviation"
            elif "TEST PROCEDURE" in up:
                mode = "procedure"
            elif "MONITORING PARAMETER" in up:
                mode = "monitoring"
            else:
                mode = None
            continue

        # non-heading paragraph
        low = text.lower()
        if low.startswith("functional check is conducted"):
            if "sop_reference" not in used:
                used.add("sop_reference")
                p.text = "Functional check is conducted as per SOP reference number: {{ sop_reference }}."
                section["items"].append({"type": "field", "key": "sop_reference",
                                         "label": "SOP Reference Number", "input": "text"})
            continue
        if mode == "deviation" and text:
            p.text = "{{ deviation }}"
            section["items"].append({"type": "textarea", "key": "deviation", "label": "Deviation"})
            mode = "done_dev"
            continue
        if mode == "procedure" and text:
            p.text = "{{ test_procedure }}"
            section["items"].append({"type": "textarea", "key": "test_procedure", "label": "Test Procedure"})
            mode = "strip_proc"
            continue
        if mode == "strip_proc" and text:
            p._p.getparent().remove(p._p)  # drop remaining boilerplate procedure paras
            continue
        if mode == "monitoring" and text.startswith("<<"):
            p.text = "{{ monitoring_parameters }}"
            section["items"].append({"type": "textarea", "key": "monitoring_parameters",
                                     "label": "Monitoring Parameters"})
            mode = None
            continue

        # image captions -> image placeholder ABOVE, caption BELOW, no spacing
        if re.match(r"(Photo|Figure)\s*\d+\s*:", text):
            img_n += 1
            ikey = slugify(text.split(":")[0], used, "img_")
            new = p._p.makeelement(W + "p", {})
            p._p.addprevious(new)   # image goes on top; the caption stays underneath it
            np = Paragraph(new, p._parent)
            np.add_run("{{ " + ikey + " }}")
            for _tight in (np, p):  # tight block: no spacing around image/caption
                _tight.paragraph_format.space_before = Pt(0)
                _tight.paragraph_format.space_after = Pt(0)
            section["items"].append({"type": "image", "key": ikey, "label": text[:80]})
            continue

        # strip leftover angle-bracket hints so they don't appear literally
        if "<" in text and ">" in text:
            for run in p.runs:
                if "<" in run.text and ">" in run.text:
                    run.text = re.sub(r"<[^>]*>", "", run.text)

    # remove manual page breaks (avoid stray blank pages)
    for br in list(doc.element.body.iter(W + "br")):
        if br.get(W + "type") == "page":
            br.getparent().remove(br)

    os.makedirs(TPL_DIR, exist_ok=True)
    os.makedirs(SCHEMA_DIR, exist_ok=True)
    doc.save(os.path.join(TPL_DIR, f"{code}.docx"))
    with open(os.path.join(SCHEMA_DIR, f"{code}.json"), "w", encoding="utf-8") as fh:
        json.dump(schema, fh, ensure_ascii=False, indent=2)
    n_fields = sum(len(s.get("items", [])) for s in schema["sections"])
    return n_fields


def main():
    for code in GENERIC_CODES:
        try:
            n = build_one(code)
            print(f"OK {code:14s} -> {code}.docx + {code}.json  ({n} items)")
        except Exception as exc:  # noqa: BLE001
            print(f"FAIL {code}: {exc}")


if __name__ == "__main__":
    main()
