"""Render any test's docxtpl template (by code) with a context + fitted images."""
import copy
import functools
import os
import re

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .generator import strip_trailing_blank_paragraphs, _add_image_borders, _PIC_NS
from .layout import (polish_layout, page_break_before_top_sections, enforce_arial_fonts,
                     enforce_arial_procedure, enforce_body_arial, shrink_wide_obs_tables,
                     paginate_generic_datasheet)

TPL_DIR = os.path.join(os.path.dirname(__file__), "word_templates")


def _box(key, code=None):
    k = key.lower()
    if "sign" in k:
        return (40, 20)              # signatures stay small
    if "img_fc" in k:
        return (140, 52)             # functional-check captures: short so all 3 fit on one page
    # All test-setup photos and measurement/emission plots default to
    # 15.92 cm (W) x 9.5 cm (H) = 159.2 x 95 mm; editable per-image in the form.
    return (159.2, 95)


def _fit(tpl, path, box, exact=False):
    bw, bh = box
    if exact:
        # The user set an exact size in the image editor (Word Picture Format -> Size).
        # Honour it precisely: set BOTH width and height so the image is stretched /
        # squeezed to fill exactly bw x bh, matching the "Set image to frame" preview
        # (python-docx keeps both dimensions verbatim when both are given).
        return InlineImage(tpl, path, width=Mm(bw), height=Mm(bh))
    try:
        from PIL import Image
        with Image.open(path) as im:
            w, h = im.size
        if w and h and (w / h) > (bw / bh):
            return InlineImage(tpl, path, width=Mm(bw))
        return InlineImage(tpl, path, height=Mm(bh))
    except Exception:
        return InlineImage(tpl, path, width=Mm(bw))


def _prune_empty_limit_tables(doc):
    """Remove any RE Test-Limit table whose row-loop produced no data (header only)
    together with its 'Maximum permissible...' intro paragraph. Empty tables occur
    when that (family x band) combination doesn't apply — docxtpl leaves just the
    header row, which we then drop so only the applicable limit tables print."""
    from docx.oxml.ns import qn
    for tbl in list(doc.tables):
        hdr = " ".join(c.text for c in tbl.rows[0].cells).lower()
        is_limit = ("quasi-peak limit" in hdr) or ("peak limit" in hdr and "average limit" in hdr)
        if not is_limit or len(tbl.rows) > 1:
            continue
        prev = tbl._tbl.getprevious()
        while prev is not None and prev.tag != qn("w:p"):
            prev = prev.getprevious()
        if prev is not None:
            ptext = "".join(prev.itertext()).lower()
            if "permissible" in ptext or "as per" in ptext:
                prev.getparent().remove(prev)
        tbl._tbl.getparent().remove(tbl._tbl)


def _re_pageprop(p, tag):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn(tag)) is None:
        pPr.insert(0, OxmlElement(tag))


def _re_clearprop(p, tag):
    from docx.oxml.ns import qn
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        for el in pPr.findall(qn(tag)):
            pPr.remove(el)


_RE_CAPTION_PREFIX = ("FIGURE", "PHOTO", "TABLE")


def strip_manual_page_breaks(doc):
    """Remove manual run-level page breaks (<w:br w:type="page"/>) from the template.
    This prevents double page breaks when heading page-break properties are added."""
    from docx.oxml.ns import qn
    for br in list(doc.element.body.iter(qn("w:br"))):
        if br.get(qn("w:type")) == "page":
            br.getparent().remove(br)


def _re_paginate(doc):
    """Lay the RE datasheet out page-by-page like the intended structure:
      * page break before 1.4 FUNCTIONAL CHECK, 2.2 DEVIATION (so 2.1 TEST
        SPECIFICATION sits alone), 2.5 MEASUREMENT DATA, 2.6 TEST SETUP PICTURES,
        and 2.7 TEST EQUIPMENT USED (so 2.7/2.8/2.9 share the last page);
      * keep every image with its caption so a page break never orphans a label.
    Two 90 mm plots + captions + a heading fit one page naturally."""
    from docx.oxml.ns import qn
    # 2.4 TEST PROCEDURE is several long paragraphs: keep-with-next alone cannot stop
    # Word splitting it when it starts low on a page, so it begins on a fresh page and
    # therefore always fits in one piece.
    BREAK_HEADINGS = ("FUNCTIONAL CHECK", "DEVIATION FROM THE STANDARD", "TEST PROCEDURE",
                      "MEASUREMENT DATA", "TEST SETUP PICTURES", "TEST EQUIPMENT USED")
    pm = {p._p: p for p in doc.paragraphs}
    tm = {t._tbl: t for t in doc.tables}
    in_meas = False
    for ch in doc.element.body.iterchildren():
        if ch.tag == qn("w:p"):
            p = pm.get(ch)
            if p is None:
                continue
            t = p.text.strip()
            if p.style.name.startswith("Heading"):
                up = t.upper()
                in_meas = ("MEASUREMENT DATA" in up)
                if any(b in up for b in BREAK_HEADINGS):
                    _re_pageprop(p, "w:pageBreakBefore")
                continue
            is_caption = t.upper().startswith(_RE_CAPTION_PREFIX)
            if p._p.findall(".//" + qn("w:drawing")):
                _re_pageprop(p, "w:keepNext")          # image stays with its caption
            elif is_caption:
                # a Figure/Photo/Table caption belongs to the item ABOVE it — it must
                # NOT be glued to whatever follows (polish_layout glues pre-table
                # paras), or the caption+table block gets pushed to the next page.
                _re_clearprop(p, "w:keepNext")
        elif ch.tag == qn("w:tbl"):
            tb = tm.get(ch)
            if tb is not None and in_meas:
                hdr = " ".join(c.text for c in tb.rows[0].cells).lower()
                if "polarization" in hdr and "eut angle" in hdr:
                    # Keep the whole data table together so it moves as a unit
                    # instead of splitting across pages.
                    rows = tb.rows
                    for r in rows[:-1]:
                        for cell in r.cells:
                            for cp in cell.paragraphs:
                                cp.paragraph_format.keep_with_next = True


def _eft_insert_observation(doc, power, signal):
    """Insert the EFT dynamic observation table(s) right after the 'Power Line:' /
    'Signal Line:' heading paragraphs. Columns are variable (built from the selected
    test voltage), so the table is created here rather than templated."""
    def _find(text):
        for p in doc.paragraphs:
            if p.text.strip() == text:
                return p
        return None

    def _build(data):
        if not data or not data.get("cols"):
            return None
        cols, rows = data["cols"], data.get("rows", [])
        t = doc.add_table(rows=1 + len(rows), cols=1 + len(cols))
        try:
            t.style = "Table Grid"
        except Exception:
            pass
        hdr = t.rows[0].cells
        hdr[0].text = "Coupling path / line"
        for j, c in enumerate(cols):
            hdr[1 + j].text = c
        for i, row in enumerate(rows):
            cs = t.rows[1 + i].cells
            cs[0].text = row.get("label", "")
            for j, v in enumerate(row.get("cells", [])):
                if 1 + j < len(cs):
                    cs[1 + j].text = v
        el = t._tbl
        el.getparent().remove(el)          # detach from the end; re-inserted at the marker
        return el

    for text, data in (("Power Line:", power), ("Signal Line:", signal)):
        marker = _find(text)
        if marker is None:
            continue
        el = _build(data)
        if el is not None:
            marker._p.addnext(el)
        else:
            marker._p.getparent().remove(marker._p)   # port not tested -> drop the dangling heading


def _eft_insert_legend(doc, legend):
    """Replace the template's static 'A: ... / B: ...' observation legend with one
    '<code>: <description>' paragraph per unique code the engineer entered. The
    static legend paragraphs are short 'code:' lines (A:/B:/C:/D:); if the user
    entered no codes, the static legend is left as-is."""
    import re
    if not legend:
        return
    statics = [p for p in doc.paragraphs
               if re.match(r"^\s*[A-Za-z0-9]{1,3}\s*:\s", p.text or "")
               and not p.text.strip().lower().startswith(("power line", "signal line"))]
    if not statics:
        return
    anchor = statics[0]
    style = anchor.style
    for entry in legend:
        code = (entry.get("code") or "").strip()
        desc = (entry.get("desc") or "").strip()
        anchor.insert_paragraph_before(("%s: %s" % (code, desc)).rstrip(), style=style)
    for p in statics:
        p._p.getparent().remove(p._p)


def _surge_power_matrix(doc, data):
    """Build one AC/DC Power-Line observation table. Columns = a fixed CM/DM x line x
    phase grid parsed from the 'CM L→PE 0°'-style meta; rows = the selected test-voltage
    +/- pairs. A 3-row merged header mirrors the reference (Common/Differential Mode ->
    coupling line -> phase). Returns the detached <w:tbl> element, or None if no data."""
    if not data or not data.get("cols"):
        return None
    cols, rows = data["cols"], data.get("rows", [])
    ncol = len(cols)
    parsed = []                                   # [(mode, line, phase), ...]
    for c in cols:
        parts = c.split(" ")
        mode = parts[0] if parts else ""
        phase = parts[-1] if len(parts) >= 2 else ""
        line = " ".join(parts[1:-1]) if len(parts) >= 3 else (parts[1] if len(parts) == 2 else "")
        parsed.append((mode, line, phase))

    def _groups(depth):                            # consecutive cols equal on parsed[:depth]
        out, start = [], 0
        for i in range(1, ncol + 1):
            if i == ncol or parsed[i][:depth] != parsed[start][:depth]:
                out.append((start, i - 1))
                start = i
        return out

    t = doc.add_table(rows=3 + len(rows), cols=1 + ncol)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    lbl = t.cell(0, 0).merge(t.cell(2, 0))         # "Test Level (kV)" spans the 3 header rows
    lbl.text = "Test Level (kV)"
    MODE = {"CM": "Common Mode", "DM": "Differential Mode"}
    for a, b in _groups(1):                         # row 0: mode bands
        cell = t.cell(0, 1 + a).merge(t.cell(0, 1 + b)) if b > a else t.cell(0, 1 + a)
        cell.text = MODE.get(parsed[a][0], parsed[a][0])
    for a, b in _groups(2):                         # row 1: coupling lines
        cell = t.cell(1, 1 + a).merge(t.cell(1, 1 + b)) if b > a else t.cell(1, 1 + a)
        cell.text = parsed[a][1]
    for i in range(ncol):                           # row 2: phases
        t.cell(2, 1 + i).text = parsed[i][2]
    for r, row in enumerate(rows):                  # data rows
        t.cell(3 + r, 0).text = row.get("label", "")
        for j, v in enumerate(row.get("cells", [])):
            if 1 + j <= ncol:
                t.cell(3 + r, 1 + j).text = v
    el = t._tbl
    el.getparent().remove(el)
    return el


def _surge_signal_table(doc, data):
    """Build the Signal-Line observation table (row per signal line, columns from the
    posted CM/DM +/- meta). Empty rows are dropped; if the port was tested but nothing
    entered the blank body is kept. Returns the detached element or None."""
    if not data or not data.get("cols"):
        return None
    cols, rows = data["cols"], data.get("rows", [])
    kept = [r for r in rows
            if (r.get("label") or "").strip() or any((c or "").strip() for c in r.get("cells", []))]
    if not kept:
        kept = rows
    ncol = len(cols)
    t = doc.add_table(rows=1 + len(kept), cols=1 + ncol)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    hdr = t.rows[0].cells
    hdr[0].text = "Name of the signal Line"
    for j, c in enumerate(cols):
        hdr[1 + j].text = c
    for r, row in enumerate(kept):
        t.cell(1 + r, 0).text = row.get("label", "")
        for j, v in enumerate(row.get("cells", [])):
            if 1 + j <= ncol:
                t.cell(1 + r, 1 + j).text = v
    el = t._tbl
    el.getparent().remove(el)
    return el


def _surge_insert_observation(doc, ac, dc, signal):
    """Insert the Surge dynamic observation tables after the 'AC Power Line:' /
    'DC Power Line:' / 'Signal Line:' markers. A marker whose port was not tested
    (no data) is removed so no empty heading dangles."""
    def _find(text):
        for p in doc.paragraphs:
            if p.text.strip() == text:
                return p
        return None

    plan = (("AC Power Line:", ac, "power"),
            ("DC Power Line:", dc, "power"),
            ("Signal Line:", signal, "signal"))
    for text, data, kind in plan:
        marker = _find(text)
        if marker is None:
            continue
        el = _surge_power_matrix(doc, data) if kind == "power" else _surge_signal_table(doc, data)
        if el is not None:
            marker._p.addnext(el)
        else:
            marker._p.getparent().remove(marker._p)


_EMU_PER_CM = 360000
_A_NS_URI = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _tc_spans(tr):
    """[(tc, start_grid, span), ...] for one table row."""
    out, g = [], 0
    for tc in tr.findall(qn("w:tc")):
        gs = tc.find(qn("w:tcPr") + "/" + qn("w:gridSpan"))
        span = int(gs.get(qn("w:val"))) if gs is not None else 1
        out.append((tc, g, span))
        g += span
    return out


def _set_span(tc, span):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = tc.makeelement(qn("w:tcPr"), {})
        tc.insert(0, tcPr)
    gs = tcPr.find(qn("w:gridSpan"))
    if span <= 1:
        if gs is not None:
            tcPr.remove(gs)
        return
    if gs is None:
        gs = tcPr.makeelement(qn("w:gridSpan"), {})
        tcPr.append(gs)
    gs.set(qn("w:val"), str(span))


def _write_tc_text(tc, text):
    """Replace a cell's text, keeping the first run's formatting."""
    ps = tc.findall(qn("w:p"))
    if not ps:
        return
    for extra in ps[1:]:
        tc.remove(extra)
    p = ps[0]
    runs = p.findall(qn("w:r"))
    if not runs:
        r = p.makeelement(qn("w:r"), {})
        t = p.makeelement(qn("w:t"), {})
        r.append(t)
        p.append(r)
        runs = [r]
    for extra in runs[1:]:
        p.remove(extra)
    ts = runs[0].findall(qn("w:t"))
    if not ts:
        t = runs[0].makeelement(qn("w:t"), {})
        runs[0].append(t)
        ts = [t]
    ts[0].text = text or ""
    ts[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    for extra in ts[1:]:
        extra.getparent().remove(extra)


def _grid_widths(tb):
    g = tb._tbl.find(qn("w:tblGrid"))
    return [int(c.get(qn("w:w")) or 0) for c in (g.findall(qn("w:gridCol")) if g is not None else [])]


def _sync_row_widths(tb):
    """Re-derive every cell's preferred width (w:tcW) from the table grid + its
    gridSpan. Word lays a table out from tcW, so a span change that leaves a stale
    tcW makes it recompute the columns — which squeezes the label column and pushes
    the table onto a second page."""
    widths = _grid_widths(tb)
    if not widths:
        return
    for tr in tb._tbl.findall(qn("w:tr")):
        g = 0
        for tc in tr.findall(qn("w:tc")):
            gs = tc.find(qn("w:tcPr") + "/" + qn("w:gridSpan"))
            span = int(gs.get(qn("w:val"))) if gs is not None else 1
            total = sum(widths[g:g + span])
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = tc.makeelement(qn("w:tcPr"), {})
                tc.insert(0, tcPr)
            w = tcPr.find(qn("w:tcW"))
            if w is None:
                w = tcPr.makeelement(qn("w:tcW"), {})
                tcPr.append(w)
            w.set(qn("w:w"), str(total))
            w.set(qn("w:type"), "dxa")
            g += span


def _re_spec_table(doc):
    """The RE 2.1 TEST SPECIFICATION table (identified by its 'Product Standard' row)."""
    for tb in doc.tables:
        try:
            if (tb.rows[0].cells[0].text or "").strip().lower().startswith("product standard"):
                return tb
        except (IndexError, AttributeError):
            continue
    return None


def _re_row_by_label(tb, *needles):
    for idx, row in enumerate(tb.rows):
        lbl = (row.cells[0].text or "").strip().lower()
        if any(n in lbl for n in needles):
            return idx
    return None


def _re_set_row_sections(tb, row_idx, values):
    """Split a spec row's value area (grid cols 1..N) into len(values) sections."""
    if row_idx is None or not values:
        return
    tr = tb._tbl.findall(qn("w:tr"))[row_idx]
    cells = _tc_spans(tr)
    if len(cells) < 2:
        return
    total = sum(span for _, _, span in cells[1:])          # grid width of the value area
    n = min(len(values), total)
    base, extra = divmod(total, n)
    spans = [base + (1 if i < extra else 0) for i in range(n)]
    proto = cells[1][0]
    for tc, _, _ in cells[2:]:                             # drop the old value cells
        tr.remove(tc)
    anchor = proto
    _set_span(proto, spans[0])
    _write_tc_text(proto, values[0])
    for i in range(1, n):
        new = copy.deepcopy(proto)
        _set_span(new, spans[i])
        _write_tc_text(new, values[i])
        anchor.addnext(new)
        anchor = new


# The template's value area is 4 uneven grid columns (one is a 168-twip sliver), so a
# 3-way split would land on it and wrap. Re-grid the value area into 12 equal columns
# (LCM of 1/2/3) mapped from the originals, so every 1/2/3 split is exact and each
# section stays on one line. The label column keeps its original width.
_RE_NEW_COLS = (3, 2, 1, 6)          # new columns per original value column -> 12


def _set_cell_sections(tb, row_idx, cell_idx, values):
    """Split ONE value cell of a spec row into len(values) side-by-side sections.

    _re_set_row_sections splits a row's whole value area, which is right for RE and CE
    where a row has a single value cell. RS_RI's rows carry one cell per frequency band, so
    splitting the whole area would collapse the bands together; this replaces just the cell
    at `cell_idx` (1 = first value cell) and leaves its neighbours alone.
    """
    if row_idx is None or cell_idx is None or not values:
        return False
    tr = tb._tbl.findall(qn("w:tr"))[row_idx]
    cells = _tc_spans(tr)
    if cell_idx >= len(cells):
        return False
    proto, _, span = cells[cell_idx]
    n = min(len(values), span)                 # cannot make more sections than grid columns
    if n < 1:
        return False
    base, extra = divmod(span, n)
    spans = [base + (1 if i < extra else 0) for i in range(n)]
    _set_span(proto, spans[0])
    _write_tc_text(proto, values[0])
    anchor = proto
    for i in range(1, n):
        new = copy.deepcopy(proto)
        _set_span(new, spans[i])
        _write_tc_text(new, values[i])
        anchor.addnext(new)
        anchor = new
    return True


def _re_regrid_value_area(tb, new_cols=_RE_NEW_COLS, label_cols=1):
    """Split the value area into equal grid columns, remapping existing spans.

    new_cols gives how many new columns each ORIGINAL value column becomes, so its sum must
    be a multiple of every section count wanted (RE: 12; CE and RS_RI: (3, 3) -> 6).

    label_cols is how many leading grid columns the LABEL occupies - 1 for RE/CE/RS_RI, but
    2 for SURGE, whose label cell carries gridSpan=2. Those columns keep their widths and
    the label keeps its span; only the value area is re-gridded.
    """
    widths = _grid_widths(tb)
    if len(widths) != label_cols + len(new_cols):
        return False                  # unexpected geometry -> leave the table alone
    keep = widths[:label_cols]
    value_total = sum(widths[label_cols:])
    n = sum(new_cols)
    base, rem = divmod(value_total, n)
    new_widths = [base + (1 if i < rem else 0) for i in range(n)]

    for tr in tb._tbl.findall(qn("w:tr")):
        for tc, g0, span in _tc_spans(tr):
            if g0 < label_cols:
                # Inside the label region. With one label column, normalise to 1 exactly as
                # before. With more, leave the span ALONE: SURGE's Test Voltage rows divide
                # the 2-column label region into 1 + 1 ('Test Voltage(kV)' and 'Common Mode
                # (CM)'), and forcing each to 2 made those rows 10 columns wide on an
                # 8-column grid, which is what broke the table.
                if label_cols == 1:
                    _set_span(tc, 1)
                continue
            start = g0 - label_cols
            _set_span(tc, sum(new_cols[start:start + span]) or 1)

    grid = tb._tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        grid.remove(gc)
    for w in keep + new_widths:
        gc = grid.makeelement(qn("w:gridCol"), {})
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    return True


#: Fallback average glyph advance for Arial, as a fraction of the font size, used only
#: when the real font metrics can't be read. It over-estimates mixed-case text badly
#: ('Kondababu Arjilli' measures 82.5pt at 11pt but estimates 97.2pt), so it is a last
#: resort - _text_width_em() measures the actual glyph advances instead.
_ARIAL_AVG_EM = 0.52
_CELL_MARGIN_PT = 10.8          # Word's default 0.19 cm left+right cell margins
#: Word wraps only when the text exceeds the usable width exactly, so this stays tiny -
#: just enough to absorb metric rounding. A larger value would shrink text that really
#: does fit, which is the bug this replaced.
_FIT_HEADROOM_PT = 0.25

_ARIAL_FONT_PATHS = (
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\Arial.ttf",
    "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",   # metric-compatible
    "/Library/Fonts/Arial.ttf",
)


@functools.lru_cache(maxsize=1)
def _arial_font():
    """Arial loaded at 1000 pt, so getlength(text)/1000 is the width in em units and the
    width at N pt is simply em * N. None when no Arial-metric font is installed."""
    try:
        from PIL import ImageFont
    except Exception:
        return None
    for path in _ARIAL_FONT_PATHS:
        try:
            if os.path.exists(path):
                return ImageFont.truetype(path, 1000)
        except Exception:
            continue
    return None


@functools.lru_cache(maxsize=512)
def _text_width_em(text):
    """Width of text in em units (multiply by the point size to get points)."""
    font = _arial_font()
    if font is not None:
        try:
            return font.getlength(text) / 1000.0
        except Exception:
            pass
    return len(text) * _ARIAL_AVG_EM


def _set_cell_font_pt(tc, pt):
    """Force every run in a cell to a given point size (w:sz is in half-points)."""
    half = str(int(round(pt * 2)))
    for r in tc.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = r.makeelement(qn("w:rPr"), {})
            r.insert(0, rPr)
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = rPr.makeelement(qn(tag), {})
                rPr.append(el)
            el.set(qn("w:val"), half)


def _re_fit_row_font(tb, needle, max_pt=11.0, min_pt=7.0, step=0.5):
    """Keep a row's values on ONE line. Try the normal 11 pt first; only when a value
    is too wide for its (possibly 3-way split) section, step down to the closest size
    that fits — never below min_pt. Returns the size applied, or None if 11 pt was fine."""
    r = _re_row_by_label(tb, needle)
    if r is None:
        return None
    cells = _tc_spans(tb._tbl.findall(qn("w:tr"))[r])[1:]
    if not cells:
        return None
    demand = []
    for tc, _, _ in cells:
        w = tc.find(qn("w:tcPr") + "/" + qn("w:tcW"))
        avail = ((int(w.get(qn("w:w"))) / 20.0) - _CELL_MARGIN_PT) if w is not None else 0.0
        text = "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
        if text:
            demand.append((_text_width_em(text), max(1.0, avail)))
    if not demand:
        return None

    def fits(pt):
        return all(em * pt + _FIT_HEADROOM_PT <= avail for em, avail in demand)

    if fits(max_pt):
        return None                                  # 11 pt is fine, leave it alone
    pt, chosen = max_pt, None
    while pt - step >= min_pt:
        pt -= step
        if fits(pt):
            chosen = pt
            break
    if chosen is None:
        # Nothing down to min_pt keeps it on one line, so shrinking would make the text
        # unreadable AND still wrap. Leave the row at its normal size and let Word wrap.
        return None
    for tc, _, _ in cells:
        _set_cell_font_pt(tc, chosen)
    return chosen


#: Rows whose two value cells are OPTION pairs, not frequency-range columns
#: (Group|Class and Tabletop|Floor standing). They must survive the range collapse.
_RE_OPTION_PAIR_ROWS = ("classification", "eut configuration")


def _re_collapse_single_range(tb, use_second):
    """Only one Frequency Range is selected, so the spec table keeps ONE value column
    spanning the whole value area. use_second=True carries the 1GHz-6GHz column's
    content over (the 30MHz-1GHz column is the one that gets dropped)."""
    for tr in tb._tbl.findall(qn("w:tr")):
        cells = _tc_spans(tr)
        if len(cells) < 2:
            continue
        label = "".join(t.text or "" for t in cells[0][0].iter(qn("w:t"))).strip().lower()
        if any(k in label for k in _RE_OPTION_PAIR_ROWS):
            continue        # keep both checkbox cells (e.g. Tabletop | Floor standing)
        value_cells = cells[1:]
        total = sum(span for _, _, span in value_cells)
        keep = value_cells[0][0]
        if use_second and len(value_cells) > 1:
            src = value_cells[-1][0]
            for p in keep.findall(qn("w:p")):
                keep.remove(p)
            for p in src.findall(qn("w:p")):
                keep.append(copy.deepcopy(p))
        for tc, _, _ in value_cells[1:]:
            tr.remove(tc)
        _set_span(keep, total)


def _re_delete_grid_cols(tb, start, count):
    """Remove grid columns [start, start+count) from a table, shrinking spans."""
    if count <= 0:
        return
    end = start + count
    for tr in tb._tbl.findall(qn("w:tr")):
        for tc, g0, span in _tc_spans(tr):
            g1 = g0 + span
            overlap = max(0, min(g1, end) - max(g0, start))
            if overlap <= 0:
                continue
            if overlap >= span:
                tr.remove(tc)
            else:
                _set_span(tc, span - overlap)
    grid = tb._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        gcs = grid.findall(qn("w:gridCol"))
        for gc in gcs[start:end]:
            grid.remove(gc)


# Tables that stay LEFT-aligned in the reference datasheet (identified by the label
# in their first cell): 1.1 EUT DETAILS, 2.1 TEST SPECIFICATION and the sign-off
# block. Every other table is centered.
_RE_LEFT_TABLES = ("job number", "product standard", "tested by")


def _re_align_tables(doc):
    """Mirror the reference datasheet's table alignment: data/limit/equipment tables
    centered, the EUT-details / test-specification / sign-off tables left-aligned
    (their '-' placeholder cells stay centered), all cells vertically centered."""
    for tb in doc.tables:
        try:
            first = (tb.rows[0].cells[0].text or "").strip().lower()
        except (IndexError, AttributeError):
            first = ""
        left = any(first.startswith(k) for k in _RE_LEFT_TABLES)
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = (p.text or "").strip()
                    if left:
                        # '-' placeholders are centered even in a left-aligned table
                        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if txt in ("-", "–", "—")
                                       else WD_ALIGN_PARAGRAPH.LEFT)
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = cell._tc.makeelement(qn("w:tcPr"), {})
                    cell._tc.insert(0, tcPr)
                va = tcPr.find(qn("w:vAlign"))
                if va is None:
                    va = tcPr.makeelement(qn("w:vAlign"), {})
                    tcPr.append(va)
                va.set(qn("w:val"), "center")


def _re_center_signoff_title(doc):
    """Centre the 'Tested By' / 'Reviewed By' heading row of the sign-off block. The
    rest of that table (Name / Signature / Date and their values) stays left-aligned,
    so this runs after _re_align_tables, which left-aligns the whole table."""
    n = 0
    for tb in doc.tables:
        try:
            rows = tb.rows
            first = (rows[0].cells[0].text or "").strip().lower()
        except (IndexError, AttributeError):
            continue
        # Identified by shape: a heading row above Name / Signature / Date.
        if not first.startswith(("tested by", "reviewed by", "approved by")):
            continue
        labels = [(rows[i].cells[0].text or "").strip().lower() for i in range(1, len(rows))]
        if "name" not in labels:
            continue
        for cell in rows[0].cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            n += 1
    return n


def _re_tighten_image_spacing(doc):
    """An image paragraph inherits the document default 8pt space-after (docDefaults
    w:after=160) plus 1.08 line spacing, so a figure ends up further from its caption
    than a table does - a table contributes no paragraph spacing at all. Zero it so
    image->caption spacing matches table->caption spacing everywhere."""
    n = 0
    for p in doc.paragraphs:
        if p._p.findall(".//" + qn("w:drawing")):
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1
            n += 1
    return n


def _re_caption_spacing(doc):
    """Captions under graphs/tables: single line spacing, centred, and NO space above -
    so every caption sits the same distance below its image or table (the template put
    6pt above some photo captions, which made those gaps look bigger)."""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.upper().startswith(_RE_CAPTION_PREFIX):
            pf = p.paragraph_format
            pf.line_spacing = 1
            pf.space_before = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _re_subsection_spacing(doc, pt=11):
    """Font-size-11 worth of space between sub-sections (2.1, 2.2, ...)."""
    first = True
    for p in doc.paragraphs:
        if p.style.name == "Heading 2":
            p.paragraph_format.space_before = Pt(0 if first else pt)
            first = False


def _re_procedure_heading_gap(doc, pt=11):
    """One blank line between the TEST PROCEDURE title and the procedure text.

    The Heading 2 style carries space_after=0, so the procedure started immediately under
    its title with nothing between them. `pt` is one 11pt line, i.e. the single Enter the
    reference layout shows. Scoped to that one heading - the other sub-sections open with a
    table or a short line and read correctly as they are."""
    n = 0
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and (p.text or "").strip().upper() == "TEST PROCEDURE":
            p.paragraph_format.space_after = Pt(pt)
            n += 1
    return n


def _re_keep_subsections(doc):
    """Keep each sub-section together so a short block is pushed whole to the next
    page rather than splitting. Long blocks (which cannot fit a page) are skipped."""
    paras = doc.paragraphs
    starts = [i for i, p in enumerate(paras) if p.style.name == "Heading 2"]
    for si, s in enumerate(starts):
        end = starts[si + 1] if si + 1 < len(starts) else len(paras)
        block = paras[s:end]
        if len(block) > 9:                     # too long to guarantee one page
            continue
        for p in block[:-1]:
            if not (p.text or "").strip().upper().startswith(_RE_CAPTION_PREFIX):
                p.paragraph_format.keep_with_next = True


def _re_fix_signature(doc, max_h_cm=1.5):
    """Signature image: no border, height capped at 1.5 cm (aspect preserved)."""
    tables = doc.tables
    if not tables:
        return
    for tb in tables[-2:]:                      # RESULT / sign-off block
        for row in tb.rows:
            for cell in row.cells:
                for inline in cell._tc.iter(qn("wp:inline")):
                    ext = inline.find(qn("wp:extent"))
                    if ext is None:
                        continue
                    cx, cy = int(ext.get("cx") or 0), int(ext.get("cy") or 0)
                    cap = int(max_h_cm * _EMU_PER_CM)
                    if cy > cap and cy:
                        ext.set("cx", str(max(1, int(cx * cap / cy))))
                        ext.set("cy", str(cap))
                for spPr in cell._tc.iter("{%s}spPr" % _PIC_NS):
                    for ln in spPr.findall("{%s}ln" % _A_NS_URI):
                        spPr.remove(ln)


def _re_prune_photos(doc, show_30m_1g, show_1g_6g, keep=(), custom_range=""):
    """Drop the Test-Setup photo slots that do not apply.

    Standard ranges are pruned by the band named in the caption. A CUSTOM range is ONE
    band, so no caption names a band to match on - there the template's slots are pruned
    by POSITION instead, keeping the first Vertical/Horizontal pair.

    'keep' holds the captions of pictures the engineer added by hand; those are never
    pruned and never counted as one of the template's two slots.
    """
    protected = {c.strip() for c in keep if (c or "").strip()}
    custom = bool((custom_range or "").strip())
    if not custom and show_30m_1g and show_1g_6g:
        return                                   # 'Both': every slot applies
    drop = () if custom else (("1GHz - 6GHz",) if not show_1g_6g else ("30MHz - 1GHz",))

    body = doc.element.body
    paras = list(doc.paragraphs)
    removed = []
    kept_template_slots = 0
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t.upper().startswith("PHOTO "):
            continue
        if t in protected:
            continue                             # an engineer-added picture
        if custom:
            # one band -> the first pair is the band's Vertical/Horizontal
            kept_template_slots += 1
            if kept_template_slots <= 2:
                continue
        elif not any(d in t for d in drop):
            continue
        removed.append(p)
        if i > 0 and paras[i - 1]._p.findall(".//" + qn("w:drawing")):
            removed.append(paras[i - 1])
        elif i > 0 and not (paras[i - 1].text or "").strip():
            removed.append(paras[i - 1])
    for p in removed:
        if p._p.getparent() is not None:
            body.remove(p._p)


#: Full content width of an RE page in twips - every other table in the template uses it.
_RE_TABLE_WIDTH_TWIPS = 9016


def _re_drop_empty_upload_tables(doc, tables):
    """Remove an upload-driven table (and the blank paragraph above it) when nothing was
    uploaded, so the datasheet does not carry an empty grid. `tables` is the list of
    {'headers', 'has_data'} dicts built by collect_upload_table()."""
    empties = [
        [str(h or "").strip() for h in (t.get("headers") or [])]
        for t in (tables or []) if not t.get("has_data") and t.get("headers")
    ]
    if not empties:
        return 0
    body = doc.element.body
    removed = 0
    for tb in list(doc.tables):
        rows = tb._tbl.findall(qn("w:tr"))
        if not rows:
            continue
        header = [
            "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
            for tc in rows[0].findall(qn("w:tc"))
        ]
        if header not in empties:
            continue
        prev = tb._tbl.getprevious()
        if tb._tbl.getparent() is not None:
            body.remove(tb._tbl)
            removed += 1
        # tidy up the spacer paragraph that preceded it
        if prev is not None and prev.tag == qn("w:p") and not \
                "".join(t.text or "" for t in prev.iter(qn("w:t"))).strip():
            if prev.getparent() is not None:
                body.remove(prev)
    return removed


def _re_sync_meas_table_widths(doc, groups, extra_header_sets=None, width=None):
    """The MEASUREMENT DATA tables are generated with a column loop, so their column
    COUNT comes from the data. docxtpl rebuilds w:tblGrid for the rendered count, but it
    re-splits the template's width and leaves each cell's w:tcW at the template value -
    so the total drifts a few twips and Word re-lays the table out from the stale widths.
    Rebuild the grid as N exactly-equal columns summing to the full page width, then
    re-derive every cell width from it.

    Tables are matched by their header row equalling a group's table_headers, so renamed
    or added columns still match and no other RE table is touched.

    `width` overrides the total table width in twips, so CE (whose measurement tables are
    9024 wide, not RE's 9016) can share this pass.
    """
    total = int(width or _RE_TABLE_WIDTH_TWIPS)
    wanted = []
    for g in (groups or []):
        # a group can hold several tables (30MHz-1GHz quasi-peak, plus Peak AND Average
        # above 1GHz), each with its own headers
        for t in (g.get("tables") or [{"headers": g.get("table_headers")}]):
            hdrs = [str(h or "").strip() for h in (t.get("headers") or [])]
            if hdrs and hdrs not in wanted:
                wanted.append(hdrs)
    for extra in (extra_header_sets or []):
        hdrs = [str(h or "").strip() for h in (extra or [])]
        if hdrs and hdrs not in wanted:
            wanted.append(hdrs)
    if not wanted:
        return 0

    fixed = 0
    for tb in doc.tables:
        rows = tb._tbl.findall(qn("w:tr"))
        if not rows:
            continue
        header = [
            "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
            for tc in rows[0].findall(qn("w:tc"))
        ]
        if header not in wanted:
            continue
        n = len(header)
        if n < 1:
            continue
        grid = tb._tbl.find(qn("w:tblGrid"))
        if grid is None:
            continue
        base, rem = divmod(total, n)
        new_widths = [base + (1 if i < rem else 0) for i in range(n)]
        for gc in grid.findall(qn("w:gridCol")):
            grid.remove(gc)
        for w in new_widths:
            gc = grid.makeelement(qn("w:gridCol"), {})
            gc.set(qn("w:w"), str(w))
            grid.append(gc)
        _sync_row_widths(tb)
        fixed += 1
    return fixed


def _re_tidy_caption_whitespace(doc):
    """Last-pass tidy of every Photo/Figure/Table caption: collapse runs of spaces and
    drop a stray space straight after an underscore ('RE plot_ Vertical' -> 'RE plot_
    Vertical' becomes 'RE plot_Vertical').

    The defaults no longer produce these, but a caption can also arrive from a saved
    draft or be typed by hand, so the document is normalised here rather than trusting
    every upstream source."""
    import re as _re
    fixed = 0
    for p in doc.paragraphs:
        t = (p.text or "")
        if not t.strip().upper().startswith(("PHOTO ", "FIGURE ", "TABLE ")):
            continue
        new = _re.sub(r"_[ \t]+", "_", t)        # '_ Vertical' -> '_Vertical'
        new = _re.sub(r"[ \t]{2,}", " ", new)    # collapse double spaces
        new = new.strip()
        if new != t:
            _write_para_text(p, new)
            fixed += 1
    return fixed


#: Standard band spellings that appear in RE's fixed template/schema caption text.
#: '1-6GHz' is the short form used in the 1GHz-6GHz TABLE captions.
_RE_CAPTION_BANDS = ("30MHz - 1GHz", "1GHz - 6GHz", "30MHz-1GHz", "1GHz-6GHz", "1-6GHz")


def _re_relabel_captions_custom_range(doc, custom_range, bands=None):
    """Rename every Photo/Figure/Table caption that still carries a STANDARD band label
    to the custom Frequency Range.

    Measurement plot and table captions are already built with the custom label by the
    service. The ones that need fixing here come from fixed template/schema text - the
    Test Setup photo captions and the Functional Check ambient plots
    ('Figure 1: RE_Ambient_plot_ Vertical_Peak_30MHz - 1GHz') - because a custom range is
    a single band and none of the standard band names apply to it. A caption the engineer
    retyped himself no longer contains a standard band name, so it is left alone.

    `bands` overrides the spellings to look for, so CE can pass its own
    ('0.15MHz - 30MHz', ...) and share this pass.
    """
    if not custom_range:
        return 0
    renamed = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if ":" not in t or not t.upper().startswith(("PHOTO ", "FIGURE ", "TABLE ")):
            continue
        head, rest = t.split(":", 1)
        new = rest
        for band in (bands or _RE_CAPTION_BANDS):
            if band in new:
                new = new.replace(band, custom_range)
        if new != rest:
            _write_para_text(p, "%s:%s" % (head, new))
            renamed += 1
    return renamed


def _re_drop_captionless_photos(doc):
    """Remove a Test-Setup photo caption when nothing was uploaded for that slot. The
    template pairs each '{{ img_photo_N }}' paragraph with its caption, so an empty slot
    would otherwise print 'Photo N: ...' over blank space. Runs before renumbering so the
    survivors stay 1..N."""
    body = doc.element.body
    paras = list(doc.paragraphs)
    removed = []
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not (t.upper().startswith("PHOTO ") and ":" in t):
            continue
        prev = paras[i - 1] if i > 0 else None
        if prev is None:
            continue
        if prev._p.findall(".//" + qn("w:drawing")):
            continue                        # has a picture -> keep
        if (prev.text or "").strip():
            continue                        # not this caption's (empty) image paragraph
        removed.extend([p, prev])
    for p in removed:
        if p._p.getparent() is not None:
            body.remove(p._p)
    return len(removed) // 2


def _re_renumber_tables(doc):
    """Number every 'Table N:' caption 1..N in document order.

    CE needs this because its table captions are generated per plot group inside the
    measurement-record loop: without it every Test restarts at 'Table 1', and a group that
    printed nothing (an Average plot the engineer skipped) would leave a gap."""
    n = 1
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.upper().startswith("TABLE ") and ":" in t:
            _write_para_text(p, "Table %d:%s" % (n, t.split(":", 1)[1]))
            n += 1
    return n - 1


def _re_renumber_figures(doc):
    """Number every 'Figure N:' caption 1..N in document order. The document runs one
    continuous sequence (Functional Check plots, then the measurement plots of each
    group), and a slot with no upload is dropped, so without this the numbering would
    show gaps wherever an empty slot had reserved a number."""
    n = 1
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.upper().startswith("FIGURE ") and ":" in t:
            _write_para_text(p, "Figure %d:%s" % (n, t.split(":", 1)[1]))
            n += 1
    return n - 1


def _re_renumber_photos(doc):
    """Number every surviving Test-Setup photo caption 'Photo 1..N' in document order.
    Runs unconditionally (not only after pruning) so pictures the engineer appended
    continue the sequence after the standard slots that actually printed."""
    n = 1
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.upper().startswith("PHOTO ") and ":" in t:
            _write_para_text(p, "Photo %d:%s" % (n, t.split(":", 1)[1]))
            n += 1
    return n - 1


def _re_rebuild_procedure(doc, text):
    """Replace the template's hard-coded 2.4 TEST PROCEDURE body with the actual
    procedure text (already range-filtered and mapped), one paragraph per block with a
    one-line gap between them. The template carries only the 30MHz-1GHz wording and no
    {{ test_procedure }} placeholder, so the body has to be rebuilt here."""
    import re as _re
    from docx.text.paragraph import Paragraph
    blocks = [b.strip() for b in _re.split(r"\n\s*\n", (text or "").strip()) if b.strip()]
    if not blocks:
        return False
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if (p.style.name or "").startswith("Heading") and "TEST PROCEDURE" in (p.text or "").upper():
            start = i
            break
    if start is None:
        return False
    end = start + 1
    while end < len(paras) and not (paras[end].style.name or "").startswith("Heading"):
        end += 1
    body = [p for p in paras[start + 1:end]]
    if not body:
        return False

    proto = body[0]
    anchor = proto._p
    new_els = []
    for blk in blocks:
        el = copy.deepcopy(proto._p)
        np = Paragraph(el, proto._parent)
        _write_para_text(np, blk)
        pf = np.paragraph_format
        pf.line_spacing = 1
        pf.space_after = Pt(11)          # one blank line between paragraphs
        pf.space_before = Pt(0)
        new_els.append(el)
    for el in reversed(new_els):
        anchor.addnext(el)
    for p in body:                        # drop the template's hard-coded wording
        if p._p.getparent() is not None:
            p._p.getparent().remove(p._p)
    return True


def _re_fill_missing_na(doc, header_needle="calibration due", value="NA"):
    """Any blank Calibration Due cell on a populated equipment row prints 'NA' rather
    than an empty box. Rows that are entirely empty (spare rows) are left alone."""
    filled = 0
    for tb in doc.tables:
        try:
            hdr = [(c.text or "").strip().lower() for c in tb.rows[0].cells]
        except (IndexError, AttributeError):
            continue
        idx = next((i for i, h in enumerate(hdr) if header_needle in h), None)
        if idx is None:
            continue
        for row in tb.rows[1:]:
            cells = row.cells
            if idx >= len(cells):
                continue
            if (cells[idx].text or "").strip():
                continue
            if any((c.text or "").strip() for i, c in enumerate(cells) if i != idx):
                _write_tc_text(cells[idx]._tc, value)
                filled += 1
    return filled


def _re_fix_procedure_standard(doc, basic_standard):
    """The RE template hard-codes the whole 2.4 TEST PROCEDURE text, including the
    reference document's standard names (there is no {{ test_procedure }} placeholder).
    Rewrite just the opening line so it names THIS datasheet's Basic Standard."""
    std = (basic_standard or "").strip()
    if not std:
        return False
    lead = "the test procedure was in accordance with"
    for p in doc.paragraphs:
        if (p.text or "").strip().lower().startswith(lead):
            _write_para_text(p, "The test procedure was in accordance with %s" % std)
            return True
    return False


def _write_para_text(p, text):
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)


#: RS_RI's spec value area is 2 grid columns (one per frequency band); 3 equal columns from
#: each gives 6, which divides by 1, 2 and 3, so every per-day split lands on a grid
#: boundary and stays on one line.
_RS_RI_VALUE_REGRID = (3, 3)


def _rs_ri_unjustify(doc, *needles):
    """Left-align the free-text blocks the template justifies.

    MONITORING PARAMETERS is one paragraph of soft-broken lines with w:jc='both', so Word
    stretches every line to the full column width - 'No    Error    Message' spread right
    across the page. Left-aligning restores normal word spacing. Matched by the heading
    above the block so only that text is touched."""
    hit = 0
    want = False
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        style = (p.style.name or "") if p.style is not None else ""
        if style.startswith("Heading"):
            want = any(n in t.upper() for n in needles)
            continue
        if want and t:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hit += 1
    return hit


def _rs_ri_center_placeholders(doc, header_needle="modification state", text="-"):
    """Centre the placeholder dashes in the EUT Modification Record.

    The engineer leaves 'Modification fitted by' / 'Date modification fitted' empty on the
    initial-state row and a '-' stands in; left-aligned it read as stray punctuation."""
    n = 0
    for tb in doc.tables:
        try:
            hdr = " ".join((c.text or "").strip().lower() for c in tb.rows[0].cells)
        except (IndexError, AttributeError):
            continue
        if header_needle not in hdr:
            continue
        for row in tb.rows[1:]:
            for cell in row.cells:
                if (cell.text or "").strip() != text:
                    continue
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                n += 1
    return n


def _rs_ri_blacken_observation(doc):
    """Force black text in the TEST OBSERVATION table.

    The template ships the '1000 to 6000' band label in blue (0070C0), left over from
    drafting; every other row is black, so it read as a defect."""
    n = 0
    for tb in doc.tables:
        try:
            hdr = " ".join((c.text or "").strip().lower() for c in tb.rows[0].cells)
        except (IndexError, AttributeError):
            continue
        if "dwell time" not in hdr:
            continue
        for color in tb._tbl.iter(qn("w:color")):
            if (color.get(qn("w:val")) or "").lower() not in ("000000", "auto"):
                color.set(qn("w:val"), "000000")
                n += 1
    return n


def _blacken_all_tables(doc):
    """Force every table run to black. Several templates carry drafting colours (SURGE's
    observation matrices and a few spec values are blue), which read as defects next to the
    black text around them."""
    n = 0
    for tb in doc.tables:
        for color in tb._tbl.iter(qn("w:color")):
            if (color.get(qn("w:val")) or "").lower() not in ("000000", "auto"):
                color.set(qn("w:val"), "000000")
                n += 1
    return n


#: SURGE's spec table puts the label in TWO grid columns (gridSpan=2) and the value in the
#: other two; 3 equal columns from each value column gives 6, which divides by 1, 2 and 3.
_SURGE_VALUE_REGRID = (3, 3)
_SURGE_LABEL_COLS = 2


def _bullet_monitoring_parameters(doc, heading="MONITORING PARAMETERS", bullet="•"):
    """Turn the Monitoring Parameters block into one bulleted line per point.

    The value arrives as a single paragraph of soft-broken lines, each already numbered by
    whoever typed the Test Request ('1.', '3.', '4.', ...) - numbering that often skips.
    Split it into a paragraph per point, drop that leading number and prefix a bullet with a
    hanging indent, so the list reads as a list.
    """
    from docx.shared import Pt as _Pt
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    pm = {p._p: p for p in doc.paragraphs}
    want = False
    made = 0
    for el in list(body.iterchildren()):
        if el.tag != qn("w:p"):
            continue
        p = pm.get(el)
        if p is None:
            continue
        style = (p.style.name or "") if p.style is not None else ""
        if style.startswith("Heading"):
            want = heading in (p.text or "").upper()
            continue
        if not want:
            continue
        raw = (p.text or "").strip()
        if not raw:
            continue
        # one point per line; a line may be numbered, or bulleted already
        points = [ln.strip() for ln in re.split(r"[\r\n\v]+", raw) if ln.strip()]
        points = [re.sub(r"^(?:\d+\s*[.)]\s*|[•\-–]\s*)", "", pt).strip() for pt in points]
        points = [pt for pt in points if pt]
        if not points:
            continue
        anchor = el
        for i, pt in enumerate(points):
            if i == 0:
                target = p
            else:
                new = copy.deepcopy(el)
                anchor.addnext(new)
                anchor = new
                target = Paragraph(new, p._parent)
            _write_para_text(target, "%s\t%s" % (bullet, pt))
            target.paragraph_format.left_indent = _Pt(18)
            target.paragraph_format.first_line_indent = _Pt(-18)
            target.paragraph_format.space_after = _Pt(2)
            target.alignment = WD_ALIGN_PARAGRAPH.LEFT
            made += 1
        want = False        # only the first text block under the heading
    return made


#: Lines inside the Test Procedure that label a port and should stand out as headings.
_SURGE_PORT_HEADINGS = ("power line", "signal line", "power lines", "signal lines")


def _surge_bold_port_headings(doc, heading="TEST PROCEDURE"):
    """Bold the 'Power Line:' / 'Signal Line:' lines inside the Test Procedure.

    The procedure arrives as ONE run holding every line separated by soft breaks, so those
    labels cannot be styled on their own as it stands. The paragraph is rebuilt as a run per
    line - each a copy of the original, so the font is preserved - with the port labels bold
    and the soft breaks put back between them. The matching observation markers further down
    are already bold in the template, so this brings the two into line.
    """
    from docx.oxml.ns import qn as _qn
    pm = {p._p: p for p in doc.paragraphs}
    want = False
    done = 0
    for el in list(doc.element.body.iterchildren()):
        if el.tag != _qn("w:p"):
            continue
        p = pm.get(el)
        if p is None:
            continue
        style = (p.style.name or "") if p.style is not None else ""
        if style.startswith("Heading"):
            if want:
                break
            want = heading in (p.text or "").upper()
            continue
        if not want:
            continue
        runs = el.findall(_qn("w:r"))
        if not runs:
            continue
        # the paragraph's text, with soft breaks as newlines
        lines = re.split(r"\n", (p.text or ""))
        if not any(l.strip().lower().rstrip(":") in _SURGE_PORT_HEADINGS for l in lines):
            continue
        proto = copy.deepcopy(runs[0])
        for r in runs:
            el.remove(r)
        for i, line in enumerate(lines):
            r = copy.deepcopy(proto)
            for t in r.findall(_qn("w:t")):
                r.remove(t)
            for b in r.findall(_qn("w:br")):
                r.remove(b)
            if i:                                    # soft break before every line but the first
                r.append(r.makeelement(_qn("w:br"), {}))
            te = r.makeelement(_qn("w:t"), {})
            te.text = line
            te.set(_qn("xml:space"), "preserve")
            r.append(te)
            if line.strip().lower().rstrip(":") in _SURGE_PORT_HEADINGS:
                rPr = r.find(_qn("w:rPr"))
                if rPr is None:
                    rPr = r.makeelement(_qn("w:rPr"), {})
                    r.insert(0, rPr)
                if rPr.find(_qn("w:b")) is None:
                    rPr.append(rPr.makeelement(_qn("w:b"), {}))
                done += 1
            el.append(r)
        want = False        # only the procedure's own text block
    return done


def _surge_drop_untested_port(doc, ports):
    """Remove the observation block of a port that was not tested.

    'AC Power Line:' / 'DC Power Line:' belong to the power port and 'Signal Line:' to the
    signal port; with only one port applicable the other's heading (and the matrix inserted
    after it) should not appear at all. Previously both were printed regardless."""
    if not ports:
        return 0
    plan = (("ac power line:", "power"), ("dc power line:", "power"), ("signal line:", "signal"))
    removed = 0
    for marker_text, port in plan:
        if ports.get(port, True):
            continue                      # tested -> keep
        for p in list(doc.paragraphs):
            if (p.text or "").strip().lower() != marker_text:
                continue
            nxt = p._p.getnext()
            if nxt is not None and nxt.tag == qn("w:tbl"):
                nxt.getparent().remove(nxt)     # the matrix inserted for this port
            p._p.getparent().remove(p._p)
            removed += 1
    return removed


#: HARMONIC's spec value area is 2 grid columns; 3 from each gives 6, which divides by
#: 1, 2 and 3 so every per-day split lands on a boundary.
_HARMONIC_VALUE_REGRID = (3, 3)


def _harmonic_finalize(doc, context):
    """HARMONIC reference-format corrections.

      * Ambient Temperature / Relative Humidity / Test Date / Tested by split into the 1-3
        per-day sections the engineer chose, on their existing row.
      * A blank Calibration Due on a populated equipment row prints 'NA'.
      * EUT MODIFICATION RECORD, TEST EQUIPMENT USED and SOFTWARE USED read as centred.
    """
    from .layout import _ce_center_table, _ce_table_header
    meta = (context or {}).get("_harmonic_meta") or {}
    splits = meta.get("row_splits") or []

    tb = _re_spec_table(doc)
    if tb is not None and splits:
        if _re_regrid_value_area(tb, _HARMONIC_VALUE_REGRID):
            for row in splits:
                vals = row.get("values") or []
                if vals:
                    _set_cell_sections(tb, _re_row_by_label(tb, row["needle"]), 1, vals)
            _sync_row_widths(tb)
            for needle in dict.fromkeys(r["needle"] for r in splits):
                _re_fit_row_font(tb, needle)

    _re_fill_missing_na(doc, header_needle="calibration due", value="NA")
    for tb2 in doc.tables:
        hdr = _ce_table_header(tb2)
        if ("equipment name" in hdr and "calibration" in hdr) or \
           ("software name" in hdr and "software version" in hdr) or \
           ("modification state" in hdr and "description" in hdr):
            _ce_center_table(tb2)
    # Extra pictures continue the sequence after the standard slot.
    _re_renumber_photos(doc)


#: VOLTAGEFLICKER's spec value area is 2 grid columns, same as HARMONIC's; 3 from each
#: gives 6, which divides by 1, 2 and 3 so every per-day split lands on a boundary.
_FLICKER_VALUE_REGRID = (3, 3)


def _strip_trailing_empty_paragraphs(doc):
    """Remove the empty paragraphs left at the very END of the body.

    The templates finish with a couple of spare paragraphs after the sign-off table. When
    the last table already reaches the bottom of the page those paragraphs spill over and
    Word shows an extra, blank page. Only genuinely empty ones go, and only from the tail -
    a paragraph holding an image, a page break or the final sectPr is left alone.
    """
    body = doc.element.body
    pm = {p._p: p for p in doc.paragraphs}
    removed = 0
    for el in reversed(list(body.iterchildren())):
        if el.tag == qn("w:sectPr"):
            continue                          # the final section properties
        if el.tag != qn("w:p"):
            break                             # hit real content (a table) - stop
        p = pm.get(el)
        if p is not None and (p.text or "").strip():
            break
        pPr = el.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            break                             # carries a section break
        if el.findall(".//" + qn("w:drawing")) or el.findall(".//" + qn("w:br")):
            break                             # an image or an explicit break
        el.getparent().remove(el)
        removed += 1
    return removed


def _drop_grid_columns(tb, drop):
    """Remove whole grid columns from a table, span-aware.

    A cell covering only removed columns goes; one that straddles keeps the columns it
    still covers with its span reduced. The freed width is handed back to the surviving
    columns in proportion, so the table stays exactly as wide as it was.
    """
    widths = _grid_widths(tb)
    drop = sorted({i for i in drop if 0 <= i < len(widths)})
    if not drop or len(drop) >= len(widths):
        return False
    dropset = set(drop)
    for tr in tb._tbl.findall(qn("w:tr")):
        for tc, g0, span in list(_tc_spans(tr)):
            gone = sum(1 for c in range(g0, g0 + span) if c in dropset)
            if not gone:
                continue
            if gone == span:
                tr.remove(tc)
            else:
                _set_span(tc, span - gone)
    keep = [i for i in range(len(widths)) if i not in dropset]
    total = sum(widths)
    kept = [widths[i] for i in keep]
    scale = total / float(sum(kept)) if sum(kept) else 1.0
    new = [int(w * scale) for w in kept]
    new[-1] += total - sum(new)                 # keep the total exact
    grid = tb._tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        grid.remove(gc)
    for w in new:
        gc = grid.makeelement(qn("w:gridCol"), {})
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    _sync_row_widths(tb)
    return True


#: PFMF's spec value area is 2 grid columns; 3 from each gives 6, so 1/2/3 sections all
#: land on a boundary.
_PFMF_VALUE_REGRID = (3, 3)

#: PFMF TEST OBSERVATION grid: 0 Field Strength, 1 Test Frequency, then the coil columns.
_PFMF_METHOD_COLUMNS = {"proximity": (2, 3, 4, 5),      # 0deg / 90deg / 180deg / 270deg
                        "immersion": (6, 7, 8)}         # X / Y / Z


def _pfmf_finalize(doc, context):
    """PFMF reference-format corrections.

      * TEST OBSERVATION keeps only the coil columns of the selected Test Method
        (Proximity / Immersion / Both), and its three header rows read as bold.
      * Ambient / Humidity / Test Date / Tested by split into the engineer's sections.
      * 1.2 EUT MODIFICATION RECORD, TEST EQUIPMENT USED and SOFTWARE USED read as
        centred; a blank Calibration Due prints 'NA'.
      * 1.4 MONITORING PARAMETERS becomes a bulleted list and the signature loses its
        border.
    """
    from .layout import _ce_center_table, _ce_table_header
    meta = (context or {}).get("_pfmf_meta") or {}
    methods = meta.get("methods") or {"proximity", "immersion"}
    splits = meta.get("row_splits") or []

    tb = _re_spec_table(doc)
    if tb is not None and splits:
        if _re_regrid_value_area(tb, _PFMF_VALUE_REGRID):
            for row in splits:
                vals = row.get("values") or []
                if vals:
                    _set_cell_sections(tb, _re_row_by_label(tb, row["needle"]), 1, vals)
            _sync_row_widths(tb)
            for needle in dict.fromkeys(r["needle"] for r in splits):
                _re_fit_row_font(tb, needle)

    for obs in doc.tables:
        hdr = _ce_table_header(obs)
        if "field strength" not in hdr or "test frequency" not in hdr:
            continue
        drop = []
        for name, cols in _PFMF_METHOD_COLUMNS.items():
            if name not in methods:
                drop.extend(cols)
        if drop:
            _drop_grid_columns(obs, drop)
        # the three stacked heading rows (Coil Orientation / method / orientation)
        for ri in range(min(3, len(obs.rows))):
            _bold_row(obs, ri)
        break

    _bullet_monitoring_parameters(doc)
    _re_fill_missing_na(doc, header_needle="calibration due", value="NA")
    for tb2 in doc.tables:
        hdr = _ce_table_header(tb2)
        if ("equipment name" in hdr and "calibration" in hdr) or \
           ("software name" in hdr and "software version" in hdr) or \
           ("modification state" in hdr and "description" in hdr):
            _ce_center_table(tb2)
    _re_renumber_photos(doc)
    _re_fix_signature(doc)
    # the template's spare paragraphs after the sign-off table would spill to a blank page
    _strip_trailing_empty_paragraphs(doc)


def _crf_prune_photos(doc, ports):
    """Keep only the Test Setup picture for the Test Port(s) actually selected.

    The template ships both slots ('Photo 1: CRF test setup - Power Line' and
    'Photo 2: ... - Signal Line'); the caption names the port, so the unselected one is
    dropped together with its image paragraph. With no port selected nothing is removed,
    so a half-filled draft still renders what it has."""
    if not ports:
        return 0
    wanted = [p.lower() for p in ports]
    body = doc.element.body
    paras = {p._p: p for p in doc.paragraphs}
    removed = 0
    for el in list(body.iterchildren()):
        if el.tag != qn("w:p"):
            continue
        p = paras.get(el)
        if p is None:
            continue
        txt = " ".join((p.text or "").split()).lower()
        if "test setup" not in txt or "photo" not in txt:
            continue
        if any(w in txt for w in wanted):
            continue
        # this caption names the other port: drop it and the image just above it
        prev = el.getprevious()
        el.getparent().remove(el)
        removed += 1
        while prev is not None and prev.tag == qn("w:p"):
            pp = paras.get(prev)
            has_img = bool(prev.findall(".//" + qn("w:drawing")))
            if not has_img and (pp is None or (pp.text or "").strip()):
                break
            gone = prev
            prev = prev.getprevious()
            gone.getparent().remove(gone)
            removed += 1
            if has_img:
                break
    return removed


#: CRF's spec value area is 2 grid columns; 3 from each gives 6, which divides by 1, 2
#: and 3 so every per-day split lands on a boundary.
_CRF_VALUE_REGRID = (3, 3)


def _crf_prune_observation(doc, ports):
    """In TEST OBSERVATION keep only the selected Test Port's block.

    The template ships a 'Power Line:' heading and a 'Signal Line:' heading with the
    observation table under the second, so the unselected port's heading (and whatever
    sits under it, up to the next heading or the table) is removed. With no port selected
    nothing is touched."""
    if not ports:
        return 0
    keep = {p.lower() for p in ports}
    body = doc.element.body
    pm = {p._p: p for p in doc.paragraphs}
    # the run of elements belonging to TEST OBSERVATION, up to its table
    els = list(body.iterchildren())
    try:
        start = next(i for i, el in enumerate(els)
                     if el.tag == qn("w:p") and pm.get(el) is not None
                     and " ".join((pm[el].text or "").split()).upper() == "TEST OBSERVATION")
    except StopIteration:
        return 0
    end = next((i for i in range(start + 1, len(els)) if els[i].tag == qn("w:tbl")), len(els))

    def port_of(el):
        if el.tag != qn("w:p"):
            return None
        p = pm.get(el)
        t = " ".join((p.text or "").split()).lower() if p is not None else ""
        for name in ("power line", "signal line"):
            if t.startswith(name) and t.endswith(":"):
                return name
        return None

    removed = 0
    i = start + 1
    while i < end:
        name = port_of(els[i])
        if name is None or name in keep:
            i += 1
            continue
        # drop this heading and everything under it until the next port heading / the table
        j = i + 1
        while j < end and port_of(els[j]) is None:
            j += 1
        for el in els[i:j]:
            if el.getparent() is not None:
                el.getparent().remove(el)
                removed += 1
        i = j
    return removed


def _crf_finalize(doc, context):
    """CRF reference-format corrections.

      * 1.4 MONITORING PARAMETERS reads as a bulleted list.
      * TEST PROCEDURE starts a line below its heading, with the 'Power Line:' /
        'Signal Line:' lead-ins in bold.
      * Only the selected Test Port's Test Setup picture is kept, and the photos are
        renumbered so the survivor is 'Photo 1'.
      * TEST EQUIPMENT USED and SOFTWARE USED read as centred; a blank Calibration Due
        prints 'NA'. The signature image loses its border.
    """
    from .layout import _ce_center_table, _ce_table_header, _ce_fill_empty_cells
    meta = (context or {}).get("_crf_meta") or {}
    ports = meta.get("ports") or []
    splits = meta.get("row_splits") or []

    tb = _re_spec_table(doc)
    if tb is not None and splits:
        if _re_regrid_value_area(tb, _CRF_VALUE_REGRID):
            for row in splits:
                vals = row.get("values") or []
                if vals:
                    _set_cell_sections(tb, _re_row_by_label(tb, row["needle"]), 1, vals)
            _sync_row_widths(tb)
            for needle in dict.fromkeys(r["needle"] for r in splits):
                _re_fit_row_font(tb, needle)

    _bullet_monitoring_parameters(doc)
    _re_procedure_heading_gap(doc)
    _surge_bold_port_headings(doc)

    _crf_prune_photos(doc, ports)
    _re_renumber_photos(doc)
    _crf_prune_observation(doc, ports)

    _re_fill_missing_na(doc, header_needle="calibration due", value="NA")
    # 1.2 EUT MODIFICATION RECORD: '-' where a cell was left blank, then centred.
    _ce_fill_empty_cells(doc, ("modification state", "description"), "-")
    for tb2 in doc.tables:
        hdr = _ce_table_header(tb2)
        if ("equipment name" in hdr and "calibration" in hdr) or \
           ("software name" in hdr and "software version" in hdr) or \
           ("modification state" in hdr and "description" in hdr) or \
           ("frequency range" in hdr and "coupling method" in hdr):   # 2.4 TEST OBSERVATION
            _ce_center_table(tb2)

    _re_fix_signature(doc)


def _bold_row(tb, row_idx):
    """Bold every run in a table row. Returns the number of runs changed."""
    try:
        row = tb.rows[row_idx]
    except IndexError:
        return 0
    n = 0
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if not run.bold:
                    run.bold = True
                    n += 1
    return n


def _center_row_values(tb, row_idx, skip_label=True):
    """Centre a row's VALUE cells (everything after the label column).

    _ce_center_table() would centre the whole table; here only one row is wanted, so the
    same paragraph alignment is applied to just that row."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    try:
        row = tb.rows[row_idx]
    except IndexError:
        return 0
    n = 0
    for ci, cell in enumerate(row.cells):
        if skip_label and ci == 0:
            continue
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            n += 1
    return n


def _vdips_finalize(doc, context):
    """VOLTAGEDIPS reference-format corrections.

      * Ambient / Humidity / Test Date / Tested by split into the engineer's 1-3 sections.
        The value cell is divided IN PLACE rather than re-gridding the table: unlike the
        other datasheets, this spec table's value area carries the Test Level sub-columns
        (1240/1240/205/1035/1650 twips), and an equal-width re-grid would move them.
      * 1.2 EUT MODIFICATION RECORD, TEST OBSERVATION, TEST EQUIPMENT USED and
        SOFTWARE USED read as centred; a blank Calibration Due prints 'NA'.
      * 1.4 MONITORING PARAMETERS becomes a bulleted list.
      * RESULT: the heading row and the % row are bold, the Met Performance Criteria
        values are centred.
      * The signature image loses its border, and the TEST SETUP PICTURES block loses the
        doubled blank line under the heading.
    """
    from .layout import (_ce_center_table, _ce_table_header, _ce_tune_plot_spacing,
                         collapse_blank_runs)
    meta = (context or {}).get("_vdips_meta") or {}

    tb = _re_spec_table(doc)
    if tb is not None:
        for row in (meta.get("row_splits") or []):
            vals = row.get("values") or []
            if vals:
                _set_cell_sections(tb, _re_row_by_label(tb, row["needle"]), 1, vals)
        _sync_row_widths(tb)
        for needle in dict.fromkeys(r["needle"] for r in (meta.get("row_splits") or [])):
            _re_fit_row_font(tb, needle)

    _re_fill_missing_na(doc, header_needle="calibration due", value="NA")

    for tb2 in doc.tables:
        hdr = _ce_table_header(tb2)
        if ("equipment name" in hdr and "calibration" in hdr) or \
           ("software name" in hdr and "software version" in hdr) or \
           ("modification state" in hdr and "description" in hdr) or \
           ("test level" in hdr and "observation" in hdr):        # TEST OBSERVATION
            _ce_center_table(tb2)
        elif "required performance criteria" in hdr:
            # r0 headings (Voltage Dips / Interruption) and r1, the % row, are bold;
            # the Met Performance Criteria row's values are centred.
            _bold_row(tb2, 0)
            _bold_row(tb2, 1)
            for ri, row in enumerate(tb2.rows):
                if "met performance criteria" in " ".join((row.cells[0].text or "").split()).lower():
                    _center_row_values(tb2, ri)

    _bullet_monitoring_parameters(doc)
    _re_fix_signature(doc)          # no border on the signature image
    collapse_blank_runs(doc)        # the doubled blank line under TEST SETUP PICTURES
    _ce_tune_plot_spacing(doc)      # image -> caption gap, as on the other datasheets
    _re_renumber_photos(doc)


def _bold_header_cells(tb, *labels):
    """Bold the header cells whose text matches one of `labels` (case-insensitive, exact
    after whitespace collapsing).

    Used where a template's header row bolds only its first column: TEST LIMITS bolds
    'Flicker Measurement' but not 'Limit', so the rest are brought up to match. Returns
    the number of cells changed."""
    changed = 0
    try:
        head = tb.rows[0]
    except IndexError:
        return 0
    wanted = {l.strip().lower() for l in labels}
    for cell in head.cells:
        if " ".join((cell.text or "").split()).lower() not in wanted:
            continue
        for para in cell.paragraphs:
            for run in para.runs:
                if not run.bold:
                    run.bold = True
                    changed += 1
    return changed


def _flicker_finalize(doc, context):
    """VOLTAGEFLICKER reference-format corrections.

      * Ambient Temperature / Relative Humidity / Test Date / Tested by split into the 1-3
        per-day sections the engineer chose, on their existing row.
      * A blank Calibration Due on a populated equipment row prints 'NA'.
      * EUT MODIFICATION RECORD, MEASUREMENT UNCERTAINITY, TEST EQUIPMENT USED and
        SOFTWARE USED read as centred.
    """
    from .layout import _ce_center_table, _ce_table_header
    meta = (context or {}).get("_flicker_meta") or {}
    splits = meta.get("row_splits") or []

    tb = _re_spec_table(doc)
    if tb is not None and splits:
        if _re_regrid_value_area(tb, _FLICKER_VALUE_REGRID):
            for row in splits:
                vals = row.get("values") or []
                if vals:
                    _set_cell_sections(tb, _re_row_by_label(tb, row["needle"]), 1, vals)
            _sync_row_widths(tb)
            for needle in dict.fromkeys(r["needle"] for r in splits):
                _re_fit_row_font(tb, needle)

    _re_fill_missing_na(doc, header_needle="calibration due", value="NA")
    for tb2 in doc.tables:
        hdr = _ce_table_header(tb2)
        if ("equipment name" in hdr and "calibration" in hdr) or \
           ("software name" in hdr and "software version" in hdr) or \
           ("modification state" in hdr and "description" in hdr) or \
           ("name of the test" in hdr):          # MEASUREMENT UNCERTAINITY
            _ce_center_table(tb2)
        # Both flicker tables bold only their first heading in the template; the remaining
        # column headings are brought up to match. MEASUREMENT DATA is the one carrying
        # 'Measured Value', TEST LIMITS the one without it.
        elif "flicker measurement" in hdr and "measured value" in hdr:
            _bold_header_cells(tb2, "Measured Value", "Limit")
        elif "flicker measurement" in hdr and "limit" in hdr:
            _bold_header_cells(tb2, "Limit")


def _surge_finalize(doc, context):
    """SURGE reference-format corrections.

      * Ambient Temperature / Relative Humidity / Test Date / Tested by split into the 1-3
        per-day sections the engineer chose, on their existing row.
      * A blank Calibration Due on a populated equipment row prints 'NA'.
      * EUT MODIFICATION RECORD, TEST EQUIPMENT USED and SOFTWARE USED read as centred.
      * Every table prints in black.
    """
    from .layout import _ce_center_table, _ce_table_header
    meta = (context or {}).get("_surge_meta") or {}
    splits = meta.get("row_splits") or []

    tb = _re_spec_table(doc)
    if tb is not None and splits:
        if _re_regrid_value_area(tb, _SURGE_VALUE_REGRID, label_cols=_SURGE_LABEL_COLS):
            for row in splits:
                vals = row.get("values") or []
                if vals:
                    # one value cell per row here, so it is always cell index 1
                    _set_cell_sections(tb, _re_row_by_label(tb, row["needle"]), 1, vals)
            _sync_row_widths(tb)
            for needle in dict.fromkeys(r["needle"] for r in splits):
                _re_fit_row_font(tb, needle)

    _re_fill_missing_na(doc, header_needle="calibration due", value="NA")
    for tb2 in doc.tables:
        hdr = _ce_table_header(tb2)
        if ("equipment name" in hdr and "calibration" in hdr) or \
           ("software name" in hdr and "software version" in hdr) or \
           ("modification state" in hdr and "description" in hdr):
            _ce_center_table(tb2)
    _bullet_monitoring_parameters(doc)
    _surge_bold_port_headings(doc)
    _surge_drop_untested_port(doc, (context or {}).get("_surge_ports"))
    # Extra pictures continue the sequence: 'Photo 3:', 'Photo 4:', ... whatever label the
    # engineer typed. Runs after the untested-port pass so a dropped photo leaves no gap.
    _re_renumber_photos(doc)
    _blacken_all_tables(doc)
    # A sub-section that no longer fits moves whole to the next page rather than breaking
    # across one (TEST SETUP PICTURES was splitting). Same passes CE and RS_RI use, but the
    # blank spacers go first: three of them in the picture block were ~13mm, the difference
    # between its two photos fitting a page and spilling onto a second.
    from .layout import (_ce_keep_subsections, _ce_break_overflowing_subsections,
                         _ce_strip_blanks_before_breaks, collapse_blank_runs,
                         fit_picture_block)
    collapse_blank_runs(doc)
    # Two full-size photos plus the heading and captions come to just over one page, so trim
    # the images by the few percent that keeps the section whole.
    fit_picture_block(doc, "TEST SETUP PICTURES")
    _ce_keep_subsections(doc)
    _ce_break_overflowing_subsections(doc, stop_heading=None)
    _ce_strip_blanks_before_breaks(doc)


def _rs_ri_finalize(doc, context):
    """RS_RI reference-format corrections.

      * Ambient Temperature / Relative Humidity / Test Date can each be split into 1-3
        per-day sections INDEPENDENTLY for the 80M-1G and the 1G-6G column, and 'Tested by'
        across the pair - all staying on their existing row, with the font shrunk per row
        only if a value would otherwise wrap.
      * A blank Calibration Due on a populated equipment row prints 'NA'.
      * MONITORING PARAMETERS is left-aligned instead of justified.
      * The EUT Modification Record's placeholder dashes are centred.
      * TEST OBSERVATION prints in black throughout.
      * No sub-section splits across a page break.
    """
    meta = (context or {}).get("_rs_ri_meta") or {}
    splits = meta.get("row_splits") or []

    tb = _re_spec_table(doc)
    if tb is not None and splits:
        if _re_regrid_value_area(tb, _RS_RI_VALUE_REGRID):
            # Right-to-left: splitting a cell inserts new cells after it, which would shift
            # the index of every cell to its right.
            for row in sorted(splits, key=lambda r: -int(r.get("cell") or 1)):
                vals = row.get("values") or []
                if not vals:
                    continue
                _set_cell_sections(tb, _re_row_by_label(tb, row["needle"]),
                                   int(row.get("cell") or 1), vals)
            # Cell widths must agree with the grid or Word re-lays the table out from the
            # stale values and the label column collapses.
            _sync_row_widths(tb)
            # Six sections across the value area leaves each one narrow; shrink just the
            # rows that were split, and only when a value really does not fit.
            for needle in dict.fromkeys(r["needle"] for r in splits):
                _re_fit_row_font(tb, needle)

    _re_fill_missing_na(doc, header_needle="calibration due", value="NA")
    _rs_ri_unjustify(doc, "MONITORING PARAMETERS")
    _rs_ri_center_placeholders(doc)
    _rs_ri_blacken_observation(doc)
    # 2.6 TEST EQUIPMENT USED and 2.7 SOFTWARE USED read as centred grids, as on CE.
    from .layout import _ce_center_table, _ce_table_header
    for _tb in doc.tables:
        _hdr = _ce_table_header(_tb)
        if ("equipment name" in _hdr and "calibration" in _hdr) or \
           ("software name" in _hdr and "software version" in _hdr):
            _ce_center_table(_tb)
    # One blank line between the TEST PROCEDURE title and the procedure text.
    _re_procedure_heading_gap(doc)
    # A sub-section that no longer fits moves whole to the next page rather than breaking
    # across one (2.3 TEST PROCEDURE and TEST SETUP PICTURES were both splitting). Shares
    # CE's passes: the keep rules first, then measured page breaks for the blocks whose body
    # is a table, which keep-with-next alone cannot hold together.
    from .layout import (_ce_keep_subsections, _ce_break_overflowing_subsections,
                         _ce_strip_blanks_before_breaks)
    _ce_keep_subsections(doc)
    _ce_break_overflowing_subsections(doc, stop_heading=None)   # every sub-section, to the end
    _ce_strip_blanks_before_breaks(doc)


def _re_finalize(doc, context):
    """Apply the RE reference-format corrections that can't be templated."""
    meta = (context or {}).get("_re_meta") or {}
    s30 = meta.get("show_30m_1g", True)
    s16 = meta.get("show_1g_6g", True)
    rot = meta.get("rotation_steps") or []
    days = meta.get("days") or []

    tb = _re_spec_table(doc)
    if tb is not None:
        # 1) Re-grid the value area to 12 equal columns so any 1/2/3 split is exact.
        _re_regrid_value_area(tb)
        # 2) Only the selected frequency range keeps a value column ('Both' keeps two).
        if s30 != s16:
            _re_collapse_single_range(tb, use_second=s16)
        # 3) Turn-table rotation step: one section per applicable family (15deg CISPR /
        #    22.5deg CFR), repeated for each frequency column that is shown, so both
        #    columns read identically (15deg | 22.5deg  |  15deg | 22.5deg).
        if rot:
            per_range = list(rot) * (2 if (s30 and s16) else 1)
            _re_set_row_sections(tb, _re_row_by_label(tb, "rotation step"), per_range)
        # 4) Per-row splits the engineer chose: Ambient Temperature, Relative Humidity,
        #    Test Date and Tested by each get 1, 2 or 3 equal sections, independently.
        #    Applied even for a single section, so the row always shows the engineer's
        #    value (these cells no longer come from template placeholders).
        for row in (meta.get("row_splits") or []):
            vals = row.get("values") or []
            if vals:
                _re_set_row_sections(tb, _re_row_by_label(tb, row["needle"]), vals)
        # 5) Cell widths must match the grid, otherwise Word re-lays the table out and
        #    the label column collapses (pushing the table onto a second page).
        _sync_row_widths(tb)
        # 6) Names are long: if a 'Tested by' value can't fit its section at 11 pt,
        #    shrink JUST that row to the closest size that keeps it on one line.
        _re_fit_row_font(tb, "tested by")

    # 2.4 TEST PROCEDURE: rebuild from the (range-filtered, mapped) text, then make
    # sure the opening line names this datasheet's Basic Standard.
    _re_rebuild_procedure(doc, (context or {}).get("test_procedure"))
    _re_fix_procedure_standard(doc, meta.get("basic_standard"))
    # Prune first (a custom range prunes by position, standard ranges by band name), THEN
    # rename what survived. Renaming first would erase the band names the standard prune
    # matches on; pruning first by band name would delete a stale caption instead of
    # renaming it - hence the position rule for a custom range.
    _re_prune_photos(doc, s30, s16,
                     keep=[p.get("caption") for p in (context or {}).get("re_extra_photos") or []],
                     custom_range=meta.get("custom_range"))
    # A custom Frequency Range is ONE band: rename every surviving caption that still
    # holds a standard band name (Test Setup photos, Functional Check ambient plots, and
    # any stale caption carried in from an older draft).
    _re_relabel_captions_custom_range(doc, meta.get("custom_range"))
    # Upload-driven tables (RE Functional Check): drop the ones with no data, then treat
    # the rest like the measurement tables when re-deriving column widths.
    _upload_tables = [v for v in (context or {}).values()
                      if isinstance(v, dict) and "has_data" in v and "headers" in v]
    _re_drop_empty_upload_tables(doc, _upload_tables)
    # dynamic measurement columns: rebuild the grid to the rendered column count
    _re_sync_meas_table_widths(
        doc, (context or {}).get("measurement_groups"),
        extra_header_sets=[t.get("headers") for t in _upload_tables if t.get("has_data")])
    _re_drop_captionless_photos(doc)  # no upload -> no 'Photo N:' caption over blank space
    _re_renumber_photos(doc)          # 'Photo 1..N' in document order, extras included
    _re_renumber_figures(doc)         # 'Figure 1..N' with no gaps from dropped plot slots
    _re_tidy_caption_whitespace(doc)  # no '_ Vertical' / double spaces in any caption
    _re_fill_missing_na(doc)          # blank Calibration Due -> NA
    _re_align_tables(doc)
    # Only the sign-off block's heading row is centred. The 2.1 TEST SPECIFICATION
    # 'Tested by' row label stays left, like every other label in that table.
    # Runs AFTER _re_align_tables, which would otherwise reset it to left.
    _re_center_signoff_title(doc)
    _re_tighten_image_spacing(doc)     # image->caption gap == table->caption gap
    _re_caption_spacing(doc)
    _re_subsection_spacing(doc)
    # After _re_subsection_spacing, which owns space_BEFORE on every Heading 2 and would
    # not touch this; here we set space_AFTER on the one heading that needs it.
    _re_procedure_heading_gap(doc)
    _re_keep_subsections(doc)
    _re_fix_signature(doc)


def render(code, context, img_keys, img_paths, output_path):
    tpl = DocxTemplate(os.path.join(TPL_DIR, f"{code}.docx"))
    _img_boxes = context.get("_img_boxes") or {}
    for k in img_keys:
        p = img_paths.get(k)
        custom = _img_boxes.get(k)
        box = custom or _box(k, code)
        context[k] = _fit(tpl, p, box, exact=bool(custom)) if (p and os.path.exists(p)) else ""
    if code == "RE":
        for group in context.get("measurement_groups") or []:
            # One entry per plot: the standard Vertical/Horizontal pair for each band the
            # group covers, plus any extra plots the engineer added. A slot with nothing
            # uploaded is dropped so no caption prints over an empty frame.
            kept = []
            for img in group.get("images") or []:
                p = img_paths.get(img["key"])
                if not (p and os.path.exists(p)):
                    continue
                custom = _img_boxes.get(img["key"])
                img["img"] = _fit(tpl, p, custom or _box("img_vertical", code), exact=bool(custom))
                kept.append(img)
            group["images"] = kept
    # Extra test-setup pictures the engineer added: an empty slot (label typed but nothing
    # uploaded, or the row left untouched) prints nothing at all. Outside the RE branch -
    # SURGE uses the same slots, and without this its extras rendered as captions with no
    # image above them.
    if context.get("re_extra_photos"):
        _extras = []
        for photo in context["re_extra_photos"]:
            p = img_paths.get(photo["key"])
            if not (p and os.path.exists(p)):
                continue
            custom = _img_boxes.get(photo["key"])
            photo["img"] = _fit(tpl, p, custom or _box(photo["key"], code), exact=bool(custom))
            _extras.append(photo)
        context["re_extra_photos"] = _extras
    tpl.render(context, autoescape=True)
    
    if code == "RE":
        # Strip manual template breaks after rendering content/applying layouts
        strip_manual_page_breaks(tpl.docx)
        _prune_empty_limit_tables(tpl.docx)   # drop CISPR/FCC limit tables that don't apply
        polish_layout(tpl.docx)
        enforce_arial_fonts(tpl.docx)              # force Arial on all table cell runs (override Calibri)
        enforce_arial_procedure(tpl.docx)          # force Arial on the Test Procedure body text
        enforce_body_arial(tpl.docx)               # body paragraphs + Normal style -> Arial 11
        page_break_before_top_sections(tpl.docx)   # each top-level (Heading 1) section on a new page
        _re_paginate(tpl.docx)                # runs LAST so it wins over polish_layout's keep-with-next
        strip_trailing_blank_paragraphs(tpl.docx)
    else:
        # For HARMONIC and other generic templates, preserve the manual layout/breaks of the template
        enforce_arial_fonts(tpl.docx)
        enforce_arial_procedure(tpl.docx)          # force Arial on the Test Procedure body text
        enforce_body_arial(tpl.docx)               # body paragraphs + Normal style -> Arial 11

    if code == "EFT":
        _eft_insert_observation(tpl.docx, context.get("eft_obs_power"), context.get("eft_obs_signal"))
        _eft_insert_legend(tpl.docx, context.get("eft_obs_legend"))

    if code == "SURGE":
        _surge_insert_observation(tpl.docx, context.get("surge_obs_ac"),
                                  context.get("surge_obs_dc"), context.get("surge_obs_signal"))
        _eft_insert_legend(tpl.docx, context.get("surge_obs_legend"))

    if code == "PFMF":
        # Replace the static A/B/C/D observation legend with one line per unique
        # value the engineer actually used in the observation grid.
        _eft_insert_legend(tpl.docx, context.get("pfmf_obs_legend"))

    if code in ("RS_RI", "ESD", "CRF", "VOLTAGEDIPS"):
        # Same replacement using the generic per-code legend the form posts
        # (obs_legend_code[]/obs_legend_desc[]); left as-is when no code was selected.
        _eft_insert_legend(tpl.docx, context.get("obs_legend"))

    # Pagination polish for the rebuilt immunity datasheets (no manual breaks in
    # their templates): rows never split across a page, small tables stay whole,
    # long tables repeat their header, and section headings never dangle at a page
    # bottom. Runs after EFT's observation table is inserted so it's covered too.
    if code in ("EFT", "VOLTAGEDIPS", "SURGE"):
        polish_layout(tpl.docx)

    if code == "SURGE":
        enforce_arial_fonts(tpl.docx)                # Arial on the freshly-inserted observation cells
        shrink_wide_obs_tables(tpl.docx)             # 17-col matrices can't be 11pt on a portrait page

    # Every major section starts on a new page, and the final block (2.6 TEST
    # EQUIPMENT USED / 2.7 SOFTWARE USED / 2.8 RESULT) is kept together on the last
    # page. RE has its own paginator (_re_paginate) so it is excluded.
    if code != "RE":
        paginate_generic_datasheet(tpl.docx)

    _add_image_borders(tpl.docx)                     # thin black border on every image
    if code == "RE":
        # RE reference-format corrections (runs LAST: it strips the signature border
        # that _add_image_borders adds, and owns the spec-table column/section layout).
        _re_finalize(tpl.docx, context)
    elif code == "RS_RI":
        _rs_ri_finalize(tpl.docx, context)
    elif code == "SURGE":
        _surge_finalize(tpl.docx, context)
    elif code == "HARMONIC":
        _harmonic_finalize(tpl.docx, context)
    elif code == "VOLTAGEFLICKER":
        _flicker_finalize(tpl.docx, context)
    elif code == "VOLTAGEDIPS":
        _vdips_finalize(tpl.docx, context)
    elif code == "CRF":
        _crf_finalize(tpl.docx, context)
    elif code == "PFMF":
        _pfmf_finalize(tpl.docx, context)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    tpl.save(output_path)
    return output_path
