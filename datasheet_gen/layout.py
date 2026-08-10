"""Post-render layout polish + human-looking checkbox runs for generated datasheets.

Why this exists (all observed on real generated documents):
  * section headings were orphaned at a page bottom while their image landed on
    the next page;
  * the sign-off table split across pages leaving a lone "Date" row on an
    otherwise empty last page;
  * images were left-aligned while their captions are centered;
  * justified paragraphs containing soft line-breaks (textarea "\n") stretched
    words across the whole line;
  * dynamic measurement tables had wildly uneven column widths.

`polish_layout(doc)` fixes those on the rendered python-docx Document, and
`human_checkbox(...)` renders checkboxes with a crossed box (☒) on the
selected option for `{{r ... }}` placeholders.
"""
import re
from xml.sax.saxutils import escape

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

CAPTION_RE = re.compile(r"^\s*(Photo|Figure)\s*\d*\s*:", re.I)


def _is_heading(text):
    """Section-heading heuristic: short, all letters uppercase (e.g.
    '2. SURGE IMMUNITY TEST', 'TEST SETUP PICTURES', '2.9. RESULT')."""
    t = text.strip()
    if not 3 < len(t) < 80:
        return False
    letters = [c for c in t if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


# --------------------------------------------------------------------------
# Human-ticked checkboxes ({{r ... }} placeholders)
# --------------------------------------------------------------------------

class RunsXml:
    """Raw WordprocessingML runs, safe under Jinja autoescape (docxtpl {{r }})."""

    def __init__(self, xml=""):
        self.xml = xml

    def add(self, xml):
        self.xml += xml
        return self

    def __str__(self):
        return self.xml

    def __html__(self):  # keeps Jinja autoescape from escaping the XML
        return self.xml


def _match(value, option):
    """Tolerant match: 'A' == 'Class A' == 'class a'."""
    v = (value or "").strip().lower()
    o = option.strip().lower()
    if not v:
        return False
    return v == o or v in o or o in v or v == o.split()[-1] or o.split()[-1] == v


def _box_run(checked, size=22):
    """The checkbox glyph. Uses a single run with ballot box characters."""
    color = '<w:color w:val="000000"/>' if checked else ""
    char = "☒" if checked else "☐"
    return (
        '<w:r><w:rPr>'
        '<w:rFonts w:ascii="Segoe UI Symbol" w:hAnsi="Segoe UI Symbol" w:cs="Segoe UI Symbol"/>'
        f'{color}<w:sz w:val="{size}"/><w:szCs w:val="{size}"/>'
        f'</w:rPr><w:t xml:space="preserve">{char}</w:t></w:r>'
    )


def _label_run(text, size=22):
    return (
        f'<w:r><w:rPr><w:sz w:val="{size}"/><w:szCs w:val="{size}"/></w:rPr>'
        f'<w:t xml:space="preserve">{escape(text)}</w:t></w:r>'
    )


def human_checkbox(value, options, size=22):
    """Render options as checkboxes with a human-looking cross on the selected one.

    Returns a RunsXml for a `{{r key }}` placeholder. With no/unknown value all
    boxes stay empty (like the blank paper form).
    """
    rt = RunsXml()
    for i, opt in enumerate(options):
        checked = _match(value, str(opt))
        rt.add(_box_run(checked, size))
        sep = "    " if i < len(options) - 1 else ""
        rt.add(_label_run(" " + str(opt) + sep, size))
    return rt


def exact_checkbox(value, options, size=22):
    """Render options as checkboxes, ticking the ones LISTED in `value`.

    `value` is a separated list ('0°, 90°, 270°'); each entry must equal an option exactly.
    human_checkbox() cannot be used for these: it matches on substrings, so a value of just
    '0°' would also tick '90°' and '180°' because '0°' is a substring of both.
    """
    picked = {p.strip().lower() for p in re.split(r"[,;/|]|\band\b", str(value or "")) if p.strip()}
    rt = RunsXml()
    for i, opt in enumerate(options):
        rt.add(_box_run(str(opt).strip().lower() in picked, size))
        sep = "    " if i < len(options) - 1 else ""
        rt.add(_label_run(" " + str(opt) + sep, size))
    return rt


def cumulative_checkbox(level, options, size=22):
    """Tick every option up to and including `level` (options given in ascending
    order). Used for Surge / EFT test-voltage rows where the level is DERIVED from
    the standard and selecting e.g. ±1 kV implies ±0.5 kV was also applied — so both
    boxes are ticked. A 'Custom' level ticks only the Custom box; a blank/unknown
    level ticks nothing. Returns a RunsXml for a `{{r key }}` placeholder.
    """
    rt = RunsXml()
    sel = -1
    for i, opt in enumerate(options):
        if _match(level, str(opt)):
            sel = i
            break
    sel_is_custom = sel >= 0 and "custom" in str(options[sel]).strip().lower()
    for i, opt in enumerate(options):
        opt_is_custom = "custom" in str(opt).strip().lower()
        if sel_is_custom:
            checked = (i == sel)                      # only the Custom box
        else:
            checked = (sel >= 0 and i <= sel and not opt_is_custom)
        rt.add(_box_run(checked, size))
        sep = "    " if i < len(options) - 1 else ""
        rt.add(_label_run(" " + str(opt) + sep, size))
    return rt


# --------------------------------------------------------------------------
# Post-render layout polish
# --------------------------------------------------------------------------

def _text(p):
    return "".join(t.text or "" for t in p._p.iter(qn("w:t")))


def _has_image(p):
    return bool(p._p.findall(".//" + qn("w:drawing")) or p._p.findall(".//" + qn("w:pict")))


def _soft_breaks(p):
    return [b for b in p._p.findall(".//" + qn("w:br")) if b.get(qn("w:type")) != "page"]


def _keep_with_next(p):
    p.paragraph_format.keep_with_next = True


def _row_cant_split(tr):
    trPr = tr._tr.get_or_add_trPr()
    if not trPr.findall(qn("w:cantSplit")):
        trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))


def _row_is_header(tr):
    """Heuristic: a data-table header row has 2+ bold runs."""
    bold = 0
    for tc in tr._tr.findall(qn("w:tc")):
        for b in tc.findall(".//" + qn("w:b")):
            bold += 1
    return bold >= 2


def _mark_header_row(tr):
    trPr = tr._tr.get_or_add_trPr()
    if not trPr.findall(qn("w:tblHeader")):
        trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))


def _equalize_columns(tbl):
    """Give a plain (merge-free) wide table equal column widths."""
    el = tbl._tbl
    if el.findall(".//" + qn("w:gridSpan")) or el.findall(".//" + qn("w:vMerge")):
        return
    grid = el.find(qn("w:tblGrid"))
    if grid is None:
        return
    cols = grid.findall(qn("w:gridCol"))
    if len(cols) < 6:
        return
    widths = [int(c.get(qn("w:w")) or 0) for c in cols]
    total = sum(widths)
    if not total:
        return
    each = total // len(cols)
    for c in cols:
        c.set(qn("w:w"), str(each))
    for tc in el.findall(".//" + qn("w:tc")):
        tcW = tc.find(qn("w:tcPr") + "/" + qn("w:tcW"))
        if tcW is not None:
            tcW.set(qn("w:w"), str(each))
            tcW.set(qn("w:type"), "dxa")


def _keep_row(tr_obj):
    for cell in tr_obj.cells:
        for cp in cell.paragraphs:
            _keep_with_next(cp)


def fit_picture_block(doc, heading, floor_mm=70, safety=0.97):
    """Shrink a picture section's images just enough that the whole section fits one page.

    Two 95mm photos plus a heading and two captions come to slightly MORE than the text area,
    so Word splits the section - the heading and first photo on one page, the second photo on
    the next. Keeping them together needs less content, and the only compressible part is the
    images. Each is scaled by the same factor (aspect preserved) and never below `floor_mm`
    tall, so a section that could only fit by becoming unreadable is left to split instead.

    `safety` trims the target a little below the measured text area, because the estimate is
    a few mm optimistic - Word's real area is smaller than page-minus-margins-minus-chrome.
    Returns the scale applied (1.0 when nothing was needed).
    """
    section = doc.sections[0]
    usable = int(_ce_usable_height(section) * safety)
    width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
    pm = {p._p: p for p in doc.paragraphs}

    block, want = [], False
    for el in doc.element.body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        p = pm.get(el)
        if p is None:
            continue
        style = (p.style.name or "") if p.style is not None else ""
        if style.startswith("Heading"):
            if want:
                break
            want = heading.upper() in (p.text or "").upper()
            if want:
                block.append(p)
            continue
        if want:
            block.append(p)
    if not block:
        return 1.0

    exts = [e for p in block for e in p._p.iter(qn("wp:extent"))]
    if not exts:
        return 1.0
    total = sum(_ce_para_height(p, width) for p in block)
    if total <= usable:
        return 1.0

    img_total = sum(int(e.get("cy") or 0) // _EMU_PER_TWIP for e in exts)
    other = total - img_total
    if img_total <= 0:
        return 1.0
    scale = (usable - other) / float(img_total)
    floor_twips = int(floor_mm * 56.7)
    smallest = min(int(e.get("cy") or 0) // _EMU_PER_TWIP for e in exts)
    if scale <= 0 or smallest * scale < floor_twips:
        return 1.0                       # cannot fit without going too small; let it split
    for e in exts:
        cx, cy = int(e.get("cx") or 0), int(e.get("cy") or 0)
        e.set("cx", str(max(1, int(cx * scale))))
        e.set("cy", str(max(1, int(cy * scale))))
        # the shape's own extent must track the drawing's, or Word crops it
        parent = e.getparent()
        for sp in parent.iter(qn("a:ext")) if parent is not None else ():
            sp.set("cx", str(max(1, int(int(sp.get("cx") or 0) * scale))))
            sp.set("cy", str(max(1, int(int(sp.get("cy") or 0) * scale))))
    return scale


def collapse_blank_runs(doc):
    """Collapse runs of consecutive empty paragraphs to a single one.

    Template spacers add up: SURGE's TEST SETUP PICTURES carried three of them, ~13mm, which
    was the difference between its two photos fitting on one page and spilling onto a second.
    Table-aware - a table between two blanks breaks the run, so the single spacer that
    separates a table from the next heading survives. (CE does this inline in
    ce_finalize_layout; this is the same rule for the schema-driven datasheets.)"""
    body = doc.element.body
    w_p = qn("w:p")
    prev_blank = False
    removed = 0
    for el in list(body):
        if el.tag != w_p:
            prev_blank = False
            continue
        txt = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        has_img = el.findall(".//" + qn("w:drawing")) or el.findall(".//" + qn("w:pict"))
        has_sectpr = el.find(".//" + qn("w:sectPr")) is not None
        blank = (not txt) and (not has_img) and (not has_sectpr)
        if blank and prev_blank:
            body.remove(el)
            removed += 1
        else:
            prev_blank = blank
    return removed


def _ce_strip_blanks_before_breaks(doc):
    """Delete the empty spacer paragraphs that sit immediately before a forced page break.

    On a fresh page such a spacer contributes nothing, but when the PREVIOUS page is
    exactly full it does not fit on it either - so it flows onto a page of its own, and the
    forced break then starts the real content on the page after that. The reader sees a
    blank sheet. That is what put an empty page 4 between 1.5 AMBIENT (which now fills page
    3 exactly) and 2 CONDUCTED EMISSION TEST.

    Must run AFTER every page_break_before has been assigned, so it sees them all.
    """
    removed = 0
    for p in list(doc.paragraphs):
        if p.paragraph_format.page_break_before:
            before = len(p._p.getparent().findall(qn("w:p")))
            _remove_blank_spacers_before(p)
            removed += before - len(p._p.getparent().findall(qn("w:p")))
    return removed


def _ce_tune_plot_spacing(doc, label_gap_pt=0, caption_gap_pt=10):
    """Tighten the gap above a plot and open one below its caption.

    Two spacing problems the reference layout does not have:

      * a "Line:" / "Neutral:" label sat ~22pt above its plot, because the label inherits
        the document default spacing (~8pt after, 1.08 line spacing). The label belongs TO
        the image directly beneath it, so its space_after goes to `label_gap_pt`.
      * a "Figure N:" caption sat flush against the table below it - something in the
        pipeline zeroes space_after on caption paragraphs, overriding the Caption style's
        own 10pt. Restored to `caption_gap_pt` wherever a table follows the caption.

    Only paragraphs in those two positions are touched, so nothing else moves.
    """
    from docx.shared import Pt as _Pt
    body = doc.element.body
    pm = {p._p: p for p in doc.paragraphs}
    label = re.compile(r"^(line|neutral)\s*:$", re.I)
    seq = list(body.iterchildren())
    tightened = opened = 0
    for i, el in enumerate(seq):
        if el.tag != qn("w:p"):
            continue
        p = pm.get(el)
        if p is None:
            continue
        nxt = seq[i + 1] if i + 1 < len(seq) else None
        if nxt is None:
            continue
        txt = (_text(p) or "").strip()
        nxt_p = pm.get(nxt) if nxt.tag == qn("w:p") else None
        # label immediately above its plot -> pull the plot up to it
        if label.match(txt) and nxt_p is not None and _has_image(nxt_p):
            p.paragraph_format.space_after = _Pt(label_gap_pt)
            nxt_p.paragraph_format.space_before = _Pt(0)
            tightened += 1
        # "Figure N:" caption immediately above its data table -> give the table room
        elif nxt.tag == qn("w:tbl") and txt.upper().startswith("FIGURE ") and ":" in txt:
            p.paragraph_format.space_after = _Pt(caption_gap_pt)
            opened += 1
    return tightened, opened


def _ce_table_header(tbl):
    """The table's header row as one lowercased string, for matching."""
    try:
        return " ".join((c.text or "").strip() for c in tbl.rows[0].cells).lower()
    except (IndexError, AttributeError):
        return ""


def _ce_center_table(tbl, vertical=True):
    """Centre every cell's content horizontally (and vertically, to match the rest of the
    document, whose cells already carry vAlign=center)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not vertical:
                continue
            tcPr = cell._tc.get_or_add_tcPr()
            if tcPr.find(qn("w:vAlign")) is None:
                va = tcPr.makeelement(qn("w:vAlign"), {})
                va.set(qn("w:val"), "center")
                tcPr.append(va)


#: Which CE tables get centred content, matched on their header row. "frequency (mhz)"
#: alone is ambiguous - it opens BOTH the Test Limits grid and the Measurement Data grid -
#: so each entry names a second word that only its own table carries.
_CE_CENTERED_TABLES = (
    ("voltage limits",),                 # 2.3 TEST LIMITS
    ("q-peak", "margin"),                # 2.5 MEASUREMENT DATA (Line and Neutral)
    ("equipment name", "calibration"),   # 2.7 TEST EQUIPMENT USED
    ("software name", "software version"),  # 2.8 SOFTWARE USED
    ("modification state",),             # 1.2 EUT MODIFICATION RECORD
)


def _ce_center_tables(doc):
    """Centre the content of the CE tables that the reference layout shows centred."""
    n = 0
    for tbl in doc.tables:
        hdr = _ce_table_header(tbl)
        if not hdr:
            continue
        if any(all(w in hdr for w in words) for words in _CE_CENTERED_TABLES):
            _ce_center_table(tbl)
            n += 1
    return n


def _ce_fill_empty_cells(doc, header_words, placeholder="-"):
    """Put `placeholder` in every empty cell of the matched table, so a partly-filled row
    reads as "nothing recorded" rather than looking unfinished.

    Only rows that carry SOME content are touched: a wholly empty spare row is left alone
    rather than filled with a line of dashes. The header row is never touched.
    """
    filled = 0
    for tbl in doc.tables:
        hdr = _ce_table_header(tbl)
        if not hdr or not all(w in hdr for w in header_words):
            continue
        for row in tbl.rows[1:]:
            cells = list(row.cells)
            if not any((c.text or "").strip() for c in cells):
                continue                       # untouched blank row
            for cell in cells:
                if (cell.text or "").strip():
                    continue
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                if p.runs:
                    p.runs[0].text = placeholder
                    for extra in p.runs[1:]:
                        extra.text = ""
                else:
                    p.add_run(placeholder)
                filled += 1
    return filled


def _ce_drop_empty_image_slots(doc):
    """Remove the placeholder paragraph of a 1.4/1.5 plot slot that had no upload.

    The template pairs each "Line:" / "Neutral:" label with a paragraph holding
    {{ func_line }} / {{ ambient_neutral }} / ... . When the engineer uploads nothing those
    render as an empty string, leaving a full-height empty paragraph behind. Four of them
    (Functional Check + Ambient) are enough to push the tail of section 1 past the bottom
    of the page, producing a page that LOOKS blank - and because section 2 forces its own
    page break, the reader sees a blank sheet between 1.5 AMBIENT and 2 CONDUCTED EMISSION
    TEST. Matched only where the paragraph directly above is a bare "Line:"/"Neutral:"
    label, so a slot that DID get an image (its paragraph holds a drawing, hence is not
    empty) and ordinary spacing elsewhere are both untouched.
    """
    body = doc.element.body
    pm = {p._p: p for p in doc.paragraphs}
    label = re.compile(r"^(line|neutral)\s*:$", re.I)
    prev_label = False
    dropped = 0
    for el in list(body.iterchildren()):
        if el.tag != qn("w:p"):
            prev_label = False
            continue
        p = pm.get(el)
        if p is None:
            continue
        txt = (_text(p) or "").strip()
        blank = (not txt) and not _has_image(p)
        if blank and prev_label:
            body.remove(el)
            dropped += 1
            prev_label = False       # only the one slot paragraph per label
            continue
        prev_label = bool(label.match(txt))
    return dropped


#: One 11pt single-spaced line, in twips.
_CE_LINE_TWIPS = 240
#: Vertical cell padding a table row adds on top of its text.
_CE_CELL_PAD_TWIPS = 60


def _ce_text_twips(text, size_pt=11):
    """Rendered width of `text` in twips, from real Arial metrics."""
    if not text:
        return 0
    from .generic_generator import _text_width_em
    return int(_text_width_em(text) * size_pt * 20)


#: EMU per twip (914400 EMU/inch / 1440 twips/inch).
_EMU_PER_TWIP = 635


def _para_spacing(p):
    """(space_before, space_after) in twips, falling back to the paragraph's STYLE.

    Most paragraphs set no spacing of their own and inherit it - the Caption style carries
    10pt after. Reading only the direct formatting counted those as 0 and under-estimated a
    picture block by ~7mm, which was enough to make a section that does not fit look as
    though it does."""
    out = []
    for attr in ("space_before", "space_after"):
        v = getattr(p.paragraph_format, attr, None)
        if v is None:
            style = p.style
            seen = 0
            while style is not None and seen < 5:      # guard against a style cycle
                v = getattr(style.paragraph_format, attr, None)
                if v is not None:
                    break
                style = getattr(style, "base_style", None)
                seen += 1
        out.append(v.twips if v is not None else 0)
    return out[0], out[1]


def _ce_para_height(p, width_twips):
    """Estimated laid-out height of a body paragraph, including wrapping and spacing.

    A paragraph holding an inline image is as tall as the image, not as a line of text -
    without this a 95mm plot is counted as ~4mm and any section containing one is judged to
    fit when it cannot."""
    before, after = _para_spacing(p)

    img = 0
    for ext in p._p.iter(qn("wp:extent")):
        img = max(img, int(ext.get("cy") or 0) // _EMU_PER_TWIP)
    if img:
        return img + before + after

    txt = (_text(p) or "").strip()
    lines = 1
    if txt and width_twips > 0:
        w = _ce_text_twips(txt)
        lines = max(1, -(-w // width_twips))          # ceil
    lines += len(_soft_breaks(p))                     # manual line breaks add lines
    if (p.style.name or "").startswith("Heading") and not after:
        after = 120                                   # heading styles carry ~6pt after
    return lines * _CE_LINE_TWIPS + before + after


def _ce_cell_size_pt(tc, default=11.0):
    """Font size of a cell's text in points, read from w:sz (half-points).

    Rows whose values were shrunk to fit (RE and RS_RI drop the split rows to as little as
    7pt) are correspondingly shorter; assuming 11pt everywhere over-estimated a 21-row spec
    table by ~30%, which was enough to make it look taller than a page when it is not."""
    sizes = [int(s.get(qn("w:val")) or 0) for s in tc.iter(qn("w:sz"))]
    sizes = [s / 2.0 for s in sizes if s]
    return max(sizes) if sizes else default


def _ce_table_height(tbl):
    """Estimated laid-out height of a table: per row, the larger of its declared height
    and the tallest wrapped cell, measured at each cell's real font size."""
    el = tbl._tbl
    grid = el.find(qn("w:tblGrid"))
    cols = [int(g.get(qn("w:w")) or 0) for g in grid.findall(qn("w:gridCol"))] if grid is not None else []
    total = 0
    for tr in el.findall(qn("w:tr")):
        trPr = tr.find(qn("w:trPr"))
        declared = 0
        if trPr is not None:
            h = trPr.find(qn("w:trHeight"))
            if h is not None:
                declared = int(h.get(qn("w:val")) or 0)
        tallest = _CE_LINE_TWIPS
        gi = 0
        for tc in tr.findall(qn("w:tc")):
            span = 1
            gs = tc.find(qn("w:tcPr") + "/" + qn("w:gridSpan"))
            if gs is not None:
                span = max(1, int(gs.get(qn("w:val")) or 1))
            cw = sum(cols[gi:gi + span]) if cols else 0
            gi += span
            txt = "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
            if not txt:
                continue
            pt = _ce_cell_size_pt(tc)
            lines = 1
            if cw > 0:
                lines = max(1, -(-_ce_text_twips(txt, pt) // cw))
            tallest = max(tallest, int(lines * pt * 20))
        total += max(declared, tallest + _CE_CELL_PAD_TWIPS)
    return total


def _ce_usable_height(section):
    """Height available to BODY text on one page, in twips.

    Not simply page minus margins: this template's running header (logo + title block) is
    far taller than the top margin, so Word pushes body text down to
    header_distance + header_height and shrinks the text area accordingly. Measured in
    Word, body text starts ~94pt below the top margin - ignoring that over-estimates the
    usable height by ~33mm and the caller then fails to break where it must.
    """
    page = section.page_height.twips
    top = section.top_margin.twips
    bottom = section.bottom_margin.twips
    width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips

    def part_height(part, distance):
        if part is None:
            return 0
        h = 0
        try:
            for p in part.paragraphs:
                h += _ce_para_height(p, width)
            for tb in part.tables:
                h += _ce_table_height(tb)
        except (AttributeError, KeyError):
            return 0
        return (distance.twips if distance is not None else 0) + h

    top_used = max(top, part_height(getattr(section, "header", None), section.header_distance))
    bottom_used = max(bottom, part_height(getattr(section, "footer", None), section.footer_distance))
    return max(1, page - top_used - bottom_used)


def _ce_break_overflowing_subsections(doc, stop_heading="test procedure", slack=1.10):
    """Give a sub-section its own page when it would otherwise straddle a page break.

    keep-with-next is not enough for a sub-section whose body is a TABLE: Word will break a
    table BETWEEN rows no matter what keepNext says on the cell paragraphs (2.3 TEST LIMITS
    was landing with its heading, its sentence and the two header rows on one page and the
    three data rows on the next). The only reliable remedy is an explicit page break, so
    this measures each sub-section and forces one when it will not fit in the space left.

    Runs from the last top-level section up to `stop_heading`; pass None to cover every
    sub-section to the end of the document (CE stops early because its tail already carries
    explicit breaks; RS_RI has no such tail and wants the lot). `slack` over-estimates
    heights slightly, so a borderline block breaks early rather than splitting.
    """
    section = doc.sections[0]
    usable = int(_ce_usable_height(section) / slack)
    width = int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)

    body = doc.element.body
    pm = {p._p: p for p in doc.paragraphs}
    tm = {t._tbl: t for t in doc.tables}

    # the span: from the LAST Heading 1 that starts a page, to stop_heading
    seq = list(body.iterchildren())
    start = None
    for i, el in enumerate(seq):
        if el.tag == qn("w:p"):
            p = pm.get(el)
            if p is not None and (p.style.name or "").startswith("Heading 1"):
                start = i
    if start is None:
        return 0
    end = len(seq)
    if stop_heading:
        for i in range(start + 1, len(seq)):
            el = seq[i]
            if el.tag == qn("w:p"):
                p = pm.get(el)
                if p is not None and (_text(p) or "").strip().lower() == stop_heading:
                    end = i
                    break

    # group [heading, ...members] at each Heading 2; anything before the first one belongs
    # to the section heading itself
    groups, cur = [], []
    for el in seq[start:end]:
        if el.tag == qn("w:p"):
            p = pm.get(el)
            if p is not None and (p.style.name or "").startswith("Heading 2"):
                groups.append(cur)
                cur = [el]
                continue
        cur.append(el)
    groups.append(cur)

    def height(group):
        h = 0
        for el in group:
            if el.tag == qn("w:p"):
                p = pm.get(el)
                if p is not None:
                    h += _ce_para_height(p, width)
            elif el.tag == qn("w:tbl"):
                tb = tm.get(el)
                if tb is not None:
                    h += _ce_table_height(tb)
        return h

    used, breaks = 0, 0
    for gi, group in enumerate(groups):
        h = height(group)
        head = pm.get(group[0]) if group and group[0].tag == qn("w:p") else None
        is_sub = bool(head and (head.style.name or "").startswith("Heading 2"))
        # gi == 0 is the section heading's own group: it always starts the page
        if gi and is_sub and used + h > usable:
            if h <= usable:
                head.paragraph_format.page_break_before = True
                used = h
                breaks += 1
                continue
            used = (used + h) % usable      # too tall for any page: let it split, resync
            continue
        used += h
    return breaks


def _ce_keep_subsections(doc, max_lines=14):
    """Keep each Heading-2 sub-section whole: if it cannot fit in the space left on the
    page, Word moves the ENTIRE block to the next page instead of splitting it.

    Same intent as RE's _re_keep_subsections, but a CE sub-section is often mostly TABLE
    (2.3 TEST LIMITS is a heading, one sentence and a 4-row grid), and paragraph
    keep-with-next says nothing about a table. So for every table in the block this also:
      * marks each row cantSplit, so no single row breaks across pages, and
      * sets keepNext on every row but the last, so the table cannot break BETWEEN rows.

    Blocks taller than `max_lines` are skipped: gluing something that cannot fit a page
    makes Word give up and split it anyway, having first wasted the page above it.
    Counts a table row as one line, which is what a one-line cell occupies.
    """
    body = doc.element.body
    pm = {p._p: p for p in doc.paragraphs}
    tm = {t._tbl: t for t in doc.tables}

    # split the body into [heading, ...members] blocks at every Heading 1/2
    blocks, cur = [], None
    for el in body.iterchildren():
        if el.tag == qn("w:p"):
            p = pm.get(el)
            if p is None:
                continue
            style = p.style.name or ""
            if style.startswith("Heading 1"):
                cur = None                      # a top-level section starts a fresh page anyway
                continue
            if style.startswith("Heading 2"):
                cur = [el]
                blocks.append(cur)
                continue
        if cur is not None and el.tag in (qn("w:p"), qn("w:tbl")):
            cur.append(el)

    glued = 0
    for block in blocks:
        # estimated height, and bail out on anything that owns a whole page already
        lines, has_img = 0, False
        for el in block:
            if el.tag == qn("w:p"):
                p = pm.get(el)
                if p is not None and _has_image(p):
                    has_img = True
                lines += 1
            else:
                tb = tm.get(el)
                lines += len(tb.rows) if tb is not None else 3
        if has_img or lines > max_lines:
            continue                            # image blocks are paginated separately

        for i, el in enumerate(block):
            last = (i == len(block) - 1)
            if el.tag == qn("w:p"):
                p = pm.get(el)
                if p is None:
                    continue
                p.paragraph_format.keep_together = True
                if not last:
                    p.paragraph_format.keep_with_next = True
            else:
                tb = tm.get(el)
                if tb is None:
                    continue
                rows = tb.rows
                for r in rows:
                    _row_cant_split(r)
                # bind row->row; the final row is left free unless the block continues
                for r in (rows[:-1] if last else rows):
                    _keep_row(r)
        glued += 1
    return glued


def page_break_before_top_sections(doc):
    """Force every top-level section (a 'Heading 1' paragraph) after the first to
    begin on a new page. Uses the paragraph's page-break-before property, which is
    reflow-safe (the heading stays pinned to the top of a fresh page) rather than a
    manual break run that could drift. The first top-level section is left in place
    so the document does not open with a blank page."""
    first = True
    for p in doc.paragraphs:
        name = (p.style.name if p.style is not None else "") or ""
        if name.strip().lower() in ("heading 1", "heading1"):
            if first:
                first = False
            else:
                p.paragraph_format.page_break_before = True


def _remove_blank_spacers_before(para):
    """Delete empty paragraphs immediately preceding `para` (template spacers used
    to push a section onto a new page — redundant once a page break is forced, and
    a cause of blank pages). Stops at the first paragraph that has text, an image,
    or section properties (a sectPr must never be removed)."""
    prev = para._p.getprevious()
    while prev is not None and prev.tag == qn("w:p"):
        txt = "".join(t.text or "" for t in prev.iter(qn("w:t"))).strip()
        has_img = prev.find(".//" + qn("w:drawing")) is not None
        has_sectpr = prev.find(".//" + qn("w:sectPr")) is not None
        if txt or has_img or has_sectpr:
            break
        nxt = prev.getprevious()
        prev.getparent().remove(prev)
        prev = nxt


def paginate_generic_datasheet(doc, last_block_heading="TEST EQUIPMENT USED"):
    """Pagination for the generic immunity datasheets (RS_RI / ESD / CRF / PFMF /
    HARMONIC / VOLTAGEFLICKER / VOLTAGEDIPS / EFT / SURGE).

      (1) Every major section (a 'Heading 1' paragraph) after the first starts on
          a NEW page, even if the previous page has room — so section 2 never
          continues under section 1.
      (3) The final subsections (from `last_block_heading` to the end: TEST
          EQUIPMENT USED / SOFTWARE USED / RESULT, i.e. 2.6 / 2.7 / 2.8) are pushed
          onto a fresh page and glued together, so they always land together on
          the LAST page instead of drifting up under the setup photos.

    Reflow-safe: uses page-break-before + keep-with-next + cantSplit only. Any
    manual break runs and the template's blank 'spacer' paragraphs that used to
    push a section down are removed first, so forcing the break never leaves an
    empty page in between."""
    # Remove manual run-level page breaks so they cannot double up with the
    # page-break-before set below (a cause of blank pages).
    for br in list(doc.element.body.iter(qn("w:br"))):
        if br.get(qn("w:type")) == "page":
            br.getparent().remove(br)

    # (1) each major section (Heading 1 after the first) begins a new page; drop the
    #     blank spacer paragraphs before it so no empty page appears in between.
    first = True
    for p in list(doc.paragraphs):
        name = (p.style.name if p.style is not None else "") or ""
        if name.strip().lower() in ("heading 1", "heading1"):
            if first:
                first = False
            else:
                _remove_blank_spacers_before(p)
                p.paragraph_format.page_break_before = True

    # (3) push the final block onto a fresh (last) page, dropping any spacer
    #     paragraphs before it, then keep the whole block together.
    target = last_block_heading.strip().upper()
    block_head = None
    for p in doc.paragraphs:
        if _text(p).strip().upper().startswith(target):
            block_head = p
            break
    if block_head is None:
        return
    _remove_blank_spacers_before(block_head)
    block_head.paragraph_format.page_break_before = True

    # recompute element positions AFTER the removals for the keep-together pass
    body = list(doc.element.body)
    pos = {el: i for i, el in enumerate(body)}
    start = pos.get(block_head._p, -1)
    # glue every paragraph from here to the end to what follows (except the very
    # last), so the 2.6/2.7/2.8 block cannot break across pages
    tail_paras = [p for p in doc.paragraphs if pos.get(p._p, -1) >= start]
    for p in tail_paras[:-1]:
        _keep_with_next(p)
    # tables in the tail (equipment / software / result): never split a row, and
    # keep their rows together so the whole small table stays on the page
    for tbl in doc.tables:
        if pos.get(tbl._tbl, -1) >= start:
            rows = tbl.rows
            for tr in rows:
                _row_cant_split(tr)
            for tr in rows[:-1]:
                _keep_row(tr)


def polish_layout(doc):
    """Apply all layout fixes to a rendered document (body only; headers/footers
    are left untouched)."""
    paras = doc.paragraphs  # body-level paragraphs (not inside tables)

    for i, p in enumerate(paras):
        if _has_image(p):
            # image centered like its caption, and glued to the caption below
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _keep_with_next(p)
            # glue the preceding heading (skipping blanks) to the image so a
            # section title can never be orphaned at a page bottom
            j = i - 1
            while j >= 0 and not _text(paras[j]).strip():
                _keep_with_next(paras[j])
                j -= 1
            if j >= 0 and not CAPTION_RE.match(_text(paras[j]).strip()):
                _keep_with_next(paras[j])
        elif _soft_breaks(p):
            # justified paragraphs + soft line-breaks = words stretched across
            # the page; left-align them (textarea content: procedure, deviation)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if _is_heading(_text(p)):
            # a section heading must never be stranded at a page bottom: glue it
            # (and any blank spacers under it) to the first following content
            _keep_with_next(p)
            j = i + 1
            while j < len(paras) and not _text(paras[j]).strip():
                _keep_with_next(paras[j])
                j += 1

    # glue each heading/label (and blanks under it) to the table that follows,
    # so "SOFTWARE USED"-style titles are never orphaned at a page bottom
    body = list(doc.element.body)
    para_by_el = {p._p: p for p in paras}
    for idx, el in enumerate(body):
        if not el.tag.endswith("}tbl"):
            continue
        j = idx - 1
        while j >= 0 and body[j].tag.endswith("}p"):
            p = para_by_el.get(body[j])
            if p is None:
                break
            _keep_with_next(p)
            if _text(p).strip():          # the heading itself — stop here
                break
            j -= 1                        # blank spacer — keep gluing upward

    for tbl in doc.tables:
        rows = tbl.rows
        for tr in rows:
            _row_cant_split(tr)          # never split a row across pages
        if len(rows) <= 6:
            # small block tables (sign-off, limits, software...) stay on one page
            for tr in rows[:-1]:
                _keep_row(tr)
        else:
            # long tables: don't strand 1-2 rows at a page bottom (orphans) or
            # leave a lone last row on the next page (widows)
            _keep_row(rows[0])
            _keep_row(rows[1])
            _keep_row(rows[-2])
            if _row_is_header(rows[0]):
                _mark_header_row(rows[0])    # repeat header on continuation pages
        _equalize_columns(tbl)


def ce_finalize_layout(doc):
    """CE-only pagination + caption placement (per the DS504 reference layout).

    Applied after polish_layout + page_break_before_top_sections, from the CE
    generator only (generic forms are untouched). It:

      1. Moves each "Table N:" caption to sit BELOW its table (the source
         template puts the caption above; the reference shows image-then-caption
         and table-then-caption -- caption always below the object).
      2. Forces page breaks so the datasheet paginates like the reference:
           * "TEST PROCEDURE"      -> new page: procedure + Measurement-Data
                                       figure 1 & table 1 sit together.
           * Figure 2 block        -> new page: figure 2 & table 2 on their own page.
           * "TEST SETUP PICTURES" -> new page: photo (2.6) + equipment (2.7) together.
           * "SOFTWARE USED"       -> new page: software (2.8) + result (2.9) last page.
      3. Keeps each image -> caption -> table block from splitting across a page.
    """
    # 0) collapse runs of >1 consecutive blank paragraphs to a single blank. The
    #    source template carries several spacer lines (e.g. 4 blanks before
    #    "Measurement Data") that waste ~2cm and push the Table 1 caption onto the
    #    next page; one blank is enough separation.
    #    Table-aware: a table (or any non-paragraph) between two blanks breaks the
    #    run, so the single spacer between a table and the next heading is kept.
    # 0a) An un-uploaded 1.4/1.5 plot slot leaves a full-height empty paragraph; enough of
    #     them spill the end of section 1 onto a sheet with nothing visible on it.
    _ce_drop_empty_image_slots(doc)

    body = doc.element.body
    w_p = qn("w:p")
    prev_blank = False
    for el in list(body):
        if el.tag == w_p:
            txt = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
            has_img = el.findall(".//" + qn("w:drawing")) or el.findall(".//" + qn("w:pict"))
            blank = (not txt) and (not has_img)
            if blank and prev_blank:
                body.remove(el)
            else:
                prev_blank = blank
        else:
            prev_blank = False   # table / sectPr / etc. ends the blank run

    # 1) table caption -> below its table
    for p in list(doc.paragraphs):
        t = _text(p).strip()
        if t[:6].lower() == "table " and ":" in t[:14]:
            nxt = p._p.getnext()
            if nxt is not None and nxt.tag == qn("w:tbl"):
                nxt.addnext(p._p)   # relocate the caption to after the table

    # 2) reset all keep-with-next (polish_layout glues nearly every paragraph,
    #    which would otherwise chain the whole document into one un-fittable block).
    for p in doc.paragraphs:
        p.paragraph_format.keep_with_next = False

    # 2b) Bind every short sub-section (heading + text + its table) into one unit, so a
    #     block that no longer fits moves whole to the next page rather than splitting -
    #     2.3 TEST LIMITS used to break between its heading and the middle of its grid.
    #     Runs after the reset above (which only clears) and before the explicit tail
    #     groups below (which only add), so neither fights the other.
    _ce_keep_subsections(doc)
    # 2c) keep-with-next cannot stop Word breaking a TABLE between rows, so measure each
    #     sub-section of section 2 and force a page break where one would straddle.
    _ce_break_overflowing_subsections(doc)

    # 3) Group the tail of the datasheet into discrete keep-together blocks, each
    #    of which fits on one page, and let Word FLOW them (no forced page breaks,
    #    which double up into blank pages when the previous page is full). A block
    #    that doesn't fit the remaining space moves whole to the next page, giving:
    #      p1 section 1 | p2 2.1-2.3 | p3 procedure+fig1/table1 | p4 fig2/table2
    #      p5 photo(2.6)+equipment(2.7) | p6 software(2.8)+result(2.9)
    paras = doc.paragraphs
    texts = [_text(p).strip().lower() for p in paras]

    def find(pred, start=0):
        for k in range(start, len(paras)):
            if pred(texts[k]):
                return k
        return None

    a0 = find(lambda t: t == "test procedure")                       # procedure (2.4)
    c0 = find(lambda t: t == "test setup pictures")                  # photo (2.6) + equipment (2.7)
    d0 = find(lambda t: t == "software used")                        # software (2.8) + result (2.9)

    def glue(start, end):
        """Bind [start, end) into one keep-together block: keep_together stops any
        single paragraph (e.g. the procedure text) splitting across pages, and
        keep_with_next chains the members. The last member is left free so the
        next block can begin on a fresh page."""
        if start is None:
            return
        end = len(paras) if end is None else end
        last = min(end, len(paras))
        for k in range(start, last):
            paras[k].paragraph_format.keep_together = True
        for k in range(start, last - 1):
            paras[k].paragraph_format.keep_with_next = True

    # 2.5 Measurement Data, one record per Test.
    #
    # This used to assume each record held exactly TWO plots and paginate by pairing the
    # "Figure 1" and "Figure 2" captions, giving each pair its own page. A record now
    # carries up to four (Quasi-peak + Average, for Line and for Neutral) plus any extra
    # images, so that pairing matched the wrong paragraphs - the forced break landed on the
    # Line Average plot instead of the Neutral one and stranded a lone "Table 2:" caption
    # on a page of its own.
    #
    # Detector-agnostic instead: force a page break only at each RECORD, glue every image
    # to the caption beneath it, and let the images flow. A figure whose caption will not
    # fit below it moves down as a unit, so nothing is ever orphaned, whatever the count.
    m0 = find(lambda t: t == "measurement data")      # 2.5 heading
    end_meas = c0 if c0 is not None else len(paras)
    start_meas = m0 if m0 is not None else 0

    def _is_caption(t):
        return bool(re.match(r"^(figure|table|photo)\s+\d+\s*:", t))

    # A record begins at its label: the only non-empty, non-caption, non-heading, image-free
    # body paragraph in the measurement region.
    labels = [k for k in range(start_meas + 1, end_meas)
              if texts[k] and not _is_caption(texts[k])
              and not (paras[k].style.name or "").startswith("Heading")
              and not _has_image(paras[k])]

    if a0 is not None:
        glue(a0, start_meas if m0 is not None else (labels[0] if labels else end_meas))
    if m0 is not None:
        paras[m0].paragraph_format.page_break_before = True   # 2.5 starts a fresh page
        paras[m0].paragraph_format.keep_with_next = True      # ... with its first record
    for j, lk in enumerate(labels):
        if j:                                  # the first record shares 2.5's page
            paras[lk].paragraph_format.page_break_before = True
        paras[lk].paragraph_format.keep_with_next = True      # label stays with its plot
    # Every plot keeps the caption that follows it; the caption itself stays free so the
    # next figure may start a new page.
    for k in range(start_meas, end_meas):
        if _has_image(paras[k]):
            paras[k].paragraph_format.keep_together = True
            paras[k].paragraph_format.keep_with_next = True
        elif _is_caption(texts[k]):
            paras[k].paragraph_format.keep_together = True

    glue(c0, d0)                # Group C: 2.6 photo + 2.7 equipment
    glue(d0, None)              # Group D: 2.8 software + 2.9 result
    if c0 is not None:
        paras[c0].paragraph_format.page_break_before = True   # 2.6 setup + 2.7 equipment on their own page

    # 2.8 Software + 2.9 Result belong together on the final page. Group C (photo +
    # equipment) leaves the page part-empty, so software would otherwise flow up onto
    # it; a page break here moves the pair down cleanly (Group C is not full, so this
    # does not create a blank page).
    if d0 is not None:
        paras[d0].paragraph_format.page_break_before = True

    # 4) 1.4 Functional Check plots (Line + Neutral): put them on their own page with
    #    each bold heading directly ABOVE its image, and keep each heading glued to its
    #    image. Detected as the image paragraphs that appear BEFORE the "Conducted
    #    Emission Test" section (the 2.5/2.6 plots live after it).
    sec2 = find(lambda t: t == "conducted emission test")
    limit = sec2 if sec2 is not None else len(paras)
    func_imgs = [k for k in range(limit) if _has_image(paras[k])]
    if func_imgs:
        def _heading_above(k):
            j = k - 1
            while j > 0 and not _text(paras[j]).strip():
                j -= 1
            return j
        paras[_heading_above(func_imgs[0])].paragraph_format.page_break_before = True
        for k in func_imgs:
            hk = _heading_above(k)
            paras[hk].paragraph_format.keep_with_next = True
            paras[hk].paragraph_format.keep_together = True
            paras[k].paragraph_format.keep_together = True

    # 1.5 Ambient always begins a new page, so the first page ends after 1.4 Functional
    # Check. Outside the `if func_imgs` above on purpose: with no plots uploaded anywhere in
    # section 1 there are no image paragraphs to detect, and Ambient still has to start its
    # own page. This no longer risks a blank sheet, because the empty placeholder paragraphs
    # that used to overflow the page are dropped in step 0a.
    amb = find(lambda t: t == "ambient")
    if amb is not None:
        amb_end = sec2 if sec2 is not None else len(paras)
        paras[amb].paragraph_format.page_break_before = True
        paras[amb].paragraph_format.keep_with_next = True
        for k in range(amb, amb_end):
            paras[k].paragraph_format.keep_together = True

    # 5) Glue each table to the "Table N:" / "Figure N:" caption that sits BELOW it.
    #    A table has no keep-with-next of its own, so when it lands at the bottom of a
    #    page Word can strand its caption on the next page (caption orphaned from its
    #    table). Setting keepNext on the table's cell paragraphs keeps the whole table
    #    with the caption that follows it. Scoped to tables followed by such a caption
    #    (the small measurement plot/data tables), so large tables are unaffected.
    from docx.oxml import OxmlElement
    for tbl in body.findall(qn("w:tbl")):
        nxt = tbl.getnext()
        if nxt is None or nxt.tag != w_p:
            continue
        cap = "".join(t.text or "" for t in nxt.iter(qn("w:t"))).strip().lower()
        if not cap.startswith("table "):   # data table -> its "Table N:" caption below
            continue
        for cp in tbl.iter(w_p):
            pPr = cp.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr"); cp.insert(0, pPr)
            if pPr.find(qn("w:keepNext")) is None:
                pPr.append(OxmlElement("w:keepNext"))

    # 6) LAST: every page_break_before is now assigned, so drop the empty spacers sitting
    #    just before them. A spacer that will not fit on a full page flows onto one of its
    #    own, and the forced break then pushes the real content to the page after - a blank
    #    sheet in the middle of the document.
    _ce_strip_blanks_before_breaks(doc)


# --------------------------------------------------------------------------
# Font enforcement — force Arial on every table cell run
# --------------------------------------------------------------------------

def enforce_arial_procedure(doc):
    """Force Arial on the TEST PROCEDURE body text.

    The procedure text is entered/edited on the web form and rendered into a
    template paragraph whose style is 'Normal (Web)' (Times New Roman) or the
    document default — so the run inherits Times New Roman. This walks the body
    paragraphs between the 'TEST PROCEDURE' heading and the next heading and sets
    every run's font to Arial (symbol/checkbox runs are skipped)."""
    from docx.oxml import OxmlElement
    ARIAL = "Arial"
    SKIP_FONTS = {"Segoe UI Symbol", "Symbol", "Wingdings"}
    in_proc = False
    for p in doc.paragraphs:
        name = (p.style.name if p.style is not None else "") or ""
        if name.strip().lower().startswith("heading"):
            in_proc = "test procedure" in (p.text or "").strip().lower()
            continue
        if not in_proc or not (p.text or "").strip():
            continue
        # The procedure is one paragraph with soft line-breaks; if the template
        # justifies it, every line is stretched with big word gaps. Left-align it
        # (matches polish_layout's textarea handling) so the text reads normally.
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                current = rFonts.get(qn("w:ascii")) or rFonts.get(qn("w:hAnsi")) or ""
                if current in SKIP_FONTS:
                    continue
                rFonts.set(qn("w:ascii"), ARIAL)
                rFonts.set(qn("w:hAnsi"), ARIAL)
                rFonts.set(qn("w:cs"), ARIAL)
            else:
                rf = OxmlElement("w:rFonts")
                rf.set(qn("w:ascii"), ARIAL)
                rf.set(qn("w:hAnsi"), ARIAL)
                rf.set(qn("w:cs"), ARIAL)
                rPr.insert(0, rf)
            run.font.name = ARIAL


def enforce_arial_fonts(doc):
    """Walk every paragraph in every table cell and force the font to Arial.

    This overrides any Calibri runs that come from:
      * docxtpl Jinja-rendered measurement-data rows (inherit Normal/Calibri style)
      * any cell whose run-level rFonts was left unset (Word defaults to Calibri)

    Body paragraphs (headings, procedure text) are left untouched because
    they already carry explicit Arial formatting via the template styles.
    The checkbox runs (Segoe UI Symbol) are skipped so ballot boxes render
    correctly.
    """
    SKIP_FONTS = {"Segoe UI Symbol", "Symbol", "Wingdings"}
    ARIAL = "Arial"

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(11)
                        rPr = run._r.get_or_add_rPr()
                        # Skip special symbol runs (checkboxes)
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is not None:
                            current = (
                                rFonts.get(qn("w:ascii")) or
                                rFonts.get(qn("w:hAnsi")) or ""
                            )
                            if current in SKIP_FONTS:
                                continue
                            # Update existing rFonts element
                            rFonts.set(qn("w:ascii"), ARIAL)
                            rFonts.set(qn("w:hAnsi"), ARIAL)
                            rFonts.set(qn("w:cs"), ARIAL)
                        else:
                            # Create new rFonts element
                            from docx.oxml import OxmlElement
                            rf = OxmlElement("w:rFonts")
                            rf.set(qn("w:ascii"), ARIAL)
                            rf.set(qn("w:hAnsi"), ARIAL)
                            rf.set(qn("w:cs"), ARIAL)
                            rPr.insert(0, rf)
                        # Also set run.font.name so python-docx's own cache is consistent
                        run.font.name = ARIAL


def enforce_body_arial(doc, size=11):
    """Force Arial on every body (non-table) paragraph and set <size>pt on every
    non-heading paragraph (Heading/Title styles keep their heading size). Also pins
    the Normal / Normal (Web) styles to Arial <size> so any run that merely INHERITS
    comes out Arial instead of Word's Calibri fallback (the root cause of the stray
    Calibri: those base styles carry no explicit font). Symbol/checkbox runs keep
    their font. Companion to enforce_arial_fonts (which covers table cells)."""
    from docx.oxml import OxmlElement
    SKIP_FONTS = {"Segoe UI Symbol", "Symbol", "Wingdings"}
    ARIAL = "Arial"
    for style_name in ("Normal", "Normal (Web)"):
        try:
            stl = doc.styles[style_name]
            stl.font.name = ARIAL
            stl.font.size = Pt(size)
        except KeyError:
            pass
    for para in doc.paragraphs:
        nm = (para.style.name if para.style is not None else "") or ""
        is_heading = nm.strip().lower().startswith(("heading", "title"))
        for run in para.runs:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                current = rFonts.get(qn("w:ascii")) or rFonts.get(qn("w:hAnsi")) or ""
                if current not in SKIP_FONTS:
                    rFonts.set(qn("w:ascii"), ARIAL)
                    rFonts.set(qn("w:hAnsi"), ARIAL)
                    rFonts.set(qn("w:cs"), ARIAL)
                    run.font.name = ARIAL
            else:
                rf = OxmlElement("w:rFonts")
                rf.set(qn("w:ascii"), ARIAL)
                rf.set(qn("w:hAnsi"), ARIAL)
                rf.set(qn("w:cs"), ARIAL)
                rPr.insert(0, rf)
                run.font.name = ARIAL
            if not is_heading:
                run.font.size = Pt(size)


def shrink_wide_obs_tables(doc, size=8):
    """Re-shrink the wide Surge observation matrices (17 columns) after the global
    Arial pass — 11pt cannot fit 17 columns on a portrait page. Keeps the (already
    Arial) font, only reduces the point size and centres each cell. Matched by the
    header (Common/Differential Mode) or the Signal-line label."""
    for t in doc.tables:
        try:
            hdr = " ".join(c.text for c in t.rows[0].cells)
            first = t.rows[0].cells[0].text.strip()
        except Exception:
            continue
        if ("Common Mode" in hdr and "Differential Mode" in hdr) or first.startswith("Name of the signal"):
            for row in t.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.font.size = Pt(size)
