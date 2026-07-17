# -*- coding: utf-8 -*-
"""Chunk a generated datasheet .docx *by section* for embedding.

Generated datasheets are short, highly structured forms (not prose): section
headers use ``Heading 1`` / ``Heading 2`` styles and most real content lives in
tables. So the natural unit is one section per chunk = a heading plus every
block (paragraphs + tables, in true document order) up to the next heading.

Only oversized sections (big measurement matrices) are split, on row
boundaries, with the section label re-emitted so each piece stays self-labeled.
Checkbox glyphs (U+2610 / U+2612) are kept - they encode PASS/FAIL, class, etc.
"""
import os
import re

MAX_CHUNK_CHARS = 1200
OVERLAP_CHARS = 150

_JINJA_RE = re.compile(r"\{\{.*?\}\}|\{%.*?%\}", re.S)   # strip template tokens if fed a raw template


def _iter_block_items(doc):
    """Yield Paragraph and Table objects in true document order."""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _table_rows(tbl):
    rows = []
    for row in tbl.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        line = " | ".join(cells).strip()
        if line.strip(" |"):
            rows.append(line)
    return rows


def _clean(text):
    text = _JINJA_RE.sub("", text or "")
    return re.sub(r"[ \t]+", " ", text).strip()


def _split_oversized(label, lines):
    """Pack `lines` (already section-labeled) into <=MAX_CHUNK_CHARS pieces,
    splitting only on line boundaries and repeating the label on each piece."""
    pieces, buf, size = [], [], 0
    header = label
    budget = max(200, MAX_CHUNK_CHARS - len(header) - 1)

    def _wrap(s):
        # a single monolithic line (e.g. one giant table row) must still be
        # broken so no emitted chunk exceeds the budget (the embedder would
        # otherwise silently truncate the tail at 8k chars).
        return [s[j:j + budget] for j in range(0, len(s), budget)] if len(s) > budget else [s]

    flat = [seg for ln in lines for seg in _wrap(ln)]
    for ln in flat:
        add = len(ln) + 1
        if buf and size + add > MAX_CHUNK_CHARS:
            pieces.append("\n".join([header] + buf))
            # carry a little overlap (last line or two) for context
            keep, ksz = [], 0
            for prev in reversed(buf):
                if ksz + len(prev) > OVERLAP_CHARS:
                    break
                keep.insert(0, prev)
                ksz += len(prev) + 1
            buf, size = list(keep), ksz
        buf.append(ln)
        size += add
    if buf:
        pieces.append("\n".join([header] + buf))
    return pieces


def chunk_datasheet(path):
    """Return a list of {"section": label, "text": chunk} for one .docx.
    Never raises for a readable docx; returns [] if it cannot be parsed."""
    from docx import Document
    try:
        doc = Document(path)
    except Exception:
        return []

    sections = []          # (label, [lines])
    cur_title, cur_h1, buf = None, "", []

    def flush():
        lines = [ln for ln in buf if ln]
        if cur_title is None:
            # content before any heading (or a doc with NO heading styles) must
            # still be indexed - emit it under a document-level label rather than
            # silently dropping it.
            if not lines:
                return
            label = title_hint(path) or os.path.basename(path or "") or "(document)"
        else:
            label = ("%s > %s" % (cur_h1, cur_title)) if cur_h1 and cur_h1 != cur_title else cur_title
        sections.append((label, lines))

    for blk in _iter_block_items(doc):
        cls = blk.__class__.__name__
        if cls == "Paragraph":
            style = (getattr(blk.style, "name", "") or "")
            text = _clean(blk.text)
            if style.startswith("Heading"):
                flush()
                buf = []
                cur_title = text or "(section)"
                if style == "Heading 1":
                    cur_h1 = text
            elif text:
                buf.append(text)
        else:  # Table
            for ln in _table_rows(blk):
                cleaned = _clean(ln)
                if cleaned:
                    buf.append(cleaned)
    flush()

    chunks = []
    for label, lines in sections:
        body_len = sum(len(x) + 1 for x in lines)
        if not lines and not label:
            continue
        if body_len <= MAX_CHUNK_CHARS:
            text = "\n".join([label] + lines).strip()
            if len(text) > len(label):     # keep only sections with real content
                chunks.append({"section": label, "text": text})
        else:
            for piece in _split_oversized(label, lines):
                chunks.append({"section": label, "text": piece.strip()})
    return chunks


def title_hint(path):
    """A short document title (first Heading 1) for metadata, best-effort."""
    from docx import Document
    try:
        doc = Document(path)
        for p in doc.paragraphs:
            if (getattr(p.style, "name", "") or "") == "Heading 1" and p.text.strip():
                return p.text.strip()
    except Exception:
        pass
    return os.path.basename(path or "")
