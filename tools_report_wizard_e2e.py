# -*- coding: utf-8 -*-
"""Does what the admin typed into the report wizard actually reach the report?

    clear the draft -> build a BASELINE report -> post all 21 fields through the
    real wizard endpoint -> check both stores -> readiness() -> build again
    through the real generate endpoint -> open the .docx and look for every value

Every value is posted through ``POST /report/wizard/15/eut`` (multipart, four
real PNGs) and the document is built by ``POST /api/test-requests/15/generate-
test-report``. Nothing is written straight into report_draft and no builder
function is called directly, because the question is whether the wired-up path
works, not whether the pieces can.

WHY THE BASELINE BUILD IS NOT OPTIONAL
--------------------------------------
This project has shipped harnesses that passed for the wrong reason - a check
that could not fail, a scorer that took a fallback for an answer. So every check
here is run TWICE: once against a report built with the draft cleared and the six
request columns nulled, and once against the report built from the full wizard
post. A check that already passes on the baseline proves nothing, and is printed
as WEAK rather than PASS. Three of the twenty-two would have been WEAK if the
posted values had been chosen carelessly:

    test_location      the cover defaults to "Permanent" when nobody answers,
                       so the test posts "Onsite"
    report_issue_date  fill_cover writes TODAY, so the test posts a date that is
                       not today
    ULR NO             the template placeholder is <TC14704YY0XXXXXXXXF> and the
                       value written over it is TC14704YY0XXXXXXXXF - the same
                       string. Searching for the value alone cannot fail, so the
                       check is that the ANGLE-BRACKETED form is gone.

TWO MODES
---------
    python tools_report_wizard_e2e.py                    # the app as it ships
    python tools_report_wizard_e2e.py --wire-draft-fill  # + the documented edit

report_gen/draft_fill.py is written but NOT called from builder.build_report -
the integration is a comment block at the end of that file. --wire-draft-fill
installs exactly that edit in this process only (it wraps
builder.tick_decision_rules, which build_report calls once, on the line the
comment says to insert before). No source file is modified either way.

WHAT THIS TOUCHES
-----------------
The six store="request" columns on request 15, its report_draft row, and the
statuses the generate endpoint moves. All of them are snapshotted at the start
and restored at the end, so the request is left as it was found. The generated
.docx files stay on disk as evidence.
"""
import argparse
import datetime
import io
import os
import struct
import sys
import time
import zipfile
import zlib

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:                                     # the report prints degree signs and NA
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:                        # noqa: BLE001 - older stdout, not fatal
    pass

import app as app_module                                        # noqa: E402
from models import db                                           # noqa: E402
from sqlalchemy import text                                     # noqa: E402

from docx import Document                                       # noqa: E402
from docx.table import Table                                    # noqa: E402
from docx.text.paragraph import Paragraph                       # noqa: E402

from report_gen import builder as B                              # noqa: E402
from report_gen import docx_tools as T                           # noqa: E402
from report_gen import draft as DRAFT                            # noqa: E402
from report_gen import wizard_fields as WF                       # noqa: E402

REQUEST_ID = 15
SEC = "EUT INFORMATION"

# The placeholder the template ships in the running header. The value written
# over it is the same string WITHOUT the angle brackets, which is why the check
# is for the disappearance of this exact form.
ULR_PLACEHOLDER = "<%s>" % WF.ULR_NO


# ==========================================================================
# the values posted, and why each one is what it is
# ==========================================================================
# Every value differs from what the request/template already holds, so a check
# cannot pass on stale or defaulted data. The three noted below are the ones that
# would otherwise be unfalsifiable.
def payload(today):
    not_today = today - datetime.timedelta(days=7)
    return {
        "condition_on_receipt": "Received sealed, no visible transit damage",
        "date_of_receipt": (today - datetime.timedelta(days=11)).isoformat(),
        # NOT "Permanent": _COVER_DEFAULTS guesses Permanent, so posting it would
        # make the check pass whether or not the wizard value was read.
        "test_location": "Onsite",
        # NOT today: fill_cover writes today's date into this row.
        "report_issue_date": not_today.isoformat(),
        "issued_to": "Ravi Kulkarni, Quality Lead\nThermo Fisher Scientific\n+91 80 4567 8900",
        # exactly representable in binary32, so a FLOAT column cannot turn a
        # correct write into a rounding failure
        "length": "412.5",
        "width": "305",
        "height": "268",
        "dimension_unit": "cm",              # the request currently says mm
        "weight": "13.75",
        "operating_frequency": "60 Hz",      # the request currently says 50 Hz
        "power_rating": "725 W",
        "measured_current": "3.15 A",
        "software_firmware": "ThermoCtl firmware v4.2.1, host application v10.3.2",
        "eut_configuration": "Bench top, mains powered, Ethernet and USB connected",
        "modes_of_operation": "Mode A: idle, display on\nMode B: continuous scan at full rate",
        "monitoring_parameters": "Sample block temperature and lid status logged once per second via ThermoLink v3",
    }


IMAGE_COLORS = {"img_block_diagram": (18, 84, 168),
                "img_eut_photo": (196, 64, 32),
                "img_eut_label": (32, 148, 96),
                "img_monitoring": (148, 96, 196)}


def png_bytes(width, height, rgb):
    """A real, minimally valid RGB PNG - no Pillow needed to make one.

    Hand-built rather than Pillow-drawn so the test does not depend on an
    optional package, and 32x24 rather than 1x1 so python-docx has a sane aspect
    ratio to scale into the picture box.
    """
    raw = b"".join(b"\x00" + bytes(rgb) * width for _ in range(height))

    def chunk(tag, data):
        body = tag + data
        return (struct.pack(">I", len(data)) + body
                + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF))

    return (b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
            + chunk(b"IDAT", zlib.compress(raw))
            + chunk(b"IEND", b""))


# ==========================================================================
# reading the document - deliberately not through the code under test
# ==========================================================================
# Rows are located by a label match written here rather than by builder._find_row,
# so a bug in the finder cannot make the assertions agree with it. The merged-cell
# readers (T.distinct_cells / T.full_text) ARE reused - they answer "what does
# this cell say", which is not what is being tested, and reimplementing vMerge
# handling is how a probe ends up reading the wrong cell.
def _norm(s):
    return "".join(ch for ch in (s or "").upper() if ch.isalnum())


def _row_value(table, label):
    """(value_text, matched_label) for the row whose label contains ``label``."""
    want = _norm(label)
    for row in table.rows:
        cells = T.distinct_cells(row)
        if len(cells) < 2:
            continue
        lab = T.full_text(cells[0]).strip()
        if want and want in _norm(lab):
            return T.full_text(cells[-1]).strip(), lab
    return None, None


def _find_table_with_row(tables, label):
    for tb in tables:
        v, lab = _row_value(tb, label)
        if lab is not None:
            return tb
    return None


def _all_tables(blocks):
    return [b for b in blocks if isinstance(b, Table)]


def _sub_text(outline, sub):
    """The body paragraphs of a section-2 subsection, heading and captions out."""
    out = []
    for b in outline.sub_blocks(SEC, sub)[1:]:
        if isinstance(b, Table):
            break
        if not isinstance(b, Paragraph):
            continue
        if T.style_name(b) == "Caption":
            continue
        txt = T.text_of(b)
        if txt:
            out.append(txt)
    return out


def _drawings(block):
    el = block._tbl if isinstance(block, Table) else block._p
    return len(el.findall(".//" + T.qn("w:drawing")))


def _drawings_before_caption(outline, sub, index):
    """Pictures sitting in the slot above the index-th Figure/Photo caption.

    Walks backwards from the caption over empty paragraphs only, which is what
    "in this slot" means - a drawing further up belongs to the previous caption.
    """
    caps = [b for b in outline.sub_blocks(SEC, sub)
            if isinstance(b, Paragraph) and T.style_name(b) == "Caption"
            and (T.text_of(b) or "").strip()[:6].lower().startswith(("figure", "photo"))]
    if len(caps) <= index:
        return None, None
    cap = caps[index]
    n = 0
    prev = cap._p.getprevious()
    while prev is not None and prev.tag == T.qn("w:p"):
        p = Paragraph(prev, cap._parent)
        d = _drawings(p)
        if d:
            n += d
        elif T.text_of(p):
            break
        prev = prev.getprevious()
    return n, T.text_of(cap)


def _header_text(doc):
    out = []
    for section in doc.sections:
        for name in ("header", "first_page_header"):
            hf = getattr(section, name, None)
            if hf is None:
                continue
            for tbl in hf.tables:
                for row in tbl.rows:
                    for cell in T.distinct_cells(row):
                        t = T.full_text(cell).strip()
                        if t:
                            out.append(t)
            for p in hf.paragraphs:
                if T.text_of(p):
                    out.append(T.text_of(p))
    return "\n".join(out)


def read_document(path):
    """Everything the checks need, read once, as plain strings and counts."""
    doc = Document(path)
    ol = B.Outline(doc)
    s, e = ol.section_span(SEC)
    sec_blocks = ol.blocks[s:e] if s is not None else []

    cover = _find_table_with_row(_all_tables(ol.blocks[:12]),
                                 "CONDITION OF EUT ON RECEIPT")
    details = _find_table_with_row(_all_tables(sec_blocks), "Size of the EUT")

    out = {"path": path, "cover_found": cover is not None,
           "details_found": details is not None, "cover": {}, "details": {},
           "labels": {}}

    for label in ("CONDITION OF EUT ON RECEIPT", "DATE OF RECEIPT OF EUT",
                  "LOCATION OF PERFORMANCE OF TEST", "TEST REPORT ISSUE DATE",
                  "ISSUED TO"):
        v, lab = (_row_value(cover, label) if cover is not None else (None, None))
        out["cover"][label] = v
        out["labels"]["cover:" + label] = lab

    for label in ("Size of the EUT", "Weight of the EUT",
                  "EUT Operating Frequency", "EUT Power Rating",
                  "Measured EUT Current"):
        v, lab = (_row_value(details, label) if details is not None else (None, None))
        out["details"][label] = v
        out["labels"]["2.1:" + label] = lab

    out["text"] = {sub: _sub_text(ol, sub) for sub in (
        "SOFTWARE AND FIRMWARE DETAILS", "EUT CONFIGURATION DURING TEST",
        "EUT MODES OF OPERATION", "EUT MONITORING PARAMETERS")}

    out["figure1"] = _drawings_before_caption(ol, "EUT SETUP DETAILS", 0)
    out["photo1"] = _drawings_before_caption(
        ol, "EUT AND ACCESSORIES PICTURES", 0)
    out["photo2"] = _drawings_before_caption(
        ol, "EUT AND ACCESSORIES PICTURES", 1)

    mon_blocks = ol.sub_blocks(SEC, "EUT MONITORING PARAMETERS")
    out["monitoring_drawings"] = sum(_drawings(b) for b in mon_blocks)
    out["monitoring_captions"] = [T.text_of(b) for b in mon_blocks
                                  if isinstance(b, Paragraph)
                                  and T.style_name(b) == "Caption" and T.text_of(b)]
    out["section2_drawings"] = sum(_drawings(b) for b in sec_blocks)
    out["header"] = _header_text(doc)
    with zipfile.ZipFile(path) as z:
        out["media"] = sorted(n for n in z.namelist() if n.startswith("word/media/"))
    return out


# ==========================================================================
# the checks
# ==========================================================================
class Check(object):
    """One field, one destination, one verdict - and whether it could have failed.

    ``weak`` is the whole point: a check whose BASELINE value already satisfies it
    is not evidence of anything, and calling it PASS is exactly the failure mode
    this file is written to avoid.
    """

    def __init__(self, key, where, expect, test):
        self.key, self.where, self.expect, self.test = key, where, expect, test
        self.before = self.after = None
        self.ok = self.weak = False
        self.note = ""

    def run(self, before, after):
        self.before, self.after = self._read(before), self._read(after)
        self.ok = bool(self.test(after))
        self.weak = bool(before is not None and self.test(before))
        return self

    def _read(self, doc):
        return None if doc is None else self.reader(doc)

    @property
    def verdict(self):
        if not self.ok:
            return "FAIL"
        return "WEAK" if self.weak else "PASS"


def build_checks(posted, today):
    """Twenty-two checks: 5 cover + 5 in 2.1 + 4 text + 4 images + ULR + 3 no-NA.

    Expected strings are composed HERE from the posted values, not by calling the
    formatter the builder uses. service.fmt_date's "%d %b %Y" and
    service._dimensions' "L x W x H unit" are restated as literal expectations,
    so a change of format is a failure rather than something both sides agree on.
    """
    checks = []

    def add(key, where, expect, reader, test, note=""):
        c = Check(key, where, expect, test)
        c.reader = reader
        c.note = note
        checks.append(c)
        return c

    def cover(label):
        return lambda d: d["cover"].get(label)

    def detail(label):
        return lambda d: d["details"].get(label)

    def contains(needle):
        return lambda d, n=needle: n in (Check.__dict__ and "") or False   # replaced below

    # ---------------- cover ----------------
    add("condition_on_receipt", "cover CONDITION OF EUT ON RECEIPT",
        posted["condition_on_receipt"], cover("CONDITION OF EUT ON RECEIPT"),
        lambda d: posted["condition_on_receipt"] in (d["cover"].get("CONDITION OF EUT ON RECEIPT") or ""))

    recv = datetime.date.fromisoformat(posted["date_of_receipt"]).strftime("%d %b %Y")
    add("date_of_receipt", "cover DATE OF RECEIPT OF EUT", recv,
        cover("DATE OF RECEIPT OF EUT"),
        lambda d: recv in (d["cover"].get("DATE OF RECEIPT OF EUT") or ""),
        "was 'NA' - a value nobody chose")

    add("test_location", "cover LOCATION OF PERFORMANCE OF TEST",
        posted["test_location"], cover("LOCATION OF PERFORMANCE OF TEST"),
        lambda d: posted["test_location"] in (d["cover"].get("LOCATION OF PERFORMANCE OF TEST") or ""),
        "'Onsite' on purpose: the default guess is 'Permanent'")

    iss = datetime.date.fromisoformat(posted["report_issue_date"]).strftime("%d %b %Y")
    add("report_issue_date", "cover TEST REPORT ISSUE DATE", iss,
        cover("TEST REPORT ISSUE DATE"),
        lambda d: iss in (d["cover"].get("TEST REPORT ISSUE DATE") or ""),
        "not today's date on purpose: fill_cover writes today")

    issued_first = posted["issued_to"].splitlines()[0]
    add("issued_to", "cover ISSUED TO", issued_first, cover("ISSUED TO"),
        lambda d: issued_first in (d["cover"].get("ISSUED TO") or ""),
        "overrides the customer contact fill_cover writes")

    # ---------------- 2.1 EUT DETAILS ----------------
    size = "%g x %g x %g %s" % (float(posted["length"]), float(posted["width"]),
                                float(posted["height"]), posted["dimension_unit"])
    add("length+width+height+dimension_unit", "2.1 Size of the EUT (L x W x H)",
        size, detail("Size of the EUT"),
        lambda d: size in (d["details"].get("Size of the EUT") or ""))

    wt = "%g kg" % float(posted["weight"])
    add("weight", "2.1 Weight of the EUT", wt, detail("Weight of the EUT"),
        lambda d: wt in (d["details"].get("Weight of the EUT") or ""))

    add("operating_frequency", "2.1 EUT Operating Frequency",
        posted["operating_frequency"], detail("EUT Operating Frequency"),
        lambda d: posted["operating_frequency"] in (d["details"].get("EUT Operating Frequency") or ""))

    add("power_rating", "2.1 EUT Power Rating", posted["power_rating"],
        detail("EUT Power Rating"),
        lambda d: posted["power_rating"] in (d["details"].get("EUT Power Rating") or ""))

    add("measured_current", "2.1 Measured EUT Current", posted["measured_current"],
        detail("Measured EUT Current"),
        lambda d: posted["measured_current"] in (d["details"].get("Measured EUT Current") or ""))

    # ---------------- 2.3 / 2.5 / 2.7 / 2.8 text ----------------
    for key, sub in (("software_firmware", "SOFTWARE AND FIRMWARE DETAILS"),
                     ("eut_configuration", "EUT CONFIGURATION DURING TEST"),
                     ("monitoring_parameters", "EUT MONITORING PARAMETERS")):
        val = posted[key]
        add(key, "2.x %s" % sub, val[:46] + "...",
            lambda d, s=sub: " | ".join(d["text"].get(s) or []),
            lambda d, s=sub, v=val: v in "\n".join(d["text"].get(s) or []))

    modes = posted["modes_of_operation"].splitlines()
    add("modes_of_operation", "2.7 EUT MODES OF OPERATION",
        "%d lines, each its own paragraph" % len(modes),
        lambda d: " | ".join(d["text"].get("EUT MODES OF OPERATION") or []),
        lambda d: all(any(m in line for line in
                          (d["text"].get("EUT MODES OF OPERATION") or []))
                      for m in modes))

    # ---------------- the four pictures ----------------
    add("img_block_diagram", "2.6 EUT SETUP DETAILS / Figure 1",
        "1 drawing above the caption",
        lambda d: "%s above %r" % (d["figure1"][0], (d["figure1"][1] or "")[:34]),
        lambda d: (d["figure1"][0] or 0) >= 1)

    add("img_eut_photo", "2.x EUT AND ACCESSORIES PICTURES / Photo 1",
        "1 drawing above the caption",
        lambda d: "%s above %r" % (d["photo1"][0], (d["photo1"][1] or "")[:34]),
        lambda d: (d["photo1"][0] or 0) >= 1)

    add("img_eut_label", "2.x EUT AND ACCESSORIES PICTURES / Photo 2",
        "1 drawing above the caption",
        lambda d: "%s above %r" % (d["photo2"][0], (d["photo2"][1] or "")[:34]),
        lambda d: (d["photo2"][0] or 0) >= 1)

    add("img_monitoring", "2.8 EUT MONITORING PARAMETERS / appended Photo",
        "1 drawing inside 2.8",
        lambda d: "%d drawing(s), captions %s" % (d["monitoring_drawings"],
                                                  d["monitoring_captions"]),
        lambda d: d["monitoring_drawings"] >= 1)

    add("(all four)", "section 2 drawings", "4 or more",
        lambda d: d["section2_drawings"],
        lambda d: d["section2_drawings"] >= 4)

    add("(all four)", "embedded media files", "4 more than the baseline",
        lambda d: len(d["media"]), lambda d: True)          # filled in by run()

    # ---------------- ULR NO ----------------
    add("ulr_no", "running header ULR NO",
        "%s, and NOT %s" % (WF.ULR_NO, ULR_PLACEHOLDER),
        lambda d: [l for l in d["header"].splitlines() if "ULR" in l.upper()],
        lambda d: (ULR_PLACEHOLDER not in d["header"]
                   and WF.ULR_NO in d["header"]),
        "the value equals the placeholder minus its <>, so the check is that "
        "the bracketed form is gone")

    # ---------------- and none of the four still says NA ----------------
    add("date_of_receipt", "cover DATE OF RECEIPT is not 'NA'", "not NA",
        cover("DATE OF RECEIPT OF EUT"),
        lambda d: (d["cover"].get("DATE OF RECEIPT OF EUT") or "").strip() != "NA")
    for key, sub in (("software_firmware", "SOFTWARE AND FIRMWARE DETAILS"),
                     ("eut_configuration", "EUT CONFIGURATION DURING TEST"),
                     ("monitoring_parameters", "EUT MONITORING PARAMETERS")):
        add(key, "2.x %s is not 'NA'" % sub, "not NA",
            lambda d, s=sub: " | ".join(d["text"].get(s) or []),
            lambda d, s=sub: [t.strip() for t in (d["text"].get(s) or [])] != ["NA"])
    return checks


# ==========================================================================
# the run
# ==========================================================================
def snapshot(app):
    """Everything this test changes, so it can be put back."""
    with app.app_context():
        cols = [f[0] for f in WF.by_store("request")]
        row = db.session.execute(text(
            "SELECT status, %s FROM iec_emc_requests WHERE id=:r"
            % ", ".join("`%s`" % c for c in cols)), {"r": REQUEST_ID}).first()
        d = db.session.execute(text(
            "SELECT form_json, images_json, page_reached, updated_by_user_id "
            "FROM report_draft WHERE test_request_id=:r"), {"r": REQUEST_ID}).first()
        entries = db.session.execute(text(
            "SELECT id, status FROM planner_entries WHERE test_request_id=:r"),
            {"r": REQUEST_ID}).fetchall()
        types = {c: db.session.execute(text(
            "SELECT COLUMN_TYPE FROM information_schema.columns WHERE "
            "table_schema=DATABASE() AND table_name='iec_emc_requests' AND "
            "COLUMN_NAME=:c"), {"c": c}).scalar() for c in cols}
    return {"request": dict(zip(["status"] + cols, row)) if row else None,
            "draft": tuple(d) if d else None,
            "entries": [tuple(e) for e in entries], "types": types}


def restore(app, snap):
    with app.app_context():
        r = snap["request"]
        if r:
            cols = [k for k in r if k != "status"]
            db.session.execute(text(
                "UPDATE iec_emc_requests SET status=:status, %s WHERE id=:r"
                % ", ".join("`%s`=:%s" % (c, c) for c in cols)),
                dict(r, r=REQUEST_ID))
        db.session.execute(text("DELETE FROM report_draft WHERE test_request_id=:r"),
                           {"r": REQUEST_ID})
        if snap["draft"]:
            db.session.execute(text(
                "INSERT INTO report_draft (test_request_id, form_json, images_json,"
                " page_reached, updated_by_user_id) VALUES (:r,:f,:i,:p,:u)"),
                {"r": REQUEST_ID, "f": snap["draft"][0], "i": snap["draft"][1],
                 "p": snap["draft"][2], "u": snap["draft"][3]})
        for eid, status in snap["entries"]:
            db.session.execute(text(
                "UPDATE planner_entries SET status=:s WHERE id=:i"),
                {"s": status, "i": eid})
        db.session.commit()


def null_request_columns(app):
    """The baseline needs the six columns empty, which is how the DB shipped."""
    with app.app_context():
        cols = [f[0] for f in WF.by_store("request")]
        db.session.execute(text("UPDATE iec_emc_requests SET %s WHERE id=:r"
                                % ", ".join("`%s`=NULL" % c for c in cols)),
                           {"r": REQUEST_ID})
        db.session.commit()


def wire_draft_fill(sink):
    """Install the integration edit documented at the end of draft_fill.py.

    In this process only. build_report calls tick_decision_rules exactly once,
    immediately after fill_eut_information + outline.refresh() and well before
    cleanup_instructions - which is precisely the insertion point the comment
    block specifies, so wrapping it runs apply_draft where the edit would.
    """
    from report_gen import draft_fill as DF
    original = B.tick_decision_rules
    if getattr(original, "_e2e_wrapped", False):
        return

    def wrapped(outline, meta):
        sink.append(DF.apply_draft(outline, outline.doc, REQUEST_ID, meta=meta))
        return original(outline, meta)

    wrapped._e2e_wrapped = True
    B.tick_decision_rules = wrapped


def generate(client, login, label):
    login()
    t0 = time.time()
    r = client.post("/api/test-requests/%d/generate-test-report" % REQUEST_ID,
                    json={"comments": "wizard end-to-end test (%s)" % label})
    body = r.get_json() or {}
    print("   %-9s HTTP %s in %5.1fs  %s" % (
        label, r.status_code, time.time() - t0,
        os.path.basename(body.get("file_path") or "") or str(body)[:120]))
    if r.status_code != 200 or not body.get("file_path"):
        return None
    path = body["file_path"]
    if not os.path.isabs(path):
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), path)
    return path if os.path.exists(path) else body["file_path"]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--wire-draft-fill", action="store_true",
                    help="install the documented builder.py edit in this process")
    ap.add_argument("--no-baseline", action="store_true",
                    help="skip the baseline build (loses the WEAK detection)")
    ap.add_argument("--keep", action="store_true",
                    help="do not restore request 15 afterwards")
    args = ap.parse_args()

    today = datetime.date.today()
    posted = payload(today)

    print("=" * 78)
    print("REPORT WIZARD END TO END - request %d" % REQUEST_ID)
    print("mode: %s" % ("builder.py edit WIRED IN (in-process)"
                        if args.wire_draft_fill else "the app exactly as it ships"))
    print("=" * 78)

    app = app_module.create_app("default")
    app.config["WTF_CSRF_ENABLED"] = False
    app.config["PROPAGATE_EXCEPTIONS"] = True
    app.login_manager.session_protection = None

    with app.app_context():
        uid = db.session.execute(text(
            "SELECT id FROM users WHERE role='admin' AND is_active=1 "
            "ORDER BY id LIMIT 1")).scalar()
    if uid is None:
        print("no active admin user - cannot drive the wizard")
        return 1

    client = app.test_client()

    def login():
        with client.session_transaction() as s:
            s["_user_id"] = str(uid)
            s["_fresh"] = True

    snap = snapshot(app)
    print("\n0. SNAPSHOT")
    print("   request status %r, six columns %s" % (
        (snap["request"] or {}).get("status"),
        {k: v for k, v in (snap["request"] or {}).items() if k != "status"}))
    print("   report_draft row: %s" % ("present" if snap["draft"] else "none"))
    print("   column types: %s" % snap["types"])

    fails = []
    baseline_doc = None
    sink = []
    try:
        # ---------------- 1. clear, and a baseline build ----------------
        print("\n1. CLEAR THE DRAFT")
        with app.app_context():
            cleared = DRAFT.clear(REQUEST_ID)
            left = DRAFT.load(REQUEST_ID)
        null_request_columns(app)
        print("   draft.clear -> %s; reloaded form=%s images=%s exists=%s"
              % (cleared, left["form"], left["images"], left["exists"]))
        print("   the six request columns set to NULL")
        if left["form"] or left["images"]:
            fails.append("draft.clear() left data behind")

        if args.wire_draft_fill:
            wire_draft_fill(sink)
            print("   builder.tick_decision_rules wrapped -> apply_draft will run")

        if not args.no_baseline:
            print("\n2. BASELINE BUILD (nothing entered - the 'before')")
            p = generate(client, login, "baseline")
            if p:
                baseline_doc = read_document(p)

        # ---------------- 3. post the whole page ----------------
        print("\n3. POST THE WIZARD PAGE  (POST /report/wizard/%d/eut)" % REQUEST_ID)
        data = dict(posted)
        for key, rgb in IMAGE_COLORS.items():
            data[key] = (io.BytesIO(png_bytes(32, 24, rgb)), "%s.png" % key)
        login()
        r = client.post("/report/wizard/%d/eut" % REQUEST_ID, data=data,
                        content_type="multipart/form-data")
        body = r.get_json() or {}
        print("   HTTP %s  saved_to_request=%s images_saved=%s outstanding=%s"
              % (r.status_code, body.get("saved_to_request"),
                 body.get("images_saved"), body.get("outstanding")))
        print("   message: %s" % body.get("message"))
        if body.get("rejected"):
            print("   REJECTED: %s" % body["rejected"])
            fails.append("the wizard rejected values: %s" % body["rejected"])
        if r.status_code != 200:
            fails.append("wizard POST returned HTTP %s" % r.status_code)
            print("   body: %s" % str(body)[:400])

        # ---------------- 4. did the two stores take it ----------------
        print("\n4. WHERE THE VALUES LANDED")
        print("   %-22s %-10s %-26s %-26s %s"
              % ("field", "store", "posted", "read back", "ok"))
        store_fail = 0
        with app.app_context():
            cols = [f[0] for f in WF.by_store("request")]
            req_row = dict(db.session.execute(text(
                "SELECT %s FROM iec_emc_requests WHERE id=:r"
                % ", ".join("`%s`" % c for c in cols)),
                {"r": REQUEST_ID}).mappings().first() or {})
            d = DRAFT.load(REQUEST_ID)
        for key, _lab, kind, store, _loc, _h in WF.FIELDS:
            if kind == "image":
                got = (d["images"] or {}).get(key)
                ok = bool(got) and os.path.exists(str(got))
                shown = os.path.basename(str(got)) if got else None
                exp = "a saved PNG"
            elif store == "request":
                got = req_row.get(key)
                exp = posted[key]
                if key in WF.NUMERIC:
                    ok = isinstance(got, float) and abs(got - float(exp)) < 1e-6
                else:
                    ok = str(got or "") == exp
                shown = "%r (%s)" % (got, type(got).__name__)
            else:
                got = (d["form"] or {}).get(key)
                exp = posted[key]
                ok = str(got or "") == exp
                shown = repr(got)
            if not ok:
                store_fail += 1
            print("   %-22s %-10s %-26s %-26s %s"
                  % (key, store, repr(exp)[:26], str(shown)[:26],
                     "ok" if ok else "MISMATCH"))
        if store_fail:
            fails.append("%d field(s) did not land in their store" % store_fail)

        # ---------------- 5. readiness ----------------
        print("\n5. readiness()")
        from report_gen import wizard_review as WR
        with app.app_context():
            state = WR.readiness(REQUEST_ID)
        print("   ready=%s  outstanding=%d  filled=%d/%d  blockers=%s"
              % (state.get("ready"), len(state.get("outstanding") or []),
                 state.get("filled"), state.get("total"), state.get("blockers")))
        for w in state.get("warnings") or []:
            print("   warning: %s" % w)
        if not state.get("ready"):
            fails.append("readiness() says not ready: %s" % state.get("blockers"))
        if state.get("outstanding"):
            fails.append("readiness() still lists %d outstanding: %s"
                         % (len(state["outstanding"]),
                            [o["key"] for o in state["outstanding"]]))

        # ---------------- 6. build it for real ----------------
        print("\n6. GENERATE THE REPORT  (POST /api/test-requests/%d/generate-test-report)"
              % REQUEST_ID)
        final = generate(client, login, "filled")
        if not final:
            fails.append("the generate endpoint did not produce a report")
            print("\nCANNOT CHECK THE DOCUMENT - no report was produced")
            return report(fails, [], None, None, sink)
        after = read_document(final)

        # ---------------- 7. the document ----------------
        checks = build_checks(posted, today)
        for c in checks:
            c.run(baseline_doc, after)
        # the media check needs both documents, so it is settled here
        for c in checks:
            if c.where == "embedded media files":
                if baseline_doc is None:
                    c.ok, c.weak = after and len(after["media"]) > 0, True
                    c.note = "no baseline build to compare against"
                else:
                    gained = len(after["media"]) - len(baseline_doc["media"])
                    c.ok = gained >= 4
                    c.weak = False
                    c.expect = "baseline+4"
                    c.note = "%+d media files" % gained
        return report(fails, checks, baseline_doc, after, sink)
    finally:
        if args.keep:
            print("\n--keep: request %d left holding the test values" % REQUEST_ID)
        else:
            restore(app, snap)
            print("\nrequest %d restored to the snapshot taken at the start"
                  % REQUEST_ID)


def report(fails, checks, before, after, sink):
    print("\n" + "=" * 78)
    print("PER-FIELD RESULT   (WEAK = the baseline already satisfied it, so the")
    print("                    check proves nothing about the wizard)")
    print("=" * 78)
    print("%-7s %-24s %-44s" % ("", "field", "destination"))
    for c in checks:
        print("%-7s %-24s %-44s" % (c.verdict, c.key[:24], c.where[:44]))
        print("        expected : %s" % str(c.expect)[:120])
        print("        baseline : %s" % str(c.before)[:120])
        print("        built    : %s" % str(c.after)[:120])
        if c.note:
            print("        note     : %s" % c.note[:120])

    n = len(checks)
    p = sum(1 for c in checks if c.verdict == "PASS")
    w = sum(1 for c in checks if c.verdict == "WEAK")
    f = sum(1 for c in checks if c.verdict == "FAIL")
    print("\n" + "-" * 78)
    print("DOCUMENT CHECKS: %d PASS, %d WEAK, %d FAIL, out of %d" % (p, w, f, n))
    if before and after:
        print("media files: %d baseline -> %d built (%+d)"
              % (len(before["media"]), len(after["media"]),
                 len(after["media"]) - len(before["media"])))
        print("section 2 drawings: %d baseline -> %d built"
              % (before["section2_drawings"], after["section2_drawings"]))
    if sink:
        d = sink[-1]
        print("\napply_draft reported: %d written, %d images, %d missing"
              % (len(d["written"]), len(d["images"]), len(d["missing"])))
        print("  missing: %s" % d["missing"])
        for note in d["notes"]:
            print("  note: %s" % note)
    if fails:
        print("\nPIPELINE FAILURES BEFORE THE DOCUMENT WAS EVEN OPENED:")
        for x in fails:
            print("  - %s" % x)
    if after:
        print("\nreport: %s" % after["path"])
    print("row labels matched: %s" % (after or {}).get("labels"))
    return 0 if (f == 0 and not fails) else 1


if __name__ == "__main__":
    sys.exit(main())
