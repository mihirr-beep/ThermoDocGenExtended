import os
import re
import json
import logging
from typing import Dict, List, Optional, Any
from datetime import datetime

try:
    from spire.doc import Document as SpireDocument
    from spire.doc import FileFormat
    SPIRE_AVAILABLE = True
except ImportError:
    SPIRE_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

if not SPIRE_AVAILABLE:
    logger.warning("Spire.Doc not available")
if not DOCX_AVAILABLE:
    logger.warning("python-docx not available")


class EnhancedDocumentProcessor:
    """Enhanced document processor using Spire.Doc for text extraction and template population."""

    def __init__(self, upload_folder: str, template_folder: str, output_folder: str):
        """Initialize the enhanced document processor.

        Args:
            upload_folder: Path to the upload folder
            template_folder: Path to the template folder
            output_folder: Path to the output folder
        """
        self.upload_folder = upload_folder
        self.template_folder = template_folder
        self.output_folder = output_folder

    def extract_text_with_spire(self, file_path: str) -> str:
        """Extract text from a DOCX file using Spire.Doc.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text as string

        Raises:
            Exception: If text extraction fails
        """
        if not SPIRE_AVAILABLE:
            raise Exception("Spire.Doc is not available")

        try:
            # Load document using Spire.Doc
            doc = SpireDocument()
            doc.LoadFromFile(file_path)

            # Extract text
            text = doc.GetText()
            doc.Close()

            logger.info(
                f"Successfully extracted text using Spire.Doc from {file_path}")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text with Spire.Doc: {e}")
            raise Exception(f"Failed to extract text with Spire.Doc: {e}")

    def extract_text_with_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file using python-docx.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text as string

        Raises:
            Exception: If text extraction fails
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx is not available")

        try:
            # Load document using python-docx
            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            text = '\n'.join(text_parts)

            logger.info(
                f"Successfully extracted text using python-docx from {file_path}")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text with python-docx: {e}")
            raise Exception(f"Failed to extract text with python-docx: {e}")

    def parse_test_request_data(self, text: str) -> Dict[str, Any]:
        """Parse extracted text to identify test request data.

        Args:
            text: Extracted text from the document

        Returns:
            Dictionary containing parsed test request data
        """
        data = {
            'project_name': '',
            'eut_name': '',  # Equipment Under Test name
            'product_name': '',  # Product name (synonym for EUT)
            'test_objective': '',
            'test_scope': '',
            'equipment_required': [],
            'tests_to_perform': [],
            'safety_requirements': '',
            'test_parameters': {},
            'contact_info': {},
            'submission_date': '',
            'priority': 'normal',
            'estimated_duration': '',
            'special_requirements': ''
        }

        lines = text.split('\n')

        # Extract project information
        for i, line in enumerate(lines):
            line = line.strip()

            # Project name patterns
            if any(keyword in line.lower() for keyword in ['project name', 'project title', 'project']):
                if ':' in line:
                    data['project_name'] = line.split(':', 1)[1].strip()
                elif i + 1 < len(lines):
                    data['project_name'] = lines[i + 1].strip()

            # EUT (Equipment Under Test) and Product Name patterns
            elif any(keyword in line.lower() for keyword in ['equipment under test', 'eut', 'product name', 'product title', 'device under test', 'dut', 'name of the product', 'name of equipment under test']):
                if ':' in line:
                    eut_value = line.split(':', 1)[1].strip()
                    # Clean up the value - take only the first part if it's too long
                    if len(eut_value) > 100:  # If value is too long, it might be wrong
                        # Try to find a shorter, more meaningful value
                        parts = eut_value.split('.')
                        if parts:
                            eut_value = parts[0].strip()
                        # If still too long, try splitting by other delimiters
                        if len(eut_value) > 50:
                            parts = eut_value.split(',')
                            if parts:
                                eut_value = parts[0].strip()
                    data['eut_name'] = eut_value
                    data['product_name'] = eut_value
                    # If project name is empty, use EUT name as project name
                    if not data['project_name']:
                        data['project_name'] = eut_value
                elif i + 1 < len(lines):
                    eut_value = lines[i + 1].strip()
                    # Clean up the value - take only the first part if it's too long
                    if len(eut_value) > 100:  # If value is too long, it might be wrong
                        # Try to find a shorter, more meaningful value
                        parts = eut_value.split('.')
                        if parts:
                            eut_value = parts[0].strip()
                        # If still too long, try splitting by other delimiters
                        if len(eut_value) > 50:
                            parts = eut_value.split(',')
                            if parts:
                                eut_value = parts[0].strip()
                    data['eut_name'] = eut_value
                    data['product_name'] = eut_value
                    # If project name is empty, use EUT name as project name
                    if not data['project_name']:
                        data['project_name'] = eut_value

            # Test objective patterns
            elif any(keyword in line.lower() for keyword in ['objective', 'purpose', 'goal']):
                if ':' in line:
                    data['test_objective'] = line.split(':', 1)[1].strip()
                elif i + 1 < len(lines):
                    data['test_objective'] = lines[i + 1].strip()

            # Test scope patterns
            elif any(keyword in line.lower() for keyword in ['scope', 'range', 'coverage']):
                if ':' in line:
                    data['test_scope'] = line.split(':', 1)[1].strip()
                elif i + 1 < len(lines):
                    data['test_scope'] = lines[i + 1].strip()

            # Equipment patterns
            elif any(keyword in line.lower() for keyword in ['equipment', 'instrument', 'device', 'tool']):
                if ':' in line:
                    equipment = line.split(':', 1)[1].strip()
                    if equipment and equipment not in data['equipment_required']:
                        data['equipment_required'].append(equipment)

            # Test procedures patterns
            elif any(keyword in line.lower() for keyword in ['test', 'procedure', 'method', 'step']):
                if line and not any(keyword in line.lower() for keyword in ['test plan', 'test report']):
                    if line not in data['tests_to_perform']:
                        data['tests_to_perform'].append(line)

            # Safety requirements patterns
            elif any(keyword in line.lower() for keyword in ['safety', 'precaution', 'warning']):
                if ':' in line:
                    data['safety_requirements'] = line.split(':', 1)[1].strip()
                elif i + 1 < len(lines):
                    data['safety_requirements'] = lines[i + 1].strip()

        return data

    def populate_iec_template(self, extracted_data: Dict[str, Any], output_filename: str) -> str:
        """Populate the IEC-FRM-503 template with extracted data using Spire.Doc.

        Args:
            extracted_data: Extracted data from the uploaded document
            output_filename: Name for the output file

        Returns:
            Path to the generated document
        """
        if not SPIRE_AVAILABLE:
            # Create a simple text-based output if Spire.Doc is not available
            logger.warning(
                "Spire.Doc not available, creating text-based output")
            return self._create_text_based_output(extracted_data, output_filename)

        try:
            # Template path
            template_path = os.path.join(
                self.template_folder, 'IEC-FRM-503 EMI EMC Test Plan.docx')

            if not os.path.exists(template_path):
                raise FileNotFoundError(
                    f"Template file not found: {template_path}")

            # Load template using Spire.Doc
            doc = SpireDocument()
            doc.LoadFromFile(template_path)

            # Prepare replacements
            replacements = {
                '{{PROJECT_NAME}}': extracted_data.get('project_name', 'N/A'),
                '{{EUT_NAME}}': extracted_data.get('eut_name', extracted_data.get('project_name', 'N/A')),
                '{{PRODUCT_NAME}}': extracted_data.get('product_name', extracted_data.get('project_name', 'N/A')),
                '{{TEST_OBJECTIVE}}': extracted_data.get('test_objective', 'N/A'),
                '{{TEST_SCOPE}}': extracted_data.get('test_scope', 'N/A'),
                '{{PRIORITY}}': extracted_data.get('priority', 'Normal').title(),
                '{{ESTIMATED_DURATION}}': extracted_data.get('estimated_duration', 'N/A'),
                '{{SUBMISSION_DATE}}': extracted_data.get('submission_date', datetime.now().strftime('%Y-%m-%d')),
                '{{SAFETY_REQUIREMENTS}}': extracted_data.get('safety_requirements', 'Standard safety procedures apply.'),
                '{{SPECIAL_REQUIREMENTS}}': extracted_data.get('special_requirements', 'None'),
            }

            # Replace placeholders in the document using simple string replacement
            try:
                # Try the new Replace method signature first
                logger.info(
                    f"Attempting to replace placeholders with new signature")
                for placeholder, value in replacements.items():
                    logger.info(f"Replacing {placeholder} with: {value}")
                    doc.Replace(placeholder, str(value), True, True)

                # Handle equipment list
                equipment_text = '\n'.join(
                    [f"• {eq}" for eq in extracted_data.get('equipment_required', [])])
                if equipment_text:
                    doc.Replace('{{EQUIPMENT_LIST}}',
                                equipment_text, True, True)
                else:
                    doc.Replace('{{EQUIPMENT_LIST}}',
                                'To be determined', True, True)

                # Handle test procedures
                test_procedures = extracted_data.get('tests_to_perform', [])
                if test_procedures:
                    procedures_text = '\n'.join(
                        [f"{i+1}. {proc}" for i, proc in enumerate(test_procedures)])
                    doc.Replace('{{TEST_PROCEDURES}}',
                                procedures_text, True, True)
                else:
                    doc.Replace('{{TEST_PROCEDURES}}',
                                'Standard test procedures will be followed.', True, True)

            except Exception as replace_error:
                logger.warning(
                    f"New Replace method failed, trying old signature: {replace_error}")
                # Fallback to old Replace method signature
                logger.info(
                    f"Attempting to replace placeholders with old signature")
                for placeholder, value in replacements.items():
                    logger.info(f"Replacing {placeholder} with: {value}")
                    doc.Replace(placeholder, str(value))

                # Handle equipment list
                equipment_text = '\n'.join(
                    [f"• {eq}" for eq in extracted_data.get('equipment_required', [])])
                if equipment_text:
                    doc.Replace('{{EQUIPMENT_LIST}}', equipment_text)
                else:
                    doc.Replace('{{EQUIPMENT_LIST}}', 'To be determined')

                # Handle test procedures
                test_procedures = extracted_data.get('tests_to_perform', [])
                if test_procedures:
                    procedures_text = '\n'.join(
                        [f"{i+1}. {proc}" for i, proc in enumerate(test_procedures)])
                    doc.Replace('{{TEST_PROCEDURES}}', procedures_text)
                else:
                    doc.Replace('{{TEST_PROCEDURES}}',
                                'Standard test procedures will be followed.')

            # Generate output path
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"IEC_FRM_503_Test_Plan_{timestamp}.docx"
            output_path = os.path.join(self.output_folder, output_filename)

            # Save the populated document
            doc.SaveToFile(output_path, FileFormat.Docx)
            doc.Close()

            logger.info(f"Successfully populated IEC template: {output_path}")
            return output_filename

        except Exception as e:
            logger.error(f"Error populating IEC template: {e}")
            raise Exception(f"Failed to populate IEC template: {e}")

    def process_document_enhanced(self, filename: str) -> Dict[str, Any]:
        """Process a document file and extract structured data, then populate IEC template.

        Args:
            filename: Name of the uploaded file

        Returns:
            Dictionary containing processing results
        """
        file_path = os.path.join(self.upload_folder, filename)

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            # Extract text using Spire.Doc first, fallback to python-docx
            logger.info(f"Extracting text from {filename}")
            try:
                text = self.extract_text_with_spire(file_path)
                extraction_method = "Spire.Doc"
            except Exception as spire_error:
                logger.warning(
                    f"Spire.Doc failed, trying python-docx: {spire_error}")
                text = self.extract_text_with_docx(file_path)
                extraction_method = "python-docx"

            if not text or len(text.strip()) == 0:
                raise Exception("No text content found in document")

            # Parse data
            logger.info(f"Parsing data from {filename}")
            parsed_data = self.parse_test_request_data(text)

            # Validate data
            logger.info(f"Validating data from {filename}")
            validation_result = self.validate_extracted_data(parsed_data)

            # Populate IEC template
            logger.info(f"Populating IEC template with extracted data")
            output_filename = self.populate_iec_template(
                validation_result['cleaned_data'], filename)

            return {
                'success': True,
                'extracted_text': text,
                'parsed_data': validation_result['cleaned_data'],
                'validation': validation_result,
                'output_filename': output_filename,
                'extraction_method': extraction_method,
                'file_size': os.path.getsize(file_path)
            }

        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
            }

    def validate_extracted_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate and clean extracted data.

        Args:
            data: Extracted data

        Returns:
            Validation result with cleaned data
        """
        validation_result = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'cleaned_data': data.copy()
        }

        # Check required fields
        if not data.get('project_name'):
            validation_result['warnings'].append('Project name not found')

        if not data.get('test_objective'):
            validation_result['warnings'].append('Test objective not found')

        if not data.get('tests_to_perform'):
            validation_result['warnings'].append('No tests specified')

        if not data.get('equipment_required'):
            validation_result['warnings'].append('No equipment specified')

        # Clean equipment list
        cleaned_equipment = []
        for equipment in data.get('equipment_required', []):
            cleaned = equipment.strip()
            if cleaned:
                cleaned = re.sub(r'^\d+\.\s*', '', cleaned)
                cleaned_equipment.append(cleaned)

        validation_result['cleaned_data']['equipment_required'] = cleaned_equipment

        # Clean test list
        cleaned_tests = []
        for test in data.get('tests_to_perform', []):
            cleaned = test.strip()
            if cleaned:
                cleaned = re.sub(r'^\d+\.\s*', '', cleaned)
                cleaned_tests.append(cleaned)

        validation_result['cleaned_data']['tests_to_perform'] = cleaned_tests

        # Set default values for missing fields
        if not validation_result['cleaned_data'].get('priority'):
            validation_result['cleaned_data']['priority'] = 'normal'

        if not validation_result['cleaned_data'].get('submission_date'):
            validation_result['cleaned_data']['submission_date'] = datetime.now().strftime(
                '%Y-%m-%d')

        return validation_result

    def _create_text_based_output(self, extracted_data: Dict[str, Any], output_filename: str) -> str:
        """Create a simple text-based output when Spire.Doc is not available.

        Args:
            extracted_data: Extracted data from the uploaded document
            output_filename: Name for the output file

        Returns:
            Name of the generated text file
        """
        try:
            # Generate output path
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"IEC_FRM_503_Test_Plan_{timestamp}.txt"
            output_path = os.path.join(self.output_folder, output_filename)

            # Create text content
            content = []
            content.append("IEC-FRM-503 EMI EMC Test Plan")
            content.append("=" * 50)
            content.append("")
            content.append(
                f"Project Name: {extracted_data.get('project_name', 'N/A')}")
            content.append(
                f"Equipment Under Test (EUT): {extracted_data.get('eut_name', extracted_data.get('project_name', 'N/A'))}")
            content.append(
                f"Product Name: {extracted_data.get('product_name', extracted_data.get('project_name', 'N/A'))}")
            content.append(
                f"Test Objective: {extracted_data.get('test_objective', 'N/A')}")
            content.append(
                f"Test Scope: {extracted_data.get('test_scope', 'N/A')}")
            content.append(
                f"Priority: {extracted_data.get('priority', 'Normal').title()}")
            content.append(
                f"Estimated Duration: {extracted_data.get('estimated_duration', 'N/A')}")
            content.append(
                f"Submission Date: {extracted_data.get('submission_date', datetime.now().strftime('%Y-%m-%d'))}")
            content.append("")
            content.append("Safety Requirements:")
            content.append(extracted_data.get(
                'safety_requirements', 'Standard safety procedures apply.'))
            content.append("")
            content.append("Special Requirements:")
            content.append(extracted_data.get('special_requirements', 'None'))
            content.append("")

            # Equipment list
            content.append("Equipment Required:")
            equipment_list = extracted_data.get('equipment_required', [])
            if equipment_list:
                for eq in equipment_list:
                    content.append(f"• {eq}")
            else:
                content.append("• To be determined")
            content.append("")

            # Test procedures
            content.append("Test Procedures:")
            test_procedures = extracted_data.get('tests_to_perform', [])
            if test_procedures:
                for i, proc in enumerate(test_procedures, 1):
                    content.append(f"{i}. {proc}")
            else:
                content.append("Standard test procedures will be followed.")

            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))

            logger.info(
                f"Successfully created text-based output: {output_path}")
            return output_filename

        except Exception as e:
            logger.error(f"Error creating text-based output: {e}")
            raise Exception(f"Failed to create text-based output: {e}")
