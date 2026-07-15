import os
import json
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Class for generating test plans and datasheets from templates."""

    def __init__(self, template_folder: str, output_folder: str):
        """Initialize the document generator.

        Args:
            template_folder: Path to the template folder
            output_folder: Path to the output folder
        """
        self.template_folder = template_folder
        self.output_folder = output_folder

    def _safe_filename_part(self, value: Any, max_length: int = 64) -> str:
        """Return a Windows-safe filename fragment with bounded length."""
        text = str(value or '').strip()
        # Keep only conservative ASCII-safe characters.
        safe = ''.join(c for c in text if c.isalnum() or c in (' ', '-', '_'))
        safe = ' '.join(safe.split()).strip().replace(' ', '_')
        if not safe:
            safe = 'test'
        # Avoid trailing dots/spaces and limit total length.
        safe = safe.rstrip(' .')
        return safe[:max_length] or 'test'

    def create_test_plan_template(self) -> str:
        """Create a default test plan template if it doesn't exist.

        Returns:
            Path to the test plan template
        """
        template_path = os.path.join(
            self.template_folder, 'test_plan_template.docx')

        if not os.path.exists(template_path):
            doc = Document()

            # Title
            title = doc.add_heading('TEST PLAN', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Project Information
            doc.add_heading('1. PROJECT INFORMATION', level=1)

            # Project details table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Value'

            # Add project fields
            project_fields = [
                'Project Name',
                'Test Objective',
                'Test Scope',
                'Priority',
                'Estimated Duration',
                'Submission Date'
            ]

            for field in project_fields:
                row_cells = table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = f'{{{{{field.upper().replace(" ", "_")}}}}}'

            doc.add_paragraph()

            # Equipment Section
            doc.add_heading('2. EQUIPMENT REQUIRED', level=1)
            equipment_table = doc.add_table(rows=1, cols=4)
            equipment_table.style = 'Table Grid'

            eq_hdr_cells = equipment_table.rows[0].cells
            eq_hdr_cells[0].text = 'Equipment Name'
            eq_hdr_cells[1].text = 'Model'
            eq_hdr_cells[2].text = 'Serial Number'
            eq_hdr_cells[3].text = 'Status'

            # Add placeholder row
            eq_row = equipment_table.add_row().cells
            eq_row[0].text = '{{EQUIPMENT_NAME}}'
            eq_row[1].text = '{{EQUIPMENT_MODEL}}'
            eq_row[2].text = '{{EQUIPMENT_SERIAL}}'
            eq_row[3].text = '{{EQUIPMENT_STATUS}}'

            doc.add_paragraph()

            # Test Methodology
            doc.add_heading('3. TEST METHODOLOGY', level=1)
            doc.add_paragraph('{{TEST_METHODOLOGY}}')

            # Safety Requirements
            doc.add_heading('4. SAFETY REQUIREMENTS', level=1)
            doc.add_paragraph('{{SAFETY_REQUIREMENTS}}')

            # Test Procedures
            doc.add_heading('5. TEST PROCEDURES', level=1)
            doc.add_paragraph('{{TEST_PROCEDURES}}')

            # Approval Section
            doc.add_heading('6. APPROVAL', level=1)
            approval_table = doc.add_table(rows=1, cols=3)
            approval_table.style = 'Table Grid'

            app_hdr_cells = approval_table.rows[0].cells
            app_hdr_cells[0].text = 'Role'
            app_hdr_cells[1].text = 'Name'
            app_hdr_cells[2].text = 'Date'

            # Add approval roles
            approval_roles = ['Test Engineer',
                              'Project Manager', 'Quality Assurance']
            for role in approval_roles:
                app_row = approval_table.add_row().cells
                app_row[0].text = role
                app_row[1].text = ''
                app_row[2].text = ''

            # Save template
            doc.save(template_path)
            logger.info(f"Created test plan template: {template_path}")

        return template_path

    def create_datasheet_template(self) -> str:
        """Create a default datasheet template if it doesn't exist.

        Returns:
            Path to the datasheet template
        """
        template_path = os.path.join(
            self.template_folder, 'datasheet_template.docx')

        if not os.path.exists(template_path):
            doc = Document()

            # Title
            title = doc.add_heading('TEST DATASHEET', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Test Information
            doc.add_heading('1. TEST INFORMATION', level=1)

            info_table = doc.add_table(rows=1, cols=2)
            info_table.style = 'Table Grid'

            hdr_cells = info_table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Value'

            # Add test fields
            test_fields = [
                'Test Name',
                'Test Description',
                'Test Date',
                'Test Engineer',
                'Equipment Used'
            ]

            for field in test_fields:
                row_cells = info_table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = f'{{{{{field.upper().replace(" ", "_")}}}}}'

            doc.add_paragraph()

            # Test Parameters
            doc.add_heading('2. TEST PARAMETERS', level=1)
            param_table = doc.add_table(rows=1, cols=3)
            param_table.style = 'Table Grid'

            param_hdr_cells = param_table.rows[0].cells
            param_hdr_cells[0].text = 'Parameter'
            param_hdr_cells[1].text = 'Value'
            param_hdr_cells[2].text = 'Unit'

            # Add placeholder row
            param_row = param_table.add_row().cells
            param_row[0].text = '{{PARAMETER_NAME}}'
            param_row[1].text = '{{PARAMETER_VALUE}}'
            param_row[2].text = '{{PARAMETER_UNIT}}'

            doc.add_paragraph()

            # Measurement Points
            doc.add_heading('3. MEASUREMENT POINTS', level=1)
            doc.add_paragraph('{{MEASUREMENT_POINTS}}')

            # Data Recording
            doc.add_heading('4. DATA RECORDING', level=1)
            data_table = doc.add_table(rows=1, cols=4)
            data_table.style = 'Table Grid'

            data_hdr_cells = data_table.rows[0].cells
            data_hdr_cells[0].text = 'Measurement Point'
            data_hdr_cells[1].text = 'Expected Value'
            data_hdr_cells[2].text = 'Actual Value'
            data_hdr_cells[3].text = 'Status'

            # Add placeholder rows
            for i in range(5):
                data_row = data_table.add_row().cells
                data_row[0].text = f'{{MP_{i+1}_NAME}}'
                data_row[1].text = f'{{MP_{i+1}_EXPECTED}}'
                data_row[2].text = ''
                data_row[3].text = ''

            doc.add_paragraph()

            # Results and Conclusions
            doc.add_heading('5. RESULTS AND CONCLUSIONS', level=1)
            doc.add_paragraph('{{RESULTS_AND_CONCLUSIONS}}')

            # Signatures
            doc.add_heading('6. SIGNATURES', level=1)
            sig_table = doc.add_table(rows=1, cols=3)
            sig_table.style = 'Table Grid'

            sig_hdr_cells = sig_table.rows[0].cells
            sig_hdr_cells[0].text = 'Role'
            sig_hdr_cells[1].text = 'Name'
            sig_hdr_cells[2].text = 'Date'

            # Add signature roles
            sig_roles = ['Test Engineer', 'Witness', 'Approver']
            for role in sig_roles:
                sig_row = sig_table.add_row().cells
                sig_row[0].text = role
                sig_row[1].text = ''
                sig_row[2].text = ''

            # Save template
            doc.save(template_path)
            logger.info(f"Created datasheet template: {template_path}")

        return template_path

    def replace_placeholders(self, doc: Document, replacements: Dict[str, str]) -> None:
        """Replace placeholders in a document with actual values.

        Args:
            doc: Document object
            replacements: Dictionary of placeholder replacements
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        placeholder, str(value))

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(
                                    placeholder, str(value))

    def generate_test_plan(self, data: Dict[str, Any], equipment_list: List[Dict[str, Any]]) -> str:
        """Generate a test plan document.

        Args:
            data: Extracted test request data
            equipment_list: List of equipment information

        Returns:
            Path to the generated test plan file
        """
        # Ensure template exists
        template_path = self.create_test_plan_template()

        # Load template
        doc = Document(template_path)

        # Prepare replacements
        replacements = {
            '{{PROJECT_NAME}}': data.get('project_name', 'N/A'),
            '{{TEST_OBJECTIVE}}': data.get('test_objective', 'N/A'),
            '{{TEST_SCOPE}}': data.get('test_scope', 'N/A'),
            '{{PRIORITY}}': data.get('priority', 'Normal').title(),
            '{{ESTIMATED_DURATION}}': data.get('estimated_duration', 'N/A'),
            '{{SUBMISSION_DATE}}': data.get('submission_date', 'N/A'),
            '{{TEST_METHODOLOGY}}': self._generate_test_methodology(data),
            '{{SAFETY_REQUIREMENTS}}': data.get('safety_requirements', 'Standard safety procedures apply.'),
            '{{TEST_PROCEDURES}}': self._generate_test_procedures(data)
        }

        # Replace placeholders
        self.replace_placeholders(doc, replacements)

        # Handle equipment table
        self._populate_equipment_table(doc, equipment_list)

        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"test_plan_{timestamp}.docx"
        os.makedirs(self.output_folder, exist_ok=True)
        output_path = os.path.join(self.output_folder, filename)

        # Save document
        doc.save(output_path)
        logger.info(f"Generated test plan: {output_path}")

        return filename

    def generate_test_datasheet(self, test_name: str, test_data: Dict[str, Any],
                                equipment_list: List[Dict[str, Any]], test_index: int) -> str:
        """Generate a test datasheet document.

        Args:
            test_name: Name of the test
            test_data: Test-specific data
            equipment_list: List of equipment information
            test_index: Index of the test for filename

        Returns:
            Path to the generated datasheet file
        """
        # Ensure template exists
        template_path = self.create_datasheet_template()

        # Load template
        doc = Document(template_path)

        # Prepare replacements
        replacements = {
            '{{TEST_NAME}}': test_name,
            '{{TEST_DESCRIPTION}}': test_data.get('description', 'N/A'),
            '{{TEST_DATE}}': datetime.now().strftime('%Y-%m-%d'),
            '{{TEST_ENGINEER}}': 'TBD',
            '{{EQUIPMENT_USED}}': ', '.join([eq.get('name', '') for eq in equipment_list]),
            '{{MEASUREMENT_POINTS}}': self._generate_measurement_points(test_data),
            '{{RESULTS_AND_CONCLUSIONS}}': 'To be filled during testing.'
        }

        # Replace placeholders
        self.replace_placeholders(doc, replacements)

        # Handle parameter table
        self._populate_parameter_table(doc, test_data)

        # Handle data recording table
        self._populate_data_recording_table(doc, test_data)

        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_test_name = self._safe_filename_part(test_name, max_length=48)
        filename = f"datasheet_{test_index}_{safe_test_name}_{timestamp}.docx"
        os.makedirs(self.output_folder, exist_ok=True)
        output_path = os.path.join(self.output_folder, filename)

        # Save document
        doc.save(output_path)
        logger.info(f"Generated datasheet: {output_path}")

        return filename

    def _generate_test_methodology(self, data: Dict[str, Any]) -> str:
        """Generate test methodology text.

        Args:
            data: Test request data

        Returns:
            Test methodology text
        """
        methodology = []

        if data.get('test_objective'):
            methodology.append(f"Objective: {data['test_objective']}")

        if data.get('test_scope'):
            methodology.append(f"Scope: {data['test_scope']}")

        if data.get('tests_to_perform'):
            methodology.append("Test Procedures:")
            for i, test in enumerate(data['tests_to_perform'], 1):
                methodology.append(f"  {i}. {test}")

        if data.get('test_parameters'):
            methodology.append("Test Parameters:")
            for param, value in data['test_parameters'].items():
                methodology.append(f"  - {param}: {value}")

        return '\n'.join(methodology) if methodology else "Standard test procedures will be followed."

    def _generate_test_procedures(self, data: Dict[str, Any]) -> str:
        """Generate test procedures text.

        Args:
            data: Test request data

        Returns:
            Test procedures text
        """
        procedures = []

        if data.get('tests_to_perform'):
            for i, test in enumerate(data['tests_to_perform'], 1):
                procedures.append(f"{i}. {test}")
                procedures.append(
                    "   - Setup equipment according to specifications")
                procedures.append("   - Perform measurements and record data")
                procedures.append("   - Analyze results and document findings")
                procedures.append("")

        return '\n'.join(procedures) if procedures else "Follow standard test procedures."

    def _populate_equipment_table(self, doc: Document, equipment_list: List[Dict[str, Any]]) -> None:
        """Populate equipment table in test plan.

        Args:
            doc: Document object
            equipment_list: List of equipment information
        """
        for table in doc.tables:
            # Find equipment table (table with Equipment Name header)
            if table.rows and table.rows[0].cells[0].text == 'Equipment Name':
                # Clear existing rows except header
                while len(table.rows) > 1:
                    table._element.remove(table.rows[1]._element)

                # Add equipment rows
                for equipment in equipment_list:
                    row = table.add_row()
                    row.cells[0].text = equipment.get('name', 'N/A')
                    row.cells[1].text = equipment.get('model', 'N/A')
                    row.cells[2].text = equipment.get('serial_number', 'N/A')
                    row.cells[3].text = equipment.get('status', 'N/A')
                break

    def _populate_parameter_table(self, doc: Document, test_data: Dict[str, Any]) -> None:
        """Populate parameter table in datasheet.

        Args:
            doc: Document object
            test_data: Test-specific data
        """
        for table in doc.tables:
            # Find parameter table (table with Parameter header)
            if table.rows and table.rows[0].cells[0].text == 'Parameter':
                # Clear existing rows except header
                while len(table.rows) > 1:
                    table._element.remove(table.rows[1]._element)

                # Add parameter rows
                parameters = test_data.get('parameters', {})
                for param_name, param_value in parameters.items():
                    row = table.add_row()
                    row.cells[0].text = param_name
                    row.cells[1].text = str(param_value)
                    row.cells[2].text = test_data.get(
                        'units', {}).get(param_name, '')
                break

    def _populate_data_recording_table(self, doc: Document, test_data: Dict[str, Any]) -> None:
        """Populate data recording table in datasheet.

        Args:
            doc: Document object
            test_data: Test-specific data
        """
        for table in doc.tables:
            # Find data recording table (table with Measurement Point header)
            if table.rows and table.rows[0].cells[0].text == 'Measurement Point':
                # Clear existing rows except header
                while len(table.rows) > 1:
                    table._element.remove(table.rows[1]._element)

                # Add measurement point rows
                measurement_points = test_data.get('measurement_points', [])
                for i, point in enumerate(measurement_points, 1):
                    row = table.add_row()
                    row.cells[0].text = point.get('name', f'MP_{i}')
                    row.cells[1].text = str(point.get('expected', ''))
                    # Actual value (to be filled during testing)
                    row.cells[2].text = ''
                    # Status (to be filled during testing)
                    row.cells[3].text = ''
                break

    def _generate_measurement_points(self, test_data: Dict[str, Any]) -> str:
        """Generate measurement points text.

        Args:
            test_data: Test-specific data

        Returns:
            Measurement points text
        """
        points = test_data.get('measurement_points', [])
        if not points:
            return "Measurement points will be determined during test setup."

        point_texts = []
        for i, point in enumerate(points, 1):
            name = point.get('name', f'MP_{i}')
            description = point.get('description', 'N/A')
            point_texts.append(f"{i}. {name}: {description}")

        return '\n'.join(point_texts)
