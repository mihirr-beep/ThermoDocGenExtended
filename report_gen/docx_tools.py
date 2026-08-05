# -*- coding: utf-8 -*-
"""Low-level Word surgery helpers for building the IEC-FRM-516 test report.

WHY SURGERY AND NOT A TEMPLATE ENGINE
-------------------------------------
The official report (``word_templates/IEC-FRM-516_REPORT.docx``) is a blank form,
not a docxtpl template: 74 tables, heavily merged headers, literal ``U+2610``
checkbox glyphs, auto-numbered headings, and real Word ``TOC``/``SEQ``/``PAGE``
fields. Marking all of that up by hand would be error-prone and would lose the
conditional-section behaviour we need (a request only contains *some* tests).

So instead we open the official document and edit it in place:

* sections for tests that are not part of the request are **deleted whole**
  (Word then renumbers the remaining headings and Figure/Photo/Table SEQ
  numbers by itself, because both are field/numbering driven);
* values are written into existing cells while **preserving the run formatting
  already in the cell**, so the document keeps its Arial 11 / bold look;
* checkboxes are ticked by flipping the ``☐`` glyph to ``☒`` at ``w:t`` level,
  which keeps every run's formatting intact.

Everything here is deliberately generic: no report-specific knowledge lives in
this module (see ``mapping.py`` / ``builder.py`` for that).
"""
import copy
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

CHECK_OFF = "☐"        # ☐ ballot box               (U+2610, w14:uncheckedState)
CHECK_ON = "☒"         # ☒ ballot box with X        (U+2612, w14:checkedState)

DEFAULT_FONT = "Arial"
DEFAULT_SIZE_PT = 11

# Word 2010 wordml namespace - carries the checkbox content-control state.
# NOTE its attributes are w14:val, not w:val.
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"


def _w14(tag):
    return "{%s}%s" % (W14, tag)


# ==========================================================================
# Document traversal
# ==========================================================================

def iter_block_items(parent):
    """Yield the Paragraphs and Tables of ``parent`` in document order.

    python-docx exposes ``.paragraphs`` and ``.tables`` separately, which loses
    the interleaving we need to locate "the table that follows this heading".
    """
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        el = parent.element.body
    elif isinstance(parent, _Cell):
        el = parent._tc
    else:                                       # pragma: no cover - defensive
        el = parent._element
    for child in el.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def style_name(block):
    """Style name of a Paragraph, or '' for anything else."""
    try:
        return block.style.name or ""
    except Exception:
        return ""


def text_of(block):
    """Whitespace-collapsed text of a Paragraph (or '' for a Table)."""
    if isinstance(block, Paragraph):
        return re.sub(r"\s+", " ", block.text or "").strip()
    return ""


def remove(block):
    """Detach a Paragraph/Table (or raw element) from its parent."""
    el = getattr(block, "_element", None)
    if el is None:
        el = getattr(block, "_p", None) or getattr(block, "_tbl", None) or block
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_blocks(blocks):
    for b in blocks:
        remove(b)


# ==========================================================================
# Run formatting
# ==========================================================================

def _default_rpr(bold=False, size_pt=DEFAULT_SIZE_PT, font=DEFAULT_FONT):
    """A fresh ``w:rPr`` matching the report body style (Arial 11)."""
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)
    rpr.append(rfonts)
    if bold:
        rpr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(szcs)
    return rpr


def template_rpr(container, bold=None):
    """The ``w:rPr`` to reuse when writing into ``container`` (cell/paragraph).

    Copies the formatting of the first run that has any, so a value written into
    a bold cell stays bold and a value written into a plain cell stays plain.
    Falls back to Arial 11. ``bold`` forces bold on/off when not None.
    """
    el = _element_of(container)
    found = None
    for rpr in el.iter(qn("w:rPr")):
        parent = rpr.getparent()
        # skip paragraph-mark formatting (w:pPr/w:rPr) and content-control
        # properties (w:sdtPr/w:rPr) - neither is a real run's formatting
        if parent is not None and parent.tag in (qn("w:pPr"), qn("w:sdtPr")):
            continue
        found = rpr
        break
    rpr = copy.deepcopy(found) if found is not None else _default_rpr()
    if bold is not None:
        for b in rpr.findall(qn("w:b")):
            rpr.remove(b)
        for b in rpr.findall(qn("w:bCs")):
            rpr.remove(b)
        if bold:
            rpr.insert(0, OxmlElement("w:b"))
    return rpr


def make_run(text, rpr=None):
    """A ``w:r`` carrying ``text``; newlines become ``w:br`` line breaks."""
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    lines = str(text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    for i, line in enumerate(lines):
        if i:
            run.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        run.append(t)
    return run


def make_seq_field(kind="Photo", rpr=None):
    """A ``SEQ`` field run, e.g. the "10" in "Photo 10:".

    Captions appended by the generator must carry a real field, not a literal
    number: Word then numbers them in sequence with the template's own captions
    and includes them in the matching list of figures/photos/tables.
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " SEQ %s \\* ARABIC " % kind)
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.text = "1"
    run.append(t)
    fld.append(run)
    return fld


def clear_runs(paragraph):
    """Drop every run (and hyperlink/field) from a paragraph, keeping its pPr."""
    p = paragraph._p if isinstance(paragraph, Paragraph) else paragraph
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


# ==========================================================================
# Writing values
# ==========================================================================

def set_paragraph_text(paragraph, text, bold=None):
    """Replace a paragraph's content with ``text``, keeping its own formatting."""
    rpr = template_rpr(paragraph, bold)
    clear_runs(paragraph)
    p = paragraph._p if isinstance(paragraph, Paragraph) else paragraph
    p.append(make_run(text, rpr))
    return paragraph


def set_cell_text(cell, text, bold=None, align=None):
    """Replace a cell's text with ``text`` (multi-line aware).

    Keeps the cell's first paragraph (so its alignment/spacing survive), drops
    any extra paragraphs, and reuses the cell's existing run formatting.
    """
    if cell is None:
        return
    rpr = template_rpr(cell, bold)
    paras = cell.paragraphs
    for extra in paras[1:]:
        remove(extra)
    p = paras[0]
    clear_runs(p)
    p._p.append(make_run("" if text is None else text, rpr))
    if align is not None:
        p.alignment = align


def append_cell_line(cell, text, bold=None):
    """Add ``text`` as a new paragraph inside ``cell`` (keeps existing content)."""
    if cell is None:
        return
    rpr = template_rpr(cell, bold)
    p = OxmlElement("w:p")
    src = cell.paragraphs[-1]._p.find(qn("w:pPr")) if cell.paragraphs else None
    if src is not None:
        p.append(copy.deepcopy(src))
    p.append(make_run(text, rpr))
    cell._tc.append(p)


def set_paragraph_lines(paragraph, lines, bold=None):
    """Render ``lines`` as one paragraph per line, cloning ``paragraph``'s format.

    The first line reuses ``paragraph`` itself; the rest are inserted after it
    so surrounding spacing/justification is preserved. Returns the last
    paragraph written (useful as an insertion anchor).
    """
    lines = [l for l in (lines or []) if str(l).strip()] or [""]
    set_paragraph_text(paragraph, lines[0], bold)
    anchor = paragraph._p
    for line in lines[1:]:
        new_p = copy.deepcopy(paragraph._p)
        clear_runs(new_p)
        new_p.append(make_run(line, template_rpr(paragraph, bold)))
        anchor.addnext(new_p)
        anchor = new_p
    return Paragraph(anchor, paragraph._parent)


# ==========================================================================
# Checkboxes
# ==========================================================================
# In this report every checkbox is a Word **checkbox content control**:
#
#   <w:sdt>
#     <w:sdtPr>... <w14:checkbox>
#            <w14:checked w14:val="0"/>                      <- the real state
#            <w14:checkedState   w14:val="2612" .../>        <- glyph when on
#            <w14:uncheckedState w14:val="2610" .../>
#          </w14:checkbox></w:sdtPr>
#     <w:sdtContent><w:r><w:t>☐</w:t></w:r></w:sdtContent>   <- cached glyph
#   </w:sdt>
#   <w:r><w:t xml:space="preserve"> Basic </w:t></w:r>       <- its label
#
# Two consequences:
#   * ``cell.text`` from python-docx does NOT include the glyph (it lives inside
#     w:sdtContent), so use ``full_text()`` when reading these cells;
#   * ticking means setting BOTH ``w14:checked`` (what Word believes) and the
#     cached glyph (what is rendered). Setting only one leaves the control
#     inconsistent and Word can revert it on open.
#
# The label of a checkbox is the text that FOLLOWS its control, up to the next
# control - which makes option matching exact rather than positional.

def _element_of(container):
    """The underlying lxml element of a Cell/Paragraph (or the element itself).

    Written with explicit ``is None`` checks - lxml elements are falsy when they
    have no children, so ``a or b`` silently picks the wrong one.
    """
    el = getattr(container, "_tc", None)
    if el is None:
        el = getattr(container, "_p", None)
    if el is None:
        el = getattr(container, "_element", None)
    if el is None:
        el = container
    return el


def _text_nodes(container):
    """All ``w:t`` elements under a cell/paragraph, in order (incl. inside sdt)."""
    return [t for t in _element_of(container).iter(qn("w:t"))]


def full_text(container):
    """Text of a cell/paragraph **including** content-control content.

    Use instead of ``cell.text`` whenever checkbox glyphs matter.
    """
    return "".join(t.text or "" for t in _text_nodes(container))


def _norm_opt(s):
    """Normalise an option label for comparison (case/space/unit insensitive)."""
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())


def _paragraphs_of(container):
    el = _element_of(container)
    if el.tag == qn("w:p"):
        return [el]
    return list(el.iter(qn("w:p")))


def checkbox_slots(container):
    """Every checkbox content control in ``container`` as (sdt, label) pairs.

    ``label`` is the text following that control up to the next one, which is
    how the report prints its option lists ("☐ Basic  ☐ Industrial  ...").
    """
    slots = []
    for p in _paragraphs_of(container):
        current = None
        for child in p:
            if child.tag == qn("w:sdt"):
                if child.find(".//" + _w14("checkbox")) is not None:
                    current = {"sdt": child, "label": ""}
                    slots.append(current)
                    continue
            if child.tag == qn("w:r") and current is not None:
                current["label"] += "".join(t.text or "" for t in child.findall(qn("w:t")))
    return [(s["sdt"], s["label"].strip()) for s in slots]


def literal_checkbox_slots(container):
    """Checkboxes written as plain ``☐`` glyphs rather than content controls.

    The template mixes both: 175 boxes are content controls, but 2.1 EUT
    Category prints "☐ Medical ☐ Laboratory" as literal text. Returns
    (w:t node, offset within that node, label) triples, where the label is the
    text up to the next box.
    """
    nodes = [n for n in _text_nodes(container)
             if not _in_content_control(n)]
    texts = [n.text or "" for n in nodes]
    full = "".join(texts)
    if CHECK_OFF not in full and CHECK_ON not in full:
        return []
    index = [(ni, off) for ni, t in enumerate(texts) for off in range(len(t))]
    boxes = [i for i, ch in enumerate(full) if ch in (CHECK_OFF, CHECK_ON)]
    slots = []
    for bi, pos in enumerate(boxes):
        end = boxes[bi + 1] if bi + 1 < len(boxes) else len(full)
        ni, off = index[pos]
        slots.append((nodes[ni], off, full[pos + 1:end].strip()))
    return slots


def _in_content_control(node):
    anc = node.getparent()
    while anc is not None:
        if anc.tag == qn("w:sdtContent"):
            return True
        anc = anc.getparent()
    return False


def set_literal_checkbox(node, offset, checked):
    """Flip a literal ``☐``/``☒`` glyph at ``offset`` inside ``node``."""
    text = node.text or ""
    if offset >= len(text):
        return False
    node.set(qn("xml:space"), "preserve")
    node.text = text[:offset] + (CHECK_ON if checked else CHECK_OFF) + text[offset + 1:]
    return True


def has_checkboxes(container):
    """True when ``container`` holds checkboxes of either flavour.

    The builder must never overwrite such a cell with ``set_cell_text`` - that
    would delete the controls. Tick it with ``tick_checkboxes`` instead.
    """
    el = _element_of(container)
    if el.find(".//" + _w14("checkbox")) is not None:
        return True
    return bool(literal_checkbox_slots(container))


def set_checkbox(sdt, checked):
    """Set one checkbox content control's state *and* its rendered glyph."""
    cb = sdt.find(".//" + _w14("checkbox"))
    if cb is None:
        return False
    node = cb.find(_w14("checked"))
    if node is None:
        node = cb.makeelement(_w14("checked"), {})
        cb.insert(0, node)
    node.set(_w14("val"), "1" if checked else "0")
    glyph = CHECK_ON if checked else CHECK_OFF
    content = sdt.find(qn("w:sdtContent"))
    if content is not None:
        for t in content.iter(qn("w:t")):
            t.set(qn("xml:space"), "preserve")
            t.text = glyph
            break
    return True


def _label_matches(label, wanted_norm):
    """True when a printed option label corresponds to one of the wanted values."""
    norm = _norm_opt(label)
    if not norm or not wanted_norm:
        return False
    if norm in wanted_norm:
        return True
    # tolerate unit/formatting drift: '±4kV' vs '4kV', '3 V/m' vs '3V/m'
    return any(len(norm) > 2 and len(w) > 2 and (norm in w or w in norm)
               for w in wanted_norm)


def tick_checkboxes(container, selected, multi=False, clear_others=False,
                    slot_range=None):
    """Tick the checkbox(es) whose printed label matches ``selected``.

    Handles both checkbox content controls and literal ``☐`` glyphs.

    selected      : a string, or an iterable of strings for multi-select rows.
    multi         : tick every match rather than stopping at the first.
    clear_others  : explicitly untick non-matching boxes (they start unticked,
                    so this is only needed when re-filling).
    slot_range    : (start, end) to restrict which boxes are considered - used
                    for cells that print two option lists on separate lines
                    (e.g. ESD Indirect Contact Discharge HCP then VCP).
    Returns the list of labels actually ticked.
    """
    wanted = selected if isinstance(selected, (list, tuple, set)) else [selected]
    wanted = [w for w in (str(x).strip() for x in wanted) if w]
    if not wanted:
        return []
    wanted_norm = [_norm_opt(w) for w in wanted]

    cc = checkbox_slots(container)
    if cc:
        slots = [("cc", sdt, None, label) for sdt, label in cc]
    else:
        slots = [("lit", node, off, label)
                 for node, off, label in literal_checkbox_slots(container)]
    if slot_range:
        slots = slots[slot_range[0]:slot_range[1]]

    ticked = []
    for kind, ref, off, label in slots:
        hit = _label_matches(label, wanted_norm)
        if hit:
            if kind == "cc":
                set_checkbox(ref, True)
            else:
                set_literal_checkbox(ref, off, True)
            ticked.append(label)
            if not multi:
                break
        elif clear_others:
            if kind == "cc":
                set_checkbox(ref, False)
            else:
                set_literal_checkbox(ref, off, False)
    return ticked


def fill_custom_slot(container, value):
    """Replace a trailing ``Custom______`` fill-in with ``Custom <value>``.

    Used when a spec value is not one of the printed options - the report keeps
    a ruled blank next to "Custom" for exactly this.
    """
    for node in _text_nodes(container):
        if node.text and "__" in node.text:
            node.set(qn("xml:space"), "preserve")
            node.text = re.sub(r"_{2,}", " " + str(value), node.text)
            return True
    return False


def tick_or_custom(container, value, custom_label="Custom"):
    """Tick ``value``'s box, or tick "Custom" and write the value in its blank.

    Returns True when a printed option matched, False when it fell through to
    the Custom slot (mirrors the datasheet engine's behaviour for these rows).
    """
    value = str(value or "").strip()
    if not value:
        return False
    if tick_checkboxes(container, value):
        return True
    if tick_checkboxes(container, custom_label):
        fill_custom_slot(container, value)
    return False


# ==========================================================================
# Tables
# ==========================================================================

def cell_text(cell):
    return re.sub(r"\s+", " ", (cell.text or "")).strip()


def row_label(row):
    """First-cell text of a table row, whitespace-collapsed."""
    cells = row.cells
    return cell_text(cells[0]) if cells else ""


def distinct_cells(row):
    """The row's cells with horizontally-merged duplicates collapsed.

    python-docx repeats the same ``_Cell`` for every grid column a merged cell
    spans; for writing we only want each real cell once.
    """
    out, seen = [], set()
    for c in row.cells:
        key = id(c._tc)
        if key in seen:
            continue
        seen.add(key)
        out.append(c)
    return out


def _row_is_all_bold(row):
    """True when every non-empty cell of a row is bold (i.e. it is a header).

    The report's observation grids have one, two or three header rows (S. No. /
    Name of test points over +2/-2/+4..., or Test Level over Common Mode over the
    phase angles). They are all set bold, while data rows are not - so boldness
    is a reliable structural signal, and much safer than guessing by content.
    """
    seen_text = False
    for tc in row._tr.findall(qn("w:tc")):
        text = "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
        if not text:
            continue
        seen_text = True
        bold = False
        for r in tc.iter(qn("w:r")):
            rt = "".join(t.text or "" for t in r.findall(qn("w:t"))).strip()
            if not rt:
                continue
            rpr = r.find(qn("w:rPr"))
            if rpr is not None and rpr.find(qn("w:b")) is not None:
                bold = True
            else:
                return False
        if not bold:
            return False
    return seen_text


def header_row_count(table, maximum=3):
    """How many leading rows of ``table`` are header rows.

    Falls back to 1 so a grid with no bold header still gets its first row
    preserved. Capped at ``maximum`` so an all-bold table cannot be mistaken for
    having no data rows.
    """
    n = 0
    for row in table.rows[:maximum]:
        if _row_is_all_bold(row):
            n += 1
        else:
            break
    return max(1, n)


def clone_row(table, template_row_index=-1):
    """Append a copy of an existing row and return it (values left as-is)."""
    tr = copy.deepcopy(table.rows[template_row_index]._tr)
    table._tbl.append(tr)
    return table.rows[-1]


def ensure_row_count(table, count, template_row_index=-1, first_data_row=1):
    """Grow/shrink ``table`` so it has ``count`` data rows after the header."""
    have = len(table.rows) - first_data_row
    while have < count:
        clone_row(table, template_row_index)
        have += 1
    while have > count:
        remove(table.rows[-1])
        have -= 1
    return table.rows[first_data_row:]


def fill_table_rows(table, rows, first_data_row=1, clear_extra=True):
    """Write ``rows`` (list of lists) into ``table`` starting at ``first_data_row``.

    Grows the table by cloning its last row when there is more data than rows.
    When ``clear_extra`` the leftover template rows are removed, so an empty
    dataset leaves just the header.
    """
    rows = [r for r in (rows or []) if any(str(x).strip() for x in r)]
    if not rows:
        if clear_extra:
            for extra in list(table.rows[first_data_row:]):
                remove(extra)
        return 0
    ensure_row_count(table, len(rows), template_row_index=-1,
                     first_data_row=first_data_row) if clear_extra else None
    # (re)fetch after resizing
    for ri, values in enumerate(rows):
        if first_data_row + ri >= len(table.rows):
            clone_row(table)
        cells = distinct_cells(table.rows[first_data_row + ri])
        for ci, val in enumerate(values):
            if ci < len(cells):
                set_cell_text(cells[ci], val)
    return len(rows)


def find_table_after(blocks, start_index):
    """Index of the first Table at/after ``start_index`` in ``blocks``."""
    for i in range(start_index, len(blocks)):
        if isinstance(blocks[i], Table):
            return i
    return -1


TEXT_WIDTH_TWIPS = 9016          # A4 with the report's 1 inch margins


def insert_table_before(paragraph, doc, headers, rows, style="Table Grid",
                        width_twips=TEXT_WIDTH_TWIPS):
    """Build a data table and move it in front of ``paragraph``.

    The report prints "Table 3: ..." captions for the measurement grids but
    ships **without** the grids themselves (they were pasted in by hand), so the
    generator has to create them. Columns are sized equally to the text width so
    the new table matches the ones already in the document.
    """
    ncols = max(len(headers), max((len(r) for r in rows), default=0))
    if not ncols:
        return None
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = style
    except KeyError:                             # style absent from the template
        pass
    table.autofit = False
    col_w = int(width_twips / ncols)

    def _add(values, bold):
        cells = table.add_row().cells
        for i, cell in enumerate(cells):
            tcw = cell._tc.get_or_add_tcPr()
            for old in tcw.findall(qn("w:tcW")):
                tcw.remove(old)
            node = OxmlElement("w:tcW")
            node.set(qn("w:w"), str(col_w))
            node.set(qn("w:type"), "dxa")
            tcw.append(node)
            set_cell_text(cell, values[i] if i < len(values) else "", bold=bold)

    if headers:
        _add(list(headers), True)
        _mark_header_row(table.rows[0])
    for r in rows:
        _add(list(r), False)

    paragraph._p.addprevious(table._tbl)
    return table


def _mark_header_row(row):
    """Repeat this row on every page the table spans (w:tblHeader)."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


# ==========================================================================
# Images
# ==========================================================================

def insert_image_before(paragraph, image, width_mm=None, height_mm=None,
                        max_width_mm=159.2, max_height_mm=95.0):
    """Insert a centred picture paragraph immediately before ``paragraph``.

    ``image`` is a path or a file-like object. When no explicit size is given
    the picture is aspect-fitted into ``max_width_mm`` x ``max_height_mm`` (the
    document's stated 9 cm x 16 cm plot/photo slot, in landscape orientation).
    Returns the new Paragraph, or None when the image could not be read.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        if width_mm and height_mm:
            run.add_picture(image, width=Mm(width_mm), height=Mm(height_mm))
        else:
            w, h = _fit(image, max_width_mm, max_height_mm)
            run.add_picture(image, width=Mm(w), height=Mm(h))
    except Exception:
        remove(p)
        return None
    return p


def _fit(image, max_w_mm, max_h_mm):
    """Aspect-fit an image into a box, in mm. Falls back to the full box."""
    try:
        from PIL import Image
        try:
            image.seek(0)
        except AttributeError:
            pass
        with Image.open(image) as im:
            iw, ih = im.size
        try:
            image.seek(0)
        except AttributeError:
            pass
        if iw and ih:
            scale = min(max_w_mm / float(iw), max_h_mm / float(ih))
            return iw * scale, ih * scale
    except Exception:
        pass
    return max_w_mm, max_h_mm


def add_image_borders(doc, emu=6350, color="000000"):
    """Thin black border around every inline picture (matches the datasheets)."""
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    for pic in doc.element.body.iter(
            "{http://schemas.openxmlformats.org/drawingml/2006/picture}pic"):
        spPr = pic.find("{%s}spPr" % a_ns)
        if spPr is None:
            continue
        if spPr.find("{%s}ln" % a_ns) is not None:
            continue
        ln = spPr.makeelement("{%s}ln" % a_ns, {"w": str(emu)})
        fill = ln.makeelement("{%s}solidFill" % a_ns, {})
        clr = fill.makeelement("{%s}srgbClr" % a_ns, {"val": color})
        fill.append(clr)
        ln.append(fill)
        spPr.append(ln)


# ==========================================================================
# Fonts
# ==========================================================================

_BOX_GLYPHS = (CHECK_OFF, CHECK_ON, "☑")


def _is_glyph_run(r):
    """True for a run that only carries a checkbox glyph.

    Those runs are deliberately set in Segoe UI Symbol / MS Gothic - Arial has no
    dependable ballot-box glyph, so forcing Arial on them would turn the ticks
    into missing-character boxes. They are excluded from the Arial pass.
    """
    text = "".join(t.text or "" for t in r.findall(qn("w:t"))).strip()
    # must actually be a box glyph - a whitespace-only run is ordinary text
    if text and all(ch in _BOX_GLYPHS for ch in text):
        return True
    # a run inside a *checkbox* content control (not just any sdt - the table of
    # contents is an sdt too, and its entries are ordinary text)
    anc = r.getparent()
    while anc is not None:
        if anc.tag == qn("w:sdt"):
            return anc.find(".//" + _w14("checkbox")) is not None
        anc = anc.getparent()
    return False


def enforce_arial(doc, size_pt=DEFAULT_SIZE_PT):
    """Normalise the typeface to Arial without touching any font *size*.

    The spec is "Arial, 11 pt body". Size needs no enforcement: the template's
    ``docDefaults`` are already 11 pt, so every run that carries no explicit size
    renders at 11 pt, and every run that *does* carry one was deliberately set
    that way - the LIST OF FIGURES/PHOTOS/TABLES titles are 14 pt, ABBREVIATIONS
    is 12 pt, headings are 12 pt and captions 10 pt. An earlier version of this
    function forced 11 pt on anything that was not Heading-styled, which silently
    shrank those titles; hence sizes are now left alone.

    Bold/italic are never touched, so template emphasis survives, and checkbox
    glyph runs keep their symbol font (see ``_is_glyph_run``).
    """
    def _fix_run(r):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), DEFAULT_FONT)

    for p in doc.element.body.iter(qn("w:p")):
        for r in p.iter(qn("w:r")):
            if _is_glyph_run(r):
                continue
            _fix_run(r)

    # Document defaults: Arial at the body size, so any run without explicit
    # formatting inherits "Arial 11" instead of the theme font.
    try:
        rpr = doc.styles.element.find(qn("w:docDefaults")) \
            .find(qn("w:rPrDefault")).find(qn("w:rPr"))
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), DEFAULT_FONT)
        for tag in ("w:sz", "w:szCs"):
            node = rpr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                rpr.append(node)
            node.set(qn("w:val"), str(int(size_pt * 2)))
    except Exception:
        pass

    # Body-prose styles that override to 12 pt. The report's Test Procedure,
    # disclaimers and performance criteria use "Normal (Web)", so leaving it at
    # 12 pt would render body text a point too large wherever a run carries no
    # explicit size. Headings/captions are deliberately excluded.
    _normalise_body_styles(doc, size_pt)


_BODY_PROSE_STYLES = ("NormalWeb", "paragraph", "ListParagraph")


def _normalise_body_styles(doc, size_pt=DEFAULT_SIZE_PT):
    for style in doc.styles.element.findall(qn("w:style")):
        if style.get(qn("w:styleId")) not in _BODY_PROSE_STYLES:
            continue
        rpr = style.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            style.append(rpr)
        for tag in ("w:sz", "w:szCs"):
            node = rpr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                rpr.append(node)
            node.set(qn("w:val"), str(int(size_pt * 2)))


# ==========================================================================
# Fields (TOC / list of figures / SEQ / PAGE)
# ==========================================================================

def refresh_fields_on_open(doc):
    """Make Word recompute every field the next time the document is opened.

    The report's table of contents, lists of figures/photos/tables, the
    Figure/Photo/Table numbers (``SEQ``) and "Page X of Y" are all Word fields.
    We cannot lay out pages ourselves, so we set ``w:updateFields`` and let Word
    calculate the real page numbers on open - and we blank the stale cached
    results first (see ``clear_field_cache``) so no wrong number is ever shown.
    """
    settings = doc.settings.element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    node = OxmlElement("w:updateFields")
    node.set(qn("w:val"), "true")
    settings.append(node)


# ==========================================================================
# Revision markup
# ==========================================================================
# The blank form ships with
#     <w:documentProtection w:edit="trackedChanges" w:enforcement="0"/>
# which puts Word into track-changes mode. That is fine for a form somebody
# fills in by hand, but in a GENERATED report it backfires: because we ask Word
# to rebuild the table of contents and the lists of figures/photos/tables on
# open (see refresh_fields_on_open), Word records its own rebuild as revisions
# and shows a column of "Formatted: Default Paragraph Font" balloons - which
# also reserves a wide markup margin and squeezes the page content.
#
# A finished report must carry no revision state at all, so the protection flag
# is removed and any revision elements are accepted (not just hidden).

_REVISION_DROP = ("rPrChange", "pPrChange", "tblPrChange", "trPrChange",
                  "tcPrChange", "sectPrChange", "numberingChange")
_SETTINGS_DROP = ("documentProtection", "trackChanges", "revisionView")


def remove_revision_markup(doc):
    """Turn off change tracking and accept/strip any revision markup.

    Returns a dict of what was removed, for the build summary.
    """
    removed = {}

    # 1. settings: drop the protection/tracking flags
    settings = doc.settings.element
    for name in _SETTINGS_DROP:
        for node in settings.findall(qn("w:" + name)):
            settings.remove(node)
            removed[name] = removed.get(name, 0) + 1

    body = doc.element.body

    # 2. deletions: drop them outright (that is what "accept" means)
    for node in list(body.iter(qn("w:del"))):
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)
            removed["del"] = removed.get("del", 0) + 1
    for node in list(body.iter(qn("w:moveFrom"))):
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)
            removed["moveFrom"] = removed.get("moveFrom", 0) + 1

    # 3. insertions/moves: keep the content, drop the wrapper
    for tag in ("ins", "moveTo"):
        for node in list(body.iter(qn("w:" + tag))):
            parent = node.getparent()
            if parent is None:
                continue
            index = list(parent).index(node)
            for child in reversed(list(node)):
                parent.insert(index, child)
            parent.remove(node)
            removed[tag] = removed.get(tag, 0) + 1

    # 4. formatting-change records carry the OLD formatting; just delete them
    for tag in _REVISION_DROP:
        for node in list(body.iter(qn("w:" + tag))):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
                removed[tag] = removed.get(tag, 0) + 1

    # 5. a deleted paragraph mark lives in pPr/rPr as <w:del/>
    for ppr in list(body.iter(qn("w:pPr"))):
        rpr = ppr.find(qn("w:rPr"))
        if rpr is None:
            continue
        for node in rpr.findall(qn("w:del")):
            rpr.remove(node)
            removed["paraMarkDel"] = removed.get("paraMarkDel", 0) + 1

    return removed


_TOC_ENTRY_STYLES = ("tableoffigures", "toc1", "toc2", "toc3", "toc4", "toc5",
                     "toc6", "toc7", "toc8", "toc9")


def clear_toc_entries(doc):
    """Blank the cached table-of-contents / list-of-figures entry text.

    The shipped template caches the entries of the original 40-page document, so
    a report built for a subset of tests would otherwise list figures, photos and
    headings that are no longer in it (e.g. "Figure 4: RE plot_..." in a
    CE-only report). Word rebuilds all of them from the live document on open
    (see ``refresh_fields_on_open``), so the cached text carries no information -
    only the risk of showing something wrong.

    The paragraphs and the field structure (``fldChar``/``instrText``) are kept
    intact; only the visible entry text is removed. Returns the count cleared.
    """
    cleared = 0
    for p in doc.element.body.iter(qn("w:p")):
        ppr = p.find(qn("w:pPr"))
        if ppr is None:
            continue
        pstyle = ppr.find(qn("w:pStyle"))
        if pstyle is None:
            continue
        style_id = (pstyle.get(qn("w:val")) or "").replace(" ", "").lower()
        if style_id not in _TOC_ENTRY_STYLES:
            continue
        for run in p.findall(qn("w:r")):
            # keep the field machinery, drop the cached caption/heading text
            if run.find(qn("w:fldChar")) is not None or \
                    run.find(qn("w:instrText")) is not None:
                continue
            for t in run.findall(qn("w:t")):
                if (t.text or "").strip():
                    t.text = ""
                    cleared += 1
        # hyperlinks wrap the entry text in the main TOC
        for link in p.findall(qn("w:hyperlink")):
            for t in link.iter(qn("w:t")):
                if (t.text or "").strip():
                    t.text = ""
                    cleared += 1
    return cleared


def clear_field_cache(doc):
    """Blank the cached result text of TOC/PAGEREF/SEQ/PAGE fields.

    A field's cached value sits between ``fldChar separate`` and
    ``fldChar end``. The numbers in the shipped template refer to the original
    40-page document, so they must not survive into a report with a different
    set of tests. Word refills them on open (see ``refresh_fields_on_open``).
    Returns the number of cached runs cleared.
    """
    body = doc.element.body
    cleared = 0
    for p in body.iter(qn("w:p")):
        children = list(p)
        depth = 0                     # nesting depth of begin..end
        caching = False
        instr = ""
        for child in children:
            if child.tag != qn("w:r"):
                continue
            fld = child.find(qn("w:fldChar"))
            if fld is not None:
                ftype = fld.get(qn("w:fldCharType"))
                if ftype == "begin":
                    depth += 1
                    caching = False
                    instr = ""
                elif ftype == "separate":
                    if depth > 0:
                        caching = True
                elif ftype == "end":
                    depth = max(0, depth - 1)
                    caching = False
                    instr = ""
                continue
            itxt = child.find(qn("w:instrText"))
            if itxt is not None:
                instr += itxt.text or ""
                continue
            if caching and depth > 0:
                # keep TOC entry captions (they carry the heading text); only
                # numeric page results and SEQ numbers need clearing
                upper = instr.upper()
                if "PAGEREF" in upper or "SEQ " in upper or "PAGE" in upper \
                        or "NUMPAGES" in upper:
                    for t in child.findall(qn("w:t")):
                        if (t.text or "").strip():
                            t.text = ""
                            cleared += 1
    return cleared
