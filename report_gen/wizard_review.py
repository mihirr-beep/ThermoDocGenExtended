# -*- coding: utf-8 -*-
"""Read-only assembly for the wizard's review pages (2, 3 and 4).

WHY EVERY NUMBER HERE COMES FROM THE BUILDER'S OWN CALL
------------------------------------------------------
``build_report`` starts with ``S.collect(request_obj, planner_entries)``. So does
``preview_source`` below, with the planner entries ordered exactly as
``app.generate_test_report`` orders them. A review page that assembled its own
version of "what the report will say" is the two-lists defect that put 49 of 92
per-test subsections out of agreement with the template - and here it would be
worse, because the admin signs off on what this page shows them.

So: 1.1's Results come from ``builder._verdict``, its spec cell from
``builder._test_method_spec``, the kept/deleted rows from
``registry.TEST_METHOD_ROWS``, approval from ``projection.derive_status``, and
the splice decision from ``splice.region_start``. Nothing is recomputed.

THE THREE THINGS collect() DOES NOT KNOW, AND WHY THEY ARE HERE
--------------------------------------------------------------
1. An unfilled 1.1 cell is now BLANK in the document: fill_summary clears it
   rather than leaving the blank form's example ("CD: <±4kV>AD: <±8kV>") standing
   in a client-facing report. So this page shows it blank too, and there is no
   longer a "from the template" state for 1.1.
2. 1.4 falls back the same way: ``fill_summary`` writes only ``if unc.get(code)``,
   so a kept row with no ``datasheet_fixed_values.measurement_uncertainty`` ships
   the template's printed figure.
3. ``datasheet.revision_no`` is the *next, live* revision - submitting N freezes
   N and moves the row to N+1. The revision peer review actually approved is
   ``MAX(datasheet_revision.revision_no)``, so showing ``revision_no`` would
   over-report by one on every approved test.

Nothing in this module writes anything.
"""
import logging
import os
import re

from . import builder as B
from . import mapping as M
from . import registry as REG
from . import service as S
from . import draft
from . import wizard_fields as WF

log = logging.getLogger(__name__)


# ==========================================================================
# the one source both these pages and build_report() read
# ==========================================================================

def preview_source(request_id):
    """(request_obj, planner_entries, data) - exactly what build_report() sees.

    The entry ORDER matters and is copied from app.generate_test_report:
    resolve_tests() falls back to ``active[-1]`` when no entry sits at
    'datasheet_uploaded', so a review page that ordered its entries differently
    could preview a datasheet the build would not splice.
    """
    from models import PlannerEntry
    from app import _resolve_request

    req = _resolve_request(request_id)
    if req is None:
        return None, [], None
    entries = PlannerEntry.query.filter_by(test_request_id=int(request_id)).order_by(
        PlannerEntry.start_date.asc(), PlannerEntry.start_time.asc(),
        PlannerEntry.id.asc()).all()
    return req, entries, S.collect(req, entries)


def _request_values(request_obj):
    """The six store="request" fields, read off the ORM object.

    getattr rather than another SELECT: preview_source already loaded the row,
    and this database is remote - the cost that matters is round trips, not rows.
    """
    return {f[0]: getattr(request_obj, f[0], None) for f in WF.by_store("request")}


def _merged_form(request_obj, loaded_draft):
    """Draft values with the request's own columns layered on top.

    Same precedence as wizard_routes.eut_page, so the outstanding list on the
    final page cannot disagree with the one on page 1.
    """
    merged = dict((loaded_draft or {}).get("form") or {})
    row = _request_values(request_obj)
    for key, v in row.items():
        if v not in (None, ""):
            merged[key] = v
    return merged, row


# ==========================================================================
# what the blank template already prints, and what survives cleanup
# ==========================================================================

_TEMPLATE_FALLBACK = None


def _template_fallbacks():
    """{"method": {code: {spec, result}}, "uncertainty": {code: value}}.

    Read POSITIONALLY out of the pristine template: its 1.1 data rows are in
    TEST_METHOD_ROWS order and its 1.4 rows in UNCERTAINTY_ROWS order (verified).
    Matching by label instead would mean keeping a second copy of
    fill_summary's row matcher, which is exactly the drift this wizard avoids.
    A row whose label has moved is skipped rather than guessed at.

    Cached: it is one file read that never changes while the process lives.
    """
    global _TEMPLATE_FALLBACK
    if _TEMPLATE_FALLBACK is not None:
        return _TEMPLATE_FALLBACK
    out = {"method": {}, "uncertainty": {}}
    try:
        from docx import Document
        from . import docx_tools as T
        outline = B.Outline(Document(REG.TEMPLATE_PATH))

        tabs = outline.tables_in("TEST REPORT SUMMARY", "TEST METHOD")
        if tabs:
            for (label, code), row in zip(REG.TEST_METHOD_ROWS, tabs[0].rows[1:]):
                if M.norm_label(T.row_label(row))[:12] != M.norm_label(label)[:12]:
                    continue                       # template drift - do not guess
                cells = T.distinct_cells(row)
                if len(cells) >= 4:
                    out["method"][code] = {
                        "spec": T.full_text(cells[1]).strip(),
                        "result": T.full_text(cells[3]).strip()}

        tabs = outline.tables_in("TEST REPORT SUMMARY", "MEASUREMENT UNCERTAINITY")
        if tabs:
            for (label, code), row in zip(REG.UNCERTAINTY_ROWS, tabs[0].rows[1:]):
                if M.norm_label(T.row_label(row))[:12] != M.norm_label(label)[:12]:
                    continue
                cells = T.distinct_cells(row)
                if len(cells) >= 2:
                    out["uncertainty"][code] = T.full_text(cells[-1]).strip()
    except Exception as exc:  # noqa: BLE001 - a preview must not need the template
        log.info("wizard review: template fallbacks unavailable: %s", exc)
    _TEMPLATE_FALLBACK = out
    return out


def _after_cleanup(text):
    """What an untouched template cell actually prints, once cleanup has run.

    Uses the builder's own predicates - _is_instruction, _INLINE_NOISE,
    NOT_APPLICABLE - so this cannot drift from the pass that really does it
    (builder.cleanup_instructions, the cell loop).
    """
    txt = (text or "").strip()
    if not txt:
        return ""
    if B._is_instruction(txt):
        return B.NOT_APPLICABLE
    cleaned = txt
    for pattern, repl in B._INLINE_NOISE:
        cleaned = pattern.sub(repl, cleaned)
    if cleaned == txt:
        return txt
    return re.sub(r"[ \t]{2,}", " ", cleaned).strip() or B.NOT_APPLICABLE


def _sourced(value, fallback=None):
    """(text, source) for one 1.1 cell: "datasheet", or empty.

    THERE IS NO LONGER A "template" SOURCE. This used to report the blank form's
    own example - "CD: <±4kV>AD: <±8kV>" for ESD - because fill_summary left it
    standing when no datasheet supplied the value, and the page had to show what
    the document would really print.

    fill_summary now clears those cells instead, so an unsourced cell is empty in
    the document and empty here. ``fallback`` is still accepted so the 1.4 caller,
    where the template's printed uncertainty IS the laboratory's real figure, does
    not have to change shape.
    """
    if value:
        return value, "datasheet"
    return "", ""


def _sourced_or_template(value, fallback):
    """(text, source) where the template's own printed value DOES still survive.

    Only 1.4 uses this. fill_summary writes an uncertainty only when the fixed
    values carry one, so the figure the blank form prints - the laboratory's own
    +/- dB - is what a reader gets otherwise. That is a real value the lab stands
    behind, not a placeholder like 1.1's "<±4kV>", so it is shown, marked as
    coming from the template rather than from a datasheet.
    """
    if value:
        return value, "datasheet"
    shipped = _after_cleanup(fallback)
    if shipped:
        return shipped, "template"
    return "", ""


# ==========================================================================
# page 1: the cover page, every row of it
# ==========================================================================

# The address block printed under ISSUED BY. It is the laboratory's own name and
# address, identical on every report this system will ever produce, and it is
# already in the template - this copy exists so the wizard can SHOW the admin
# what that row will say. builder.py does not read it: nothing writes that row,
# because the template already carries it.
ISSUED_BY = (
    "Thermo Fisher Scientific Product Testing Laboratory\n"
    "Warehouse No 1A Plot No 6, Survey No 315 434/1,\n"
    "All Cargo Logistics & Industrial Park,\n"
    "Ramachandrapuram Mandal, Velmula Village,\n"
    "Sangareddy 502300, Telangana, India"
)


def _span_text(meta):
    """The DATES ON WHICH TESTS WERE PERFORMED row, as fill_cover composes it."""
    frm, to = meta.get("tests_from"), meta.get("tests_to")
    if not frm:
        return ""
    if to and to != frm:
        return "From %s to %s" % (frm, to)
    return "On %s" % frm


def cover_preview(request_id, request_obj=None, data=None, loaded_draft=None):
    """Every row of the cover table, in the template's order, with its source.

    WHY THE READ-ONLY ROWS ARE HERE AND NOT JUST IN THE DOCUMENT
    -----------------------------------------------------------
    The wizard used to show the five rows the admin can type and nothing else,
    so the front page of a client-facing report was two-thirds invisible until
    the .docx existed. An admin cannot notice that the serial number is wrong on
    a page they were never shown.

    So every row is listed. ``editable`` says whether this page can change it;
    the rest are shown with where the value comes from, and are changed by fixing
    the test request or the datasheets rather than here.

    Values match what build_report will print because they come from the same
    ``S.collect`` meta, with the draft layered on exactly where draft_fill layers
    it.
    """
    if data is None or request_obj is None:
        request_obj, _entries, data = preview_source(request_id)
    if data is None:
        return []
    meta = data["meta"]
    form = (loaded_draft or draft.load(request_id)).get("form") or {}

    def drafted(key, fallback="", is_date=False):
        v = str(form.get(key) or "").strip()
        if v:
            return S.fmt_date(v) if is_date else v
        return fallback

    return [
        {"label": "MANUFACTURER", "value": meta["manufacturer"], "editable": None},
        {"label": "ADDRESS", "value": meta["manufacturer_address"], "editable": None},
        {"label": "EUT NAME", "value": meta["eut_name"], "editable": None},
        {"label": "EUT MODEL/SKU NUMBER", "value": meta["eut_model"], "editable": None},
        {"label": "EUT SERIAL NUMBER", "value": meta["eut_serial"], "editable": None},
        {"label": "CONDITION OF EUT ON RECEIPT",
         "value": drafted("condition_on_receipt", meta["sample_condition"]),
         "editable": "condition_on_receipt"},
        {"label": "DATE OF RECEIPT OF EUT",
         "value": drafted("date_of_receipt", meta["sample_received"], is_date=True),
         "editable": "date_of_receipt"},
        {"label": "DATES ON WHICH TESTS WERE PERFORMED",
         "value": _span_text(meta), "editable": None},
        {"label": "LOCATION OF PERFORMANCE OF TEST",
         "value": drafted("test_location", "Permanent"), "editable": "test_location"},
        {"label": "TEST REPORT ISSUE DATE",
         "value": drafted("report_issue_date", meta["issue_date"], is_date=True),
         "editable": "report_issue_date"},
        {"label": "ISSUED TO: NAME AND CONTACT INFORMATION OF CUSTOMER",
         "value": drafted("issued_to"), "editable": "issued_to"},
        {"label": "ISSUED BY: NAME AND ADDRESS OF TEST LABORATORY",
         "value": ISSUED_BY, "editable": None},
    ]


def signature_preview(request_id, request_obj=None, data=None, loaded_draft=None):
    """The three signature columns: who signs, on what date, with which picture.

    Name comes from the database - whoever submitted the datasheets, whoever peer
    reviewed them, whoever manages the laboratory. Date and Signature are the
    admin's, because they record an act rather than a fact already stored.
    """
    if data is None or request_obj is None:
        request_obj, _entries, data = preview_source(request_id)
    if data is None:
        return []
    meta = data["meta"]
    d = loaded_draft or draft.load(request_id)
    form, images = d.get("form") or {}, d.get("images") or {}
    cols = (("Prepared By", meta["prepared_by"], "prepared"),
            ("Reviewed By", meta["reviewed_by"], "reviewed"),
            ("Authorized Signatory", meta["lab_manager_name"], "authorized"))
    return [{"column": col, "name": name,
             "date_key": "sign_date_%s" % suffix,
             "date": S.fmt_date(str(form.get("sign_date_%s" % suffix) or "").strip()),
             "image_key": "img_sign_%s" % suffix,
             "has_image": bool(str(images.get("img_sign_%s" % suffix) or "").strip())}
            for col, name, suffix in cols]


# ==========================================================================
# page 1: section 2, every subsection of it
# ==========================================================================

# The 2.1 rows that no wizard field touches. Size, weight, operating frequency,
# power rating and measured current are inputs on this page and are deliberately
# not repeated here.
_EUT_DETAIL_READONLY = (
    ("Manufacturer", "manufacturer"),
    ("EUT Name", "eut_name"),
    ("EUT Model/SKU Number", "eut_model"),
    ("EUT Serial Number", "eut_serial"),
    ("Number of Test Samples", "test_samples"),
    ("EUT Operating Voltage", "operating_voltage"),
)


def eut_detail_rows(data):
    """The 2.1 rows that come from the request and are not editable here."""
    meta = data["meta"]
    rows = [{"label": label, "value": meta.get(key) or "", "source": "test request"}
            for label, key in _EUT_DETAIL_READONLY]
    cats = ", ".join(meta.get("categories") or [])
    rows.append({"label": "EUT Category", "value": cats, "source": "test request"})
    rows.append({"label": "Type of Equipment", "value": meta.get("product_type") or "",
                 "source": "test request"})
    return rows


# Centre a column only when the WHOLE column is short - an index, a state number,
# a "Power"/"Signal". Deciding per cell put "FB900 Monitor" in the middle of a
# column whose other value, "EMC32 Measurement Suite", sat on the left.
_NARROW = 12


def _col_align(headers, rows):
    """One CSS class per column, from the widest value in that column."""
    out = []
    for i in range(len(headers or [])):
        vals = [str(r[i]).strip() for r in (rows or [])
                if i < len(r) and str(r[i]).strip()]
        widest = max((len(v) for v in vals), default=0)
        out.append("rw-mid" if 0 < widest <= _NARROW else "")
    return out


def _drop_empty_rows(rows):
    """Rows where every cell is blank say nothing; the gap is the message."""
    return [r for r in (rows or []) if any(str(c).strip() for c in r)]


def section2_preview(request_id, request_obj=None, data=None, loaded_draft=None):
    """2.2 to 2.9, in document order, with the value each will actually print.

    WHY THIS REPLACED FOUR TEXTAREAS
    --------------------------------
    The wizard used to offer free-text boxes for 2.3, 2.5, 2.7 and 2.8 under a
    heading saying all four print as NA because nothing supplies them. Three of
    those four statements had stopped being true:

      2.3 is a table built from every test's datasheet_software rows.
      2.7 is the request's functional modes, labelled Mode A/B/C by position.
      2.5 and 2.8 are request columns, filled on ten of the thirty requests.

    So an admin was being asked to retype content the report already had, and
    the completeness banner counted four already-answered questions as
    outstanding. What is derived is now SHOWN, with its real value; what is
    typed is still typed, and lands on the request rather than in a private copy.
    """
    if data is None or request_obj is None:
        request_obj, _entries, data = preview_source(request_id)
    if data is None:
        return []
    meta = data["meta"]
    d = loaded_draft or draft.load(request_id)
    images = d.get("images") or {}

    def has(key):
        return bool(str(images.get(key) or "").strip())

    mods = _drop_empty_rows(meta.get("modifications") or [])
    software = _drop_empty_rows(meta.get("software_rows") or [])
    accessories = _drop_empty_rows(meta.get("accessories") or [])
    cables = _drop_empty_rows(meta.get("cables") or [])
    # "pics" is the uniform shape for image slots, so the template renders the
    # real upload control - file input, Edit, document-shaped thumbnail - inside
    # the subsection the picture belongs to. It used to print a status line and a
    # link to an Images card at the bottom of the page, which made 2.6 read as
    # broken: the one subsection whose entire content IS a picture showed no way
    # to supply one.
    return [
        {"no": "2.2", "title": "DESCRIPTION OF EUT", "kind": "text",
         "value": meta.get("description") or ""},
        {"no": "2.3", "title": "SOFTWARE AND FIRMWARE DETAILS", "kind": "table",
         "caption": "Software and firmware recorded on each test's datasheet",
         "headers": ["Test", "Software / Firmware", "Version"],
         "rows": software,
         "align": _col_align(["Test", "Software / Firmware", "Version"], software)},
        {"no": "2.4", "title": "EUT MODIFICATION RECORD", "kind": "table",
         "caption": "Modification states across every test on this request",
         "headers": ["State", "Description of modification still fitted",
                     "Fitted by", "Date fitted"],
         "rows": mods,
         "align": _col_align(["State", "Description of modification still fitted",
                              "Fitted by", "Date fitted"], mods)},
        {"no": "2.5", "title": "EUT CONFIGURATION DURING TEST", "kind": "text",
         "value": meta.get("configuration") or "",
         "editable": "test_configuration"},
        {"no": "2.6", "title": "EUT SETUP DETAILS", "kind": "image",
         "pics": [{"key": "img_block_diagram",
                   "label": "Figure 1 - block diagram of the EUT setup",
                   "has_image": has("img_block_diagram")}]},
        # "lines", not "items": Jinja resolves b.items to dict.items before it
        # ever looks for a key of that name, so the template rendered a builtin
        # and raised "object is not iterable".
        {"no": "2.7", "title": "EUT MODES OF OPERATION", "kind": "list",
         "lines": meta.get("modes") or []},
        {"no": "2.8", "title": "EUT MONITORING PARAMETERS", "kind": "text",
         "value": meta.get("monitoring") or "",
         "editable": "monitoring_parameters",
         "pics": [{"key": "img_monitoring",
                   "label": "Screenshot of the monitoring software (optional)",
                   "has_image": has("img_monitoring")}]},
        {"no": "2.9", "title": "ACCESSORIES, CABLES AND PICTURES", "kind": "table",
         "caption": "Table 1 - List of accessories used for testing",
         "headers": ["S. No.", "Accessory", "Make", "Model No.", "Serial No."],
         "rows": accessories,
         "align": _col_align(["S. No.", "Accessory", "Make", "Model No.",
                              "Serial No."], accessories),
         "extra": {"caption": "Table 2 - List of cables connected to the EUT",
                   "headers": ["S. No.", "Cable", "Length (m)", "Power/Signal",
                               "Shielded"],
                   "rows": cables,
                   "align": _col_align(["S. No.", "Cable", "Length (m)",
                                        "Power/Signal", "Shielded"], cables)},
         "pics": [{"key": "img_eut_photo", "label": "Photo 1 - the EUT",
                   "has_image": has("img_eut_photo")},
                  {"key": "img_eut_label",
                   "label": "Photo 2 - model / serial label",
                   "has_image": has("img_eut_label")}]},
    ]


# ==========================================================================
# page 2: section 1
# ==========================================================================

def method_rows(data):
    """1.1 TEST METHOD - exactly the rows fill_summary() will keep, in its order."""
    meta, tests = data["meta"], data["tests"]
    by_code = {t["code"]: t for t in tests}
    tpl = _template_fallbacks()["method"]
    rows = []
    for label, code in REG.TEST_METHOD_ROWS:
        t = by_code.get(code)
        if t is None:
            continue                        # fill_summary removes this row entirely
        fb = tpl.get(code) or {}
        try:
            spec = B._test_method_spec(code, t["form"], meta)
        except Exception as exc:  # noqa: BLE001 - one odd datasheet, not the page
            log.info("wizard review: 1.1 spec failed for %s: %s", code, exc)
            spec = ""
        try:
            verdict = B._verdict(code, t["form"])
        except Exception as exc:  # noqa: BLE001
            log.info("wizard review: 1.1 verdict failed for %s: %s", code, exc)
            verdict = ""
        spec_text, spec_src = _sourced(spec, fb.get("spec"))
        res_text, res_src = _sourced(verdict, fb.get("result"))
        rows.append({"code": code, "test": label,
                     "spec": spec_text, "spec_source": spec_src,
                     "port": REG.TEST_METHOD_PORT.get(code, ""),
                     "result": res_text, "result_source": res_src})
    return rows


def omitted_method_rows(data):
    """The 1.1 rows fill_summary() will delete, so the admin sees the shrink."""
    codes = {t["code"] for t in data["tests"]}
    return [{"code": code, "test": label}
            for label, code in REG.TEST_METHOD_ROWS if code not in codes]


def standards(data):
    """1.2 APPLICABLE STANDARDS as (product_standards, basic_standards).

    _fill_standards() computes and writes in one function, so there is no helper
    to call for the list alone; this repeats its split-and-dedup loop. The clean
    fix is to extract standards_lists() from _fill_standards and have both call
    it - out of scope here, this phase adds no edits to builder.py.
    """
    meta, tests = data["meta"], data["tests"]
    products = list(meta.get("product_standards") or [])
    basics, seen = [], set()
    for t in tests:
        for part in re.split(r"\s*[&;]\s*", M._val(t["form"], "basic_standard")):
            part = part.strip()
            if part and part.lower() not in seen:
                seen.add(part.lower())
                basics.append(part)
    return products, basics


def uncertainty_rows(data):
    """1.4 MEASUREMENT UNCERTAINITY - kept rows only, with their printed value."""
    codes = {t["code"] for t in data["tests"]}
    unc, tpl = data["uncertainty"], _template_fallbacks()["uncertainty"]
    rows = []
    for label, code in REG.UNCERTAINTY_ROWS:
        if code not in codes:
            continue                        # fill_summary removes this row
        value, source = _sourced_or_template(unc.get(code, ""), tpl.get(code))
        rows.append({"code": code, "test": label, "value": value,
                     "source": source})
    return rows


_SECTION1_STATIC = None


def section1_static():
    """The FIXED text of section 1, read out of the blank template.

    1.1 carries five disclaimer paragraphs under its table - what the results
    relate to, that the report may not be reproduced except in full, that the
    laboratory is not responsible for customer-supplied data or for sampling, and
    that manufacturer/configuration/criteria come from the request. 1.3 is the
    laboratory's NABL accreditation and TC number. None of it varies by request.

    Read from the template rather than pasted here. This page exists so an admin
    can see what the report will say, and a second copy of a legal disclaimer is a
    copy that can disagree with the document the moment the lab edits the form.

    Cached: one file read that cannot change while the process lives.
    """
    global _SECTION1_STATIC
    if _SECTION1_STATIC is not None:
        return _SECTION1_STATIC
    out = {"disclaimers": [], "accreditation": {"intro": "", "rows": []}}
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from . import docx_tools as T
        outline = B.Outline(Document(REG.TEMPLATE_PATH))

        # 1.1: the paragraphs AFTER the table. Everything before it is the
        # heading; the table itself is previewed row by row already.
        seen_table = False
        for b in outline.sub_blocks("TEST REPORT SUMMARY", "TEST METHOD"):
            if isinstance(b, Table):
                seen_table = True
                continue
            if seen_table and isinstance(b, Paragraph):
                txt = T.text_of(b).strip()
                if txt:
                    out["disclaimers"].append(txt)

        for b in outline.sub_blocks("TEST REPORT SUMMARY", "ACCREDITATION DETAILS"):
            if isinstance(b, Paragraph):
                txt = T.text_of(b).strip()
                if txt and not txt.upper().startswith("ACCREDITATION DETAILS"):
                    out["accreditation"]["intro"] = txt
            elif isinstance(b, Table):
                for row in b.rows:
                    cells = [T.full_text(c).strip() for c in T.distinct_cells(row)]
                    if any(cells):
                        out["accreditation"]["rows"].append(cells)
    except Exception as exc:  # noqa: BLE001 - a preview must not need the template
        log.info("wizard review: section 1 static text unavailable: %s", exc)
    _SECTION1_STATIC = out
    return out


def summary_preview(request_id):
    """Everything section 1 will contain, as the document will contain it."""
    req, entries, data = preview_source(request_id)
    if req is None:
        return {"ok": False, "reason": "Request not found"}
    products, basics = standards(data)
    rows = method_rows(data)
    unc = uncertainty_rows(data)
    # These are read by a test engineer, not by whoever wrote the builder. The
    # first version of them named a private function (_fill_standards) and
    # described the blank form's internals, which tells the reader nothing they
    # can act on. Each note now says what the report will look like and what to do
    # about it.
    notes = []
    if not rows:
        notes.append("No test on this request has any datasheet data yet, so 1.1 "
                     "would have no rows. Generating will be refused until at "
                     "least one datasheet is complete.")
    if any(r["spec_source"] == "template" or r["result_source"] == "template"
           for r in rows):
        notes.append("Some cells below are marked \"not from a datasheet\". No "
                     "datasheet supplies them, so the report will print the blank "
                     "form's example text or \"NA\" there. Fix those on the "
                     "datasheet - they cannot be typed on this page.")
    if not products and not basics:
        notes.append("No standards are recorded on the request or on any datasheet, "
                     "so 1.2 will print the blank form's example standards. Add them "
                     "to the test request.")
    if not unc:
        notes.append("1.4 lists measurement uncertainty for emission tests only "
                     "(Conducted, Radiated, Harmonic, Flicker). This request has "
                     "none of those, so 1.4 will print its heading and lead-in "
                     "sentence with no table.")
    return {
        "ok": True,
        "tco_id": S._s(getattr(req, "tco_id", "")),
        "product": S._s(getattr(req, "product_name", "")),
        "method": rows,
        "omitted": omitted_method_rows(data),
        "standards": {
            "products": products, "basics": basics,
            # The cells builder._fill_standards will write, from that function's
            # own pairing - so the page shows the table the report will have,
            # including any cell the pairing leaves empty. Two loose lists could
            # not show that, and an empty cell in a client-facing table is
            # something the admin should see here rather than in the .docx.
            "grid": B.standards_cells(products, basics),
            "blanks": sum(1 for row in B.standards_cells(products, basics)
                          for c in row if not c),
        },
        "uncertainty": unc,
        "codes": data["codes"],
        "skipped": data["skipped"],
        "notes": notes,
        "entries": len(entries),
        # 1.1's five disclaimer paragraphs and 1.3's accreditation table, from
        # the template. Fixed on every report, and shown because "read-only" is
        # not a reason to hide what the page will carry.
        "static": section1_static(),
    }


# ==========================================================================
# page 3: section 3 (the criteria note) and sections 4..N (one per test)
# ==========================================================================

def _lifecycle(tests):
    """{planner_entry_id: datasheet projection row} in ONE query.

    One round trip for every test rather than one per test, and best-effort: the
    `datasheet` table is a queryable mirror of form_json, not the truth, so a
    review page must still render when it is missing or stale.
    """
    ids = [getattr(t["entry"], "id", None) for t in tests
           if t.get("entry") is not None]
    ids = [i for i in ids if i]
    if not ids:
        return {}
    from sqlalchemy import text
    from models import db
    sql = ("SELECT d.id, d.planner_entry_id, d.status, d.revision_no, d.result, "
           "d.decided_at, d.peer_reviewer_name, "
           "(SELECT MAX(v.revision_no) FROM datasheet_revision v "
           " WHERE v.datasheet_id = d.id) AS approved_revision "
           "FROM `datasheet` d WHERE d.planner_entry_id IN (%s)"
           % ", ".join(":p%d" % i for i in range(len(ids))))
    try:
        rows = db.session.execute(
            text(sql), {"p%d" % i: v for i, v in enumerate(ids)}).mappings()
        return {r["planner_entry_id"]: dict(r) for r in rows}
    except Exception as exc:  # noqa: BLE001
        log.info("wizard review: datasheet projection unavailable: %s", exc)
        return {}


def _splice_check(path, code):
    """(will_splice, note). Cheap pre-check of the one failure extract_region raises.

    build_report splices when datasheet_path is set and silently falls back to
    fill_test_section on ANY exception, so without this the page could promise
    "your approved pages will be used" and be wrong.
    """
    from . import splice as SPL
    if not path:
        return False, "no generated datasheet on disk"
    try:
        from docx import Document
        if SPL.region_start(Document(path), code) is not None:
            return True, ""
        return False, ("no %r heading in the file - the report falls back to its "
                       "own blank template section"
                       % SPL.TEST_HEADING.get(code, code))
    except Exception as exc:  # noqa: BLE001
        return False, "the file cannot be read (%s)" % exc


def criteria_note(data):
    """Section 3 - what tick_decision_rules() will and will not mark.

    3.1 PERFORMANCE CRITERIA is static prose and no code touches it; 3.2 DECISION
    RULE is ticked from meta["decision_rules"], and tick_decision_rules() returns
    immediately on an empty list. Saying so is this page's only real job for
    section 3 - an untricked 3.2 is otherwise invisible until someone reads the
    finished PDF.
    """
    selected = list(data["meta"].get("decision_rules") or [])
    labels = B._DECISION_RULE_LABELS
    ticked = [labels.get(r, r) for r in selected]
    ticked_norm = {M.norm_label(x) for x in ticked}
    return {
        "ticked": ticked,
        "unticked": [v for v in labels.values()
                     if M.norm_label(v) not in ticked_norm],
        "none": not ticked,
        "criteria_letters": "A, B and C",
        # The template has no paragraph for criterion D, yet _verdict prints
        # whatever the datasheet recorded - request 15's EFT reports D.
        "no_d_paragraph": True,
    }


def tests_preview(request_id):
    """Per test: datasheet, revision, result, approval, splice-or-fallback."""
    req, _entries, data = preview_source(request_id)
    if req is None:
        return {"ok": False, "reason": "Request not found"}
    from datasheet_gen import projection as P

    lifecycle = _lifecycle(data["tests"])
    out = []
    for t in data["tests"]:
        entry, rec = t.get("entry"), (t.get("record") or {})
        life = lifecycle.get(getattr(entry, "id", None), {})
        path = t.get("datasheet_path")
        will_splice, note = _splice_check(path, t["code"])
        try:
            status = P.derive_status(entry, rec)
        except Exception as exc:  # noqa: BLE001
            log.info("wizard review: derive_status failed for %s: %s", t["code"], exc)
            status = "Unknown"
        required = M._val(t["form"], "required_performance_criteria")
        met = M._val(t["form"], "met_performance_criteria")
        try:
            verdict = B._verdict(t["code"], t["form"])
        except Exception:  # noqa: BLE001
            verdict = ""
        out.append({
            "code": t["code"],
            "section": t["section"],
            "name": t["name"],
            "form_no": t.get("form_no"),
            "datasheet": os.path.basename(path) if path else None,
            "datasheet_path": path,
            # the frozen revision peer review saw, not datasheet.revision_no
            "revision": life.get("approved_revision"),
            "live_revision": life.get("revision_no"),
            "result": verdict,
            "recorded_result": S._s(life.get("result") or rec.get("result") or ""),
            "approved": status == "Approved",
            "review_status": status,
            "entry_status": getattr(entry, "status", None),
            "reviewer": S._s(life.get("peer_reviewer_name") or ""),
            "decided_at": life.get("decided_at"),
            "has_data": t["has_data"],
            "images": len(t.get("images") or {}),
            "will_splice": will_splice,
            "source": ("the approved datasheet's own pages are spliced in"
                       if will_splice
                       else "filled into the report's blank template section"),
            "note": note,
            "required_criterion": required,
            "met_criterion": met,
            # Nothing in the builder compares these two, so a test that met a
            # weaker criterion than it required still prints its letter in 1.1
            # with nothing marking it as a miss.
            "criterion_mismatch": bool(required and met
                                       and required.strip().upper() != met.strip().upper()),
        })
    return {"ok": True,
            "tco_id": S._s(getattr(req, "tco_id", "")),
            "product": S._s(getattr(req, "product_name", "")),
            "tests": out,
            "skipped": data["skipped"],
            "criteria": criteria_note(data)}


# ==========================================================================
# page 4: can the existing endpoint actually succeed?
# ==========================================================================

def latest_report(request_id, entries=None):
    """The most recent generated report .docx for this request, or None.

    Two sources because either can be the newer one: the planner entries carry
    report_file_path from the last generate, and the directory holds every run.
    """
    from flask import current_app
    candidates = []
    for e in entries or []:
        p = getattr(e, "report_file_path", None)
        if p and os.path.exists(p):
            candidates.append(p)
    try:
        folder = os.path.join(
            current_app.config.get("UPLOAD_FOLDER", "uploads"), "reports",
            str(int(request_id)))
        if os.path.isdir(folder):
            candidates.extend(os.path.join(folder, n) for n in os.listdir(folder)
                              if n.lower().endswith(".docx"))
    except Exception as exc:  # noqa: BLE001 - a missing folder is not an error
        log.info("wizard review: report folder unreadable for %s: %s",
                 request_id, exc)
    if not candidates:
        return None
    return max(set(candidates), key=lambda p: os.path.getmtime(p))


def readiness(request_id):
    """{"ready", "blockers", "outstanding"} + what the final page needs to say.

    blockers are the generate endpoint's OWN preconditions, evaluated without
    calling it - no planner entries (404), no submitted datasheet record (400),
    no resolvable tests (build_report raises ValueError -> 400). outstanding is
    wizard_fields.outstanding(): field gaps, which do NOT block a build, they
    just come out blank or NA. Conflating the two would either refuse to generate
    a report the lab is entitled to, or promise a complete one that is not.
    """
    from . import render as R

    req, entries, data = preview_source(request_id)
    if req is None:
        return {"ready": False, "blockers": ["Request not found"],
                "outstanding": [], "ok": False, "reason": "Request not found"}

    from datasheet_gen import records as DR
    active = [e for e in entries
              if str(getattr(e, "status", "") or "").strip().lower() != S.CANCELLED]
    with_data = []
    for e in active:
        try:
            if DR.get_record_for_assignment(e.id):
                with_data.append(e)
        except Exception as exc:  # noqa: BLE001
            log.info("wizard review: record lookup failed for entry %s: %s",
                     e.id, exc)

    blockers = []
    if not entries:
        blockers.append("This request has no planner entries, so there is nothing "
                        "to report on (the endpoint returns 404).")
    if not with_data:
        blockers.append("No submitted datasheet data was found. Complete and "
                        "approve the datasheets first (the endpoint returns 400).")
    if not data["tests"]:
        blockers.append("No test resolved to a section, so build_report raises "
                        "\"no completed tests\" (the endpoint returns 400).")

    d = draft.load(request_id)
    merged, row = _merged_form(req, d)
    left = WF.outstanding(merged, row, d.get("images"))

    warnings = []
    unapproved = []
    from datasheet_gen import projection as P
    for t in data["tests"]:
        try:
            if P.derive_status(t.get("entry"), t.get("record") or {}) != "Approved":
                unapproved.append(t["code"])
        except Exception:  # noqa: BLE001
            unapproved.append(t["code"])
    if unapproved:
        warnings.append("Not peer-approved yet: %s. The report will still be built "
                        "from whatever those datasheets currently hold."
                        % ", ".join(unapproved))
    no_data = [t["code"] for t in data["tests"] if not t["has_data"]]
    if no_data:
        warnings.append("No saved datasheet data for %s - those sections ship as "
                        "the blank form." % ", ".join(no_data))
    no_file = [t["code"] for t in data["tests"] if not t.get("datasheet_path")]
    if no_file:
        warnings.append("No generated datasheet on disk for %s, so the report "
                        "fills its own template section instead of splicing the "
                        "approved pages." % ", ".join(no_file))
    if data["skipped"]:
        warnings.append("Dropped from the report: %s."
                        % "; ".join("%s (%s)" % (s["code"], s["reason"])
                                    for s in data["skipped"]))
    if not (data["meta"].get("decision_rules") or []):
        warnings.append("The request selected no decision rule, so 3.2 DECISION "
                        "RULE ships with nothing ticked.")

    already = [e for e in entries if getattr(e, "report_file_path", None)]
    existing = latest_report(request_id, entries)

    return {
        "ok": True,
        "ready": not blockers,
        "blockers": blockers,
        "outstanding": left,
        "warnings": warnings,
        "tco_id": S._s(getattr(req, "tco_id", "")),
        "product": S._s(getattr(req, "product_name", "")),
        "filled": WF.filled_count(merged, row, d.get("images")),
        "total": len(WF.FIELDS),
        "images": {k: os.path.basename(v)
                   for k, v in (d.get("images") or {}).items() if v},
        "tests": len(data["tests"]),
        "entries": len(entries),
        "with_data": len(with_data),
        # The endpoint is repeatable and has no status gate of its own, and
        # request 15 already holds three reports from repeat calls. Say so rather
        # than letting Generate look like a first run.
        "already_generated": bool(already),
        "existing_report": os.path.basename(existing) if existing else None,
        # The basename is what a reader recognises, but the download endpoint
        # takes a path - so carry both rather than making the template
        # reconstruct one from the other.
        "existing_report_path": existing or None,
        "pdf_backend": R.backend(),
        "pdf_available": R.available(),
        "pdf_note": None if R.available() else R.unavailable_note(),
        "page_reached": d.get("page", 1),
    }
