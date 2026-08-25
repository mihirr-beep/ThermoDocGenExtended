# -*- coding: utf-8 -*-
"""Compute the report's fields on the server, so Word never has to.

WHY
---
The builder blanks the template's cached table-of-contents entries (they belong
to a different 40-page document) and sets ``w:updateFields`` so Word rebuilds
them on open. That works, and it has three costs the reader pays:

  * "This document contains fields that may refer to other files" on open;
  * "Update Table of Figures - page numbers only / entire table", four times,
    where the harmless-looking default leaves the lists BLANK because the
    cached entries were deliberately cleared;
  * and the one that actually breaks the deliverable: asking Word to modify the
    document on open means the modification can be RECORDED. With track changes
    active the rebuild lands as revisions - a contents page in red underline, and
    "Field Code Changed" balloons in a markup margin down the side of the page.

No amount of cleaning the file prevents that third one, because the file is
clean: the balloons are created after opening, by Word, out of the rebuild we
asked for. The fix is to stop asking. Compute the fields here, drop the
``w:updateFields`` flag, and the document that reaches the reader is finished -
no prompts, no rebuild, nothing to record.

HOW
---
Something with a LAYOUT ENGINE does the work, because knowing that a heading
falls on page 31 means laying the document out and Python cannot. Word over COM
on a dev box; LibreOffice over its UNO bridge on the Linux server, which is what
makes production produce the same finished document rather than a chatty one.
Neither present: the entry text is written in pure Python and the page numbers
are left to the reader's Word.

Best-effort at every tier - a failure leaves the file exactly as the builder
wrote it, because a report with un-computed fields is worse to read but still
correct, and losing the report entirely is not an acceptable trade for tidier
page numbers. Never raises into the request. ``finalise()`` is the entry point
and its docstring carries the ordering constraint between the tiers.
"""
import logging
import os
import shutil
import subprocess
import tempfile

log = logging.getLogger(__name__)

# Word constants (we avoid importing the type library just for these)
_WD_ALERTS_NONE = 0
_WD_DO_NOT_SAVE = 0
_WD_FORMAT_XML_DOCUMENT = 12       # .docx


def available():
    """True when this host can drive Word.

    REPORT_DISABLE_WORD=1 makes this say no on a machine that HAS Word, so a
    Windows dev box can produce exactly what the Linux production host will:
    the Python fallback, updateFields left in place, page numbers left to the
    reader. Without it there is no way to see the deployed behaviour short of
    deploying, and the two outputs differ in the one respect nobody notices
    until a reader opens the file.
    """
    if os.environ.get("REPORT_DISABLE_WORD"):
        return False
    if os.name != "nt":
        return False
    try:
        import win32com.client  # noqa: F401
        return True
    except Exception:  # noqa: BLE001
        return False


# Said once per process, not once per report, so it is visible in the boot log
# without burying every generation behind it.
_warned = [False]


def warn_if_unavailable():
    """Say which engine will finish the reports here, and warn if none will.

    Computing page numbers needs a LAYOUT ENGINE - knowing that a heading falls
    on page 31 means laying the document out, and Python has none. Word is one;
    LibreOffice is the other, and it is why the Linux server is not stuck. So the
    only bad case is a host with neither, and it deserves a loud line in the boot
    log rather than a silent False: that made it look like the problem had been
    fixed everywhere when it had been fixed on one developer's laptop.

    On a host with neither, the report keeps w:updateFields, every reader gets the
    "update the fields?" prompt and the four "page numbers only / entire table"
    prompts, and if their Word has track changes on, its rebuild lands as
    revisions - the red contents page and the "Field Code Changed" balloons.

    LIBREOFFICE ON DEBIAN/UBUNTU also needs ``python3-uno``. The soffice binary
    alone is not enough: without the bridge this falls through to the Python
    writer, which looks like it worked until somebody opens the file.
    """
    if _warned[0]:
        return
    _warned[0] = True
    if available():
        return
    if libreoffice_available():
        log.info("no Word on this host (%s) - report fields will be computed in "
                 "LibreOffice instead, which on this report gives the same "
                 "pagination Word does. See report_gen/finalise.py.", os.name)
        return
    log.warning(
        "report fields will NOT be computed on this host (%s): no Word, and no "
        "LibreOffice with a working uno bridge (on Debian/Ubuntu: apt install "
        "python3-uno). Reports keep updateFields, so readers get the field "
        "prompts and, with track changes on, a markup margin. See "
        "report_gen/finalise.py.", os.name)


def compute_fields(path, timeout_hint=120):
    """Open ``path`` in Word, compute every field, save, close.

    Returns True when the document was finalised. On any failure the file is
    left exactly as the builder wrote it - a report with un-computed fields is
    worse to read but still correct, and losing the report entirely is not an
    acceptable trade for tidier page numbers.
    """
    if not available():
        return False

    path = os.path.abspath(path)
    if not os.path.exists(path):
        return False

    import pythoncom
    import win32com.client as win32

    pythoncom.CoInitialize()
    word = doc = None
    try:
        # DispatchEx, not Dispatch: a private instance, so we never attach to -
        # or quit - a Word the user has open with unsaved work in it.
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = _WD_ALERTS_NONE
        # Stop Word recording OUR field rebuild as somebody's edits. This is the
        # setting that produced the red contents page and the balloon margin.
        try:
            word.Options.SaveNormalPrompt = False
        except Exception:  # noqa: BLE001 - not present on every build
            pass

        doc = word.Documents.Open(path, ReadOnly=False, AddToRecentFiles=False,
                                  Visible=False, ConfirmConversions=False)

        doc.TrackRevisions = False
        try:
            if doc.Revisions.Count:
                doc.Revisions.AcceptAll()
        except Exception:  # noqa: BLE001
            pass
        # documentProtection would refuse the edits below; the builder already
        # strips it, this is belt and braces for a hand-edited template.
        try:
            if doc.ProtectionType != -1:      # wdNoProtection
                doc.Unprotect()
        except Exception:  # noqa: BLE001
            pass

        # Order matters. Fields first so captions and SEQ numbers settle,
        # then the tables that INDEX them, then repaginate so PAGE / NUMPAGES
        # and every PAGEREF resolve against the final layout, then fields once
        # more to pick those page numbers up.
        doc.Fields.Update()
        for i in range(1, doc.TablesOfContents.Count + 1):
            doc.TablesOfContents(i).Update()
        for i in range(1, doc.TablesOfFigures.Count + 1):
            doc.TablesOfFigures(i).Update()
        doc.Repaginate()
        doc.Fields.Update()

        pages = int(doc.ComputeStatistics(2))          # wdStatisticPages
        doc.Save()
        doc.Close(SaveChanges=_WD_DO_NOT_SAVE)
        doc = None
        log.info("report finalised in Word: %s (%d pages)",
                 os.path.basename(path), pages)
        return True
    except Exception as exc:  # noqa: BLE001
        log.warning("could not finalise %s in Word (%s) - leaving the "
                    "update-on-open flag in place", os.path.basename(path), exc)
        return False
    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=_WD_DO_NOT_SAVE)
        except Exception:  # noqa: BLE001
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:  # noqa: BLE001
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:  # noqa: BLE001
            pass


def clear_update_on_open(path):
    """Remove ``w:updateFields`` once the fields are already computed.

    Without this Word still asks "do you want to update the fields" on every
    open, and answering yes starts the whole rebuild - and the recording - over
    again.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(path)
        settings = doc.settings.element
        removed = 0
        for node in settings.findall(qn("w:updateFields")):
            settings.remove(node)
            removed += 1
        if removed:
            doc.save(path)
        return removed
    except Exception as exc:  # noqa: BLE001
        log.warning("could not clear updateFields on %s: %s",
                    os.path.basename(path), exc)
        return 0


def unlink_body_fields(path):
    """Freeze the BODY's fields into their computed text. Headers stay live.

    Removing w:updateFields stops Word RECALCULATING on open; it does not stop it
    ASKING. "This document contains fields that may refer to other files" is
    triggered by the presence of TOC fields, which Word classes as potentially
    external because a table of contents can pull entries from other documents
    with an RD switch. Ours does not, but the class is what Word checks.

    So once Word has computed everything there is nothing left for a field to do,
    and unlinking them - the same thing Ctrl+Shift+F9 does by hand - turns the
    result into plain text. No fields, no prompt, and an issued report is the
    right place for a frozen contents page: it should say what was checked and
    approved, not recompute itself on somebody's machine.

    BODY ONLY. The header's PAGE and NUMPAGES fields must stay live or every one
    of the 53 pages reads "Page 1 of 53" - they live in header2.xml, so leaving
    the header parts alone is the whole safeguard.

    Only ever called after compute_fields has succeeded. On a host without Word
    the numbers are not computed, and freezing a blank page number would make it
    blank permanently.
    """
    try:
        from docx import Document
        doc = Document(path)
        removed = 0
        body = doc.element.body
        # w:fldSimple carries its instruction as an attribute and its result as
        # children - unwrap it so the result survives without the field.
        for simple in list(body.iter(_q("fldSimple"))):
            parent = simple.getparent()
            if parent is None:
                continue
            at = list(parent).index(simple)
            for child in reversed(list(simple)):
                parent.insert(at, child)
            parent.remove(simple)
            removed += 1
        # a complex field is a run holding fldChar or instrText; every visible
        # result sits in ordinary runs between them, so dropping those runs
        # leaves exactly the computed text behind
        for run in list(body.iter(_q("r"))):
            if run.find(_q("fldChar")) is None and run.find(_q("instrText")) is None:
                continue
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)
                removed += 1
        doc.save(path)
        log.info("froze %d field run(s) in the body of %s; header pagination "
                 "left live", removed, os.path.basename(path))
        return removed
    except Exception as exc:  # noqa: BLE001
        log.warning("could not unlink fields in %s: %s", os.path.basename(path), exc)
        return 0


# ==========================================================================
# The LibreOffice path: what Word does, on a host that has no Word
# ==========================================================================
# Measured on this project's own 60-page report, and the reason this exists:
#
#   populate_lists alone   updateFields left set, so the reader is prompted; the
#                          entries carry no page numbers; and the template's
#                          stale PAGEREF bookmarks (_Toc226113467, from the
#                          40-page document the template was cut from) are still
#                          there, so a reader who says Yes gets EVERY page number
#                          resolved to 1 and the raw field code
#                          'TOC \\o "1-3" \\h \\z \\u' printed as body text. With
#                          track changes on, all of it lands as revisions - the
#                          "Field Code Changed" balloons down the margin.
#
#   LibreOffice refresh    122 entries, every one paginated, updateFields
#                          cleared, and after unlink_body_fields zero fields
#                          left - structurally identical to the Word output.
#
# AND THE PAGE NUMBERS AGREE. Built both ways from the same request and compared
# entry by entry: 60 pages each, 123 entries each, and all 122 paginated entries
# on the SAME page as Word put them. This is not "close enough on the report we
# tried"; the two engines were also checked against each other on one file and
# laid it out identically, sheet for sheet.
#
# Getting there needed the two-pass update inside the script below - a single
# pass had every entry at +1, and see the comment there for why. Worth knowing
# that the failure looked exactly like an unfixable font-metric divergence
# between two layout engines, and was not: LibreOffice's file put the heading on
# sheet 9 with a header reading "PAGE: 9" while its own contents said 10, and a
# number contradicting its own document is a bug, not a difference of opinion.
#
# --convert-to docx does NOT do any of this. Measured: the contents came back
# byte for byte as it went in, and updateFields was stripped on the way - the
# worst combination, a blank contents page with nothing left to offer to rebuild
# it. Refreshing the indexes over the UNO bridge is what makes the difference.

_UNO_SCRIPT = r'''
import os, subprocess, sys, time
import uno
from com.sun.star.beans import PropertyValue

def _pv(n, v):
    p = PropertyValue(); p.Name = n; p.Value = v; return p

def connect(port, tries=45):
    local = uno.getComponentContext()
    resolver = local.ServiceManager.createInstanceWithContext(
        "com.sun.star.bridge.UnoUrlResolver", local)
    url = "uno:socket,host=127.0.0.1,port=%d;urp;StarOffice.ComponentContext" % port
    started = False
    for i in range(tries):
        try:
            return resolver.resolve(url)
        except Exception:
            if not started:
                started = True
                exe = os.environ.get("SOFFICE_BIN") or "soffice"
                subprocess.Popen([exe, "--headless", "--norestore", "--invisible",
                                  "--nologo", "--nodefault",
                                  "--accept=socket,host=127.0.0.1,port=%d;urp;" % port])
            time.sleep(1.0)
    raise SystemExit("no soffice listener on %d" % port)

src, dst, port = sys.argv[1], sys.argv[2], int(sys.argv[3])
ctx = connect(port)
desktop = ctx.ServiceManager.createInstanceWithContext(
    "com.sun.star.frame.Desktop", ctx)
doc = desktop.loadComponentFromURL(uno.systemPathToFileUrl(src), "_blank", 0,
                                   (_pv("Hidden", True), _pv("UpdateDocMode", 1)))
if doc is None:
    raise SystemExit("load failed")
try:
    # Fields first, then indexes. An index entry CONTAINS PAGEREF fields, so
    # refreshing the indexes first computes the page numbers and the field
    # refresh then wipes them straight back out.
    try: doc.refresh()
    except Exception: pass
    try:
        f = doc.getTextFields()
        if f is not None: f.refresh()
    except Exception: pass
    idx = doc.getDocumentIndexes()
    n = idx.getCount()
    # The COLLECTION has no refresh(); each index has update(). Calling
    # refresh() on the collection raises a bare "refresh" and leaves every page
    # number exactly as it was - which reads like success in a log.
    #
    # TWICE, and this is not superstition - it was measured. Rebuilding the four
    # lists CHANGES THE LENGTH OF THE FRONT MATTER (the template's surplus
    # placeholder lines go away), so page numbers computed during the first pass
    # describe a document a page longer than the one that gets saved. One pass
    # put every entry at +1 against the file's own layout: the heading printed on
    # sheet 9, header "PAGE: 9", contents entry saying 10. The second pass runs
    # against the settled layout and agrees with it. This is the same shape as
    # the Word path above, which updates, repaginates, then updates again.
    for _pass in (1, 2):
        for i in range(n):
            idx.getByIndex(i).update()
        try: doc.refresh()
        except Exception: pass
    doc.storeToURL(uno.systemPathToFileUrl(dst),
                   (_pv("FilterName", "MS Word 2007 XML"), _pv("Overwrite", True)))
    print("INDEXES %d" % n)
finally:
    try: doc.close(True)
    except Exception: pass
'''


def _uno_python():
    """An interpreter that can ``import uno``, or None.

    LibreOffice bundles one beside soffice on Windows and in the .deb/.rpm
    builds; Debian's own libreoffice-writer package does not, and needs
    python3-uno installed for the system python to carry the bridge. Both are
    tried rather than assumed, because the cost of assuming is a report that
    silently ships with no page numbers.
    """
    from .render import _soffice_path
    exe = _soffice_path()
    cands = []
    if exe:
        here = os.path.dirname(exe)
        cands += [os.path.join(here, "python.exe"), os.path.join(here, "python"),
                  os.path.join(here, "python3")]
    cands += ["python3", "python"]
    for cand in cands:
        try:
            r = subprocess.run([cand, "-c", "import uno"], timeout=60,
                               stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if r.returncode == 0:
                return cand
        except (OSError, subprocess.SubprocessError):
            continue
    return None


def libreoffice_available():
    """True when this host can rebuild the report's indexes with LibreOffice."""
    if os.environ.get("REPORT_DISABLE_LIBREOFFICE"):
        return False
    from .render import _soffice_path
    return bool(_soffice_path()) and bool(_uno_python())


def compute_fields_libreoffice(path, port=2002, timeout=420):
    """Rebuild every index in ``path`` with LibreOffice. True when it worked.

    Writes to a temp file and only replaces the original on success, so a failure
    leaves the report exactly as the builder wrote it.
    """
    if os.environ.get("REPORT_DISABLE_LIBREOFFICE"):
        return False
    from .render import _soffice_path
    exe = _soffice_path()
    if not exe:
        return False
    py = _uno_python()
    if not py:
        log.warning("LibreOffice is installed but no interpreter here can "
                    "'import uno' - on Debian/Ubuntu install python3-uno. "
                    "Falling back to the Python list writer.")
        return False
    path = os.path.abspath(path)
    if not os.path.exists(path):
        return False

    work = tempfile.mkdtemp(prefix="lo_finalise_")
    script = os.path.join(work, "refresh_uno.py")
    out = os.path.join(work, "out.docx")
    try:
        with open(script, "w", encoding="utf-8") as fh:
            fh.write(_UNO_SCRIPT)
        # The caller's venv leaks into LibreOffice's OWN interpreter through
        # these three, and it then cannot find its own standard library.
        env = {k: v for k, v in os.environ.items()
               if k not in ("PYTHONHOME", "PYTHONPATH", "PYTHONEXECUTABLE")}
        env["SOFFICE_BIN"] = exe
        r = subprocess.run([py, script, path, out, str(port)], timeout=timeout,
                           stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=env)
        if r.returncode != 0 or not os.path.exists(out):
            log.warning("LibreOffice index refresh failed for %s: %s",
                        os.path.basename(path),
                        (r.stderr or b"").decode("utf-8", "replace")[:400])
            return False
        shutil.copy(out, path)
        log.info("report finalised in LibreOffice: %s (%s)",
                 os.path.basename(path),
                 (r.stdout or b"").decode("utf-8", "replace").strip() or "no count")
        return True
    except (OSError, subprocess.SubprocessError) as exc:
        log.warning("LibreOffice index refresh raised for %s: %s",
                    os.path.basename(path), exc)
        return False
    finally:
        shutil.rmtree(work, ignore_errors=True)


def finalise(path):
    """Finish the report as well as this host can. Returns what was done.

    Three tiers, best first, and the first two both produce a FINISHED document -
    no prompt on open, no page numbers left blank, and no fields left for a
    reader's Word to rebuild and record as revisions:

      word          Windows dev box. Word does the layout, so the page numbers
                    are the ones Word would compute - by definition exact.
      libreoffice   The Linux server. LibreOffice does the layout instead, and
                    on this report produced the same 60 pages with all 122
                    contents entries on the same pages as Word. Measured, not
                    assumed - see the comment above compute_fields_libreoffice.
      python        Neither engine present. Entry text and caption numbers only;
                    updateFields is LEFT in place so a reader who says yes gets
                    page numbers, and one who says no gets a readable list.

    ORDER MATTERS, and not in the obvious way: populate_lists must NOT run before
    the LibreOffice attempt. It writes the entry text in FRONT of the still-live
    TOC field, which is right when nothing will ever compute the index and wrong
    when LibreOffice is about to - measured, LibreOffice then found 1 of the 4
    indexes and rebuilt it alongside the text already written, so every entry
    appeared twice. LibreOffice gets the raw document or none.
    """
    warn_if_unavailable()
    if compute_fields(path):
        clear_update_on_open(path)
        frozen = unlink_body_fields(path)
        return {"engine": "word", "page_numbers": True, "fields_frozen": frozen}
    if compute_fields_libreoffice(path):
        clear_update_on_open(path)
        frozen = unlink_body_fields(path)
        return {"engine": "libreoffice", "page_numbers": True,
                "fields_frozen": frozen}
    info = populate_lists(path)
    info.update({"engine": "python", "page_numbers": False})
    return info


# ==========================================================================
# The Linux path: everything that does not need a layout engine
# ==========================================================================
# compute_fields needs Word, so on the production host it does nothing and the
# report ships with blanked contents lists - clear_toc_entries empties the
# template's stale entries and, without Word, nothing puts real ones back.
#
# Page numbers genuinely cannot be computed here. Knowing that a heading falls on
# page 31 means laying the document out, and Python has no engine for that.
# Everything ELSE about those lists is countable from the document itself: which
# headings exist, at what level, in what order, and how many figures precede a
# given caption.
#
# So this writes the entry TEXT and the caption numbers and leaves only the page
# number blank. Two things follow, and the second is the point:
#
#   * a reader who clicks Yes gets a perfect document, exactly as before;
#   * a reader who clicks No now gets a READABLE contents page instead of a blank
#     one - and because Word then never modifies the document, there is nothing
#     for it to record. The red contents page and the balloon margin cannot
#     happen on that path at all.
_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _q(tag):
    return _W + tag


def _text_of(el):
    return "".join(t.text or "" for t in el.iter(_q("t"))).strip()


def _style_of(el):
    ppr = el.find(_q("pPr"))
    if ppr is None:
        return ""
    st = ppr.find(_q("pStyle"))
    return (st.get(_q("val")) or "") if st is not None else ""


def _collect_entries(doc):
    """What each of the four lists should hold, in document order.

    Heading numbers are computed from level counters rather than read off the
    paragraph: the body numbers its headings through list definitions, so the
    text does not carry "4.2." anywhere.
    """
    out = {"TOC": [], "Figure": [], "Photo": [], "Table": []}
    counters = [0, 0, 0]
    seen = {"Figure": 0, "Photo": 0, "Table": 0}
    for el in doc.element.body.iter(_q("p")):
        style = _style_of(el).lower()
        text = _text_of(el)
        if not text:
            continue
        if style.startswith("heading") and style[-1:].isdigit():
            lvl = int(style[-1])
            if 1 <= lvl <= 3:
                counters[lvl - 1] += 1
                for deeper in range(lvl, 3):
                    counters[deeper] = 0
                number = ".".join(str(counters[i]) for i in range(lvl)) + "."
                out["TOC"].append((number, text))
            continue
        if style == "caption":
            for label in ("Figure", "Photo", "Table"):
                if not text.startswith(label):
                    continue
                seen[label] += 1
                tail = text[len(label):].lstrip()
                tail = tail.split(":", 1)[1].strip() if ":" in tail else tail
                out[label].append(("%s %d:" % (label, seen[label]), tail))
                break
    return out


def number_seq_fields(doc):
    """Set each SEQ field's cached number by counting them in document order.

    Word would compute these. Without Word the cached value is whatever the
    source document carried, so a spliced "Photo 1" stays 1 even when it is the
    ninth photo in the report.
    """
    counts, changed = {}, 0
    for p in doc.element.body.iter(_q("p")):
        runs = list(p)
        for i, r in enumerate(runs):
            instr = "".join(t.text or "" for t in r.iter(_q("instrText")))
            if "SEQ" not in instr:
                continue
            parts = instr.split()
            if "SEQ" not in parts:
                continue
            label = parts[parts.index("SEQ") + 1] if len(parts) > parts.index("SEQ") + 1 else ""
            if not label:
                continue
            counts[label] = counts.get(label, 0) + 1
            for later in runs[i:]:
                fc = later.find(_q("fldChar"))
                if fc is not None and fc.get(_q("fldCharType")) == "end":
                    break
                done = False
                for t in later.iter(_q("t")):
                    if (t.text or "").strip().isdigit():
                        t.text = str(counts[label])
                        changed += 1
                        done = True
                        break
                if done:
                    break
    return changed


def _regions(doc):
    """[(kind, [paragraphs of the field])] for the four TOC-family fields."""
    paras = list(doc.element.body.iter(_q("p")))
    out, depth, cur = [], 0, None
    for i, el in enumerate(paras):
        instr = " ".join("".join(t.text or "" for t in el.iter(_q("instrText"))).split())
        for fc in el.iter(_q("fldChar")):
            typ = fc.get(_q("fldCharType"))
            if typ == "begin":
                depth += 1
                if depth == 1 and instr.startswith("TOC"):
                    kind = "TOC"
                    for label in ("Figure", "Photo", "Table"):
                        if '"%s"' % label in instr:
                            kind = label
                    cur = (kind, i)
            elif typ == "end":
                if depth == 1 and cur is not None:
                    out.append((cur[0], paras[cur[1]:i + 1]))
                    cur = None
                depth = max(0, depth - 1)
    return out


def _write_entry(p, number, text):
    """Rewrite one contents paragraph as number, tab, text, tab.

    The field machinery - the outer TOC chars and the nested PAGEREF - is kept
    exactly as it was, so Word can still rebuild the list and fill the page in.
    Only the visible text is replaced, which is the same surgery
    docx_tools.clear_toc_entries already does to blank it.
    """
    import copy
    from docx.oxml import OxmlElement
    rpr = None
    for r in list(p):
        if r.tag != _q("r"):
            continue
        if r.find(_q("fldChar")) is not None or r.find(_q("instrText")) is not None:
            continue
        if rpr is None:
            found = r.find(_q("rPr"))
            if found is not None:
                rpr = copy.deepcopy(found)
        p.remove(r)
    for link in p.findall(_q("hyperlink")):
        for t in link.iter(_q("t")):
            t.text = ""

    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(rpr)
    t = OxmlElement("w:t")
    t.set(_XML_SPACE, "preserve")
    t.text = "%s\t%s\t" % (number, text)
    run.append(t)

    ppr = p.find(_q("pPr"))
    if ppr is not None:
        ppr.addnext(run)
    else:
        p.insert(0, run)


def populate_lists(path):
    """Write the contents entries and caption numbers without Word.

    Best-effort: any failure leaves the file exactly as the builder wrote it,
    which is the behaviour this replaces.
    """
    try:
        from docx import Document
        doc = Document(path)
        seq = number_seq_fields(doc)
        entries = _collect_entries(doc)
        written = dropped = 0
        for kind, paras in _regions(doc):
            want = entries.get(kind) or []
            # only paragraphs with no field machinery may be rewritten or removed;
            # the first and last of a region carry the field and must survive
            slots = [p for p in paras
                     if p.find(_q("fldChar")) is None
                     and p.find(_q("instrText")) is None]
            for i, p in enumerate(slots):
                if i < len(want):
                    _write_entry(p, want[i][0], want[i][1])
                    written += 1
                else:
                    parent = p.getparent()
                    if parent is not None:
                        parent.remove(p)
                        dropped += 1
        doc.save(path)
        log.info("report lists written without Word: %d entries, %d caption "
                 "numbers, %d surplus lines removed (page numbers left to Word)",
                 written, seq, dropped)
        return {"entries_written": written, "captions_numbered": seq,
                "surplus_removed": dropped}
    except Exception as exc:  # noqa: BLE001
        log.warning("could not write the report lists for %s: %s",
                    os.path.basename(path), exc)
        return {"entries_written": 0, "captions_numbered": 0, "surplus_removed": 0}
