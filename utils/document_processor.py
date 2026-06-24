import os
import re
import json
import logging
from typing import Dict, List, Optional, Any
from datetime import datetime
import docx2txt
from docx import Document

logger = logging.getLogger(__name__)

try:
    from spire.doc import Document as SpireDocument  # type: ignore
    from spire.doc import FileFormat  # type: ignore
    SPIRE_AVAILABLE = True
except ImportError:
    SPIRE_AVAILABLE = False
    logger.warning("Spire.Doc not available - will use fallback methods only")


class DocumentProcessor:
    """Class for processing uploaded Word documents and extracting structured data."""

    def __init__(self, upload_folder: str):
        """Initialize the document processor.

        Args:
            upload_folder: Path to the upload folder
        """
        self.upload_folder = upload_folder

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file using multiple methods.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text as string

        Raises:
            Exception: If text extraction fails
        """
        try:
            # Try docx2txt first (more reliable for simple documents)
            text = docx2txt.process(file_path)
            if text and len(text.strip()) > 0:
                return text.strip()

            # Fallback to python-docx
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text.strip())

            return '\n'.join(text)

        except Exception as e:
            logger.error(
                f"Error extracting text with docx2txt/python-docx: {e}")

            # Final fallback to Spire.Doc (if available)
            if SPIRE_AVAILABLE:
                try:
                    doc = SpireDocument()
                    doc.LoadFromFile(file_path)
                    text = doc.GetText()
                    doc.Close()
                    return text.strip()
                except Exception as spire_error:
                    logger.error(
                        f"Error extracting text with Spire.Doc: {spire_error}")

            raise Exception(f"Failed to extract text from document: {e}")

    def parse_test_request_data(self, text: str) -> Dict[str, Any]:
        """Parse extracted text to identify test request data.

        Args:
            text: Extracted text from the document

        Returns:
            Dictionary containing parsed test request data
        """
        data = {
            'project_name': '',
            'test_objective': '',
            'test_scope': '',
            'equipment_required': [],
            'tests_to_perform': [],
            'safety_requirements': '',
            'test_parameters': {},
            'contact_info': {},
            'submission_date': '',
            'priority': 'normal',
            'estimated_duration': '',
            'special_requirements': ''
        }

        lines = text.split('\n')

        # Extract project information
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()

            # Project name
            if any(keyword in line_lower for keyword in ['project', 'test name', 'title']):
                if ':' in line:
                    data['project_name'] = line.split(':', 1)[1].strip()
                elif i + 1 < len(lines):
                    data['project_name'] = lines[i + 1].strip()

            # Test objective
            elif any(keyword in line_lower for keyword in ['objective', 'purpose', 'goal']):
                if ':' in line:
                    data['test_objective'] = line.split(':', 1)[1].strip()
                else:
                    # Collect multiple lines for objective
                    objective_lines = []
                    j = i + 1
                    while j < len(lines) and lines[j].strip() and not any(keyword in lines[j].lower() for keyword in ['scope', 'equipment', 'test', 'safety']):
                        objective_lines.append(lines[j].strip())
                        j += 1
                    data['test_objective'] = ' '.join(objective_lines)

            # Test scope
            elif 'scope' in line_lower:
                if ':' in line:
                    data['test_scope'] = line.split(':', 1)[1].strip()
                else:
                    # Collect multiple lines for scope
                    scope_lines = []
                    j = i + 1
                    while j < len(lines) and lines[j].strip() and not any(keyword in lines[j].lower() for keyword in ['equipment', 'test', 'safety', 'parameter']):
                        scope_lines.append(lines[j].strip())
                        j += 1
                    data['test_scope'] = ' '.join(scope_lines)

            # Equipment required
            elif any(keyword in line_lower for keyword in ['equipment', 'instrument', 'tool']):
                if ':' in line:
                    equipment_text = line.split(':', 1)[1].strip()
                    if equipment_text:
                        data['equipment_required'].extend(
                            [eq.strip() for eq in equipment_text.split(',')])
                else:
                    # Collect equipment from subsequent lines
                    j = i + 1
                    while j < len(lines) and lines[j].strip() and not any(keyword in lines[j].lower() for keyword in ['test', 'safety', 'parameter', 'procedure']):
                        equipment_line = lines[j].strip()
                        if equipment_line:
                            data['equipment_required'].extend(
                                [eq.strip() for eq in equipment_line.split(',')])
                        j += 1

            # Tests to perform
            elif any(keyword in line_lower for keyword in ['test', 'procedure', 'method']):
                if ':' in line:
                    test_text = line.split(':', 1)[1].strip()
                    if test_text:
                        data['tests_to_perform'].extend(
                            [test.strip() for test in test_text.split(',')])
                else:
                    # Collect tests from subsequent lines
                    j = i + 1
                    while j < len(lines) and lines[j].strip() and not any(keyword in lines[j].lower() for keyword in ['safety', 'parameter', 'contact', 'date']):
                        test_line = lines[j].strip()
                        if test_line:
                            data['tests_to_perform'].extend(
                                [test.strip() for test in test_line.split(',')])
                        j += 1

            # Safety requirements
            elif any(keyword in line_lower for keyword in ['safety', 'hazard', 'precaution']):
                if ':' in line:
                    data['safety_requirements'] = line.split(':', 1)[1].strip()
                else:
                    # Collect safety requirements from subsequent lines
                    safety_lines = []
                    j = i + 1
                    while j < len(lines) and lines[j].strip() and not any(keyword in lines[j].lower() for keyword in ['contact', 'date', 'parameter']):
                        safety_lines.append(lines[j].strip())
                        j += 1
                    data['safety_requirements'] = ' '.join(safety_lines)

            # Test parameters
            elif any(keyword in line_lower for keyword in ['parameter', 'condition', 'setting']):
                if ':' in line:
                    param_text = line.split(':', 1)[1].strip()
                    if '=' in param_text:
                        key, value = param_text.split('=', 1)
                        data['test_parameters'][key.strip()] = value.strip()

            # Contact information
            elif any(keyword in line_lower for keyword in ['contact', 'email', 'phone', 'name']):
                if ':' in line:
                    key = line.split(':', 1)[0].strip().lower()
                    value = line.split(':', 1)[1].strip()
                    if 'email' in key:
                        data['contact_info']['email'] = value
                    elif 'phone' in key:
                        data['contact_info']['phone'] = value
                    elif 'name' in key:
                        data['contact_info']['name'] = value

            # Submission date
            elif any(keyword in line_lower for keyword in ['date', 'submitted', 'requested']):
                # Look for date patterns
                date_pattern = r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}'
                date_match = re.search(date_pattern, line)
                if date_match:
                    data['submission_date'] = date_match.group()

            # Priority
            elif 'priority' in line_lower:
                if 'high' in line_lower:
                    data['priority'] = 'high'
                elif 'low' in line_lower:
                    data['priority'] = 'low'
                else:
                    data['priority'] = 'normal'

            # Estimated duration
            elif any(keyword in line_lower for keyword in ['duration', 'time', 'estimate']):
                # Look for time patterns
                time_pattern = r'\d+\s*(hour|day|week|month)s?'
                time_match = re.search(time_pattern, line_lower)
                if time_match:
                    data['estimated_duration'] = time_match.group()

        # Clean up lists
        data['equipment_required'] = [
            eq for eq in data['equipment_required'] if eq]
        data['tests_to_perform'] = [
            test for test in data['tests_to_perform'] if test]

        return data

    def validate_extracted_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate and clean extracted data.

        Args:
            data: Extracted data dictionary

        Returns:
            Validated and cleaned data dictionary
        """
        validation_result = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'cleaned_data': data.copy()
        }

        # Check required fields
        if not data.get('project_name'):
            validation_result['warnings'].append('Project name not found')

        if not data.get('test_objective'):
            validation_result['warnings'].append('Test objective not found')

        if not data.get('tests_to_perform'):
            validation_result['warnings'].append('No tests specified')

        if not data.get('equipment_required'):
            validation_result['warnings'].append('No equipment specified')

        # Clean equipment list
        cleaned_equipment = []
        for equipment in data.get('equipment_required', []):
            # Remove common prefixes/suffixes
            cleaned = equipment.strip()
            if cleaned:
                # Remove numbering (e.g., "1. Multimeter" -> "Multimeter")
                cleaned = re.sub(r'^\d+\.\s*', '', cleaned)
                cleaned_equipment.append(cleaned)

        validation_result['cleaned_data']['equipment_required'] = cleaned_equipment

        # Clean test list
        cleaned_tests = []
        for test in data.get('tests_to_perform', []):
            cleaned = test.strip()
            if cleaned:
                # Remove numbering
                cleaned = re.sub(r'^\d+\.\s*', '', cleaned)
                cleaned_tests.append(cleaned)

        validation_result['cleaned_data']['tests_to_perform'] = cleaned_tests

        # Set default values for missing fields
        if not validation_result['cleaned_data'].get('priority'):
            validation_result['cleaned_data']['priority'] = 'normal'

        if not validation_result['cleaned_data'].get('submission_date'):
            validation_result['cleaned_data']['submission_date'] = datetime.now().strftime(
                '%Y-%m-%d')

        return validation_result

    def process_document(self, filename: str) -> Dict[str, Any]:
        """Process a document file and extract structured data.

        Args:
            filename: Name of the uploaded file

        Returns:
            Dictionary containing processing results
        """
        file_path = os.path.join(self.upload_folder, filename)

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            # Extract text
            logger.info(f"Extracting text from {filename}")
            text = self.extract_text_from_docx(file_path)

            if not text or len(text.strip()) == 0:
                raise Exception("No text content found in document")

            # Parse data
            logger.info(f"Parsing data from {filename}")
            parsed_data = self.parse_test_request_data(text)

            # Validate data
            logger.info(f"Validating data from {filename}")
            validation_result = self.validate_extracted_data(parsed_data)

            return {
                'success': True,
                'extracted_text': text,
                'parsed_data': validation_result['cleaned_data'],
                'validation': validation_result,
                'file_size': os.path.getsize(file_path)
            }

        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
            }
