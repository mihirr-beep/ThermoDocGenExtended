"""Consolidate a TCO's approved datasheets into one 'Test Report' .docx.

Each EMC test of a TCO produces its own datasheet .docx (PlannerEntry.datasheet_file_path).
This merges them, in order, into a single document with a cover page + a contents list,
each datasheet starting on its own page. Uses docxcompose so images, tables, styles and
numbering from every datasheet are preserved.
"""
import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer


def _prepend_page_break(doc):
    """Insert a hard page break before the document's first paragraph, so each
    appended datasheet starts on a fresh page."""
    if not doc.paragraphs:
        doc.add_paragraph()
    first = doc.paragraphs[0]
    p = first.insert_paragraph_before()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _build_cover(meta, datasheets):
    """Return a python-docx Document holding the Test Report cover page + contents."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Test Report")
    run.bold = True
    run.font.size = Pt(30)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("Consolidated EMC Test Datasheets")
    srun.italic = True
    srun.font.size = Pt(13)

    doc.add_paragraph()

    info_rows = [
        ("TCO ID", meta.get("tco_id")),
        ("Job Number", meta.get("job_number")),
        ("Product", meta.get("product_name")),
        ("Manufacturer", meta.get("manufacturer")),
        ("Model Number", meta.get("model_number")),
        ("Generated On", meta.get("date")),
        ("Datasheets Included", str(len(datasheets))),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in info_rows:
        if value in (None, ""):
            continue
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = str(value)

    doc.add_paragraph()
    contents = doc.add_paragraph()
    contents.add_run("Contents").bold = True
    for i, ds in enumerate(datasheets, start=1):
        doc.add_paragraph(
            f"{i}.  {ds.get('title') or os.path.basename(ds['path'])}"
        )
    return doc


def build_consolidated_test_report(datasheets, output_path, meta=None):
    """Merge the given datasheet .docx files into one Test Report at output_path.

    datasheets : ordered list of {"path": <abs .docx path>, "title": <test name>}
    output_path: destination .docx path
    meta       : dict with tco_id/job_number/product_name/manufacturer/model_number/date
    Returns (output_path, merged_count). Raises if nothing could be merged.
    """
    meta = meta or {}
    usable = [
        ds for ds in (datasheets or [])
        if ds.get("path") and os.path.exists(ds["path"])
        and ds["path"].lower().endswith(".docx")
    ]
    if not usable:
        raise ValueError("No datasheet .docx files available to consolidate.")

    master = _build_cover(meta, usable)
    composer = Composer(master)

    merged = 0
    for ds in usable:
        try:
            sub = Document(ds["path"])
        except Exception:
            continue  # skip a corrupt/unreadable datasheet rather than fail the whole report
        _prepend_page_break(sub)
        composer.append(sub)
        merged += 1

    if merged == 0:
        raise ValueError("None of the datasheet documents could be read for consolidation.")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    composer.save(output_path)
    return output_path, merged
