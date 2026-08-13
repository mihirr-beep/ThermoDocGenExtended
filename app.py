"""
Flask application for EMI/EMC test plan creation and document generation.

This module provides a web interface for creating test plans by uploading test request documents,
processing them to extract test requirements, and generating comprehensive test plans
and datasheets based on the extracted information.

"""

import json
import logging
import os
import re
import base64
from collections import Counter
from copy import deepcopy
from io import BytesIO
from datetime import datetime, date, time, timezone, timedelta
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.utils import formataddr
from email.header import Header

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from flask import Flask, current_app, render_template, request, jsonify, send_file, flash, redirect, url_for, send_from_directory, abort, session
from flask_login import LoginManager, login_required, current_user
from sqlalchemy import inspect, text
from sqlalchemy.orm import joinedload

from mysql_config import config
from models import db, User, TestRequest, TestPlan, TestDatasheet, Equipment, Maintenance, EquipmentHistory, EMCRequest, PlannerEntry
# Forms imported but not used directly in this file (used in auth_routes)
# from forms import LoginForm, RegistrationForm, ForgotPasswordForm, ResetPasswordForm, ChangePasswordForm
from utils.document_generator import DocumentGenerator
from utils.document_processor import DocumentProcessor
from utils.enhanced_document_processor import EnhancedDocumentProcessor
from utils.equipment_manager import EquipmentManager
from utils.emc_request_repository import get_request_by_legacy_or_normalized_id, get_request_by_tco_id
from utils.normalized_emc_request_service import populate_emc_request_from_form
from utils.upload_routes import create_upload_routes

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# IST timezone (UTC+5:30)
IST = timezone(timedelta(hours=5, minutes=30))

# Single source of truth for the visible application version.
APP_VERSION = '1.5'

# Change this when you want to switch the whole app between testing/production
# without setting environment variables in the shell.
DEFAULT_APP_ENV = 'testing'


def get_ist_now():
    """Get current datetime in IST (Indian Standard Time)."""
    return datetime.now(IST)


def _normalize_planner_total_hours(value):
    """Normalize planner hour values to two decimals."""
    try:
        if value in (None, '', []):
            return None
        return round(float(value), 2)
    except (TypeError, ValueError):
        return None


def _natural_sort_parts(value):
    """Build comparable parts for natural sorting of mixed text/number values."""
    normalized = str(value or '').strip().casefold()
    if not normalized:
        return [(1, '')]

    parts = []
    for chunk in re.split(r'(\d+)', normalized):
        if not chunk:
            continue
        if chunk.isdigit():
            parts.append((0, int(chunk)))
        else:
            parts.append((1, chunk))
    return parts or [(1, normalized)]


def _extract_request_job_number(item):
    """Read a request/job number from ORM rows, dicts, or plain objects."""
    if isinstance(item, dict):
        if 'job_number' in item:
            return item.get('job_number') or ''
        return (
            item.get('job_id')
            or item.get('tco_id')
            or ''
        )

    if hasattr(item, 'job_number'):
        return getattr(item, 'job_number', None) or ''

    return (
        getattr(item, 'job_id', None)
        or getattr(item, 'tco_id', None)
        or ''
    )


def _job_number_sort_key(item):
    """Sort requests by job number, keeping blanks at the end."""
    job_number = str(_extract_request_job_number(item) or '').strip()
    is_blank = 1 if not job_number else 0

    if isinstance(item, dict):
        fallback_value = (
            item.get('created_at')
            or item.get('submitted_date')
            or item.get('id')
            or ''
        )
    else:
        fallback_value = (
            getattr(item, 'created_at', None)
            or getattr(item, 'submitted_at', None)
            or getattr(item, 'id', None)
            or ''
        )

    return (
        is_blank,
        _natural_sort_parts(job_number),
        str(fallback_value or ''),
    )


#: Statuses that are finished/closed. Requests in one of these sort to the END of every
#: list so active work stays at the top.
TERMINAL_REQUEST_STATUSES = ('completed', 'cancelled', 'rejected')


def _is_terminal_request_status(status):
    return str(status or '').strip().lower() in TERMINAL_REQUEST_STATUSES


def _terminal_last(items, status_of=lambda item: getattr(item, 'status', None)):
    """Re-order a list so terminal (completed/cancelled/rejected) items come last,
    preserving the existing relative order inside each group (stable partition)."""
    active, terminal = [], []
    for item in items:
        (terminal if _is_terminal_request_status(status_of(item)) else active).append(item)
    return active + terminal


def _planner_filters(request_id, tco_id):
    """Dual filter for the planner entries belonging to one parent request: match on
    test_request_id OR tco_id. Entries created through different paths link by only one
    of the two, and querying a single column silently missed them - which is what let
    parent request status drift out of sync with its planner entries.

    Returns a list of criteria for ``PlannerEntry.query.filter(db.or_(*...))``.
    """
    filters = [PlannerEntry.test_request_id == request_id]
    if tco_id:
        filters.append(PlannerEntry.tco_id == tco_id)
    return filters


REQUEST_IDENTIFIER_SORT_FIELDS = {'tco_id', 'job_number'}


def _normalize_request_sort_args(sort_by, sort_dir):
    """Normalize supported request-list sort parameters."""
    normalized_sort_by = str(sort_by or '').strip().lower()
    normalized_sort_dir = str(sort_dir or 'asc').strip().lower()

    if normalized_sort_by not in REQUEST_IDENTIFIER_SORT_FIELDS:
        normalized_sort_by = ''
    if normalized_sort_dir not in {'asc', 'desc'}:
        normalized_sort_dir = 'asc'

    return normalized_sort_by, normalized_sort_dir


def _extract_request_identifier_value(item, field_name: str) -> str:
    """Extract a comparable TCO/job identifier from ORM rows, dicts, or objects."""
    if field_name == 'job_number':
        if isinstance(item, dict):
            value = item.get('job_number') or item.get('job_id') or ''
        else:
            value = (
                getattr(item, 'job_number', None)
                or getattr(item, 'job_id', None)
                or ''
            )
        return str(value or '').strip()

    if isinstance(item, dict):
        return str(item.get(field_name) or '').strip()

    return str(getattr(item, field_name, '') or '').strip()


def _apply_request_identifier_sort(items, sort_by, sort_dir='asc'):
    """Sort request-like objects by TCO or Job number with blanks kept last."""
    normalized_sort_by, normalized_sort_dir = _normalize_request_sort_args(
        sort_by, sort_dir
    )
    if not normalized_sort_by:
        return list(items)

    non_blank_items = []
    blank_items = []
    for item in items:
        value = _extract_request_identifier_value(item, normalized_sort_by)
        if value:
            non_blank_items.append(item)
        else:
            blank_items.append(item)

    non_blank_items.sort(
        key=lambda item: _natural_sort_parts(
            _extract_request_identifier_value(item, normalized_sort_by)
        ),
        reverse=(normalized_sort_dir == 'desc')
    )
    return non_blank_items + blank_items


def _validate_planner_time_window(start_time_obj, end_time_obj):
    """Validate planner times are provided and ordered logically."""
    if not start_time_obj or not end_time_obj:
        return 'start_time and end_time must be provided in HH:MM format'

    business_start = time(9, 0, 0)
    business_end = time(18, 0, 0)
    if end_time_obj <= start_time_obj:
        return 'end_time must be after start_time'
    if min(end_time_obj, business_end) <= max(start_time_obj, business_start):
        return 'Schedule must overlap working hours between 09:00 and 18:00'
    return None


NON_TEST_EVENT_TYPES = {
    'meeting',
    'training',
    'out-of-office',
    'general',
    'other'
}


def _normalize_planner_conflict_test_key(value):
    """Normalize planner test names so equivalent labels compare consistently."""
    normalized = re.sub(r'[^A-Za-z0-9_]+', '', str(value or '').upper())
    if not normalized:
        return None

    alias_map = {
        'CE': 'CE',
        'CONDUCTEDEMISSION': 'CE',
        'RE': 'RE',
        'RADIATEDEMISSION': 'RE',
        'HARMONIC': 'HARMONIC',
        'HARMONICCURRENTEMISSION': 'HARMONIC',
        'VOLTAGEFLICKER': 'VOLTAGEFLICKER',
        'FLICKER': 'VOLTAGEFLICKER',
        'VOLTAGECHANGES': 'VOLTAGEFLICKER',
        'ESD': 'ESD',
        'ELECTROSTATICDISCHARGE': 'ESD',
        'RS': 'RS_RI',
        'RI': 'RS_RI',
        'RSRI': 'RS_RI',
        'RS_RI': 'RS_RI',
        'RADIATEDSUSCEPTIBILITY': 'RS_RI',
        'RSINTERIM': 'RS_RI_INTERIM',
        'RS_RI_INTERIM': 'RS_RI_INTERIM',
        'RSRIINTERIM': 'RS_RI_INTERIM',
        'EFT': 'EFT',
        'SURGE': 'SURGE',
        'CRF': 'CRF',
        'PFMF': 'PFMF',
        'POWER': 'PFMF',
        'POWERFREQUENCYMAGNETICFIELD': 'PFMF',
        'POWERFREQUENCYMAGNETICFIELDIMMUNITY': 'PFMF',
        'VOLTAGE': 'VOLTAGEDIPS',
        'VOLTAGEDIP': 'VOLTAGEDIPS',
        'VOLTAGEDIPS': 'VOLTAGEDIPS',
        'VOLTAGEDIPSSHORTINTERRUPTIONS': 'VOLTAGEDIPS',
    }
    return alias_map.get(normalized, normalized)


def _extract_selected_test_shortcuts(value) -> set:
    """Normalize selected_tests payload into canonical test shortcuts."""
    def _collect_tokens(raw_value):
        if raw_value is None:
            return []

        if isinstance(raw_value, list):
            tokens = []
            for item in raw_value:
                tokens.extend(_collect_tokens(item))
            return tokens

        if isinstance(raw_value, dict):
            tokens = []
            for key in ('shortcut', 'short_code', 'code', 'test', 'test_name', 'name', 'value'):
                if key in raw_value and raw_value.get(key) not in (None, ''):
                    tokens.extend(_collect_tokens(raw_value.get(key)))

            for key, state in raw_value.items():
                normalized = str(state).strip().lower()
                if state is True or normalized in ('true', 'yes', '1', 'on', 'selected'):
                    tokens.append(str(key))

            return tokens

        if isinstance(raw_value, str):
            text = raw_value.strip()
            if not text:
                return []
            try:
                parsed = json.loads(text)
                return _collect_tokens(parsed)
            except (json.JSONDecodeError, TypeError, ValueError):
                return [token.strip() for token in re.split(r'[,/|]', text) if token.strip()]

        return [str(raw_value)]

    alias_map = {
        'CE': 'CE',
        'CONDUCTEDEMISSION': 'CE',
        'RE': 'RE',
        'RADIATEDEMISSION': 'RE',
        'HARMONIC': 'HARMONIC',
        'HARMONICCURRENTEMISSION': 'HARMONIC',
        'VOLTAGEFLICKER': 'VOLTAGEFLICKER',
        'FLICKER': 'VOLTAGEFLICKER',
        'VOLTAGECHANGES': 'VOLTAGEFLICKER',
        'ESD': 'ESD',
        'ELECTROSTATICDISCHARGE': 'ESD',
        'RS': 'RS_RI',
        'RI': 'RS_RI',
        'RSRI': 'RS_RI',
        'RS_RI': 'RS_RI',
        'RADIATEDSUSCEPTIBILITY': 'RS_RI',
        'RSINTERIM': 'RS_RI_INTERIM',
        'RS_RI_INTERIM': 'RS_RI_INTERIM',
        'RSRIINTERIM': 'RS_RI_INTERIM',
        'EFT': 'EFT',
        'SURGE': 'SURGE',
        'CRF': 'CRF',
        'PFMF': 'PFMF',
        'POWER': 'PFMF',
        'POWERFREQUENCYMAGNETICFIELD': 'PFMF',
        'POWERFREQUENCYMAGNETICFIELDIMMUNITY': 'PFMF',
        'VOLTAGE': 'VOLTAGEDIPS',
        'VOLTAGEDIP': 'VOLTAGEDIPS',
        'VOLTAGEDIPS': 'VOLTAGEDIPS',
        'VOLTAGEDIPSSHORTINTERRUPTIONS': 'VOLTAGEDIPS',
    }

    shortcuts = set()
    for token in _collect_tokens(value):
        normalized = re.sub(r'[^A-Za-z0-9_]+', '', str(token).upper())
        if not normalized:
            continue
        canonical = alias_map.get(normalized)
        if canonical:
            shortcuts.add(canonical)
    return shortcuts


def _normalize_assignment_test_key(value) -> str | None:
    """Return a stable key for matching selected tests with assignment rows."""
    shortcuts = _extract_selected_test_shortcuts(value)
    if shortcuts:
        return sorted(shortcuts)[0]

    normalized = re.sub(r'[^A-Za-z0-9_]+', '', str(value or '').upper())
    return normalized or None


def _planner_entry_uses_test_capacity(event_type):
    """Return True when the planner row should block the same test slot."""
    return (str(event_type or '').strip().lower() not in NON_TEST_EVENT_TYPES)


def _coerce_planner_conflict_date(value):
    """Coerce planner date-like values into a date object."""
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text_value = str(value).strip()
    if not text_value:
        return None
    try:
        return datetime.strptime(text_value[:10], '%Y-%m-%d').date()
    except ValueError:
        return None


def _coerce_planner_conflict_time(value):
    """Coerce planner time-like values into a time object."""
    if value is None:
        return None
    if isinstance(value, time):
        return value
    if isinstance(value, datetime):
        return value.time()
    text_value = str(value).strip()
    if not text_value:
        return None
    for fmt in ('%H:%M:%S', '%H:%M'):
        try:
            return datetime.strptime(text_value, fmt).time()
        except ValueError:
            continue
    return None


def _build_planner_conflict_window(
    start_date_value,
    end_date_value,
    start_time_value=None,
    end_time_value=None,
    all_day=False
):
    """Build comparable datetimes for planner overlap checks."""
    intervals = _build_planner_conflict_intervals(
        start_date_value,
        end_date_value,
        start_time_value,
        end_time_value,
        all_day=all_day
    )
    if not intervals:
        return None, None
    return intervals[0][0], intervals[-1][1]


def _iter_planner_workdays(start_date_obj, end_date_obj):
    """Yield working dates between two dates, excluding Saturday and Sunday."""
    start_date_obj = _coerce_planner_conflict_date(start_date_obj)
    end_date_obj = _coerce_planner_conflict_date(end_date_obj)
    if not start_date_obj or not end_date_obj or end_date_obj < start_date_obj:
        return

    current_date = start_date_obj
    while current_date <= end_date_obj:
        if current_date.weekday() < 5:
            yield current_date
        current_date += timedelta(days=1)


def _build_planner_conflict_intervals(
    start_date_value,
    end_date_value,
    start_time_value=None,
    end_time_value=None,
    all_day=False
):
    """Split a planner schedule into weekday-only intervals for conflict checks."""
    start_date_obj = _coerce_planner_conflict_date(start_date_value)
    end_date_obj = _coerce_planner_conflict_date(end_date_value)
    if not start_date_obj or not end_date_obj:
        return []

    start_time_obj = _coerce_planner_conflict_time(start_time_value)
    end_time_obj = _coerce_planner_conflict_time(end_time_value)
    if all_day or not start_time_obj or not end_time_obj:
        interval_start_time = time(0, 0, 0)
        interval_end_time = time(23, 59, 59)
    else:
        if end_time_obj <= start_time_obj:
            return []
        interval_start_time = start_time_obj
        interval_end_time = end_time_obj

    intervals = []
    for workday in _iter_planner_workdays(start_date_obj, end_date_obj):
        interval_start = datetime.combine(workday, interval_start_time)
        interval_end = datetime.combine(workday, interval_end_time)
        if interval_start < interval_end:
            intervals.append((interval_start, interval_end))
    return intervals


def _calculate_planner_schedule_hours(
    start_date_value,
    end_date_value,
    start_time_value,
    end_time_value
):
    """Calculate schedule hours across weekdays within 09:00-18:00 only."""
    start_time_obj = _coerce_planner_conflict_time(start_time_value)
    end_time_obj = _coerce_planner_conflict_time(end_time_value)
    if not start_time_obj or not end_time_obj or end_time_obj <= start_time_obj:
        return 0

    business_start = time(9, 0, 0)
    business_end = time(18, 0, 0)
    effective_start = max(start_time_obj, business_start)
    effective_end = min(end_time_obj, business_end)
    if effective_end <= effective_start:
        return 0

    start_dt = datetime.combine(date.today(), effective_start)
    end_dt = datetime.combine(date.today(), effective_end)
    daily_hours = (end_dt - start_dt).total_seconds() / 3600
    if daily_hours <= 0:
        return 0

    workday_count = sum(1 for _ in _iter_planner_workdays(start_date_value, end_date_value))
    if workday_count <= 0:
        return 0

    return round(daily_hours * workday_count, 2)


def _build_planner_conflict_snapshot(record):
    """Normalize DB rows and in-memory planner records for conflict checks."""
    def _safe_int(value):
        if value in (None, '', []):
            return None
        try:
            return int(value)
        except (TypeError, ValueError):
            return None

    if isinstance(record, dict):
        record_id = record.get('id')
        test_request_id = record.get('test_request_id')
        engineer_user_id = record.get('engineer_user_id', record.get('engineer_id'))
        engineer_name = record.get('test_person_name', record.get('engineer_name'))
        test_name = record.get('test_name')
        tco_id = record.get('tco_id')
        start_date_value = record.get('start_date')
        end_date_value = record.get('end_date')
        start_time_value = record.get('start_time')
        end_time_value = record.get('end_time')
        is_all_day = bool(record.get('is_all_day'))
        status = record.get('status')
        event_type = record.get('event_type')
    else:
        record_id = getattr(record, 'id', None)
        test_request_id = getattr(record, 'test_request_id', None)
        engineer_user_id = getattr(record, 'engineer_user_id', None)
        engineer_name = getattr(record, 'test_person_name', None)
        test_name = getattr(record, 'test_name', None)
        tco_id = getattr(record, 'tco_id', None)
        start_date_value = getattr(record, 'start_date', None)
        end_date_value = getattr(record, 'end_date', None)
        start_time_value = getattr(record, 'start_time', None)
        end_time_value = getattr(record, 'end_time', None)
        is_all_day = bool(getattr(record, 'is_all_day', False))
        status = getattr(record, 'status', None)
        event_type = getattr(record, 'event_type', None)

    schedule_intervals = _build_planner_conflict_intervals(
        start_date_value,
        end_date_value,
        start_time_value,
        end_time_value,
        all_day=is_all_day
    )
    start_dt, end_dt = _build_planner_conflict_window(
        start_date_value,
        end_date_value,
        start_time_value,
        end_time_value,
        all_day=is_all_day
    )

    return {
        'id': record_id,
        'test_request_id': test_request_id,
        'engineer_user_id': _safe_int(engineer_user_id),
        'engineer_name': str(engineer_name or '').strip() or None,
        'test_name': str(test_name or '').strip() or None,
        'test_key': _normalize_planner_conflict_test_key(test_name),
        'tco_id': str(tco_id or '').strip() or None,
        'start_dt': start_dt,
        'end_dt': end_dt,
        'schedule_intervals': schedule_intervals,
        'status': str(status or '').strip().lower() or None,
        'event_type': str(event_type or '').strip().lower() or None,
        'uses_test_capacity': _planner_entry_uses_test_capacity(event_type),
    }


def _find_schedule_conflicts_in_snapshots(
    existing_snapshots,
    candidate_snapshot,
    *,
    exclude_entry_id=None,
    ignore_request_id=None,
    ignore_tco_id=None
):
    """Compare a candidate planner window against normalized snapshots."""
    conflicts = {
        'engineer_conflicts': [],
        'test_conflicts': []
    }

    candidate_intervals = candidate_snapshot.get('schedule_intervals') or []
    candidate_start = candidate_snapshot.get('start_dt')
    candidate_end = candidate_snapshot.get('end_dt')
    if not candidate_intervals or not candidate_start or not candidate_end or candidate_start >= candidate_end:
        return conflicts

    for snapshot in existing_snapshots:
        if not snapshot:
            continue
        if snapshot.get('status') == 'cancelled':
            continue
        if exclude_entry_id is not None and snapshot.get('id') == exclude_entry_id:
            continue
        if ignore_request_id is not None and snapshot.get('test_request_id') == ignore_request_id:
            continue
        if ignore_tco_id and snapshot.get('tco_id') == ignore_tco_id:
            continue

        existing_start = snapshot.get('start_dt')
        existing_end = snapshot.get('end_dt')
        existing_intervals = snapshot.get('schedule_intervals') or []
        if (
            not existing_start or
            not existing_end or
            existing_start >= existing_end or
            not existing_intervals
        ):
            continue
        if not (candidate_start < existing_end and candidate_end > existing_start):
            continue

        has_interval_overlap = any(
            candidate_interval_start < existing_interval_end and
            candidate_interval_end > existing_interval_start
            for candidate_interval_start, candidate_interval_end in candidate_intervals
            for existing_interval_start, existing_interval_end in existing_intervals
        )
        if not has_interval_overlap:
            continue

        if (
            candidate_snapshot.get('engineer_user_id') is not None
            and snapshot.get('engineer_user_id') == candidate_snapshot.get('engineer_user_id')
        ):
            conflicts['engineer_conflicts'].append(snapshot)

        if (
            candidate_snapshot.get('test_key')
            and candidate_snapshot.get('uses_test_capacity')
            and snapshot.get('uses_test_capacity')
            and snapshot.get('test_key') == candidate_snapshot.get('test_key')
        ):
            conflicts['test_conflicts'].append(snapshot)

    return conflicts


def _find_db_schedule_conflicts(
    candidate_snapshot,
    *,
    exclude_entry_id=None,
    ignore_request_id=None,
    ignore_tco_id=None
):
    """Query planner rows that overlap a candidate assignment/event."""
    start_dt = candidate_snapshot.get('start_dt')
    end_dt = candidate_snapshot.get('end_dt')
    if not start_dt or not end_dt:
        return {'engineer_conflicts': [], 'test_conflicts': []}

    overlapping_entries = PlannerEntry.query.filter(
        PlannerEntry.status != 'cancelled',
        PlannerEntry.start_date <= end_dt.date(),
        PlannerEntry.end_date >= start_dt.date()
    ).order_by(
        PlannerEntry.start_date.asc(),
        PlannerEntry.start_time.asc()
    ).all()

    snapshots = [
        _build_planner_conflict_snapshot(entry)
        for entry in overlapping_entries
    ]
    return _find_schedule_conflicts_in_snapshots(
        snapshots,
        candidate_snapshot,
        exclude_entry_id=exclude_entry_id,
        ignore_request_id=ignore_request_id,
        ignore_tco_id=ignore_tco_id
    )


def _format_planner_conflict_window(snapshot):
    """Render a compact planner conflict time window."""
    start_dt = snapshot.get('start_dt')
    end_dt = snapshot.get('end_dt')
    if not start_dt or not end_dt:
        return 'the selected schedule'
    return f"{start_dt.strftime('%Y-%m-%d %H:%M')} to {end_dt.strftime('%Y-%m-%d %H:%M')}"


def _has_schedule_conflicts(conflicts):
    """Return True when either engineer or test conflicts were detected."""
    return bool(conflicts.get('engineer_conflicts') or conflicts.get('test_conflicts'))


def _format_schedule_conflict_message(test_name, engineer_name, conflicts):
    """Build a user-facing message describing the blocking schedule conflict."""
    messages = []

    engineer_conflicts = conflicts.get('engineer_conflicts') or []
    if engineer_conflicts:
        first_conflict = engineer_conflicts[0]
        conflicting_test = first_conflict.get('test_name') or 'another schedule'
        conflicting_engineer = first_conflict.get('engineer_name') or engineer_name or 'Selected engineer'
        messages.append(
            f'Engineer "{conflicting_engineer}" is not available during '
            f'{_format_planner_conflict_window(first_conflict)} because of "{conflicting_test}".'
        )

    test_conflicts = conflicts.get('test_conflicts') or []
    if test_conflicts:
        first_conflict = test_conflicts[0]
        assigned_to = first_conflict.get('engineer_name')
        tco_display = first_conflict.get('tco_id')
        detail_parts = []
        if assigned_to:
            detail_parts.append(f'assigned to {assigned_to}')
        if tco_display:
            detail_parts.append(f'TCO {tco_display}')
        detail_text = f" ({', '.join(detail_parts)})" if detail_parts else ''
        messages.append(
            f'Test "{test_name}" is not available during '
            f'{_format_planner_conflict_window(first_conflict)}{detail_text}.'
        )

    return ' '.join(messages) or 'Schedule conflict detected.'


def _resolve_request(request_id):
    """Load a migrated EMC request by normalized or legacy-facing id."""
    try:
        normalized_request_id = int(request_id)
    except (TypeError, ValueError):
        return None
    return get_request_by_legacy_or_normalized_id(normalized_request_id)


def _get_request_or_404(request_id):
    """Resolve a migrated EMC request or raise HTTP 404."""
    request_obj = _resolve_request(request_id)
    if request_obj is None:
        abort(404)
    return request_obj


def _extract_service_types(request_obj):
    """Normalize service_types for legacy UI code during the schema cutover."""
    raw_value = getattr(request_obj, 'service_types', None)
    if not raw_value:
        return []
    if isinstance(raw_value, (list, tuple)):
        values = []
        for item in raw_value:
            if isinstance(item, str):
                text = item.strip()
            else:
                text = str(getattr(item, 'service_type', '') or '').strip()
            if text:
                values.append(text)
        return values
    if isinstance(raw_value, str):
        try:
            parsed = json.loads(raw_value)
        except (TypeError, ValueError):
            parsed = raw_value
        if isinstance(parsed, list):
            return [str(item).strip() for item in parsed if str(item).strip()]
        if isinstance(parsed, str) and parsed.strip():
            return [parsed.strip()]
    return []


def _request_skips_report_review(request_obj) -> bool:
    """Return True when the request should bypass the draft-report review flow."""
    return any(
        'development' in service_type.casefold() and 'assist' in service_type.casefold()
        for service_type in _extract_service_types(request_obj)
    )


def _calculate_calibration_status(calibration_due_date, calibration_required):
    """Calculate calibration status based on calibration due date.

    Args:
        calibration_due_date: Date object or None
        calibration_required: String 'Yes' or 'No' or None

    Returns:
        String: 'In Calibration' if due date is today or future, 
                'Out of Calibration' if due date is past,
                None if calibration not required or no due date
    """
    if calibration_required != 'Yes' or not calibration_due_date:
        return None

    today = get_ist_now().date()

    if calibration_due_date < today:
        return 'Out of Calibration'
    else:
        return 'In Calibration'


def parse_date_field(date_str):
    """Helper function to parse date string to date object."""
    if not date_str:
        return None
    try:
        return datetime.strptime(date_str, '%Y-%m-%d').date()
    except ValueError:
        return None


def log_equipment_history(equipment_id, action_type, old_values=None, new_values=None, notes=None):
    """Log equipment changes to history."""
    try:
        # Calculate changes
        changes = {}
        if old_values and new_values:
            for key in set(list(old_values.keys()) + list(new_values.keys())):
                old_val = old_values.get(key)
                new_val = new_values.get(key)
                if old_val != new_val:
                    changes[key] = {'old': old_val, 'new': new_val}
        elif new_values:
            # For created action, store all new values
            changes = new_values

        history = EquipmentHistory(
            equipment_id=equipment_id,
            action_type=action_type,
            changed_by_user_id=current_user.id if current_user.is_authenticated else None,
            notes=notes
        )
        if changes:
            history.set_changes(changes)
        db.session.add(history)
    except Exception as e:
        logger.error("Error logging equipment history: %s", e)


def _is_draft_status(status_value) -> bool:
    """Return True when a request status should be treated as private draft."""
    return (status_value or '').strip().lower() == 'draft'


def _is_request_owner(test_request: EMCRequest) -> bool:
    """Return True when the current user owns the given request."""
    return bool(
        test_request
        and test_request.user_id
        and test_request.user_id == current_user.id
    )


def _can_access_iec_request(
    test_request: EMCRequest,
    *,
    allow_lab_engineer: bool = False,
    require_assigned_lab_engineer: bool = False
) -> bool:
    """Return True when current user can access this IEC request.

    Draft requests are always owner-only.
    """
    if not test_request or not current_user.is_authenticated:
        return False

    if _is_request_owner(test_request):
        return True

    if _is_draft_status(getattr(test_request, 'status', None)):
        return False

    if current_user.role == 'admin':
        return True

    if current_user.role == 'lab_engineer':
        if not allow_lab_engineer:
            return False
        if not require_assigned_lab_engineer:
            return True
        return bool(
            getattr(test_request, 'assigned_engineer_id', None) == current_user.id
        )

    return False


def _can_access_review_thread(test_request: EMCRequest) -> bool:
    """Return True when the current user can view or post request comments."""
    return _can_access_iec_request(
        test_request,
        allow_lab_engineer=True,
        require_assigned_lab_engineer=True
    )


def _parse_review_comment_thread(test_request: EMCRequest) -> list:
    """Normalize review_comments storage into a list of comment dicts."""
    raw = test_request.review_comments
    if not raw:
        return []

    comments = []
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            comments = parsed
        elif isinstance(parsed, dict):
            comments = [parsed]
        else:
            comments = [str(parsed)]
    except (json.JSONDecodeError, TypeError, ValueError):
        comments = [raw]

    normalized = []
    for item in comments:
        if isinstance(item, dict):
            text_value = (item.get('comment') or '').strip()
            if not text_value:
                continue
            normalized_item = {
                'comment': text_value,
                'username': item.get('username') or test_request.reviewed_by or 'Unknown',
                'created_at': item.get('created_at') or (
                    test_request.reviewed_at.strftime('%d %b %Y, %I:%M %p')
                    if test_request.reviewed_at else ''
                )
            }
            if item.get('role'):
                normalized_item['role'] = item.get('role')
            normalized.append(normalized_item)
            continue

        text_value = str(item).strip()
        if text_value:
            normalized.append({
                'comment': text_value,
                'username': test_request.reviewed_by or 'Unknown',
                'created_at': test_request.reviewed_at.strftime('%d %b %Y, %I:%M %p')
                if test_request.reviewed_at else ''
            })

    return normalized


def _append_review_comment_entry(
    test_request: EMCRequest,
    comment: str,
    username: str,
    role=None
) -> dict:
    """Append a comment to review thread and update review metadata."""
    now = get_ist_now()
    thread = _parse_review_comment_thread(test_request)
    entry = {
        'comment': comment,
        'username': username or 'Unknown',
        'created_at': now.strftime('%d %b %Y, %I:%M %p')
    }
    if role:
        entry['role'] = role

    thread.append(entry)
    test_request.review_comments = json.dumps(thread)
    test_request.reviewed_by = username
    test_request.reviewed_at = now
    return entry


def _get_active_selected_test_labels(test_request: EMCRequest) -> list[str]:
    """Return selected test labels excluding tests already cancelled at request level."""
    selected_tests = []
    for request_test in getattr(test_request, 'tests', []) or []:
        if not getattr(request_test, 'is_selected', False):
            continue
        if str(getattr(request_test, 'workflow_status', '') or '').strip().lower() == 'cancelled':
            continue
        label = EMCRequest._legacy_test_key(getattr(request_test, 'test_code', None))
        if label and label not in selected_tests:
            selected_tests.append(label)
    return selected_tests


def _reconcile_removed_tests_after_admin_edit(
    test_request: EMCRequest,
    previous_selected_codes: set[str],
    edited_by_user: User
) -> dict:
    """Cancel linked planner rows when an admin removes tests from a TCO."""
    current_selected_codes = {
        str(getattr(test_row, 'test_code', '') or '').strip().upper()
        for test_row in (getattr(test_request, 'tests', []) or [])
        if getattr(test_row, 'is_selected', False)
    }
    removed_codes = {
        str(code or '').strip().upper()
        for code in previous_selected_codes
        if str(code or '').strip().upper()
    } - current_selected_codes

    if not removed_codes:
        return {
            'removed_codes': set(),
            'removed_labels': [],
            'cancelled_planner_entries': 0
        }

    removed_labels = []
    removed_match_keys = set()
    for code in sorted(removed_codes):
        label = EMCRequest._legacy_test_key(code) or code
        if label not in removed_labels:
            removed_labels.append(label)
        removed_match_keys.add(_normalize_assignment_test_key(code))
        removed_match_keys.add(_normalize_assignment_test_key(label))

    now = get_ist_now()
    cancel_reason = (
        f"Removed from TCO during admin edit by {edited_by_user.username}."
    )
    planner_filters = [PlannerEntry.test_request_id == test_request.id]
    if test_request.tco_id:
        planner_filters.append(PlannerEntry.tco_id == test_request.tco_id)

    cancelled_planner_entries = 0
    for planner_entry in PlannerEntry.query.filter(db.or_(*planner_filters)).all():
        planner_key = _normalize_assignment_test_key(
            getattr(planner_entry, 'test_name', None)
        )
        if planner_key not in removed_match_keys:
            continue
        planner_status = str(getattr(planner_entry, 'status', '') or '').strip().lower()
        if planner_status in {'cancelled', 'completed'}:
            continue

        planner_entry.status = 'cancelled'
        planner_entry.cancel_reason = cancel_reason
        planner_entry.cancelled_at = now
        planner_entry.cancelled_by = edited_by_user.id
        planner_entry.updated_at = now
        cancelled_planner_entries += 1

    _append_review_comment_entry(
        test_request=test_request,
        comment=f"[Admin Edit] Removed tests: {', '.join(removed_labels)}",
        username=edited_by_user.username,
        role=edited_by_user.role
    )

    return {
        'removed_codes': removed_codes,
        'removed_labels': removed_labels,
        'cancelled_planner_entries': cancelled_planner_entries
    }


def _cancel_tco_request(
    test_request: EMCRequest,
    cancel_reason: str,
    cancelled_by_user: User
) -> dict:
    """Cancel a parent TCO and cascade the same reason to linked tests."""
    reason = str(cancel_reason or '').strip()
    if not reason:
        raise ValueError('Cancellation reason is required')

    now = get_ist_now()
    test_request.status = 'Cancelled'
    test_request.updated_at = now
    setattr(test_request, 'rejection_reason', reason)
    setattr(test_request, 'rejected_by', cancelled_by_user.username)
    setattr(test_request, 'rejected_at', now)
    _append_review_comment_entry(
        test_request=test_request,
        comment=f'[TCO Cancelled] {reason}',
        username=cancelled_by_user.username,
        role=cancelled_by_user.role
    )

    cancelled_request_tests = 0
    for request_test in getattr(test_request, 'tests', []) or []:
        if str(getattr(request_test, 'workflow_status', '') or '').strip().lower() != 'cancelled':
            cancelled_request_tests += 1
        request_test.workflow_status = 'cancelled'

    planner_filters = [PlannerEntry.test_request_id == test_request.id]
    if test_request.tco_id:
        planner_filters.append(PlannerEntry.tco_id == test_request.tco_id)

    planner_entries = PlannerEntry.query.filter(
        db.or_(*planner_filters)
    ).all()

    cancelled_planner_entries = 0
    seen_entry_ids = set()
    for planner_entry in planner_entries:
        if planner_entry.id in seen_entry_ids:
            continue
        seen_entry_ids.add(planner_entry.id)

        if str(planner_entry.status or '').strip().lower() != 'cancelled':
            cancelled_planner_entries += 1
        planner_entry.status = 'cancelled'
        planner_entry.cancel_reason = reason
        planner_entry.cancelled_at = now
        planner_entry.cancelled_by = cancelled_by_user.id
        planner_entry.updated_at = now

    return {
        'cancelled_request_tests': cancelled_request_tests,
        'cancelled_planner_entries': cancelled_planner_entries
    }


def _get_service_type_filter_options(service_type_lists=None):
    """Extract unique service type options from request data for filtering UI.

    Canonicalizes service type aliases so filters do not show duplicates such as:
    - Developmental Assistance
    - Developmental Assitance (DA)
    """
    ordered_service_types = [
        'Developmental Assistance (DA)',
        'Pre-Compliance',
        'Compliance',
    ]

    def _canonicalize_service_type(service_type):
        text = str(service_type or '').strip()
        if not text:
            return ''

        normalized = ' '.join(text.casefold().split())
        if (
            'developmental' in normalized and
            ('assistance' in normalized or 'assitance' in normalized)
        ):
            return 'Developmental Assistance (DA)'
        if normalized == 'pre-compliance':
            return 'Pre-Compliance'
        if normalized == 'compliance':
            return 'Compliance'
        return ''

    if not service_type_lists:
        return ordered_service_types

    unique_types = set()
    for type_list in service_type_lists:
        if isinstance(type_list, (list, tuple)):
            for service_type in type_list:
                canonical = _canonicalize_service_type(service_type)
                if canonical:
                    unique_types.add(canonical)
        elif type_list:
            canonical = _canonicalize_service_type(type_list)
            if canonical:
                unique_types.add(canonical)

    return [service_type for service_type in ordered_service_types if service_type in unique_types] or ordered_service_types


def _matches_service_type_filter(plan_service_types, selected_filter):
    """Check if a plan's service types match the selected filter."""
    if not selected_filter:
        return True

    if not plan_service_types:
        return False

    def _canonicalize_for_match(service_type):
        text = str(service_type or '').strip()
        if not text:
            return ''

        normalized = ' '.join(text.casefold().split())
        if (
            'developmental' in normalized and
            ('assistance' in normalized or 'assitance' in normalized)
        ):
            return 'developmental assistance (da)'
        if normalized == 'pre-compliance':
            return 'pre-compliance'
        if normalized == 'compliance':
            return 'compliance'
        return normalized

    normalized_filter = _canonicalize_for_match(selected_filter)
    for service_type in plan_service_types:
        if _canonicalize_for_match(service_type) == normalized_filter:
            return True

    return False


def _report_access_requires_feedback(test_request):
    """Determine if report access requires feedback submission."""
    if not test_request:
        return False
    
    # Check if request has 'Developmental' or 'Assist' service type
    service_types = _extract_service_types(test_request)
    return any(
        'development' in st.lower() or 'assist' in st.lower()
        for st in service_types
    )


def _report_is_approvable_status(status) -> bool:
    """Return whether a report status represents a draft awaiting approval."""
    normalized_status = ' '.join(
        str(status or '').strip().casefold().replace('_', ' ').split()
    )
    return normalized_status in {'draft report', 'report uploaded'}


def _has_report_access_grant(request_id, planner_entry):
    """Check if user has been granted access to view/download report."""
    if not planner_entry or not request_id:
        return False
    
    # Check if a report access grant exists for this entry
    access_granted = getattr(planner_entry, 'report_access_granted', False)
    return bool(access_granted)


def _grant_report_access(request_id, planner_entry):
    """Grant report access for a planner entry after feedback is submitted."""
    if planner_entry:
        planner_entry.report_access_granted = True
        planner_entry.report_access_granted_at = get_ist_now()
        planner_entry.updated_at = get_ist_now()


def _get_assignment_status_filter_options(test_plans):
    """Unique assignment statuses from the test plans, as {value, label} dicts for the
    status-filter dropdown. `value` is the raw status (the client-side filter matches it
    case-insensitively against each row's data-status); `label` is display-friendly.

    NOTE: the template renders `status_option.value` / `status_option.label`, so this
    MUST return dicts — returning bare strings makes every <option> render blank.
    """
    if not test_plans:
        return []

    unique_statuses = set()
    for plan in test_plans:
        if plan and isinstance(plan, dict):
            status = str(plan.get('status', '') or '').strip()
            if status:
                unique_statuses.add(status)

    def _label(s):
        # 'in_progress' -> 'In Progress'; 'Peer Review' -> 'Peer Review'
        return s.replace('_', ' ').strip().title() if s else s

    return [{'value': s, 'label': _label(s)}
            for s in sorted(unique_statuses, key=str.lower)]


def _parse_comment_timestamp(value) -> datetime:
    """Best-effort parser for serialized comment timestamps."""
    raw = (value or '').strip() if isinstance(value, str) else (str(value).strip() if value else '')
    if not raw:
        return datetime.min

    known_formats = [
        '%d %b %Y, %I:%M %p',
        '%Y-%m-%d %H:%M:%S',
        '%Y-%m-%dT%H:%M:%S',
        '%Y-%m-%dT%H:%M:%S.%f',
    ]
    for fmt in known_formats:
        try:
            return datetime.strptime(raw, fmt)
        except ValueError:
            continue

    try:
        # Handle ISO 8601 timestamps with timezone suffixes.
        parsed = datetime.fromisoformat(raw.replace('Z', '+00:00'))
        if parsed.tzinfo is not None:
            return parsed.astimezone(timezone.utc).replace(tzinfo=None)
        return parsed
    except ValueError:
        return datetime.min


def _get_combined_review_comment_thread(test_request: EMCRequest) -> list:
    """Return merged review comments for the same TCO (fallback: request-only thread)."""
    if not test_request:
        return []

    tco_id = (test_request.tco_id or '').strip()
    related_requests = [test_request]

    if tco_id:
        query = EMCRequest.query.filter_by(tco_id=tco_id)
        if current_user.is_authenticated and current_user.role not in ['admin', 'lab_engineer']:
            query = query.filter_by(user_id=test_request.user_id)

        related_requests = query.order_by(
            EMCRequest.updated_at.asc(),
            EMCRequest.created_at.asc(),
            EMCRequest.id.asc()
        ).all() or [test_request]

    merged = []
    seen = set()
    for req in related_requests:
        for item in _parse_review_comment_thread(req):
            comment_text = (item.get('comment') or '').strip()
            username = (item.get('username') or 'Unknown').strip()
            created_at = (item.get('created_at') or '').strip()
            role = (item.get('role') or '').strip()

            if not comment_text:
                continue

            key = (comment_text, username, created_at, role)
            if key in seen:
                continue
            seen.add(key)

            normalized = {
                'comment': comment_text,
                'username': username,
                'created_at': created_at
            }
            if role:
                normalized['role'] = role
            merged.append(normalized)

    merged.sort(key=lambda entry: _parse_comment_timestamp(entry.get('created_at')))
    return merged


def ensure_equipment_document_link_column():
    """Ensure document_link column exists in equipment table."""
    try:
        inspector = inspect(db.engine)
    except Exception as exc:
        logger.warning(
            'Skipping equipment document_link column check: %s', exc)
        return

    try:
        existing_tables = inspector.get_table_names()
    except Exception as exc:
        logger.warning('Unable to list tables for equipment check: %s', exc)
        return

    if 'equipment' not in existing_tables:
        logger.info(
            'Equipment table not found yet; document_link column check skipped.')
        return

    existing_columns = {
        column['name']: column for column in inspector.get_columns('equipment')
    }

    if 'document_link' not in existing_columns:
        try:
            logger.info(
                'Executing: ALTER TABLE equipment ADD COLUMN document_link VARCHAR(500) NULL')
            db.session.execute(
                text('ALTER TABLE equipment ADD COLUMN document_link VARCHAR(500) NULL')
            )
            db.session.commit()
            logger.info(
                'Successfully added document_link column to equipment table')
        except Exception as clause_exc:
            db.session.rollback()
            error_msg = str(clause_exc).lower()
            if 'duplicate' in error_msg or 'already exists' in error_msg or 'duplicate column' in error_msg:
                logger.info(
                    'Column document_link already exists; skipping addition')
            else:
                logger.warning(
                    'Error adding document_link column: %s', clause_exc)
    logger.info('Completed equipment document_link column check')


def ensure_equipment_test_name_column():
    """Ensure test_name column exists in equipment table."""
    try:
        inspector = inspect(db.engine)
    except Exception as exc:
        logger.warning(
            'Skipping equipment test_name column check: %s', exc)
        return

    try:
        existing_tables = inspector.get_table_names()
    except Exception as exc:
        logger.warning('Unable to list tables for equipment check: %s', exc)
        return

    if 'equipment' not in existing_tables:
        logger.info(
            'Equipment table not found yet; test_name column check skipped.')
        return

    existing_columns = {
        column['name']: column for column in inspector.get_columns('equipment')
    }

    if 'test_name' not in existing_columns:
        try:
            logger.info(
                'Executing: ALTER TABLE equipment ADD COLUMN test_name VARCHAR(500) NULL')
            db.session.execute(
                text('ALTER TABLE equipment ADD COLUMN test_name VARCHAR(500) NULL COMMENT "Test Name (comma-separated values or single value)"')
            )
            db.session.commit()
            logger.info(
                'Successfully added test_name column to equipment table')
        except Exception as clause_exc:
            db.session.rollback()
            error_msg = str(clause_exc).lower()
            if 'duplicate' in error_msg or 'already exists' in error_msg or 'duplicate column' in error_msg:
                logger.info(
                    'Column test_name already exists; skipping addition')
            else:
                logger.warning(
                    'Error adding test_name column: %s', clause_exc)
    logger.info('Completed equipment test_name column check')


def ensure_planner_table():
    """Ensure planner_entries table exists with required columns."""
    try:
        inspector = inspect(db.engine)
    except Exception as exc:
        logger.warning('Skipping planner table check: %s', exc)
        return

    try:
        existing_tables = inspector.get_table_names()
    except Exception as exc:
        logger.warning('Unable to list tables for planner check: %s', exc)
        return

    if 'planner_entries' not in existing_tables:
        try:
            db.session.execute(text("""
                CREATE TABLE IF NOT EXISTS planner_entries (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    test_person_name VARCHAR(200) NOT NULL,
                    engineer_user_id INT NULL,
                    created_by_user_id INT NULL,
                    test_name VARCHAR(200) NOT NULL,
                    tco_id VARCHAR(50) NOT NULL,
                    test_request_id INT NULL,
                    start_date DATE NOT NULL,
                    end_date DATE NOT NULL,
                    start_time TIME NULL,
                    end_time TIME NULL,
                    total_hours FLOAT NULL,
                    event_description TEXT NULL,
                    event_type VARCHAR(50) NULL,
                    recurrence VARCHAR(20) NULL,
                    recurrence_end_date DATE NULL,
                    is_all_day TINYINT(1) NOT NULL DEFAULT 0,
                    datasheet_file_path VARCHAR(500) NULL,
                    datasheet_uploaded_at DATETIME NULL,
                    datasheet_uploaded_by INT NULL,
                    peer_reviewer_user_id INT NULL,
                    peer_review_assigned_at DATETIME NULL,
                    datasheet_comments TEXT NULL,
                    completion_date DATE NULL,
                    cancel_reason TEXT NULL,
                    cancelled_at DATETIME NULL,
                    cancelled_by INT NULL,
                    report_file_path VARCHAR(500) NULL,
                    report_comments TEXT NULL,
                    report_uploaded_at DATETIME NULL,
                    report_uploaded_by INT NULL,
                    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    INDEX idx_planner_test_person (test_person_name),
                    INDEX idx_planner_tco (tco_id),
                    INDEX idx_planner_start (start_date),
                    INDEX idx_planner_end (end_date),
                    INDEX idx_planner_engineer (engineer_user_id),
                    INDEX idx_planner_peer_reviewer (peer_reviewer_user_id),
                    INDEX idx_planner_request (test_request_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """))
            db.session.commit()
            logger.info('Created planner_entries table')
        except Exception as exc:
            db.session.rollback()
            logger.error('Failed to create planner_entries table: %s', exc)
        return

    required_columns = {
        'test_person_name': 'ADD COLUMN test_person_name VARCHAR(200) NOT NULL DEFAULT \'\' AFTER id',
        'engineer_user_id': 'ADD COLUMN engineer_user_id INT NULL AFTER test_person_name',
        'created_by_user_id': 'ADD COLUMN created_by_user_id INT NULL AFTER engineer_user_id',
        'test_name': 'ADD COLUMN test_name VARCHAR(200) NOT NULL DEFAULT \'\' AFTER test_person_name',
        'tco_id': 'ADD COLUMN tco_id VARCHAR(50) NOT NULL DEFAULT \'\' AFTER test_name',
        'test_request_id': 'ADD COLUMN test_request_id INT NULL AFTER tco_id',
        'start_date': 'ADD COLUMN start_date DATE NOT NULL AFTER test_request_id',
        'end_date': 'ADD COLUMN end_date DATE NOT NULL AFTER start_date',
        'start_time': 'ADD COLUMN start_time TIME NULL AFTER end_date',
        'end_time': 'ADD COLUMN end_time TIME NULL AFTER start_time',
        'total_hours': 'ADD COLUMN total_hours FLOAT NULL AFTER end_time',
        'event_description': 'ADD COLUMN event_description TEXT NULL AFTER total_hours',
        'event_type': 'ADD COLUMN event_type VARCHAR(50) NULL AFTER event_description',
        'recurrence': 'ADD COLUMN recurrence VARCHAR(20) NULL AFTER event_type',
        'recurrence_end_date': 'ADD COLUMN recurrence_end_date DATE NULL AFTER recurrence',
        'is_all_day': 'ADD COLUMN is_all_day TINYINT(1) NOT NULL DEFAULT 0 AFTER recurrence_end_date',
        'datasheet_file_path': 'ADD COLUMN datasheet_file_path VARCHAR(500) NULL AFTER is_all_day',
        'datasheet_uploaded_at': 'ADD COLUMN datasheet_uploaded_at DATETIME NULL AFTER datasheet_file_path',
        'datasheet_uploaded_by': 'ADD COLUMN datasheet_uploaded_by INT NULL AFTER datasheet_uploaded_at',
        'peer_reviewer_user_id': 'ADD COLUMN peer_reviewer_user_id INT NULL AFTER datasheet_uploaded_by',
        'peer_review_assigned_at': 'ADD COLUMN peer_review_assigned_at DATETIME NULL AFTER peer_reviewer_user_id',
        'datasheet_comments': 'ADD COLUMN datasheet_comments TEXT NULL AFTER peer_review_assigned_at',
        'completion_date': 'ADD COLUMN completion_date DATE NULL AFTER datasheet_comments',
        'cancel_reason': 'ADD COLUMN cancel_reason TEXT NULL AFTER completion_date',
        'cancelled_at': 'ADD COLUMN cancelled_at DATETIME NULL AFTER cancel_reason',
        'cancelled_by': 'ADD COLUMN cancelled_by INT NULL AFTER cancelled_at',
        'report_file_path': 'ADD COLUMN report_file_path VARCHAR(500) NULL AFTER cancelled_by',
        'report_comments': 'ADD COLUMN report_comments TEXT NULL AFTER report_file_path',
        'report_uploaded_at': 'ADD COLUMN report_uploaded_at DATETIME NULL AFTER report_comments',
        'report_uploaded_by': 'ADD COLUMN report_uploaded_by INT NULL AFTER report_uploaded_at',
        'created_at': 'ADD COLUMN created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP',
        'updated_at': 'ADD COLUMN updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP'
    }

    existing_columns = {
        column['name']: column for column in inspector.get_columns('planner_entries')
    }

    for column_name, clause in required_columns.items():
        if column_name in existing_columns:
            continue
        try:
            db.session.execute(
                text(f'ALTER TABLE planner_entries {clause}')
            )
            db.session.commit()
            logger.info(
                'Added missing column %s to planner_entries', column_name)
        except Exception as exc:
            db.session.rollback()
            error_msg = str(exc).lower()
            if 'duplicate' in error_msg or 'already exists' in error_msg:
                logger.info(
                    'Planner column %s already exists (race condition)', column_name)
                continue
            logger.error(
                'Failed to add planner column %s: %s', column_name, exc)

    # Ensure index on test_request_id for quick lookups
    try:
        planner_indexes = {idx['name']
                           for idx in inspector.get_indexes('planner_entries')}
        if 'idx_planner_request' not in planner_indexes:
            db.session.execute(
                text(
                    'CREATE INDEX idx_planner_request ON planner_entries (test_request_id)')
            )
            db.session.commit()
            logger.info('Added idx_planner_request index to planner_entries')
        if 'idx_planner_engineer' not in planner_indexes:
            db.session.execute(
                text(
                    'CREATE INDEX idx_planner_engineer ON planner_entries (engineer_user_id)')
            )
            db.session.commit()
            logger.info('Added idx_planner_engineer index to planner_entries')
        if 'idx_planner_peer_reviewer' not in planner_indexes:
            db.session.execute(
                text(
                    'CREATE INDEX idx_planner_peer_reviewer ON planner_entries (peer_reviewer_user_id)')
            )
            db.session.commit()
            logger.info('Added idx_planner_peer_reviewer index to planner_entries')
    except Exception as exc:
        db.session.rollback()
        error_msg = str(exc).lower()
        if 'duplicate' in error_msg or 'already exists' in error_msg:
            logger.info('Planner request index already exists')
        else:
            logger.warning('Unable to ensure planner request index: %s', exc)


def create_app(config_name='default'):
    """Application factory function."""
    flask_app = Flask(__name__)
    resolved_config_name = (
        os.environ.get('APP_ENV')
        or os.environ.get('FLASK_ENV')
        or config_name
        or DEFAULT_APP_ENV
    )
    resolved_config_name = str(resolved_config_name).strip().lower() or 'default'
    if resolved_config_name not in config:
        resolved_config_name = 'default'

    # SMTP configuration (Thermo Fisher relay)
    flask_app.config.setdefault('SMTP_SERVER', 'SMTPRELAY1.THERMOFISHER.COM')
    flask_app.config.setdefault('SMTP_PORT', 25)
    flask_app.config.setdefault(
        'SMTP_FROM_EMAIL', 'noreply-pcbacoe@thermofisher.com')
    flask_app.config.setdefault('SMTP_FROM_NAME', 'noreply-pcbacoe')

    def _send_smtp_notification(recipients, subject: str, body: str) -> None:
        """Utility to send notification email via SMTP relay."""
        if not recipients:
            return

        try:
            msg = MIMEText(body, 'plain', 'utf-8')
            msg['Subject'] = subject
            msg['From'] = formataddr(
                (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
            msg['To'] = ', '.join(recipients)

            # Send using SMTP relay (no auth on internal relay)
            with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                server.sendmail(
                    flask_app.config['SMTP_FROM_EMAIL'], recipients, msg.as_string())
            logger.info(
                'Notification email sent to lab engineers: %s', recipients)
        except Exception as e:
            logger.error('Failed to send SMTP notification: %s', e)

    # Note: These functions are kept for potential future use
    # They may be used by other modules or for different notification scenarios
    def send_lab_engineer_notification(subject: str, body: str) -> None:  # noqa: F841
        """Send notification email to all lab engineers via SMTP relay."""
        recipients = [
            u.email for u in User.query.filter_by(
                role='lab_engineer', is_active=True).all() if u.email
        ]
        if not recipients:
            logger.info(
                'No lab engineer recipients found; skipping email notification.')
            return
        _send_smtp_notification(recipients, subject, body)

    def send_admin_submission_notification(subject: str, body: str) -> None:  # noqa: F841
        """Send notification email to all admins when a request is submitted."""
        recipients = [
            u.email for u in User.query.filter_by(
                role='admin', is_active=True).all() if u.email
        ]
        if not recipients:
            logger.info(
                'No admin recipients found; skipping admin notification.')
            return
        _send_smtp_notification(recipients, subject, body)

    def send_submission_notification(test_request: EMCRequest, requester_name: str, requester_email: str = '') -> None:
        """Send professional HTML email notification when a test request is submitted."""
        try:
            # Get recipients
            lab_engineer_recipients = [
                u.email for u in User.query.filter_by(
                    role='lab_engineer', is_active=True).all() if u.email
            ]
            admin_recipients = [
                u.email for u in User.query.filter_by(
                    role='admin', is_active=True).all() if u.email
            ]

            if not lab_engineer_recipients and not admin_recipients:
                logger.info(
                    'No recipients found; skipping submission notification.')
                return

            # Format selected tests
            selected_tests = []
            if test_request.selected_tests:
                try:
                    if isinstance(test_request.selected_tests, str):
                        selected_tests = json.loads(
                            test_request.selected_tests)
                    else:
                        selected_tests = test_request.selected_tests
                except Exception:
                    selected_tests = []

            tests_display = ', '.join(
                selected_tests) if selected_tests else 'No tests selected'

            # Brief product description (truncate if too long)
            product_desc = test_request.product_description or 'No description provided.'
            if len(product_desc) > 300:
                product_desc = product_desc[:300] + '...'

            # Create HTML email body (matching assignment email format)
            html_body = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        color: #333333;
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background: linear-gradient(135deg, #10B981 0%, #059669 100%);
                        color: white;
                        padding: 30px;
                        text-align: center;
                        border-radius: 8px 8px 0 0;
                    }}
                    .content {{
                        background: #ffffff;
                        padding: 30px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                    .info-box {{
                        background: #f9fafb;
                        border-left: 4px solid #10B981;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .detail-row {{
                        display: flex;
                        padding: 10px 0;
                        border-bottom: 1px solid #e5e7eb;
                    }}
                    .detail-label {{
                        font-weight: bold;
                        width: 150px;
                        color: #6b7280;
                    }}
                    .detail-value {{
                        flex: 1;
                        color: #111827;
                    }}
                    .footer {{
                        background: #f9fafb;
                        padding: 20px;
                        text-align: center;
                        color: #6b7280;
                        font-size: 12px;
                        border-radius: 0 0 8px 8px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1 style="margin: 0; font-size: 24px;">New Test Request Submitted</h1>
                    <p style="margin: 10px 0 0 0; opacity: 0.9;">A new EMI/EMC test request requires review</p>
                </div>
                
                <div class="content">
                    <p>Dear Team,</p>
                    
                    <p>A new EMI/EMC test request has been submitted and is awaiting review:</p>
                    
                    <div class="info-box">
                        <div class="detail-row">
                            <div class="detail-label">TCO ID:</div>
                            <div class="detail-value"><strong>{test_request.tco_id or f'REQ-{test_request.id}'}</strong></div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Product Name:</div>
                            <div class="detail-value">{test_request.product_name or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Manufacturer:</div>
                            <div class="detail-value">{test_request.manufacturer or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Model Number:</div>
                            <div class="detail-value">{test_request.model_number or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Requester:</div>
                            <div class="detail-value">{requester_name or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Selected Tests:</div>
                            <div class="detail-value">{tests_display}</div>
                        </div>
                    </div>

                    <h3 style="color: #111827; margin-top: 30px;">Product Description</h3>
                    <p style="background: #f9fafb; padding: 15px; border-radius: 6px; border-left: 4px solid #10B981;">
                        {product_desc}
                    </p>

                    <p style="margin-top: 30px;">
                        Please log in to the system to review the complete test request details and proceed with approval.
                    </p>
                </div>
                
                <div class="footer">
                    <p style="margin: 0;">This is an automated notification from the Test Request Management System.</p>
                    <p style="margin: 5px 0 0 0;">Please do not reply to this email.</p>
                </div>
            </body>
            </html>
            """

            # Create plain text version
            text_body = f"""
New Test Request Submitted

Dear Team,

A new EMI/EMC test request has been submitted and is awaiting review:

TCO ID: {test_request.tco_id or f'REQ-{test_request.id}'}
Product Name: {test_request.product_name or 'N/A'}
Manufacturer: {test_request.manufacturer or 'N/A'}
Model Number: {test_request.model_number or 'N/A'}
Requester: {requester_name or 'N/A'}
Selected Tests: {tests_display}

Product Description:
{product_desc}

Please log in to the system to review the complete test request details and proceed with approval.

---
This is an automated notification from the Test Request Management System.
Please do not reply to this email.
            """

            # Send to lab engineers
            if lab_engineer_recipients:
                msg_lab = MIMEMultipart('alternative')
                msg_lab[
                    'Subject'] = f'New Test Request Submitted: {test_request.product_name or "Unnamed Product"} (TCO: {test_request.tco_id or f"REQ-{test_request.id}"})'
                msg_lab['From'] = formataddr(
                    (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
                msg_lab['To'] = ', '.join(lab_engineer_recipients)

                part1_lab = MIMEText(text_body, 'plain', 'utf-8')
                part2_lab = MIMEText(html_body, 'html', 'utf-8')
                msg_lab.attach(part1_lab)
                msg_lab.attach(part2_lab)

                with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                    server.sendmail(
                        flask_app.config['SMTP_FROM_EMAIL'], lab_engineer_recipients, msg_lab.as_string())
                logger.info(
                    'Submission notification email sent to lab engineers: %s', lab_engineer_recipients)

            # Send to admins with additional info
            if admin_recipients:
                admin_html_body = html_body.replace(
                    '</div>',
                    f'''
                    <div class="info-box" style="background: #fef3c7; border-left-color: #f59e0b; margin-top: 20px;">
                        <h3 style="margin-top: 0; color: #92400e;">Submission Details</h3>
                        <div class="detail-row">
                            <div class="detail-label">Submitted By:</div>
                            <div class="detail-value">{current_user.username if current_user.is_authenticated else 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">User Email:</div>
                            <div class="detail-value">{current_user.email if current_user.is_authenticated and current_user.email else 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Department:</div>
                            <div class="detail-value">{test_request.requester_department or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Submission Time:</div>
                            <div class="detail-value">{get_ist_now().strftime('%Y-%m-%d %H:%M:%S')} IST</div>
                        </div>
                    </div>
                    </div>
                    ''',
                    1
                )

                admin_text_body = text_body + f'''

Submission Details:
Submitted By: {current_user.username if current_user.is_authenticated else 'N/A'}
User Email: {current_user.email if current_user.is_authenticated and current_user.email else 'N/A'}
Department: {test_request.requester_department or 'N/A'}
Submission Time: {get_ist_now().strftime('%Y-%m-%d %H:%M:%S')} IST
                '''

                msg_admin = MIMEMultipart('alternative')
                msg_admin[
                    'Subject'] = f'New Test Request Submitted: {test_request.product_name or "Unnamed Product"} (TCO: {test_request.tco_id or f"REQ-{test_request.id}"})'
                msg_admin['From'] = formataddr(
                    (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
                msg_admin['To'] = ', '.join(admin_recipients)

                part1_admin = MIMEText(admin_text_body, 'plain', 'utf-8')
                part2_admin = MIMEText(admin_html_body, 'html', 'utf-8')
                msg_admin.attach(part1_admin)
                msg_admin.attach(part2_admin)

                with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                    server.sendmail(
                        flask_app.config['SMTP_FROM_EMAIL'], admin_recipients, msg_admin.as_string())
                logger.info(
                    'Submission notification email sent to admins: %s', admin_recipients)

        except Exception as e:
            logger.error('Failed to send submission notification email: %s', e)

    def send_plan_update_notification(test_request: EMCRequest, comments: str, requested_by: str) -> None:
        """Notify the assigned engineer that the admin requested updates to the test schedule."""
        if not test_request.assigned_engineer_id:
            logger.warning(
                'Plan update requested but no assigned engineer is set for request %s', test_request.id)
            return

        assigned_engineer = db.session.get(
            User, test_request.assigned_engineer_id)
        if not assigned_engineer or not assigned_engineer.email:
            logger.warning(
                'Plan update notification skipped because assigned engineer email is missing for request %s', test_request.id)
            return

        admin_recipients = [
            u.email for u in User.query.filter_by(role='admin', is_active=True).all()
            if u.email
        ]

        subject = f"Update Plan Requested: {test_request.tco_id or f'REQ-{test_request.id}'}"
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = formataddr(
            (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
        msg['To'] = assigned_engineer.email
        if admin_recipients:
            msg['Cc'] = ', '.join(admin_recipients)

        product_name = test_request.product_name or 'Unnamed Product'
        tco_display = test_request.tco_id or f'REQ-{test_request.id}'
        plain_body = f"""
An admin has requested changes to the planned tests for {product_name} ({tco_display}).

Requested By: {requested_by}
Comments:
{comments}

Please revisit the "Assign Tests to Engineers" form, adjust the schedule, and resubmit the plan for approval.

This is an automated notification from the Test Request Management System.
"""

        html_body = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #111827;
            max-width: 640px;
            margin: 0 auto;
            padding: 24px;
        }}
        .header {{
            background: linear-gradient(135deg, #7c3aed 0%, #2563eb 100%);
            color: #fff;
            padding: 24px;
            border-radius: 12px 12px 0 0;
            text-align: center;
        }}
        .header h1 {{
            margin: 0;
            font-size: 22px;
        }}
        .content {{
            border: 1px solid #e5e7eb;
            border-top: none;
            padding: 24px;
            background: #fff;
        }}
        .comment-box {{
            background: #eef2ff;
            border-left: 4px solid #6366f1;
            padding: 16px;
            border-radius: 8px;
            margin: 18px 0;
        }}
        .footer {{
            text-align: center;
            font-size: 12px;
            color: #6b7280;
            margin-top: 20px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Update Plan Requested</h1>
        <p>{product_name} Â· {tco_display}</p>
    </div>
    <div class="content">
        <p>Hello <strong>{assigned_engineer.username}</strong>,</p>
        <p><strong>{requested_by}</strong> asked for updates to the scheduled tests. Please review the comments below, adjust the "Assign Tests to Engineers" details, and resubmit the plan.</p>
        <div class="comment-box">
            <p style="margin:0;white-space:pre-line;">{comments}</p>
        </div>
        <p>You can reopen the request from your dashboard under Assigned Requests.</p>
        <p>Regards,<br/>Test Request Management System</p>
    </div>
    <div class="footer">
        Automated notification Â· Please do not reply
    </div>
</body>
</html>
"""

        msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
        msg.attach(MIMEText(html_body, 'html', 'utf-8'))

        recipients = [assigned_engineer.email] + admin_recipients
        try:
            with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                server.sendmail(
                    flask_app.config['SMTP_FROM_EMAIL'],
                    recipients,
                    msg.as_string()
                )
            logger.info('Plan update notification sent to %s',
                        assigned_engineer.email)
        except Exception as exc:
            logger.error('Failed to send plan update notification: %s', exc)

    def send_completion_notification(test_request: EMCRequest, completed_by: str) -> None:
        """Notify the requester (cc all admins) that their TCO is complete.

        Best-effort: never raises into the caller, so a mail failure cannot roll back or
        break the completion itself - it is only logged.
        """
        requester_email = (test_request.requester_email or '').strip()
        if not requester_email:
            logger.warning(
                'No requester email on request %s; skipping completion notification.',
                test_request.id)
            return

        # cc the admins, but never the requester twice (a requester can also be an admin)
        _seen_emails = {requester_email.lower()}
        admin_recipients = []
        for _admin in User.query.filter_by(role='admin', is_active=True).all():
            _admin_email = (_admin.email or '').strip()
            if _admin_email and _admin_email.lower() not in _seen_emails:
                _seen_emails.add(_admin_email.lower())
                admin_recipients.append(_admin_email)

        try:
            selected_tests = _get_active_selected_test_labels(test_request)
            tests_display = ', '.join(selected_tests) if selected_tests else 'No tests selected'

            tco_display = test_request.tco_id or f'REQ-{test_request.id}'
            completed_on = get_ist_now().strftime('%Y-%m-%d %H:%M:%S')

            rows = [
                ('TCO ID', tco_display),
                ('Job Number', test_request.job_number or 'N/A'),
                ('Product Name', test_request.product_name or 'N/A'),
                ('Manufacturer', test_request.manufacturer or 'N/A'),
                ('Model Number', test_request.model_number or 'N/A'),
                ('Tests', tests_display),
                ('Completed By', completed_by or 'Laboratory Administrator'),
                ('Completed On', f'{completed_on} IST'),
            ]
            detail_rows = ''.join(
                '<tr>'
                f'<td style="padding:8px 12px;color:#6b7280;font-weight:bold;width:170px;">{label}</td>'
                f'<td style="padding:8px 12px;color:#111827;">{value}</td>'
                '</tr>'
                for label, value in rows
            )

            html_body = f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"></head>
<body style="font-family:Arial,sans-serif;line-height:1.6;color:#333;max-width:640px;margin:0 auto;padding:24px;">
    <div style="background:linear-gradient(135deg,#10b981 0%,#059669 100%);color:#fff;padding:24px;border-radius:12px 12px 0 0;text-align:center;">
        <h1 style="margin:0;font-size:22px;">Test Completed Successfully</h1>
        <p style="margin:8px 0 0;opacity:.9;">Your EMI/EMC test request has been completed</p>
    </div>
    <div style="background:#fff;padding:24px;border:1px solid #e5e7eb;border-top:none;">
        <p>Dear {test_request.requester_name or 'User'},</p>
        <p>All tests for your EMI/EMC request have been executed, the report has been
           reviewed and approved, and the request is now closed.</p>
        <table style="width:100%;border-collapse:collapse;background:#f0fdf4;border-left:4px solid #10b981;margin:16px 0;">
            {detail_rows}
        </table>
        <p style="margin-bottom:4px;"><strong>Next steps</strong></p>
        <ul style="margin-top:0;">
            <li>The final report is available in the system.</li>
            <li>You can download and review the complete test documentation.</li>
            <li>Contact the lab if you have questions about the results.</li>
        </ul>
        <p>Thank you for choosing our laboratory for your EMI/EMC testing.</p>
    </div>
    <div style="background:#f9fafb;padding:16px;text-align:center;color:#6b7280;font-size:12px;border:1px solid #e5e7eb;border-top:none;border-radius:0 0 12px 12px;">
        Automated notification from the Test Request Management System. Please do not reply.
    </div>
</body>
</html>
"""
            plain_lines = [
                'Test Completed Successfully', '',
                'All tests for your EMI/EMC request have been executed, the report has been',
                'reviewed and approved, and the request is now closed.', '',
            ]
            plain_lines += [f'{label}: {value}' for label, value in rows]
            plain_lines += [
                '', 'Next steps:',
                '- The final report is available in the system.',
                '- You can download and review the complete test documentation.',
                '- Contact the lab if you have questions about the results.',
                '', 'Automated notification from the Test Request Management System.',
            ]
            plain_body = '\n'.join(plain_lines)

            msg = MIMEMultipart('alternative')
            msg['Subject'] = (
                f'Test Completed: {test_request.product_name or "Unnamed Product"} '
                f'(TCO: {tco_display})'
            )
            msg['From'] = formataddr(
                (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
            msg['To'] = requester_email
            if admin_recipients:
                msg['Cc'] = ', '.join(admin_recipients)

            msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
            msg.attach(MIMEText(html_body, 'html', 'utf-8'))

            recipients = [requester_email] + admin_recipients
            with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                server.sendmail(
                    flask_app.config['SMTP_FROM_EMAIL'], recipients, msg.as_string())

            logger.info(
                'Completion notification sent for request %s to %s and %d admin(s)',
                test_request.id, requester_email, len(admin_recipients))
        except Exception as exc:  # noqa: BLE001 - completion must not fail on mail errors
            logger.error(
                'Failed to send completion notification for request %s: %s',
                test_request.id, exc)

    def send_more_info_notification(test_request: EMCRequest, comments: str, reviewed_by: str) -> None:
        """Send professional HTML email notification to requester when more information is needed."""
        requester_email = test_request.requester_email
        if not requester_email or not requester_email.strip():
            logger.warning(
                'No requester email found; skipping more info notification.')
            return

        try:
            # Get admin and assigned engineer emails for CC
            admin_recipients = [
                u.email for u in User.query.filter_by(
                    role='admin', is_active=True).all() if u.email
            ]

            assigned_engineer_email = None
            if test_request.assigned_engineer_id:
                assigned_engineer = db.session.get(User,
                                                   test_request.assigned_engineer_id)
                if assigned_engineer and assigned_engineer.email:
                    assigned_engineer_email = assigned_engineer.email

            # Format selected tests
            selected_tests = []
            if test_request.selected_tests:
                try:
                    if isinstance(test_request.selected_tests, str):
                        selected_tests = json.loads(
                            test_request.selected_tests)
                    else:
                        selected_tests = test_request.selected_tests
                except:
                    selected_tests = []

            tests_display = ', '.join(
                selected_tests) if selected_tests else 'No tests selected'

            # Brief product description (truncate if too long)
            product_desc = test_request.product_description or 'No description provided.'
            if len(product_desc) > 300:
                product_desc = product_desc[:300] + '...'

            # Create HTML email body (matching submission email format but with yellow/orange theme)
            html_body = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        color: #333333;
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background: linear-gradient(135deg, #F59E0B 0%, #D97706 100%);
                        color: white;
                        padding: 30px;
                        text-align: center;
                        border-radius: 8px 8px 0 0;
                    }}
                    .content {{
                        background: #ffffff;
                        padding: 30px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                    .info-box {{
                        background: #f9fafb;
                        border-left: 4px solid #F59E0B;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .more-info-box {{
                        background: #FFFBEB;
                        border: 2px solid #F59E0B;
                        border-left: 4px solid #D97706;
                        padding: 20px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .detail-row {{
                        display: flex;
                        padding: 10px 0;
                        border-bottom: 1px solid #e5e7eb;
                    }}
                    .detail-label {{
                        font-weight: bold;
                        width: 150px;
                        color: #6b7280;
                    }}
                    .detail-value {{
                        flex: 1;
                        color: #111827;
                    }}
                    .footer {{
                        background: #f9fafb;
                        padding: 20px;
                        text-align: center;
                        color: #6b7280;
                        font-size: 12px;
                        border-radius: 0 0 8px 8px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1 style="margin: 0; font-size: 24px;">Additional Information Required</h1>
                    <p style="margin: 10px 0 0 0; opacity: 0.9;">Your EMI/EMC test request needs more details</p>
                </div>
                
                <div class="content">
                    <p>Dear {test_request.requester_name or 'User'},</p>
                    
                    <p>We have reviewed your EMI/EMC test request and need additional information to proceed. Please review the details below:</p>
                    
                    <div class="info-box">
                        <div class="detail-row">
                            <div class="detail-label">TCO ID:</div>
                            <div class="detail-value"><strong>{test_request.tco_id or f'REQ-{test_request.id}'}</strong></div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Product Name:</div>
                            <div class="detail-value">{test_request.product_name or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Manufacturer:</div>
                            <div class="detail-value">{test_request.manufacturer or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Model Number:</div>
                            <div class="detail-value">{test_request.model_number or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Status:</div>
                            <div class="detail-value"><strong style="color: #D97706;">Need More Information</strong></div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Selected Tests:</div>
                            <div class="detail-value">{tests_display}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Reviewed By:</div>
                            <div class="detail-value">{reviewed_by or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Reviewed On:</div>
                            <div class="detail-value">{get_ist_now().strftime('%Y-%m-%d %H:%M:%S')} IST</div>
                        </div>
                    </div>

                    <div class="more-info-box">
                        <h3 style="color: #D97706; margin-top: 0; font-size: 18px;">Information Required</h3>
                        <p style="color: #92400E; white-space: pre-line; margin: 0;">{comments}</p>
                    </div>

                    <h3 style="color: #111827; margin-top: 30px;">Product Description</h3>
                    <p style="background: #f9fafb; padding: 15px; border-radius: 6px; border-left: 4px solid #F59E0B;">
                        {product_desc}
                    </p>

                    <p style="margin-top: 30px;">
                        Please log in to the system to provide the requested information and resubmit your test request.
                    </p>
                </div>
                
                <div class="footer">
                    <p style="margin: 0;">This is an automated notification from the Test Request Management System.</p>
                    <p style="margin: 5px 0 0 0;">Please do not reply to this email.</p>
                </div>
            </body>
            </html>
            """

            # Create plain text version
            text_body = f"""
Additional Information Required

Dear {test_request.requester_name or 'User'},

We have reviewed your EMI/EMC test request and need additional information to proceed. Please review the details below:

TCO ID: {test_request.tco_id or f'REQ-{test_request.id}'}
Product Name: {test_request.product_name or 'N/A'}
Manufacturer: {test_request.manufacturer or 'N/A'}
Model Number: {test_request.model_number or 'N/A'}
Status: Need More Information
Selected Tests: {tests_display}
Reviewed By: {reviewed_by or 'N/A'}
Reviewed On: {get_ist_now().strftime('%Y-%m-%d %H:%M:%S')} IST

Information Required:
{comments}

Product Description:
{product_desc}

Please log in to the system to provide the requested information and resubmit your test request.

---
This is an automated notification from the Test Request Management System.
Please do not reply to this email.
            """

            # Prepare recipients (To: requester, CC: admin and assigned engineer)
            recipients = [requester_email]
            cc_recipients = []

            if admin_recipients:
                cc_recipients.extend(admin_recipients)

            if assigned_engineer_email and assigned_engineer_email not in cc_recipients:
                cc_recipients.append(assigned_engineer_email)

            # Create email message
            msg = MIMEMultipart('alternative')
            msg['Subject'] = f'Additional Information Required: {test_request.product_name or "Unnamed Product"} (TCO: {test_request.tco_id or f"REQ-{test_request.id}"})'
            msg['From'] = formataddr(
                (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
            msg['To'] = requester_email
            if cc_recipients:
                msg['Cc'] = ', '.join(cc_recipients)

            part1 = MIMEText(text_body, 'plain', 'utf-8')
            part2 = MIMEText(html_body, 'html', 'utf-8')
            msg.attach(part1)
            msg.attach(part2)

            # Send email to all recipients (To + CC)
            all_recipients = recipients + cc_recipients

            with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                server.sendmail(
                    flask_app.config['SMTP_FROM_EMAIL'], all_recipients, msg.as_string())
            logger.info(
                f'More info notification email sent to {requester_email} with CC to {cc_recipients}')

        except Exception as e:
            logger.error('Failed to send more info notification email: %s', e)

    def send_rejection_notification(test_request: EMCRequest, rejection_reason: str, rejected_by: str) -> None:
        """Send professional HTML email notification to requester when test request is rejected."""
        requester_email = test_request.requester_email
        if not requester_email or not requester_email.strip():
            logger.warning(
                'No requester email found; skipping rejection notification.')
            return

        try:
            # Format selected tests
            selected_tests = []
            if test_request.selected_tests:
                try:
                    if isinstance(test_request.selected_tests, str):
                        selected_tests = json.loads(
                            test_request.selected_tests)
                    else:
                        selected_tests = test_request.selected_tests
                except:
                    selected_tests = []

            tests_display = ', '.join(
                selected_tests) if selected_tests else 'No tests selected'

            # Brief product description (truncate if too long)
            product_desc = test_request.product_description or 'No description provided.'
            if len(product_desc) > 300:
                product_desc = product_desc[:300] + '...'

            # Create HTML email body (matching submission email format but with red theme)
            html_body = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        color: #333333;
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background: linear-gradient(135deg, #EF4444 0%, #DC2626 100%);
                        color: white;
                        padding: 30px;
                        text-align: center;
                        border-radius: 8px 8px 0 0;
                    }}
                    .content {{
                        background: #ffffff;
                        padding: 30px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                    .info-box {{
                        background: #f9fafb;
                        border-left: 4px solid #EF4444;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .rejection-box {{
                        background: #FEF2F2;
                        border: 2px solid #EF4444;
                        border-left: 4px solid #DC2626;
                        padding: 20px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .detail-row {{
                        display: flex;
                        padding: 10px 0;
                        border-bottom: 1px solid #e5e7eb;
                    }}
                    .detail-label {{
                        font-weight: bold;
                        width: 150px;
                        color: #6b7280;
                    }}
                    .detail-value {{
                        flex: 1;
                        color: #111827;
                    }}
                    .footer {{
                        background: #f9fafb;
                        padding: 20px;
                        text-align: center;
                        color: #6b7280;
                        font-size: 12px;
                        border-radius: 0 0 8px 8px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1 style="margin: 0; font-size: 24px;">Test Request Rejected</h1>
                    <p style="margin: 10px 0 0 0; opacity: 0.9;">Your EMI/EMC test request has been reviewed</p>
                </div>
                
                <div class="content">
                    <p>Dear {test_request.requester_name or 'User'},</p>
                    
                    <p>We regret to inform you that your EMI/EMC test request has been rejected. Please review the details below:</p>
                    
                    <div class="info-box">
                        <div class="detail-row">
                            <div class="detail-label">TCO ID:</div>
                            <div class="detail-value"><strong>{test_request.tco_id or f'REQ-{test_request.id}'}</strong></div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Product Name:</div>
                            <div class="detail-value">{test_request.product_name or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Manufacturer:</div>
                            <div class="detail-value">{test_request.manufacturer or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Model Number:</div>
                            <div class="detail-value">{test_request.model_number or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Status:</div>
                            <div class="detail-value"><strong style="color: #DC2626;">Rejected</strong></div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Selected Tests:</div>
                            <div class="detail-value">{tests_display}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Rejected By:</div>
                            <div class="detail-value">{rejected_by or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Rejected On:</div>
                            <div class="detail-value">{get_ist_now().strftime('%Y-%m-%d %H:%M:%S')} IST</div>
                        </div>
                    </div>

                    <div class="rejection-box">
                        <h3 style="color: #DC2626; margin-top: 0; font-size: 18px;">Rejection Reason</h3>
                        <p style="color: #991B1B; white-space: pre-line; margin: 0;">{rejection_reason}</p>
                    </div>

                    <h3 style="color: #111827; margin-top: 30px;">Product Description</h3>
                    <p style="background: #f9fafb; padding: 15px; border-radius: 6px; border-left: 4px solid #EF4444;">
                        {product_desc}
                    </p>

                    <p style="margin-top: 30px;">
                        If you have any questions or would like to discuss this decision, please contact the lab administration team.
                    </p>
                </div>
                
                <div class="footer">
                    <p style="margin: 0;">This is an automated notification from the Test Request Management System.</p>
                    <p style="margin: 5px 0 0 0;">Please do not reply to this email.</p>
                </div>
            </body>
            </html>
            """

            # Create plain text version
            text_body = f"""
Test Request Rejected

Dear {test_request.requester_name or 'User'},

We regret to inform you that your EMI/EMC test request has been rejected. Please review the details below:

TCO ID: {test_request.tco_id or f'REQ-{test_request.id}'}
Product Name: {test_request.product_name or 'N/A'}
Manufacturer: {test_request.manufacturer or 'N/A'}
Model Number: {test_request.model_number or 'N/A'}
Status: Rejected
Selected Tests: {tests_display}
Rejected By: {rejected_by or 'N/A'}
Rejected On: {get_ist_now().strftime('%Y-%m-%d %H:%M:%S')} IST

Rejection Reason:
{rejection_reason}

Product Description:
{product_desc}

If you have any questions or would like to discuss this decision, please contact the lab administration team.

---
This is an automated notification from the Test Request Management System.
Please do not reply to this email.
            """

            # Send email to requester
            msg = MIMEMultipart('alternative')
            msg['Subject'] = f'Test Request Rejected: {test_request.product_name or "Unnamed Product"} (TCO: {test_request.tco_id or f"REQ-{test_request.id}"})'
            msg['From'] = formataddr(
                (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
            msg['To'] = requester_email

            part1 = MIMEText(text_body, 'plain', 'utf-8')
            part2 = MIMEText(html_body, 'html', 'utf-8')
            msg.attach(part1)
            msg.attach(part2)

            with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                server.sendmail(
                    flask_app.config['SMTP_FROM_EMAIL'], [requester_email], msg.as_string())
            logger.info(
                'Rejection notification email sent to requester: %s', requester_email)

        except Exception as e:
            logger.error('Failed to send rejection notification email: %s', e)

    def send_assignment_notification(engineer_email: str, admin_email: str, test_request: EMCRequest, engineer_name: str, priority: str, due_date: str, notes: str) -> None:
        """Send professional HTML email notification to assigned engineer with admin in CC."""
        if not engineer_email or not engineer_email.strip():
            logger.warning(
                'No engineer email provided; skipping assignment notification.')
            return

        # Ensure admin_email is a string (can be empty)
        admin_email = admin_email or ''

        try:
            # Create HTML email
            msg = MIMEMultipart('alternative')

            # FIX: Properly encode subject with UTF-8
            subject_text = f'New Assignment - TCO: {test_request.tco_id or f"REQ-{test_request.id}"}'
            msg['Subject'] = Header(subject_text, 'utf-8')

            msg['From'] = formataddr(
                (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
            msg['To'] = engineer_email
            if admin_email:
                msg['Cc'] = admin_email

            # Format due date
            due_date_str = 'Not specified'
            if due_date:
                try:
                    due_date_obj = datetime.strptime(
                        due_date, '%Y-%m-%d').date()
                    due_date_str = due_date_obj.strftime('%B %d, %Y')
                except (ValueError, AttributeError):
                    due_date_str = due_date

            # Priority badge color
            priority_colors = {
                'normal': '#3B82F6',  # Blue
                'high': '#F59E0B',    # Amber
                'urgent': '#EF4444'   # Red
            }
            priority_color = priority_colors.get(priority.lower(), '#3B82F6')
            priority_label = priority.capitalize()

            # Brief product description (truncate if too long)
            product_desc = test_request.product_description or 'No description provided.'
            if len(product_desc) > 300:
                product_desc = product_desc[:300] + '...'

            # Create HTML email body - EMOJIS REPLACED WITH HTML ENTITIES
            html_body = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        color: #333333;
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background: linear-gradient(135deg, #10B981 0%, #059669 100%);
                        color: white;
                        padding: 30px;
                        text-align: center;
                        border-radius: 8px 8px 0 0;
                    }}
                    .content {{
                        background: #ffffff;
                        padding: 30px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                    .info-box {{
                        background: #f9fafb;
                        border-left: 4px solid #10B981;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .detail-row {{
                        display: flex;
                        padding: 10px 0;
                        border-bottom: 1px solid #e5e7eb;
                    }}
                    .detail-label {{
                        font-weight: bold;
                        width: 150px;
                        color: #6b7280;
                    }}
                    .detail-value {{
                        flex: 1;
                        color: #111827;
                    }}
                    .priority-badge {{
                        display: inline-block;
                        padding: 4px 12px;
                        border-radius: 12px;
                        font-size: 12px;
                        font-weight: bold;
                        color: white;
                        background-color: {priority_color};
                    }}
                    .notes-section {{
                        background: #fef3c7;
                        border: 1px solid #fbbf24;
                        border-radius: 6px;
                        padding: 15px;
                        margin: 20px 0;
                    }}
                    .footer {{
                        background: #f9fafb;
                        padding: 20px;
                        text-align: center;
                        color: #6b7280;
                        font-size: 12px;
                        border-radius: 0 0 8px 8px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                    .button {{
                        display: inline-block;
                        padding: 12px 24px;
                        background: #10B981;
                        color: white;
                        text-decoration: none;
                        border-radius: 6px;
                        margin: 20px 0;
                        font-weight: bold;
                    }}
                    .highlight {{
                        background: #FEF3C7;
                        padding: 2px 6px;
                        border-radius: 4px;
                        font-weight: bold;
                    }}
                    .icon {{
                        display: inline-block;
                        font-size: 18px;
                        margin-right: 5px;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1 style="margin: 0; font-size: 24px;">Test Request Assigned</h1>
                    <p style="margin: 10px 0 0 0; opacity: 0.9;">You have been assigned a new test request</p>
                </div>
                
                
                
                <div class="content">
                    <p>Dear <strong>{engineer_name}</strong>,</p>
                    
                    <p>You have been assigned a new test request. Please review the details below:</p>
                    
                    <div class="info-box">
                        <div class="detail-row">
                            <div class="detail-label">TCO ID:</div>
                            <div class="detail-value"><strong>{test_request.tco_id or f'REQ-{test_request.id}'}</strong></div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Product Name:</div>
                            <div class="detail-value">{test_request.product_name or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Manufacturer:</div>
                            <div class="detail-value">{test_request.manufacturer or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Model Number:</div>
                            <div class="detail-value">{test_request.model_number or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Requester:</div>
                            <div class="detail-value">{test_request.requester_name or 'N/A'}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Priority:</div>
                            <div class="detail-value"><span class="priority-badge">{priority_label}</span></div>
                        </div>
                        <div class="detail-row" style="border-bottom: none;">
                            <div class="detail-label">Due Date:</div>
                            <div class="detail-value">{due_date_str}</div>
                        </div>
                    </div>

                    <h3 style="color: #111827; margin-top: 30px;">Product Description</h3>
                    <p style="background: #f9fafb; padding: 15px; border-radius: 6px; border-left: 4px solid #10B981;">
                        {product_desc}
                    </p>

                    {f'''
                    <div class="notes-section">
                        <h3 style="margin-top: 0; color: #92400e;">&#9888; Assignment Notes</h3>
                        <p style="margin-bottom: 0; white-space: pre-wrap;">{notes}</p>
                    </div>
                    ''' if notes else ''}

                    <p style="margin-top: 30px;">
                        Please log in to the system to view the complete test request details and begin working on the assignment.
                    </p>
                    
                    <p style="background: #EFF6FF; border-left: 4px solid #3B82F6; padding: 12px; border-radius: 4px; margin-top: 20px;">
                        <strong>&#128204; Important:</strong> Please reference <strong>TCO ID: {test_request.tco_id or f'REQ-{test_request.id}'}</strong> in all communications regarding this test request.
                    </p>
                </div>
                
                <div class="footer">
                    <p style="margin: 0;">This is an automated notification from the Test Request Management System.</p>
                    <p style="margin: 5px 0 0 0;">Please do not reply to this email.</p>
                </div>
            </body>
            </html>
            """

            # Create plain text version - NO EMOJIS
            text_body = f"""
    ===============================================================
                        TEST REQUEST ASSIGNED
    ===============================================================

    Dear {engineer_name},

    You have been assigned a new test request. Please review the details below:

    ---------------------------------------------------------------
    ASSIGNMENT DETAILS
    ---------------------------------------------------------------

    TCO ID:          {test_request.tco_id or f'REQ-{test_request.id}'}
    Product Name:    {test_request.product_name or 'N/A'}
    Manufacturer:    {test_request.manufacturer or 'N/A'}
    Model Number:    {test_request.model_number or 'N/A'}
    Requester:       {test_request.requester_name or 'N/A'}
    Priority:        {priority_label}
    Due Date:        {due_date_str}

    ---------------------------------------------------------------
    PRODUCT DESCRIPTION
    ---------------------------------------------------------------

    {product_desc}

    {f'''---------------------------------------------------------------
    ASSIGNMENT NOTES
    ---------------------------------------------------------------

    {notes}

    ''' if notes else ''}---------------------------------------------------------------


    Please log in to the system to view the complete test request details 
    and begin working on the assignment.

    ===============================================================
    This is an automated notification from the Test Request 
    Management System. Please do not reply to this email.
    ===============================================================
            """

            # FIX: Attach both versions with explicit UTF-8 encoding
            part1 = MIMEText(text_body, 'plain', 'utf-8')
            part2 = MIMEText(html_body, 'html', 'utf-8')
            msg.attach(part1)
            msg.attach(part2)

            # Prepare recipients (To + CC)
            recipients = [engineer_email]
            if admin_email:
                recipients.append(admin_email)

            # FIX: Send email with UTF-8 encoding
            with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                # Convert message to string with UTF-8 encoding
                msg_string = msg.as_string()
                server.sendmail(
                    flask_app.config['SMTP_FROM_EMAIL'],
                    recipients,
                    msg_string
                )

            logger.info(
                'Assignment notification email sent to engineer %s (CC: %s) for TCO: %s',
                engineer_email, admin_email or 'None', test_request.tco_id or f'REQ-{test_request.id}'
            )
        except Exception as e:
            logger.error('Failed to send assignment notification email: %s', e)
            import traceback
            logger.error('Traceback: %s', traceback.format_exc())

    def _format_assignment_datetime(date_str: str | None, time_str: str | None) -> str:
        if date_str and time_str:
            try:
                dt = datetime.strptime(
                    f'{date_str} {time_str}', '%Y-%m-%d %H:%M')
                return dt.strftime('%B %d, %Y %I:%M %p')
            except ValueError:
                pass
        if date_str:
            try:
                date_obj = datetime.strptime(date_str, '%Y-%m-%d').date()
                return date_obj.strftime('%B %d, %Y')
            except ValueError:
                return date_str
        return 'N/A'

    def send_test_assignment_emails(
        test_request: EMCRequest,
        assignment_payloads: list[dict],
        admin_emails: list[str],
        *,
        status_label: str = 'Test Plan To Approve',
        send_to_admins_first: bool = True
    ) -> None:
        """Notify stakeholders about scheduled tests using the shared status template."""
        if not assignment_payloads:
            return

        admin_cc = ', '.join(admin_emails) if admin_emails else ''

        # Get admin User objects for "Test Plan To Approve" status to include names in email
        admin_users = []
        if status_label == 'Test Plan To Approve' and admin_emails:
            admin_users = [
                User.query.filter_by(
                    email=email, role='admin', is_active=True).first()
                for email in admin_emails
            ]
            # Remove None values
            admin_users = [admin for admin in admin_users if admin]

        # Format admin names for greeting
        admin_names_text = ''
        if status_label == 'Test Plan To Approve' and admin_users:
            admin_names = [
                admin.username for admin in admin_users if admin.username]
            if len(admin_names) == 1:
                admin_names_text = admin_names[0]
            elif len(admin_names) == 2:
                admin_names_text = f"{admin_names[0]} and {admin_names[1]}"
            elif len(admin_names) > 2:
                admin_names_text = f"{', '.join(admin_names[:-1])}, and {admin_names[-1]}"

        # For engineer-facing statuses, consolidate tests by engineer and send one email per engineer.
        if status_label in ('Assigned Lab Engineer', 'Test Plan Approved'):
            try:
                # Group assignments by engineer
                engineer_groups = {}
                for payload in assignment_payloads:
                    engineer = payload.get('engineer')
                    assignment = payload.get('assignment') or {}
                    if not engineer or not engineer.email:
                        continue

                    engineer_email = engineer.email
                    if engineer_email not in engineer_groups:
                        engineer_groups[engineer_email] = {
                            'engineer': engineer,
                            'assignments': []
                        }
                    engineer_groups[engineer_email]['assignments'].append(
                        assignment)

                # Send one consolidated email per engineer
                for engineer_email, group_data in engineer_groups.items():
                    engineer = group_data['engineer']
                    assignments = group_data['assignments']
                    is_approved = status_label == 'Test Plan Approved'

                    # Build table rows for all assignments for this engineer
                    table_rows_html = []
                    table_rows_plain = []

                    for assignment in assignments:
                        test_name = assignment.get('test_name', 'N/A')
                        start_window = _format_assignment_datetime(
                            assignment.get('start_date'), assignment.get('start_time'))
                        end_window = _format_assignment_datetime(
                            assignment.get('end_date'), assignment.get('end_time'))
                        schedule = f"{start_window} â†’ {end_window}"
                        total_hours = assignment.get('total_hours')
                        total_hours_text = f"{total_hours} hrs" if total_hours not in (
                            None, '', []) else 'N/A'
                        product_tco = f"{test_request.product_name or 'Unnamed Product'} / {test_request.tco_id or f'REQ-{test_request.id}'}"

                        # HTML table row
                        table_rows_html.append(f"""
                        <tr>
                            <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{product_tco}</td>
                            <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{test_name}</td>
                            <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{schedule}</td>
                            <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{total_hours_text}</td>
                            <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;"><strong style="color:#059669;">{status_label}</strong></td>
                        </tr>
                        """)

                        # Plain text row
                        table_rows_plain.append(
                            f"{product_tco} | {test_name} | {schedule} | {total_hours_text} | {status_label}")

                    # Create consolidated email for this engineer
                    subject = f"{status_label}: {len(assignments)} Test(s) - {test_request.tco_id or f'REQ-{test_request.id}'}"
                    msg = MIMEMultipart('alternative')
                    msg['Subject'] = subject
                    msg['From'] = formataddr(
                        (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
                    msg['To'] = engineer.email
                    if admin_cc:
                        msg['Cc'] = admin_cc

                    greeting_plain = f"Hello {engineer.username},"
                    greeting_html = f"Hello <strong>{engineer.username}</strong>,"
                    intro_plain = (
                        f"Your test plan has been approved for {len(assignments)} test(s). "
                        "Please review the test schedule details below and begin planning your lab time."
                    ) if is_approved else (
                        f"You have been assigned {len(assignments)} test(s). "
                        "Please review the test schedule details below and begin planning your lab time."
                    )
                    intro_html = (
                        f"Your test plan has been approved for <strong>{len(assignments)} test(s)</strong>. "
                        "Please review the test schedule details below and begin planning your lab time."
                    ) if is_approved else (
                        f"You have been assigned <strong>{len(assignments)} test(s)</strong>. "
                        "Please review the test schedule details below and begin planning your lab time."
                    )

                    plain_body = f"""
{status_label}

{greeting_plain}

{intro_plain}

Test Schedule Summary:
{'=' * 100}
Product / TCO | Test Name | Schedule | Estimated Effort | Status
{'-' * 100}
{chr(10).join(table_rows_plain)}
{'=' * 100}

Please log in to the Test Request Management System to review the full plan and attachments.

Regards,
Test Request Management System

---
This is an automated notification. Please do not reply.
"""

                    html_body = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #111827;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
        }}
        .header {{
            background: linear-gradient(135deg, #059669 0%, #10b981 100%);
            color: #fff;
            padding: 28px;
            border-radius: 12px 12px 0 0;
            text-align: center;
        }}
        .header h1 {{
            margin: 0;
            font-size: 24px;
        }}
        .header p {{
            margin: 8px 0 0 0;
            opacity: 0.85;
            font-size: 14px;
        }}
        .content {{
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-top: none;
            padding: 28px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background: #ffffff;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        th {{
            background: #059669;
            color: #fff;
            padding: 12px;
            text-align: left;
            font-weight: 600;
            border-bottom: 2px solid #047857;
        }}
        td {{
            padding: 12px;
            border-bottom: 1px solid #e5e7eb;
        }}
        tr:hover {{
            background-color: #f9fafb;
        }}
        .footer {{
            background: #f9fafb;
            border: 1px solid #e5e7eb;
            border-top: none;
            border-radius: 0 0 12px 12px;
            padding: 18px;
            text-align: center;
            color: #6b7280;
            font-size: 12px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{status_label}</h1>
        <p>{len(assignments)} test schedule(s) {'approved' if is_approved else 'assigned'}</p>
    </div>
    <div class="content">
        <p>{greeting_html}</p>
        <p>{intro_html}</p>
        <table>
            <thead>
                <tr>
                    <th>Product / TCO</th>
                    <th>Test Name</th>
                    <th>Schedule</th>
                    <th>Estimated Effort</th>
                    <th>Status</th>
                </tr>
            </thead>
            <tbody>
                {''.join(table_rows_html)}
            </tbody>
        </table>
        <p style="margin-top:20px;">Log in to the Test Request Management System for full procedures, attachments, and planner availability.</p>
        <p style="margin-top:20px;">Regards,<br/>Test Request Management System</p>
    </div>
    <div class="footer">
        Automated notification Â· Please do not reply
    </div>
</body>
</html>
"""

                    msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
                    msg.attach(MIMEText(html_body, 'html', 'utf-8'))

                    recipients = [engineer.email]
                    if admin_emails:
                        recipients.extend(admin_emails)
                    recipients = list(set(recipients))  # Remove duplicates

                    with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                        server.sendmail(
                            flask_app.config['SMTP_FROM_EMAIL'],
                            recipients,
                            msg.as_string()
                        )
                    logger.info(
                        'Sent consolidated "%s" email to engineer %s with %d test(s)',
                        status_label, engineer.email, len(assignments))

                return  # Exit early for consolidated emails
            except Exception as exc:
                logger.error(
                    'Failed to send consolidated "%s" emails: %s', status_label, exc)
                # Fall through to individual emails if consolidation fails

        # For "Test Plan To Approve" status, consolidate all tests into one email with a table
        if status_label == 'Test Plan To Approve' and admin_emails:
            try:
                # Build table rows for all assignments
                table_rows_html = []
                table_rows_plain = []

                for payload in assignment_payloads:
                    assignment = payload.get('assignment') or {}
                    engineer = payload.get('engineer')

                    engineer_name = engineer.username if engineer else 'Unassigned'
                    test_name = assignment.get('test_name', 'N/A')
                    start_window = _format_assignment_datetime(
                        assignment.get('start_date'), assignment.get('start_time'))
                    end_window = _format_assignment_datetime(
                        assignment.get('end_date'), assignment.get('end_time'))
                    schedule = f"{start_window} â†’ {end_window}"
                    total_hours = assignment.get('total_hours')
                    total_hours_text = f"{total_hours} hrs" if total_hours not in (
                        None, '', []) else 'N/A'
                    product_tco = f"{test_request.product_name or 'Unnamed Product'} / {test_request.tco_id or f'REQ-{test_request.id}'}"

                    # HTML table row
                    table_rows_html.append(f"""
                    <tr>
                        <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{engineer_name}</td>
                        <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{product_tco}</td>
                        <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{test_name}</td>
                        <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{schedule}</td>
                        <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;">{total_hours_text}</td>
                        <td style="padding: 12px; border-bottom: 1px solid #e5e7eb;"><strong style="color:#059669;">{status_label}</strong></td>
                    </tr>
                    """)

                    # Plain text row
                    table_rows_plain.append(
                        f"{engineer_name} | {product_tco} | {test_name} | {schedule} | {total_hours_text} | {status_label}")

                # Create consolidated email
                subject = f"{status_label}: {len(assignment_payloads)} Test(s) - {test_request.tco_id or f'REQ-{test_request.id}'}"
                msg = MIMEMultipart('alternative')
                msg['Subject'] = subject
                msg['From'] = formataddr(
                    (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
                msg['To'] = admin_cc

                greeting_plain = f"Hello {admin_names_text}," if admin_names_text else "Hello Admins,"
                greeting_html = f"Hello <strong>{admin_names_text}</strong>," if admin_names_text else "Hello <strong>Admins</strong>,"

                plain_body = f"""
{status_label}

{greeting_plain}

A test schedule has been proposed for {len(assignment_payloads)} test(s). Please review the test schedule details below and approve or request changes as needed.

Product / TCO ID: {test_request.product_name or 'Unnamed Product'} / {test_request.tco_id or f'REQ-{test_request.id}'}

Test Schedule Summary:
{'=' * 100}
Name | Product / TCO | Test Name | Schedule | Estimated Effort | Status
{'-' * 100}
{chr(10).join(table_rows_plain)}
{'=' * 100}

Please log in to the Test Request Management System to review the full plan and attachments.

Regards,
Test Request Management System

---
This is an automated notification. Please do not reply.
"""

                html_body = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #111827;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
        }}
        .header {{
            background: linear-gradient(135deg, #312e81 0%, #2563eb 100%);
            color: #fff;
            padding: 28px;
            border-radius: 12px 12px 0 0;
            text-align: center;
        }}
        .header h1 {{
            margin: 0;
            font-size: 24px;
        }}
        .header p {{
            margin: 8px 0 0 0;
            opacity: 0.85;
            font-size: 14px;
        }}
        .content {{
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-top: none;
            padding: 28px;
        }}
        .info-box {{
            background: #f9fafb;
            border-left: 4px solid #2563eb;
            padding: 18px;
            margin: 22px 0;
            border-radius: 10px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background: #ffffff;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        th {{
            background: #312e81;
            color: #fff;
            padding: 12px;
            text-align: left;
            font-weight: 600;
            border-bottom: 2px solid #1e1b4b;
        }}
        td {{
            padding: 12px;
            border-bottom: 1px solid #e5e7eb;
        }}
        tr:hover {{
            background-color: #f9fafb;
        }}
        .footer {{
            background: #f9fafb;
            border: 1px solid #e5e7eb;
            border-top: none;
            border-radius: 0 0 12px 12px;
            padding: 18px;
            text-align: center;
            color: #6b7280;
            font-size: 12px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{status_label}</h1>
        <p>{len(assignment_payloads)} test schedule(s) proposed for review</p>
    </div>
    <div class="content">
        <p>{greeting_html}</p>
        <p>A test schedule has been proposed for <strong>{len(assignment_payloads)} test(s)</strong>. Please review the test schedule details below and approve or request changes as needed.</p>
        <div class="info-box">
            <div style="margin-bottom: 15px;">
                <strong>Product / TCO ID:</strong> {test_request.product_name or 'Unnamed Product'} / {test_request.tco_id or f'REQ-{test_request.id}'}
            </div>
        </div>
        <table>
            <thead>
                <tr>
                    <th>Name</th>
                    <th>Product / TCO</th>
                    <th>Test Name</th>
                    <th>Schedule</th>
                    <th>Estimated Effort</th>
                    <th>Status</th>
                </tr>
            </thead>
            <tbody>
                {''.join(table_rows_html)}
            </tbody>
        </table>
        <p style="margin-top:20px;">Log in to the Test Request Management System for full procedures, attachments, and planner availability.</p>
        <p style="margin-top:20px;">Regards,<br/>Test Request Management System</p>
    </div>
    <div class="footer">
        Automated notification Â· Please do not reply
    </div>
</body>
</html>
"""

                msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
                msg.attach(MIMEText(html_body, 'html', 'utf-8'))

                recipients = list({email for email in (admin_emails or [])})
                with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                    server.sendmail(
                        flask_app.config['SMTP_FROM_EMAIL'],
                        recipients,
                        msg.as_string()
                    )
                logger.info(
                    'Sent consolidated test assignment email to admins: %s', ', '.join(admin_emails))
                return  # Exit early for consolidated email
            except Exception as exc:
                logger.error(
                    'Failed to send consolidated test assignment email: %s', exc)
                # Fall through to individual emails if consolidation fails

        # Original logic for individual emails (for other statuses)
        for payload in assignment_payloads:
            engineer = payload.get('engineer')
            assignment = payload.get('assignment') or {}
            if not engineer or not engineer.email:
                continue

            try:
                subject = f"{status_label}: {assignment.get('test_name', 'Lab Test')} ({test_request.tco_id or f'REQ-{test_request.id}'})"
                msg = MIMEMultipart('alternative')
                msg['Subject'] = subject
                msg['From'] = formataddr(
                    (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))

                # When requested, send TO all admins and (optionally) CC the assigned engineer
                if send_to_admins_first and admin_cc:
                    msg['To'] = admin_cc
                    # For "Test Plan To Approve" status, do NOT CC the assigned engineer
                    if status_label != 'Test Plan To Approve':
                        msg['Cc'] = engineer.email
                else:
                    msg['To'] = engineer.email
                    if admin_cc:
                        msg['Cc'] = admin_cc

                start_window = _format_assignment_datetime(
                    assignment.get('start_date'), assignment.get('start_time'))
                end_window = _format_assignment_datetime(
                    assignment.get('end_date'), assignment.get('end_time'))
                total_hours = assignment.get('total_hours')
                total_hours_text = f"{total_hours} hrs" if total_hours not in (
                    None, '', []) else 'N/A'

                # Customize greeting based on status
                if status_label == 'Test Plan To Approve' and admin_names_text:
                    greeting_plain = f"Hello {admin_names_text},"
                    greeting_html = f"Hello <strong>{admin_names_text}</strong>,"
                    body_text_plain = f"A test schedule has been proposed for: {assignment.get('test_name')}"
                    body_text_html = "Please review the test schedule details and approve or request changes as needed."
                else:
                    greeting_plain = f"Hello {engineer.username},"
                    greeting_html = f"Hello <strong>{engineer.username}</strong>,"
                    body_text_plain = f"You have been scheduled for: {assignment.get('test_name')}"
                    body_text_html = "Please review the assignment details and begin planning your lab time."

                plain_body = f"""
{status_label}

{greeting_plain}

{body_text_plain}
Product / TCO ID: {test_request.product_name or 'Unnamed Product'} / {test_request.tco_id or f'REQ-{test_request.id}'}
Schedule: {start_window} â†’ {end_window}
Estimated Effort: {total_hours_text}

Please log in to the Test Request Management System to review the full plan and attachments.

Regards,
Test Request Management System

---
This is an automated notification. Please do not reply.
"""

                html_body = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #111827;
            max-width: 620px;
            margin: 0 auto;
            padding: 20px;
        }}
        .header {{
            background: linear-gradient(135deg, #312e81 0%, #2563eb 100%);
            color: #fff;
            padding: 28px;
            border-radius: 12px 12px 0 0;
            text-align: center;
        }}
        .header h1 {{
            margin: 0;
            font-size: 24px;
        }}
        .header p {{
            margin: 8px 0 0 0;
            opacity: 0.85;
            font-size: 14px;
        }}
        .content {{
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-top: none;
            padding: 28px;
        }}
        .info-box {{
            background: #f9fafb;
            border-left: 4px solid #2563eb;
            padding: 18px;
            margin: 22px 0;
            border-radius: 10px;
        }}
        .detail-row {{
            display: flex;
            padding: 10px 0;
            border-bottom: 1px solid #e5e7eb;
        }}
        .detail-label {{
            width: 150px;
            font-weight: 600;
            color: #6b7280;
        }}
        .detail-value {{
            flex: 1;
            color: #0f172a;
        }}
        .footer {{
            background: #f9fafb;
            border: 1px solid #e5e7eb;
            border-top: none;
            border-radius: 0 0 12px 12px;
            padding: 18px;
            text-align: center;
            color: #6b7280;
            font-size: 12px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{status_label}</h1>
        <p>A test schedule has been proposed for {assignment.get('test_name')}</p>
    </div>
    <div class="content">
        <p>{greeting_html}</p>
        <p>{body_text_html}</p>
        <div class="info-box">
            <div class="detail-row">
                <div class="detail-label">Product / TCO</div>
                <div class="detail-value"><strong>{test_request.product_name or 'Unnamed Product'}</strong> Â· {test_request.tco_id or f'REQ-{test_request.id}'}</div>
            </div>
            <div class="detail-row">
                <div class="detail-label">Test Name</div>
                <div class="detail-value">{assignment.get('test_name')}</div>
            </div>
            <div class="detail-row">
                <div class="detail-label">Schedule</div>
                <div class="detail-value">{start_window} â†’ {end_window}</div>
            </div>
            <div class="detail-row">
                <div class="detail-label">Estimated Effort</div>
                <div class="detail-value">{total_hours_text}</div>
            </div>
            <div class="detail-row">
                <div class="detail-label">Status</div>
                <div class="detail-value"><strong style="color:#059669;">{status_label}</strong></div>
            </div>
        </div>
        <p style="margin-top:20px;">Log in to the Test Request Management System for full procedures, attachments, and planner availability.</p>
        <p style="margin-top:20px;">Regards,<br/>Test Request Management System</p>
    </div>
    <div class="footer">
        Automated notification Â· Please do not reply
    </div>
</body>
</html>
"""

                msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
                msg.attach(MIMEText(html_body, 'html', 'utf-8'))

                with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                    # Ensure all intended recipients receive the email
                    # For "Test Plan To Approve", only send to admins (not engineer)
                    if status_label == 'Test Plan To Approve':
                        recipients = list(
                            {email for email in (admin_emails or [])})
                    else:
                        recipients = list({email for email in (
                            [engineer.email] if engineer.email else []) + (admin_emails or [])})
                    server.sendmail(
                        flask_app.config['SMTP_FROM_EMAIL'],
                        recipients,
                        msg.as_string()
                    )
                if status_label == 'Test Plan To Approve':
                    logger.info(
                        'Sent test assignment email to admins: %s', ', '.join(admin_emails))
                else:
                    logger.info('Sent test assignment email to %s',
                                engineer.email)
            except Exception as exc:
                logger.error('Failed to send test assignment email to %s: %s',
                             getattr(engineer, 'email', 'unknown'), exc)

    def send_equipment_reminder_emails():
        """Send email reminders for equipment due dates (Calibration, IC, Maintenance) based on EOU status.

        Reminder schedule:
        - EOU: 2 months, 1 month, 15 days, 1 week before due date
        - Non EOU: 1 month, 15 days, 1 week before due date
        """
        try:
            today = get_ist_now().date()
            lab_engineer_recipients = [
                u.email for u in User.query.filter_by(
                    role='lab_engineer', is_active=True).all() if u.email
            ]
            admin_recipients = [
                u.email for u in User.query.filter_by(
                    role='admin', is_active=True).all() if u.email
            ]

            if not lab_engineer_recipients and not admin_recipients:
                logger.info(
                    'No lab engineer or admin recipients found; skipping equipment reminders.')
                return

            def _is_yes(value):
                return str(value or '').strip().lower() == 'yes'

            def _is_not_no(value):
                return str(value or '').strip().lower() != 'no'

            def _reminder_days_for_equipment(equipment):
                eou_status = getattr(equipment, 'eou_status', None) or ''
                if eou_status.strip().upper() == 'EOU':
                    return [60, 30, 15, 7]
                return [30, 15, 7]

            reminders = []
            seen_reminders = set()

            def _append_reminder(reminder_type, equipment, due_date, days_remaining, **extra):
                if not equipment or not due_date or days_remaining <= 0:
                    return
                reminder_days = _reminder_days_for_equipment(equipment)
                if days_remaining not in reminder_days:
                    return
                dedupe_key = (
                    reminder_type,
                    getattr(equipment, 'id', None),
                    due_date.isoformat(),
                )
                if dedupe_key in seen_reminders:
                    return
                seen_reminders.add(dedupe_key)
                reminder = {
                    'type': reminder_type,
                    'equipment': equipment,
                    'due_date': due_date,
                    'days_remaining': days_remaining,
                }
                reminder.update(extra)
                reminders.append(reminder)

            for equipment in Equipment.query.filter(Equipment.calibration_due_date.isnot(None)).all():
                if not _is_yes(equipment.calibration_required):
                    continue
                due_date = equipment.calibration_due_date
                if not due_date:
                    continue
                _append_reminder(
                    'Calibration',
                    equipment,
                    due_date,
                    (due_date - today).days,
                )

            for equipment in Equipment.query.filter(Equipment.ic_due_date.isnot(None)).all():
                if not _is_yes(equipment.ic_required):
                    continue
                due_date = equipment.ic_due_date
                if not due_date:
                    continue
                _append_reminder(
                    'IC (Intermediate Check)',
                    equipment,
                    due_date,
                    (due_date - today).days,
                )

            for equipment in Equipment.query.filter(Equipment.maintenance_due_date.isnot(None)).all():
                if not _is_not_no(equipment.maintenance_required):
                    continue
                due_date = equipment.maintenance_due_date
                if not due_date:
                    continue
                _append_reminder(
                    'Maintenance',
                    equipment,
                    due_date,
                    (due_date - today).days,
                    source='equipment',
                )

            maintenance_records = Maintenance.query.filter(
                Maintenance.maintenance_due_date.isnot(None)
            ).all()
            for maintenance in maintenance_records:
                due_date = maintenance.maintenance_due_date
                equipment = maintenance.equipment
                if not equipment or not due_date:
                    continue
                if not _is_not_no(maintenance.maintenance_required):
                    continue
                _append_reminder(
                    'Maintenance',
                    equipment,
                    due_date,
                    (due_date - today).days,
                    maintenance_id=maintenance.id,
                    source='maintenance_record',
                )

            if reminders:
                calibration_reminders = [
                    r for r in reminders if r['type'] == 'Calibration']
                ic_reminders = [
                    r for r in reminders if r['type'] == 'IC (Intermediate Check)']
                maintenance_reminders = [
                    r for r in reminders if r['type'] == 'Maintenance']

                if calibration_reminders:
                    _send_equipment_reminder_email(
                        'Calibration Due Date Reminder',
                        calibration_reminders,
                        lab_engineer_recipients,
                        admin_recipients
                    )

                if ic_reminders:
                    _send_equipment_reminder_email(
                        'IC (Intermediate Check) Due Date Reminder',
                        ic_reminders,
                        lab_engineer_recipients,
                        admin_recipients
                    )

                if maintenance_reminders:
                    _send_equipment_reminder_email(
                        'Maintenance Due Date Reminder',
                        maintenance_reminders,
                        lab_engineer_recipients,
                        admin_recipients
                    )

                logger.info(
                    'Sent equipment reminder emails: calibration=%s, ic=%s, maintenance=%s, total=%s',
                    len(calibration_reminders),
                    len(ic_reminders),
                    len(maintenance_reminders),
                    len(reminders),
                )
            else:
                logger.info('No equipment reminders to send today.')

        except Exception as e:
            logger.error(f'Error sending equipment reminder emails: {e}')
            import traceback
            logger.error(traceback.format_exc())

    def _send_equipment_reminder_email(subject: str, reminders: list, lab_engineer_recipients: list, admin_recipients: list):
        """Send HTML email with equipment reminder information."""
        try:
            # Format days remaining text
            def format_days(days):
                if days == 60:
                    return '2 months'
                elif days == 30:
                    return '1 month'
                elif days == 15:
                    return '15 days'
                elif days == 7:
                    return '1 week'
                else:
                    return f'{days} days'

            # Build reminder items HTML
            reminder_items_html = ''
            for reminder in reminders:
                equipment = reminder['equipment']
                due_date = reminder['due_date']
                days_remaining = reminder['days_remaining']
                reminder_type = reminder['type']

                eou_status = getattr(
                    equipment, 'eou_status', None) or 'Not specified'
                asset_id = equipment.asset_id or 'N/A'
                name = equipment.name or 'N/A'
                make = equipment.make or 'N/A'
                model_no = equipment.model_no or 'N/A'
                serial_no = equipment.serial_no or 'N/A'
                location = equipment.location or 'N/A'

                reminder_items_html += f"""
                <div style="background: #f9fafb; border-left: 4px solid #3b82f6; padding: 15px; margin: 15px 0; border-radius: 4px;">
                    <div style="font-weight: bold; color: #1f2937; margin-bottom: 10px; font-size: 16px;">
                        {name}
                    </div>
                    <div style="display: grid; grid-template-columns: 150px 1fr; gap: 8px; font-size: 14px;">
                        <div style="color: #6b7280; font-weight: bold;">Asset ID:</div>
                        <div style="color: #111827;">{asset_id}</div>
                        <div style="color: #6b7280; font-weight: bold;">Make:</div>
                        <div style="color: #111827;">{make}</div>
                        <div style="color: #6b7280; font-weight: bold;">Model No:</div>
                        <div style="color: #111827;">{model_no}</div>
                        <div style="color: #6b7280; font-weight: bold;">Serial No:</div>
                        <div style="color: #111827;">{serial_no}</div>
                        <div style="color: #6b7280; font-weight: bold;">Location:</div>
                        <div style="color: #111827;">{location}</div>
                        <div style="color: #6b7280; font-weight: bold;">EOU Status:</div>
                        <div style="color: #111827;">{eou_status}</div>
                        <div style="color: #6b7280; font-weight: bold;">Due Date:</div>
                        <div style="color: #dc2626; font-weight: bold;">{due_date.strftime('%B %d, %Y')}</div>
                        <div style="color: #6b7280; font-weight: bold;">Time Remaining:</div>
                        <div style="color: #dc2626; font-weight: bold;">{format_days(days_remaining)}</div>
                    </div>
                </div>
                """

            # Get reminder type from first reminder for email body
            reminder_type_name = reminders[0]['type'].lower(
            ) if reminders else 'reminder'

            # Create HTML email body
            html_body = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        color: #333333;
                        max-width: 700px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
                        color: white;
                        padding: 30px;
                        text-align: center;
                        border-radius: 8px 8px 0 0;
                    }}
                    .content {{
                        background: #ffffff;
                        padding: 30px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                    .info-box {{
                        background: #fef3c7;
                        border-left: 4px solid #f59e0b;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .footer {{
                        background: #f9fafb;
                        padding: 20px;
                        text-align: center;
                        color: #6b7280;
                        font-size: 12px;
                        border-radius: 0 0 8px 8px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1 style="margin: 0; font-size: 24px;">Equipment {subject}</h1>
                </div>
                <div class="content">
                    <div class="info-box">
                        <strong>Reminder:</strong> The following equipment items have {reminder_type_name} due dates approaching.
                        Please ensure necessary actions are taken before the due dates.
                    </div>
                    
                    <h2 style="color: #1f2937; margin-top: 30px; margin-bottom: 20px;">Equipment Requiring Attention ({len(reminders)} item(s)):</h2>
                    
                    {reminder_items_html}
                    
                    <p style="margin-top: 30px; color: #6b7280; font-size: 14px;">
                        Please review the equipment listed above and take appropriate action to ensure compliance.
                    </p>
                </div>
                <div class="footer">
                    <p style="margin: 5px 0 0 0;">This is an automated reminder from the Equipment Management System.</p>
                    <p style="margin: 5px 0 0 0;">Please do not reply to this email.</p>
                </div>
            </body>
            </html>
            """

            # Create plain text version
            plain_body = f"""
Equipment {subject}

Reminder: The following equipment items have {reminder_type_name} due dates approaching.
Please ensure necessary actions are taken before the due dates.

Equipment Requiring Attention ({len(reminders)} item(s)):

"""
            for reminder in reminders:
                equipment = reminder['equipment']
                due_date = reminder['due_date']
                days_remaining = reminder['days_remaining']

                plain_body += f"""
- {equipment.name or 'N/A'}
  Asset ID: {equipment.asset_id or 'N/A'}
  Make: {equipment.make or 'N/A'}
  Model No: {equipment.model_no or 'N/A'}
  Serial No: {equipment.serial_no or 'N/A'}
  Location: {equipment.location or 'N/A'}
  EOU Status: {getattr(equipment, 'eou_status', None) or 'Not specified'}
  Due Date: {due_date.strftime('%B %d, %Y')}
  Time Remaining: {format_days(days_remaining)}

"""

            plain_body += """
Please review the equipment listed above and take appropriate action to ensure compliance.

This is an automated reminder from the Equipment Management System.
Please do not reply to this email.
"""

            # Create email message
            msg = MIMEMultipart('alternative')
            msg['Subject'] = subject
            msg['From'] = formataddr(
                (flask_app.config['SMTP_FROM_NAME'], flask_app.config['SMTP_FROM_EMAIL']))
            msg['To'] = ', '.join(lab_engineer_recipients)
            if admin_recipients:
                msg['Cc'] = ', '.join(admin_recipients)

            # Attach both plain text and HTML versions
            msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
            msg.attach(MIMEText(html_body, 'html', 'utf-8'))

            # Send email to all recipients (To + CC)
            all_recipients = lab_engineer_recipients + admin_recipients

            with smtplib.SMTP(flask_app.config['SMTP_SERVER'], flask_app.config['SMTP_PORT'], timeout=10) as server:
                server.sendmail(
                    flask_app.config['SMTP_FROM_EMAIL'],
                    all_recipients,
                    msg.as_string()
                )

            logger.info(
                f'Equipment reminder email sent to {len(lab_engineer_recipients)} lab engineer(s) and {len(admin_recipients)} admin(s)')

        except Exception as e:
            logger.error(f'Failed to send equipment reminder email: {e}')
            import traceback
            logger.error(traceback.format_exc())

    # Expose key helpers for scripts/tests
    # type: ignore[attr-defined]
    flask_app.send_test_assignment_emails = send_test_assignment_emails
    # type: ignore[attr-defined]
    flask_app.send_equipment_reminder_emails = send_equipment_reminder_emails

    flask_app.config.from_object(config[resolved_config_name])
    config[resolved_config_name].init_app(flask_app)
    flask_app.config['APP_ENVIRONMENT'] = resolved_config_name
    flask_app.config['SHOW_ENV_MARKER'] = resolved_config_name != 'production'
    flask_app.config['ENV_MARKER_LABEL'] = (
        'TESTING' if resolved_config_name == 'testing' else 'NON-PRODUCTION'
    )

    @flask_app.context_processor
    def inject_environment_marker():
        return {
            'app_version': APP_VERSION,
            'app_environment': flask_app.config.get('APP_ENVIRONMENT', 'default'),
            'show_env_marker': flask_app.config.get('SHOW_ENV_MARKER', True),
            'env_marker_label': flask_app.config.get('ENV_MARKER_LABEL', 'NON-PRODUCTION'),
        }

    # Initialize extensions
    db.init_app(flask_app)

    with flask_app.app_context():
        ensure_planner_table()
        ensure_equipment_document_link_column()
        ensure_equipment_test_name_column()

    # Initialize Flask-Login
    login_manager = LoginManager()
    login_manager.init_app(flask_app)
    login_manager.login_view = 'auth.login'  # type: ignore
    login_manager.login_message = 'Please log in to access this page.'
    login_manager.login_message_category = 'info'
    login_manager.session_protection = 'strong'

    @login_manager.user_loader
    def load_user(user_id):
        try:
            normalized_user_id = int(user_id)
        except (TypeError, ValueError):
            return None

        user = db.session.get(User, normalized_user_id)
        if user is None or not getattr(user, 'is_active', True):
            return None
        return user

    @login_manager.unauthorized_handler
    def unauthorized():
        """Handle unauthorized access without next parameter."""
        # Check if this is an API request (JSON expected)
        if request.path.startswith('/api/'):
            return jsonify({
                'success': False,
                'error': 'Authentication required. Please log in.'
            }), 401
        # For regular page requests, redirect to login
        flash('Please log in to access this page.', 'info')
        return redirect(url_for('auth.login'))

    @flask_app.after_request
    def add_auth_cache_headers(response):
        """Prevent authenticated pages from being cached across users."""
        if request.path.startswith('/static/'):
            return response

        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, private, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'

        existing_vary = response.headers.get('Vary')
        if existing_vary:
            vary_parts = [part.strip() for part in existing_vary.split(',') if part.strip()]
            if 'Cookie' not in vary_parts:
                vary_parts.append('Cookie')
                response.headers['Vary'] = ', '.join(vary_parts)
        else:
            response.headers['Vary'] = 'Cookie'

        return response

    # Initialize managers
    # Note: document_processor and enhanced_document_processor may be used by create_upload_routes
    document_processor = DocumentProcessor(flask_app.config['UPLOAD_FOLDER'])  # noqa: F841
    enhanced_document_processor = EnhancedDocumentProcessor(  # noqa: F841
        flask_app.config['UPLOAD_FOLDER'],
        flask_app.config['TEMPLATE_FOLDER'],
        flask_app.config['OUTPUT_FOLDER']
    )
    document_generator = DocumentGenerator(
        flask_app.config['TEMPLATE_FOLDER'], flask_app.config['OUTPUT_FOLDER'])
    equipment_manager_instance = EquipmentManager(
        flask_app.config['EQUIPMENT_DATA_FILE'])

    # Register enhanced test plan creation routes
    create_upload_routes(flask_app)

    # Datasheet generation module (new, modular - see datasheet_gen/)
    from datasheet_gen import register_datasheet_gen
    register_datasheet_gen(flask_app)

    # NL search over lab data (admin "Ask the Lab Data" tool - see nlp_search/)
    from nlp_search import register_nlp_search
    register_nlp_search(flask_app)

    # Import and register authentication routes
    from auth_routes import auth_bp
    flask_app.register_blueprint(auth_bp)

    @flask_app.route('/')
    @login_required
    def index():
        """Main page with test plan creation form."""

        # Status filters shown on requester index page
        status_options = [
            'Draft',
            'Draft Report',
            'Proceed Report',
            'In Progress',
            'Need More Information',
            'Completed',
            'Rejected',
        ]

        # Collect all service types from all requests for filter dropdown
        all_requests = EMCRequest.query.all()
        all_service_types = [
            req.service_types for req in all_requests
            if hasattr(req, 'service_types') and req.service_types
        ]

        return render_template(
            'index.html',
            status_options=status_options,
            service_type_options=_get_service_type_filter_options(all_service_types),
        )

    @flask_app.route('/help')
    def help_page():
        """Application help and usage guidelines page."""
        return render_template('help.html')

    def _safe_json_dump(value):
        if value is None:
            return None
        if isinstance(value, str):
            return value  # already a string (json or plain)
        if isinstance(value, (dict, list)):
            return json.dumps(value, default=str)
        return json.dumps(value, default=str)

    def _safe_float(value):
        if value in (None, ''):
            return None
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    def _normalize_dimensions_to_mm(dimension_unit, length, width, height):
        unit = (dimension_unit or 'mm').strip().lower()
        conversion_factors = {
            'mm': 1.0,
            'millimeter': 1.0,
            'millimeters': 1.0,
            'in': 25.4,
            'inch': 25.4,
            'inches': 25.4,
            'ft': 304.8,
            'foot': 304.8,
            'feet': 304.8
        }
        factor = conversion_factors.get(unit, 1.0)

        def convert(value):
            parsed = _safe_float(value)
            if parsed is None:
                return None
            return parsed * factor

        return 'mm', convert(length), convert(width), convert(height)

    def _resolve_test_standard_values(raw_standards, test_code: str, product_standards=None) -> list:
        """Return canonical test-standard labels with fallback inference."""
        options_by_test = {
            'CE': [
                ('CISPR 11', 'cispr 11'),
                ('EN 55011', 'en 55011'),
                ('ANSI C63.4', 'ansi c63.4', 'ansi c63-4'),
            ],
            'RE': [
                ('CISPR 11', 'cispr 11'),
                ('EN 55011', 'en 55011'),
                ('ANSI C63.4', 'ansi c63.4', 'ansi c63-4'),
            ],
            'HARMONIC': [
                ('IEC 61000-3-2', 'iec 61000-3-2'),
                ('EN 61000-3-2', 'en 61000-3-2'),
            ],
            'FLICKER': [
                ('IEC 61000-3-3', 'iec 61000-3-3'),
                ('EN 61000-3-3', 'en 61000-3-3'),
            ],
            'ESD': [
                ('IEC 61000-4-2', 'iec 61000-4-2'),
                ('EN 61000-4-2', 'en 61000-4-2'),
            ],
            'RS': [
                ('IEC 61000-4-3', 'iec 61000-4-3'),
                ('EN 61000-4-3', 'en 61000-4-3'),
            ],
            'RS_INTERIM': [
                ('IEC 61000-4-3', 'iec 61000-4-3'),
                ('EN 61000-4-3', 'en 61000-4-3'),
            ],
            'EFT': [
                ('IEC 61000-4-4', 'iec 61000-4-4'),
                ('EN 61000-4-4', 'en 61000-4-4'),
            ],
            'SURGE': [
                ('IEC 61000-4-5', 'iec 61000-4-5'),
                ('EN 61000-4-5', 'en 61000-4-5'),
            ],
            'CRF': [
                ('IEC 61000-4-6', 'iec 61000-4-6'),
                ('EN 61000-4-6', 'en 61000-4-6'),
            ],
            'POWER_FREQ': [
                ('IEC 61000-4-8', 'iec 61000-4-8'),
                ('EN 61000-4-8', 'en 61000-4-8'),
            ],
            'VOLTAGE_DIPS': [
                ('IEC 61000-4-11', 'iec 61000-4-11'),
                ('EN 61000-4-11', 'en 61000-4-11'),
            ],
        }

        normalized_test_code = str(test_code or '').strip().upper()
        options = options_by_test.get(normalized_test_code, [])
        if not options:
            return []

        selected_set = _selected_option_set(raw_standards)
        resolved = []
        for label, *aliases in options:
            if _selection_contains(selected_set, label, *aliases):
                resolved.append(label)
        if resolved:
            return resolved

        # Fallback for newer form payloads where standards are not sent directly:
        # infer standards from selected product standards.
        product_selected = _selected_option_set(product_standards)
        has_iec_family = any(
            token.startswith('iec ') or 'iec ' in token or 'cispr' in token
            for token in product_selected
        )
        has_en_family = any(
            token.startswith('en ') or 'en ' in token
            for token in product_selected
        )
        has_ansi_family = _selection_contains(
            product_selected,
            'ansi c63.4',
            'ansi c63-4',
            'fcc subpart 15b',
            'ices-001'
        )

        inferred = []
        if normalized_test_code in ('CE', 'RE'):
            if has_iec_family:
                inferred.append('CISPR 11')
            if has_en_family:
                inferred.append('EN 55011')
            if has_ansi_family:
                inferred.append('ANSI C63.4')
        elif normalized_test_code == 'HARMONIC':
            if has_iec_family:
                inferred.append('IEC 61000-3-2')
            if has_en_family:
                inferred.append('EN 61000-3-2')
        else:
            iec_label = next((label for label, *_ in options if label.startswith('IEC ')), None)
            en_label = next((label for label, *_ in options if label.startswith('EN ')), None)
            if has_iec_family and iec_label:
                inferred.append(iec_label)
            if has_en_family and en_label:
                inferred.append(en_label)

        return inferred

    def _populate_iec_request_from_form(test_request_obj, form_data):
        """Map submitted form data into the normalized EMCRequest model."""
        return populate_emc_request_from_form(
            test_request_obj,
            form_data,
            normalize_dimensions_to_mm=_normalize_dimensions_to_mm,
            resolve_test_standard_values=_resolve_test_standard_values,
            today_date=get_ist_now().date(),
        )

    @flask_app.route('/api/save-draft', methods=['POST'])
    @login_required
    def save_draft():
        """Save form data as a normalized EMCRequest draft."""
        try:
            data = request.get_json()
            if not data:
                return jsonify({
                    'success': False,
                    'error': 'No data provided'
                }), 400

            form_data = data.get('form_data', {})
            # For updating existing draft
            draft_id = data.get('draft_id', None)

            if draft_id:
                test_request = _resolve_request(draft_id)
                if (
                    not test_request
                    or test_request.user_id != current_user.id
                    or test_request.status != 'Draft'
                ):
                    return jsonify({
                        'success': False,
                        'error': 'Draft not found or you do not have permission to edit it'
                    }), 404
            else:
                test_request = EMCRequest(
                    user_id=current_user.id,
                    status='Draft'
                )

            _populate_iec_request_from_form(test_request, form_data)
            test_request.status = 'Draft'

            if not draft_id:
                db.session.add(test_request)

            db.session.commit()

            return jsonify({
                'success': True,
                'message': 'Draft saved successfully',
                'data': {
                    'id': test_request.id,
                    'legacy_request_id': getattr(test_request, 'legacy_request_id', None),
                    'tco_id': test_request.tco_id,
                    'status': test_request.status,
                    'created_at': test_request.created_at.isoformat() if test_request.created_at else None,
                    'updated_at': test_request.updated_at.isoformat() if test_request.updated_at else None
                }
            })

        except Exception as e:
            db.session.rollback()
            logger.error(f'Error saving draft: {e}')
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @flask_app.route('/api/test-requests', methods=['GET'])
    @login_required
    def get_test_requests():
        """Get all test requests for the current user with filtering, search, and pagination."""
        try:
            logger.info(
                f'Fetching test requests for user_id: {current_user.id}')

            # ------------------------------------------------------------------ #
            #  Query parameters                                                    #
            # ------------------------------------------------------------------ #
            status_filter = (request.args.get('status')
                             or 'all').strip() or 'all'
            search_query = (request.args.get('search') or '').strip()
            service_type_filter = (request.args.get('service_type') or '').strip()
            sort_by, sort_dir = _normalize_request_sort_args(
                request.args.get('sort_by'),
                request.args.get('sort_dir')
            )
            peer_review_filter = request.args.get(
                'peer_review', '').lower() == 'true'

            try:
                page = max(int(request.args.get('page', 1)), 1)
            except (TypeError, ValueError):
                page = 1

            try:
                per_page = min(
                    max(int(request.args.get('per_page', 10)), 1), 50)
            except (TypeError, ValueError):
                per_page = 10

            logger.info(
                f'Query params: status={status_filter}, search={search_query}, '
                f'service_type={service_type_filter}, sort_by={sort_by}, '
                f'sort_dir={sort_dir}, '
                f'peer_review={peer_review_filter}, page={page}, per_page={per_page}'
            )

            # ------------------------------------------------------------------ #
            #  Base query â€” scoped by role                                         #
            # ------------------------------------------------------------------ #
            try:
                if current_user.role == 'admin':
                    non_draft_filter = db.or_(
                        EMCRequest.status.is_(None),
                        db.func.lower(db.func.trim(EMCRequest.status)) != 'draft'
                    )
                    base_query = EMCRequest.query.filter(non_draft_filter)
                    logger.info(
                        f'Admin user {current_user.id} â€” showing all submitted requests and own drafts')
                else:
                    base_query = EMCRequest.query.filter_by(
                        user_id=current_user.id)
                    logger.info(
                        f'User {current_user.id} (role: {current_user.role}) '
                        f'â€” showing only own test requests'
                    )
            except Exception as e:
                logger.error(f'Error creating base query: {e}')
                raise

            # ------------------------------------------------------------------ #
            #  Status mapping definitions                                           #
            # ------------------------------------------------------------------ #
            STATUS_MAPPING = {
                'draft': [
                    'draft',
                ],
                'draft_report': [
                    'draft report',
                ],
                'in_progress': [
                    'at review',
                    'assigned lab engineer',
                    'test schedule in progress',
                    'test plan to approve',
                    'test plan approved',
                    'engineer assigned for test',
                    'update plan',
                    'in progress',
                    'datasheet uploaded',
                    'proceed report',
                    'admin sign off',
                    'approved',
                    'assigned',
                ],
                'completed': [
                    'completed',
                ],
                'cancelled': [
                    'cancelled',
                    'rejected',
                ],
                'peer_review': [
                    'peer review',
                    'in peer review',
                    'awaiting peer review',
                    'under peer review',
                ],
            }

            has_review_messages_expr = db.case(
                (
                    db.and_(
                        EMCRequest.review_comments.isnot(None),
                        db.func.trim(EMCRequest.review_comments) != '',
                        db.func.trim(EMCRequest.review_comments) != '[]',
                        db.func.trim(EMCRequest.review_comments) != '{}',
                        db.func.lower(db.func.trim(EMCRequest.review_comments)) != 'null',
                    ),
                    True,
                ),
                else_=False,
            ).label('has_review_messages')

            list_query = base_query.with_entities(
                EMCRequest.id.label('id'),
                EMCRequest.tco_id.label('tco_id'),
                EMCRequest.job_number.label('job_number'),
                EMCRequest.product_name.label('product_name'),
                EMCRequest.manufacturer.label('manufacturer'),
                EMCRequest.model_number.label('model_number'),
                EMCRequest.requester_name.label('requester_name'),
                EMCRequest.assigned_engineer_name.label('assigned_engineer_name'),
                EMCRequest.status.label('status'),
                EMCRequest.submitted_at.label('submitted_at'),
                EMCRequest.created_at.label('created_at'),
                has_review_messages_expr,
            )

            # ------------------------------------------------------------------ #
            #  Apply filters                                                     #
            # ------------------------------------------------------------------ #
            query = list_query

            # Peer Review Filter
            if peer_review_filter:
                query = query.filter(
                    db.func.lower(EMCRequest.status).in_(
                        [s.lower() for s in STATUS_MAPPING['peer_review']]
                    )
                )
                logger.info(f'Applied peer review filter')
            elif status_filter != 'all':
                filter_lower = status_filter.lower().replace(' ', '_')
                if filter_lower in STATUS_MAPPING:
                    query = query.filter(
                        db.func.lower(EMCRequest.status).in_(
                            [s.lower() for s in STATUS_MAPPING[filter_lower]]
                        )
                    )
                else:
                    query = query.filter(
                        db.func.lower(
                            EMCRequest.status) == status_filter.lower()
                    )
                logger.info(f'Applied status filter: {status_filter}')

            # Search filter â€” case-insensitive across key text fields
            if search_query:
                pattern = f'%{search_query}%'
                query = query.filter(
                    db.or_(
                        EMCRequest.product_name.ilike(pattern),
                        EMCRequest.tco_id.ilike(pattern),
                        EMCRequest.model_number.ilike(pattern),
                        EMCRequest.manufacturer.ilike(pattern),
                        EMCRequest.requester_name.ilike(pattern),
                        EMCRequest.assigned_engineer_name.ilike(pattern),
                        EMCRequest.status.ilike(pattern),
                    )
                )
                logger.info(f'Applied search filter: {search_query}')

            # ------------------------------------------------------------------ #
            #  Ordering & pagination                                               #
            # ------------------------------------------------------------------ #
            filtered_request_rows = query.all()

            request_service_types_by_request_id = {}
            if service_type_filter:
                candidate_request_ids = [
                    tr.id for tr in filtered_request_rows if tr.id is not None
                ]
                if candidate_request_ids:
                    candidate_requests = EMCRequest.query.options(
                        joinedload(EMCRequest.service_types)
                    ).filter(
                        EMCRequest.id.in_(candidate_request_ids)
                    ).all()
                    request_service_types_by_request_id = {
                        request_obj.id: _extract_service_types(request_obj)
                        for request_obj in candidate_requests
                    }
                    filtered_request_rows = [
                        tr for tr in filtered_request_rows
                        if _matches_service_type_filter(
                            request_service_types_by_request_id.get(tr.id, []),
                            service_type_filter
                        )
                    ]
                    logger.info(f'Applied service type filter: {service_type_filter}')

            if sort_by:
                sorted_requests = _apply_request_identifier_sort(
                    filtered_request_rows,
                    sort_by,
                    sort_dir
                )
            else:
                sorted_requests = sorted(filtered_request_rows, key=_job_number_sort_key)
            # Finished work (completed / cancelled / rejected) goes to the end of the list,
            # after whichever sort was applied, so active requests stay on top.
            sorted_requests = _terminal_last(sorted_requests)
            total_sorted_requests = len(sorted_requests)
            start_index = (page - 1) * per_page
            end_index = start_index + per_page
            test_requests = sorted_requests[start_index:end_index]

            total_pages = (
                (total_sorted_requests + per_page - 1) // per_page
                if total_sorted_requests else 0
            )

            logger.info(
                f'Found {len(test_requests)} test requests '
                f'(page {page}/{total_pages}, total: {total_sorted_requests})'
            )

            page_request_ids = [tr.id for tr in test_requests if tr.id is not None]
            report_entries_by_request_id = {}
            prefetched_service_types_by_request_id = dict(request_service_types_by_request_id)
            request_service_types_by_request_id = {}

            try:
                if page_request_ids:
                    page_requests = EMCRequest.query.options(
                        joinedload(EMCRequest.service_types)
                    ).filter(
                        EMCRequest.id.in_(page_request_ids)
                    ).all()

                    request_service_types_by_request_id = {
                        request_obj.id: _extract_service_types(request_obj)
                        for request_obj in page_requests
                    }
                    request_service_types_by_request_id.update(
                        prefetched_service_types_by_request_id
                    )

                    report_entries = PlannerEntry.query.with_entities(
                        PlannerEntry.test_request_id.label('test_request_id'),
                        PlannerEntry.report_comments.label('report_comments'),
                    ).filter(
                        PlannerEntry.test_request_id.in_(page_request_ids),
                        PlannerEntry.report_file_path.isnot(None)
                    ).order_by(
                        PlannerEntry.test_request_id.asc(),
                        PlannerEntry.updated_at.desc(),
                        PlannerEntry.id.desc()
                    ).all()

                    for entry in report_entries:
                        if entry.test_request_id not in report_entries_by_request_id:
                            report_entries_by_request_id[entry.test_request_id] = entry
            except Exception as prefetch_error:
                logger.warning(
                    f'Unable to prefetch planner/report data for index page: {prefetch_error}')

            # ------------------------------------------------------------------ #
            #  Serialise rows                                                       #
            # ------------------------------------------------------------------ #
            result = []
            for tr in test_requests:
                try:
                    report_entry = report_entries_by_request_id.get(tr.id)
                    has_report = report_entry is not None

                    fallback_draft_label = f'DRAFT-{tr.id}'
                    raw_status = (tr.status or 'Draft').strip() or 'Draft'

                    result.append({
                        'id':              tr.id,
                        'tco_id':          tr.tco_id or fallback_draft_label,
                        'job_number':      tr.job_number or '',
                        'product_name':    tr.product_name or 'Unnamed Product',
                        'manufacturer':    tr.manufacturer or 'N/A',
                        'model_number':    tr.model_number or 'N/A',
                        'requester_name':  tr.requester_name or 'Unknown',
                        'assigned_engineer_name': tr.assigned_engineer_name or '',
                        'assigned_engineer_display': tr.assigned_engineer_name or '',
                        'status':          raw_status,
                        'created_at':      tr.created_at.strftime('%Y-%m-%d') if tr.created_at else None,
                        'submitted_date':  (
                            (getattr(tr, 'submitted_at', None) or tr.created_at).strftime('%Y-%m-%d %H:%M')
                            if (getattr(tr, 'submitted_at', None) or tr.created_at) else None
                        ),
                        'service_types':   request_service_types_by_request_id.get(tr.id, []),
                        'has_report':      has_report,
                        'has_review_messages': bool(getattr(tr, 'has_review_messages', False)),
                    })
                except Exception as e:
                    logger.error(
                        f'Error serialising test request {tr.id}: {e}')
                    import traceback
                    logger.error(traceback.format_exc())
                    continue

            logger.info(f'Serialised {len(result)} test requests')

            # ------------------------------------------------------------------ #
            #  Statistics counts                                                    #
            # ------------------------------------------------------------------ #
            status_category_lookup = {}
            for category, raw_statuses in STATUS_MAPPING.items():
                for raw_status in raw_statuses:
                    status_category_lookup[raw_status.lower()] = category

            grouped_status_counts = (
                base_query.with_entities(
                    db.func.lower(db.func.coalesce(EMCRequest.status, 'draft')).label('status_key'),
                    db.func.count(EMCRequest.id).label('count')
                )
                .group_by('status_key')
                .all()
            )

            counts_by_category = {
                'draft': 0,
                'draft_report': 0,
                'in_progress': 0,
                'completed': 0,
                'cancelled': 0,
                'peer_review': 0,
            }
            total_count = 0

            for grouped_row in grouped_status_counts:
                raw_status = (grouped_row.status_key or 'draft').strip().lower()
                row_count = int(grouped_row.count or 0)
                total_count += row_count
                mapped_category = status_category_lookup.get(raw_status)
                if mapped_category in counts_by_category:
                    counts_by_category[mapped_category] += row_count

            draft_count = counts_by_category['draft']
            draft_report_count = counts_by_category['draft_report']
            in_progress_count = counts_by_category['in_progress']
            completed_count = counts_by_category['completed']
            cancelled_count = counts_by_category['cancelled']
            peer_review_count = counts_by_category['peer_review']

            logger.info(
                f'Statistics: total={total_count}, draft={draft_count}, '
                f'in_progress={in_progress_count}, completed={completed_count}, '
                f'cancelled={cancelled_count}, peer_review={peer_review_count}, '
                f'draft_report={draft_report_count}'
            )

            # ------------------------------------------------------------------ #
            #  Pagination metadata                                                  #
            # ------------------------------------------------------------------ #
            if total_sorted_requests == 0:
                start_item = end_item = 0
            else:
                start_item = start_index + 1
                end_item = start_item + len(test_requests) - 1

            # ------------------------------------------------------------------ #
            #  Build response                                                       #
            # ------------------------------------------------------------------ #
            response_data = {
                'success': True,
                'data': result,
                'counts': {
                    'total':        total_count,
                    'draft':        draft_count,
                    'in_progress':  in_progress_count,
                    'completed':    completed_count,
                    'cancelled':    cancelled_count,
                    'peer_review':  peer_review_count,
                    'draft_report': draft_report_count,  # â† NEW
                },
                'pagination': {
                    'page':          page,
                    'per_page':      per_page,
                    'total_pages':   total_pages,
                    'total_items':   total_sorted_requests,
                    'start_item':    start_item,
                    'end_item':      end_item,
                    'items_on_page': len(test_requests),
                },
            }

            logger.info(f'Returning {len(result)} items, total: {total_count}')
            return jsonify(response_data)

        # ---------------------------------------------------------------------- #
        #  Global error handler                                                    #
        # ---------------------------------------------------------------------- #
        except Exception as e:
            import traceback
            tb = traceback.format_exc()
            logger.error(f'Error fetching test requests: {e}')
            logger.error(f'Traceback: {tb}')
            return jsonify({
                'success': False,
                'error':   str(e),
                'data':    [],
                'counts': {
                    'total':        0,
                    'draft':        0,
                    'in_progress':  0,
                    'completed':    0,
                    'cancelled':    0,
                    'peer_review':  0,
                    'draft_report': 0,  # â† NEW
                },
                'pagination': {
                    'page':          1,
                    'per_page':      10,
                    'total_pages':   0,
                    'total_items':   0,
                    'start_item':    0,
                    'end_item':      0,
                    'items_on_page': 0,
                },
                'traceback': tb if logger.level <= logging.DEBUG else None,
            }), 500

    def _append_datasheet_peer_review_comment(
        entry: PlannerEntry,
        comment: str,
        username: str,
        action_label: str | None = None
    ) -> None:
        """Append a peer-review note to datasheet comments while preserving prior notes."""
        cleaned_comment = str(comment or '').strip()
        if not cleaned_comment:
            return

        timestamp = get_ist_now().strftime('%d %b %Y, %I:%M %p')
        prefix_parts = [timestamp]
        if username:
            prefix_parts.append(username)
        if action_label:
            prefix_parts.append(action_label)

        formatted_entry = f"[{' | '.join(prefix_parts)}]\n{cleaned_comment}"
        existing_comments = str(entry.datasheet_comments or '').strip()
        entry.datasheet_comments = (
            f"{existing_comments}\n\n{formatted_entry}"
            if existing_comments else
            formatted_entry
        )

    def _record_datasheet_transition(
        entry: PlannerEntry,
        to_status: str,
        comment: str = ''
    ) -> None:
        """Record a review decision in the datasheet audit trail.

        The note appended to datasheet_comments above stays - it is what the UI
        renders - but it is free text, so "which datasheets were rejected, by
        whom, and why" could not be answered from it. This writes the same
        decision as a row. Best-effort: the review action has already happened.
        """
        try:
            # land the review decision first. record_transition rolls its own
            # session back if the audit write fails, and the decision is still
            # pending in that same session - committing here keeps a failed
            # audit write from quietly undoing an approval.
            db.session.commit()
            from datasheet_gen.projection import record_transition
            record_transition(
                entry.id, to_status, actor=current_user, comment=comment,
                decided=to_status in ('Approved', 'Rejected'))
        except Exception as exc:  # noqa: BLE001
            logger.warning(
                'Datasheet transition not recorded for entry %s: %s', entry.id, exc)

    def _generate_final_datasheet_after_peer_review(
        entry: PlannerEntry
    ) -> tuple[bool, str, str | None]:
        """Generate the final datasheet file from the saved submitted form."""
        from datasheet_gen import records as datasheet_records

        record = datasheet_records.get_record_for_assignment(entry.id)
        if not record:
            return False, 'No submitted datasheet form found for this entry', None

        form_data = datasheet_records.draft_form(entry.id)
        if not form_data:
            return False, 'No submitted datasheet data found for this entry', None

        raw_code = str(
            record.get('test_code')
            or entry.test_name
            or ''
        ).strip().upper()
        if not raw_code:
            return False, 'Unable to determine datasheet type for this entry', None

        try:
            if raw_code == 'CE':
                from datasheet_gen.routes import _render_ce_docx

                output_path, _, filename = _render_ce_docx(
                    entry,
                    form_data,
                    entry.tco_id,
                    None
                )
                record_code = 'CE'
            elif raw_code == 'SURGE':
                parent_request = db.session.get(EMCRequest, entry.test_request_id)
                if not parent_request:
                    return False, 'Parent test request not found for this Surge datasheet', None
                if not os.path.exists(SURGE_DATASHEET_TEMPLATE_PATH):
                    return False, 'Surge datasheet template not found', None

                os.makedirs(UPLOAD_FOLDER, exist_ok=True)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                safe_tco_id = secure_filename(str(entry.tco_id or parent_request.tco_id or 'TCO'))
                safe_test_name = secure_filename(str(entry.test_name or 'Surge'))
                filename = f"{safe_tco_id}_{safe_test_name}_{timestamp}.docx"
                output_path = os.path.join(UPLOAD_FOLDER, filename)

                _build_surge_datasheet_docx(
                    SURGE_DATASHEET_TEMPLATE_PATH,
                    output_path,
                    parent_request,
                    entry,
                    form_data
                )
                record_code = 'SURGE'
            else:
                from datasheet_gen.generic_routes import _render_datasheet_docx
                from datasheet_gen.registry import load_schema, normalize_code

                record_code = normalize_code(raw_code)
                schema = load_schema(record_code)
                output_path, _, filename = _render_datasheet_docx(
                    record_code,
                    schema,
                    entry,
                    form_data,
                    entry.tco_id
                )
        except Exception as exc:
            logger.error(
                'Failed to generate final datasheet during peer review approval '
                'for planner entry %s: %s',
                entry.id,
                exc
            )
            import traceback
            logger.error(traceback.format_exc())
            return False, 'Failed to generate the final datasheet for this entry', None

        entry.datasheet_file_path = output_path
        db.session.execute(
            text(
                """
                UPDATE datasheet_records
                SET test_code = :test_code,
                    generated_file_path = :generated_file_path,
                    updated_at = :updated_at
                WHERE planner_entry_id = :planner_entry_id
                """
            ),
            {
                'test_code': record_code,
                'generated_file_path': output_path,
                'updated_at': get_ist_now(),
                'planner_entry_id': entry.id,
            }
        )
        return True, filename, output_path

    def _apply_peer_review_action(
        entry: PlannerEntry,
        action: str,
        comment: str = ''
    ) -> tuple[bool, str, int]:
        """Apply a peer-review action to a planner entry."""
        normalized_action = str(action or '').strip().lower()
        if normalized_action not in {'approve', 'reject', 'comment'}:
            return False, 'Invalid peer review action', 400

        if entry.status != 'Peer Review':
            return False, (
                f'Entry is not in Peer Review status. Current status: {entry.status}'
            ), 400

        if current_user.role != 'admin':
            if not entry.peer_reviewer_user_id:
                return False, 'No peer reviewer is assigned to this datasheet yet', 400
            if entry.peer_reviewer_user_id != current_user.id:
                return False, 'This datasheet is assigned to another peer reviewer', 403
            if entry.datasheet_uploaded_by and entry.datasheet_uploaded_by == current_user.id:
                return False, 'You cannot review a datasheet that you uploaded', 403

        comment_text = str(comment or '').strip()
        action_label = None

        if normalized_action == 'approve':
            if not entry.datasheet_file_path:
                generation_success, generation_message, _ = _generate_final_datasheet_after_peer_review(entry)
                if not generation_success:
                    return False, generation_message, 400
            action_label = 'APPROVED'
            entry.status = 'datasheet_uploaded'
            _append_datasheet_peer_review_comment(
                entry,
                comment_text or 'Datasheet approved during peer review.',
                current_user.username,
                action_label
            )
            _update_parent_request_datasheet_status(entry)
            _record_datasheet_transition(
                entry, 'Approved',
                comment_text or 'Datasheet approved during peer review.')
            return True, 'Peer review approved successfully', 200

        if normalized_action == 'reject':
            if not comment_text:
                return False, 'Comments are required when rejecting a datasheet', 400
            action_label = 'REJECTED'
            entry.status = 'in_progress'
            _append_datasheet_peer_review_comment(
                entry,
                comment_text,
                current_user.username,
                action_label
            )
            _update_parent_request_datasheet_status(entry)
            _record_datasheet_transition(entry, 'Rejected', comment_text)
            return True, 'Peer review rejected and sent back to engineer', 200

        _append_datasheet_peer_review_comment(
            entry,
            comment_text,
            current_user.username,
            'COMMENT'
        )
        _record_datasheet_transition(entry, 'Peer Review', comment_text)
        return True, 'Peer review comment added successfully', 200

    @flask_app.route('/api/planner/<int:planner_id>/peer-review-action', methods=['POST'])
    @login_required
    def peer_review_action(planner_id):
        """Apply a peer-review action to a planner entry."""
        logger.info(f'Applying peer review action for planner entry {planner_id}')

        try:
            if current_user.role not in ['admin', 'lab_engineer']:
                return jsonify({'success': False, 'error': 'Not authorized'}), 403

            entry = db.session.get(PlannerEntry, planner_id)
            if not entry:
                return jsonify({'success': False, 'error': 'Planner entry not found'}), 404

            payload = request.get_json(silent=True) or {}
            action = payload.get('action')
            comment = payload.get('comment', '')

            success, message, status_code = _apply_peer_review_action(
                entry,
                action,
                comment
            )
            if not success:
                db.session.rollback()
                return jsonify({'success': False, 'error': message}), status_code

            entry.updated_at = get_ist_now()
            db.session.commit()

            return jsonify({
                'success': True,
                'message': message,
                'data': {
                    'id': entry.id,
                    'status': entry.status,
                    'test_name': entry.test_name,
                    'test_request_id': entry.test_request_id,
                    'datasheet_comments': entry.datasheet_comments,
                }
            }), status_code

        except Exception as e:
            db.session.rollback()
            logger.error(
                f'Error applying peer review action for entry {planner_id}: {e}')
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @flask_app.route('/api/planner/<int:planner_id>/peer-review-approve', methods=['POST'])
    @login_required
    def approve_peer_review(planner_id):
        """Approve a planner entry's datasheet and mark it as uploaded."""
        logger.info(f'Approving peer review for planner entry {planner_id}')

        try:
            # Check permission
            if current_user.role not in ['admin', 'lab_engineer']:
                return jsonify({'success': False, 'error': 'Not authorized'}), 403

            # Fetch the planner entry
            entry = db.session.get(PlannerEntry, planner_id)
            if not entry:
                return jsonify({'success': False, 'error': 'Planner entry not found'}), 404

            success, message, status_code = _apply_peer_review_action(
                entry,
                'approve',
                ''
            )
            if not success:
                db.session.rollback()
                return jsonify({
                    'success': False,
                    'error': message
                }), status_code

            entry.updated_at = get_ist_now()

            db.session.commit()

            logger.info(
                f'Peer review approved for planner entry {planner_id} '
                f'by user {current_user.id} ({current_user.username})'
            )

            return jsonify({
                'success': True,
                'message': message,
                'data': {
                    'id': entry.id,
                    'status': entry.status,
                    'test_name': entry.test_name,
                    'test_request_id': entry.test_request_id
                }
            })

        except Exception as e:
            db.session.rollback()
            logger.error(
                f'Error approving peer review for entry {planner_id}: {e}')
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @flask_app.route('/api/planner/peer-review', methods=['GET'])
    @login_required
    def get_peer_review_entries():
        """Get planner entries currently waiting for peer review."""
        logger.info('Fetching peer review planner entries')
        try:
            if current_user.role not in ['admin', 'lab_engineer']:
                return jsonify({'success': False, 'error': 'Not authorized'}), 403

            # Query with explicit join to get correct TCO ID from test_request
            query = db.session.query(
                PlannerEntry,
                EMCRequest.tco_id.label('actual_tco_id'),
                EMCRequest.job_number.label('job_number'),
                EMCRequest.product_name,
                EMCRequest.requester_name,
                EMCRequest.id.label('actual_request_id'),
                EMCRequest.status.label('parent_status')
            ).join(
                EMCRequest,
                PlannerEntry.test_request_id == EMCRequest.id
            ).filter(
                PlannerEntry.status == 'Peer Review'
            ).order_by(
                EMCRequest.tco_id,
                PlannerEntry.test_name,
                PlannerEntry.completion_date.desc()
            )

            if current_user.role == 'lab_engineer':
                query = query.filter(
                    PlannerEntry.peer_reviewer_user_id == current_user.id
                )

            all_results = query.all()

            # Entries whose PARENT request is finished sort to the end of the list.
            results = _terminal_last(all_results, status_of=lambda row: row.parent_status)
            _terminal_count = len(all_results) - len(
                [r for r in all_results if not _is_terminal_request_status(r.parent_status)])

            logger.info(
                f'Found {len(results)} peer review entries '
                f'(active: {len(results) - _terminal_count}, finished: {_terminal_count})')

            # Serialize the entries
            result = []
            for entry, actual_tco_id, job_number, product_name, requester_name, actual_request_id, parent_status in results:
                # Log for debugging
                logger.info(f'Peer review entry: id={entry.id}, '
                            f'planner_tco={entry.tco_id}, '
                            f'actual_tco={actual_tco_id}, '
                            f'test={entry.test_name}, '
                            f'engineer={entry.test_person_name}')

                # Get uploaded_by user name if available
                uploaded_by_name = None
                if entry.datasheet_uploaded_by:
                    uploaded_by_user = db.session.get(User, 
                        entry.datasheet_uploaded_by)
                    if uploaded_by_user:
                        uploaded_by_name = uploaded_by_user.username

                peer_reviewer_name = None
                if entry.peer_reviewer_user_id:
                    peer_reviewer_user = db.session.get(User, entry.peer_reviewer_user_id)
                    if peer_reviewer_user:
                        peer_reviewer_name = peer_reviewer_user.username

                datasheet_record = None
                datasheet_record_view_url = None
                try:
                    from datasheet_gen import records as datasheet_records
                    datasheet_record = datasheet_records.get_record_for_assignment(entry.id)
                    if datasheet_record:
                        datasheet_record_view_url = url_for(
                            'datasheet_records.record_detail',
                            record_id=datasheet_record['id']
                        )
                except Exception as exc:
                    logger.warning(
                        'Unable to load saved datasheet record for peer review entry %s: %s',
                        entry.id,
                        exc
                    )

                result.append({
                    'id': entry.id,
                    'test_name': entry.test_name,
                    'tco_id': actual_tco_id,  # USE THE ACTUAL TCO ID FROM TEST REQUEST
                    'job_number': job_number or '',
                    'planner_tco_id': entry.tco_id,  # Include for debugging
                    'test_request_id': entry.test_request_id,
                    'actual_request_id': actual_request_id,
                    'product_name': product_name,
                    'requester_name': requester_name,
                    'test_person_name': entry.test_person_name,
                    'engineer_user_id': entry.engineer_user_id,
                    # parent request status - the peer-review table sorts finished TCOs last
                    'parent_status': parent_status or '',
                    'start_date': entry.start_date.isoformat() if entry.start_date else None,
                    'end_date': entry.end_date.isoformat() if entry.end_date else None,
                    'total_hours': entry.total_hours,
                    'status': entry.status,
                    'datasheet_file_path': entry.datasheet_file_path,
                    'datasheet_uploaded_at': entry.datasheet_uploaded_at.isoformat() if entry.datasheet_uploaded_at else None,
                    'datasheet_uploaded_by': entry.datasheet_uploaded_by,
                    'peer_reviewer_user_id': entry.peer_reviewer_user_id,
                    'peer_review_assigned_at': entry.peer_review_assigned_at.isoformat() if entry.peer_review_assigned_at else None,
                    'peer_reviewer_name': peer_reviewer_name,
                    'datasheet_comments': entry.datasheet_comments,
                    'completion_date': entry.completion_date.isoformat() if entry.completion_date else None,
                    'uploaded_by_name': uploaded_by_name,
                    'has_datasheet_record': bool(datasheet_record),
                    'datasheet_record_id': datasheet_record.get('id') if datasheet_record else None,
                    'datasheet_record_view_url': datasheet_record_view_url
                })

            return jsonify({'success': True, 'data': result})

        except Exception as e:
            logger.error(f'Error fetching peer review entries: {e}')
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @flask_app.route('/api/files/download', methods=['GET'])
    @login_required
    def download_file_by_path():
        """Download a file by path"""
        try:
            file_path = request.args.get('path', '')

            if not file_path:
                return jsonify({'error': 'No file path provided'}), 400

            # Security check - prevent directory traversal
            if '..' in file_path:
                return jsonify({'error': 'Invalid file path'}), 400

            # Normalize the path (handles both / and \ separators)
            file_path = os.path.normpath(file_path)

            # ---- REPAIR BROKEN PATHS FROM OLD DB RECORDS ----
            # Fixes missing separator e.g. '...test_datasheetsIEC-EMC-019_...'
            # corrects it to  '...test_datasheets\IEC-EMC-019_...'
            FOLDER_MARKER = 'test_datasheets'
            if FOLDER_MARKER in file_path:
                idx = file_path.find(FOLDER_MARKER)
                after_marker = file_path[idx + len(FOLDER_MARKER):]
                if after_marker and after_marker[0] not in (os.sep, '/', '\\'):
                    file_path = file_path[:idx +
                                          len(FOLDER_MARKER)] + os.sep + after_marker
                    file_path = os.path.normpath(file_path)
                    logger.info(f'Repaired broken path to: {file_path}')
            # ---- END REPAIR ----

            # If path is already absolute, use it directly
            # If relative, make it absolute relative to app root
            if os.path.isabs(file_path):
                absolute_path = file_path
            else:
                absolute_path = os.path.normpath(
                    os.path.join(flask_app.root_path, file_path)
                )

            logger.info(f'Attempting to download file: {absolute_path}')

            if not os.path.exists(absolute_path):
                logger.error(f'File not found: {absolute_path}')
                return jsonify({'error': 'File not found'}), 404

            # Security - verify file is within the uploads directory
            if os.path.isabs(UPLOAD_FOLDER):
                allowed_dir = os.path.normpath(UPLOAD_FOLDER)
            else:
                allowed_dir = os.path.normpath(
                    os.path.join(flask_app.root_path, UPLOAD_FOLDER)
                )

            # Also allow parent 'uploads' folder for report files
            uploads_dir = os.path.normpath(
                os.path.join(flask_app.root_path, 'uploads')
            )

            if not (absolute_path.startswith(allowed_dir) or
                    absolute_path.startswith(uploads_dir)):
                logger.warning(
                    f'Attempted access outside uploads dir: {absolute_path}'
                )
                return jsonify({'error': 'Access denied'}), 403

            directory = os.path.dirname(absolute_path)
            filename = os.path.basename(absolute_path)

            logger.info(f'Sending file: {filename} from {directory}')

            # PDFs and images open in browser, other files trigger download
            extension = filename.rsplit(
                '.', 1)[-1].lower() if '.' in filename else ''
            as_attachment = extension not in (
                'pdf', 'png', 'jpg', 'jpeg', 'gif', 'webp')

            return send_from_directory(
                directory,
                filename,
                as_attachment=as_attachment,
                download_name=filename
            )

        except Exception as e:
            logger.error(f'Error downloading file: {e}')
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'error': str(e)}), 500

    @flask_app.route('/api/planner/<int:planner_id>', methods=['GET'])
    @login_required
    def get_planner_entry(planner_id):
        """Get detailed information for a single planner entry"""
        logger.info(f'Fetching planner entry {planner_id}')
        try:
            entry = db.session.get(PlannerEntry, planner_id)
            if not entry:
                return jsonify({'success': False, 'error': 'Planner entry not found'}), 404

            # Get uploaded_by user name if available
            uploaded_by_name = None
            if entry.datasheet_uploaded_by:
                uploaded_by_user = db.session.get(User, entry.datasheet_uploaded_by)
                if uploaded_by_user:
                    uploaded_by_name = uploaded_by_user.username

            peer_reviewer_name = None
            if entry.peer_reviewer_user_id:
                peer_reviewer_user = db.session.get(User, entry.peer_reviewer_user_id)
                if peer_reviewer_user:
                    peer_reviewer_name = peer_reviewer_user.username

            datasheet_record = None
            datasheet_record_view_url = None
            try:
                from datasheet_gen import records as datasheet_records
                datasheet_record = datasheet_records.get_record_for_assignment(entry.id)
                if datasheet_record:
                    datasheet_record_view_url = url_for(
                        'datasheet_records.record_detail',
                        record_id=datasheet_record['id']
                    )
            except Exception as exc:
                logger.warning(
                    'Unable to load saved datasheet record for planner entry %s: %s',
                    entry.id,
                    exc
                )

            result = {
                'id': entry.id,
                'test_name': entry.test_name,
                'tco_id': entry.tco_id,
                'test_person_name': entry.test_person_name,
                'engineer_user_id': entry.engineer_user_id,
                'start_date': entry.start_date.isoformat() if entry.start_date else None,
                'end_date': entry.end_date.isoformat() if entry.end_date else None,
                'start_time': str(entry.start_time) if entry.start_time else None,
                'end_time': str(entry.end_time) if entry.end_time else None,
                'total_hours': entry.total_hours,
                'status': entry.status,
                'datasheet_file_path': entry.datasheet_file_path,
                'datasheet_uploaded_at': entry.datasheet_uploaded_at.isoformat() if entry.datasheet_uploaded_at else None,
                'datasheet_uploaded_by': entry.datasheet_uploaded_by,
                'peer_reviewer_user_id': entry.peer_reviewer_user_id,
                'peer_review_assigned_at': entry.peer_review_assigned_at.isoformat() if entry.peer_review_assigned_at else None,
                'peer_reviewer_name': peer_reviewer_name,
                'datasheet_comments': entry.datasheet_comments,
                'completion_date': entry.completion_date.isoformat() if entry.completion_date else None,
                'uploaded_by_name': uploaded_by_name,
                'has_datasheet_record': bool(datasheet_record),
                'datasheet_record_id': datasheet_record.get('id') if datasheet_record else None,
                'datasheet_record_view_url': datasheet_record_view_url,
                'event_description': entry.event_description
            }

            return jsonify({'success': True, 'data': result})

        except Exception as e:
            logger.error(f'Error fetching planner entry: {e}')
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @flask_app.route('/api/test-requests/<int:request_id>', methods=['GET', 'DELETE'])
    @login_required
    def get_test_request(request_id):
        """Get or delete a single test request by ID."""
        try:
            # Get the test request
            test_request = _get_request_or_404(request_id)

            if not test_request or not _can_access_iec_request(
                test_request,
                allow_lab_engineer=True,
                require_assigned_lab_engineer=True
            ):
                return jsonify({
                    'success': False,
                    'error': 'Test request not found or you do not have permission to view it'
                }), 404

            if request.method == 'DELETE':
                # Only draft requests can be deleted from index actions.
                status_normalized = (test_request.status or '').strip().lower()
                if status_normalized != 'draft':
                    return jsonify({
                        'success': False,
                        'error': 'Only Draft requests can be deleted.'
                    }), 400

                # Drafts are private to the owner.
                if test_request.user_id != current_user.id:
                    return jsonify({
                        'success': False,
                        'error': 'Access denied. You can only delete your own draft requests.'
                    }), 403

                tco_display = test_request.tco_id or f'REQ-{test_request.id}'

                # Remove planner rows linked by request id or matching TCO.
                PlannerEntry.query.filter(
                    db.or_(
                        PlannerEntry.test_request_id == test_request.id,
                        PlannerEntry.tco_id == tco_display
                    )
                ).delete(synchronize_session=False)

                db.session.delete(test_request)
                db.session.commit()

                logger.info(
                    'Deleted draft test request id=%s, tco=%s by user=%s',
                    test_request.id,
                    tco_display,
                    current_user.username
                )

                return jsonify({
                    'success': True,
                    'message': 'Draft test request deleted successfully.'
                })

            request_data = _build_request_payload(test_request)

            # DEBUG: Log CE fields to verify data
            logger.info(f"CE Data for request {request_id}:")
            logger.info(f"  ce_standard: {request_data.get('ce_standard')}")
            logger.info(
                f"  ce_voltage_freq: {request_data.get('ce_voltage_freq')}")
            logger.info(
                f"  ce_freq_range: {request_data.get('ce_freq_range')}")
            logger.info(f"  ce_class: {request_data.get('ce_class')}")
            logger.info(
                f"  selected_tests: {request_data.get('selected_tests')}")

            return jsonify({
                'success': True,
                'data': request_data
            })
        except Exception as e:
            logger.error(f'Error fetching test request: {e}')
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @flask_app.route('/api/test-requests/<int:request_id>/cancel-unassigned-test', methods=['POST'])
    @login_required
    def cancel_unassigned_test(request_id):
        """Cancel a selected test before it has been assigned to the planner."""
        try:
            if current_user.role not in ['admin', 'lab_engineer']:
                return jsonify({
                    'success': False,
                    'error': 'You do not have permission to cancel tests.'
                }), 403

            test_request = _get_request_or_404(request_id)
            if not test_request or not _can_access_iec_request(
                test_request,
                allow_lab_engineer=True,
                require_assigned_lab_engineer=True
            ):
                return jsonify({
                    'success': False,
                    'error': 'Test request not found or you do not have permission to edit it.'
                }), 404

            data = request.get_json() or {}
            test_name = str(data.get('test_name') or '').strip()
            cancel_reason = str(data.get('cancel_reason') or '').strip()

            if not test_name:
                return jsonify({
                    'success': False,
                    'error': 'Test name is required.'
                }), 400

            if not cancel_reason:
                return jsonify({
                    'success': False,
                    'error': 'Cancellation reason is required.'
                }), 400

            requested_key = _normalize_assignment_test_key(test_name)
            request_test = next((
                row for row in (getattr(test_request, 'tests', []) or [])
                if getattr(row, 'is_selected', False)
                and _normalize_assignment_test_key(
                    EMCRequest._legacy_test_key(getattr(row, 'test_code', None))
                ) == requested_key
            ), None)

            if not request_test:
                return jsonify({
                    'success': False,
                    'error': 'Selected test was not found on this TCO.'
                }), 404

            if str(getattr(request_test, 'workflow_status', '') or '').strip().lower() == 'cancelled':
                return jsonify({
                    'success': False,
                    'error': 'This test is already cancelled.'
                }), 400

            has_request_level_assignment = any([
                getattr(request_test, 'assigned_engineer_id', None),
                str(getattr(request_test, 'assigned_engineer_name', '') or '').strip(),
                getattr(request_test, 'planned_start_date', None),
                getattr(request_test, 'planned_end_date', None),
            ])

            planner_filters = [PlannerEntry.test_request_id == test_request.id]
            if test_request.tco_id:
                planner_filters.append(PlannerEntry.tco_id == test_request.tco_id)

            matching_planner_entry = None
            for planner_entry in PlannerEntry.query.filter(db.or_(*planner_filters)).all():
                planner_key = _normalize_assignment_test_key(getattr(planner_entry, 'test_name', None))
                if planner_key == requested_key and str(getattr(planner_entry, 'status', '') or '').strip().lower() != 'cancelled':
                    matching_planner_entry = planner_entry
                    break

            if has_request_level_assignment or matching_planner_entry:
                return jsonify({
                    'success': False,
                    'error': 'This test is already assigned. Cancel it from the Assigned Tests page.'
                }), 400

            request_test.workflow_status = 'cancelled'
            request_test.assigned_engineer_id = None
            request_test.assigned_engineer_name = None
            request_test.planned_start_date = None
            request_test.planned_end_date = None
            test_request.updated_at = get_ist_now()

            _append_review_comment_entry(
                test_request=test_request,
                comment=f'[Test Cancelled: {test_name}] {cancel_reason}',
                username=current_user.username,
                role=current_user.role
            )

            db.session.commit()

            return jsonify({
                'success': True,
                'message': 'Test cancelled successfully before assignment.',
                'data': {
                    'test_name': test_name,
                    'remaining_tests': _get_active_selected_test_labels(test_request)
                }
            })
        except Exception as exc:
            db.session.rollback()
            logger.error('Error cancelling unassigned test for request %s: %s', request_id, exc)
            return jsonify({
                'success': False,
                'error': 'Failed to cancel the selected test.'
            }), 500

    @flask_app.route('/api/test-requests/<int:request_id>/job-number', methods=['PATCH'])
    @login_required
    def update_test_request_job_number(request_id):
        """Update the manual job number for a request from review/assigned views."""
        try:
            if current_user.role not in ['admin', 'lab_engineer']:
                return jsonify({
                    'success': False,
                    'error': 'You do not have permission to edit the job number.'
                }), 403

            test_request = _get_request_or_404(request_id)
            if not test_request or not _can_access_iec_request(
                test_request,
                allow_lab_engineer=True,
                require_assigned_lab_engineer=True
            ):
                return jsonify({
                    'success': False,
                    'error': 'Test request not found or you do not have permission to edit it.'
                }), 404

            data = request.get_json() or {}
            job_number = str(data.get('job_number') or '').strip()
            if len(job_number) > 100:
                return jsonify({
                    'success': False,
                    'error': 'Job Number must be 100 characters or fewer.'
                }), 400

            test_request.job_number = job_number or None
            test_request.updated_at = get_ist_now()
            db.session.commit()

            logger.info(
                'Job Number updated for request %s by %s to "%s"',
                request_id,
                current_user.username,
                test_request.job_number or ''
            )

            return jsonify({
                'success': True,
                'message': 'Job Number updated successfully.',
                'data': {
                    'request_id': test_request.id,
                    'job_number': test_request.job_number or '',
                    'tco_id': test_request.tco_id or f'REQ-{test_request.id}'
                }
            })
        except Exception as e:
            db.session.rollback()
            logger.error('Error updating Job Number for request %s: %s', request_id, e)
            return jsonify({
                'success': False,
                'error': 'Error updating Job Number.'
            }), 500

    def _export_field_label(field_name: str) -> str:
        """Convert snake_case field names into readable labels."""
        return (field_name or '').replace('_', ' ').strip().title()

    def _looks_like_embedded_binary(value: str) -> bool:
        """Detect large/base64-like strings so Word export stays readable."""
        if not isinstance(value, str):
            return False
        trimmed = value.strip()
        if not trimmed:
            return False
        if trimmed.startswith('data:'):
            return True
        if len(trimmed) < 300:
            return False
        sample = trimmed[:1200]
        return bool(re.fullmatch(r'[A-Za-z0-9+/=\r\n]+', sample))

    def _format_export_value(value) -> str:
        """Normalize request values into display text for the DOCX export."""
        if value is None:
            return 'N/A'
        if isinstance(value, bool):
            return 'Yes' if value else 'No'
        if isinstance(value, (int, float)):
            return str(value)

        if isinstance(value, str):
            text_value = value.strip()
            if not text_value:
                return 'N/A'
            if _looks_like_embedded_binary(text_value):
                return '[Attachment content stored in application]'
            if (
                (text_value.startswith('{') and text_value.endswith('}'))
                or (text_value.startswith('[') and text_value.endswith(']'))
            ):
                try:
                    return _format_export_value(json.loads(text_value))
                except (json.JSONDecodeError, TypeError, ValueError):
                    pass
            if len(text_value) > 5000:
                return text_value[:5000] + '\n...[truncated]'
            return text_value

        if isinstance(value, list):
            if not value:
                return 'N/A'

            if all(not isinstance(item, (dict, list)) for item in value):
                normalized = [
                    _format_export_value(item)
                    for item in value
                    if _format_export_value(item) != 'N/A'
                ]
                return ', '.join(normalized) if normalized else 'N/A'

            lines = []
            for index, item in enumerate(value, start=1):
                if isinstance(item, dict):
                    parts = []
                    for key, inner_value in item.items():
                        formatted_inner = _format_export_value(inner_value)
                        if formatted_inner and formatted_inner != 'N/A':
                            parts.append(
                                f"{_export_field_label(str(key))}: {formatted_inner}")
                    lines.append(
                        f"{index}. " + ('; '.join(parts) if parts else 'N/A')
                    )
                else:
                    lines.append(f"{index}. {_format_export_value(item)}")
            return '\n'.join(lines) if lines else 'N/A'

        if isinstance(value, dict):
            if not value:
                return 'N/A'
            lines = []
            for key, inner_value in value.items():
                formatted_inner = _format_export_value(inner_value)
                if formatted_inner and formatted_inner != 'N/A':
                    lines.append(
                        f"{_export_field_label(str(key))}: {formatted_inner}")
            return '\n'.join(lines) if lines else 'N/A'

        return str(value)

    def _add_export_section(document, section_title: str, rows: list) -> None:
        """Append a section heading and a key/value table to the DOCX."""
        document.add_heading(section_title, level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Field'
        table.rows[0].cells[1].text = 'Value'

        for label, raw_value in rows:
            formatted = _format_export_value(raw_value)
            if formatted == 'N/A':
                continue
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = formatted

    def _coerce_export_text(value) -> str:
        """Convert any export value to safe plain text (empty when unavailable)."""
        formatted = _format_export_value(value)
        return '' if formatted == 'N/A' else formatted

    def _coerce_export_list(value) -> list:
        """Convert value to a list when possible."""
        if isinstance(value, list):
            return value
        if isinstance(value, str):
            trimmed = value.strip()
            if not trimmed:
                return []
            try:
                parsed = json.loads(trimmed)
                return parsed if isinstance(parsed, list) else []
            except (json.JSONDecodeError, TypeError, ValueError):
                return []
        return []

    def _parse_export_dict(value) -> dict:
        """Convert export values into a dictionary when possible."""
        if isinstance(value, dict):
            return value
        if isinstance(value, str):
            trimmed = value.strip()
            if not trimmed:
                return {}
            try:
                parsed = json.loads(trimmed)
                return parsed if isinstance(parsed, dict) else {}
            except (json.JSONDecodeError, TypeError, ValueError):
                return {}
        return {}

    def _extract_custom_range_text(spec_obj) -> str:
        """Build a readable custom frequency-range string from nested spec objects."""
        if not isinstance(spec_obj, dict) or not spec_obj:
            return ''

        custom_range = (
            spec_obj.get('customSpecRange')
            or spec_obj.get('custom_spec_range')
            or spec_obj.get('customFrequencyRange')
            or spec_obj.get('custom_frequency_range')
            or spec_obj.get('customRange')
            or spec_obj.get('custom_range')
        )

        if not custom_range:
            return ''

        if not isinstance(custom_range, dict):
            return _coerce_export_text(custom_range)

        normalized = custom_range
        for nested_key in ('frequency', 'frequency_range', 'frequencyRange', 'range'):
            nested_value = custom_range.get(nested_key)
            if isinstance(nested_value, dict):
                normalized = nested_value
                break

        from_value = _coerce_export_text(
            normalized.get('from') if isinstance(normalized, dict) else ''
        )
        from_unit = _coerce_export_text(
            normalized.get('fromUnit') if isinstance(normalized, dict) else normalized.get('from_unit')
        )
        to_value = _coerce_export_text(
            normalized.get('to') if isinstance(normalized, dict) else ''
        )
        to_unit = _coerce_export_text(
            normalized.get('toUnit') if isinstance(normalized, dict) else normalized.get('to_unit')
        )

        if from_value and to_value:
            return ' '.join(
                part for part in [from_value, from_unit, 'to', to_value, to_unit] if part
            )
        if from_value:
            return ' '.join(part for part in [from_value, from_unit] if part)
        if to_value:
            return ' '.join(part for part in [to_value, to_unit] if part)
        return ''

    def _extract_frequency_range_display(request_data: dict, *keys: str) -> str:
        """Resolve frequency-range text from nested custom-spec structures or flat fallback keys."""
        if not isinstance(request_data, dict):
            return ''

        for key in keys:
            raw_value = request_data.get(key)
            parsed_spec = _parse_export_dict(raw_value)
            if parsed_spec:
                for selector_key in ('spec', 'frequency', 'frequencySpec', 'frequency_spec'):
                    selector_value = _coerce_export_text(parsed_spec.get(selector_key))
                    if not selector_value:
                        continue
                    if 'custom' in selector_value.lower():
                        custom_range_text = _extract_custom_range_text(parsed_spec)
                        if custom_range_text:
                            return custom_range_text
                    else:
                        return selector_value

                custom_range_text = _extract_custom_range_text(parsed_spec)
                if custom_range_text:
                    return custom_range_text
            else:
                text_value = _coerce_export_text(raw_value)
                if text_value:
                    return text_value

        return ''

    def _select_truthy_option_labels(value) -> str:
        """Render selected keys from dictionary-style option groups."""
        if not isinstance(value, dict):
            return _coerce_export_text(value)

        selected = []
        for key, state in value.items():
            normalized = str(state).strip().lower()
            if state is True or normalized in ('true', 'yes', '1', 'on', 'selected'):
                selected.append(str(key))
        return ', '.join(selected)

    def _pick_item_value(item: dict, *keys) -> str:
        """Pick the first non-empty value from a dict using alternate keys."""
        if not isinstance(item, dict):
            return ''
        for key in keys:
            if key in item and item.get(key) not in (None, ''):
                return _coerce_export_text(item.get(key))
        return ''

    def _iso_date_text(value) -> str:
        """Normalize a date-like value into YYYY-MM-DD text."""
        if value in (None, ''):
            return ''
        if isinstance(value, datetime):
            return value.date().isoformat()
        if isinstance(value, date):
            return value.isoformat()
        text_value = str(value).strip()
        if not text_value:
            return ''
        if 'T' in text_value:
            return text_value.split('T', 1)[0]
        return text_value[:10] if len(text_value) >= 10 and text_value[4:5] == '-' else text_value

    def _format_duration_text(value) -> str:
        """Normalize duration values for consistent display/export."""
        if value in (None, ''):
            return ''
        try:
            numeric_value = round(float(value), 2)
        except (TypeError, ValueError):
            return str(value).strip()
        if numeric_value <= 0:
            return ''
        if float(numeric_value).is_integer():
            return str(int(numeric_value))
        return f'{numeric_value:.2f}'.rstrip('0').rstrip('.')

    def _format_signature_export_value(value, fallback_text: str = '[Signature in system]') -> str:
        """Render typed signatures as text and image signatures as a placeholder."""
        text_value = _coerce_export_text(value).strip()
        if not text_value:
            return ''
        if text_value.startswith('data:image/') or text_value.startswith('http'):
            return fallback_text
        if len(text_value) > 100 and re.fullmatch(r'[A-Za-z0-9+/=\s]+', text_value):
            return fallback_text
        return text_value

    def _get_equipment_for_test_keyword(test_keyword: str) -> list[dict]:
        """Return equipment rows whose test_name contains the provided keyword."""
        normalized_keyword = (test_keyword or '').strip()
        if not normalized_keyword:
            return []

        try:
            equipment_rows = Equipment.query.filter(
                Equipment.test_name.isnot(None),
                Equipment.test_name.ilike(f'%{normalized_keyword}%')
            ).order_by(
                Equipment.sl_no.asc(),
                Equipment.name.asc()
            ).all()
            return [row.to_dict() for row in equipment_rows]
        except Exception as exc:
            logger.error(
                'Error fetching equipment for test keyword %s: %s',
                normalized_keyword,
                exc
            )
            return []

    def _get_software_defaults_for_test_keyword(test_keyword: str) -> list[dict]:
        """Return datasheet software rows for the provided test keyword."""
        normalized_keyword = (test_keyword or '').strip()
        if not normalized_keyword:
            return []

        try:
            software_rows = Equipment.query.filter(
                Equipment.type.isnot(None),
                db.func.lower(Equipment.type).in_(['software', 'application', 'tool']),
                Equipment.test_name.isnot(None),
                Equipment.test_name.ilike(f'%{normalized_keyword}%')
            ).order_by(
                Equipment.sl_no.asc(),
                Equipment.name.asc()
            ).all()

            normalized_rows = []
            for row in software_rows:
                used = str(row.name or row.make or '').strip()
                version = str(row.model_no or row.serial_no or '').strip()
                if used or version:
                    normalized_rows.append({
                        'software_used': used,
                        'software_version': version
                    })

            if normalized_rows:
                return normalized_rows
        except Exception as exc:
            logger.error(
                'Error fetching software defaults for test keyword %s: %s',
                normalized_keyword,
                exc
            )

        fallback_defaults = {
            'surge': [
                {
                    'software_used': 'IEC',
                    'software_version': '10.3.2'
                }
            ]
        }
        return fallback_defaults.get(normalized_keyword.lower(), [])

    def _build_request_payload(test_request, include_review_thread: bool = True) -> dict:
        """Build a single enriched request payload for view/download flows."""
        request_data = test_request.to_dict()

        planner_entries = PlannerEntry.query.filter(
            db.or_(
                PlannerEntry.test_request_id == test_request.id,
                PlannerEntry.tco_id == (test_request.tco_id or '')
            )
        ).order_by(
            PlannerEntry.start_date.asc(),
            PlannerEntry.start_time.asc()
        ).all()
        request_data['planner_entries'] = [entry.to_dict() for entry in planner_entries]

        if include_review_thread:
            review_thread = _get_combined_review_comment_thread(test_request)
            request_data['review_comment_thread'] = review_thread
            if review_thread:
                latest_comment = review_thread[-1]
                request_data['review_comments'] = latest_comment.get('comment')
                request_data['reviewed_by'] = latest_comment.get('username') or request_data.get('reviewed_by')
                request_data['reviewed_at'] = latest_comment.get('created_at') or request_data.get('reviewed_at')

        earliest_start_date = min(
            (entry.start_date for entry in planner_entries if entry.start_date),
            default=None
        )
        latest_end_date = max(
            (entry.end_date for entry in planner_entries if entry.end_date),
            default=None
        )
        total_planner_hours = round(sum(
            float(entry.total_hours or 0) for entry in planner_entries
            if entry.total_hours not in (None, '')
        ), 2)

        canonical_job_number = (
            request_data.get('job_number')
            or getattr(test_request, 'job_number', None)
            or ''
        )
        request_data['job_number'] = canonical_job_number
        if not request_data.get('sample_received_date'):
            request_received = getattr(test_request, 'submitted_at', None) or getattr(test_request, 'created_at', None)
            request_data['sample_received_date'] = _iso_date_text(request_received)
        if not request_data.get('requester_signature') and request_data.get('submitted_at'):
            request_data['requester_signature'] = (
                request_data.get('requester_name')
                or getattr(test_request, 'requester_name', None)
                or ''
            )
        if not request_data.get('test_commencement_date') and earliest_start_date:
            request_data['test_commencement_date'] = _iso_date_text(earliest_start_date)
        if not request_data.get('test_completion_date') and latest_end_date:
            request_data['test_completion_date'] = _iso_date_text(latest_end_date)
        if not request_data.get('test_duration') and total_planner_hours > 0:
            request_data['test_duration'] = _format_duration_text(total_planner_hours)
        if not request_data.get('lab_manager_date') and getattr(test_request, 'lab_manager_signed_at', None):
            request_data['lab_manager_date'] = _iso_date_text(test_request.lab_manager_signed_at)

        request_data['selected_tests'] = _get_active_selected_test_labels(test_request)

        selected_test_shortcuts = _extract_selected_test_shortcuts(
            request_data.get('selected_tests')
        )
        planner_test_names = [
            str(getattr(entry, 'test_name', '') or '').strip().lower()
            for entry in planner_entries
        ]
        surge_is_applicable = (
            'SURGE' in selected_test_shortcuts
            or any('surge' in name for name in planner_test_names)
        )
        request_data['surge_equipment_master'] = (
            _get_equipment_for_test_keyword('Surge')
            if surge_is_applicable else []
        )
        request_data['surge_software_master'] = (
            _get_software_defaults_for_test_keyword('Surge')
            if surge_is_applicable else []
        )

        # Calculate remaining_tests: selected tests that don't have assignments yet
        selected_tests = list(request_data.get('selected_tests') or [])

        assigned_test_keys = set()
        for entry in planner_entries:
            # A CANCELLED run is not an assignment.
            #
            # The review page has always skipped these (see the
            # active_entries_by_test filter), which is why it shows such a test as
            # "Not Assigned". This payload did not, so the same test was
            # simultaneously unassigned on the page and assigned to the dialog
            # that offers unassigned tests - and the dialog, having nothing left
            # to offer, said everything was already assigned.
            #
            # Measured on IEC-EMC-004: EFT is selected on the request, its only
            # planner entry (id 10) is 'cancelled', and it could not be
            # re-assigned from the UI at all.
            if str(getattr(entry, 'status', '') or '').strip().lower() == 'cancelled':
                continue
            test_name = getattr(entry, 'test_name', None)
            if test_name:
                assigned_key = _normalize_assignment_test_key(test_name)
                if assigned_key:
                    assigned_test_keys.add(assigned_key)

        remaining_tests = []
        for selected_test in selected_tests:
            selected_key = _normalize_assignment_test_key(selected_test)
            if selected_key and selected_key in assigned_test_keys:
                continue
            remaining_tests.append(selected_test)

        request_data['remaining_tests'] = remaining_tests

        return request_data

    def _selected_option_set(value) -> set:
        """Normalize selection input into a lowercase set of selected labels."""
        selected = set()

        if isinstance(value, list):
            for item in value:
                text = str(item).strip().lower()
                if text:
                    selected.add(text)
            return selected

        if isinstance(value, dict):
            for key, state in value.items():
                normalized = str(state).strip().lower()
                if state is True or normalized in ('true', 'yes', '1', 'on', 'selected'):
                    text = str(key).strip().lower()
                    if text:
                        selected.add(text)
                    continue

                if isinstance(state, str):
                    if normalized and normalized not in ('false', 'no', '0', 'off', 'none', 'null'):
                        selected.add(normalized)
                    continue

                if isinstance(state, (list, dict)):
                    selected.update(_selected_option_set(state))
            return selected

        if isinstance(value, str):
            text_value = value.strip()
            if not text_value:
                return selected
            try:
                parsed = json.loads(text_value)
                return _selected_option_set(parsed)
            except (json.JSONDecodeError, TypeError, ValueError):
                pass

            for token in re.split(r'[,/|]', text_value):
                token_clean = token.strip().lower()
                if token_clean:
                    selected.add(token_clean)
            return selected

        return selected

    def _selection_contains(selected_set: set, *needles: str) -> bool:
        """Return True when any selected option contains one of the provided needles."""
        for item in selected_set:
            normalized_item = str(item).strip().lower()
            for needle in needles:
                if needle and needle.lower() in normalized_item:
                    return True
        return False

    def _checkbox_options_text(options: list, selected_set: set) -> str:
        """Render checkbox options text with selected values checked."""
        checked = '\u2611'
        unchecked = '\u2610'
        parts = []
        for label, *aliases in options:
            alias_pool = [label] + list(aliases)
            is_selected = _selection_contains(selected_set, *alias_pool)
            parts.append(f"{checked if is_selected else unchecked} {label}")
        return '    '.join(parts)

    def _get_export_empty_placeholder(source) -> str:
        """Return the export placeholder configured for a document instance."""
        return getattr(source, '_empty_export_placeholder', '')

    def _set_table_cell_text(doc, table_idx: int, row_idx: int, col_idx: int, value) -> None:
        """Safely set a docx table cell's text."""
        try:
            table = doc.tables[table_idx]
            if row_idx >= len(table.rows):
                return
            row = table.rows[row_idx]
            if col_idx >= len(row.cells):
                return
            text_value = _coerce_export_text(value)
            row.cells[col_idx].text = text_value if text_value else _get_export_empty_placeholder(doc)
        except Exception as exc:
            logger.debug(
                'Skipping table write (table=%s row=%s col=%s): %s',
                table_idx,
                row_idx,
                col_idx,
                exc
            )

    def _prevent_table_rows_splitting(doc, table_idx: int, row_indexes=None) -> None:
        """Prevent selected table rows from splitting across pages."""
        try:
            table = doc.tables[table_idx]
            if row_indexes is None:
                target_rows = table.rows
            else:
                target_rows = [
                    table.rows[row_index]
                    for row_index in row_indexes
                    if 0 <= row_index < len(table.rows)
                ]

            for row in target_rows:
                tr = row._tr
                tr_pr = tr.get_or_add_trPr()
                cant_split = OxmlElement('w:cantSplit')
                tr_pr.append(cant_split)
        except Exception as exc:
            logger.debug(
                'Unable to enforce non-splitting rows for table %s: %s',
                table_idx,
                exc
            )

    def _ensure_table_repeated_rows(doc, table_idx: int, template_row_idx: int, required_count: int) -> None:
        """Resize repeated table rows to exactly match the requested data-row count."""
        try:
            table = doc.tables[table_idx]
            if template_row_idx < 0 or template_row_idx >= len(table.rows):
                return

            template_tr = table.rows[template_row_idx]._tr
            existing_count = len(table.rows) - template_row_idx
            if required_count < existing_count:
                for row_idx in range(len(table.rows) - 1, template_row_idx + required_count - 1, -1):
                    table._tbl.remove(table.rows[row_idx]._tr)
                existing_count = len(table.rows) - template_row_idx

            rows_to_add = required_count - existing_count

            for _ in range(max(rows_to_add, 0)):
                table._tbl.append(deepcopy(template_tr))
        except Exception as exc:
            logger.debug(
                'Unable to expand repeated rows for table %s from row %s: %s',
                table_idx,
                template_row_idx,
                exc
            )

    def _insert_paragraph_after(paragraph, text: str = ''):
        """Insert and return a new paragraph immediately after the given paragraph."""
        new_paragraph_xml = OxmlElement('w:p')
        paragraph._p.addnext(new_paragraph_xml)
        new_paragraph = Paragraph(new_paragraph_xml, paragraph._parent)
        if text:
            new_paragraph.add_run(text)
        return new_paragraph

    def _set_following_paragraph_text(doc, label_snippet: str, value) -> bool:
        """Set the paragraph immediately following a paragraph containing the label snippet."""
        normalized_label = ' '.join(str(label_snippet or '').split()).upper()
        if not normalized_label:
            return False
        for idx, paragraph in enumerate(doc.paragraphs):
            paragraph_text = ' '.join((paragraph.text or '').split()).upper()
            if normalized_label in paragraph_text and idx + 1 < len(doc.paragraphs):
                text_value = _coerce_export_text(value)
                doc.paragraphs[idx + 1].text = text_value if text_value else _get_export_empty_placeholder(doc)
                return True
        return False

    def _load_block_diagram_image_bytes(block_diagram_value):
        """Extract block diagram image bytes from data-url/base64/path/json."""
        if block_diagram_value in (None, ''):
            return None

        if isinstance(block_diagram_value, dict):
            for key in ('data', 'base64', 'content', 'path', 'file_path', 'url'):
                if key in block_diagram_value and block_diagram_value.get(key):
                    return _load_block_diagram_image_bytes(block_diagram_value.get(key))
            return None

        raw_text = str(block_diagram_value).strip()
        if not raw_text:
            return None

        # Handle serialized JSON payloads.
        if (
            (raw_text.startswith('{') and raw_text.endswith('}'))
            or (raw_text.startswith('[') and raw_text.endswith(']'))
        ):
            try:
                parsed = json.loads(raw_text)
                return _load_block_diagram_image_bytes(parsed)
            except (json.JSONDecodeError, TypeError, ValueError):
                pass

        # Handle data URL format.
        data_url_match = re.match(
            r'^data:image\/[a-zA-Z0-9.+-]+;base64,(.+)$',
            raw_text,
            re.DOTALL
        )
        if data_url_match:
            try:
                encoded = data_url_match.group(1).strip()
                return base64.b64decode(encoded)
            except Exception:
                return None

        # Handle filesystem paths.
        candidate_paths = [raw_text]
        if not os.path.isabs(raw_text):
            candidate_paths.append(
                os.path.join(
                    flask_app.root_path,
                    raw_text.lstrip('/\\')
                )
            )
        for candidate in candidate_paths:
            if os.path.exists(candidate) and os.path.isfile(candidate):
                try:
                    with open(candidate, 'rb') as image_file:
                        return image_file.read()
                except Exception:
                    continue

        # Handle plain base64 without data-url prefix.
        try:
            compact = re.sub(r'\s+', '', raw_text)
            if len(compact) > 120 and re.fullmatch(r'[A-Za-z0-9+/=]+', compact):
                return base64.b64decode(compact)
        except Exception:
            return None

        return None

    def _insert_block_diagram_in_template(doc, block_diagram_value) -> None:
        """Insert block diagram image under the block diagram section in template."""
        image_bytes = _load_block_diagram_image_bytes(block_diagram_value)
        if not image_bytes:
            return

        anchor_paragraph = None
        for paragraph in doc.paragraphs:
            text_upper = (paragraph.text or '').upper()
            if 'BLOCK DIAGRAM OF THE TEST SETUP' in text_upper:
                anchor_paragraph = paragraph
                break

        if not anchor_paragraph:
            return

        try:
            # Place image in a dedicated paragraph immediately BELOW the anchor text line.
            image_paragraph = _insert_paragraph_after(anchor_paragraph, '')
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_run = image_paragraph.add_run()
            image_run.add_picture(BytesIO(image_bytes), width=Inches(2.8))
        except Exception as exc:
            logger.warning('Unable to insert block diagram image into TRF template: %s', exc)

    def _format_effective_date_value(value) -> str:
        """Format a request date value for document header display."""
        if value is None:
            return ''
        if isinstance(value, datetime):
            return value.date().isoformat()
        text_value = str(value).strip()
        if not text_value:
            return ''
        try:
            normalized = text_value.replace('Z', '+00:00')
            parsed = datetime.fromisoformat(normalized)
            return parsed.date().isoformat()
        except Exception:
            if 'T' in text_value:
                return text_value.split('T', 1)[0]
            return text_value

    def _resolve_request_effective_date(request_data: dict) -> str:
        """Resolve the request raised date for export header."""
        for key in ('requester_date', 'submitted_at', 'created_at'):
            formatted = _format_effective_date_value(request_data.get(key))
            if formatted:
                return formatted
        return ''

    def _format_effective_date_for_template_header(value: str) -> str:
        """Format date for template header cell, e.g. '18 Jun 2024'."""
        if not value:
            return ''
        try:
            parsed = datetime.fromisoformat(value)
            return parsed.strftime('%d %b %Y')
        except Exception:
            try:
                parsed = datetime.fromisoformat(str(value).replace('Z', '+00:00'))
                return parsed.strftime('%d %b %Y')
            except Exception:
                return str(value)

    def _format_test_duration_days_only(value) -> str:
        """Return duration text as day-count only (no hour/day suffix)."""
        raw = _coerce_export_text(value)
        if not raw:
            return ''
        text = str(raw).strip()
        if not text:
            return ''
        # Keep only the numeric portion when unit text is present.
        match = re.search(r'\d+(?:\.\d+)?', text)
        return match.group(0) if match else text

    def _set_template_effective_date_in_header(doc, effective_date_text: str) -> None:
        """Write effective date into the template header metadata table."""
        if not effective_date_text:
            return
        try:
            for section in doc.sections:
                header = section.header
                for table in header.tables:
                    for row in table.rows:
                        if len(row.cells) < 3:
                            continue
                        label = (row.cells[1].text or '').strip().lower()
                        if 'effective date' in label:
                            row.cells[2].text = effective_date_text
                            return
        except Exception as exc:
            logger.debug('Unable to set template header effective date: %s', exc)

    def _remove_template_helper_text(doc) -> None:
        """Remove helper text snippets from template labels."""
        helper_patterns = [
            r'\s*\(Please choose applicable\)',
        ]

        def _clean_text(value: str) -> str:
            text = str(value or '')
            for pattern in helper_patterns:
                text = re.sub(pattern, '', text, flags=re.IGNORECASE)
            return text

        try:
            for paragraph in doc.paragraphs:
                cleaned = _clean_text(paragraph.text)
                if cleaned != paragraph.text:
                    paragraph.text = cleaned

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cleaned = _clean_text(cell.text)
                        if cleaned != cell.text:
                            cell.text = cleaned

            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    cleaned = _clean_text(paragraph.text)
                    if cleaned != paragraph.text:
                        paragraph.text = cleaned
                for paragraph in section.footer.paragraphs:
                    cleaned = _clean_text(paragraph.text)
                    if cleaned != paragraph.text:
                        paragraph.text = cleaned
        except Exception as exc:
            logger.debug('Unable to remove helper text from template: %s', exc)

    def _get_reference_trf_template_path() -> str:
        """Legacy wrapper kept for callers still referencing the old helper name."""
        return _get_active_trf_template_path()

    def _build_test_request_word_export_from_template(request_data: dict):
        """Legacy wrapper that now delegates to the active Rev1 template exporter."""
        return _build_test_request_word_export_rev1(request_data)

        try:
            doc = DocxDocument(template_path)
            doc._empty_export_placeholder = '--------'
            _remove_template_helper_text(doc)

            effective_date = _resolve_request_effective_date(request_data)
            if effective_date:
                _set_template_effective_date_in_header(
                    doc,
                    _format_effective_date_for_template_header(effective_date)
                )

            # Table 0: service type selection
            service_types = [
                str(item).strip().lower()
                for item in _coerce_export_list(request_data.get('service_types'))
            ]
            if len(doc.tables) > 0 and len(doc.tables[0].rows) > 0:
                row0 = doc.tables[0].rows[0]
                service_checks = [
                    any('development' in item for item in service_types),
                    any('pre' in item and 'compliance' in item for item in service_types),
                    any('compliance' in item and 'pre' not in item for item in service_types),
                ]
                for idx, is_selected in enumerate(service_checks):
                    if idx < len(row0.cells):
                        current_label = row0.cells[idx].text.strip() or (
                            'Developmental Assistance (DA)' if idx == 0 else
                            'Pre-Compliance' if idx == 1 else
                            'Compliance'
                        )
                        row0.cells[idx].text = f"{'â˜‘' if is_selected else 'â˜'} {current_label}"

            # Table 1: product identity
            serial_number_display = _coerce_export_text(
                request_data.get('serial_number'))
            if not serial_number_display:
                serial_number_display = ', '.join(
                    str(x) for x in _coerce_export_list(request_data.get('serial_numbers'))
                )
            _set_table_cell_text(doc, 1, 0, 2, request_data.get('product_name'))
            _set_table_cell_text(doc, 1, 1, 2, request_data.get('manufacturer'))
            _set_table_cell_text(doc, 1, 2, 2, request_data.get('model_number'))
            _set_table_cell_text(doc, 1, 3, 2, serial_number_display)

            # Table 2: product details
            _set_table_cell_text(doc, 2, 0, 2, request_data.get('test_samples'))
            _set_table_cell_text(doc, 2, 1, 3, request_data.get('length'))
            _set_table_cell_text(doc, 2, 1, 6, request_data.get('width'))
            _set_table_cell_text(doc, 2, 1, 9, request_data.get('height'))
            _set_table_cell_text(doc, 2, 2, 2, request_data.get('weight'))
            _set_table_cell_text(doc, 2, 3, 2, request_data.get('operating_frequency'))

            category_selected = _selected_option_set(request_data.get('category'))
            category_checkbox_text = (
                f"{'â˜‘' if 'medical' in category_selected else 'â˜'} Medical    "
                f"{'â˜‘' if 'laboratory' in category_selected else 'â˜'} Laboratory    "
                f"{'â˜‘' if 'custom' in category_selected else 'â˜'} Custom"
            )
            _set_table_cell_text(doc, 2, 4, 2, category_checkbox_text)

            product_type_raw = _coerce_export_text(request_data.get('type')).lower()
            type_other = _coerce_export_text(request_data.get('type_others'))
            is_floor_standing = 'floor' in product_type_raw
            is_tabletop = 'table' in product_type_raw
            is_others = 'other' in product_type_raw and not (
                is_floor_standing or is_tabletop
            )

            floor_text = f"{'â˜‘' if is_floor_standing else 'â˜'} Floor Standing"
            tabletop_text = f"{'â˜‘' if is_tabletop else 'â˜'} Tabletop"
            others_text = f"{'â˜‘' if is_others else 'â˜'} others: specify"
            if is_others and type_other:
                others_text = f"{others_text} {type_other}"

            # Table template uses merged cells; write only once per merged block.
            _set_table_cell_text(doc, 2, 5, 2, floor_text)
            _set_table_cell_text(doc, 2, 5, 4, tabletop_text)
            _set_table_cell_text(doc, 2, 5, 7, others_text)

            # Table 3: accessories
            accessories = _coerce_export_list(request_data.get('accessories'))
            for idx, accessory in enumerate(accessories[:4], start=1):
                _set_table_cell_text(doc, 3, idx, 1, _pick_item_value(
                    accessory, 'equipmentName', 'equipment_name', 'name'))
                _set_table_cell_text(doc, 3, idx, 2, _pick_item_value(
                    accessory, 'make', 'manufacturer'))
                _set_table_cell_text(doc, 3, idx, 3, _pick_item_value(
                    accessory, 'modelNo', 'model_no', 'model'))
                _set_table_cell_text(doc, 3, idx, 4, _pick_item_value(
                    accessory, 'serialNo', 'serial_no', 'serialNumber'))

            # Table 4: cables
            cables = _coerce_export_list(request_data.get('cables'))
            for idx, cable in enumerate(cables[:6], start=1):
                _set_table_cell_text(doc, 4, idx, 1, _pick_item_value(
                    cable, 'cableName', 'cable_name', 'name'))
                _set_table_cell_text(doc, 4, idx, 2, _pick_item_value(
                    cable, 'length', 'length_m'))
                _set_table_cell_text(doc, 4, idx, 3, _pick_item_value(
                    cable, 'powerSignal', 'power_signal', 'type'))
                _set_table_cell_text(doc, 4, idx, 4, _pick_item_value(
                    cable, 'shielded', 'shielded_unshielded', 'shielding'))

            # Table 5: operating condition (use first EUT spec row)
            eut_specs = _coerce_export_list(request_data.get('eut_specs'))
            eut_spec = eut_specs[0] if eut_specs and isinstance(
                eut_specs[0], dict) else {}
            _set_table_cell_text(doc, 5, 1, 1, _pick_item_value(
                eut_spec, 'acVoltageRange', 'ac_voltage_range'))
            _set_table_cell_text(doc, 5, 1, 2, _pick_item_value(
                eut_spec, 'acVoltageNominal', 'ac_voltage_nominal'))
            _set_table_cell_text(doc, 5, 1, 3, _pick_item_value(
                eut_spec, 'acFreqRange', 'freqRange', 'freq_range'))
            _set_table_cell_text(doc, 5, 1, 4, _pick_item_value(
                eut_spec, 'acFreqNominal', 'freqNominal', 'freq_nominal'))
            _set_table_cell_text(doc, 5, 1, 5, _pick_item_value(
                eut_spec, 'dcVoltageRange', 'dc_voltage_range'))
            _set_table_cell_text(doc, 5, 1, 6, _pick_item_value(
                eut_spec, 'dcVoltageNominal', 'dc_voltage_nominal'))
            _set_table_cell_text(doc, 5, 1, 7, _pick_item_value(
                eut_spec, 'acDcInputCurrent', 'acInputCurrent', 'inputCurrent', 'input_current', 'ac_input_current'))
            _set_table_cell_text(doc, 5, 1, 8, _pick_item_value(
                eut_spec, 'ratedPower', 'rated_power'))

            # Table 6: wireless interfaces
            wireless_rows = _coerce_export_list(request_data.get('wireless'))
            for idx, wireless in enumerate(wireless_rows[:2], start=1):
                _set_table_cell_text(doc, 6, idx, 1, _pick_item_value(
                    wireless, 'type', 'interfaceType', 'interface_type'))
                _set_table_cell_text(doc, 6, idx, 2, _pick_item_value(
                    wireless, 'carrierFrequency', 'carrier_frequency'))
                _set_table_cell_text(doc, 6, idx, 3, _pick_item_value(
                    wireless, 'maxOutputPower', 'max_output_power'))
                _set_table_cell_text(doc, 6, idx, 4, _pick_item_value(
                    wireless, 'clockFrequencies', 'clock_frequencies'))
            _insert_block_diagram_in_template(
                doc,
                request_data.get('block_diagram')
            )

            # Table 7: description/configuration
            _set_table_cell_text(doc, 7, 1, 2, request_data.get('product_description'))
            _set_table_cell_text(doc, 7, 3, 2, request_data.get('test_configuration'))
            _set_table_cell_text(doc, 7, 4, 2, request_data.get('operation_modes'))
            _set_table_cell_text(doc, 7, 5, 2, request_data.get('monitoring_parameters'))
            _set_table_cell_text(doc, 7, 6, 2, request_data.get('additional_info'))
            _prevent_table_rows_splitting(doc, 7, row_indexes=[4, 5, 6])

            # Table 8: test duration per selected test
            test_hours = request_data.get('test_hours')
            if isinstance(test_hours, str):
                try:
                    test_hours = json.loads(test_hours)
                except (json.JSONDecodeError, TypeError, ValueError):
                    test_hours = {}
            if not isinstance(test_hours, dict):
                test_hours = {}

            def _find_test_hour(*keys):
                for key in keys:
                    if key in test_hours and test_hours.get(key) not in (None, ''):
                        return test_hours.get(key)
                return ''

            test_hours_mapping = [
                (1, ('CE', 'Conducted Emission')),
                (2, ('RE', 'Radiated Emission')),
                (3, ('Harmonic', 'HARMONIC', 'Harmonic Current Emission')),
                (4, ('VoltageFlicker', 'Voltage Flicker', 'Flicker')),
                (5, ('ESD', 'Electrostatic Discharge Immunity')),
                (6, ('RS_RI', 'RS', 'Radiated Susceptibility/Immunity')),
                (7, ('EFT', 'Electrical fast transient/burst Immunity')),
                (8, ('Surge', 'SURGE')),
                (9, ('CRF', 'Conducted RF Disturbance Immunity')),
                (10, ('PFMF', 'POWER', 'Power Frequency Magnetic Field Immunity')),
                (11, ('VoltageDips', 'VOLTAGE', 'Voltage Dips, Short Interruptions Immunity')),
            ]
            for row_idx, keys in test_hours_mapping:
                _set_table_cell_text(doc, 8, row_idx, 1, _find_test_hour(*keys))

            # Table 9: standards/environment/group/class
            standards_selected = _selected_option_set(
                request_data.get('product_standards')
            )
            standards_text = (
                f"{'â˜‘' if _selection_contains(standards_selected, 'iec 61326-1') else 'â˜'} IEC 61326-1    "
                f"{'â˜‘' if _selection_contains(standards_selected, 'en 61326-1') else 'â˜'} EN 61326-1    "
                f"{'â˜‘' if _selection_contains(standards_selected, 'iec 60601-1-2') else 'â˜'} IEC 60601-1-2\n"
                f"{'â˜‘' if _selection_contains(standards_selected, 'en 60601-1-2') else 'â˜'} EN 60601-1-2    "
                f"{'â˜‘' if _selection_contains(standards_selected, 'ansi c63.4', 'ansi c63-4') else 'â˜'} ANSI C63.4    "
                f"{'â˜‘' if _selection_contains(standards_selected, 'other') else 'â˜'} Other: please specify"
            )
            _set_table_cell_text(doc, 9, 0, 1, standards_text)

            environment_selected = _selected_option_set(
                request_data.get('product_environment')
            )
            environment_text = (
                f"{'â˜‘' if _selection_contains(environment_selected, 'basic') else 'â˜'} Basic Electromagnetic    "
                f"{'â˜‘' if _selection_contains(environment_selected, 'industrial') else 'â˜'} Industrial    "
                f"{'â˜‘' if _selection_contains(environment_selected, 'controlled') else 'â˜'} Controlled    "
                f"{'â˜‘' if _selection_contains(environment_selected, 'na', 'n/a') else 'â˜'} NA"
            )
            _set_table_cell_text(doc, 9, 1, 1, environment_text)

            group_selected = _selected_option_set(request_data.get('group'))
            _set_table_cell_text(
                doc, 9, 2, 1,
                f"{'â˜‘' if _selection_contains(group_selected, 'group 1') else 'â˜'} Group 1"
            )
            _set_table_cell_text(
                doc, 9, 2, 2,
                f"{'â˜‘' if _selection_contains(group_selected, 'group 2') else 'â˜'} Group 2"
            )

            class_selected = _selected_option_set(request_data.get('class_type'))
            _set_table_cell_text(
                doc, 9, 3, 1,
                f"{'â˜‘' if _selection_contains(class_selected, 'class a') else 'â˜'} Class A"
            )
            _set_table_cell_text(
                doc, 9, 3, 2,
                f"{'â˜‘' if _selection_contains(class_selected, 'class b') else 'â˜'} Class B"
            )

            # Table 10: test-specific details (populate selected tests and entered values)
            selected_tests = set()
            selected_test_shortcuts = set()

            def _accumulate_selected_tests(source_value) -> None:
                if source_value in (None, '', [], {}):
                    return
                selected_tests.update(_selected_option_set(source_value))
                selected_test_shortcuts.update(
                    _extract_selected_test_shortcuts(source_value)
                )

            # Primary sources (parsed values from DB model dictionary)
            _accumulate_selected_tests(request_data.get('selected_tests'))
            _accumulate_selected_tests(request_data.get('selected_tests_for_development'))
            _accumulate_selected_tests(request_data.get('testSelected'))

            # Raw DB fallbacks (covers legacy rows where selected_tests was not stored as list)
            _accumulate_selected_tests(request_data.get('_selected_tests_db_raw'))
            _accumulate_selected_tests(
                request_data.get('_selected_tests_for_development_db_raw')
            )

            # Last resort fallback: infer selected tests from stored test-hours keys
            test_hours_for_selection = request_data.get('test_hours')
            if isinstance(test_hours_for_selection, str):
                try:
                    test_hours_for_selection = json.loads(test_hours_for_selection)
                except (json.JSONDecodeError, TypeError, ValueError):
                    test_hours_for_selection = {}
            if isinstance(test_hours_for_selection, dict) and test_hours_for_selection:
                _accumulate_selected_tests(list(test_hours_for_selection.keys()))

            def _is_test_selected(*aliases: str) -> bool:
                for alias in aliases:
                    alias_shortcuts = _extract_selected_test_shortcuts(alias)
                    if alias_shortcuts and (alias_shortcuts & selected_test_shortcuts):
                        return True
                return _selection_contains(selected_tests, *aliases)

            # Table 10: section header checkboxes in first column
            test_block_rows = [
                (0, 'Conducted Emission', ('CE', 'conducted emission')),
                (6, 'Radiated Emission', ('RE', 'radiated emission')),
                (11, 'Harmonic Current Emission', ('Harmonic', 'harmonic current emission')),
                (14, 'Voltage Changes, Voltage Fluctuations and Flicker Emission',
                 ('VoltageFlicker', 'Flicker', 'voltage changes')),
                (17, 'Electrostatic Discharge Immunity', ('ESD', 'electrostatic discharge')),
                (21, 'Radiated Susceptibility/Immunity', ('RS_RI', 'RS', 'radiated susceptibility')),
                (26, 'Radiated Susceptibility/Immunity (Interim Method)',
                 ('RS_RI_Interim', 'RS_Interim', 'interim method')),
                (31, 'Electrical fast transient/burst Immunity', ('EFT', 'fast transient')),
                (36, 'Surge Immunity', ('Surge',)),
                (42, 'Conducted RF Disturbance Immunity', ('CRF', 'conducted rf')),
                (48, 'Power Frequency Magnetic Field Immunity', ('PFMF', 'POWER')),
                (52, 'Voltage Dips, Short Interruptions and Voltage Variations Immunity',
                 ('VoltageDips', 'VOLTAGE', 'voltage dips')),
            ]

            def _set_table10_block_checkbox(label: str, aliases: tuple, fallback_row_idx: int) -> None:
                checked_symbol = '\u2611' if _is_test_selected(*aliases) else '\u2610'
                target_text = f'{checked_symbol} {label}'
                updated = False

                # Prefer label-based update to handle template row shifts/merged-cell behavior.
                try:
                    table10 = doc.tables[10] if len(doc.tables) > 10 else None
                    if table10:
                        for row in table10.rows:
                            if not row.cells:
                                continue
                            cell0 = row.cells[0]
                            existing = (cell0.text or '').strip().lower()
                            if not existing:
                                continue
                            label_norm = label.strip().lower()
                            if label_norm in existing:
                                cell0.text = target_text
                                updated = True
                                break
                except Exception as exc:
                    logger.debug('Unable to set table10 block checkbox by label (%s): %s', label, exc)

                # Fallback to fixed row index.
                if not updated:
                    _set_table_cell_text(doc, 10, fallback_row_idx, 0, target_text)

            for row_idx, label, aliases in test_block_rows:
                _set_table10_block_checkbox(label, aliases, row_idx)

            # Table 8: Selected tests list with checkbox marks
            test_name_rows = [
                (1, 'Conducted Emission (CE)', ('CE', 'conducted emission')),
                (2, 'Radiated Emission (RE)', ('RE', 'radiated emission')),
                (3, 'Harmonic Current Emission', ('Harmonic', 'harmonic current emission')),
                (4, 'Voltage Changes, Voltage Fluctuations and Flicker Emission',
                 ('VoltageFlicker', 'Flicker', 'voltage changes')),
                (5, 'Electrostatic Discharge Immunity (ESD)', ('ESD', 'electrostatic discharge')),
                (6, 'Radiated Susceptibility/Immunity (RS/RI)', ('RS_RI', 'RS', 'radiated susceptibility')),
                (7, 'Electrical fast transient/burst Immunity (EFT)', ('EFT', 'fast transient')),
                (8, 'Surge Immunity', ('Surge',)),
                (9, 'Conducted RF Disturbance Immunity (CRF)', ('CRF', 'conducted rf')),
                (10, 'Power Frequency Magnetic Field Immunity', ('PFMF', 'POWER')),
                (11, 'Voltage Dips, Short Interruptions Immunity', ('VoltageDips', 'VOLTAGE')),
            ]
            for row_idx, label, aliases in test_name_rows:
                checked_symbol = 'â˜‘' if _is_test_selected(*aliases) else 'â˜'
                _set_table_cell_text(doc, 8, row_idx, 0, f'{checked_symbol} {label}')

            def _set_test_row_if_selected(row_idx: int, value, *aliases: str) -> None:
                if not _is_test_selected(*aliases):
                    return
                text_value = _coerce_export_text(value)
                if text_value:
                    _set_table_cell_text(doc, 10, row_idx, 2, text_value)

            def _set_standard_row_if_selected(row_idx: int, standards_value, aliases: tuple, options: list) -> None:
                if not _is_test_selected(*aliases):
                    return
                selected_standards = _selected_option_set(standards_value)
                _set_table_cell_text(
                    doc,
                    10,
                    row_idx,
                    2,
                    _checkbox_options_text(options, selected_standards)
                )

            # CE
            ce_standards_for_export = _resolve_test_standard_values(
                request_data.get('ce_standard'),
                'CE',
                request_data.get('product_standards')
            )
            _set_standard_row_if_selected(
                0,
                ce_standards_for_export,
                ('ce', 'conducted emission'),
                [('CISPR 11', 'cispr 11'), ('EN 55011', 'en 55011'),
                 ('ANSI C63.4', 'ansi c63.4', 'ansi c63-4')]
            )
            _set_test_row_if_selected(1, request_data.get('ce_voltage_freq'), 'ce', 'conducted emission')
            _set_test_row_if_selected(
                2,
                _extract_frequency_range_display(request_data, 'ce_custom_spec', 'ce', 'ce_freq_range'),
                'ce',
                'conducted emission'
            )
            _set_test_row_if_selected(3, f"Power: {_coerce_export_text(request_data.get('ce_cables'))}",
                                      'ce', 'conducted emission')
            _set_test_row_if_selected(4, request_data.get('ce_class'), 'ce', 'conducted emission')

            # RE
            re_standards_for_export = _resolve_test_standard_values(
                request_data.get('re_standard'),
                'RE',
                request_data.get('product_standards')
            )
            _set_standard_row_if_selected(
                6,
                re_standards_for_export,
                ('re', 'radiated emission'),
                [('CISPR 11', 'cispr 11'), ('EN 55011', 'en 55011'),
                 ('ANSI C63.4', 'ansi c63.4', 'ansi c63-4')]
            )
            _set_test_row_if_selected(7, request_data.get('re_voltage_freq'), 're', 'radiated emission')
            _set_test_row_if_selected(
                8,
                _extract_frequency_range_display(request_data, 're_custom_spec', 're', 're_freq_range'),
                're',
                'radiated emission'
            )
            _set_test_row_if_selected(9, request_data.get('re_class'), 're', 'radiated emission')

            # Harmonic
            harmonic_standards_for_export = _resolve_test_standard_values(
                request_data.get('harmonic_standard'),
                'HARMONIC',
                request_data.get('product_standards')
            )
            _set_standard_row_if_selected(
                11,
                harmonic_standards_for_export,
                ('harmonic', 'harmonic current emission'),
                [('IEC 61000-3-2', 'iec 61000-3-2'),
                 ('EN 61000-3-2', 'en 61000-3-2')]
            )
            harmonic_voltage = _coerce_export_text(request_data.get('harmonic_voltage_freq'))
            harmonic_class = _coerce_export_text(request_data.get('harmonic_class'))
            harmonic_detail = harmonic_voltage
            if harmonic_class:
                harmonic_detail = f"{harmonic_detail} | Class: {harmonic_class}" if harmonic_detail else f"Class: {harmonic_class}"
            _set_test_row_if_selected(12, harmonic_detail, 'harmonic', 'harmonic current emission')

            # Flicker / Voltage Fluctuation
            _set_standard_row_if_selected(
                14,
                request_data.get('flicker_standard'),
                ('voltageflicker', 'flicker', 'voltage changes'),
                [('IEC 61000-3-3', 'iec 61000-3-3'),
                 ('EN 61000-3-3', 'en 61000-3-3')]
            )
            _set_test_row_if_selected(15, request_data.get('flicker_voltage_freq'),
                                      'voltageflicker', 'flicker', 'voltage changes')

            # ESD
            _set_standard_row_if_selected(
                17,
                request_data.get('esd_standard'),
                ('esd', 'electrostatic discharge'),
                [('IEC 61000-4-2', 'iec 61000-4-2'),
                 ('EN 61000-4-2', 'en 61000-4-2')]
            )
            _set_test_row_if_selected(18, request_data.get('esd_voltage_freq'), 'esd', 'electrostatic discharge')
            esd_contact = _coerce_export_text(request_data.get('esd_contact'))
            esd_air = _coerce_export_text(request_data.get('esd_air'))
            esd_level = ''
            if esd_contact:
                esd_level += f"Contact: {esd_contact}"
            if esd_air:
                esd_level += (' | ' if esd_level else '') + f"Air: {esd_air}"
            _set_test_row_if_selected(19, esd_level, 'esd', 'electrostatic discharge')

            # RS/RI
            _set_standard_row_if_selected(
                21,
                request_data.get('rs_standard'),
                ('rs_ri', 'rs', 'radiated susceptibility'),
                [('IEC 61000-4-3', 'iec 61000-4-3'),
                 ('EN 61000-4-3', 'en 61000-4-3')]
            )
            _set_test_row_if_selected(22, request_data.get('rs_voltage_freq'),
                                      'rs_ri', 'rs', 'radiated susceptibility')
            _set_test_row_if_selected(23,
                                      _extract_frequency_range_display(request_data, 'rs_ri_custom_spec', 'rs_ri', 'rs_freq_range'),
                                      'rs_ri', 'rs', 'radiated susceptibility')
            rs_field = _coerce_export_text(request_data.get('rs_field_strength1'))
            rs_field2 = _coerce_export_text(request_data.get('rs_field_strength2'))
            rs_field3 = _coerce_export_text(request_data.get('rs_field_strength3'))
            rs_field_text = ' / '.join([x for x in [rs_field, rs_field2, rs_field3] if x])
            _set_test_row_if_selected(24, rs_field_text, 'rs_ri', 'rs', 'radiated susceptibility')

            # RS/RI Interim
            _set_standard_row_if_selected(
                26,
                request_data.get('rs_interim_standard'),
                ('rs_ri_interim', 'rs_interim', 'interim method'),
                [('IEC 61000-4-3', 'iec 61000-4-3'),
                 ('EN 61000-4-3', 'en 61000-4-3')]
            )
            _set_test_row_if_selected(27, request_data.get('rs_interim_voltage_freq'),
                                      'rs_ri_interim', 'rs_interim', 'interim method')
            _set_test_row_if_selected(28,
                                      _extract_frequency_range_display(request_data, 'rs_ri_interim_custom_spec', 'rs_ri_interim', 'rs_interim_freq_range'),
                                      'rs_ri_interim', 'rs_interim', 'interim method')
            rs_i_field = _coerce_export_text(request_data.get('rs_interim_field_strength1'))
            rs_i_field2 = _coerce_export_text(request_data.get('rs_interim_field_strength2'))
            rs_i_field3 = _coerce_export_text(request_data.get('rs_interim_field_strength3'))
            rs_i_text = ' / '.join([x for x in [rs_i_field, rs_i_field2, rs_i_field3] if x])
            _set_test_row_if_selected(29, rs_i_text, 'rs_ri_interim', 'rs_interim', 'interim method')

            # EFT
            _set_standard_row_if_selected(
                31,
                request_data.get('eft_standard'),
                ('eft', 'fast transient'),
                [('IEC 61000-4-4', 'iec 61000-4-4'),
                 ('EN 61000-4-4', 'en 61000-4-4')]
            )
            _set_test_row_if_selected(32, request_data.get('eft_voltage_freq'), 'eft', 'fast transient')
            eft_cables = []
            eft_power = _coerce_export_text(request_data.get('eft_cables_power'))
            eft_signal = _coerce_export_text(request_data.get('eft_cables_signal'))
            if eft_power:
                eft_cables.append(f"Power: {eft_power}")
            if eft_signal:
                eft_cables.append(f"Signal: {eft_signal}")
            _set_test_row_if_selected(33, ' | '.join(eft_cables), 'eft', 'fast transient')
            eft_level1 = _coerce_export_text(request_data.get('eft_test_level1'))
            eft_level2 = _coerce_export_text(request_data.get('eft_test_level2'))
            eft_levels = ' / '.join([x for x in [eft_level1, eft_level2] if x])
            _set_test_row_if_selected(34, eft_levels, 'eft', 'fast transient')

            # Surge
            _set_standard_row_if_selected(
                36,
                request_data.get('surge_standard'),
                ('surge',),
                [('IEC 61000-4-5', 'iec 61000-4-5'),
                 ('EN 61000-4-5', 'en 61000-4-5')]
            )
            _set_test_row_if_selected(37, request_data.get('surge_voltage_freq'), 'surge')
            surge_cables = []
            surge_power = _coerce_export_text(request_data.get('surge_cables_power'))
            surge_signal = _coerce_export_text(request_data.get('surge_cables_signal'))
            if surge_power:
                surge_cables.append(f"Power: {surge_power}")
            if surge_signal:
                surge_cables.append(f"Signal: {surge_signal}")
            _set_test_row_if_selected(38, ' | '.join(surge_cables), 'surge')
            surge_cm1 = _coerce_export_text(request_data.get('surge_cm1'))
            surge_cm2 = _coerce_export_text(request_data.get('surge_cm2'))
            surge_dm1 = _coerce_export_text(request_data.get('surge_dm1'))
            surge_dm2 = _coerce_export_text(request_data.get('surge_dm2'))
            _set_test_row_if_selected(39, ' / '.join([x for x in [surge_cm1, surge_cm2] if x]), 'surge')
            _set_test_row_if_selected(40, ' / '.join([x for x in [surge_dm1, surge_dm2] if x]), 'surge')

            # CRF
            _set_standard_row_if_selected(
                42,
                request_data.get('crf_standard'),
                ('crf', 'conducted rf'),
                [('IEC 61000-4-6', 'iec 61000-4-6'),
                 ('EN 61000-4-6', 'en 61000-4-6')]
            )
            _set_test_row_if_selected(43, request_data.get('crf_voltage_freq'), 'crf', 'conducted rf')
            _set_test_row_if_selected(
                44,
                _extract_frequency_range_display(request_data, 'crf_custom_spec', 'crf', 'crf_freq_range'),
                'crf',
                'conducted rf'
            )
            crf_cables = []
            crf_power = _coerce_export_text(request_data.get('crf_cables_power'))
            crf_signal = _coerce_export_text(request_data.get('crf_cables_signal'))
            if crf_power:
                crf_cables.append(f"Power: {crf_power}")
            if crf_signal:
                crf_cables.append(f"Signal: {crf_signal}")
            _set_test_row_if_selected(45, ' | '.join(crf_cables), 'crf', 'conducted rf')
            crf_level1 = _coerce_export_text(request_data.get('crf_test_level1'))
            crf_level2 = _coerce_export_text(request_data.get('crf_test_level2'))
            _set_test_row_if_selected(46, ' / '.join([x for x in [crf_level1, crf_level2] if x]),
                                      'crf', 'conducted rf')

            # PFMF
            _set_standard_row_if_selected(
                48,
                request_data.get('power_freq_standard'),
                ('pfmf', 'power', 'power frequency magnetic field'),
                [('IEC 61000-4-8', 'iec 61000-4-8'),
                 ('EN 61000-4-8', 'en 61000-4-8')]
            )
            _set_test_row_if_selected(49, request_data.get('power_freq_voltage_freq'),
                                      'pfmf', 'power', 'power frequency magnetic field')
            _set_test_row_if_selected(50, request_data.get('power_freq_test_level'),
                                      'pfmf', 'power', 'power frequency magnetic field')

            # Voltage Dips
            _set_standard_row_if_selected(
                52,
                request_data.get('voltage_dips_standard'),
                ('voltagedips', 'voltage', 'voltage dips'),
                [('IEC 61000-4-11', 'iec 61000-4-11'),
                 ('EN 61000-4-11', 'en 61000-4-11')]
            )
            vd_min = _coerce_export_text(request_data.get('voltage_dips_min'))
            vd_max = _coerce_export_text(request_data.get('voltage_dips_max'))
            vd_range = ' - '.join([x for x in [vd_min, vd_max] if x])
            _set_test_row_if_selected(53, vd_range, 'voltagedips', 'voltage', 'voltage dips')
            _set_test_row_if_selected(54, request_data.get('voltage_dips_voltage_freq'),
                                      'voltagedips', 'voltage', 'voltage dips')
            vd_levels = ' / '.join([
                x for x in [
                    _coerce_export_text(request_data.get('voltage_dips_voltage_dip1')),
                    _coerce_export_text(request_data.get('voltage_dips_voltage_dip2')),
                    _coerce_export_text(request_data.get('voltage_dips_voltage_dip3')),
                    _coerce_export_text(request_data.get('voltage_dips_interruption'))
                ] if x
            ])
            _set_test_row_if_selected(55, vd_levels, 'voltagedips', 'voltage', 'voltage dips')
            vd_times = ' / '.join([
                x for x in [
                    _coerce_export_text(request_data.get('voltage_dips_time1')),
                    _coerce_export_text(request_data.get('voltage_dips_time2')),
                    _coerce_export_text(request_data.get('voltage_dips_time3')),
                    _coerce_export_text(request_data.get('voltage_dips_time4'))
                ] if x
            ])
            _set_test_row_if_selected(57, vd_times, 'voltagedips', 'voltage', 'voltage dips')

            # Table 12: requester details
            _set_table_cell_text(doc, 12, 0, 1, request_data.get('requester_name'))
            requester_dept = _coerce_export_text(
                request_data.get('requester_department'))
            requester_div = _coerce_export_text(request_data.get('requester_division'))
            _set_table_cell_text(
                doc, 12, 0, 3,
                f"{requester_dept}{f' / {requester_div}' if requester_div else ''}"
            )
            _set_table_cell_text(doc, 12, 1, 1, request_data.get('requester_email'))
            _set_table_cell_text(doc, 12, 1, 3, request_data.get('requester_contact'))
            _set_table_cell_text(doc, 12, 2, 1, request_data.get('requester_designation'))
            _set_table_cell_text(doc, 12, 2, 3, request_data.get('requester_date'))
            _set_table_cell_text(doc, 12, 3, 1, _format_signature_export_value(
                request_data.get('requester_signature')))

            # Table 13: laboratory use
            _set_table_cell_text(doc, 13, 0, 1, request_data.get('job_number') or request_data.get('job_id'))
            _set_table_cell_text(doc, 13, 1, 1, request_data.get('sample_condition'))
            _set_table_cell_text(doc, 13, 2, 1, request_data.get('capability_available'))
            _set_table_cell_text(doc, 13, 3, 1, request_data.get('sample_received_date'))
            _set_table_cell_text(
                doc,
                13,
                5,
                2,
                _format_test_duration_days_only(request_data.get('test_duration'))
            )
            _set_table_cell_text(doc, 13, 6, 2, request_data.get('test_commencement_date'))
            _set_table_cell_text(doc, 13, 7, 2, request_data.get('test_completion_date'))

            # Table 14: lab manager sign-off
            _set_table_cell_text(doc, 14, 0, 1, request_data.get('lab_manager_name'))
            _set_table_cell_text(doc, 14, 0, 3, request_data.get('lab_manager_date'))
            _set_table_cell_text(doc, 14, 1, 1, _format_signature_export_value(
                request_data.get('lab_manager_signature')))

            output_stream = BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            return output_stream
        except Exception as exc:
            logger.warning(
                'TRF template export failed, falling back to generic export: %s',
                exc
            )
            return None

    def _build_generic_test_request_word_export(request_data: dict):
        """Build a self-contained DOCX when the branded template cannot be used."""
        doc = DocxDocument()
        request_title = request_data.get('tco_id') or f"Request {request_data.get('id', 'N/A')}"
        effective_date = _resolve_request_effective_date(request_data)
        if effective_date:
            doc.add_paragraph(f"Effective Date: {effective_date}")
        doc.add_heading('EMI/EMC Test Request Form', level=1)
        doc.add_paragraph(f"Reference: {request_title}")
        doc.add_paragraph(f"Generated On: {get_ist_now().strftime('%Y-%m-%d %H:%M:%S')} IST")

        sections = [
            (
                'Basic Information',
                [
                    ('Request ID', request_data.get('id')),
                    ('TCO ID', request_data.get('tco_id')),
                    ('Status', request_data.get('status')),
                    ('Requester Status', request_data.get('requester_status')),
                    ('Created At', request_data.get('created_at')),
                    ('Updated At', request_data.get('updated_at')),
                    ('Submitted At', request_data.get('submitted_at')),
                    ('Job ID', request_data.get('job_id')),
                    ('Job Number', request_data.get('job_number')),
                ]
            ),
            (
                'Type of Service Requested',
                [
                    ('Service Types', request_data.get('service_types')),
                ]
            ),
            (
                'Product Identity',
                [
                    ('Product Name', request_data.get('product_name')),
                    ('Manufacturer', request_data.get('manufacturer')),
                    ('Manufacturer Address', request_data.get('manufacturer_address')),
                    ('Model Number', request_data.get('model_number')),
                    ('Serial Number', request_data.get('serial_number')),
                    ('Serial Numbers', request_data.get('serial_numbers')),
                    ('Additional Models', request_data.get('additional_models')),
                    ('Test Samples', request_data.get('test_samples')),
                    ('Samples Available In Lab', request_data.get('samples_available_in_lab')),
                    ('Has Model Variance', request_data.get('has_model_variance')),
                    ('Model Variance', request_data.get('model_variance')),
                    ('Project Details/Intent', request_data.get('project_details_intent')),
                ]
            ),
            (
                'Product Details',
                [
                    ('Category', request_data.get('category')),
                    ('Type', request_data.get('type')),
                    ('Type Others', request_data.get('type_others')),
                    ('Dimension Unit', request_data.get('dimension_unit')),
                    ('Length', request_data.get('length')),
                    ('Width', request_data.get('width')),
                    ('Height', request_data.get('height')),
                    ('Weight', request_data.get('weight')),
                    ('Operating Frequency', request_data.get('operating_frequency')),
                    ('Accessories', request_data.get('accessories')),
                    ('Cables', request_data.get('cables')),
                    ('EUT Specs', request_data.get('eut_specs')),
                    ('Supply Voltage/Frequency', request_data.get('supply_vf')),
                    ('Has Wireless Interface', request_data.get('has_wireless_interface')),
                    ('Wireless Details', request_data.get('wireless')),
                ]
            ),
            (
                'Description and Test Scope',
                [
                    ('Product Description', request_data.get('product_description')),
                    ('Test Configuration', request_data.get('test_configuration')),
                    ('Operation Modes', request_data.get('operation_modes')),
                    ('Monitoring Parameters', request_data.get('monitoring_parameters')),
                    ('Additional Information', request_data.get('additional_info')),
                ]
            ),
            (
                'Tests Required',
                [
                    ('Product Standards', request_data.get('product_standards')),
                    ('Product Environment', request_data.get('product_environment')),
                    ('Product Environment Other', request_data.get('product_environment_other')),
                    ('Group', request_data.get('group')),
                    ('Class Type', request_data.get('class_type')),
                    ('Selected Tests', request_data.get('selected_tests')),
                    ('Selected Tests (Development)', request_data.get('selected_tests_for_development')),
                    ('Test Hours', request_data.get('test_hours')),
                    ('Test Remarks', request_data.get('test_remarks')),
                ]
            ),
            (
                'Test-Specific Details',
                [
                    ('CE Voltage/Frequency', request_data.get('ce_voltage_freq')),
                    ('CE Frequency Range', request_data.get('ce_freq_range')),
                    ('CE Cables', request_data.get('ce_cables')),
                    ('CE Class', request_data.get('ce_class')),
                    ('RE Frequency Range', request_data.get('re_freq_range')),
                    ('RE Class', request_data.get('re_class')),
                    ('ESD Contact', request_data.get('esd_contact')),
                    ('ESD Air', request_data.get('esd_air')),
                    ('Harmonic Class', request_data.get('harmonic_class')),
                    ('RS Frequency Range', request_data.get('rs_freq_range')),
                    ('RS Field Strength', request_data.get('rs_field_strength1')),
                ]
            ),
            (
                'Conformity and Functional Modes',
                [
                    ('Continue Testing', request_data.get('continue_testing')),
                    ('Test Report Required', request_data.get('test_report_required')),
                    ('Uncertainty Required', request_data.get('uncertainty_required')),
                    ('Test Witness', request_data.get('test_witness')),
                    ('Conformity Required', request_data.get('conformity_required')),
                    ('Conformity Statement', request_data.get('conformity_statement')),
                    ('Decision Rule', request_data.get('decision_rule')),
                    ('Number of Modes', request_data.get('number_of_modes')),
                    ('Functional Modes', request_data.get('functional_modes')),
                ]
            ),
            (
                'Scheduling and Assignment',
                [
                    ('Assigned Engineer Name', request_data.get('assigned_engineer_name')),
                    ('Assignment Priority', request_data.get('assignment_priority')),
                    ('Assignment Due Date', request_data.get('assignment_due_date')),
                    ('Assignment Notes', request_data.get('assignment_notes')),
                    ('Test Assignments', request_data.get('test_assignments')),
                    ('Planner Entries', request_data.get('planner_entries')),
                    ('Plan Update History', request_data.get('plan_update_history')),
                    ('Sample Condition', request_data.get('sample_condition')),
                    ('Capability Available', request_data.get('capability_available')),
                    ('Sample Received Date', request_data.get('sample_received_date')),
                    ('Test Duration', request_data.get('test_duration')),
                    ('Test Commencement Date', request_data.get('test_commencement_date')),
                    ('Test Completion Date', request_data.get('test_completion_date')),
                ]
            ),
            (
                'Requester and Lab Manager',
                [
                    ('Requester Name', request_data.get('requester_name')),
                    ('Requester Department', request_data.get('requester_department')),
                    ('Requester Group', request_data.get('requester_group')),
                    ('Requester Division', request_data.get('requester_division')),
                    ('Requester Site', request_data.get('requester_site')),
                    ('Requester Email', request_data.get('requester_email')),
                    ('Requester Contact', request_data.get('requester_contact')),
                    ('Requester Designation', request_data.get('requester_designation')),
                    ('Requester Date', request_data.get('requester_date')),
                    ('Expected Completion Date', request_data.get('requester_expected_completion_date')),
                    ('Lab Manager Name', request_data.get('lab_manager_name')),
                    ('Lab Manager Date', request_data.get('lab_manager_date')),
                    ('Lab Manager Signed At', request_data.get('lab_manager_signed_at')),
                ]
            ),
            (
                'Review and Sign-Off',
                [
                    ('Review Comments', request_data.get('review_comments')),
                    ('Reviewed By', request_data.get('reviewed_by')),
                    ('Reviewed At', request_data.get('reviewed_at')),
                    ('Rejection Reason', request_data.get('rejection_reason')),
                    ('Rejected By', request_data.get('rejected_by')),
                    ('Rejected At', request_data.get('rejected_at')),
                    ('Report Comments', request_data.get('report_comments')),
                    ('Report Uploaded At', request_data.get('report_uploaded_at')),
                ]
            ),
            (
                'Attachments',
                [
                    ('Model Variance Document', request_data.get('model_variance_document')),
                    ('Block Diagram', request_data.get('block_diagram')),
                    ('Requester Signature', request_data.get('requester_signature')),
                    ('Lab Manager Signature', request_data.get('lab_manager_signature')),
                    ('Report File Path', request_data.get('report_file_path')),
                    ('Generated Files', request_data.get('generated_files')),
                ]
            ),
        ]

        for section_title, section_rows in sections:
            _add_export_section(doc, section_title, section_rows)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def _get_active_trf_template_path() -> str:
        """Resolve the active Rev1 TRF template path."""
        candidate_paths = [
            flask_app.config.get('TRF_TEMPLATE_PATH'),
            os.environ.get('TRF_TEMPLATE_PATH'),
            os.path.join(flask_app.root_path, 'word_templates', 'EMI EMC TRF_Draft - Rev1.docx'),
            r'c:\Users\saimounik.chandavolu\Downloads\EMI EMC TRF_Draft - Rev1.docx',
        ]
        for path in candidate_paths:
            if path and os.path.isfile(path):
                return path
        return ''

    def _build_test_request_word_export_rev1(request_data: dict):
        """Build the request export using the Rev1 TRF template."""
        template_path = _get_active_trf_template_path()
        if not template_path:
            return None

        checked = '\u2611'
        unchecked = '\u2610'

        def _box(label: str, selected: bool) -> str:
            return f"{checked if selected else unchecked} {label}"

        def _parse_map(value) -> dict:
            if isinstance(value, dict):
                return value
            if isinstance(value, str):
                try:
                    parsed = json.loads(value)
                    return parsed if isinstance(parsed, dict) else {}
                except (json.JSONDecodeError, TypeError, ValueError):
                    return {}
            return {}

        def _norm_yes_no(value) -> str:
            if isinstance(value, bool):
                return 'yes' if value else 'no'
            text = _coerce_export_text(value).strip().lower()
            if text in ('yes', 'y', 'true', '1', 'good'):
                return 'yes'
            if text in ('no', 'n', 'false', '0', 'bad'):
                return 'no'
            return ''

        def _fmt_date(value) -> str:
            text = _coerce_export_text(value)
            if not text:
                return ''
            try:
                return datetime.fromisoformat(text.replace('Z', '+00:00')).strftime('%Y-%m-%d')
            except Exception:
                return text.split('T', 1)[0] if 'T' in text else text

        def _join(parts, sep='\n') -> str:
            values = [_coerce_export_text(part) for part in parts]
            return sep.join(item for item in values if item)

        def _line(label: str, value) -> str:
            text = _coerce_export_text(value)
            return f"{label}: {text}" if text else ''

        def _test_selected(*aliases: str) -> bool:
            for alias in aliases:
                alias_shortcuts = _extract_selected_test_shortcuts(alias)
                if alias_shortcuts and (alias_shortcuts & selected_test_shortcuts):
                    return True
            return _selection_contains(selected_tests, *aliases)

        def _map_value(source: dict, *keys):
            for key in keys:
                if source.get(key) not in (None, ''):
                    return source.get(key)
            return ''

        try:
            doc = DocxDocument(template_path)
            doc._empty_export_placeholder = '--------'
            _remove_template_helper_text(doc)

            effective_date = _resolve_request_effective_date(request_data)
            if effective_date:
                _set_template_effective_date_in_header(
                    doc,
                    _format_effective_date_for_template_header(effective_date)
                )

            service_types = [str(item).strip().lower() for item in _coerce_export_list(request_data.get('service_types'))]
            service_flags = [
                any('development' in item for item in service_types),
                any('pre' in item and 'compliance' in item for item in service_types),
                any('compliance' in item and 'pre' not in item for item in service_types),
            ]
            for idx, label in enumerate(['Developmental Assistance (DA)', 'Pre-Compliance', 'Compliance']):
                _set_table_cell_text(doc, 0, 0, idx, _box(label, service_flags[idx]))

            serials = _coerce_export_text(request_data.get('serial_number')) or ', '.join(str(x) for x in _coerce_export_list(request_data.get('serial_numbers')))
            identity_rows = {
                1: request_data.get('product_name'),
                2: request_data.get('manufacturer'),
                3: request_data.get('manufacturer_address'),
                4: request_data.get('model_number'),
                5: request_data.get('test_samples'),
                7: serials,
                8: request_data.get('requester_group'),
                9: request_data.get('requester_division'),
                10: request_data.get('requester_site'),
                12: request_data.get('model_variance'),
                13: request_data.get('project_details_intent'),
            }
            for row_idx, value in identity_rows.items():
                _set_table_cell_text(doc, 1, row_idx, 1, value)
            samples_lab = _norm_yes_no(request_data.get('samples_available_in_lab'))
            _set_table_cell_text(doc, 1, 6, 1, f"{_box('Yes', samples_lab == 'yes')}    {_box('No', samples_lab == 'no')}")
            model_variance = _norm_yes_no(request_data.get('has_model_variance'))
            _set_table_cell_text(doc, 1, 11, 1, f"{_box('Yes', model_variance == 'yes')}    {_box('No', model_variance == 'no')}")

            _set_table_cell_text(doc, 2, 1, 2, request_data.get('length'))
            _set_table_cell_text(doc, 2, 1, 5, request_data.get('width'))
            _set_table_cell_text(doc, 2, 1, 8, request_data.get('height'))
            _set_table_cell_text(doc, 2, 2, 1, request_data.get('weight'))
            _set_table_cell_text(doc, 2, 3, 1, request_data.get('operating_frequency'))
            cat = _selected_option_set(request_data.get('category'))
            _set_table_cell_text(doc, 2, 4, 1, '    '.join([
                _box('Medical', _selection_contains(cat, 'medical')),
                _box('Laboratory', _selection_contains(cat, 'laboratory')),
                _box('Custom', _selection_contains(cat, 'custom')),
            ]))
            type_raw = _coerce_export_text(request_data.get('type')).lower()
            type_other = _coerce_export_text(request_data.get('type_others'))
            _set_table_cell_text(doc, 2, 5, 1, _box('Floor Standing', 'floor' in type_raw))
            _set_table_cell_text(doc, 2, 5, 3, _box('Tabletop', 'table' in type_raw))
            _set_table_cell_text(doc, 2, 5, 6, _box(f"others: specify {type_other}".strip(), 'other' in type_raw and 'floor' not in type_raw and 'table' not in type_raw))

            accessories = [
                item for item in _coerce_export_list(request_data.get('accessories'))
                if isinstance(item, dict) and any([
                    _pick_item_value(item, 'equipmentName', 'equipment_name', 'name'),
                    _pick_item_value(item, 'make', 'manufacturer'),
                    _pick_item_value(item, 'modelNo', 'model_no', 'model'),
                    _pick_item_value(item, 'serialNo', 'serial_no', 'serialNumber'),
                    _pick_item_value(item, 'functionPurpose', 'function_purpose', 'purpose', 'function'),
                ])
            ]
            _ensure_table_repeated_rows(doc, 3, 2, len(accessories))
            for idx, accessory in enumerate(accessories, start=2):
                row_idx = idx
                _set_table_cell_text(doc, 3, row_idx, 0, _pick_item_value(accessory, 'equipmentName', 'equipment_name', 'name'))
                _set_table_cell_text(doc, 3, row_idx, 1, _pick_item_value(accessory, 'make', 'manufacturer'))
                _set_table_cell_text(doc, 3, row_idx, 2, _pick_item_value(accessory, 'modelNo', 'model_no', 'model'))
                _set_table_cell_text(doc, 3, row_idx, 3, _pick_item_value(accessory, 'serialNo', 'serial_no', 'serialNumber'))
                _set_table_cell_text(doc, 3, row_idx, 4, _pick_item_value(accessory, 'functionPurpose', 'function_purpose', 'purpose', 'function'))

            cables = [
                item for item in _coerce_export_list(request_data.get('cables'))
                if isinstance(item, dict) and any([
                    _pick_item_value(item, 'cableName', 'cable_name', 'name'),
                    _pick_item_value(item, 'length', 'length_m'),
                    _pick_item_value(item, 'powerSignal', 'power_signal', 'type'),
                    _pick_item_value(item, 'shielded', 'shielded_unshielded', 'shielding'),
                    _pick_item_value(item, 'functionPurpose', 'function_purpose', 'purpose', 'function'),
                ])
            ]
            _ensure_table_repeated_rows(doc, 4, 2, len(cables))
            for idx, cable in enumerate(cables, start=2):
                row_idx = idx
                _set_table_cell_text(doc, 4, row_idx, 0, _pick_item_value(cable, 'cableName', 'cable_name', 'name'))
                _set_table_cell_text(doc, 4, row_idx, 1, _pick_item_value(cable, 'length', 'length_m'))
                _set_table_cell_text(doc, 4, row_idx, 2, _pick_item_value(cable, 'powerSignal', 'power_signal', 'type'))
                _set_table_cell_text(doc, 4, row_idx, 3, _pick_item_value(cable, 'shielded', 'shielded_unshielded', 'shielding'))
                _set_table_cell_text(doc, 4, row_idx, 4, _pick_item_value(cable, 'functionPurpose', 'function_purpose', 'purpose', 'function'))

            eut = (_coerce_export_list(request_data.get('eut_specs')) or [{}])[0]
            _set_table_cell_text(doc, 5, 2, 0, _pick_item_value(eut, 'acVoltageRange', 'ac_voltage_range'))
            _set_table_cell_text(doc, 5, 2, 1, _pick_item_value(eut, 'acFreqRange', 'freqRange', 'freq_range'))
            _set_table_cell_text(doc, 5, 2, 2, _pick_item_value(eut, 'acDcInputCurrent', 'acInputCurrent', 'inputCurrent', 'input_current', 'ac_input_current'))
            _set_table_cell_text(doc, 5, 2, 3, _pick_item_value(eut, 'dcVoltageRange', 'dc_voltage_range'))
            _set_table_cell_text(doc, 5, 2, 4, _pick_item_value(eut, 'dcInputCurrent', 'dc_input_current'))
            _set_table_cell_text(doc, 5, 2, 5, _pick_item_value(eut, 'ratedPower', 'rated_power'))

            supply_rows = [
                item for item in _coerce_export_list(request_data.get('supply_vf'))
                if isinstance(item, dict) and any([
                    _pick_item_value(item, 'supplyVoltage', 'voltage', 'supply_voltage'),
                    _pick_item_value(item, 'supplyFrequency', 'frequency', 'supply_frequency'),
                    _pick_item_value(item, 'remarks', 'notes', 'differentVoltageFrequency', 'different_voltage_frequency'),
                ])
            ]
            _ensure_table_repeated_rows(doc, 6, 2, len(supply_rows))
            for idx, supply in enumerate(supply_rows, start=2):
                _set_table_cell_text(doc, 6, idx, 0, _pick_item_value(supply, 'supplyVoltage', 'voltage', 'supply_voltage'))
                _set_table_cell_text(doc, 6, idx, 1, _pick_item_value(supply, 'supplyFrequency', 'frequency', 'supply_frequency'))
                _set_table_cell_text(doc, 6, idx, 2, _pick_item_value(supply, 'remarks', 'notes', 'differentVoltageFrequency', 'different_voltage_frequency'))

            has_wireless = _norm_yes_no(request_data.get('has_wireless_interface'))
            _set_table_cell_text(doc, 7, 1, 1, f"{_box('Yes', has_wireless == 'yes')}    {_box('No', has_wireless == 'no')}")
            wireless_rows = _coerce_export_list(request_data.get('wireless'))
            _set_table_cell_text(doc, 7, 3, 0, _join([_pick_item_value(item, 'type', 'interfaceType', 'interface_type') for item in wireless_rows]))
            _set_table_cell_text(doc, 7, 3, 1, _join([_pick_item_value(item, 'carrierFrequency', 'carrier_frequency') for item in wireless_rows]))
            _set_table_cell_text(doc, 7, 3, 2, _join([_pick_item_value(item, 'maxOutputPower', 'max_output_power') for item in wireless_rows]))
            _set_table_cell_text(doc, 7, 3, 3, _join([_pick_item_value(item, 'clockFrequencies', 'clock_frequencies') for item in wireless_rows]))
            _set_table_cell_text(doc, 7, 3, 4, _join([_pick_item_value(item, 'functionPurpose', 'function_purpose', 'purpose', 'function') for item in wireless_rows]))

            _insert_block_diagram_in_template(doc, request_data.get('block_diagram'))

            _set_following_paragraph_text(
                doc,
                'Brief Description about the EUT/its application/environment where it is used:',
                request_data.get('product_description')
            )
            _set_following_paragraph_text(
                doc,
                'Explain how the EUT configure during the test',
                request_data.get('test_configuration')
            )

            standards = _selected_option_set(request_data.get('product_standards'))
            _set_table_cell_text(doc, 8, 0, 1, '    '.join([
                _box('IEC 61326-1:2020', _selection_contains(standards, 'iec 61326-1')),
                _box('EN 61326-1:2021', _selection_contains(standards, 'en 61326-1')),
                _box('IEC 60601-1-2: 2014+A1:2020', _selection_contains(standards, 'iec 60601-1-2')),
                _box('EN 60601-1-2: 2015+A1:2021', _selection_contains(standards, 'en 60601-1-2')),
                _box('FCC Subpart 15B: 2024', _selection_contains(standards, 'fcc')),
                _box('ICES-001 Issue 5: 2020', _selection_contains(standards, 'ices')),
                _box('Other', _selection_contains(standards, 'other')),
            ]))
            env = _selected_option_set(request_data.get('product_environment'))
            _set_table_cell_text(doc, 8, 2, 1, '    '.join([
                _box('Basic Electromagnetic', _selection_contains(env, 'basic')),
                _box('Industrial', _selection_contains(env, 'industrial')),
                _box('Controlled', _selection_contains(env, 'controlled')),
                _box('NA', _selection_contains(env, 'na', 'n/a')),
            ]))
            _set_table_cell_text(doc, 8, 3, 1, '    '.join([
                _box('Home Healthcare', _selection_contains(env, 'home healthcare', 'home')),
                _box('Professional Healthcare', _selection_contains(env, 'professional healthcare', 'professional')),
            ]))
            grp = _selected_option_set(request_data.get('group'))
            _set_table_cell_text(doc, 8, 4, 1, _box('Group 1', _selection_contains(grp, 'group 1')))
            _set_table_cell_text(doc, 8, 4, 2, _box('Group 2', _selection_contains(grp, 'group 2')))
            cls = _selected_option_set(request_data.get('class_type'))
            _set_table_cell_text(doc, 8, 5, 1, _box('Class A', _selection_contains(cls, 'class a')))
            _set_table_cell_text(doc, 8, 5, 2, _box('Class B', _selection_contains(cls, 'class b')))

            selected_tests = set()
            selected_test_shortcuts = set()
            for source in [
                request_data.get('selected_tests'),
                request_data.get('selected_tests_for_development'),
                request_data.get('testSelected'),
                request_data.get('_selected_tests_db_raw'),
                request_data.get('_selected_tests_for_development_db_raw'),
            ]:
                if source not in (None, '', [], {}):
                    selected_tests.update(_selected_option_set(source))
                    selected_test_shortcuts.update(_extract_selected_test_shortcuts(source))
            test_hours = _parse_map(request_data.get('test_hours'))
            test_remarks = _parse_map(request_data.get('test_remarks'))
            if test_hours:
                selected_tests.update(_selected_option_set(list(test_hours.keys())))
                selected_test_shortcuts.update(_extract_selected_test_shortcuts(list(test_hours.keys())))

            ce_freq_range_display = _extract_frequency_range_display(
                request_data, 'ce_custom_spec', 'ce', 'ce_freq_range'
            )
            re_freq_range_display = _extract_frequency_range_display(
                request_data, 're_custom_spec', 're', 're_freq_range'
            )
            rs_freq_range_display = _extract_frequency_range_display(
                request_data, 'rs_ri_custom_spec', 'rs_ri', 'rs_freq_range'
            )
            rs_interim_freq_range_display = _extract_frequency_range_display(
                request_data, 'rs_ri_interim_custom_spec', 'rs_ri_interim', 'rs_interim_freq_range'
            )
            crf_freq_range_display = _extract_frequency_range_display(
                request_data, 'crf_custom_spec', 'crf', 'crf_freq_range'
            )

            test_blocks = [
                (0, 'Conducted Emission (CE)', ('CE',), [
                    _join([
                        ce_freq_range_display,
                        _line('Voltage/Frequency', request_data.get('ce_voltage_freq')),
                        _line('Cables', request_data.get('ce_cables')),
                        _line('Class', request_data.get('ce_class')),
                    ]),
                    _map_value(test_hours, 'CE', 'Conducted Emission'),
                    _map_value(test_remarks, 'CE', 'Conducted Emission'),
                ]),
                (4, 'Radiated Emission (RE)', ('RE',), [
                    _join([
                        re_freq_range_display,
                        _line('Voltage/Frequency', request_data.get('re_voltage_freq')),
                        _line('Class', request_data.get('re_class')),
                    ]),
                    _map_value(test_hours, 'RE', 'Radiated Emission'),
                    _map_value(test_remarks, 'RE', 'Radiated Emission'),
                ]),
                (8, 'Harmonic Current Emission', ('Harmonic', 'HARMONIC'), [_join([_line('Equipment Class', request_data.get('harmonic_class')), _line('Voltage/Frequency', request_data.get('harmonic_voltage_freq'))]), _map_value(test_hours, 'Harmonic', 'HARMONIC'), _map_value(test_remarks, 'Harmonic', 'HARMONIC')]),
                (12, 'Voltage Changes, Voltage Fluctuations and Flicker Emission', ('VoltageFlicker', 'Flicker'), [_join([_line('Specification', request_data.get('flicker_voltage_freq'))]), _map_value(test_hours, 'VoltageFlicker', 'Flicker'), _map_value(test_remarks, 'VoltageFlicker', 'Flicker')]),
                (16, 'Electrostatic Discharge Immunity', ('ESD',), [_join([_line('Voltage/Frequency', request_data.get('esd_voltage_freq')), _line('Contact', request_data.get('esd_contact')), _line('Air', request_data.get('esd_air')), _line('Indirect', request_data.get('esd_indirect'))]), _map_value(test_hours, 'ESD'), _map_value(test_remarks, 'ESD')]),
                (20, 'Radiated Susceptibility/Immunity', ('RS_RI', 'RS'), [_join([_line('Frequency Range', rs_freq_range_display)]), _join([_line('Field Strength 1', request_data.get('rs_field_strength1')), _line('Field Strength 2', request_data.get('rs_field_strength2')), _line('Field Strength 3', request_data.get('rs_field_strength3')), _line('Voltage/Frequency', request_data.get('rs_voltage_freq'))]), _map_value(test_hours, 'RS_RI', 'RS'), _map_value(test_remarks, 'RS_RI', 'RS')]),
                (25, 'Radiated Susceptibility/Immunity (Interim Method)', ('RS_RI_Interim', 'RS_Interim'), [_join([_line('Frequency Range', rs_interim_freq_range_display)]), _join([_line('Field Strength 1', request_data.get('rs_interim_field_strength1')), _line('Field Strength 2', request_data.get('rs_interim_field_strength2')), _line('Field Strength 3', request_data.get('rs_interim_field_strength3')), _line('Voltage/Frequency', request_data.get('rs_interim_voltage_freq'))]), _map_value(test_hours, 'RS_RI_Interim', 'RS_Interim'), _map_value(test_remarks, 'RS_RI_Interim', 'RS_Interim')]),
                (30, 'Electrical fast transient/burst Immunity', ('EFT',), [_join([_line('Voltage/Frequency', request_data.get('eft_voltage_freq')), _line('Test Level 1', request_data.get('eft_test_level1')), _line('Test Level 2', request_data.get('eft_test_level2'))]), _join([_line('Power lines', request_data.get('eft_cables_power')), _line('Signal lines', request_data.get('eft_cables_signal'))]), _join([_line('Signal Line 1', request_data.get('eft_signal_line1')), _line('Signal Line 2', request_data.get('eft_signal_line2'))]), _map_value(test_hours, 'EFT'), _map_value(test_remarks, 'EFT')]),
                (36, 'Surge Immunity', ('Surge', 'SURGE'), [_join([_line('Voltage/Frequency', request_data.get('surge_voltage_freq')), _line('Common Mode', request_data.get('surge_cm1')), _line('Differential Mode', request_data.get('surge_dm1'))]), _join([_line('Power lines', request_data.get('surge_cables_power')), _line('Signal lines', request_data.get('surge_cables_signal'))]), _join([_line('Signal Line 1', request_data.get('surge_signal_line1')), _line('Signal Line 2', request_data.get('surge_signal_line2'))]), _map_value(test_hours, 'Surge', 'SURGE'), _map_value(test_remarks, 'Surge', 'SURGE')]),
                (42, 'Conducted RF Disturbance Immunity', ('CRF',), [_join([_line('Voltage/Frequency', request_data.get('crf_voltage_freq')), _line('Test Level 1', request_data.get('crf_test_level1')), _line('Test Level 2', request_data.get('crf_test_level2'))]), _join([_line('Power lines', request_data.get('crf_cables_power')), _line('Signal lines', request_data.get('crf_cables_signal'))]), _join([_line('Signal Line 1', request_data.get('crf_signal_line1')), _line('Signal Line 2', request_data.get('crf_signal_line2'))]), _join([_line('Frequency Range', crf_freq_range_display)]), _map_value(test_hours, 'CRF'), _map_value(test_remarks, 'CRF')]),
                (49, 'Power Frequency Magnetic Field Immunity', ('PFMF', 'POWER'), [_join([_line('Voltage/Frequency', request_data.get('power_freq_voltage_freq')), _line('Field Strength', request_data.get('power_freq_test_level'))]), _map_value(test_hours, 'PFMF', 'POWER'), _map_value(test_remarks, 'PFMF', 'POWER')]),
                (53, 'Voltage Dips, Short Interruptions and Voltage Variations Immunity', ('VoltageDips', 'VOLTAGE'), [_join([_line('Voltage Dip 1', request_data.get('voltage_dips_voltage_dip1')), _line('Voltage Dip 2', request_data.get('voltage_dips_voltage_dip2')), _line('Voltage Dip 3', request_data.get('voltage_dips_voltage_dip3'))]), _join([_line('Voltage Variations', request_data.get('voltage_dips_voltage_freq')), _line('Time 1', request_data.get('voltage_dips_time1')), _line('Time 2', request_data.get('voltage_dips_time2'))]), _join([_line('Short Interruption', request_data.get('voltage_dips_interruption')), _line('Time 3', request_data.get('voltage_dips_time3')), _line('Time 4', request_data.get('voltage_dips_time4'))]), _map_value(test_hours, 'VoltageDips', 'VOLTAGE'), _map_value(test_remarks, 'VoltageDips', 'VOLTAGE')]),
            ]
            for start_row, label, aliases, values in test_blocks:
                selected = _test_selected(*aliases)
                _set_table_cell_text(doc, 9, start_row, 0, _box(label, selected))
                for offset, value in enumerate(values):
                    _set_table_cell_text(doc, 9, start_row + offset, 2, value if selected else '')

            for row_idx, field in enumerate(['continue_testing', 'test_report_required', 'uncertainty_required', 'test_witness'], start=0):
                choice = _norm_yes_no(request_data.get(field))
                _set_table_cell_text(doc, 10, row_idx, 1, _box('YES', choice == 'yes'))
                _set_table_cell_text(doc, 10, row_idx, 2, _box('YES', choice == 'yes'))
                _set_table_cell_text(doc, 10, row_idx, 3, _box('NO', choice == 'no'))
            conformity = _norm_yes_no(request_data.get('conformity_statement') or request_data.get('conformity_required'))
            _set_table_cell_text(doc, 10, 4, 1, _box('YES', conformity == 'yes'))
            _set_table_cell_text(doc, 10, 4, 2, _join(_coerce_export_list(request_data.get('decision_rule'))) or doc.tables[10].rows[4].cells[2].text)
            _set_table_cell_text(doc, 10, 4, 3, _box('NO', conformity == 'no'))
            _set_table_cell_text(doc, 10, 5, 1, _box('NO', conformity == 'no'))
            _set_table_cell_text(doc, 10, 5, 2, 'We will report the measured results and the uncertainty.' if conformity == 'no' else '')

            modes = _coerce_export_list(request_data.get('functional_modes'))
            _set_table_cell_text(doc, 11, 0, 1, request_data.get('number_of_modes'))
            _set_table_cell_text(doc, 11, 1, 1, '\n'.join(f"Mode {chr(65 + idx)}: {_coerce_export_text(mode)}" for idx, mode in enumerate(modes) if _coerce_export_text(mode)))
            _set_table_cell_text(doc, 11, 2, 1, request_data.get('monitoring_parameters'))
            _set_table_cell_text(doc, 11, 3, 1, request_data.get('additional_info'))

            _set_table_cell_text(doc, 12, 0, 1, request_data.get('requester_name'))
            _set_table_cell_text(doc, 12, 0, 3, _join([request_data.get('requester_department'), request_data.get('requester_division')], ' / '))
            _set_table_cell_text(doc, 12, 1, 1, request_data.get('requester_email'))
            _set_table_cell_text(doc, 12, 1, 3, request_data.get('requester_contact'))
            _set_table_cell_text(doc, 12, 2, 1, _fmt_date(request_data.get('requester_date')))
            _set_table_cell_text(doc, 12, 2, 3, _format_signature_export_value(
                request_data.get('requester_signature')))

            sample_condition = _coerce_export_text(request_data.get('sample_condition')).lower()
            _set_table_cell_text(doc, 13, 0, 1, request_data.get('job_number') or request_data.get('job_id'))
            _set_table_cell_text(doc, 13, 1, 1, f"{_box('Good', sample_condition == 'good')}    {_box('Bad', sample_condition == 'bad')}")
            cap = _norm_yes_no(request_data.get('capability_available'))
            _set_table_cell_text(doc, 13, 2, 1, f"{_box('YES', cap == 'yes')}    {_box('NO', cap == 'no')}")
            _set_table_cell_text(doc, 13, 3, 1, _fmt_date(request_data.get('sample_received_date')))
            _set_table_cell_text(doc, 13, 5, 2, _format_test_duration_days_only(request_data.get('test_duration')))
            _set_table_cell_text(doc, 13, 6, 2, _fmt_date(request_data.get('test_commencement_date')))
            _set_table_cell_text(doc, 13, 7, 2, _fmt_date(request_data.get('test_completion_date')))

            _set_table_cell_text(doc, 14, 0, 1, request_data.get('lab_manager_name'))
            _set_table_cell_text(doc, 14, 0, 3, _fmt_date(request_data.get('lab_manager_date')))
            _set_table_cell_text(doc, 14, 1, 1, _format_signature_export_value(
                request_data.get('lab_manager_signature')))

            output_stream = BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            return output_stream
        except Exception as exc:
            logger.exception('Rev1 TRF export failed for request %s', request_data.get('id'))
            logger.warning('Rev1 TRF export detail: %s', exc)
            return None

    def _build_test_request_word_export_with_fallback(
        request_data: dict,
        request_id=None,
    ):
        """Use the branded TRF export, falling back to a generic DOCX."""
        document_stream = _build_test_request_word_export_rev1(request_data)
        if document_stream is not None:
            return document_stream

        logger.warning(
            'Falling back to the generic test request DOCX export '
            '(request_id=%s) because the active TRF template is unavailable '
            'or could not be populated.',
            request_id if request_id is not None else request_data.get('id')
        )
        return _build_generic_test_request_word_export(request_data)

    @flask_app.route('/api/test-requests/<int:request_id>/download-form-docx', methods=['GET'])
    @login_required
    def download_test_request_form_docx(request_id):
        """Download a filled test request form as a Word document."""
        try:
            test_request = _get_request_or_404(request_id)
            if not test_request or not _can_access_iec_request(
                test_request,
                allow_lab_engineer=True,
                require_assigned_lab_engineer=True
            ):
                return jsonify({
                    'success': False,
                    'error': 'Test request not found or you do not have permission to view it'
                }), 404

            request_data = _build_request_payload(test_request, include_review_thread=False)
            request_data['_selected_tests_db_raw'] = getattr(
                test_request, 'selected_tests', None
            )
            request_data['_selected_tests_for_development_db_raw'] = getattr(
                test_request, 'selected_tests_for_development', None
            )

            document_stream = _build_test_request_word_export_with_fallback(
                request_data,
                request_id
            )
            filename_seed = test_request.tco_id or f"REQ-{test_request.id}"
            safe_filename_seed = re.sub(
                r'[^A-Za-z0-9._-]+', '_', filename_seed).strip('_')
            if not safe_filename_seed:
                safe_filename_seed = f"request_{test_request.id}"
            download_name = f"{safe_filename_seed}_test_request_form.docx"

            return send_file(
                document_stream,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=download_name
            )
        except Exception as e:
            logger.exception(
                'Error exporting test request form (request_id=%s): %s',
                request_id,
                e
            )
            return jsonify({
                'success': False,
                'error': 'Unable to export the test request form'
            }), 500

    @flask_app.route('/api/test-requests/tco/<tco_id>', methods=['GET'])
    @login_required
    def get_test_request_by_tco(tco_id):
        """Look up a test request by TCO number to PREFILL a new request form.

        Any logged-in user may look up any TCO, not just the person who raised it -
        engineers routinely start a new request from an existing TCO's product details.
        To keep that safe, this is a prefill payload only:
          * tco_id / job_number are stripped, so submitting creates fresh identifiers
            instead of writing into someone else's TCO;
          * requester fields are blanked, so they come from the logged-in user;
          * the review comment thread is never included.
        """
        try:
            normalized_tco_id = (tco_id or '').strip()
            if not normalized_tco_id:
                return jsonify({
                    'success': False,
                    'error': 'TCO ID is required'
                }), 400

            # Open lookup by TCO number (no per-user permission check - see docstring).
            test_request = get_request_by_tco_id(normalized_tco_id)

            if not test_request:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            request_data = _build_request_payload(
                test_request,
                include_review_thread=False
            )

            # Drop the identifiers so the form creates a NEW TCO/Job on submit.
            request_data.pop('tco_id', None)
            request_data.pop('job_number', None)

            # Clear requester details; the form fills these from the logged-in user.
            for _field in ('requester_name', 'requester_email', 'requester_phone',
                           'requester_contact', 'company_name'):
                if _field in request_data:
                    request_data[_field] = ''

            logger.debug(
                'Fetched request by TCO=%s (id=%s, selected_tests=%s)',
                normalized_tco_id,
                test_request.id,
                request_data.get('selected_tests')
            )

            return jsonify({
                'success': True,
                'data': request_data
            })
        except Exception as e:
            logger.exception(
                'Error fetching test request by TCO ID: %s',
                tco_id
            )
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @flask_app.route('/api/lab-engineers', methods=['GET'])
    @login_required
    def get_lab_engineers():
        """Get list of lab engineers."""
        try:
            if current_user.role not in ['admin', 'lab_engineer']:
                return jsonify({
                    'success': False,
                    'error': 'Unauthorized access'
                }), 403

            return jsonify({
                'success': True,
                'data': _serialize_peer_reviewer_candidates()
            })
        except Exception as e:
            logger.error(f'Error fetching lab engineers: {e}')
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    def generate_next_job_id():
        """Generate the next Job ID in format TFS-EMC-YYYY-001, TFS-EMC-YYYY-002, etc."""
        # Get current year
        current_year = datetime.now().year

        # Find the highest existing Job ID for current year
        job_id_pattern = f'TFS-EMC-{current_year}-%'

        last_job = EMCRequest.query.filter(
            EMCRequest.job_id.isnot(None),
            EMCRequest.job_id.like(job_id_pattern)
        ).order_by(EMCRequest.job_id.desc()).first()

        if last_job and last_job.job_id:
            try:
                # Extract the sequence number part from "TFS-EMC-2026-005"
                last_number = int(last_job.job_id.split('-')[-1])
                next_number = last_number + 1
            except (ValueError, IndexError):
                logger.warning(
                    f"Failed to parse job_id: {last_job.job_id}, starting from 1")
                next_number = 1
        else:
            # First job ID for this year
            next_number = 1

        # Format as TFS-EMC-YYYY-XXX (3 digits with leading zeros)
        new_job_id = f"TFS-EMC-{current_year}-{next_number:03d}"
        logger.info(f"Generated new Job ID: {new_job_id}")
        return new_job_id

    @flask_app.route('/api/test-requests/<int:request_id>/assign-tests', methods=['POST'])
    @login_required
    def assign_tests_to_engineers(request_id):
        """Assign individual tests to lab engineers with start and end dates."""
        try:
            # Check permissions
            if current_user.role not in ['admin', 'lab_engineer']:
                return jsonify({
                    'success': False,
                    'error': 'Unauthorized access'
                }), 403

            # Get test request
            test_request = _get_request_or_404(request_id)
            if not test_request:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            # Check if user has permission (lab engineer can only assign tests to requests assigned to them)
            if current_user.role == 'lab_engineer' and test_request.assigned_engineer_id != current_user.id:
                return jsonify({
                    'success': False,
                    'error': 'You can only assign tests to requests assigned to you'
                }), 403

            # Get request data
            data = request.get_json()
            if not data or 'test_assignments' not in data:
                return jsonify({
                    'success': False,
                    'error': 'No test assignments provided'
                }), 400

            test_assignments = data.get('test_assignments', [])
            if not test_assignments:
                return jsonify({
                    'success': False,
                    'error': 'At least one test assignment is required'
                }), 400

            merge_existing_assignments = bool(
                data.get('merge_existing_assignments'))
            laboratory_details_raw = data.get('laboratory_details')

            def _clean_optional_text(value, max_length=None):
                if value is None:
                    return None
                cleaned = str(value).strip()
                if not cleaned:
                    return None
                if max_length:
                    return cleaned[:max_length]
                return cleaned

            normalized_laboratory_details = None
            has_laboratory_details = isinstance(laboratory_details_raw, dict) and any(
                _clean_optional_text(value) for value in laboratory_details_raw.values()
            )
            if has_laboratory_details:
                sample_condition_value = _clean_optional_text(
                    laboratory_details_raw.get('sample_condition'), 200)
                if not sample_condition_value:
                    return jsonify({
                        'success': False,
                        'error': 'Sample condition is required'
                    }), 400
                sample_condition_value = sample_condition_value.lower()
                if sample_condition_value not in ('good', 'bad'):
                    return jsonify({
                        'success': False,
                        'error': 'Sample condition must be either "good" or "bad"'
                    }), 400

                capability_available_value = _clean_optional_text(
                    laboratory_details_raw.get('capability_available'), 200)
                if not capability_available_value:
                    return jsonify({
                        'success': False,
                        'error': 'Capability availability is required'
                    }), 400
                capability_available_value = capability_available_value.lower()
                if capability_available_value not in ('yes', 'no'):
                    return jsonify({
                        'success': False,
                        'error': 'Capability availability must be either "yes" or "no"'
                    }), 400

                sample_received_date_raw = _clean_optional_text(
                    laboratory_details_raw.get('sample_received_date'))
                if not sample_received_date_raw:
                    return jsonify({
                        'success': False,
                        'error': 'Sample received date is required'
                    }), 400
                sample_received_date_value = _parse_iso_date(sample_received_date_raw)
                if not sample_received_date_value:
                    return jsonify({
                        'success': False,
                        'error': 'Invalid sample received date format'
                    }), 400

                test_commencement_date_raw = _clean_optional_text(
                    laboratory_details_raw.get('test_commencement_date'))
                if not test_commencement_date_raw:
                    return jsonify({
                        'success': False,
                        'error': 'Test commencement date is required'
                    }), 400
                test_commencement_date_value = _parse_iso_date(test_commencement_date_raw)
                if not test_commencement_date_value:
                    return jsonify({
                        'success': False,
                        'error': 'Invalid test commencement date format'
                    }), 400

                test_completion_date_raw = _clean_optional_text(
                    laboratory_details_raw.get('test_completion_date'))
                if not test_completion_date_raw:
                    return jsonify({
                        'success': False,
                        'error': 'Test completion date is required'
                    }), 400
                test_completion_date_value = _parse_iso_date(test_completion_date_raw)
                if not test_completion_date_value:
                    return jsonify({
                        'success': False,
                        'error': 'Invalid test completion date format'
                    }), 400

                if test_commencement_date_value > test_completion_date_value:
                    return jsonify({
                        'success': False,
                        'error': 'Test completion date cannot be earlier than test commencement date'
                    }), 400

                test_duration_raw = _clean_optional_text(
                    laboratory_details_raw.get('test_duration'), 100)
                if not test_duration_raw:
                    return jsonify({
                        'success': False,
                        'error': 'Test duration is required'
                    }), 400
                try:
                    duration_number = float(test_duration_raw)
                    if duration_number <= 0:
                        raise ValueError()
                    if duration_number.is_integer():
                        test_duration_raw = str(int(duration_number))
                    else:
                        test_duration_raw = str(round(duration_number, 2))
                except (TypeError, ValueError):
                    return jsonify({
                        'success': False,
                        'error': 'Test duration must be a positive number'
                    }), 400

                normalized_laboratory_details = {
                    'sample_condition': sample_condition_value,
                    'capability_available': capability_available_value,
                    'sample_received_date': sample_received_date_value,
                    'test_duration': test_duration_raw,
                    'test_commencement_date': test_commencement_date_value,
                    'test_completion_date': test_completion_date_value
                }

            ensure_planner_table()
            ensure_equipment_test_name_column()

            tco_display = test_request.tco_id or f'REQ-{test_request.id}'
            engineer_cache: dict[object, User | None] = {}
            existing_serialized_assignments = []
            existing_assignment_keys = set()
            existing_schedule_records = []
            fallback_existing_schedule_snapshots = []
            serialized_assignments = []
            planner_records = []
            notification_payloads = []
            submitted_assignment_keys = set()
            submitted_schedule_snapshots = []

            def _resolve_assignable_engineer(engineer_id, engineer_name=None):
                normalized_engineer_id = None
                if engineer_id not in (None, '', []):
                    try:
                        normalized_engineer_id = int(engineer_id)
                    except (TypeError, ValueError):
                        normalized_engineer_id = None

                if normalized_engineer_id is not None and normalized_engineer_id not in engineer_cache:
                    engineer_cache[normalized_engineer_id] = User.query.filter(
                        User.id == normalized_engineer_id,
                        User.role.in_(['lab_engineer', 'admin']),
                        User.is_active.is_(True)
                    ).first()

                engineer_user = (
                    engineer_cache.get(normalized_engineer_id)
                    if normalized_engineer_id is not None else None
                )
                normalized_engineer_name = str(engineer_name or '').strip()

                if not engineer_user and normalized_engineer_name:
                    engineer_name_cache_key = f'username:{normalized_engineer_name.casefold()}'
                    if engineer_name_cache_key not in engineer_cache:
                        engineer_cache[engineer_name_cache_key] = User.query.filter(
                            User.username == normalized_engineer_name,
                            User.role.in_(['lab_engineer', 'admin']),
                            User.is_active.is_(True)
                        ).first()
                    engineer_user = engineer_cache.get(engineer_name_cache_key)

                    if engineer_user and normalized_engineer_id is not None and engineer_user.id != normalized_engineer_id:
                        logger.warning(
                            'Remapped stale engineer id %s to active user %s (%s) while assigning request %s.',
                            engineer_id,
                            engineer_user.username,
                            engineer_user.id,
                            request_id
                        )

                return engineer_user

            def _record_schedule_snapshot(start_date, end_date, total_hours):
                existing_schedule_records.append({
                    'start_date': start_date,
                    'end_date': end_date,
                    'total_hours': total_hours
                })

            def _remember_existing_assignment(payload):
                if not isinstance(payload, dict):
                    return

                test_name = str(payload.get('test_name') or '').strip()
                engineer_id = payload.get('engineer_id')
                engineer_name = str(payload.get('engineer_name') or '').strip()
                start_date = str(payload.get('start_date') or '').strip()
                end_date = str(payload.get('end_date') or '').strip()
                start_time = str(payload.get('start_time') or '').strip()
                end_time = str(payload.get('end_time') or '').strip()

                engineer_user = _resolve_assignable_engineer(engineer_id, engineer_name)
                if engineer_user:
                    engineer_id = engineer_user.id
                    engineer_name = engineer_user.username
                else:
                    logger.warning(
                        'Ignoring stale existing assignment for request %s test "%s" because engineer id=%s name="%s" is no longer valid.',
                        request_id,
                        test_name,
                        engineer_id,
                        engineer_name
                    )
                    engineer_id = None

                if not all([test_name, engineer_id, engineer_name, start_date, end_date]):
                    return

                total_hours_raw = payload.get('total_hours')
                try:
                    total_hours = round(float(total_hours_raw), 2) if total_hours_raw not in (None, '', []) else None
                except (TypeError, ValueError):
                    total_hours = None

                normalized_payload = {
                    'test_name': test_name,
                    'engineer_id': engineer_id,
                    'engineer_name': engineer_name,
                    'start_date': start_date,
                    'end_date': end_date,
                    'start_time': start_time or None,
                    'end_time': end_time or None,
                    'total_hours': total_hours
                }
                existing_serialized_assignments.append(normalized_payload)

                assignment_key = _normalize_assignment_test_key(test_name)
                if assignment_key:
                    existing_assignment_keys.add(assignment_key)

                try:
                    start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
                    end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
                except ValueError:
                    return

                _record_schedule_snapshot(start_date_obj, end_date_obj, total_hours)

            if merge_existing_assignments:
                existing_planner_entries = PlannerEntry.query.filter(
                    db.or_(
                        PlannerEntry.test_request_id == test_request.id,
                        PlannerEntry.tco_id == tco_display
                    )
                ).order_by(
                    PlannerEntry.start_date.asc(),
                    PlannerEntry.start_time.asc()
                ).all()

                if existing_planner_entries:
                    for entry in existing_planner_entries:
                        _remember_existing_assignment({
                            'test_name': entry.test_name,
                            'engineer_id': entry.engineer_user_id,
                            'engineer_name': entry.test_person_name,
                            'start_date': entry.start_date.isoformat() if entry.start_date else '',
                            'end_date': entry.end_date.isoformat() if entry.end_date else '',
                            'start_time': entry.start_time.strftime('%H:%M') if entry.start_time else '',
                            'end_time': entry.end_time.strftime('%H:%M') if entry.end_time else '',
                            'total_hours': entry.total_hours
                        })
                elif test_request.test_assignments:
                    try:
                        raw_existing_assignments = json.loads(
                            test_request.test_assignments)
                    except (TypeError, ValueError):
                        raw_existing_assignments = []

                    for existing_assignment in raw_existing_assignments:
                        _remember_existing_assignment(existing_assignment)
                        fallback_existing_schedule_snapshots.append(
                            _build_planner_conflict_snapshot(existing_assignment)
                        )

            for assignment in test_assignments:
                missing_keys = [
                    key for key in ['test_name', 'engineer_id',
                                    'engineer_name', 'start_date', 'end_date', 'start_time', 'end_time']
                    if key not in assignment or assignment[key] in (None, '')
                ]
                if missing_keys:
                    return jsonify({
                        'success': False,
                        'error': f"Missing required fields ({', '.join(missing_keys)}) for one of the tests."
                    }), 400

                test_name = assignment['test_name'].strip()
                engineer_id = assignment['engineer_id']
                assignment_key = _normalize_assignment_test_key(test_name)

                if assignment_key:
                    if assignment_key in submitted_assignment_keys:
                        return jsonify({
                            'success': False,
                            'error': f'Duplicate test assignment submitted for test: {test_name}'
                        }), 400
                    if merge_existing_assignments and assignment_key in existing_assignment_keys:
                        return jsonify({
                            'success': False,
                            'error': f'Test "{test_name}" is already assigned. Select only the remaining tests.'
                        }), 400
                    submitted_assignment_keys.add(assignment_key)

                engineer_user = _resolve_assignable_engineer(
                    engineer_id,
                    assignment.get('engineer_name')
                )
                if not engineer_user:
                    return jsonify({
                        'success': False,
                        'error': f'Engineer with ID {engineer_id} not found'
                    }), 400

                engineer_id = engineer_user.id

                # Validate dates
                try:
                    start_date_obj = datetime.strptime(
                        assignment['start_date'], '%Y-%m-%d').date()
                    end_date_obj = datetime.strptime(
                        assignment['end_date'], '%Y-%m-%d').date()
                except ValueError:
                    return jsonify({
                        'success': False,
                        'error': f'Invalid date format for test: {test_name}'
                    }), 400

                if start_date_obj > end_date_obj:
                    return jsonify({
                        'success': False,
                        'error': f'End date must be after start date for test: {test_name}'
                    }), 400

                start_time_str = assignment.get('start_time')
                end_time_str = assignment.get('end_time')
                start_time_obj = _parse_iso_time(start_time_str)
                end_time_obj = _parse_iso_time(end_time_str)

                if not start_time_str or not start_time_obj:
                    return jsonify({
                        'success': False,
                        'error': f'Invalid start time for test: {test_name}'
                    }), 400
                if not end_time_str or not end_time_obj:
                    return jsonify({
                        'success': False,
                        'error': f'Invalid end time for test: {test_name}'
                    }), 400

                total_hours_val = assignment.get('total_hours')
                total_hours_float = None
                if total_hours_val not in (None, '', []):
                    try:
                        total_hours_float = round(float(total_hours_val), 2)
                    except (ValueError, TypeError):
                        return jsonify({
                            'success': False,
                            'error': f'Invalid total hours value for test: {test_name}'
                        }), 400

                if total_hours_float is None and start_time_obj and end_time_obj:
                    diff_hours = _calculate_planner_schedule_hours(
                        start_date_obj,
                        end_date_obj,
                        start_time_obj,
                        end_time_obj
                    )
                    if diff_hours > 0:
                        total_hours_float = round(diff_hours, 2)

                candidate_snapshot = _build_planner_conflict_snapshot({
                    'test_request_id': test_request.id,
                    'test_person_name': engineer_user.username,
                    'engineer_user_id': engineer_id,
                    'test_name': test_name,
                    'tco_id': tco_display,
                    'start_date': start_date_obj,
                    'end_date': end_date_obj,
                    'start_time': start_time_obj,
                    'end_time': end_time_obj,
                    'event_type': 'test',
                    'status': 'scheduled'
                })

                db_conflicts = _find_db_schedule_conflicts(
                    candidate_snapshot,
                    ignore_request_id=None if merge_existing_assignments else test_request.id,
                    ignore_tco_id=None if merge_existing_assignments else tco_display
                )
                if _has_schedule_conflicts(db_conflicts):
                    return jsonify({
                        'success': False,
                        'error': _format_schedule_conflict_message(
                            test_name,
                            engineer_user.username,
                            db_conflicts
                        )
                    }), 409

                in_memory_conflicts = _find_schedule_conflicts_in_snapshots(
                    fallback_existing_schedule_snapshots + submitted_schedule_snapshots,
                    candidate_snapshot
                )
                if _has_schedule_conflicts(in_memory_conflicts):
                    return jsonify({
                        'success': False,
                        'error': _format_schedule_conflict_message(
                            test_name,
                            engineer_user.username,
                            in_memory_conflicts
                        )
                    }), 409

                serialized_assignment = {
                    'test_name': test_name,
                    'engineer_id': engineer_id,
                    'engineer_name': engineer_user.username,
                    'start_date': start_date_obj.isoformat(),
                    'end_date': end_date_obj.isoformat(),
                    'start_time': start_time_obj.strftime('%H:%M') if start_time_obj else None,
                    'end_time': end_time_obj.strftime('%H:%M') if end_time_obj else None,
                    'total_hours': total_hours_float
                }
                serialized_assignments.append(serialized_assignment)

                planner_records.append({
                    'test_request_id': test_request.id,
                    'test_person_name': engineer_user.username,
                    'engineer_user_id': engineer_id,
                    'created_by_user_id': current_user.id,
                    'test_name': test_name,
                    'tco_id': tco_display,
                    'start_date': start_date_obj,
                    'end_date': end_date_obj,
                    'start_time': start_time_obj,
                    'end_time': end_time_obj,
                    'total_hours': total_hours_float
                })

                notification_payloads.append({
                    'engineer': engineer_user,
                    'assignment': serialized_assignment
                })
                submitted_schedule_snapshots.append(candidate_snapshot)

            combined_schedule_records = existing_schedule_records + planner_records
            earliest_start_date = min(
                (record['start_date'] for record in combined_schedule_records),
                default=None
            )
            latest_end_date = max(
                (record['end_date'] for record in combined_schedule_records),
                default=None
            )
            total_assigned_hours = round(sum(
                record.get('total_hours') or 0
                for record in combined_schedule_records
            ), 2)

            if total_assigned_hours > 0:
                if float(total_assigned_hours).is_integer():
                    computed_test_duration = str(int(total_assigned_hours))
                else:
                    computed_test_duration = f'{total_assigned_hours:.2f}'.rstrip(
                        '0').rstrip('.')
            else:
                computed_test_duration = None

            request_received_dt = test_request.submitted_at or test_request.created_at
            if isinstance(request_received_dt, datetime):
                computed_sample_received_date = request_received_dt.date()
            else:
                computed_sample_received_date = request_received_dt

            combined_serialized_assignments = (
                existing_serialized_assignments + serialized_assignments
                if merge_existing_assignments else serialized_assignments
            )
            test_request.test_assignments = json.dumps(
                combined_serialized_assignments)

            if not merge_existing_assignments:
                test_request.review_comments = None
                test_request.reviewed_by = None
                test_request.reviewed_at = None

            if normalized_laboratory_details is not None:
                test_request.sample_condition = normalized_laboratory_details['sample_condition']
                test_request.capability_available = normalized_laboratory_details['capability_available']
                test_request.sample_received_date = normalized_laboratory_details['sample_received_date']
                test_request.test_duration = normalized_laboratory_details['test_duration']
                test_request.test_commencement_date = normalized_laboratory_details['test_commencement_date']
                test_request.test_completion_date = normalized_laboratory_details['test_completion_date']
                test_request.lab_manager_name = None
                test_request.lab_manager_date = None
                test_request.lab_manager_signature = None
                test_request.lab_manager_signed_at = None
            else:
                if computed_test_duration:
                    test_request.test_duration = computed_test_duration
                if earliest_start_date:
                    test_request.test_commencement_date = earliest_start_date
                if latest_end_date:
                    test_request.test_completion_date = latest_end_date

            # Fall back to computed values only when form values are missing.
            if not test_request.sample_received_date and computed_sample_received_date:
                test_request.sample_received_date = computed_sample_received_date
            if not test_request.test_duration and computed_test_duration:
                test_request.test_duration = computed_test_duration
            if not test_request.test_commencement_date and earliest_start_date:
                test_request.test_commencement_date = earliest_start_date
            if not test_request.test_completion_date and latest_end_date:
                test_request.test_completion_date = latest_end_date

            if merge_existing_assignments:
                if not test_request.status:
                    test_request.status = 'Assigned Lab Engineer'
            else:
                test_request.status = 'Test Plan To Approve'
            test_request.updated_at = get_ist_now()

            if not merge_existing_assignments:
                PlannerEntry.query.filter(
                    PlannerEntry.test_request_id == test_request.id
                ).delete(synchronize_session=False)

                # Also delete any orphaned entries with matching TCO but no request_id
                PlannerEntry.query.filter(
                    PlannerEntry.test_request_id.is_(None),
                    PlannerEntry.tco_id == tco_display
                ).delete(synchronize_session=False)
            else:
                # âœ… FIX: When merging assignments, cancel old entries for tests being rescheduled
                # This prevents duplicate entries and allows datasheet uploads
                test_names_being_added = set()
                for record in planner_records:
                    if record.get('test_name'):
                        test_names_being_added.add(record['test_name'].strip())
                
                if test_names_being_added:
                    old_entries_to_cancel = PlannerEntry.query.filter(
                        db.or_(
                            PlannerEntry.test_request_id == test_request.id,
                            PlannerEntry.tco_id == tco_display
                        ),
                        PlannerEntry.test_name.in_(list(test_names_being_added)),
                        PlannerEntry.status != 'cancelled'
                    ).all()
                    
                    for old_entry in old_entries_to_cancel:
                        old_entry.status = 'cancelled'
                        old_entry.cancel_reason = 'Rescheduled with new dates'
                        old_entry.cancelled_at = get_ist_now()
                        old_entry.cancelled_by = current_user.id
                        logger.info(
                            f"Cancelled old entry {old_entry.id} for test '{old_entry.test_name}' "
                            f"as it is being rescheduled for request {test_request.id}"
                        )

            # Insert new planner entries directly to database
            for record in planner_records:
                db.session.add(PlannerEntry(**record))

            if merge_existing_assignments:
                logger.info(
                    'Successfully appended %s planner entries for request %s.',
                    len(planner_records),
                    request_id
                )
            else:
                logger.info(
                    f'Successfully added {len(planner_records)} planner entries for request {request_id}. '
                    f'Status updated to "Test Plan To Approve".'
                )

            # Commit all changes to database
            db.session.commit()

            if merge_existing_assignments:
                logger.info(
                    'Additional tests scheduled for request %s by %s',
                    request_id,
                    current_user.username
                )
            else:
                logger.info(
                    f'Tests scheduled and submitted for admin approval for request {request_id} by {current_user.username}')

            try:
                admin_emails = [
                    admin.email for admin in User.query.filter_by(role='admin').all()
                    if admin.email
                ]
                # Include the assigner in CC
                if current_user.email and current_user.email not in admin_emails:
                    admin_emails.append(current_user.email)

                send_test_assignment_emails(
                    test_request=test_request,
                    assignment_payloads=notification_payloads,
                    admin_emails=admin_emails,
                    status_label='Assigned Lab Engineer' if merge_existing_assignments else 'Test Plan To Approve',
                    send_to_admins_first=not merge_existing_assignments
                )
            except Exception as email_exc:
                logger.error(
                    'Failed to send test assignment emails: %s', email_exc)

            return jsonify({
                'success': True,
                'message': 'Additional tests assigned successfully' if merge_existing_assignments else 'Tests assigned successfully and sent for admin approval',
                'data': {
                    'id': test_request.id,
                    'status': test_request.status,
                    'job_id': test_request.job_id,
                    'test_assignments': combined_serialized_assignments,
                    'planner_entries_count': len(planner_records)
                }
            })
        except Exception as e:
            db.session.rollback()
            import traceback
            error_trace = traceback.format_exc()
            logger.error(
                f'Error assigning tests to request {request_id}: {e}\n{error_trace}')
            return jsonify({
                'success': False,
                'error': f'Error assigning tests: {str(e)}'
            }), 500

    @flask_app.route('/api/test-requests/<int:request_id>/review', methods=['POST'])
    @login_required
    def submit_review(request_id):
        """Submit review for a test request (approve, reject, or request more information)."""
        try:
            # Check permissions
            if current_user.role not in ['admin', 'lab_engineer']:
                return jsonify({
                    'success': False,
                    'error': 'Unauthorized access'
                }), 403

            # Get test request
            test_request = _get_request_or_404(request_id)
            if not test_request:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            # Check if user has permission (lab engineer can only review requests assigned to them)
            if current_user.role == 'lab_engineer' and test_request.assigned_engineer_id != current_user.id:
                return jsonify({
                    'success': False,
                    'error': 'You can only review requests assigned to you'
                }), 403

            # Get request data
            data = request.get_json()
            if not data:
                return jsonify({
                    'success': False,
                    'error': 'No data provided'
                }), 400

            action = data.get('action')
            comments = data.get('comments', '').strip()
            action_mode_raw = (data.get('action_mode') or '').strip().lower()
            action_mode = 'comments' if action_mode_raw in [
                'comments', 'comment', 'add_comment'] else 'need_more_info'
            original_status = test_request.status
            should_send_more_info_email = False

            if not action:
                return jsonify({
                    'success': False,
                    'error': 'Action is required'
                }), 400

            # Validate comments for non-approve actions
            if action != 'approve' and not comments:
                return jsonify({
                    'success': False,
                    'error': 'Comments are required for this action'
                }), 400

            # Update status and comments based on action
            if action == 'more_info':
                _append_review_comment_entry(
                    test_request=test_request,
                    comment=comments,
                    username=current_user.username,
                    role=current_user.role
                )
                # Only "Submit Review" should move status to Need More Information.
                if action_mode != 'comments':
                    test_request.status = 'Need More Information'
                    should_send_more_info_email = True
            elif action == 'reject':
                test_request.status = 'Rejected'
                _append_review_comment_entry(
                    test_request=test_request,
                    comment=comments,
                    username=current_user.username,
                    role=current_user.role
                )
                # Also set rejection fields for consistency
                test_request.rejection_reason = comments
                test_request.rejected_by = current_user.username
                test_request.rejected_at = get_ist_now()
            elif action == 'approve':
                # Approval might need different handling, but for now just update comments if provided
                if comments:
                    _append_review_comment_entry(
                        test_request=test_request,
                        comment=comments,
                        username=current_user.username,
                        role=current_user.role
                    )
                # Status might be updated elsewhere for approval (e.g., through assign endpoint)
            else:
                return jsonify({
                    'success': False,
                    'error': 'Invalid action'
                }), 400

            # Update timestamps
            test_request.updated_at = get_ist_now()

            db.session.commit()

            logger.info(
                f'Review submitted for request {request_id} by {current_user.username}: '
                f'action={action}, action_mode={action_mode}, original_status={original_status}, '
                f'new_status={test_request.status}'
            )

            # Send email only for true "Need More Information" review submissions.
            if action == 'more_info' and should_send_more_info_email:
                try:
                    send_more_info_notification(
                        test_request=test_request,
                        comments=comments,
                        reviewed_by=current_user.username
                    )
                except Exception as email_error:
                    logger.error(
                        f'Failed to send more info notification email: {email_error}')
                    # Don't fail the review if email fails

            return jsonify({
                'success': True,
                'message': 'Review submitted successfully',
                'data': {
                    'id': test_request.id,
                    'status': test_request.status,
                    'action': action,
                    'reviewed_by': test_request.reviewed_by,
                    'reviewed_at': test_request.reviewed_at.isoformat() if test_request.reviewed_at else None
                }
            })
        except Exception as e:
            db.session.rollback()
            import traceback
            error_trace = traceback.format_exc()
            logger.error(
                f'Error submitting review for request {request_id}: {e}\n{error_trace}')
            return jsonify({
                'success': False,
                'error': f'Error submitting review: {str(e)}'
            }), 500

    def process_test_configurations(selected_tests, test_hours, test_configurations):
        """Process test configurations into structured data."""
        processed_data = {
            'tests': []
        }

        # Add logging
        logger.info("=" * 80)
        logger.info("Processing test configurations")
        logger.info(f"Selected tests: {selected_tests}")
        logger.info(f"Test configurations received: {test_configurations}")
        logger.info("=" * 80)

        for test_name in selected_tests:
            test_info = {
                'name': test_name,
                'hours': test_hours.get(test_name, ''),
                'configuration': {}
            }
            config = test_configurations.get(test_name, {})

            # Log each test's config
            logger.info(f"Processing {test_name} config: {config}")

            if test_name == 'CE':
                test_info['configuration'] = {
                    'spec_type': config.get('spec', 'As per the standard'),
                    'custom_range': config.get('customSpecRange'),
                    'cables': {
                        'power': config.get('cables', {}).get('power', ''),
                        'power_count': config.get('cablesPowerCount', ''),
                        'signal': config.get('cables', {}).get('signal', ''),
                        'signal_count': config.get('cablesSignalCount', '')
                    },
                    'signal_line_types': config.get('signalLineTypes', [])
                }
            elif test_name == 'RE':
                test_info['configuration'] = {
                    'spec_type': config.get('spec', 'As per the standard'),
                    'custom_range': config.get('customSpecRange')
                }
            elif test_name == 'Harmonic':
                test_info['configuration'] = {
                    'equipment_class': config.get('class', '')
                }
            elif test_name == 'VoltageFlicker':  # FIXED: Keep consistent field names
                spec_type = config.get('specificationType', config.get(
                    'specification_type', 'standard'))
                custom_spec = config.get(
                    'customSpecification', config.get('custom_specification', ''))

                test_info['configuration'] = {
                    'specificationType': spec_type,  # Keep original field name
                    'customSpecification': custom_spec  # Keep original field name
                }

                # Add detailed logging for Voltage Flicker
                logger.info(f"VoltageFlicker configuration processed:")
                logger.info(f"  - specificationType: {spec_type}")
                logger.info(f"  - customSpecification: {custom_spec}")

            elif test_name == 'ESD':  # Electrostatic Discharge
                test_info['configuration'] = {
                    'spec_type': config.get('spec', 'As per the standard'),
                    'custom_contact': config.get('customContact', ''),
                    'custom_contact_kv': config.get('customContactKV', ''),
                    'custom_air': config.get('customAir', ''),
                    'custom_air_kv': config.get('customAirKV', ''),
                    'custom_indirect': config.get('customIndirect', ''),
                    'custom_indirect_kv': config.get('customIndirectKV', '')
                }

            elif test_name == 'RS_RI':  # Radiated Susceptibility/Immunity
                test_info['configuration'] = {
                    'frequency_spec': config.get('frequency', 'As per the standard'),
                    'custom_frequency_range': config.get('customSpecRange'),
                    'test_level': config.get('testLevel', 'As per the standard'),
                    'custom_field_strength_vm': config.get('testLevelCustomVm', '')
                }

            elif test_name == 'RS_RI_Interim':
                test_info['configuration'] = {
                    'frequency_spec': config.get('frequency', 'As per the standard'),
                    'custom_frequency_range': config.get('customSpecRange'),
                    'test_level': config.get('testLevel', 'As per the standard'),
                    'custom_field_strength_vm': config.get('testLevelCustomVm', '')
                }

            elif test_name == 'EFT':  # Electrical Fast Transient
                test_info['configuration'] = {
                    'cables': {
                        'power': config.get('cables', {}).get('power', ''),
                        'power_count': config.get('cablesPowerCount', ''),
                        'signal': config.get('cables', {}).get('signal', ''),
                        'signal_count': config.get('cablesSignalCount', '')
                    },
                    'test_level': config.get('testLevel', 'As per the standard'),
                    'custom_voltage_kv': config.get('testLevelCustomKv', '')
                }

            elif test_name == 'Surge':  # Surge Immunity
                test_info['configuration'] = {
                    'cables': {
                        'power': config.get('cables', {}).get('power', ''),
                        'power_count': config.get('cablesPowerCount', ''),
                        'signal': config.get('cables', {}).get('signal', ''),
                        'signal_count': config.get('cablesSignalCount', '')
                    },
                    'test_level': config.get('testLevel', 'As per the standard'),
                    'custom_common': config.get('customCommon', ''),
                    'custom_common_kv': config.get('customCommonKV', ''),
                    'custom_differential': config.get('customDifferential', ''),
                    'custom_differential_kv': config.get('customDifferentialKV', '')
                }

            elif test_name == 'CRF':  # Conducted RF Disturbance
                test_info['configuration'] = {
                    'frequency_spec': config.get('frequency', 'As per the standard'),
                    'custom_frequency_range': config.get('customSpecRange'),
                    'cables': {
                        'power': config.get('cables', {}).get('power', ''),
                        'power_count': config.get('cablesPowerCount', ''),
                        'signal': config.get('cables', {}).get('signal', ''),
                        'signal_count': config.get('cablesSignalCount', '')
                    },
                    'test_level': config.get('testLevel', 'As per the standard'),
                    'custom_field_strength_vrms': config.get('testLevelCustomVrms', '')
                }

            elif test_name == 'PFMF':  # Power Frequency Magnetic Field
                test_info['configuration'] = {
                    'test_level': config.get('testLevel', 'As per the standard'),
                    'custom_field_strength_am': config.get('testLevelCustomAm', '')
                }

            elif test_name == 'VoltageDips':  # Voltage Dips
                test_info['configuration'] = {
                    'voltage_dip': config.get('voltageDip', 'As per the standard'),
                    'voltage_dip_custom': config.get('voltageDipCustom', ''),
                    'voltage_variations': config.get('voltageVariations', 'As per the standard'),
                    'voltage_variations_custom': config.get('voltageVariationsCustom', ''),
                    'short_interruption': config.get('shortInterruption', 'As per the standard'),
                    'short_interruption_custom': config.get('shortInterruptionCustom', '')
                }

            processed_data['tests'].append(test_info)

        logger.info(f"Processed data: {processed_data}")
        logger.info("=" * 80)

        return processed_data

    @flask_app.route('/create-test-plan', methods=['POST'])
    @login_required
    def create_test_plan():
        """Handle test plan creation from form data."""
        try:
            def safe_file_size(path: str) -> int:
                """Return file size safely without failing request flow."""
                try:
                    return os.path.getsize(path)
                except OSError as exc:
                    logger.warning('Unable to read file size for %s: %s', path, exc)
                    return 0

            # Get JSON data from request
            data = request.get_json()
            if not data:
                return jsonify({
                    'success': False,
                    'error': 'No data provided'
                }), 400

            form_data = data.get('form_data', {})
            draft_id = data.get('draft_id')
            selected_equipment = data.get('selected_equipment', [])

            # NEW: Get TCO update parameters
            tco_id_from_request = data.get('tco_id')
            is_update = data.get('is_update', False)

            logger.info(f"=== CREATE/UPDATE TEST PLAN ===")
            logger.info(f"User: {current_user.id}")
            logger.info(f"Draft ID: {draft_id}")
            logger.info(f"TCO ID from request: {tco_id_from_request}")
            logger.info(f"Is Update: {is_update}")

            # Validate required form fields
            required_fields = ['productName', 'manufacturer',
                               'modelNumber', 'testSamples', 'operatingFrequency',
                               'requesterDepartment']
            missing_fields = []
            for field in required_fields:
                if not form_data.get(field, '').strip():
                    missing_fields.append(field)

            if missing_fields:
                return jsonify({
                    'success': False,
                    'error': f'Missing required fields: {", ".join(missing_fields)}'
                }), 400

            service_types = form_data.get('serviceTypes', [])
            product_standards = form_data.get('productStandards', [])
            selected_tests = form_data.get('selectedTests', [])
            test_hours = form_data.get('testHours', {}) or {}
            test_remarks = form_data.get('testRemarks', {}) or {}
            test_configurations = form_data.get('testConfigurations', {}) or {}

            selected_test_tokens = {
                str(test_value or '').strip().upper()
                for test_value in (selected_tests if isinstance(selected_tests, list) else [selected_tests])
                if str(test_value or '').strip()
            }
            selected_standard_tokens = {
                str(standard_value or '').strip().lower()
                for standard_value in (product_standards if isinstance(product_standards, list) else [product_standards])
                if str(standard_value or '').strip()
            }
            rs_interim_allowed = any(
                token.startswith('iec 60601-1-2') or token.startswith('en 60601-1-2')
                for token in selected_standard_tokens
            )
            if 'RS_RI_INTERIM' in selected_test_tokens and not rs_interim_allowed:
                return jsonify({
                    'success': False,
                    'error': 'RS (Interim method) can only be selected when IEC 60601-1-2 or EN 60601-1-2 is chosen as a product standard.'
                }), 400

            logger.info(f"Received form submission:")
            logger.info(f"  Service Types: {service_types}")
            logger.info(f"  Selected Tests: {selected_tests}")
            logger.info(f"  Test Hours: {test_hours}")
            logger.info(f"  Test Remarks: {test_remarks}")
            logger.info(f"  Test Configurations: {test_configurations}")

            # Server-side validation: only enforce hours when Development Assistance is requested
            requires_test_hours = any(
                'development' in (stype or '').strip().lower()
                and 'assist' in (stype or '').strip().lower()
                for stype in service_types
            )

            if requires_test_hours:
                if isinstance(test_hours, str):
                    try:
                        test_hours = json.loads(test_hours)
                    except Exception:
                        test_hours = {}
                missing_or_invalid_hours = []
                for test_name in selected_tests:
                    raw_value = None
                    try:
                        raw_value = test_hours.get(test_name, '')
                        value = float(str(raw_value).strip()) if str(
                            raw_value).strip() != '' else 0.0
                    except Exception:
                        value = 0.0
                    if value <= 0:
                        missing_or_invalid_hours.append(test_name)
                if missing_or_invalid_hours:
                    return jsonify({
                        'success': False,
                        'error': 'Please provide hours (> 0) for selected tests: ' + ', '.join(missing_or_invalid_hours)
                    }), 400

            processed_test_data = process_test_configurations(
                selected_tests, test_hours, test_configurations)
            logger.info(f"Processed test data: {processed_test_data}")

            # Helper function to generate next TCO ID
            def generate_next_tco_id():
                """Generate the next TCO ID in format IEC-EMC-001, IEC-EMC-002, etc."""
                # Find the highest existing TCO ID
                last_tco = TestRequest.query.filter(
                    TestRequest.tco_id.isnot(None),
                    TestRequest.tco_id.like('IEC-EMC-%')
                ).order_by(TestRequest.tco_id.desc()).first()

                if last_tco and last_tco.tco_id:
                    # Extract the number part
                    try:
                        last_number = int(last_tco.tco_id.split('-')[-1])
                        next_number = last_number + 1
                    except (ValueError, IndexError):
                        next_number = 1
                else:
                    next_number = 1

                # Format as IEC-EMC-XXX (3 digits with leading zeros)
                return f"IEC-EMC-{next_number:03d}"

            # ============================================================================
            # HANDLE UPDATE OF EXISTING TCO (Need More Information or Draft scenario)
            # ============================================================================
            if is_update and tco_id_from_request:
                logger.info(
                    f"=== UPDATING EXISTING TCO: {tco_id_from_request} ===")

                # Find the existing EMC request by TCO ID
                existing_request = EMCRequest.query.filter_by(
                    tco_id=tco_id_from_request,
                    user_id=current_user.id
                ).first()

                # If not found for this user, check if admin
                if not existing_request and current_user.role == 'admin':
                    existing_request = EMCRequest.query.filter_by(
                        tco_id=tco_id_from_request
                    ).first()

                if not existing_request:
                    logger.error(
                        f"TCO {tco_id_from_request} not found for update")
                    return jsonify({
                        'success': False,
                        'error': f'Test request with TCO {tco_id_from_request} not found or you do not have permission to update it'
                    }), 404

                logger.info(
                    f"Found existing request: ID={existing_request.id}, Status={existing_request.status}")

                if (
                    _is_draft_status(existing_request.status)
                    and current_user.role != 'admin'
                    and not _is_request_owner(existing_request)
                ):
                    logger.warning(
                        "Blocked update attempt on private draft. request_id=%s actor_id=%s",
                        existing_request.id,
                        current_user.id
                    )
                    return jsonify({
                        'success': False,
                        'error': 'Draft not found or you do not have permission to edit it'
                    }), 403

                # Verify status allows editing
                if (
                    current_user.role != 'admin'
                    and existing_request.status not in ['Draft', 'Need More Information']
                ):
                    logger.error(
                        f"Cannot update TCO {tco_id_from_request} with status {existing_request.status}")
                    return jsonify({
                        'success': False,
                        'error': f'Cannot update test request with status "{existing_request.status}". Only Draft or Need More Information requests can be edited.'
                    }), 400

                # Store the old status
                old_status = existing_request.status
                previous_selected_codes = {
                    str(getattr(test_row, 'test_code', '') or '').strip().upper()
                    for test_row in (getattr(existing_request, 'tests', []) or [])
                    if getattr(test_row, 'is_selected', False)
                }

                # Update the existing request with new form data
                logger.info(
                    f"Updating EMC request ID {existing_request.id} with new data")
                _populate_iec_request_from_form(existing_request, form_data)

                # Change status to "At Review" if it was "Need More Information"
                if old_status == 'Need More Information':
                    has_assigned_engineer = bool(
                        getattr(existing_request, 'assigned_engineer_id', None))
                    has_test_assignments = False
                    try:
                        assignments_raw = getattr(
                            existing_request, 'test_assignments', None)
                        if assignments_raw:
                            parsed_assignments = json.loads(
                                assignments_raw) if isinstance(assignments_raw, str) else assignments_raw
                            has_test_assignments = isinstance(
                                parsed_assignments, list) and len(parsed_assignments) > 0
                    except Exception:
                        has_test_assignments = bool(
                            getattr(existing_request, 'test_assignments', None))

                    if has_assigned_engineer or has_test_assignments:
                        logger.info(
                            "Changing status from 'Need More Information' to 'Assigned Lab Engineer' for re-review")
                        existing_request.status = 'Assigned Lab Engineer'
                    else:
                        logger.info(
                            "Changing status from 'Need More Information' to 'At Review' (no engineer assignment found)")
                        existing_request.status = 'At Review'
                    existing_request.submitted_at = get_ist_now()
                    # Preserve comment thread for two-way requester/reviewer communication.
                elif old_status == 'Draft':
                    logger.info(f"Promoting Draft to 'At Review'")
                    existing_request.status = 'At Review'
                    existing_request.submitted_at = get_ist_now()
                elif current_user.role == 'admin':
                    logger.info(
                        "Admin edit preserved workflow status '%s' for TCO %s",
                        old_status,
                        tco_id_from_request
                    )
                    removed_test_sync = _reconcile_removed_tests_after_admin_edit(
                        test_request=existing_request,
                        previous_selected_codes=previous_selected_codes,
                        edited_by_user=current_user
                    )
                    if removed_test_sync['removed_labels']:
                        logger.info(
                            "Admin removed tests from TCO %s: %s (cancelled planner entries=%d)",
                            tco_id_from_request,
                            ', '.join(removed_test_sync['removed_labels']),
                            removed_test_sync['cancelled_planner_entries']
                        )

                # Keep the same TCO ID
                tco_id = existing_request.tco_id
                iec_request = existing_request

                # Commit the update
                db.session.commit()
                logger.info(
                    f"Successfully updated EMC request ID {iec_request.id} with TCO {tco_id}")

                # Get or create TestRequest for file tracking
                test_request = TestRequest.query.filter_by(
                    tco_id=tco_id).first()
                if not test_request:
                    timestamp = get_ist_now().strftime('%Y%m%d_%H%M%S')
                    filename = f"{timestamp}_form_based_test_plan.json"
                    test_request = TestRequest(
                        user_id=current_user.id,
                        tco_id=tco_id,
                        filename=filename,
                        original_filename="Form-based Test Request (Updated)",
                        file_size=0,
                        user_ip=request.remote_addr,
                        status='created'
                    )
                    db.session.add(test_request)
                    db.session.commit()

                # Send email notification about the update
                try:
                    requester_name = form_data.get(
                        'requesterName', '') or current_user.username
                    requester_email = form_data.get('requesterEmail', '') or (
                        current_user.email if current_user.is_authenticated else '')
                    send_submission_notification(
                        iec_request, requester_name, requester_email)
                except Exception as e:
                    logger.error(
                        'Failed to send update notification email: %s', e)

            # ============================================================================
            # HANDLE NEW REQUEST OR DRAFT PROMOTION (Original logic)
            # ============================================================================
            else:
                logger.info("=== CREATING NEW REQUEST OR PROMOTING DRAFT ===")

                # If the user is submitting an existing draft, promote that draft instead of
                # creating a separate EMC request so the old draft row disappears
                existing_draft = None
                if draft_id:
                    try:
                        draft_id_int = int(draft_id)
                    except (TypeError, ValueError):
                        draft_id_int = None

                    if draft_id_int is not None:
                        existing_draft = EMCRequest.query.filter_by(
                            id=draft_id_int,
                            user_id=current_user.id,
                            status='Draft'
                        ).first()

                # Determine or generate TCO ID
                if existing_draft and existing_draft.tco_id and existing_draft.tco_id.startswith('IEC-EMC-'):
                    tco_id = existing_draft.tco_id
                    logger.info(f"Using existing draft TCO ID: {tco_id}")
                else:
                    tco_id = generate_next_tco_id()
                    logger.info(f"Generated new TCO ID: {tco_id}")

                # Create a TestRequest record for traceability / file tracking
                timestamp = get_ist_now().strftime('%Y%m%d_%H%M%S')
                filename = f"{timestamp}_form_based_test_plan.json"

                test_request = TestRequest(
                    user_id=current_user.id,
                    tco_id=tco_id,
                    filename=filename,
                    original_filename="Form-based Test Request",
                    file_size=0,
                    user_ip=request.remote_addr,
                    status='created'
                )
                db.session.add(test_request)
                db.session.commit()
                logger.info(
                    f"Created TestRequest record with ID: {test_request.id}")

                # Use existing EMC draft when available, otherwise create new
                if existing_draft:
                    logger.info(
                        f"Promoting existing draft ID: {existing_draft.id}")
                    iec_request = existing_draft
                    iec_request.status = 'At Review'
                    iec_request.tco_id = tco_id
                    iec_request.submitted_at = get_ist_now()
                    _populate_iec_request_from_form(iec_request, form_data)
                else:
                    logger.info("Creating new IEC EMC Test Request")
                    iec_request = EMCRequest(
                        user_id=current_user.id,
                        status='At Review',
                        tco_id=tco_id,
                        submitted_at=get_ist_now()
                    )
                    _populate_iec_request_from_form(iec_request, form_data)

                db.session.add(iec_request)

                # Check for dict/list fields before commit
                bad = []
                for col in iec_request.__table__.columns:
                    v = getattr(iec_request, col.name)
                    if isinstance(v, (dict, list)):
                        bad.append((col.name, type(v).__name__))
                logger.info("DICT/LIST FIELDS BEFORE COMMIT: %s", bad)

                db.session.commit()
                logger.info(
                    f"IEC EMC Test Request saved with ID: {iec_request.id}")

                # Send email notification to lab engineers and admins
                try:
                    requester_name = form_data.get(
                        'requesterName', '') or current_user.username
                    requester_email = form_data.get('requesterEmail', '') or (
                        current_user.email if current_user.is_authenticated else '')
                    send_submission_notification(
                        iec_request, requester_name, requester_email)
                except Exception as e:
                    logger.error(
                        'Failed to send submission notification email: %s', e)

            # ============================================================================
            # STORE EXTRACTED SUBMISSION DATA (no automatic document generation)
            # ============================================================================

            # Prepare extracted data from form
            extracted_data = {
                # Service and Product Information
                'service_types': service_types,
                'product_name': form_data.get('productName', ''),
                'manufacturer': form_data.get('manufacturer', ''),
                'model_number': form_data.get('modelNumber', ''),
                'serial_number': form_data.get('serialNumber', ''),
                'test_samples': form_data.get('testSamples', ''),
                'weight': form_data.get('weight', ''),
                'operating_frequency': form_data.get('operatingFrequency', ''),
                'length': form_data.get('length', ''),
                'width': form_data.get('width', ''),
                'height': form_data.get('height', ''),
                'category': form_data.get('category', ''),
                'type': form_data.get('type', ''),
                'type_others': form_data.get('typeOthers', ''),

                # Accessories and Cables
                'accessories': form_data.get('accessories', []),
                'cables': form_data.get('cables', []),
                'eut_specs': form_data.get('eutSpecs', []),
                'supply_vf': form_data.get('supplyVf', []),

                # Operating Conditions
                'ac_voltage_range': form_data.get('acVoltageRange', ''),
                'ac_voltage_nominal': form_data.get('acVoltageNominal', ''),
                'freq_range': form_data.get('freqRange', ''),
                'freq_nominal': form_data.get('freqNominal', ''),
                'dc_voltage_range': form_data.get('dcVoltageRange', ''),
                'dc_voltage_nominal': form_data.get('dcVoltageNominal', ''),
                'input_current': form_data.get('inputCurrent', ''),
                'rated_power': form_data.get('ratedPower', ''),

                # Wireless Interface
                'has_wireless_interface': form_data.get('hasWirelessInterface', ''),
                'wireless': form_data.get('wireless', []),
                'wireless_interface': form_data.get('wirelessInterface', ''),
                'carrier_frequency': form_data.get('carrierFrequency', ''),
                'max_output_power': form_data.get('maxOutputPower', ''),
                'clock_frequencies': form_data.get('clockFrequencies', ''),

                # Product Description and Configuration
                'product_description': form_data.get('productDescription', ''),
                'test_configuration': form_data.get('testConfiguration', ''),
                'operation_modes': form_data.get('operationModes', ''),
                'monitoring_parameters': form_data.get('monitoringParameters', ''),
                'additional_info': form_data.get('additionalInfo', ''),

                # Test Duration (for Development Assistance)
                'test_hours': form_data.get('testHours', {}),
                'test_configurations': test_configurations,
                'processed_test_data': processed_test_data,

                # Tests Required
                'product_standards': product_standards,
                'product_environment': form_data.get('productEnvironment', {}),
                'group': form_data.get('group', ''),
                'class': form_data.get('class', ''),
                'selected_tests': selected_tests,
                'selected_tests_for_development': form_data.get('selectedTestsForDevelopment', []),

                # Test-specific data
                'ce_standard': form_data.get('ce_standard', []),
                'ce_voltage_freq': form_data.get('ce_voltageFreq', ''),
                'ce_freq_range': form_data.get('ce_freqRange', ''),
                'ce_cables': form_data.get('ce_cables', ''),
                'ce_class': form_data.get('ce_class', ''),

                're_standard': form_data.get('re_standard', []),
                're_voltage_freq': form_data.get('re_voltageFreq', ''),
                're_freq_range': form_data.get('re_freqRange', ''),
                're_class': form_data.get('re_class', ''),

                'esd_standard': form_data.get('esd_standard', []),
                'esd_voltage_freq': form_data.get('esd_voltageFreq', ''),
                'esd_contact': form_data.get('esd_contact', ''),
                'esd_air': form_data.get('esd_air', ''),

                'harmonic_standard': form_data.get('harmonic_standard', []),
                'harmonic_voltage_freq': form_data.get('harmonic_voltageFreq', ''),

                'flicker_standard': form_data.get('flicker_standard', []),
                'flicker_voltage_freq': form_data.get('flicker_voltageFreq', ''),

                'rs_standard': form_data.get('rs_standard', []),
                'rs_voltage_freq': form_data.get('rs_voltageFreq', ''),
                'rs_freq_range': form_data.get('rs_freqRange', ''),
                'rs_field_strength_1': form_data.get('rs_fieldStrength1', ''),
                'rs_field_strength_2': form_data.get('rs_fieldStrength2', ''),
                'rs_field_strength_3': form_data.get('rs_fieldStrength3', ''),

                # Additional test sections
                'rs_interim_standard': form_data.get('rs_interim_standard', []),
                'rs_interim_voltage_freq': form_data.get('rs_interim_voltageFreq', ''),
                'rs_interim_freq_range': form_data.get('rs_interim_freqRange', ''),
                'rs_interim_field_strength_1': form_data.get('rs_interim_fieldStrength1', ''),
                'rs_interim_field_strength_2': form_data.get('rs_interim_fieldStrength2', ''),
                'rs_interim_field_strength_3': form_data.get('rs_interim_fieldStrength3', ''),

                'eft_standard': form_data.get('eft_standard', []),
                'eft_voltage_freq': form_data.get('eft_voltageFreq', ''),
                'eft_cables_power': form_data.get('eft_cables_power', ''),
                'eft_cables_signal': form_data.get('eft_cables_signal', ''),
                'eft_test_level_1': form_data.get('eft_testLevel1', ''),
                'eft_test_level_2': form_data.get('eft_testLevel2', ''),

                'surge_standard': form_data.get('surge_standard', []),
                'surge_voltage_freq': form_data.get('surge_voltageFreq', ''),
                'surge_cables_power': form_data.get('surge_cables_power', ''),
                'surge_cables_signal': form_data.get('surge_cables_signal', ''),
                'surge_cm_1': form_data.get('surge_cm1', ''),
                'surge_cm_2': form_data.get('surge_cm2', ''),
                'surge_dm_1': form_data.get('surge_dm1', ''),
                'surge_dm_2': form_data.get('surge_dm2', ''),

                'crf_standard': form_data.get('crf_standard', []),
                'crf_voltage_freq': form_data.get('crf_voltageFreq', ''),
                'crf_freq_range': form_data.get('crf_freqRange', ''),
                'crf_cables_power': form_data.get('crf_cables_power', ''),
                'crf_cables_signal': form_data.get('crf_cables_signal', ''),
                'crf_test_level_1': form_data.get('crf_testLevel1', ''),
                'crf_test_level_2': form_data.get('crf_testLevel2', ''),

                'power_freq_standard': form_data.get('power_freq_standard', []),
                'power_freq_voltage_freq': form_data.get('power_freq_voltageFreq', ''),
                'power_freq_test_level': form_data.get('power_freq_testLevel', ''),

                'voltage_dips_standard': form_data.get('voltage_dips_standard', []),
                'voltage_dips_min': form_data.get('voltage_dips_min', ''),
                'voltage_dips_max': form_data.get('voltage_dips_max', ''),
                'voltage_dips_voltage_freq': form_data.get('voltage_dips_voltageFreq', ''),
                'voltage_dips_voltage_dip_1': form_data.get('voltage_dips_voltageDip1', ''),
                'voltage_dips_voltage_dip_2': form_data.get('voltage_dips_voltageDip2', ''),
                'voltage_dips_voltage_dip_3': form_data.get('voltage_dips_voltageDip3', ''),
                'voltage_dips_interruption': form_data.get('voltage_dips_interruption', ''),
                'voltage_dips_time_1': form_data.get('voltage_dips_time1', ''),
                'voltage_dips_time_2': form_data.get('voltage_dips_time2', ''),
                'voltage_dips_time_3': form_data.get('voltage_dips_time3', ''),
                'voltage_dips_time_4': form_data.get('voltage_dips_time4', ''),

                # Test Requirements and Decision Rules
                'continue_testing': form_data.get('continueTesting', ''),
                'test_report_required': form_data.get('testReportRequired', ''),
                'uncertainty_required': form_data.get('uncertaintyRequired', ''),
                'test_witness': form_data.get('testWitness', ''),
                'conformity_required': form_data.get('conformityRequired', ''),
                'decision_rule': form_data.get('decisionRule', []),

                # Requester Information
                'requester_name': form_data.get('requesterName', ''),
                'requester_department': form_data.get('requesterDepartment', ''),
                'requester_email': form_data.get('requesterEmail', ''),
                'requester_contact': form_data.get('requesterContact', ''),
                'requester_designation': form_data.get('requesterDesignation', ''),
                'requester_status': form_data.get('requesterStatus', 'At Review'),
                'requester_date': form_data.get('requesterDate', ''),
                'requester_signature': form_data.get('requesterSignature', ''),

                # Laboratory Use Only
                'job_number': form_data.get('jobNumber', ''),
                'sample_condition': form_data.get('sampleCondition', ''),
                'capability_available': form_data.get('capabilityAvailable', ''),
                'sample_received_date': form_data.get('sampleReceivedDate', ''),
                'test_duration': form_data.get('testDuration', ''),
                'test_commencement_date': form_data.get('testCommencementDate', ''),
                'test_completion_date': form_data.get('testCompletionDate', ''),

                # Lab Manager
                'lab_manager_name': form_data.get('labManagerName', ''),
                'lab_manager_date': form_data.get('labManagerDate', ''),
                'lab_manager_signature': form_data.get('labManagerSignature', '')
            }

            # Update test request with extracted data for traceability
            test_request.set_extracted_data(extracted_data)
            test_request.status = 'processed'
            test_request.processing_time = 0.1
            db.session.commit()
            logger.info(
                "Submission saved for request_id=%s (tco_id=%s). Automatic document generation skipped.",
                test_request.id,
                tco_id
            )

            return jsonify({
                'success': True,
                'message': f'Test request {"updated" if is_update else "created"} successfully',
                'data': {
                    'request_id': test_request.id,
                    'tco_id': tco_id,
                    'test_plan': None,
                    'datasheets': [],
                    'total_files': 0,
                    'was_update': is_update
                }
            })

        except Exception as e:
            logger.exception("Error creating test plan from form: %s", e)
            db.session.rollback()
            # Return more detailed error message for debugging
            error_message = str(e)
            if flask_app.config.get('DEBUG'):
                return jsonify({
                    'success': False,
                    'error': f'An error occurred while creating the test plan: {error_message}'
                }), 500
            else:
                return jsonify({
                    'success': False,
                    'error': 'An error occurred while creating the test plan. Please check the server logs for details.'
                }), 500

    @flask_app.route('/generate', methods=['POST'])
    @login_required
    def generate_documents():
        """Generate test plan and datasheets."""
        try:
            def safe_file_size(path: str) -> int:
                try:
                    return os.path.getsize(path)
                except OSError as exc:
                    logger.warning('Unable to read file size for %s: %s', path, exc)
                    return 0

            data = request.get_json()
            request_id = data.get('request_id')
            selected_tests = data.get('selected_tests', [])
            equipment_list = data.get('equipment_list', [])

            # Get test request
            test_request = db.session.get(TestRequest, request_id)
            if not test_request:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            # Get extracted data
            extracted_data = test_request.get_extracted_data()

            # Get equipment information
            equipment_objects = equipment_manager_instance.find_equipment_by_requirements(
                equipment_list)
            equipment_data = equipment_objects  # Already returns list of dictionaries

            generated_files = []

            # Generate test plan
            test_plan_filename = document_generator.generate_test_plan(
                extracted_data, equipment_data)
            generated_files.append(test_plan_filename)

            # Create test plan record
            test_plan = TestPlan(
                test_request_id=request_id,
                filename=test_plan_filename,
                title=extracted_data.get('project_name', 'Test Request'),
                test_objective=extracted_data.get('test_objective', ''),
                test_scope=extracted_data.get('test_scope', ''),
                equipment_used=json.dumps([eq['id'] for eq in equipment_data]),
                test_methodology=document_generator._generate_test_methodology(
                    extracted_data),
                safety_requirements=extracted_data.get(
                    'safety_requirements', ''),
                file_size=safe_file_size(os.path.join(
                    flask_app.config['OUTPUT_FOLDER'], test_plan_filename))
            )
            db.session.add(test_plan)
            db.session.commit()

            # Generate datasheets for selected tests
            datasheet_filenames = []
            tests_to_perform = extracted_data.get('tests_to_perform', [])

            for i, test_name in enumerate(selected_tests):
                if test_name in tests_to_perform:
                    # Create test data for datasheet
                    test_data = {
                        'description': f"Test procedure for {test_name}",
                        'parameters': extracted_data.get('test_parameters', {}),
                        'measurement_points': [
                            {'name': f'MP_{j+1}',
                                'description': f'Measurement point {j+1}', 'expected': ''}
                            for j in range(3)
                        ]
                    }

                    datasheet_filename = document_generator.generate_test_datasheet(
                        test_name, test_data, equipment_data, i + 1
                    )
                    datasheet_filenames.append(datasheet_filename)

                    # Create datasheet record
                    datasheet = TestDatasheet(
                        test_request_id=request_id,
                        test_plan_id=test_plan.id,
                        filename=datasheet_filename,
                        test_name=test_name,
                        test_description=test_data['description'],
                        test_parameters=json.dumps(test_data['parameters']),
                        measurement_points=json.dumps(
                            test_data['measurement_points']),
                        equipment_required=json.dumps(
                            [eq['id'] for eq in equipment_data]),
                        test_procedure=f"Standard procedure for {test_name}",
                        data_recording_sections=json.dumps([]),
                        file_size=safe_file_size(os.path.join(
                            flask_app.config['OUTPUT_FOLDER'], datasheet_filename))
                    )
                    db.session.add(datasheet)

            db.session.commit()
            generated_files.extend(datasheet_filenames)

            # Update test request with generated files
            test_request.add_generated_file(test_plan_filename)
            for filename in datasheet_filenames:
                test_request.add_generated_file(filename)
            test_request.status = 'completed'
            db.session.commit()

            return jsonify({
                'success': True,
                'message': 'Documents generated successfully',
                'data': {
                    'test_plan': test_plan_filename,
                    'datasheets': datasheet_filenames,
                    'total_files': len(generated_files)
                }
            })

        except Exception as e:
            logger.error("Error generating documents: %s", e)
            return jsonify({
                'success': False,
                'error': 'An error occurred while generating documents'
            }), 500

    @flask_app.route('/download/<filename>')
    @login_required
    def download_file(filename):
        """Download generated file."""
        try:
            file_path = os.path.join(
                flask_app.config['OUTPUT_FOLDER'], filename)
            if os.path.exists(file_path):
                return send_file(file_path, as_attachment=True)
            else:
                return jsonify({
                    'success': False,
                    'error': 'File not found'
                }), 404
        except Exception as e:
            logger.error("Error downloading file %s: %s", filename, e)
            return jsonify({
                'success': False,
                'error': 'An error occurred while downloading the file'
            }), 500

    @flask_app.route('/equipment', methods=['GET', 'POST'])
    @login_required
    def equipment_list():
        """Display equipment list with pagination."""
        # Check if user is admin or lab engineer
        if current_user.role not in ['admin', 'lab_engineer']:
            flash('Access denied. Admin or Lab Engineer privileges required.', 'error')
            return redirect(url_for('index'))

        try:
            # Get pagination parameters from either GET or POST
            status_filter: str
            search_term: str
            test_type_filter: str
            test_name_filter: str
            if request.method == 'POST':
                page = request.form.get('page', 1, type=int)
                per_page = request.form.get('per_page', 10, type=int)
                status_filter = request.form.get('status_filter', '') or ''
                search_term = request.form.get('search_term', '') or ''
                test_type_filter = request.form.get(
                    'test_type_filter', '') or ''
                test_name_filter = request.form.get(
                    'test_name_filter', '') or ''
            else:
                page = request.args.get('page', 1, type=int)
                per_page = request.args.get('per_page', 10, type=int)
                status_filter = request.args.get('status_filter', '') or ''
                search_term = request.args.get('search_term', '') or ''
                test_type_filter = request.args.get(
                    'test_type_filter', '') or ''
                test_name_filter = request.args.get(
                    'test_name_filter', '') or ''

            # Validate per_page to prevent abuse
            if per_page not in [5, 10, 20, 50, 100]:
                per_page = 10

            normalized_status_column = db.func.lower(db.func.trim(Equipment.status))

            def _equipment_status_values(status_key):
                status_aliases = {
                    'available': {'available'},
                    'in_use': {'in_use', 'in use', 'active'},
                    'maintenance': {'maintenance', 'under maintenance'},
                    'out_of_service': {'out_of_service', 'out of service'},
                    'calibration': {'calibration', 'needs calibration'},
                }
                return tuple(status_aliases.get(status_key, {status_key}))

            def _equipment_status_filter(status_key):
                return normalized_status_column.in_(
                    _equipment_status_values(status_key)
                )

            # Build query with status filter, test type filter, and search term
            query = Equipment.query
            if status_filter and status_filter.strip():
                normalized_status_filter = status_filter.strip()
                status_filter_key = normalized_status_filter.lower()

                if status_filter_key in {'calibration', 'needs_calibration', 'out_of_calibration'}:
                    query = query.filter(
                        db.or_(
                            Equipment.calibration_status_col == 'Out of Calibration',
                            _equipment_status_filter('calibration')
                        )
                    )
                elif status_filter_key == 'in_calibration':
                    query = query.filter(
                        Equipment.calibration_status_col == 'In Calibration'
                    )
                else:
                    query = query.filter(
                        _equipment_status_filter(status_filter_key)
                    )

            if test_type_filter and test_type_filter.strip():
                query = query.filter(Equipment.test_type == test_type_filter)

            if test_name_filter and test_name_filter.strip():
                query = query.filter(
                    Equipment.test_name.isnot(None),
                    Equipment.test_name.ilike(f'%{test_name_filter}%')
                )

            if search_term and search_term.strip():
                search_filter = f'%{search_term}%'
                query = query.filter(
                    (Equipment.asset_id.ilike(search_filter)) |
                    (Equipment.name.ilike(search_filter)) |
                    (Equipment.make.ilike(search_filter)) |
                    (Equipment.model_no.ilike(search_filter)) |
                    (Equipment.serial_no.ilike(search_filter)) |
                    (Equipment.test_type.ilike(search_filter)) |
                    (Equipment.test_name.ilike(search_filter)) |
                    (Equipment.status.ilike(search_filter))  # type: ignore
                )

            raw_test_name_rows = db.session.query(Equipment.test_name).filter(
                Equipment.test_name.isnot(None),
                db.func.trim(Equipment.test_name) != ''
            ).all()
            test_name_options = sorted(
                {
                    name_part.strip()
                    for row in raw_test_name_rows
                    for name_part in str(row[0] or '').split(',')
                    if name_part and name_part.strip()
                },
                key=str.lower
            )

            # Get equipment with pagination
            pagination = query.order_by(Equipment.name).paginate(
                page=page, per_page=per_page, error_out=False
            )

            # Calculate statistics
            total = Equipment.query.count()
            available = Equipment.query.filter(
                _equipment_status_filter('available')
            ).count()
            in_use = Equipment.query.filter(
                _equipment_status_filter('in_use')
            ).count()
            maintenance = Equipment.query.filter(
                _equipment_status_filter('maintenance')
            ).count()
            # Calculate calibration status statistics based on calibration_status_col
            in_calibration = Equipment.query.filter(
                Equipment.calibration_status_col == 'In Calibration'
            ).count()
            out_of_calibration = Equipment.query.filter(
                Equipment.calibration_status_col == 'Out of Calibration'
            ).count()
            # Keep needs_calibration for backward compatibility (based on due date)
            needs_calibration = Equipment.query.filter(
                Equipment.calibration_due_date < get_ist_now().date()
            ).count()

            statistics = {
                'total': total,
                'available': available,
                'in_use': in_use,
                'maintenance': maintenance,
                'needs_calibration': needs_calibration,
                'in_calibration': in_calibration,
                'out_of_calibration': out_of_calibration
            }

            return render_template('equipment.html', pagination=pagination,
                                   statistics=statistics, status_filter=status_filter,
                                   search_term=search_term, test_type_filter=test_type_filter,
                                   test_name_filter=test_name_filter,
                                   test_name_options=test_name_options)
        except Exception as e:
            logger.error("Error loading equipment page: %s", e)
            flash('Error loading equipment information', 'error')
            return redirect(url_for('index'))

    @flask_app.route('/api/equipment')
    @login_required
    def api_equipment():
        """API endpoint for equipment data."""
        # Check if user is admin or lab engineer
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin or Lab Engineer privileges required.'
            }), 403

        try:
            equipment = Equipment.query.order_by(Equipment.name).all()
            return jsonify({
                'success': True,
                'data': [eq.to_dict() for eq in equipment]
            })
        except Exception as e:
            logger.error("Error in equipment API: %s", e)
            return jsonify({
                'success': False,
                'error': 'Error retrieving equipment data'
            }), 500

    @flask_app.route('/api/equipment/search')
    @login_required
    def api_equipment_search():
        """API endpoint for equipment search."""
        # Check if user is admin or lab engineer
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin or Lab Engineer privileges required.'
            }), 403

        try:
            search_term = request.args.get('q', '')
            if not search_term:
                return jsonify({
                    'success': False,
                    'error': 'Search term is required'
                }), 400

            equipment = Equipment.query.filter(
                (Equipment.asset_id.ilike(f'%{search_term}%')) |
                (Equipment.name.ilike(f'%{search_term}%')) |
                (Equipment.make.ilike(f'%{search_term}%')) |
                (Equipment.model_no.ilike(f'%{search_term}%')) |
                (Equipment.serial_no.ilike(f'%{search_term}%')) |
                (Equipment.test_type.ilike(f'%{search_term}%'))
            ).all()

            return jsonify({
                'success': True,
                'data': [eq.to_dict() for eq in equipment]
            })
        except Exception as e:
            logger.error("Error in equipment search API: %s", e)
            return jsonify({
                'success': False,
                'error': 'Error searching equipment'
            }), 500

    @flask_app.route('/api/equipment/<int:equipment_id>', methods=['GET'])
    @login_required
    def api_equipment_detail(equipment_id):
        """API endpoint for equipment details."""
        # Check if user is admin or lab engineer
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin or Lab Engineer privileges required.'
            }), 403

        try:
            equipment = db.session.get(Equipment, equipment_id)

            if not equipment:
                return jsonify({
                    'success': False,
                    'error': 'Equipment not found'
                }), 404

            return jsonify({
                'success': True,
                'data': equipment.to_dict()
            })
        except Exception as e:
            logger.error("Error in equipment detail API: %s", e)
            return jsonify({
                'success': False,
                'error': 'Error retrieving equipment details'
            }), 500

    @flask_app.route('/api/equipment/<int:equipment_id>/history', methods=['GET'])
    @login_required
    def api_equipment_history(equipment_id):
        """API endpoint for equipment history."""
        # Check if user is admin or lab engineer
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin or Lab Engineer privileges required.'
            }), 403

        try:
            equipment = db.session.get(Equipment, equipment_id)

            if not equipment:
                return jsonify({
                    'success': False,
                    'error': 'Equipment not found'
                }), 404

            # Get history records ordered by most recent first
            try:
                history_records = EquipmentHistory.query.filter_by(
                    equipment_id=equipment_id
                ).options(joinedload(EquipmentHistory.changed_by)).order_by(EquipmentHistory.created_at.desc()).all()

                # Build history data with user information
                history_data = []
                for record in history_records:
                    record_dict = record.to_dict()
                    # Add user name if available
                    if record.changed_by:
                        record_dict['changed_by_username'] = record.changed_by.username
                        record_dict['changed_by_name'] = getattr(
                            record.changed_by, 'name', None) or record.changed_by.username
                    else:
                        record_dict['changed_by_username'] = None
                        record_dict['changed_by_name'] = 'System' if record.changed_by_user_id is None else 'Unknown User'
                    history_data.append(record_dict)

                return jsonify({
                    'success': True,
                    'data': history_data
                })
            except Exception as query_error:
                logger.error("Error querying equipment history: %s",
                             query_error, exc_info=True)
                # Return empty array if query fails (table might not exist or other DB issue)
                return jsonify({
                    'success': True,
                    'data': []
                })
        except Exception as e:
            logger.error("Error in equipment history API: %s",
                         e, exc_info=True)
            return jsonify({
                'success': False,
                'error': f'Error retrieving equipment history: {str(e)}'
            }), 500

    @flask_app.route('/api/equipment', methods=['POST'])
    @login_required
    def api_equipment_create():
        """API endpoint for creating equipment."""
        # Check if user is admin or lab engineer
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin or Lab Engineer privileges required.'
            }), 403

        try:
            data = request.get_json()

            # Server-side validation
            validation_errors = []
            today = get_ist_now().date()

            # Check required fields
            if not data or not data.get('asset_id', '').strip():
                validation_errors.append('Asset ID is required')
            elif len(data.get('asset_id', '').strip()) > 100:
                validation_errors.append(
                    'Asset ID must be no more than 100 characters')
            else:
                # Check if asset_id already exists
                existing = Equipment.query.filter_by(
                    asset_id=data.get('asset_id', '').strip()).first()
                if existing:
                    validation_errors.append(
                        'Asset ID already exists. Please use a unique Asset ID.')

            if not data or not data.get('type', '').strip():
                validation_errors.append('Type is required')
            elif data.get('type') not in ['Instrument', 'Equipment', 'Accessory']:
                validation_errors.append(
                    'Type must be Instrument, Equipment, or Accessory')

            if not data or not data.get('name', '').strip():
                validation_errors.append('Equipment name is required')
            elif len(data.get('name', '').strip()) < 2:
                validation_errors.append(
                    'Equipment name must be at least 2 characters')
            elif len(data.get('name', '').strip()) > 200:
                validation_errors.append(
                    'Equipment name must be no more than 200 characters')

            if data.get('make') and len(data.get('make', '').strip()) > 100:
                validation_errors.append(
                    'Make must be no more than 100 characters')

            if data.get('model_no') and len(data.get('model_no', '').strip()) > 200:
                validation_errors.append(
                    'Model number must be no more than 200 characters')

            if data.get('serial_no') and len(data.get('serial_no', '').strip()) > 100:
                validation_errors.append(
                    'Serial number must be no more than 100 characters')

            # Validate calibration_due_date only if calibration is required
            calibration_due_raw = data.get('calibration_due_date')
            if data.get('calibration_required') == 'Yes':
                if not calibration_due_raw:
                    validation_errors.append(
                        'Calibration due date is required when calibration is required')
            if calibration_due_raw:
                try:
                    cal_date = datetime.strptime(
                        calibration_due_raw, '%Y-%m-%d').date()
                    if cal_date < today:
                        validation_errors.append(
                            'Calibration due date cannot be in the past')
                except (TypeError, ValueError):
                    validation_errors.append(
                        'Invalid calibration due date format')

            if not data or not data.get('status'):
                validation_errors.append('Status is required')
            elif data.get('status') not in ['Available', 'in_use', 'maintenance',
                                            'calibration', 'out_of_service']:
                validation_errors.append('Invalid status value')

            test_type_value = data.get(
                'test_type', '').strip() if data.get('test_type') else ''
            if not test_type_value:
                validation_errors.append('Test type is required')
            elif test_type_value not in ['EMC', 'SAFETY', 'Safety/EMC', 'GENERAL']:
                validation_errors.append(
                    'Test type must be one of: EMC, SAFETY, Safety/EMC, GENERAL')

            # Validate calibration_required is mandatory
            if not data.get('calibration_required'):
                validation_errors.append('Calibration required is mandatory')
            elif data.get('calibration_required') not in ['Yes', 'No']:
                validation_errors.append(
                    'Calibration required must be Yes or No')

            # Validate date fields
            date_fields = {
                'calibration_date': 'Calibration date',
                'ic_date': 'IC date',
                'maintenance_date': 'Maintenance date'
            }
            for field, label in date_fields.items():
                if data.get(field):
                    try:
                        datetime.strptime(data.get(field), '%Y-%m-%d').date()
                    except (TypeError, ValueError):
                        validation_errors.append(f'Invalid {label} format')

            if data.get('ic_due_date'):
                try:
                    ic_due_date = datetime.strptime(
                        data.get('ic_due_date'), '%Y-%m-%d').date()
                    if ic_due_date < today:
                        validation_errors.append(
                            'IC due date cannot be in the past')
                except (TypeError, ValueError):
                    validation_errors.append('Invalid IC due date format')

            maintenance_records_payload = data.get('maintenance_records', [])
            if maintenance_records_payload and not isinstance(maintenance_records_payload, list):
                validation_errors.append('Invalid maintenance records payload')
            elif isinstance(maintenance_records_payload, list):
                for index, maint_data in enumerate(maintenance_records_payload, start=1):
                    if not isinstance(maint_data, dict):
                        validation_errors.append(
                            f'Invalid maintenance record at entry {index}')
                        continue

                    maint_date_raw = maint_data.get('maintenance_date')
                    if maint_date_raw:
                        try:
                            datetime.strptime(
                                maint_date_raw, '%Y-%m-%d').date()
                        except (TypeError, ValueError):
                            validation_errors.append(
                                f'Invalid Maintenance date format (entry {index})')

                    maint_due_raw = maint_data.get('maintenance_due_date')
                    if maint_due_raw:
                        try:
                            maint_due_date = datetime.strptime(
                                maint_due_raw, '%Y-%m-%d').date()
                            if maint_due_date < today:
                                validation_errors.append(
                                    f'Maintenance due date cannot be in the past (entry {index})')
                        except (TypeError, ValueError):
                            validation_errors.append(
                                f'Invalid Maintenance due date format (entry {index})')

            if validation_errors:
                return jsonify({
                    'success': False,
                    'error': '; '.join(validation_errors)
                }), 400

            # Create new equipment in MySQL database
            equipment = Equipment(  # type: ignore
                asset_id=data.get('asset_id', '').strip(),
                sl_no=data.get('sl_no') if data.get('sl_no') else None,
                type=data.get('type', '').strip(),
                # calibration_status_col will be set automatically below based on calibration_due_date
                name=data.get('name', '').strip(),
                make=data.get('make', '').strip(
                ) if data.get('make') else None,
                model_no=data.get('model_no', '').strip(
                ) if data.get('model_no') else None,
                serial_no=data.get('serial_no', '').strip(
                ) if data.get('serial_no') else None,
                location=data.get('location', '').strip(
                ) if data.get('location') else None,
                test_type=data.get('test_type', '').strip(),
                eou_status=data.get('eou_status', '').strip(
                ) if data.get('eou_status') else None,
                test_name=data.get('test_name', '').strip(
                ) if data.get('test_name') else None,
                calibration_required=data.get('calibration_required') or None,
                calibration_frequency=data.get(
                    'calibration_frequency') or None,
                calibration_date=parse_date_field(
                    data.get('calibration_date')),
                calibration_due_date=parse_date_field(
                    data.get('calibration_due_date')),
                # Automatically set calibration_status_col based on calibration_due_date
                calibration_status_col=_calculate_calibration_status(
                    parse_date_field(data.get('calibration_due_date')),
                    data.get('calibration_required')
                ),
                ic_required=data.get('ic_required') or None,
                ic_date=parse_date_field(data.get('ic_date')),
                ic_due_date=parse_date_field(data.get('ic_due_date')),
                manufacturer_calibration_params=data.get(
                    'manufacturer_calibration_params') or None,
                calibration_agency_params=data.get(
                    'calibration_agency_params') or None,
                document_link=data.get('document_link', '').strip(
                ) if data.get('document_link') else None,
                status=data.get('status')
            )

            # Set calibration_status_col automatically based on calibration_due_date
            calibration_due_date_parsed = parse_date_field(
                data.get('calibration_due_date'))
            equipment.calibration_status_col = _calculate_calibration_status(
                calibration_due_date_parsed,
                data.get('calibration_required')
            ) or data.get('calibration_status_col') or None

            db.session.add(equipment)
            db.session.flush()  # Flush to get equipment.id

            # Handle maintenance records
            maintenance_records = data.get('maintenance_records', [])
            first_maintenance_entry = next(
                (
                    maint_data for maint_data in maintenance_records
                    if maint_data.get('maintenance_required') or maint_data.get('maintenance_date') or maint_data.get('maintenance_due_date')
                ),
                None
            )
            equipment.maintenance_required = (
                first_maintenance_entry.get('maintenance_required') or None
                if first_maintenance_entry else None
            )
            equipment.maintenance_date = parse_date_field(
                first_maintenance_entry.get('maintenance_date')
            ) if first_maintenance_entry else None
            equipment.maintenance_due_date = parse_date_field(
                first_maintenance_entry.get('maintenance_due_date')
            ) if first_maintenance_entry else None
            if maintenance_records:
                for maint_data in maintenance_records:
                    if maint_data.get('maintenance_required') or maint_data.get('maintenance_date') or maint_data.get('maintenance_due_date'):
                        maintenance = Maintenance(
                            equipment_id=equipment.id,
                            maintenance_required=maint_data.get(
                                'maintenance_required') or None,
                            maintenance_date=parse_date_field(
                                maint_data.get('maintenance_date')),
                            maintenance_due_date=parse_date_field(
                                maint_data.get('maintenance_due_date'))
                        )
                        db.session.add(maintenance)

            # Log creation to history
            new_values = equipment.to_dict()
            log_equipment_history(
                equipment.id, 'created', new_values=new_values, notes='Equipment created')

            db.session.commit()

            return jsonify({
                'success': True,
                'data': equipment.to_dict()
            })
        except Exception as e:
            logger.error("Error in equipment create API: %s", e)
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': 'Error creating equipment'
            }), 500

    @flask_app.route('/api/equipment/<int:equipment_id>', methods=['PUT'])
    @login_required
    def api_equipment_update(equipment_id):
        """API endpoint for updating equipment."""
        # Check if user is admin or lab engineer
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin or Lab Engineer privileges required.'
            }), 403

        try:
            data = request.get_json()
            if not data:
                return jsonify({
                    'success': False,
                    'error': 'No data provided'
                }), 400

            equipment = db.session.get(Equipment, equipment_id)
            if not equipment:
                return jsonify({
                    'success': False,
                    'error': 'Equipment not found'
                }), 404

            # Server-side validation
            validation_errors = []
            today = get_ist_now().date()

            # Check required fields
            if not data.get('asset_id', '').strip():
                validation_errors.append('Asset ID is required')
            elif len(data.get('asset_id', '').strip()) > 100:
                validation_errors.append(
                    'Asset ID must be no more than 100 characters')
            else:
                # Check if asset_id already exists (excluding current equipment)
                # Only check for duplicates if the asset_id has changed
                new_asset_id = data.get('asset_id', '').strip()
                if new_asset_id != equipment.asset_id:
                    existing = Equipment.query.filter(
                        Equipment.asset_id == new_asset_id,
                        Equipment.id != equipment_id
                    ).first()
                    if existing:
                        validation_errors.append(
                            'Asset ID already exists. Please use a unique Asset ID.')

            if not data.get('type', '').strip():
                validation_errors.append('Type is required')
            elif data.get('type') not in ['Instrument', 'Equipment', 'Accessory']:
                validation_errors.append(
                    'Type must be Instrument, Equipment, or Accessory')

            if not data.get('name', '').strip():
                validation_errors.append('Equipment name is required')
            elif len(data.get('name', '').strip()) < 2:
                validation_errors.append(
                    'Equipment name must be at least 2 characters')
            elif len(data.get('name', '').strip()) > 200:
                validation_errors.append(
                    'Equipment name must be no more than 200 characters')

            if data.get('make') and len(data.get('make', '').strip()) > 100:
                validation_errors.append(
                    'Make must be no more than 100 characters')

            if data.get('model_no') and len(data.get('model_no', '').strip()) > 200:
                validation_errors.append(
                    'Model number must be no more than 200 characters')

            if data.get('serial_no') and len(data.get('serial_no', '').strip()) > 100:
                validation_errors.append(
                    'Serial number must be no more than 100 characters')

            # Validate calibration_due_date only if calibration is required
            calibration_due_raw = data.get('calibration_due_date')
            if data.get('calibration_required') == 'Yes':
                if not calibration_due_raw:
                    validation_errors.append(
                        'Calibration due date is required when calibration is required')
            if calibration_due_raw:
                try:
                    calibration_due_date = datetime.strptime(
                        calibration_due_raw, '%Y-%m-%d').date()
                    if calibration_due_date < today:
                        validation_errors.append(
                            'Calibration due date cannot be in the past')
                except (TypeError, ValueError):
                    validation_errors.append(
                        'Invalid calibration due date format')

            if not data.get('status'):
                validation_errors.append('Status is required')
            elif data.get('status') not in ['Available', 'in_use', 'maintenance',
                                            'calibration', 'out_of_service']:
                validation_errors.append('Invalid status value')

            test_type_value = data.get(
                'test_type', '').strip() if data.get('test_type') else ''
            if not test_type_value:
                validation_errors.append('Test type is required')
            elif test_type_value not in ['EMC', 'SAFETY', 'Safety/EMC', 'GENERAL']:
                validation_errors.append(
                    'Test type must be one of: EMC, SAFETY, Safety/EMC, GENERAL')

            # Validate calibration_required is mandatory
            if not data.get('calibration_required'):
                validation_errors.append('Calibration required is mandatory')
            elif data.get('calibration_required') not in ['Yes', 'No']:
                validation_errors.append(
                    'Calibration required must be Yes or No')

            # Validate date fields
            date_fields = {
                'calibration_date': 'Calibration date',
                'ic_date': 'IC date',
                'maintenance_date': 'Maintenance date'
            }
            for field, label in date_fields.items():
                if data.get(field):
                    try:
                        datetime.strptime(data.get(field), '%Y-%m-%d').date()
                    except (TypeError, ValueError):
                        validation_errors.append(f'Invalid {label} format')

            if data.get('ic_due_date'):
                try:
                    ic_due_date = datetime.strptime(
                        data.get('ic_due_date'), '%Y-%m-%d').date()
                    if ic_due_date < today:
                        validation_errors.append(
                            'IC due date cannot be in the past')
                except (TypeError, ValueError):
                    validation_errors.append('Invalid IC due date format')

            maintenance_records_payload = data.get('maintenance_records', [])
            if maintenance_records_payload and not isinstance(maintenance_records_payload, list):
                validation_errors.append('Invalid maintenance records payload')
            elif isinstance(maintenance_records_payload, list):
                for index, maint_data in enumerate(maintenance_records_payload, start=1):
                    if not isinstance(maint_data, dict):
                        validation_errors.append(
                            f'Invalid maintenance record at entry {index}')
                        continue

                    maint_date_raw = maint_data.get('maintenance_date')
                    if maint_date_raw:
                        try:
                            datetime.strptime(
                                maint_date_raw, '%Y-%m-%d').date()
                        except (TypeError, ValueError):
                            validation_errors.append(
                                f'Invalid Maintenance date format (entry {index})')

                    maint_due_raw = maint_data.get('maintenance_due_date')
                    if maint_due_raw:
                        try:
                            maint_due_date = datetime.strptime(
                                maint_due_raw, '%Y-%m-%d').date()
                            if maint_due_date < today:
                                validation_errors.append(
                                    f'Maintenance due date cannot be in the past (entry {index})')
                        except (TypeError, ValueError):
                            validation_errors.append(
                                f'Invalid Maintenance due date format (entry {index})')

            if validation_errors:
                return jsonify({
                    'success': False,
                    'error': '; '.join(validation_errors)
                }), 400

            # Store old values for history before updating
            old_values = equipment.to_dict()

            # Update equipment fields
            equipment.asset_id = data.get('asset_id', '').strip()
            if data.get('sl_no') is not None:
                equipment.sl_no = int(
                    data.get('sl_no')) if data.get('sl_no') else None
            equipment.type = data.get('type', '').strip()
            equipment.name = data.get('name', '').strip()
            equipment.make = data.get(
                'make', '').strip() if data.get('make') else None
            equipment.model_no = data.get(
                'model_no', '').strip() if data.get('model_no') else None
            equipment.serial_no = data.get(
                'serial_no', '').strip() if data.get('serial_no') else None
            equipment.location = data.get(
                'location', '').strip() if data.get('location') else None
            equipment.test_type = data.get('test_type', '').strip()
            equipment.eou_status = data.get(
                'eou_status', '').strip() if data.get('eou_status') else None
            equipment.test_name = data.get(
                'test_name', '').strip() if data.get('test_name') else None
            equipment.calibration_required = data.get(
                'calibration_required') or None
            equipment.calibration_frequency = data.get(
                'calibration_frequency') or None
            equipment.calibration_date = parse_date_field(
                data.get('calibration_date'))
            equipment.calibration_due_date = parse_date_field(
                data.get('calibration_due_date'))
            # Automatically calculate calibration_status_col based on calibration_due_date
            calibration_due_date_parsed = parse_date_field(
                data.get('calibration_due_date'))
            equipment.calibration_status_col = _calculate_calibration_status(
                calibration_due_date_parsed,
                data.get('calibration_required')
            ) or data.get('calibration_status_col') or None
            equipment.ic_required = data.get('ic_required') or None
            equipment.ic_date = parse_date_field(data.get('ic_date'))
            equipment.ic_due_date = parse_date_field(data.get('ic_due_date'))
            equipment.manufacturer_calibration_params = data.get(
                'manufacturer_calibration_params') or None
            equipment.calibration_agency_params = data.get(
                'calibration_agency_params') or None
            equipment.document_link = data.get(
                'document_link', '').strip() if data.get('document_link') else None
            equipment.status = data.get('status')

            # Handle maintenance records - delete existing and create new ones
            Maintenance.query.filter_by(equipment_id=equipment.id).delete()
            maintenance_records = data.get('maintenance_records', [])
            first_maintenance_entry = next(
                (
                    maint_data for maint_data in maintenance_records
                    if maint_data.get('maintenance_required') or maint_data.get('maintenance_date') or maint_data.get('maintenance_due_date')
                ),
                None
            )
            equipment.maintenance_required = (
                first_maintenance_entry.get('maintenance_required') or None
                if first_maintenance_entry else None
            )
            equipment.maintenance_date = parse_date_field(
                first_maintenance_entry.get('maintenance_date')
            ) if first_maintenance_entry else None
            equipment.maintenance_due_date = parse_date_field(
                first_maintenance_entry.get('maintenance_due_date')
            ) if first_maintenance_entry else None
            if maintenance_records:
                for maint_data in maintenance_records:
                    if maint_data.get('maintenance_required') or maint_data.get('maintenance_date') or maint_data.get('maintenance_due_date'):
                        maintenance = Maintenance(
                            equipment_id=equipment.id,
                            maintenance_required=maint_data.get(
                                'maintenance_required') or None,
                            maintenance_date=parse_date_field(
                                maint_data.get('maintenance_date')),
                            maintenance_due_date=parse_date_field(
                                maint_data.get('maintenance_due_date'))
                        )
                        db.session.add(maintenance)

            # Log update to history
            new_values = equipment.to_dict()
            log_equipment_history(equipment.id, 'updated', old_values=old_values,
                                  new_values=new_values, notes='Equipment updated')

            db.session.commit()

            return jsonify({
                'success': True,
                'message': 'Equipment updated successfully'
            })
        except Exception as e:
            logger.error("Error in equipment update API: %s", e, exc_info=True)
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': f'Error updating equipment: {str(e)}'
            }), 500

    @flask_app.route('/api/equipment/<int:equipment_id>', methods=['DELETE'])
    @login_required
    def api_equipment_delete(equipment_id):
        """API endpoint for deleting equipment."""
        # Check if user is admin or lab engineer
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin or Lab Engineer privileges required.'
            }), 403

        try:
            equipment = db.session.get(Equipment, equipment_id)
            if not equipment:
                return jsonify({
                    'success': False,
                    'error': 'Equipment not found'
                }), 404

            db.session.delete(equipment)
            db.session.commit()

            return jsonify({
                'success': True,
                'message': 'Equipment deleted successfully'
            })
        except Exception as e:
            logger.error("Error in equipment delete API: %s", e)
            return jsonify({
                'success': False,
                'error': 'Error deleting equipment'
            }), 500

    @flask_app.route('/users')
    @login_required
    def user_management():
        """Display user management page for admins."""
        # Check if user is admin
        if current_user.role != 'admin':
            flash('Access denied. Admin privileges required.', 'error')
            return redirect(url_for('index'))

        try:
            # Get pagination parameters
            page = request.args.get('page', 1, type=int)
            per_page = request.args.get('per_page', 10, type=int)
            search_term = request.args.get('search', '')
            role_filter = request.args.get('role', '')

            # Validate per_page to prevent abuse
            if per_page not in [5, 10, 20, 50, 100]:
                per_page = 10

            # Build query with filters
            query = User.query
            if search_term:
                search_filter = f'%{search_term}%'
                query = query.filter(
                    (User.username.ilike(search_filter)) |
                    (User.email.ilike(search_filter))
                )

            if role_filter:
                query = query.filter(User.role == role_filter)

            # Get users with pagination
            pagination = query.order_by(User.created_at.desc()).paginate(
                page=page, per_page=per_page, error_out=False
            )

            # Calculate statistics
            total_users = User.query.count()
            admin_users = User.query.filter_by(role='admin').count()
            regular_users = User.query.filter_by(role='user').count()
            active_users = User.query.filter_by(is_active=True).count()

            statistics = {
                'total': total_users,
                'admins': admin_users,
                'users': regular_users,
                'active': active_users
            }

            return render_template('users.html', pagination=pagination,
                                   statistics=statistics, search_term=search_term,
                                   role_filter=role_filter)
        except Exception as e:
            logger.error("Error loading user management page: %s", e)
            flash('Error loading user information', 'error')
            return redirect(url_for('index'))

    @flask_app.route('/api/users/<int:user_id>/role', methods=['PUT'])
    @login_required
    def api_update_user_role(user_id):
        """API endpoint for updating user role."""
        # Check if user is admin
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin privileges required.'
            }), 403

        try:
            data = request.get_json()
            new_role = data.get('role')

            if not new_role or new_role not in ['user', 'lab_engineer', 'admin']:
                return jsonify({
                    'success': False,
                    'error': 'Invalid role. Must be "user", "lab_engineer", or "admin".'
                }), 400

            # Prevent admin from changing their own role
            if user_id == current_user.id:
                return jsonify({
                    'success': False,
                    'error': 'Cannot change your own role.'
                }), 400

            user = db.session.get(User, user_id)
            if not user:
                return jsonify({
                    'success': False,
                    'error': 'User not found.'
                }), 404

            old_role = user.role
            user.role = new_role
            db.session.commit()

            logger.info(
                f"User {current_user.username} changed role of user {user.username} from {old_role} to {new_role}")

            return jsonify({
                'success': True,
                'message': f'User role updated from {old_role} to {new_role}',
                'data': {
                    'user_id': user.id,
                    'username': user.username,
                    'old_role': old_role,
                    'new_role': new_role
                }
            })
        except Exception as e:
            logger.error("Error updating user role: %s", e)
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': 'Error updating user role'
            }), 500

    @flask_app.route('/api/users/<int:user_id>/status', methods=['PUT'])
    @login_required
    def api_update_user_status(user_id):
        """API endpoint for updating user active status."""
        # Check if user is admin
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin privileges required.'
            }), 403

        try:
            data = request.get_json()
            new_status = data.get('is_active')

            if new_status is None:
                return jsonify({
                    'success': False,
                    'error': 'Status is required.'
                }), 400

            # Prevent admin from deactivating themselves
            if user_id == current_user.id:
                return jsonify({
                    'success': False,
                    'error': 'Cannot change your own status.'
                }), 400

            user = db.session.get(User, user_id)
            if not user:
                return jsonify({
                    'success': False,
                    'error': 'User not found.'
                }), 404

            old_status = user.is_active
            user.is_active = new_status
            db.session.commit()

            status_text = 'activated' if new_status else 'deactivated'
            logger.info(
                f"User {current_user.username} {status_text} user {user.username}")

            return jsonify({
                'success': True,
                'message': f'User {status_text} successfully',
                'data': {
                    'user_id': user.id,
                    'username': user.username,
                    'old_status': old_status,
                    'new_status': new_status
                }
            })
        except Exception as e:
            logger.error("Error updating user status: %s", e)
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': 'Error updating user status'
            }), 500

    @flask_app.route('/api/users/<int:user_id>', methods=['GET', 'PUT', 'DELETE'])
    @login_required
    def api_user_detail(user_id: int):
        """API endpoint for user details, update, and deletion."""
        # Check if user is admin
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Access denied. Admin privileges required.'
            }), 403

        try:
            user = db.session.get(User, user_id)
            if not user:
                return jsonify({
                    'success': False,
                    'error': 'User not found.'
                }), 404

            # Handle GET request - return user details
            if request.method == 'GET':
                return jsonify({
                    'success': True,
                    'data': {
                        'id': user.id,
                        'username': user.username,
                        'email': user.email,
                        'role': user.role,
                        'is_active': user.is_active,
                        'created_at': user.created_at.isoformat() if user.created_at else None,
                        'last_login': user.last_login.isoformat() if user.last_login else None
                    }
                })

            # Handle DELETE request
            if request.method == 'DELETE':
                # Prevent admin from deleting themselves
                if user_id == current_user.id:
                    return jsonify({
                        'success': False,
                        'error': 'Cannot delete your own account.'
                    }), 400

                username = user.username
                db.session.delete(user)
                db.session.commit()

                logger.info(
                    f"User {username} deleted by admin {current_user.username}")

                return jsonify({
                    'success': True,
                    'message': f'User {username} deleted successfully'
                })

            # Handle PUT request - update user
            if request.method == 'PUT':
                data = request.get_json()

                # Validate required fields
                username = data.get('username', '').strip()
                email = data.get('email', '').strip()
                role = data.get('role')
                is_active = data.get('is_active')

                if not username:
                    return jsonify({
                        'success': False,
                        'error': 'Username is required.'
                    }), 400

                if not email:
                    return jsonify({
                        'success': False,
                        'error': 'Email is required.'
                    }), 400

                if role not in ['user', 'lab_engineer', 'admin']:
                    return jsonify({
                        'success': False,
                        'error': 'Invalid role. Must be "user", "lab_engineer", or "admin".'
                    }), 400

                if is_active is None:
                    return jsonify({
                        'success': False,
                        'error': 'Status is required.'
                    }), 400

                # Check if username already exists (excluding current user)
                existing_user = User.query.filter(
                    User.username == username,
                    User.id != user_id
                ).first()
                if existing_user:
                    return jsonify({
                        'success': False,
                        'error': 'Username already exists.'
                    }), 400

                # Check if email already exists (excluding current user)
                existing_email = User.query.filter(
                    User.email == email,
                    User.id != user_id
                ).first()
                if existing_email:
                    return jsonify({
                        'success': False,
                        'error': 'Email already exists.'
                    }), 400

                # Update user
                old_username = user.username
                old_email = user.email
                old_role = user.role
                old_status = user.is_active

                user.username = username
                user.email = email
                user.role = role
                user.is_active = is_active

                db.session.commit()

                logger.info(
                    f"User {old_username} updated by admin {current_user.username}: "
                    f"username={old_username}->{username}, "
                    f"email={old_email}->{email}, "
                    f"role={old_role}->{role}, "
                    f"status={old_status}->{is_active}"
                )

                return jsonify({
                    'success': True,
                    'message': 'User updated successfully',
                    'data': {
                        'id': user.id,
                        'username': user.username,
                        'email': user.email,
                        'role': user.role,
                        'is_active': user.is_active
                    }
                })

            # If method is not GET, PUT, or DELETE, return method not allowed
            return jsonify({
                'success': False,
                'error': 'Method not allowed'
            }), 405

        except Exception as e:
            logger.error("Error in user API: %s", e)
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': 'Error processing user request'
            }), 500

    @flask_app.route('/api/request/<int:request_id>', methods=['GET', 'DELETE'])
    @login_required
    def api_request_detail(request_id):
        """API endpoint for test request details and deletion."""
        try:
            test_request = db.session.get(TestRequest, request_id)

            if not test_request:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            # Handle DELETE request
            if request.method == 'DELETE':
                # Check if user owns the request or is admin
                if test_request.user_id != current_user.id and current_user.role != 'admin':
                    return jsonify({
                        'success': False,
                        'error': 'Access denied. You can only delete your own requests.'
                    }), 403

                # Delete associated files
                try:
                    # Delete uploaded test plan file
                    if test_request.filename:
                        file_path = os.path.join(
                            flask_app.config['UPLOAD_FOLDER'], test_request.filename)
                        if os.path.exists(file_path):
                            os.remove(file_path)

                    # Delete generated files
                    generated_files = test_request.get_generated_files()
                    for file_name in generated_files:
                        file_path = os.path.join(
                            flask_app.config['OUTPUT_FOLDER'], file_name)
                        if os.path.exists(file_path):
                            os.remove(file_path)
                except Exception as e:
                    logger.warning(
                        "Error deleting files for request %d: %s", request_id, e)

                # Delete from database
                db.session.delete(test_request)
                db.session.commit()

                logger.info("Test request %d deleted by user %s",
                            request_id, current_user.username)

                return jsonify({
                    'success': True,
                    'message': 'Test request deleted successfully'
                })

            # Handle GET request
            return jsonify({
                'success': True,
                'data': test_request.to_dict()
            })
        except Exception as e:
            logger.error("Error in request API: %s", e)
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': 'Error processing request'
            }), 500

    # Admin API Routes

    def _parse_iso_date(value):
        if not value:
            return None
        try:
            return datetime.strptime(value, '%Y-%m-%d').date()
        except ValueError:
            return None

    def _parse_iso_time(value):
        if not value:
            return None
        try:
            return datetime.strptime(value.strip(), '%H:%M').time()
        except ValueError:
            return None

    def _format_date_display(value):
        if not value:
            return 'N/A'
        try:
            return datetime.strptime(value, '%Y-%m-%d').strftime('%b %d, %Y')
        except ValueError:
            return value

    def _format_time_display(value):
        if not value:
            return '--'
        try:
            return datetime.strptime(value, '%H:%M').strftime('%I:%M %p')
        except ValueError:
            return value

    def _get_assigned_tests_context():
        """Build test plan list and statistics for assigned test views."""
        if current_user.role not in ['admin', 'lab_engineer']:
            raise PermissionError('Not authorized to view assigned tests')

        assigned_statuses = ['Assigned Lab Engineer', 'Update plan',
                             'Need More Information', 'Test Plan Approved', 'Test Plan To Approve', 'Draft Report', 'Datasheet Uploaded', 'Proceed Report', 'Admin Sign Off', 'Completed', 'Assigned', 'Test Schedule In Progress', 'In Progress', 'Report Uploaded', 'report_uploaded', 'Peer Review']
        approved_statuses = ['Approved', 'Test Plan Approved', 'Completed']

        assigned_requests = EMCRequest.query.options(
            joinedload(EMCRequest.service_types),
            joinedload(EMCRequest.assigned_engineer)
        ).filter(
            EMCRequest.status.in_(assigned_statuses)
        ).order_by(
            EMCRequest.updated_at.desc()
        ).all()

        logger.info(f"Found {len(assigned_requests)} assigned request candidates")

        test_plans = []
        for request in assigned_requests:
            try:
                selected_tests = _get_active_selected_test_labels(request)

                assigned_display = request.assigned_engineer_name
                if not assigned_display and request.assigned_engineer:
                    assigned_display = request.assigned_engineer.username

                # Decode service types (Type of Service Requested) from JSON
                service_types = _extract_service_types(request)

                # Fetch actual PlannerEntry records for this test request.
                # Some legacy/current rows are linked only by TCO ID, so include both.
                planner_entry_filters = [PlannerEntry.test_request_id == request.id]
                if request.tco_id:
                    planner_entry_filters.append(PlannerEntry.tco_id == request.tco_id)

                planner_entries = PlannerEntry.query.filter(
                    db.or_(*planner_entry_filters)
                ).order_by(
                    PlannerEntry.start_date,
                    PlannerEntry.start_time
                ).all()

                # âœ… FIX: Filter out cancelled entries and handle rescheduled tests
                # For each test_name, keep only the most recent non-cancelled entry
                active_entries_by_test = {}
                for entry in planner_entries:
                    if entry.status == 'cancelled':
                        continue
                    test_key = (entry.test_name or '').strip()
                    if test_key:
                        # Keep the most recent entry for this test
                        if test_key not in active_entries_by_test or entry.updated_at > active_entries_by_test[test_key].updated_at:
                            active_entries_by_test[test_key] = entry
                
                # If all entries are cancelled, use the most recent one anyway
                if not active_entries_by_test and planner_entries:
                    most_recent = max(planner_entries, key=lambda e: e.updated_at)
                    active_entries_by_test[(most_recent.test_name or '').strip()] = most_recent
                
                filtered_planner_entries = list(active_entries_by_test.values())

                parsed_assignments = []
                assigned_test_keys = set()

                # Helper functions for formatting
                def safe_date_format(date_obj, format_str='%Y-%m-%d'):
                    if date_obj is None:
                        return ''
                    if isinstance(date_obj, str):
                        return date_obj
                    try:
                        return date_obj.strftime(format_str)
                    except AttributeError:
                        return str(date_obj)

                def safe_time_format(time_obj):
                    if time_obj is None:
                        return ''
                    if isinstance(time_obj, str):
                        return time_obj
                    try:
                        return time_obj.strftime('%H:%M')
                    except AttributeError:
                        return str(time_obj)

                def safe_datetime_format(dt_obj):
                    if dt_obj is None:
                        return None
                    if isinstance(dt_obj, str):
                        return dt_obj
                    try:
                        return dt_obj.isoformat()
                    except AttributeError:
                        return str(dt_obj)

                def is_meaningful_assignment(assignment_payload):
                    if not isinstance(assignment_payload, dict):
                        return False

                    test_name = str(
                        assignment_payload.get('test_name') or ''
                    ).strip()
                    engineer_name = str(
                        assignment_payload.get('engineer_name') or ''
                    ).strip()
                    engineer_id = assignment_payload.get('engineer_id')
                    start_date = str(
                        assignment_payload.get('start_date') or ''
                    ).strip()
                    end_date = str(
                        assignment_payload.get('end_date') or ''
                    ).strip()
                    total_hours = assignment_payload.get('total_hours')

                    has_assignment_details = any([
                        engineer_name,
                        engineer_id not in (None, ''),
                        start_date,
                        end_date,
                        total_hours not in (None, ''),
                    ])
                    return bool(test_name and has_assignment_details)

                def assignment_matches_current_engineer(assignment_payload):
                    if current_user.role != 'lab_engineer':
                        return True

                    engineer_id = assignment_payload.get('engineer_id')
                    if engineer_id not in (None, ''):
                        try:
                            return int(engineer_id) == current_user.id
                        except (TypeError, ValueError):
                            pass

                    engineer_name = str(
                        assignment_payload.get('engineer_name') or ''
                    ).strip()
                    current_engineer_name = str(
                        current_user.username or ''
                    ).strip()
                    return bool(
                        engineer_name and current_engineer_name and
                        engineer_name.casefold() == current_engineer_name.casefold()
                    )

                def normalize_review_status(raw_status):
                    status_text = str(raw_status or '').strip()
                    if not status_text:
                        return 'Assigned Lab Engineer'

                    status_map = {
                        'assigned': 'Assigned Lab Engineer',
                        'assigned lab engineer': 'Assigned Lab Engineer',
                        'update plan': 'Update plan',
                        'test plan approved': 'Test Plan Approved',
                        'test plan to approve': 'Test Plan To Approve',
                        'in progress': 'In Progress',
                        'in_progress': 'In Progress',
                        'test schedule in progress': 'In Progress',
                        'draft report': 'Draft Report',
                        'proceed report': 'Proceed Report',
                        'peer review': 'Peer Review',
                        'datasheet uploaded': 'Datasheet Uploaded',
                        'datasheet_uploaded': 'Datasheet Uploaded',
                        'report uploaded': 'Draft Report',
                        'report_uploaded': 'Draft Report',
                        'admin sign off': 'Admin Sign Off',
                        'need more information': 'Need More Information',
                        'completed': 'Completed',
                        'cancelled': 'Cancelled',
                        'rejected': 'Rejected',
                    }
                    return status_map.get(status_text.casefold(), status_text)

                def review_status_key(status_label):
                    if status_label == 'Need More Information':
                        return 'need_more_info'
                    if status_label == 'Completed':
                        return 'completed'
                    if status_label == 'Cancelled':
                        return 'cancelled'
                    if status_label == 'Rejected':
                        return 'rejected'
                    if status_label in {
                        'Assigned Lab Engineer',
                        'Update plan',
                        'Test Plan Approved',
                        'Test Plan To Approve',
                    }:
                        return 'assigned'
                    return 'in_progress'

                def derive_review_status(parent_status, remaining_test_names, assignment_payloads):
                    normalized_parent_status = normalize_review_status(parent_status)
                    if normalized_parent_status in {
                        'Need More Information',
                        'Completed',
                        'Cancelled',
                        'Rejected',
                    }:
                        return normalized_parent_status

                    status_priority = {
                        'Assigned Lab Engineer': 0,
                        'Update plan': 1,
                        'Test Plan Approved': 2,
                        'Test Plan To Approve': 3,
                        'In Progress': 4,
                        'Draft Report': 5,
                        'Proceed Report': 6,
                        'Peer Review': 7,
                        'Datasheet Uploaded': 8,
                        'Admin Sign Off': 9,
                        'Completed': 10,
                    }

                    candidate_statuses = []
                    if remaining_test_names:
                        candidate_statuses.append('Assigned Lab Engineer')

                    for assignment_payload in assignment_payloads:
                        normalized_assignment_status = normalize_review_status(
                            assignment_payload.get('status')
                        )
                        if normalized_assignment_status in status_priority:
                            candidate_statuses.append(normalized_assignment_status)

                    if not candidate_statuses:
                        return normalized_parent_status or 'Assigned Lab Engineer'

                    return min(
                        candidate_statuses,
                        key=lambda status_label: status_priority.get(status_label, 999)
                    )

                # Use PlannerEntry records if available, otherwise fall back to JSON
                if filtered_planner_entries:
                    for entry in filtered_planner_entries:
                        try:
                            peer_reviewer_name = None
                            peer_reviewer_user_id = getattr(
                                entry, 'peer_reviewer_user_id', None
                            )
                            if peer_reviewer_user_id:
                                peer_reviewer_user = db.session.get(User, 
                                    peer_reviewer_user_id
                                )
                                if peer_reviewer_user:
                                    peer_reviewer_name = (
                                        peer_reviewer_user.username
                                    )

                            assignment_payload = {
                                'id': entry.id,
                                'test_name': entry.test_name,
                                'engineer_name': entry.test_person_name,
                                'engineer_id': entry.engineer_user_id,
                                'start_date': safe_date_format(entry.start_date),
                                'end_date': safe_date_format(entry.end_date),
                                'start_time': safe_time_format(entry.start_time),
                                'end_time': safe_time_format(entry.end_time),
                                'total_hours': entry.total_hours,
                                'start_date_display': safe_date_format(entry.start_date) or 'N/A',
                                'end_date_display': safe_date_format(entry.end_date) or 'N/A',
                                'start_time_display': safe_time_format(entry.start_time) or 'N/A',
                                'end_time_display': safe_time_format(entry.end_time) or 'N/A',
                                'status': entry.status,
                                'datasheet_file_path': getattr(entry, 'datasheet_file_path', None),
                                'datasheet_uploaded_at': safe_datetime_format(
                                    getattr(entry, 'datasheet_uploaded_at', None)
                                ),
                                'datasheet_uploaded_by': getattr(
                                    entry, 'datasheet_uploaded_by', None
                                ),
                                'peer_reviewer_user_id': peer_reviewer_user_id,
                                'peer_review_assigned_at': safe_datetime_format(
                                    getattr(entry, 'peer_review_assigned_at', None)
                                ),
                                'peer_reviewer_name': peer_reviewer_name,
                                'datasheet_comments': getattr(
                                    entry, 'datasheet_comments', None
                                ),
                            }
                            if is_meaningful_assignment(assignment_payload):
                                assignment_key = _normalize_assignment_test_key(
                                    assignment_payload.get('test_name'))
                                if assignment_key:
                                    assigned_test_keys.add(assignment_key)
                                parsed_assignments.append(assignment_payload)
                        except Exception as e:
                            logger.error(
                                f"Error processing planner entry {entry.id}: {e}")
                            import traceback
                            logger.error(traceback.format_exc())
                            continue

                elif request.test_assignments:
                    # Fallback to JSON parsing if no PlannerEntry records exist
                    try:
                        raw_assignments = json.loads(request.test_assignments)
                    except (TypeError, ValueError) as exc:
                        logger.warning(
                            'Unable to parse test_assignments for request %s: %s', request.id, exc
                        )
                        raw_assignments = []

                    for assignment in raw_assignments:
                        assignment_payload = {
                            'id': None,
                            'test_name': assignment.get('test_name', 'Test'),
                            'engineer_name': assignment.get('engineer_name', 'Unassigned'),
                            'engineer_id': assignment.get('engineer_id'),
                            'start_date': assignment.get('start_date', ''),
                            'end_date': assignment.get('end_date', ''),
                            'start_time': assignment.get('start_time', ''),
                            'end_time': assignment.get('end_time', ''),
                            'total_hours': assignment.get('total_hours'),
                            'start_date_display': _format_date_display(assignment.get('start_date')),
                            'end_date_display': _format_date_display(assignment.get('end_date')),
                            'start_time_display': _format_time_display(assignment.get('start_time')),
                            'end_time_display': _format_time_display(assignment.get('end_time')),
                            'status': 'in_progress',
                            'datasheet_file_path': None,
                            'datasheet_uploaded_at': None,
                            'datasheet_uploaded_by': None,
                            'peer_reviewer_user_id': None,
                            'peer_review_assigned_at': None,
                            'peer_reviewer_name': None,
                            'datasheet_comments': None,
                        }
                        if is_meaningful_assignment(assignment_payload):
                            assignment_key = _normalize_assignment_test_key(
                                assignment_payload.get('test_name'))
                            if assignment_key:
                                assigned_test_keys.add(assignment_key)
                            parsed_assignments.append(assignment_payload)

                remaining_tests = []
                for selected_test in selected_tests:
                    selected_key = _normalize_assignment_test_key(selected_test)
                    if selected_key and selected_key in assigned_test_keys:
                        continue
                    remaining_tests.append(selected_test)

                has_any_assigned_tests = bool(parsed_assignments)
                has_partially_assigned_tests = (
                    has_any_assigned_tests and bool(remaining_tests)
                )

                visible_assignments = parsed_assignments
                visible_remaining_tests = (
                    remaining_tests if has_partially_assigned_tests else []
                )
                if current_user.role == 'lab_engineer':
                    visible_assignments = [
                        assignment for assignment in parsed_assignments
                        if assignment_matches_current_engineer(assignment)
                    ]
                    visible_remaining_tests = []
                    if (
                        request.assigned_engineer_id == current_user.id and
                        has_partially_assigned_tests
                    ):
                        visible_remaining_tests = remaining_tests
                    if not visible_assignments and not visible_remaining_tests:
                        continue

                report_already_uploaded = any(
                    getattr(e, 'report_file_path', None)
                    for e in filtered_planner_entries
                )

                # âœ… FIX: Get report fields from filtered_planner_entries directly
                report_file_path = next(
                    (
                        e.report_file_path
                        for e in filtered_planner_entries
                        if getattr(e, 'report_file_path', None)
                    ),
                    None
                )
                report_comments = next(
                    (
                        e.report_comments
                        for e in filtered_planner_entries
                        if getattr(e, 'report_file_path', None)
                    ),
                    None
                )
                report_uploaded_at = next(
                    (
                        e.report_uploaded_at
                        for e in planner_entries
                        if getattr(e, 'report_file_path', None)
                    ),
                    None
                )

                show_upload_report = False
                show_prepare_report = False
                if parsed_assignments:
                    all_terminal = all(
                        a['status'] in ('cancelled', 'datasheet_uploaded')
                        for a in parsed_assignments
                    )
                    # At least one datasheet_uploaded assignment that also has a datasheet file
                    has_datasheet_uploaded_with_file = any(
                        a['status'] == 'datasheet_uploaded' and bool(
                            a.get('datasheet_file_path'))
                        for a in parsed_assignments
                    )
                    # Every datasheet_uploaded assignment must have a datasheet file
                    all_datasheet_uploaded_have_file = all(
                        bool(a.get('datasheet_file_path'))
                        for a in parsed_assignments
                        if a['status'] == 'datasheet_uploaded'
                    )
                    # âœ… FIX: Only show upload button if report not already uploaded
                    # Also require that EVERY test selected on the request has an
                    # assignment: the consolidated report covers all of them, so a
                    # selected-but-never-scheduled test must not let it be generated
                    # early (remaining_tests holds exactly those).
                    show_upload_report = (
                        all_terminal
                        and has_datasheet_uploaded_with_file
                        and all_datasheet_uploaded_have_file
                        and not remaining_tests
                        and not report_already_uploaded
                    )
                    # The report WIZARD needs the datasheets finished, and nothing
                    # more. Two conditions above are deliberately relaxed:
                    #
                    #  * report_already_uploaded - the commonest moment to want the
                    #    wizard is just after reading a generated report and finding
                    #    the EUT details blank, and it is the only route that
                    #    collects them.
                    #  * all_terminal - generating a report moves every assignment
                    #    from 'datasheet_uploaded' to 'report_uploaded', which is
                    #    not in that tuple. So the button vanished the instant a
                    #    report existed, i.e. precisely when it is needed. A
                    #    reported test is at least as finished as an uploaded one.
                    wizard_terminal = all(
                        a['status'] in ('cancelled', 'datasheet_uploaded',
                                        'report_uploaded')
                        for a in parsed_assignments
                    )
                    has_any_datasheet_file = any(
                        a['status'] in ('datasheet_uploaded', 'report_uploaded')
                        and bool(a.get('datasheet_file_path'))
                        for a in parsed_assignments
                    )
                    show_prepare_report = (
                        wizard_terminal
                        and has_any_datasheet_file
                        and not remaining_tests
                    )

                aggregated_status = derive_review_status(
                    request.status,
                    remaining_tests,
                    parsed_assignments
                )

                shared_review_thread = _get_combined_review_comment_thread(request)
                latest_shared_comment = shared_review_thread[-1] if shared_review_thread else None

                test_plans.append({
                    'id': request.id,
                    'tco_id': request.tco_id or f'REQ-{request.id}',
                    'job_id': request.job_id,
                    'job_number': request.job_number or '',
                    'name': request.product_name or 'Unnamed Product',
                    'product_name': request.product_name or 'Unnamed Product',
                    'manufacturer': request.manufacturer or 'N/A',
                    'model_number': request.model_number or 'N/A',
                    'project': request.manufacturer or 'N/A',
                    'status': aggregated_status,
                    'status_key': review_status_key(aggregated_status),
                    'status_label': aggregated_status,
                    'status_display': aggregated_status,
                    'parent_status': request.status or '',
                    'requester_name': request.requester_name or 'Unknown',
                    'requester_email': request.requester_email or '',
                    'submitted_date': request.submitted_at or request.created_at,
                    'created_at': safe_datetime_format(request.created_at),
                    'assigned_engineer_display': assigned_display,
                    'assigned_engineer_id': request.assigned_engineer_id,
                    'assignment_priority': request.assignment_priority,
                    'assignment_due_date': safe_datetime_format(request.assignment_due_date),
                    'assignment_notes': request.assignment_notes,
                    'selected_tests': selected_tests,
                    'test_assignments': visible_assignments,
                    'remaining_tests': visible_remaining_tests,
                    'remaining_tests_count': len(visible_remaining_tests),
                    'has_review_messages': len(shared_review_thread) > 0,
                    'review_comments': latest_shared_comment.get('comment') if latest_shared_comment else None,
                    'reviewed_by': latest_shared_comment.get('username') if latest_shared_comment else getattr(request, 'reviewed_by', None),
                    'reviewed_at': latest_shared_comment.get('created_at') if latest_shared_comment else safe_datetime_format(getattr(request, 'reviewed_at', None)),
                    'service_types': service_types,
                    # âœ… FIX: show_upload_report now correctly hides when report is uploaded
                    'show_upload_report': show_upload_report,
                    'show_prepare_report': show_prepare_report,
                    # âœ… FIX: report fields now correctly read from PlannerEntry
                    'report_file_path': report_file_path,
                    'report_comments': report_comments,
                    'report_uploaded_at': safe_datetime_format(report_uploaded_at),
                    'sample_condition': request.sample_condition,
                    'capability_available': request.capability_available,
                    'sample_received_date': safe_date_format(request.sample_received_date),
                    'test_duration': request.test_duration,
                    'test_commencement_date': safe_date_format(request.test_commencement_date),
                    'test_completion_date': safe_date_format(request.test_completion_date),
                    'lab_manager_name': request.lab_manager_name,
                    'lab_manager_date': safe_date_format(request.lab_manager_date),
                    'lab_manager_signature': request.lab_manager_signature,
                    'lab_manager_signed_at': safe_datetime_format(request.lab_manager_signed_at),
                })

            except Exception as e:
                logger.error(f"Error processing request {request.id}: {e}")
                import traceback
                logger.error(traceback.format_exc())
                continue

        if current_user.role == 'lab_engineer':
            under_review_count = EMCRequest.query.filter_by(
                status='At Review').count()
            approved_count = EMCRequest.query.filter(
                EMCRequest.status.in_(approved_statuses)
            ).count()
            rejected_count = EMCRequest.query.filter_by(
                status='Rejected').count()
            assigned_count = len(test_plans)
        else:
            under_review_count = EMCRequest.query.filter_by(
                status='At Review').count()
            approved_count = EMCRequest.query.filter(
                EMCRequest.status.in_(approved_statuses)
            ).count()
            rejected_count = EMCRequest.query.filter_by(
                status='Rejected').count()
            assigned_count = EMCRequest.query.filter(
                EMCRequest.status.in_(assigned_statuses)
            ).count()

        statistics = {
            'under_review': under_review_count,
            'approved': approved_count,
            'rejected': rejected_count,
            'assigned': assigned_count
        }

        test_plans.sort(key=lambda plan: (
            # completed AND cancelled/rejected plans sort after everything active
            1 if _is_terminal_request_status(plan['status']) else 0,
            *_job_number_sort_key(plan),
        ))
        return test_plans, statistics

    @flask_app.route('/api/planner', methods=['GET', 'POST'])
    @login_required
    def planner_entries_api():
        """List or create planner entries stored in MySQL."""
        try:
            if request.method == 'GET':
                # Role-based filtering
                query = PlannerEntry.query

                if current_user.role == 'admin':
                    # Admin can filter by engineer_id query param, or see all
                    engineer_filter = request.args.get('engineer_id')
                    logger.info(
                        f'Admin planner filter request - engineer_id: {engineer_filter}')
                    if engineer_filter:
                        try:
                            engineer_id = int(engineer_filter)
                            query = query.filter_by(
                                engineer_user_id=engineer_id)
                            logger.info(
                                f'Filtering planner entries for engineer_id: {engineer_id}')
                        except (ValueError, TypeError) as e:
                            logger.warning(
                                f'Invalid engineer_id filter: {engineer_filter}, error: {e}')
                            # Invalid filter, show all
                elif current_user.role == 'lab_engineer':
                    # Lab engineers see only their own events
                    query = query.filter_by(engineer_user_id=current_user.id)
                else:
                    # Regular users see only events for their TCOs
                    # Get TCO IDs from test requests created by this user
                    user_tco_ids = db.session.query(EMCRequest.tco_id).filter_by(
                        user_id=current_user.id
                    ).filter(EMCRequest.tco_id.isnot(None)).all()
                    tco_id_list = [tco[0] for tco in user_tco_ids if tco[0]]
                    if tco_id_list:
                        query = query.filter(
                            PlannerEntry.tco_id.in_(tco_id_list))
                    else:
                        # No TCOs found, return empty result
                        # Impossible condition
                        query = query.filter(PlannerEntry.id == -1)

                entries = query.order_by(
                    PlannerEntry.start_date.asc(),
                    PlannerEntry.end_date.asc()
                ).all()
                response = jsonify({
                    'success': True,
                    'data': [entry.to_dict() for entry in entries]
                })
                response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
                response.headers['Pragma'] = 'no-cache'
                response.headers['Expires'] = '0'
                return response

            # POST method - Create new planner entry
            data = request.get_json() or {}

            status = data.get('status', 'in_progress').strip().lower()
            valid_statuses = ['in_progress', 'completed', 'cancelled']
            if status not in valid_statuses:
                status = 'in_progress'

            # Validate test_person_name
            test_person_name = (data.get('test_person_name') or '').strip()
            if not test_person_name:
                return jsonify({
                    'success': False,
                    'error': 'test_person_name is required'
                }), 400

            # Validate test_name
            test_name = (data.get('test_name') or '').strip()
            if not test_name:
                return jsonify({
                    'success': False,
                    'error': 'test_name is required'
                }), 400

            # Get event_type early to check if TCO ID is required
            event_type = (data.get('event_type') or 'test').strip().lower()
            # TCO ID is optional for selected non-test event types
            tco_id = (data.get('tco_id') or '').strip()
            if not tco_id and event_type not in NON_TEST_EVENT_TYPES:
                return jsonify({
                    'success': False,
                    'error': 'tco_id is required for this event type'
                }), 400

            # Validate dates
            start_date_str = data.get('start_date')
            end_date_str = data.get('end_date')
            start_date = _parse_iso_date(start_date_str)
            end_date = _parse_iso_date(end_date_str)

            if not start_date:
                return jsonify({
                    'success': False,
                    'error': 'start_date must be provided in YYYY-MM-DD format'
                }), 400

            if not end_date:
                return jsonify({
                    'success': False,
                    'error': 'end_date must be provided in YYYY-MM-DD format'
                }), 400

            if end_date < start_date:
                return jsonify({
                    'success': False,
                    'error': 'end_date cannot be earlier than start_date'
                }), 400

            # Handle all-day events and time validation
            all_day = bool(data.get('all_day'))
            start_time_obj = None
            end_time_obj = None
            if not all_day:
                start_time_str = data.get('start_time')
                end_time_str = data.get('end_time')
                start_time_obj = _parse_iso_time(start_time_str)
                end_time_obj = _parse_iso_time(end_time_str)
                time_window_error = _validate_planner_time_window(
                    start_time_obj, end_time_obj)
                if time_window_error:
                    return jsonify({
                        'success': False,
                        'error': time_window_error
                    }), 400
                if start_date == end_date and start_time_obj >= end_time_obj:
                    return jsonify({
                        'success': False,
                        'error': 'end_time must be after start_time for single-day events'
                    }), 400

            # Parse total_hours
            total_hours_value = data.get('total_hours')
            total_hours = _normalize_planner_total_hours(total_hours_value)

            # Get event description
            event_description = (data.get('description') or '').strip()

            # Handle recurrence
            recurrence_value = (data.get('recurrence') or '').strip().lower()
            if recurrence_value in ('', 'none'):
                recurrence_value = None

            recurrence_end_date = _parse_iso_date(
                data.get('recurrence_end_date'))
            if recurrence_end_date and recurrence_end_date < start_date:
                return jsonify({
                    'success': False,
                    'error': 'recurrence_end_date cannot be earlier than start_date'
                }), 400

            # Get user IDs
            engineer_user_id = data.get('engineer_user_id') or current_user.id
            try:
                engineer_user_id = int(engineer_user_id)
            except (TypeError, ValueError):
                return jsonify({
                    'success': False,
                    'error': 'engineer_user_id must be a valid integer'
                }), 400
            created_by_user_id = current_user.id

            # Handle test_request lookup when tco_id might be None or empty
            test_request = None
            test_request_id = None
            if tco_id:
                test_request = EMCRequest.query.filter_by(
                    tco_id=tco_id).first()
                test_request_id = test_request.id if test_request else None

            candidate_snapshot = _build_planner_conflict_snapshot({
                'test_request_id': test_request_id,
                'test_person_name': test_person_name,
                'engineer_user_id': engineer_user_id,
                'test_name': test_name,
                'tco_id': tco_id if tco_id else None,
                'start_date': start_date,
                'end_date': end_date,
                'start_time': start_time_obj,
                'end_time': end_time_obj,
                'is_all_day': all_day,
                'event_type': event_type,
                'status': status
            })
            conflicts = _find_db_schedule_conflicts(candidate_snapshot)
            if _has_schedule_conflicts(conflicts):
                return jsonify({
                    'success': False,
                    'error': _format_schedule_conflict_message(
                        test_name,
                        test_person_name,
                        conflicts
                    )
                }), 409

            # Create planner entry
            planner_entry = PlannerEntry(
                test_person_name=test_person_name,
                engineer_user_id=engineer_user_id,
                created_by_user_id=created_by_user_id,
                test_name=test_name,
                tco_id=tco_id if tco_id else None,  # Explicitly set to None if empty
                test_request_id=test_request_id,
                start_date=start_date,
                end_date=end_date,
                start_time=start_time_obj,
                end_time=end_time_obj,
                total_hours=total_hours,
                event_description=event_description or None,
                event_type=event_type or None,
                recurrence=recurrence_value,
                recurrence_end_date=recurrence_end_date,
                is_all_day=all_day,
                status=status  # ADD THIS LINE
            )
            db.session.add(planner_entry)
            db.session.commit()

            logger.info(
                "Planner entry created for %s (%s -> %s) with event_type=%s, tco_id=%s",
                test_person_name, start_date, end_date, event_type, tco_id or 'None', status
            )

            return jsonify({
                'success': True,
                'data': planner_entry.to_dict()
            }), 201

        except Exception as exc:
            db.session.rollback()
            logger.error('Planner API error: %s', exc)
            return jsonify({
                'success': False,
                'error': 'Failed to process planner request'
            }), 500

    @flask_app.route('/api/planner/tco-ids', methods=['GET'])
    @login_required
    def get_planner_tco_ids():
        """Get unique TCO IDs from planner entries for dropdown."""
        try:
            # Role-based filtering for TCO IDs
            query = db.session.query(PlannerEntry.tco_id).distinct()

            if current_user.role == 'admin':
                # Admin can see all TCO IDs
                pass
            elif current_user.role == 'lab_engineer':
                # Lab engineers see only their own TCO IDs
                query = query.filter_by(engineer_user_id=current_user.id)
            else:
                # Regular users see only TCO IDs from their test requests
                user_tco_ids = db.session.query(EMCRequest.tco_id).filter_by(
                    user_id=current_user.id
                ).filter(EMCRequest.tco_id.isnot(None)).all()
                tco_id_list = [tco[0] for tco in user_tco_ids if tco[0]]
                if tco_id_list:
                    query = query.filter(PlannerEntry.tco_id.in_(tco_id_list))
                else:
                    # No TCOs found, return empty
                    query = query.filter(PlannerEntry.id == -1)

            tco_ids = [tco[0] for tco in query.filter(
                PlannerEntry.tco_id.isnot(None)
            ).order_by(PlannerEntry.tco_id.asc()).all() if tco[0]]

            return jsonify({
                'success': True,
                'data': tco_ids
            })
        except Exception as exc:
            logger.error('Error fetching TCO IDs: %s', exc)
            return jsonify({
                'success': False,
                'error': 'Error fetching TCO IDs'
            }), 500

    def send_cancellation_email(planner_entry, cancelled_by_user):
        """Send professional HTML email notification when an event is cancelled."""
        try:
            # Get recipient email addresses
            recipients = []

            # Add the engineer/test person email
            engineer = db.session.get(User, planner_entry.engineer_user_id)
            if engineer and engineer.email:
                recipients.append(engineer.email)

            # Add the creator's email if different from engineer
            creator = db.session.get(User, planner_entry.created_by_user_id)
            if creator and creator.email and creator.id != planner_entry.engineer_user_id:
                recipients.append(creator.email)

            # Optionally: Add admin emails or other stakeholders
            # admins = User.query.filter_by(role='admin').all()
            # for admin in admins:
            #     if admin.email:
            #         recipients.append(admin.email)

            if not recipients:
                logger.warning(
                    f'No recipients found for cancellation email (planner_id: {planner_entry.id})')
                return

            # Format dates and times
            start_date = planner_entry.start_date.strftime('%B %d, %Y')
            end_date = planner_entry.end_date.strftime('%B %d, %Y')

            start_time = planner_entry.start_time.strftime(
                '%I:%M %p') if planner_entry.start_time else 'All Day'
            end_time = planner_entry.end_time.strftime(
                '%I:%M %p') if planner_entry.end_time else ''

            # Get cancellation reason (escape HTML)
            from markupsafe import escape
            cancel_reason = escape(
                planner_entry.cancel_reason or 'No reason provided')

            # Create email subject
            subject = f'Event Cancelled: {planner_entry.test_name}'

            # Create professional HTML email body (matching submission notification format)
            html_body = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        color: #333333;
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background: linear-gradient(135deg, #dc2626 0%, #991b1b 100%);
                        color: white;
                        padding: 30px;
                        text-align: center;
                        border-radius: 8px 8px 0 0;
                    }}
                    .content {{
                        background: #ffffff;
                        padding: 30px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                    .info-box {{
                        background: #f9fafb;
                        border-left: 4px solid #dc2626;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .detail-row {{
                        display: flex;
                        padding: 10px 0;
                        border-bottom: 1px solid #e5e7eb;
                    }}
                    .detail-row:last-child {{
                        border-bottom: none;
                    }}
                    .detail-label {{
                        font-weight: bold;
                        width: 150px;
                        color: #6b7280;
                    }}
                    .detail-value {{
                        flex: 1;
                        color: #111827;
                    }}
                    .alert-box {{
                        background: #fee2e2;
                        border-left: 4px solid #dc2626;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .reason-box {{
                        background: #fffbeb;
                        border-left: 4px solid #f59e0b;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 4px;
                    }}
                    .footer {{
                        background: #f9fafb;
                        padding: 20px;
                        text-align: center;
                        color: #6b7280;
                        font-size: 12px;
                        border-radius: 0 0 8px 8px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1 style="margin: 0; font-size: 24px;">Event Cancelled</h1>
                    <p style="margin: 10px 0 0 0; opacity: 0.9;">A scheduled event has been cancelled</p>
                </div>
                
                <div class="content">
                    <p>Dear Team,</p>
                    
                    <div class="alert-box">
                        <p style="margin: 0; color: #991b1b; font-weight: bold;">
                            âš ï¸ The following event has been <strong>CANCELLED</strong> by {cancelled_by_user.username}
                        </p>
                    </div>
                    
                    <div class="info-box">
                        <div class="detail-row">
                            <div class="detail-label">Event Name:</div>
                            <div class="detail-value"><strong>{planner_entry.test_name}</strong></div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Engineer:</div>
                            <div class="detail-value">{planner_entry.test_person_name}</div>
                        </div>
                        {f'''<div class="detail-row">
                            <div class="detail-label">TCO ID:</div>
                            <div class="detail-value">{planner_entry.tco_id}</div>
                        </div>''' if planner_entry.tco_id else ''}
                        <div class="detail-row">
                            <div class="detail-label">Date:</div>
                            <div class="detail-value">{start_date}{f' - {end_date}' if start_date != end_date else ''}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Time:</div>
                            <div class="detail-value">{start_time}{f' - {end_time}' if end_time else ''}</div>
                        </div>
                        {f'''<div class="detail-row">
                            <div class="detail-label">Duration:</div>
                            <div class="detail-value">{planner_entry.total_hours} hours</div>
                        </div>''' if planner_entry.total_hours else ''}
                        <div class="detail-row">
                            <div class="detail-label">Cancelled By:</div>
                            <div class="detail-value">{cancelled_by_user.username}{f' ({cancelled_by_user.email})' if cancelled_by_user.email else ''}</div>
                        </div>
                        <div class="detail-row">
                            <div class="detail-label">Cancellation Time:</div>
                            <div class="detail-value">{get_ist_now().strftime('%B %d, %Y at %I:%M %p IST')}</div>
                        </div>
                    </div>

                    <div class="reason-box">
                        <h3 style="margin: 0 0 10px 0; color: #92400e; font-size: 16px;">Cancellation Reason:</h3>
                        <p style="margin: 0; color: #78350f; white-space: pre-wrap;">{cancel_reason}</p>
                    </div>

                    {f'''<h3 style="color: #111827; margin-top: 30px;">Event Description</h3>
                    <p style="background: #f9fafb; padding: 15px; border-radius: 6px; border-left: 4px solid #dc2626;">
                        {planner_entry.event_description}
                    </p>''' if planner_entry.event_description else ''}

                    <p style="margin-top: 30px;">
                        If you have any questions about this cancellation, please contact {cancelled_by_user.username}{f' at {cancelled_by_user.email}' if cancelled_by_user.email else ''}.
                    </p>
                </div>
                
                <div class="footer">
                    <p style="margin: 0;">This is an automated notification from the Test Planner System.</p>
                    <p style="margin: 5px 0 0 0;">Please do not reply to this email.</p>
                </div>
            </body>
            </html>
            """

            # Create plain text version
            text_body = f"""
    Event Cancelled

    Dear Team,

    âš ï¸ The following event has been CANCELLED by {cancelled_by_user.username}:

    Event Name: {planner_entry.test_name}
    Engineer: {planner_entry.test_person_name}
    {f'TCO ID: {planner_entry.tco_id}' if planner_entry.tco_id else ''}
    Date: {start_date}{f' - {end_date}' if start_date != end_date else ''}
    Time: {start_time}{f' - {end_time}' if end_time else ''}
    {f'Duration: {planner_entry.total_hours} hours' if planner_entry.total_hours else ''}
    Cancelled By: {cancelled_by_user.username}{f' ({cancelled_by_user.email})' if cancelled_by_user.email else ''}
    Cancellation Time: {get_ist_now().strftime('%B %d, %Y at %I:%M %p IST')}

    Cancellation Reason:
    {planner_entry.cancel_reason or 'No reason provided'}

    {f'Event Description:{chr(10)}{planner_entry.event_description}{chr(10)}' if planner_entry.event_description else ''}
    If you have any questions about this cancellation, please contact {cancelled_by_user.username}{f' at {cancelled_by_user.email}' if cancelled_by_user.email else ''}.

    ---
    This is an automated notification from the Test Planner System.
    Please do not reply to this email.
            """

            # Create message
            msg = MIMEMultipart('alternative')
            msg['Subject'] = f'Event Cancelled: {planner_entry.test_name} ({start_date})'
            msg['From'] = formataddr(
                (flask_app.config.get('SMTP_FROM_NAME', 'Test Planner'),
                 flask_app.config.get('SMTP_FROM_EMAIL', 'noreply-pcbacoe@thermofisher.com')))
            msg['To'] = ', '.join(recipients)

            # Attach both plain text and HTML versions
            part1 = MIMEText(text_body, 'plain', 'utf-8')
            part2 = MIMEText(html_body, 'html', 'utf-8')
            msg.attach(part1)
            msg.attach(part2)

            # Send the email using SMTP relay
            smtp_server = flask_app.config.get(
                'SMTP_SERVER', 'SMTPRELAY1.THERMOFISHER.COM')
            smtp_port = flask_app.config.get('SMTP_PORT', 25)

            logger.info(
                f'Connecting to SMTP server: {smtp_server}:{smtp_port}')

            with smtplib.SMTP(smtp_server, smtp_port, timeout=10) as server:
                # No authentication needed for internal relay
                server.sendmail(
                    flask_app.config.get(
                        'SMTP_FROM_EMAIL', 'noreply-pcbacoe@thermofisher.com'),
                    recipients,
                    msg.as_string())

            logger.info(
                f'Cancellation email sent for planner_id {planner_entry.id} to {", ".join(recipients)}')

        except Exception as exc:
            logger.error(f'Error sending cancellation email: {exc}')
            import traceback
            logger.error(traceback.format_exc())
            raise

    @flask_app.route('/api/planner/<int:planner_id>/status', methods=['PATCH'])
    @login_required
    def update_planner_status(planner_id):
        """
        Update the status of a planner entry (e.g., mark as cancelled, completed).

        Permission rules:
        - Admin: Can update any entry
        - Lab Engineer: Can only update their own entries
        - Regular User: Can only update entries for their TCOs

        Args:
            planner_id: ID of the planner entry to update

        JSON Body:
            status (required): New status (scheduled, in_progress, completed, cancelled)
            cancel_reason (required if status=cancelled): Reason for cancellation

        Returns:
            JSON response with success status and updated entry data
        """
        try:
            # Validate request data
            data = request.get_json()
            if not data:
                return jsonify({
                    'success': False,
                    'error': 'No data provided'
                }), 400

            new_status = data.get('status')
            cancel_reason = data.get('cancel_reason')

            if not new_status:
                return jsonify({
                    'success': False,
                    'error': 'Status is required'
                }), 400

            # Validate status value
            valid_statuses = ['scheduled',
                              'in_progress', 'completed', 'cancelled']
            if new_status not in valid_statuses:
                return jsonify({
                    'success': False,
                    'error': f'Invalid status. Must be one of: {", ".join(valid_statuses)}'
                }), 400

            # Get the planner entry
            planner_entry = db.session.get(PlannerEntry, planner_id)
            if not planner_entry:
                logger.warning(f'Planner entry {planner_id} not found')
                return jsonify({
                    'success': False,
                    'error': 'Planner entry not found'
                }), 404

            # ========= PERMISSION CHECKS =========
            if current_user.role == 'admin':
                # Admin can update any entry
                logger.info(
                    f'Admin {current_user.username} updating planner entry {planner_id}')

            elif current_user.role == 'lab_engineer':
                # Lab engineers can only update their own entries
                if planner_entry.engineer_user_id != current_user.id:
                    logger.warning(
                        f'Lab engineer {current_user.username} (ID: {current_user.id}) '
                        f'attempted to update planner entry {planner_id} '
                        f'owned by engineer ID: {planner_entry.engineer_user_id}'
                    )
                    return jsonify({
                        'success': False,
                        'error': 'You do not have permission to update this entry'
                    }), 403
                logger.info(
                    f'Lab engineer {current_user.username} updating their own planner entry {planner_id}')

            else:
                # Regular users can only update entries for their TCOs
                user_tco_ids = db.session.query(EMCRequest.tco_id).filter_by(
                    user_id=current_user.id
                ).filter(EMCRequest.tco_id.isnot(None)).all()
                tco_id_list = [tco[0] for tco in user_tco_ids if tco[0]]

                if not planner_entry.tco_id or planner_entry.tco_id not in tco_id_list:
                    logger.warning(
                        f'Regular user {current_user.username} (ID: {current_user.id}) '
                        f'attempted to update planner entry {planner_id} '
                        f'with TCO ID: {planner_entry.tco_id}. '
                        f'User\'s TCO IDs: {tco_id_list}'
                    )
                    return jsonify({
                        'success': False,
                        'error': 'You do not have permission to update this entry'
                    }), 403
                logger.info(
                    f'Regular user {current_user.username} updating planner entry {planner_id} for their TCO {planner_entry.tco_id}')
            # ========= END PERMISSION CHECKS =========

            # Prevent status changes on already cancelled entries (optional business rule)
            if planner_entry.status == 'cancelled' and new_status != 'cancelled':
                logger.warning(
                    f'Attempted to change status of cancelled planner entry {planner_id}')
                return jsonify({
                    'success': False,
                    'error': 'Cannot change status of a cancelled entry'
                }), 400

            # If cancelling, require a cancellation reason
            if new_status == 'cancelled':
                if not cancel_reason or not cancel_reason.strip():
                    return jsonify({
                        'success': False,
                        'error': 'Cancellation reason is required when cancelling an entry'
                    }), 400

                # Store cancellation details
                planner_entry.cancel_reason = cancel_reason.strip()
                planner_entry.cancelled_at = get_ist_now()
                planner_entry.cancelled_by = current_user.id

                logger.info(
                    f'Planner entry {planner_id} being cancelled by user {current_user.id} '
                    # Log first 50 chars
                    f'with reason: {cancel_reason[:50]}...'
                )

                # Send cancellation email notification
                try:
                    send_cancellation_email(planner_entry, current_user)
                    logger.info(
                        f'Cancellation email sent for planner entry {planner_id}')
                except Exception as email_error:
                    logger.error(
                        f'Failed to send cancellation email for planner entry {planner_id}: {email_error}'
                    )
                    import traceback
                    logger.error(traceback.format_exc())
                    # Don't fail the request if email fails - just log the error

            # Update the status
            old_status = planner_entry.status
            planner_entry.status = new_status
            planner_entry.updated_at = get_ist_now()

            # Commit changes to database
            db.session.commit()

            logger.info(
                f'Planner entry {planner_id} status updated from "{old_status}" to "{new_status}" '
                f'by user {current_user.username} (ID: {current_user.id})'
            )

            return jsonify({
                'success': True,
                'message': f'Planner entry status updated to {new_status}',
                'data': planner_entry.to_dict()
            })

        except Exception as exc:
            db.session.rollback()
            logger.error(
                f'Error updating planner status for entry {planner_id}: {exc}')
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'error': f'Failed to update planner entry status: {str(exc)}'
            }), 500

    @flask_app.route('/api/planner/test-persons', methods=['GET'])
    @login_required
    def get_planner_test_persons():
        """Get unique test person names from planner entries for dropdown."""
        try:
            # Build query based on user role
            if current_user.role == 'admin':
                # Admin can see all test persons
                query = db.session.query(PlannerEntry.test_person_name).filter(
                    PlannerEntry.test_person_name.isnot(None)
                ).distinct()
            elif current_user.role == 'lab_engineer':
                # Lab engineers see only their own test persons
                query = db.session.query(PlannerEntry.test_person_name).filter(
                    PlannerEntry.test_person_name.isnot(None),
                    PlannerEntry.engineer_user_id == current_user.id
                ).distinct()
            else:
                # Regular users see only test persons from their test requests
                user_tco_ids = db.session.query(EMCRequest.tco_id).filter_by(
                    user_id=current_user.id
                ).filter(EMCRequest.tco_id.isnot(None)).all()
                tco_id_list = [tco[0] for tco in user_tco_ids if tco[0]]
                if tco_id_list:
                    query = db.session.query(PlannerEntry.test_person_name).filter(
                        PlannerEntry.test_person_name.isnot(None),
                        PlannerEntry.tco_id.in_(tco_id_list)
                    ).distinct()
                else:
                    # No TCOs found, return empty
                    return jsonify({
                        'success': True,
                        'data': []
                    })

            # Order and get results
            test_person_names = query.order_by(
                PlannerEntry.test_person_name.asc()).all()
            result = [name[0] for name in test_person_names if name[0]]

            return jsonify({
                'success': True,
                'data': result
            })
        except Exception as exc:
            logger.error('Error fetching test person names: %s', exc)
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'error': f'Error fetching test person names: {str(exc)}'
            }), 500

    @flask_app.route('/api/admin/approve/<int:test_plan_id>', methods=['POST'])
    @login_required
    def admin_approve_test_plan(test_plan_id):
        """Admin approves a test plan."""
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Unauthorized access'
            }), 403

        try:
            request_obj = _get_request_or_404(test_plan_id)
            if not request_obj:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            request_obj.status = 'Test Plan Approved'
            request_obj.assigned_engineer_id = None
            request_obj.assigned_engineer_name = None
            request_obj.assignment_priority = None
            request_obj.assignment_due_date = None
            request_obj.assignment_notes = None
            request_obj.updated_at = get_ist_now()

            db.session.commit()

            logger.info(
                "Admin %s approved test request %s", current_user.username, test_plan_id)

            return jsonify({
                'success': True,
                'message': 'Test request approved successfully',
                'data': request_obj.to_dict()
            })
        except Exception as e:
            logger.error(f"Error approving test plan {test_plan_id}: {e}")
            return jsonify({
                'success': False,
                'error': 'Error approving test plan'
            }), 500

    @flask_app.route('/api/planner/<int:planner_id>', methods=['PUT', 'PATCH'])
    @login_required
    def update_planner_entry(planner_id):
        """Update an existing planner entry."""
        try:
            # Get the planner entry
            planner_entry = db.session.get(PlannerEntry, planner_id)

            if not planner_entry:
                return jsonify({
                    'success': False,
                    'error': 'Planner entry not found'
                }), 404

            # Check permissions
            if current_user.role == 'admin':
                # Admin can update any entry
                pass
            elif current_user.role == 'lab_engineer':
                # Lab engineers can only update their own entries
                if planner_entry.engineer_user_id != current_user.id:
                    return jsonify({
                        'success': False,
                        'error': 'You do not have permission to update this entry'
                    }), 403
            else:
                # Regular users can only update entries they created
                if planner_entry.created_by_user_id != current_user.id:
                    return jsonify({
                        'success': False,
                        'error': 'You do not have permission to update this entry'
                    }), 403

            # Get update data
            data = request.get_json() or {}
            target_entry = planner_entry

            # Update test_person_name
            test_person_name = (data.get('test_person_name') or '').strip()
            if test_person_name:
                target_entry.test_person_name = test_person_name

            engineer_user_id = data.get('engineer_user_id')
            if engineer_user_id not in (None, '', []):
                try:
                    target_entry.engineer_user_id = int(engineer_user_id)
                except (TypeError, ValueError):
                    return jsonify({
                        'success': False,
                        'error': 'engineer_user_id must be a valid integer'
                    }), 400

            # Update test_name
            test_name = (data.get('test_name') or '').strip()
            if test_name:
                target_entry.test_name = test_name

            # Get event_type to check if TCO ID is required
            event_type = (data.get('event_type')
                          or target_entry.event_type or 'test').strip().lower()
            target_entry.event_type = event_type

            # Update TCO ID (optional for selected non-test event types)
            tco_id = (data.get('tco_id') or '').strip()
            if tco_id or event_type in NON_TEST_EVENT_TYPES:
                target_entry.tco_id = tco_id if tco_id else None
                # Update test_request_id if tco_id changed
                if tco_id:
                    test_request = EMCRequest.query.filter_by(
                        tco_id=tco_id).first()
                    target_entry.test_request_id = test_request.id if test_request else None
                else:
                    target_entry.test_request_id = None

            # Update dates
            start_date_str = data.get('start_date')
            end_date_str = data.get('end_date')
            if start_date_str:
                start_date = _parse_iso_date(start_date_str)
                if start_date:
                    target_entry.start_date = start_date

            if end_date_str:
                end_date = _parse_iso_date(end_date_str)
                if end_date:
                    target_entry.end_date = end_date

            # Validate date range
            if target_entry.end_date < target_entry.start_date:
                return jsonify({
                    'success': False,
                    'error': 'end_date cannot be earlier than start_date'
                }), 400

            # Handle all-day events and time validation
            all_day = bool(data.get('all_day'))
            target_entry.is_all_day = all_day

            if not all_day:
                start_time_str = data.get('start_time')
                end_time_str = data.get('end_time')
                start_time_obj = _parse_iso_time(start_time_str)
                end_time_obj = _parse_iso_time(end_time_str)
                time_window_error = _validate_planner_time_window(
                    start_time_obj, end_time_obj)
                if time_window_error:
                    return jsonify({
                        'success': False,
                        'error': time_window_error
                    }), 400

                target_entry.start_time = start_time_obj
                target_entry.end_time = end_time_obj

                # Validate times for single-day events
                if target_entry.start_date == target_entry.end_date and start_time_obj >= end_time_obj:
                    return jsonify({
                        'success': False,
                        'error': 'end_time must be after start_time for single-day events'
                    }), 400
            else:
                target_entry.start_time = None
                target_entry.end_time = None

            # Update total_hours
            total_hours_value = data.get('total_hours')
            target_entry.total_hours = _normalize_planner_total_hours(
                total_hours_value)

            # Update event description
            event_description = data.get('description')
            if event_description is not None:
                target_entry.event_description = event_description.strip() or None

            # Update recurrence
            recurrence_value = (data.get('recurrence') or '').strip().lower()
            target_entry.recurrence = recurrence_value if recurrence_value not in (
                '', 'none') else None

            recurrence_end_date = _parse_iso_date(
                data.get('recurrence_end_date'))
            if recurrence_end_date:
                if recurrence_end_date < target_entry.start_date:
                    return jsonify({
                        'success': False,
                        'error': 'recurrence_end_date cannot be earlier than start_date'
                    }), 400
                target_entry.recurrence_end_date = recurrence_end_date
            else:
                target_entry.recurrence_end_date = None

            # Update status if provided (but don't change if not provided)
            if 'status' in data:
                status = data.get('status', 'in_progress').strip().lower()
                valid_statuses = ['in_progress',
                                  'completed', 'cancelled', 'scheduled']
                if status in valid_statuses:
                    target_entry.status = status

            candidate_snapshot = _build_planner_conflict_snapshot(target_entry)
            conflicts = _find_db_schedule_conflicts(
                candidate_snapshot,
                exclude_entry_id=target_entry.id
            )
            if _has_schedule_conflicts(conflicts):
                return jsonify({
                    'success': False,
                    'error': _format_schedule_conflict_message(
                        target_entry.test_name,
                        target_entry.test_person_name,
                        conflicts
                    )
                }), 409

            # Update timestamps
            target_entry.updated_at = get_ist_now()

            # Commit changes
            db.session.commit()

            logger.info(
                "Planner entry %d updated by %s (%s -> %s)",
                planner_id,
                current_user.username,
                target_entry.start_date,
                target_entry.end_date
            )

            return jsonify({
                'success': True,
                'data': target_entry.to_dict(),
                'message': 'Planner entry updated successfully'
            })

        except Exception as exc:
            db.session.rollback()
            logger.error('Error updating planner entry %d: %s',
                         planner_id, exc)
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'error': f'Failed to update planner entry: {str(exc)}'
            }), 500
    @flask_app.route('/api/planner/<int:planner_id>', methods=['DELETE'])
    @login_required
    def delete_planner_entry(planner_id):
        """Permanently delete a planner entry. Admin only."""
        try:
            # Role check
            if current_user.role != 'admin':
                return jsonify({
                    'success': False,
                    'error': 'Permission denied. Only admins can delete planner entries.'
                }), 403

            # Fetch the entry
            entry = db.session.get(PlannerEntry, planner_id)
            if not entry:
                return jsonify({
                    'success': False,
                    'error': f'Planner entry {planner_id} not found.'
                }), 404

            entry_title = entry.test_name or str(planner_id)

            # Delete and commit
            db.session.delete(entry)
            db.session.commit()

            return jsonify({
                'success': True,
                'message': f'Planner entry "{entry_title}" deleted successfully.',
                'deleted_id': planner_id
            }), 200

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error deleting planner entry {planner_id}: {e}')
            return jsonify({
                'success': False,
                'error': 'An internal error occurred while deleting the entry.'
            }), 500
    @flask_app.route('/api/admin/assign/<int:test_plan_id>', methods=['POST'])
    @login_required
    def admin_assign_test_plan(test_plan_id):
        """Admin assigns a test plan to a lab engineer or admin user."""
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Unauthorized access'
            }), 403

        try:
            data = request.get_json()
            if not data:
                logger.error("No JSON data received in assign request")
                return jsonify({
                    'success': False,
                    'error': 'No data provided'
                }), 400

            logger.debug(f"Assign request data: {data}")
            assigned_engineer_id = data.get('assigned_engineer_id')
            priority = data.get('priority', 'normal')
            due_date = data.get('due_date')
            assignment_notes = data.get('assignment_notes', '')

            if not assigned_engineer_id:
                logger.error("No assigned_engineer_id provided")
                return jsonify({
                    'success': False,
                    'error': 'Assignee is required'
                }), 400

            # Verify the selected assignee exists and has an allowed role.
            assignee = User.query.filter(
                User.id == assigned_engineer_id,
                User.role.in_(['lab_engineer', 'admin']),
                User.is_active.is_(True)
            ).first()
            if not assignee:
                logger.error(
                    f"Assignee with id {assigned_engineer_id} not found or role is not allowed")
                return jsonify({
                    'success': False,
                    'error': 'Invalid assignee selected'
                }), 400

            request_obj = _get_request_or_404(test_plan_id)
            if not request_obj:
                logger.error(f"Test request {test_plan_id} not found")
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            logger.debug(
                f"Updating test request {test_plan_id} with assignment data")
            request_obj.status = 'Assigned Lab Engineer'
            request_obj.assigned_engineer_id = assignee.id
            request_obj.assigned_engineer_name = assignee.username
            request_obj.assignment_priority = priority
            request_obj.assignment_due_date = _parse_iso_date(due_date)
            request_obj.assignment_notes = assignment_notes
            request_obj.updated_at = get_ist_now()

            try:
                db.session.commit()
                logger.info(
                    f"Successfully assigned test request {test_plan_id} to {assignee.username}"
                )
            except Exception:
                db.session.rollback()
                raise

            logger.info(
                "Admin %s assigned test request %s to %s",
                current_user.username,
                test_plan_id,
                assignee.username
            )

            # Send email notification to engineer with admin in CC
            try:
                admin_email = current_user.email if current_user.email else ''
                engineer_email = assignee.email if assignee.email else ''
                if engineer_email:
                    send_assignment_notification(
                        engineer_email=engineer_email,
                        admin_email=admin_email,
                        test_request=request_obj,
                        engineer_name=assignee.username,
                        priority=priority,
                        due_date=due_date or '',
                        notes=assignment_notes
                    )
            except Exception as e:
                logger.error(
                    'Failed to send assignment notification email: %s', e)
                # Don't fail the assignment if email fails

            # Safely convert to dict, handling any serialization errors
            try:
                response_data = request_obj.to_dict()
            except Exception as dict_error:
                logger.warning(
                    'Error serializing test request to dict: %s', dict_error)
                # Return minimal data if to_dict() fails
                response_data = {
                    'id': request_obj.id,
                    'tco_id': request_obj.tco_id,
                    'status': request_obj.status,
                    'assigned_engineer_id': request_obj.assigned_engineer_id,
                    'assigned_engineer_name': request_obj.assigned_engineer_name,
                    'assignment_priority': request_obj.assignment_priority,
                    'assignment_due_date': request_obj.assignment_due_date.isoformat() if request_obj.assignment_due_date else None,
                    'assignment_notes': request_obj.assignment_notes
                }

            return jsonify({
                'success': True,
                'message': 'Test request assigned successfully',
                'data': response_data
            })
        except Exception as e:
            db.session.rollback()
            import traceback
            error_trace = traceback.format_exc()
            logger.error(
                f"Error assigning test plan {test_plan_id}: {e}\n{error_trace}")
            return jsonify({
                'success': False,
                'error': f'Error assigning test plan: {str(e)}'
            }), 500

    @flask_app.route('/api/admin/reject/<int:test_plan_id>', methods=['POST'])
    @login_required
    def admin_reject_test_plan(test_plan_id):
        """Admin rejects a test plan."""
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Unauthorized access'
            }), 403

        try:
            data = request.get_json()
            rejection_reason = data.get('rejection_reason', '').strip()

            if not rejection_reason:
                return jsonify({
                    'success': False,
                    'error': 'Rejection reason is required'
                }), 400

            # Find the test request
            test_request = _get_request_or_404(test_plan_id)
            if not test_request:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            rejected_at_value = get_ist_now()
            updated_at_value = get_ist_now()

            try:
                test_request.status = 'Rejected'
                test_request.updated_at = updated_at_value

                setattr(test_request, 'rejection_reason', rejection_reason)
                setattr(test_request, 'rejected_by', current_user.username)
                setattr(test_request, 'rejected_at', rejected_at_value)

                db.session.commit()
                logger.info(
                    "Admin %s rejected test plan %s: %s",
                    current_user.username,
                    test_plan_id,
                    rejection_reason
                )
            except Exception as db_error:
                db.session.rollback()
                raise db_error

            # Send rejection notification email to the requester
            try:
                send_rejection_notification(
                    test_request=test_request,
                    rejection_reason=rejection_reason,
                    rejected_by=current_user.username
                )
            except Exception as email_error:
                logger.error(
                    f"Failed to send rejection notification email: {email_error}")
                # Don't fail the rejection if email fails

            # Refresh to get the latest data from database
            db.session.refresh(test_request)

            # Get rejection fields using getattr in case they're not in model definition
            rejection_reason_val = getattr(
                test_request, 'rejection_reason', None)
            rejected_by_val = getattr(test_request, 'rejected_by', None)
            rejected_at_val = getattr(test_request, 'rejected_at', None)

            return jsonify({
                'success': True,
                'message': 'Test plan rejected successfully',
                'data': {
                    'id': test_request.id,
                    'status': test_request.status,
                    'rejection_reason': rejection_reason_val,
                    'rejected_by': rejected_by_val,
                    'rejected_at': rejected_at_val.isoformat() if rejected_at_val else None
                }
            })
        except Exception as e:
            db.session.rollback()
            import traceback
            error_trace = traceback.format_exc()
            logger.error(
                f"Error rejecting test plan {test_plan_id}: {e}\n{error_trace}")
            return jsonify({
                'success': False,
                'error': f'Error rejecting test plan: {str(e)}'
            }), 500

    @flask_app.route('/api/admin/test-requests/<int:request_id>/final-approval', methods=['POST'])
    @login_required
    def admin_final_approve_test_plan(request_id: int):
        """Finalize a reviewer-proposed schedule, notify engineers, and mark request as approved."""
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Unauthorized access'
            }), 403

        try:
            test_request = _get_request_or_404(request_id)
            if not test_request:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            payload = request.get_json() or {}

            def _clean_optional_text(value, max_length=None):
                if value is None:
                    return None
                cleaned = str(value).strip()
                if not cleaned:
                    return None
                if max_length:
                    return cleaned[:max_length]
                return cleaned

            lab_manager_name = _clean_optional_text(
                payload.get('lab_manager_name'), 200) or _clean_optional_text(current_user.username, 200)
            if not lab_manager_name:
                return jsonify({
                    'success': False,
                    'error': 'Lab manager name is required for final approval'
                }), 400

            lab_manager_signature = _clean_optional_text(
                payload.get('lab_manager_signature')) or lab_manager_name
            if not lab_manager_signature:
                return jsonify({
                    'success': False,
                    'error': 'Lab manager signature is required for final approval'
                }), 400

            if not test_request.test_assignments:
                return jsonify({
                    'success': False,
                    'error': 'No test assignments available to approve'
                }), 400

            try:
                assignments: list[dict] = json.loads(
                    test_request.test_assignments)
            except (ValueError, TypeError):
                logger.exception(
                    'Failed to parse test_assignments for request %s', request_id)
                return jsonify({
                    'success': False,
                    'error': 'Stored test assignments are invalid and cannot be processed'
                }), 500

            if not assignments:
                return jsonify({
                    'success': False,
                    'error': 'No test assignments available to approve'
                }), 400

            engineer_cache: dict[int, User | None] = {}
            notification_payloads: list[dict] = []

            for assignment in assignments:
                engineer_id = assignment.get('engineer_id')
                if not engineer_id:
                    continue

                if engineer_id not in engineer_cache:
                    engineer_cache[engineer_id] = User.query.filter(
                        User.id == engineer_id,
                        User.role.in_(['lab_engineer', 'admin']),
                        User.is_active.is_(True)
                    ).first()

                engineer = engineer_cache[engineer_id]
                if not engineer or not engineer.email:
                    logger.warning(
                        'Skipping final approval email for assignment %s due to missing engineer/email',
                        assignment.get('test_name')
                    )
                    continue

                notification_payloads.append({
                    'engineer': engineer,
                    'assignment': assignment
                })

            if not notification_payloads:
                return jsonify({
                    'success': False,
                    'error': 'No valid engineer assignments found to notify'
                }), 400

            signed_at = get_ist_now()
            test_request.status = 'Test Plan Approved'
            test_request.review_comments = None
            test_request.reviewed_by = None
            test_request.reviewed_at = None
            test_request.lab_manager_name = lab_manager_name
            test_request.lab_manager_signature = lab_manager_signature
            test_request.lab_manager_date = signed_at.date()
            test_request.lab_manager_signed_at = signed_at
            test_request.updated_at = get_ist_now()
            db.session.commit()

            try:
                admin_emails = [
                    admin.email for admin in User.query.filter_by(
                        role='admin').all() if admin.email
                ]
                send_test_assignment_emails(
                    test_request=test_request,
                    assignment_payloads=notification_payloads,
                    admin_emails=admin_emails,
                    status_label='Test Plan Approved',
                    send_to_admins_first=False
                )
            except Exception as email_exc:
                logger.error(
                    'Final approval emails failed for request %s: %s',
                    request_id, email_exc
                )

            return jsonify({
                'success': True,
                'message': 'Test plan approved and notifications sent',
                'data': test_request.to_dict()
            })
        except Exception as exc:
            db.session.rollback()
            logger.error('Error during final approval of %s: %s',
                         request_id, exc)
            return jsonify({
                'success': False,
                'error': 'Error approving test plan schedule'
            }), 500

    @flask_app.route('/api/admin/test-requests/<int:request_id>/request-plan-update', methods=['POST'])
    @login_required
    def admin_request_plan_update(request_id: int):
        """Send a plan change request back to the reviewer with comments."""
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Unauthorized access'
            }), 403

        try:
            payload = request.get_json() or {}
            comments = (payload.get('comments') or '').strip()
            if not comments:
                return jsonify({
                    'success': False,
                    'error': 'Comments are required to request a plan update'
                }), 400

            test_request = _get_request_or_404(request_id)
            if not test_request:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404
            if not test_request.test_assignments:
                return jsonify({
                    'success': False,
                    'error': 'No assigned tests found to update'
                }), 400

            history_entries = []
            if test_request.plan_update_history:
                try:
                    history_entries = json.loads(
                        test_request.plan_update_history)
                    if not isinstance(history_entries, list):
                        history_entries = []
                except (TypeError, ValueError):
                    history_entries = []

            history_entries.append({
                'requested_by': current_user.username,
                'requested_at': get_ist_now().isoformat(),
                'comments': comments
            })

            test_request.plan_update_history = json.dumps(history_entries)
            test_request.status = 'Update plan'
            _append_review_comment_entry(
                test_request=test_request,
                comment=comments,
                username=current_user.username,
                role=current_user.role
            )
            test_request.updated_at = get_ist_now()
            db.session.commit()

            try:
                send_plan_update_notification(
                    test_request=test_request,
                    comments=comments,
                    requested_by=current_user.username
                )
            except Exception as email_error:
                logger.error(
                    'Plan update notification failed: %s', email_error)

            return jsonify({
                'success': True,
                'message': 'Plan update requested successfully'
            })
        except Exception as exc:
            db.session.rollback()
            logger.error('Error requesting plan update for %s: %s',
                         request_id, exc)
            return jsonify({
                'success': False,
                'error': 'Error requesting plan update'
            }), 500

    @flask_app.route('/api/admin/test-requests/<int:request_id>', methods=['PATCH'])
    @login_required
    def admin_update_test_request(request_id):
        """Admin updates status and assignment details for a test request."""
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Unauthorized access'
            }), 403

        try:
            payload = request.get_json() or {}
            request_obj = _get_request_or_404(request_id)
            if not request_obj:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            allowed_statuses = {
                'At Review',
                'Approved',
                'Assigned Lab Engineer',
                'Completed',
                'Cancelled',
                'Rejected',
                'Update plan'
            }

            new_status = payload.get('status') or request_obj.status
            if new_status not in allowed_statuses:
                return jsonify({
                    'success': False,
                    'error': 'Invalid status provided'
                }), 400

            if new_status == 'Cancelled':
                cancel_reason = (payload.get('cancel_reason') or '').strip()
                if not cancel_reason:
                    return jsonify({
                        'success': False,
                        'error': 'Cancellation reason is required'
                    }), 400

                cancel_counts = _cancel_tco_request(
                    test_request=request_obj,
                    cancel_reason=cancel_reason,
                    cancelled_by_user=current_user
                )
                db.session.commit()

                logger.info(
                    "Admin %s cancelled TCO %s with %d planner entries and %d request tests updated",
                    current_user.username,
                    request_id,
                    cancel_counts['cancelled_planner_entries'],
                    cancel_counts['cancelled_request_tests']
                )

                return jsonify({
                    'success': True,
                    'message': 'TCO cancelled successfully',
                    'data': request_obj.to_dict(),
                    'cancelled_planner_entries': cancel_counts['cancelled_planner_entries'],
                    'cancelled_request_tests': cancel_counts['cancelled_request_tests']
                })

            assigned_engineer_id = payload.get('assigned_engineer_id')
            assignment_priority = payload.get('assignment_priority')
            due_date_value = payload.get('assignment_due_date') or payload.get(
                'due_date')
            assignment_notes = payload.get('assignment_notes')

            engineer = None
            if assigned_engineer_id:
                engineer = User.query.filter(
                    User.id == assigned_engineer_id,
                    User.role.in_(['lab_engineer', 'admin']),
                    User.is_active.is_(True)
                ).first()
                if not engineer:
                    return jsonify({
                        'success': False,
                        'error': 'Invalid assignee selected'
                    }), 400

            if new_status == 'Assigned Lab Engineer' and not engineer:
                return jsonify({
                    'success': False,
                    'error': 'Assignee is required for Assigned Lab Engineer status'
                }), 400

            request_obj.status = new_status

            if engineer:
                request_obj.assigned_engineer_id = engineer.id
                request_obj.assigned_engineer_name = engineer.username
            elif new_status != 'Assigned Lab Engineer':
                request_obj.assigned_engineer_id = None
                request_obj.assigned_engineer_name = None

            if assignment_priority is not None:
                request_obj.assignment_priority = assignment_priority or None

            if due_date_value is not None:
                request_obj.assignment_due_date = _parse_iso_date(
                    due_date_value)

            if assignment_notes is not None:
                request_obj.assignment_notes = assignment_notes or None

            request_obj.updated_at = get_ist_now()
            db.session.commit()

            logger.info(
                "Admin %s updated test request %s", current_user.username, request_id
            )

            return jsonify({
                'success': True,
                'data': request_obj.to_dict()
            })
        except Exception as e:
            logger.error("Error updating test request %s: %s", request_id, e)
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': 'Error updating test request'
            }), 500

    @flask_app.route('/api/admin/test-requests/<int:request_id>/reassign-owner', methods=['POST'])
    @login_required
    def admin_reassign_tco_owner(request_id):
        """Reassign only the parent TCO owner and leave test-level assignments untouched."""
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Unauthorized access'
            }), 403

        try:
            payload = request.get_json() or {}
            assigned_engineer_id = payload.get('assigned_engineer_id')
            if not assigned_engineer_id:
                return jsonify({
                    'success': False,
                    'error': 'Assignee is required'
                }), 400

            assignee = User.query.filter(
                User.id == assigned_engineer_id,
                User.role.in_(['lab_engineer', 'admin']),
                User.is_active.is_(True)
            ).first()
            if not assignee:
                return jsonify({
                    'success': False,
                    'error': 'Invalid assignee selected'
                }), 400

            request_obj = _get_request_or_404(request_id)
            if not request_obj:
                return jsonify({
                    'success': False,
                    'error': 'Test request not found'
                }), 404

            previous_assignee = request_obj.assigned_engineer_name or 'Unassigned'
            request_obj.assigned_engineer_id = assignee.id
            request_obj.assigned_engineer_name = assignee.username
            request_obj.updated_at = get_ist_now()
            db.session.commit()

            logger.info(
                'Admin %s reassigned TCO owner for request %s from %s to %s',
                current_user.username,
                request_id,
                previous_assignee,
                assignee.username
            )

            return jsonify({
                'success': True,
                'message': 'TCO owner reassigned successfully',
                'data': {
                    'id': request_obj.id,
                    'status': request_obj.status,
                    'assigned_engineer_id': request_obj.assigned_engineer_id,
                    'assigned_engineer_name': request_obj.assigned_engineer_name,
                    'assignment_priority': request_obj.assignment_priority,
                    'assignment_due_date': request_obj.assignment_due_date.isoformat() if request_obj.assignment_due_date else None,
                    'assignment_notes': request_obj.assignment_notes
                }
            })
        except Exception as e:
            logger.error("Error reassigning TCO owner %s: %s", request_id, e)
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': 'Error reassigning TCO owner'
            }), 500

    @flask_app.route('/planner')
    @login_required
    def planner():
        """Planner page for managing test schedules and planning."""
        try:
            # Get list of lab engineers for admin filter dropdown
            lab_engineers = []
            if current_user.role == 'admin':
                lab_engineers = User.query.filter_by(
                    role='lab_engineer').order_by(User.username).all()

            return render_template('planner.html',
                                   user_role=current_user.role,
                                   lab_engineers=lab_engineers)
        except Exception as e:
            logger.error("Error loading planner page: %s", e)
            flash('Error loading planner page. Please try again.', 'error')
            return redirect(url_for('index'))

    @flask_app.route('/dashboard')
    @login_required
    def dashboard():
        """Dashboard page for viewing analytics and overview (Admin only)."""
        try:
            # Check if user is admin
            if current_user.role != 'admin':
                logger.warning(
                    f'Non-admin user {current_user.username} (role: {current_user.role}) '
                    f'attempted to access admin dashboard'
                )
                flash(
                    'Access denied. This page is only available to administrators.', 'error')
                return redirect(url_for('index'))

            # Admin access granted
            logger.info(
                f'Admin user {current_user.username} accessed dashboard')

            now = get_ist_now()
            today = now.date()
            last_30_days = now - timedelta(days=30)
            previous_30_days = now - timedelta(days=60)

            requests = EMCRequest.query.options(
                joinedload(EMCRequest.tests)
            ).order_by(
                EMCRequest.updated_at.desc(),
                EMCRequest.created_at.desc()
            ).all()
            planner_entries = PlannerEntry.query.order_by(
                PlannerEntry.updated_at.desc(),
                PlannerEntry.created_at.desc()
            ).all()

            total_requests = len(requests)

            active_request_statuses = {
                'at review',
                'approved',
                'assigned',
                'assigned lab engineer',
                'test plan approved',
                'test schedule in progress',
                'test plan to approve',
                'in progress',
                'update plan',
                'need more information',
                'draft report',
                'proceed report',
                'datasheet uploaded',
                'report uploaded',
                'report_uploaded',
                'peer review',
                'admin sign off',
            }
            completed_request_statuses = {'completed'}
            rejected_request_statuses = {'rejected', 'cancelled'}
            terminal_planner_statuses = {'completed', 'cancelled'}

            def _normalize_status(value):
                return str(value or '').strip().lower()

            def _request_event_dt(request_obj):
                return (
                    request_obj.updated_at or
                    request_obj.submitted_at or
                    request_obj.created_at
                )

            def _status_badge_for_due_date(due_date_value):
                if not due_date_value:
                    return {
                        'text': 'No due date',
                        'class_name': 'bg-gray-100 text-gray-600'
                    }

                days_until_due = (due_date_value - today).days
                if days_until_due <= 0:
                    return {
                        'text': 'Due Today',
                        'class_name': 'bg-red-100 text-red-600'
                    }
                if days_until_due == 1:
                    return {
                        'text': 'Due Tomorrow',
                        'class_name': 'bg-yellow-100 text-yellow-600'
                    }
                return {
                    'text': f'Due in {days_until_due} days',
                    'class_name': 'bg-green-100 text-green-600'
                }

            def _shift_month(base_year, base_month, delta_months):
                month_index = (base_year * 12 + (base_month - 1)) + delta_months
                target_year = month_index // 12
                target_month = (month_index % 12) + 1
                return target_year, target_month

            def _initials(name_value):
                words = [word for word in str(name_value or '').split() if word]
                if not words:
                    return 'NA'
                return ''.join(word[0].upper() for word in words[:2])

            def _coerce_dashboard_datetime(value):
                if not value:
                    return None
                if isinstance(value, date) and not isinstance(value, datetime):
                    return datetime.combine(value, time.min, tzinfo=IST)
                if getattr(value, 'tzinfo', None) is None:
                    return value.replace(tzinfo=IST)
                return value.astimezone(IST)

            def _relative_time(dt_obj):
                dt_obj = _coerce_dashboard_datetime(dt_obj)
                if not dt_obj:
                    return 'No timestamp'
                delta = now - dt_obj

                seconds = max(int(delta.total_seconds()), 0)
                if seconds < 60:
                    return 'Just now'
                if seconds < 3600:
                    minutes = seconds // 60
                    return f'{minutes} min ago'
                if seconds < 86400:
                    hours = seconds // 3600
                    return f'{hours} hr ago'
                if seconds < 604800:
                    days = seconds // 86400
                    return f'{days} day ago'
                return dt_obj.strftime('%d %b %Y')

            def _count_in_window(dt_obj, start_dt, end_dt):
                dt_obj = _coerce_dashboard_datetime(dt_obj)
                start_dt = _coerce_dashboard_datetime(start_dt)
                end_dt = _coerce_dashboard_datetime(end_dt)
                if not dt_obj or not start_dt or not end_dt:
                    return False
                return start_dt <= dt_obj < end_dt

            active_requests = sum(
                1 for request_obj in requests
                if _normalize_status(request_obj.status) in active_request_statuses
            )
            pending_approval_requests = sum(
                1 for request_obj in requests
                if _normalize_status(request_obj.status) == 'at review'
            )
            completed_requests = sum(
                1 for request_obj in requests
                if _normalize_status(request_obj.status) in completed_request_statuses
            )
            rejected_requests = sum(
                1 for request_obj in requests
                if _normalize_status(request_obj.status) in rejected_request_statuses
            )
            report_ready_statuses = {
                'draft report',
                'proceed report',
                'datasheet uploaded',
                'report uploaded',
                'report_uploaded',
                'admin sign off',
            }
            ready_for_report_requests = sum(
                1 for request_obj in requests
                if _normalize_status(request_obj.status) in report_ready_statuses
            )

            active_tests = sum(
                1 for entry in planner_entries
                if _normalize_status(entry.status) not in terminal_planner_statuses
            )

            requests_last_30 = sum(
                1 for request_obj in requests
                if _count_in_window(
                    request_obj.submitted_at or request_obj.created_at,
                    last_30_days,
                    now
                )
            )
            requests_previous_30 = sum(
                1 for request_obj in requests
                if _count_in_window(
                    request_obj.submitted_at or request_obj.created_at,
                    previous_30_days,
                    last_30_days
                )
            )
            completed_last_30 = sum(
                1 for request_obj in requests
                if _normalize_status(request_obj.status) in completed_request_statuses and
                _count_in_window(request_obj.updated_at or request_obj.created_at, last_30_days, now)
            )

            completion_rate = round(
                (completed_requests / total_requests) * 100, 1
            ) if total_requests else 0.0

            equipment_total = Equipment.query.count()
            equipment_needs_calibration = Equipment.query.filter(
                Equipment.calibration_due_date < today
            ).count()
            equipment_in_calibration = Equipment.query.filter(
                Equipment.calibration_status_col == 'In Calibration'
            ).count()

            active_lab_engineers = User.query.filter(
                User.role == 'lab_engineer',
                User.is_active.is_(True)
            ).count()

            selected_test_counter = Counter()
            for request_obj in requests:
                for test_row in getattr(request_obj, 'tests', []):
                    if getattr(test_row, 'is_selected', False):
                        label = EMCRequest._legacy_test_key(test_row.test_code)
                        if label:
                            selected_test_counter[label] += 1

            total_selected_tests = sum(selected_test_counter.values())
            distribution_colors = [
                'bg-blue-500',
                'bg-green-500',
                'bg-amber-500',
                'bg-fuchsia-500',
            ]
            test_distribution = []
            for index, (label, count) in enumerate(selected_test_counter.most_common(4)):
                percent = round((count / total_selected_tests) * 100) if total_selected_tests else 0
                test_distribution.append({
                    'label': label,
                    'count': count,
                    'percent': percent,
                    'bar_class': distribution_colors[index % len(distribution_colors)]
                })

            total_hours_values = [
                float(entry.total_hours) for entry in planner_entries
                if entry.total_hours not in (None, '')
                and _normalize_status(entry.status) != 'cancelled'
            ]
            avg_hours_per_schedule = round(
                sum(total_hours_values) / len(total_hours_values), 1
            ) if total_hours_values else 0.0

            engineers_scheduled_last_30 = len({
                entry.engineer_user_id for entry in planner_entries
                if entry.engineer_user_id
                and _normalize_status(entry.status) != 'cancelled'
                and _count_in_window(entry.updated_at or entry.created_at, last_30_days, now)
            })

            recent_activity = []
            activity_color_map = {
                'completed': 'bg-emerald-500',
                'rejected': 'bg-rose-500',
                'at review': 'bg-blue-500',
                'assigned lab engineer': 'bg-amber-500',
                'test plan approved': 'bg-violet-500',
                'admin sign off': 'bg-fuchsia-500',
            }
            for request_obj in requests[:5]:
                status_text = str(request_obj.status or 'Unknown').strip() or 'Unknown'
                status_key = _normalize_status(status_text)
                event_dt = _request_event_dt(request_obj)
                recent_activity.append({
                    'title': request_obj.tco_id or f'REQ-{request_obj.id}',
                    'description': (
                        f"{request_obj.product_name or 'Unnamed Product'}: {status_text}"
                    ),
                    'time_ago': _relative_time(event_dt),
                    'dot_class': activity_color_map.get(status_key, 'bg-slate-400')
                })

            upcoming_deadlines = []
            for request_obj in requests:
                status_key = _normalize_status(request_obj.status)
                if (
                    status_key in completed_request_statuses or
                    status_key in rejected_request_statuses
                ):
                    continue

                due_date_value = request_obj.assignment_due_date or request_obj.test_completion_date
                if not due_date_value or due_date_value < today:
                    continue

                badge = _status_badge_for_due_date(due_date_value)
                upcoming_deadlines.append({
                    'title': request_obj.tco_id or f'REQ-{request_obj.id}',
                    'subtitle': request_obj.product_name or status_key.title(),
                    'badge_text': badge['text'],
                    'badge_class': badge['class_name'],
                    'due_date': due_date_value
                })

            if len(upcoming_deadlines) < 3:
                existing_deadline_keys = {
                    (item['title'], item['due_date']) for item in upcoming_deadlines
                }
                for entry in planner_entries:
                    if _normalize_status(entry.status) in terminal_planner_statuses:
                        continue
                    if not entry.end_date or entry.end_date < today:
                        continue

                    title = entry.tco_id or entry.test_name or f'Planner-{entry.id}'
                    candidate_key = (title, entry.end_date)
                    if candidate_key in existing_deadline_keys:
                        continue

                    badge = _status_badge_for_due_date(entry.end_date)
                    upcoming_deadlines.append({
                        'title': title,
                        'subtitle': entry.test_name or (entry.test_person_name or 'Scheduled task'),
                        'badge_text': badge['text'],
                        'badge_class': badge['class_name'],
                        'due_date': entry.end_date
                    })
                    existing_deadline_keys.add(candidate_key)
                    if len(upcoming_deadlines) >= 3:
                        break

            upcoming_deadlines.sort(key=lambda item: item['due_date'])
            upcoming_deadlines = upcoming_deadlines[:3]

            utilization_by_engineer = Counter()
            for entry in planner_entries:
                if _normalize_status(entry.status) in terminal_planner_statuses:
                    continue
                engineer_name = str(entry.test_person_name or '').strip()
                if not engineer_name:
                    continue
                utilization_by_engineer[engineer_name] += float(entry.total_hours or 1)

            max_utilization = max(utilization_by_engineer.values(), default=0)
            utilization_colors = ['bg-blue-500', 'bg-green-500', 'bg-yellow-500']
            lab_utilization = []
            for index, (engineer_name, total_hours) in enumerate(utilization_by_engineer.most_common(3)):
                percent = round((total_hours / max_utilization) * 100) if max_utilization else 0
                lab_utilization.append({
                    'label': engineer_name,
                    'percent': percent,
                    'hours_text': f'{total_hours:.1f} hrs',
                    'bar_class': utilization_colors[index % len(utilization_colors)]
                })

            monthly_trends = []
            for delta_months in range(-5, 1):
                year_value, month_value = _shift_month(today.year, today.month, delta_months)
                month_start = datetime(year_value, month_value, 1, tzinfo=IST)
                next_year, next_month = _shift_month(year_value, month_value, 1)
                next_month_start = datetime(next_year, next_month, 1, tzinfo=IST)
                request_count = sum(
                    1 for request_obj in requests
                    if _count_in_window(
                        request_obj.submitted_at or request_obj.created_at,
                        month_start,
                        next_month_start
                    )
                )
                monthly_trends.append({
                    'label': month_start.strftime('%b'),
                    'count': request_count
                })

            max_monthly_count = max(
                (item['count'] for item in monthly_trends),
                default=0
            )
            for item in monthly_trends:
                item['height_pct'] = round((item['count'] / max_monthly_count) * 100) if max_monthly_count else 0

            current_period_team = Counter()
            previous_period_team = Counter()
            team_name_lookup = {}
            for entry in planner_entries:
                if _normalize_status(entry.status) == 'cancelled':
                    continue
                engineer_name = str(entry.test_person_name or '').strip()
                if not engineer_name:
                    continue

                team_name_lookup[engineer_name] = engineer_name
                timestamp_value = entry.updated_at or entry.created_at
                if _count_in_window(timestamp_value, last_30_days, now):
                    current_period_team[engineer_name] += 1
                elif _count_in_window(timestamp_value, previous_30_days, last_30_days):
                    previous_period_team[engineer_name] += 1

            team_performance = []
            avatar_colors = ['bg-blue-500', 'bg-green-500', 'bg-purple-500']
            for index, (engineer_name, current_count) in enumerate(current_period_team.most_common(3)):
                previous_count = previous_period_team.get(engineer_name, 0)
                delta_count = current_count - previous_count
                if delta_count > 0:
                    change_text = f'+{delta_count}'
                    change_class = 'text-green-600'
                elif delta_count < 0:
                    change_text = str(delta_count)
                    change_class = 'text-red-600'
                else:
                    change_text = '0'
                    change_class = 'text-gray-500'

                team_performance.append({
                    'name': team_name_lookup.get(engineer_name, engineer_name),
                    'initials': _initials(engineer_name),
                    'subtitle': 'Scheduled entries in last 30 days',
                    'count': current_count,
                    'change_text': change_text,
                    'change_class': change_class,
                    'avatar_class': avatar_colors[index % len(avatar_colors)]
                })

            rejected_rate = round(
                (rejected_requests / total_requests) * 100, 1
            ) if total_requests else 0.0
            quality_metrics = {
                'completion_rate': completion_rate,
                'completion_dasharray': f'{completion_rate}, 100',
                'completion_text': f'{completion_rate:.1f}%',
                'primary_label': 'Completion Rate',
                'secondary_metrics': [
                    {
                        'value': f'{rejected_rate:.1f}%',
                        'label': 'Rejected Rate'
                    },
                    {
                        'value': f'{avg_hours_per_schedule:.1f}',
                        'label': 'Avg. Schedule Hours'
                    }
                ]
            }

            def _format_period_delta(current_value, previous_value):
                delta = current_value - previous_value
                if delta > 0:
                    return f'+{delta} vs previous 30 days'
                if delta < 0:
                    return f'{delta} vs previous 30 days'
                return 'No change vs previous 30 days'

            dashboard_data = {
                'summary': {
                    'total_requests': total_requests,
                    'requests_delta_text': _format_period_delta(requests_last_30, requests_previous_30),
                    'active_requests': active_requests,
                    'pending_approval_requests': pending_approval_requests,
                    'active_tests': active_tests,
                    'completed_requests': completed_requests,
                    'completion_rate': completion_rate,
                    'rejected_requests': rejected_requests,
                    'ready_for_report_requests': ready_for_report_requests,
                    'equipment_total': equipment_total,
                    'equipment_needs_calibration': equipment_needs_calibration,
                    'equipment_in_calibration': equipment_in_calibration,
                    'active_lab_engineers': active_lab_engineers,
                },
                'recent_activity': recent_activity,
                'test_distribution': test_distribution,
                'upcoming_deadlines': upcoming_deadlines,
                'lab_utilization': lab_utilization,
                'monthly_trends': monthly_trends,
                'team_performance': team_performance,
                'quality_metrics': quality_metrics,
                'performance_metrics': [
                    {
                        'value': requests_last_30,
                        'label': 'Requests in 30d'
                    },
                    {
                        'value': completed_last_30,
                        'label': 'Completed in 30d'
                    },
                    {
                        'value': avg_hours_per_schedule,
                        'label': 'Avg hours / schedule'
                    },
                    {
                        'value': engineers_scheduled_last_30,
                        'label': 'Engineers active in 30d'
                    }
                ]
            }

            return render_template('dashboard.html', dashboard=dashboard_data)

        except Exception as e:
            logger.error("Error loading dashboard page: %s", e)
            flash('Error loading dashboard page. Please try again.', 'error')
            return redirect(url_for('index'))

    @flask_app.route('/review')
    @login_required
    def review():
        """Display test plan review page for lab engineers and admins."""
        # Check if user has permission to access review page
        if current_user.role not in ['admin', 'lab_engineer']:
            flash('You do not have permission to access the review page.', 'error')
            return redirect(url_for('index'))

        try:
            test_plans, statistics = _get_assigned_tests_context()
            status_filter = (request.args.get('status') or '').strip().lower()
            search_query = (request.args.get('search') or '').strip().lower()
            service_type_filter = (request.args.get('service_type') or '').strip()
            sort_by, sort_dir = _normalize_request_sort_args(
                request.args.get('sort_by'),
                request.args.get('sort_dir')
            )

            def _matches_status(plan_status, selected_status):
                raw = (plan_status or '').strip().lower()
                if not selected_status:
                    return True

                # Treat workflow statuses as in-progress bucket when selected.
                in_progress_group = {
                    'assigned lab engineer',
                    'update plan',
                    'test plan approved',
                    'datasheet uploaded',
                    'proceed report',
                    'admin sign off',
                    'test schedule in progress',
                    'test plan to approve',
                    'in progress',
                }
                cancelled_group = {'rejected', 'cancelled'}

                if selected_status == 'in progress':
                    return raw in in_progress_group
                if selected_status == 'cancelled':
                    return raw in cancelled_group

                return raw == selected_status

            if status_filter:
                test_plans = [
                    plan for plan in test_plans
                    if _matches_status(plan.get('status'), status_filter) or
                    _matches_status(plan.get('status_display'), status_filter)
                ]

            if service_type_filter:
                test_plans = [
                    plan for plan in test_plans
                    if _matches_service_type_filter(
                        plan.get('service_types', []),
                        service_type_filter
                    )
                ]

            if search_query:
                def _to_text(value):
                    return (str(value) if value is not None else '').strip().lower()

                def _matches_search(plan):
                    searchable_fields = [
                        plan.get('tco_id'),
                        plan.get('name'),
                        plan.get('product_name'),
                        plan.get('project'),
                        plan.get('manufacturer'),
                        plan.get('model_number'),
                        plan.get('requester_name'),
                        plan.get('status'),
                        plan.get('status_display'),
                    ]
                    return any(search_query in _to_text(field) for field in searchable_fields)

                test_plans = [plan for plan in test_plans if _matches_search(plan)]

            if sort_by:
                test_plans = _apply_request_identifier_sort(
                    test_plans,
                    sort_by,
                    sort_dir
                )

            return render_template(
                'review.html',
                test_plans=test_plans,
                statistics=statistics,
                service_type_filter=service_type_filter,
                sort_by=sort_by,
                sort_dir=sort_dir,
                service_type_options=_get_service_type_filter_options(
                    [plan.get('service_types', []) for plan in test_plans]
                ),
            )
        except Exception as e:
            logger.error("Error loading review page: %s", e)
            flash('Error loading review page', 'error')
            return redirect(url_for('index'))

    @flask_app.route('/api/test-requests/<int:request_id>/upload-report', methods=['POST'])
    @login_required
    def upload_report(request_id):
        """Upload final test report for a completed test request."""

        # Step 1: Fetch the parent test request.
        test_request = _get_request_or_404(request_id)
        skips_report_review = _request_skips_report_review(test_request)

        # Step 2: Validate incoming data.
        if 'report_file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['report_file']
        comments = request.form.get('comments', '').strip()

        if not file.filename:
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        if not comments:
            return jsonify({'success': False, 'error': 'Comments are required'}), 400

        allowed_report_extensions = {'pdf', 'doc', 'docx'}
        ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
        if ext not in allowed_report_extensions:
            return jsonify({
                'success': False,
                'error': 'Invalid report file type. Allowed: PDF, DOC, DOCX'
            }), 400

        # Step 3: Save file to disk.
        from werkzeug.utils import secure_filename

        filename = secure_filename(file.filename)
        upload_dir = os.path.join(
            current_app.config.get('UPLOAD_FOLDER', 'uploads'),
            'reports',
            str(request_id)
        )
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, filename)
        file.save(file_path)

        # Step 4: Fetch all planner entries for this request.
        planner_entries = PlannerEntry.query.filter_by(
            test_request_id=request_id
        ).all()

        if not planner_entries:
            return jsonify({
                'success': False,
                'error': 'No planner entries found for this test request'
            }), 404

        now = get_ist_now()

        # Step 5: Update PlannerEntry rows (all non-cancelled entries).
        for entry in planner_entries:
            entry.report_file_path = file_path
            entry.report_comments = comments
            entry.report_uploaded_at = now
            entry.report_uploaded_by = current_user.id
            if entry.status != 'cancelled':
                entry.status = 'completed' if skips_report_review else 'report_uploaded'
            entry.updated_at = now

        # Step 6: Update EMC request row.
        test_request.status = 'Completed' if skips_report_review else 'Draft Report'
        test_request.updated_at = now

        # Step 7: Commit both updates in one transaction.
        try:
            db.session.commit()
            logger.info(
                f"Report uploaded for request {request_id}. "
                f"Updated {len(planner_entries)} planner entries + test request status."
            )

            # This request skips report review, so uploading the report completes the TCO -
            # tell the requester (and admins) straight away.
            if skips_report_review:
                send_completion_notification(
                    test_request=test_request,
                    completed_by=current_user.username
                )
        except Exception as e:
            db.session.rollback()
            logger.error(
                f"DB commit failed for upload_report request_id={request_id}: {e}"
            )
            return jsonify({
                'success': False,
                'error': f'Database error: {str(e)}'
            }), 500

        # Step 8: Return success.
        new_status = 'Completed' if skips_report_review else 'Draft Report'
        return jsonify({
            'success': True,
            'message': (
                'Report uploaded successfully. Developmental Assistance request completed.'
                if skips_report_review else
                'Report uploaded successfully'
            ),
            'new_status': new_status,
            'file_path': file_path,
            'updated_planner_entries': len(planner_entries)
        })

    @flask_app.route('/api/test-requests/<int:request_id>/generate-test-report', methods=['POST'])
    @login_required
    def generate_test_report(request_id):
        """Build the IEC-FRM-516 EMI EMC Test Report for this request and route it
        through the same report-approval flow as an uploaded report.

        The report is generated from the DATA of the approved datasheets
        (datasheet_records.form_json / images_json) rather than by concatenating
        their .docx files - see report_gen/."""
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({'success': False, 'error': 'Not authorized'}), 403

        from werkzeug.utils import secure_filename
        from report_gen import build_request_test_report
        from report_gen.registry import FORM_NO as REPORT_FORM_NO

        test_request = _get_request_or_404(request_id)
        skips_report_review = _request_skips_report_review(test_request)

        entries = PlannerEntry.query.filter_by(
            test_request_id=request_id
        ).order_by(
            PlannerEntry.start_date.asc(),
            PlannerEntry.start_time.asc(),
            PlannerEntry.id.asc()
        ).all()
        if not entries:
            return jsonify({'success': False, 'error': 'No planner entries found for this test request'}), 404

        # The report is built from the saved datasheet FORM DATA, so require that
        # at least one non-cancelled test has a submitted datasheet record.
        from datasheet_gen import records as datasheet_records
        active_entries = [
            entry for entry in entries
            if str(entry.status or '').strip().lower() != 'cancelled'
        ]
        with_data = [
            entry for entry in active_entries
            if datasheet_records.get_record_for_assignment(entry.id)
        ]
        if not with_data:
            return jsonify({
                'success': False,
                'error': ('No submitted datasheet data found for this request. '
                          'Complete and approve the datasheets first.')
            }), 400

        payload = request.get_json(silent=True) or {}
        comments = str(payload.get('comments') or request.form.get('comments') or '').strip()
        if not comments:
            comments = f'EMI EMC Test Report generated from {len(with_data)} approved datasheet(s).'

        now = get_ist_now()
        upload_dir = os.path.join(
            current_app.config.get('UPLOAD_FOLDER', 'uploads'), 'reports', str(request_id)
        )
        ts = now.strftime('%Y%m%d_%H%M%S')
        safe_tco = secure_filename(str(test_request.tco_id or f'REQ-{request_id}'))
        output_path = os.path.join(upload_dir, f'{safe_tco}_Test_Report_{ts}.docx')

        try:
            _, report_summary = build_request_test_report(
                test_request, entries, output_path, now=now)
        except ValueError as exc:
            return jsonify({'success': False, 'error': str(exc)}), 400
        except Exception as exc:
            logger.error('EMI EMC Test Report generation failed for request %s: %s', request_id, exc)
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': 'Failed to generate the EMI EMC Test Report'}), 500

        merged_count = len(report_summary.get('tests') or [])
        logger.info(
            'EMI EMC Test Report for request %s: tests=%s dropped=%s images=%d '
            'extra_blocks=%d skipped=%s',
            request_id, report_summary.get('tests'),
            report_summary.get('dropped_sections'), report_summary.get('images', 0),
            report_summary.get('extra_blocks', 0), report_summary.get('skipped'))

        # Route it into the SAME report-approval flow as an uploaded report.
        for entry in entries:
            entry.report_file_path = output_path
            entry.report_comments = comments
            entry.report_uploaded_at = now
            entry.report_uploaded_by = current_user.id
            if entry.status != 'cancelled':
                entry.status = 'completed' if skips_report_review else 'report_uploaded'
            entry.updated_at = now
        test_request.status = 'Completed' if skips_report_review else 'Draft Report'
        test_request.updated_at = now

        try:
            db.session.commit()
        except Exception as exc:
            db.session.rollback()
            logger.error('DB commit failed for generate_test_report request %s: %s', request_id, exc)
            return jsonify({'success': False, 'error': f'Database error: {str(exc)}'}), 500

        # Same terminal state as upload_report: a request that skips report review is
        # complete once the report exists, so notify from here too.
        if skips_report_review:
            send_completion_notification(
                test_request=test_request,
                completed_by=current_user.username
            )

        new_status = 'Completed' if skips_report_review else 'Draft Report'
        notes = []
        if report_summary.get('dropped_sections'):
            notes.append('Not tested (section omitted): '
                         + ', '.join(report_summary['dropped_sections']) + '.')
        if report_summary.get('tests_without_data'):
            notes.append('No saved datasheet data, section left blank: '
                         + ', '.join(report_summary['tests_without_data']) + '.')
        return jsonify({
            'success': True,
            'message': (
                f'EMI EMC Test Report ({REPORT_FORM_NO}) generated for '
                f'{merged_count} test(s). '
                + ('Developmental Assistance request completed. ' if skips_report_review
                   else 'It is now in the report approval flow (Draft Report). ')
                + 'Open it in Word once to let the table of contents and page '
                  'numbers calculate.'
                + ((' ' + ' '.join(notes)) if notes else '')
            ),
            'new_status': new_status,
            'file_path': output_path,
            'datasheet_count': merged_count,
            'report': {
                'tests': report_summary.get('tests'),
                'omitted': report_summary.get('dropped_sections'),
                'without_data': report_summary.get('tests_without_data'),
                'images': report_summary.get('images'),
                'extra_blocks': report_summary.get('extra_blocks'),
            },
        })

    @flask_app.route('/api/test-requests/<int:request_id>/admin-sign-off', methods=['POST'])
    @login_required
    def admin_sign_off(request_id):
        if current_user.role not in ['admin', 'lab_engineer']:
            return jsonify({'success': False, 'error': 'Only admin or lab engineer can send to Admin Sign Off'}), 403

        test_request = _get_request_or_404(request_id)
        data = request.get_json() or {}
        note = (data.get('note') or '').strip()

        try:
            if test_request.status == 'Completed':
                return jsonify({'success': False, 'error': 'Completed reports cannot be sent to Admin Sign Off'}), 400

            # Workflow: Proceed Report -> Admin Sign Off -> Completed
            if test_request.status == 'Admin Sign Off':
                return jsonify({'success': True, 'message': 'Report is already in Admin Sign Off'})

            if test_request.status != 'Proceed Report':
                # FALLBACK for a parent whose status lagged behind its planner entries:
                # if the entries themselves show the report already progressed, allow the
                # transition instead of blocking the user on a stale parent status.
                planner_entries = PlannerEntry.query.filter(
                    db.or_(*_planner_filters(test_request.id, test_request.tco_id))
                ).all()
                planner_can_proceed = any(
                    entry.status in ('Proceed Report', 'Admin Sign Off', 'Completed')
                    for entry in planner_entries
                    if entry.status != 'cancelled'
                )
                if not planner_can_proceed:
                    return jsonify({'success': False, 'error': 'Move report to Proceed Report before sending to Admin Sign Off'}), 400

                current_app.logger.warning(
                    "Admin sign-off for request %s: parent status %r did not match "
                    "'Proceed Report', but its planner entries allowed the transition. "
                    "Forcing the status update.",
                    request_id, test_request.status
                )

            test_request.status = 'Admin Sign Off'
            test_request.reviewed_by = current_user.username
            test_request.reviewed_at = get_ist_now()

            # Also update all planner entries for this request (dual filter).
            planner_entries = PlannerEntry.query.filter(
                db.or_(*_planner_filters(test_request.id, test_request.tco_id))
            ).all()
            for entry in planner_entries:
                if entry.status != 'cancelled':
                    entry.status = 'Admin Sign Off'
                entry.updated_at = get_ist_now()

            # Append sign-off note to review comment thread if provided.
            if note:
                _append_review_comment_entry(
                    test_request=test_request,
                    comment=f'[Admin Sign Off] {note}',
                    username=current_user.username,
                    role=current_user.role
                )

            db.session.commit()
            return jsonify({'success': True, 'message': 'Report signed off successfully'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500

    @flask_app.route('/api/test-requests/<int:request_id>/resync-status', methods=['POST'])
    @login_required
    def resync_request_status(request_id):
        """Force-recompute a parent request's status from its planner entries (admin only).

        Repair tool for records whose parent status is stuck out of sync with the actual
        planner entry states. Uses the same dual filter and the same derivation as
        _update_parent_request_datasheet_status, so the result is what the live flow
        would have produced.
        """
        if current_user.role != 'admin':
            return jsonify({'success': False, 'error': 'Only admin can resync request status'}), 403

        test_request = _get_request_or_404(request_id)

        try:
            all_entries = PlannerEntry.query.filter(
                db.or_(*_planner_filters(test_request.id, test_request.tco_id))
            ).all()

            old_status = test_request.status

            if not all_entries:
                return jsonify({
                    'success': True,
                    'message': 'No planner entries found. Parent status unchanged.',
                    'old_status': old_status,
                    'new_status': old_status,
                    'planner_entries_count': 0
                })

            active_entries = [e for e in all_entries if e.status != 'cancelled']
            if not active_entries:
                active_entries = all_entries

            all_terminal = all(
                entry.status in ('cancelled', 'datasheet_uploaded')
                for entry in active_entries
            )
            has_peer_review_entries = any(
                entry.status == 'Peer Review' for entry in active_entries
            )
            has_datasheet_uploaded_with_file = any(
                entry.status == 'datasheet_uploaded' and bool(entry.datasheet_file_path)
                for entry in active_entries
            )
            all_datasheet_uploaded_have_file = all(
                bool(entry.datasheet_file_path)
                for entry in active_entries
                if entry.status == 'datasheet_uploaded'
            )

            if (
                all_terminal
                and has_datasheet_uploaded_with_file
                and all_datasheet_uploaded_have_file
            ):
                test_request.status = 'Datasheet Uploaded'
            elif has_peer_review_entries:
                test_request.status = 'Peer Review'
            elif test_request.status not in (
                'Assigned Lab Engineer',
                'Update plan',
                'Test Plan Approved',
                'Draft Report',
                'Datasheet Uploaded',
                'Peer Review'
            ):
                test_request.status = 'Test Plan Approved'

            db.session.commit()

            current_app.logger.info(
                "Resync status for request %s by admin %s: %r -> %r",
                request_id, current_user.username, old_status, test_request.status
            )

            return jsonify({
                'success': True,
                'message': f'Status resynced from {len(all_entries)} planner entries.',
                'old_status': old_status,
                'new_status': test_request.status,
                'planner_entries_count': len(all_entries)
            })
        except Exception as e:
            db.session.rollback()
            current_app.logger.error('Error resyncing request %s status: %s', request_id, e)
            return jsonify({'success': False, 'error': str(e)}), 500

    @flask_app.route('/api/test-requests/<int:request_id>/report-access-feedback', methods=['POST'])
    @login_required
    def submit_report_access_feedback(request_id):
        """Require feedback + acknowledgement before enabling requester report access."""
        test_request = _get_request_or_404(request_id)
        if not _can_access_review_thread(test_request):
            return jsonify({
                'success': False,
                'error': 'You do not have permission to access this report'
            }), 403

        planner_entry = PlannerEntry.query.filter(
            PlannerEntry.test_request_id == request_id,
            PlannerEntry.report_file_path != None
        ).first()

        if not planner_entry or not planner_entry.report_file_path:
            return jsonify({'success': False, 'error': 'No report found'}), 404

        data = request.get_json() or {}
        feedback = (data.get('feedback') or '').strip()
        acknowledged = data.get('acknowledged') is True

        if not feedback:
            return jsonify({
                'success': False,
                'error': 'Feedback is required before report access'
            }), 400

        if not acknowledged:
            return jsonify({
                'success': False,
                'error': 'Please acknowledge the report before continuing'
            }), 400

        try:
            _append_review_comment_entry(
                test_request=test_request,
                comment=f'Report feedback before download: {feedback}',
                username=current_user.username,
                role=current_user.role
            )
            _grant_report_access(request_id, planner_entry)
            db.session.commit()
            return jsonify({
                'success': True,
                'message': 'Feedback submitted. View and download are now enabled.',
                'report_access_granted': True
            })
        except Exception as e:
            db.session.rollback()
            logger.error(f'Error submitting report access feedback for request {request_id}: {e}')
            return jsonify({
                'success': False,
                'error': 'Failed to submit report feedback'
            }), 500

    # Download report (triggers file download instead of opening in browser)
    @flask_app.route('/api/test-requests/<int:request_id>/download-report', methods=['GET'])
    @login_required
    def download_report(request_id):
        test_request = _get_request_or_404(request_id)
        if not _can_access_review_thread(test_request):
            return jsonify({
                'success': False,
                'error': 'You do not have permission to access this report'
            }), 403

        planner_entry = PlannerEntry.query.filter(
            PlannerEntry.test_request_id == request_id,
            PlannerEntry.report_file_path != None
        ).first()

        if not planner_entry or not planner_entry.report_file_path:
            return jsonify({'success': False, 'error': 'No report found'}), 404

        if (
            _report_access_requires_feedback(test_request)
            and not _has_report_access_grant(request_id, planner_entry)
        ):
            return (
                'Submit feedback and acknowledge the report before downloading it.',
                403,
                {'Content-Type': 'text/plain; charset=utf-8'}
            )

        from flask import send_file

        # report_file_path may be stored relative to the app root; resolve it the same way
        # view_report does, otherwise the download 404s even though the file exists.
        file_path = os.path.normpath(planner_entry.report_file_path)
        if not os.path.isabs(file_path):
            file_path = os.path.join(flask_app.root_path, file_path)

        if not os.path.exists(file_path):
            current_app.logger.error(
                'Report file not found on disk: %s (stored path: %s)',
                file_path, planner_entry.report_file_path
            )
            return jsonify({'success': False, 'error': 'File not found on disk'}), 404

        # Reports are uploaded as PDF or Word; send the matching type so a .docx is not
        # downloaded with a .pdf content type.
        _ext = os.path.splitext(file_path)[1].lower()
        _mime = ('application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                 if _ext == '.docx' else 'application/pdf')

        return send_file(
            file_path,
            mimetype=_mime,
            as_attachment=True,   # â† triggers download
            download_name=os.path.basename(file_path)
        )

    # 1. Submit a new review comment (appends to JSON array; never overwrites).

    @flask_app.route('/api/test-requests/<int:request_id>/review-comment', methods=['POST'])
    @login_required
    def submit_review_comment(request_id):
        test_request = _get_request_or_404(request_id)
        if not _can_access_review_thread(test_request):
            return jsonify({
                'success': False,
                'error': 'You do not have permission to comment on this request'
            }), 403

        data = request.get_json() or {}
        comment = (data.get('comment') or '').strip()

        if not comment:
            return jsonify({'success': False, 'error': 'Comment is required'}), 400

        try:
            _append_review_comment_entry(
                test_request=test_request,
                comment=comment,
                username=current_user.username,
                role=current_user.role
            )

            db.session.commit()
            return jsonify({'success': True, 'message': 'Comment submitted successfully'})

        except Exception as e:
            db.session.rollback()
            app.logger.error(f'Error submitting review comment: {e}')
            return jsonify({'success': False, 'error': str(e)}), 500

    # 2. Get all review comments for a request.

    @flask_app.route('/api/test-requests/<int:request_id>/review-comments', methods=['GET'])
    @login_required
    def get_review_comments(request_id):
        test_request = _get_request_or_404(request_id)
        if not _can_access_review_thread(test_request):
            return jsonify({
                'success': False,
                'error': 'You do not have permission to view comments on this request'
            }), 403

        comments = _get_combined_review_comment_thread(test_request)

        return jsonify({'success': True, 'comments': comments})

    # 3. Proceed with report (marks request as Proceed Report).

    @flask_app.route('/api/test-requests/<int:request_id>/proceed-report', methods=['POST'])
    @login_required
    def proceed_report(request_id):
        test_request = _get_request_or_404(request_id)
        if not _can_access_review_thread(test_request):
            return jsonify({
                'success': False,
                'error': 'You do not have permission to approve this report'
            }), 403
        try:
            current_status = (test_request.status or '').strip().lower()
            if not _report_is_approvable_status(current_status):
                # FALLBACK for a parent whose status lagged behind its planner entries:
                # if an entry already holds a draft report in an approvable state, allow
                # the transition rather than blocking on the stale parent status.
                planner_entries = PlannerEntry.query.filter(
                    db.or_(*_planner_filters(test_request.id, test_request.tco_id))
                ).all()
                planner_has_draft_report = any(
                    entry.report_file_path
                    and _report_is_approvable_status((entry.status or '').strip().lower())
                    for entry in planner_entries
                    if entry.status != 'cancelled'
                )
                if not planner_has_draft_report:
                    return jsonify({
                        'success': False,
                        'error': 'Report approval is allowed only when the draft report is ready.'
                    }), 400

                current_app.logger.warning(
                    "Proceed report for request %s: parent status %r was not approvable, "
                    "but its planner entries held a draft report. Forcing the status update.",
                    request_id, test_request.status
                )

            planner_entry = PlannerEntry.query.filter(
                db.or_(*_planner_filters(test_request.id, test_request.tco_id)),
                PlannerEntry.report_file_path != None
            ).first()
            if (
                planner_entry
                and _report_access_requires_feedback(test_request)
                and not _has_report_access_grant(request_id, planner_entry)
            ):
                return jsonify({
                    'success': False,
                    'error': 'Submit feedback and acknowledge the report before approving it.'
                }), 403

            test_request.status = 'Proceed Report'

            # Also update all planner entries for this request (dual filter).
            planner_entries = PlannerEntry.query.filter(
                db.or_(*_planner_filters(test_request.id, test_request.tco_id))
            ).all()
            for entry in planner_entries:
                entry.status = 'Proceed Report'
                entry.updated_at = get_ist_now()

            db.session.commit()
            return jsonify({
                'success': True,
                'message': 'Report approved successfully'
            })
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500

    # 4. View report PDF in browser.

    @flask_app.route('/api/test-requests/<int:request_id>/view-report', methods=['GET'])
    @login_required
    def view_report(request_id):
        test_request = _get_request_or_404(request_id)
        if not _can_access_review_thread(test_request):
            return jsonify({
                'success': False,
                'error': 'You do not have permission to access this report'
            }), 403

        planner_entry = PlannerEntry.query.filter(
            PlannerEntry.test_request_id == request_id,
            PlannerEntry.report_file_path != None
        ).first()

        if not planner_entry or not planner_entry.report_file_path:
            return jsonify({'success': False, 'error': 'No report found'}), 404

        if (
            _report_access_requires_feedback(test_request)
            and not _has_report_access_grant(request_id, planner_entry)
        ):
            return (
                'Submit feedback and acknowledge the report before viewing it.',
                403,
                {'Content-Type': 'text/plain; charset=utf-8'}
            )

        file_path = os.path.normpath(planner_entry.report_file_path)

        if not os.path.isabs(file_path):
            file_path = os.path.join(flask_app.root_path, file_path)

        if not os.path.exists(file_path):
            # 'app' is not defined at module scope - this used to raise NameError and
            # return a 500 instead of the intended 404.
            current_app.logger.error(
                'Report file not found on disk: %s (stored path: %s)',
                file_path, planner_entry.report_file_path
            )
            return jsonify({'success': False, 'error': 'Report file not found on disk'}), 404

        # A consolidated Test Report is a .docx, which browsers can't render inline —
        # serve it as a download; PDFs still open inline.
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.docx':
            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=os.path.basename(file_path)
            )
        return send_file(
            file_path,
            mimetype='application/pdf',
            as_attachment=False,
            download_name=os.path.basename(file_path)
        )

    @flask_app.route('/api/test-requests/<int:request_id>/view-report-data', methods=['GET'])
    @login_required
    def view_report_data(request_id):
        test_request = _get_request_or_404(request_id)
        if not _can_access_review_thread(test_request):
            return jsonify({
                'success': False,
                'error': 'You do not have permission to access this report'
            }), 403

        planner_entry = PlannerEntry.query.filter(
            PlannerEntry.test_request_id == request_id,
            PlannerEntry.report_file_path != None
        ).first()

        has_report = False
        if planner_entry and planner_entry.report_file_path:
            file_path = os.path.normpath(planner_entry.report_file_path)
            if not os.path.isabs(file_path):
                file_path = os.path.join(flask_app.root_path, file_path)
            has_report = os.path.exists(file_path)

        comments = _get_combined_review_comment_thread(test_request)
        report_access_required = _report_access_requires_feedback(test_request)
        report_access_granted = (
            not report_access_required
            or (
                planner_entry is not None
                and _has_report_access_grant(request_id, planner_entry)
            )
        )
        current_status = (test_request.status or '').strip().lower()
        can_approve_report = _report_is_approvable_status(current_status)

        entry_data = None
        if planner_entry:
            entry_data = {
                'test_name': getattr(planner_entry, 'test_name', 'Test Report') or 'Test Report',
                'report_file_path': f'/api/test-requests/{request_id}/view-report' if has_report else None,
                'report_uploaded_at': planner_entry.report_uploaded_at.strftime('%d %b %Y, %I:%M %p')
                if getattr(planner_entry, 'report_uploaded_at', None) else None,
                'report_comments': getattr(planner_entry, 'report_comments', None),
            }

        return jsonify({
            'success': True,
            'tco_id': test_request.tco_id,
            'product_name': test_request.product_name,
            'status': test_request.status,
            'service_types': _extract_service_types(test_request),
            'skip_review_flow': _request_skips_report_review(test_request),
            'require_feedback_before_report_access': report_access_required,
            'report_access_granted': report_access_granted,
            'can_approve_report': can_approve_report,
            'planner_entries': [entry_data] if entry_data else [],
            'review_comments': comments,
        })

    @flask_app.route('/assigned-tests')
    @login_required
    def assigned_test():
        """Dedicated page showing assigned test requests."""
        if current_user.role not in ['admin', 'lab_engineer']:
            flash(
                'You do not have permission to access the assigned tests page.', 'error')
            return redirect(url_for('index'))

        try:
            test_plans, statistics = _get_assigned_tests_context()
            return render_template(
                'assigned_test.html',
                test_plans=test_plans,
                statistics=statistics,
                assigned_status_options=_get_assignment_status_filter_options(
                    test_plans
                ),
                peer_reviewer_options=_serialize_peer_reviewer_candidates(
                    exclude_user_id=current_user.id
                ),
                service_type_options=_get_service_type_filter_options(
                    [plan.get('service_types', []) for plan in test_plans]
                ),
            )
        except Exception as e:
            logger.error("Error loading assigned tests page: %s", e)
            import traceback
            # This will show the full stack trace
            logger.error(traceback.format_exc())
            flash('Error loading assigned tests page', 'error')
            return redirect(url_for('index'))
    from werkzeug.utils import secure_filename

    # Configuration
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads', 'test_datasheets')
    SURGE_DATASHEET_TEMPLATE_PATH = (
        r'c:\Users\saimounik.chandavolu\OneDrive - Thermo Fisher Scientific'
        r'\Desktop\EMI EMC\Report automation inputs\Format Templates'
        r'\02_Test Datasheet\Surge\IEC-FRM-510 Surge Data Sheet.docx'
    )
    def _datasheet_text(value, default=''):
        """Normalize values before writing them into DOCX cells."""
        if value is None:
            return default
        if isinstance(value, list):
            parts = [str(item).strip() for item in value if str(item).strip()]
            return ', '.join(parts) if parts else default
        text_value = str(value).strip()
        return text_value if text_value else default

    def _datasheet_list(value):
        """Return a normalized flat string list from a form payload value."""
        if value is None:
            return []
        if isinstance(value, list):
            return [str(item).strip() for item in value if str(item).strip()]
        text_value = str(value).strip()
        return [text_value] if text_value else []

    def _get_peer_reviewer_candidates(exclude_user_id=None):
        """Return active admin/lab engineer users eligible for peer review."""
        query = User.query.filter(
            User.role.in_(['lab_engineer', 'admin']),
            User.is_active.is_(True)
        )

        if exclude_user_id is not None:
            query = query.filter(User.id != exclude_user_id)

        return query.order_by(User.role.asc(), User.username.asc()).all()

    def _serialize_peer_reviewer_candidates(exclude_user_id=None):
        reviewers = _get_peer_reviewer_candidates(exclude_user_id=exclude_user_id)
        return [
            {
                'id': reviewer.id,
                'name': reviewer.username,
                'email': reviewer.email,
                'role': reviewer.role
            }
            for reviewer in reviewers
        ]

    def _resolve_peer_reviewer(peer_reviewer_id_raw, uploader_user_id):
        """Validate peer reviewer selection for datasheet submission."""
        peer_reviewer_value = str(peer_reviewer_id_raw or '').strip()
        if not peer_reviewer_value:
            return None, 'Peer reviewer is required before submitting the datasheet.'

        try:
            peer_reviewer_id = int(peer_reviewer_value)
        except (TypeError, ValueError):
            return None, 'Selected peer reviewer is invalid.'

        if uploader_user_id and peer_reviewer_id == uploader_user_id:
            return None, 'You cannot assign yourself as the peer reviewer.'

        reviewer = User.query.filter(
            User.id == peer_reviewer_id,
            User.role.in_(['lab_engineer', 'admin']),
            User.is_active.is_(True)
        ).first()

        if not reviewer:
            return None, 'Selected peer reviewer is no longer available.'

        return reviewer, None

    def _safe_table_cell_text(table, row_idx, col_idx, value):
        """Safely write a value into a DOCX table cell."""
        try:
            table.rows[row_idx].cells[col_idx].text = _datasheet_text(value)
        except Exception as exc:
            logger.debug(
                'Skipping datasheet cell write (row=%s col=%s): %s',
                row_idx,
                col_idx,
                exc
            )

    def _safe_repeating_row_values(table, start_row, rows_to_write):
        """Populate a table section with row data while preserving template layout."""
        required_rows = start_row + len(rows_to_write)
        while len(table.rows) < required_rows:
            table.add_row()

        for offset, row_values in enumerate(rows_to_write):
            row = table.rows[start_row + offset]
            for col_idx, value in enumerate(row_values):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = _datasheet_text(value)

    def _set_paragraph_after_label(doc, label_text, value):
        """Write value into the paragraph immediately after a label heading."""
        normalized_label = ' '.join(str(label_text or '').split()).upper()
        for idx, paragraph in enumerate(doc.paragraphs):
            paragraph_label = ' '.join((paragraph.text or '').split()).upper()
            if paragraph_label == normalized_label and idx + 1 < len(doc.paragraphs):
                doc.paragraphs[idx + 1].text = _datasheet_text(value)
                return True
        return False

    def _set_paragraph_containing(doc, contains_text, value):
        """Replace a paragraph text when it contains a target snippet."""
        normalized_target = ' '.join(str(contains_text or '').split()).upper()
        for paragraph in doc.paragraphs:
            paragraph_text = ' '.join((paragraph.text or '').split()).upper()
            if normalized_target in paragraph_text:
                paragraph.text = _datasheet_text(value)
                return True
        return False

    def _insert_image_before_caption(doc, caption_text, image_value, width_inches=5.5):
        """Insert an image immediately above a matching caption paragraph."""
        image_bytes = _load_block_diagram_image_bytes(image_value)
        if not image_bytes:
            return

        normalized_caption = ' '.join(str(caption_text or '').split()).upper()
        for paragraph in doc.paragraphs:
            paragraph_text = ' '.join((paragraph.text or '').split()).upper()
            if normalized_caption in paragraph_text:
                try:
                    image_paragraph = paragraph.insert_paragraph_before('')
                    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    image_paragraph.add_run().add_picture(
                        BytesIO(image_bytes),
                        width=Inches(width_inches)
                    )
                except Exception as exc:
                    logger.warning('Unable to insert Surge image for %s: %s', caption_text, exc)
                return

    def _insert_images_after_paragraph_containing(doc, contains_text, image_values, width_inches=5.5):
        """Insert one or more images immediately after a paragraph containing the target text."""
        if not image_values:
            return

        if not isinstance(image_values, list):
            image_values = [image_values]

        normalized_target = ' '.join(str(contains_text or '').split()).upper()
        if not normalized_target:
            return

        for paragraph in doc.paragraphs:
            paragraph_text = ' '.join((paragraph.text or '').split()).upper()
            if normalized_target in paragraph_text:
                anchor_paragraph = paragraph
                for image_value in image_values:
                    image_bytes = _load_block_diagram_image_bytes(image_value)
                    if not image_bytes:
                        continue
                    try:
                        image_paragraph = _insert_paragraph_after(anchor_paragraph, '')
                        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        image_paragraph.add_run().add_picture(
                            BytesIO(image_bytes),
                            width=Inches(width_inches)
                        )
                        anchor_paragraph = image_paragraph
                    except Exception as exc:
                        logger.warning('Unable to insert Surge functional check image: %s', exc)
                return

    def _update_parent_request_datasheet_status(assignment):
        """Update parent request status after a datasheet is generated or uploaded.

        Queries planner entries by BOTH test_request_id and tco_id (dual filter), so
        entries linked to the parent by only one of the two are still considered -
        missing them is what left parent status out of sync.
        """
        db.session.flush()

        all_entries = PlannerEntry.query.filter(
            db.or_(*_planner_filters(assignment.test_request_id, assignment.tco_id))
        ).all()

        if not all_entries:
            return

        # âœ… FIX: Only consider active (non-cancelled) entries for terminal state check
        # This allows rescheduled tests (old entries cancelled, new entry active) to properly upload
        active_entries = [e for e in all_entries if e.status != 'cancelled']
        
        if not active_entries:
            # All entries are cancelled, consider it terminal
            active_entries = all_entries

        all_terminal = all(
            entry.status in ('cancelled', 'datasheet_uploaded')
            for entry in active_entries
        )
        has_peer_review_entries = any(
            entry.status == 'Peer Review'
            for entry in active_entries
        )
        has_datasheet_uploaded_with_file = any(
            entry.status == 'datasheet_uploaded' and bool(entry.datasheet_file_path)
            for entry in active_entries
        )
        all_datasheet_uploaded_have_file = all(
            bool(entry.datasheet_file_path)
            for entry in active_entries
            if entry.status == 'datasheet_uploaded'
        )

        parent_request = _resolve_request(assignment.test_request_id)
        if not parent_request:
            return

        if (
            all_terminal
            and has_datasheet_uploaded_with_file
            and all_datasheet_uploaded_have_file
        ):
            parent_request.status = 'Datasheet Uploaded'
            current_app.logger.info(
                "All active entries terminal for request %s. Parent status updated to 'Datasheet Uploaded'.",
                assignment.test_request_id
            )
        elif has_peer_review_entries:
            parent_request.status = 'Peer Review'
            current_app.logger.info(
                "Request %s has datasheets awaiting peer review. Parent status updated to 'Peer Review'.",
                assignment.test_request_id
            )
        elif parent_request.status not in (
            'Assigned Lab Engineer',
            'Update plan',
            'Test Plan Approved',
            'Draft Report',
            'Datasheet Uploaded',
            'Peer Review'
        ):
            parent_request.status = 'Test Plan Approved'
            current_app.logger.info(
                "Request %s still has pending entries. Parent status set to 'Test Plan Approved'.",
                assignment.test_request_id
            )

    def _format_checkbox_option_line(value, options):
        """Render selected values against a fixed option list using checkbox glyphs."""
        raw_text = _datasheet_text(value, '')
        if not raw_text:
            return ''

        normalized_text = str(raw_text).strip()
        if not normalized_text:
            return ''

        normalized_compact = normalized_text.lower().replace(' ', '')
        if normalized_compact in {'asperthestandard', 'as/perthestandard'}:
            return 'As per the standard'

        tokens = {
            token.strip().lower().replace(' ', '')
            for token in re.split(r'[,/|;]+', normalized_text)
            if token and token.strip()
        }

        normalized_options = [
            (option, option.lower().replace(' ', ''))
            for option in options
        ]
        matched_any = any(option_key in tokens for _, option_key in normalized_options)
        if not matched_any:
            return normalized_text

        return ' '.join(
            f"{'☒' if option_key in tokens else '☐'} {option}"
            for option, option_key in normalized_options
        )

    def _build_surge_datasheet_docx(template_path, output_path, parent_request, assignment, form_data):
        """Populate the external Surge datasheet template and save it."""
        doc = DocxDocument(template_path)
        request_data = parent_request.to_dict()

        def form_value(key, default=''):
            return _datasheet_text(form_data.get(key), default)

        def form_values(key):
            return _datasheet_list(form_data.get(key))

        product_standards = request_data.get('product_standards', [])
        surge_standards = request_data.get('surge_standard', [])
        serial_number = parent_request.serial_number or ', '.join(
            item for item in (request_data.get('serial_numbers') or [])
            if item
        )

        test_spec_table = doc.tables[0]
        specification_values = [
            form_value('job_number', parent_request.job_number or ''),
            form_value('eut_name', parent_request.product_name or ''),
            form_value('eut_model', parent_request.model_number or ''),
            form_value('eut_serial', serial_number),
            form_value('basic_standard', ', '.join(surge_standards)),
            form_value('product_standard', ', '.join(product_standards)),
            form_value('immunity_test_requirement'),
            form_value('test_port'),
            _format_checkbox_option_line(form_value('test_voltage_cm'), ['±0.5', '±1', '±2', '±4']),
            _format_checkbox_option_line(form_value('test_voltage_dm'), ['±0.5', '±1', '±2', '±4']),
            form_value('coupling_phases'),
            form_value('repetition_rate'),
            form_value('test_mode'),
            form_value('eut_configuration'),
            form_value('eut_input_voltage_freq', request_data.get('surge_voltage_freq') or ''),
            form_value('ambient_temperature'),
            form_value('relative_humidity'),
            form_value('test_date'),
            form_value('tested_by', assignment.test_person_name or parent_request.assigned_engineer_name or '')
        ]
        for row_idx, row_value in enumerate(specification_values):
            if row_idx < len(test_spec_table.rows):
                for col_idx in range(2, len(test_spec_table.rows[row_idx].cells)):
                    test_spec_table.rows[row_idx].cells[col_idx].text = _datasheet_text(row_value)

        equipment_rows = []
        equipment_names = form_values('surge_equipment_name[]')
        equipment_makes = form_values('surge_equipment_make[]')
        equipment_models = form_values('surge_equipment_model[]')
        equipment_serials = form_values('surge_equipment_serial[]')
        equipment_cal_dues = form_values('surge_equipment_cal_due[]')
        equipment_count = max(
            len(equipment_names),
            len(equipment_makes),
            len(equipment_models),
            len(equipment_serials),
            len(equipment_cal_dues),
            0
        )
        for index in range(equipment_count):
            equipment_rows.append([
                index + 1,
                equipment_names[index] if index < len(equipment_names) else '',
                equipment_makes[index] if index < len(equipment_makes) else '',
                equipment_models[index] if index < len(equipment_models) else '',
                equipment_serials[index] if index < len(equipment_serials) else '',
                equipment_cal_dues[index] if index < len(equipment_cal_dues) else '',
            ])
        if equipment_rows:
            _safe_repeating_row_values(doc.tables[1], 1, equipment_rows)

        modification_rows = []
        modification_states = form_values('surge_modification_state[]')
        modification_descriptions = form_values('surge_modification_description[]')
        modification_fitted_by = form_values('surge_modification_fitted_by[]')
        modification_dates = form_values('surge_modification_date[]')
        modification_count = max(
            len(modification_states),
            len(modification_descriptions),
            len(modification_fitted_by),
            len(modification_dates),
            0
        )
        for index in range(modification_count):
            modification_rows.append([
                modification_states[index] if index < len(modification_states) else '',
                modification_descriptions[index] if index < len(modification_descriptions) else '',
                modification_fitted_by[index] if index < len(modification_fitted_by) else '',
                modification_dates[index] if index < len(modification_dates) else '',
            ])
        if modification_rows:
            _safe_repeating_row_values(doc.tables[2], 1, modification_rows)

        def fill_power_table(table, segment_start):
            pos_field_names = [
                'surge_l1pe_0_pos[]', 'surge_l1pe_90_pos[]', 'surge_l1pe_180_pos[]', 'surge_l1pe_270_pos[]',
                'surge_npe_0_pos[]', 'surge_npe_90_pos[]', 'surge_npe_180_pos[]', 'surge_npe_270_pos[]',
                'surge_l1npe_0_pos[]', 'surge_l1npe_90_pos[]', 'surge_l1npe_180_pos[]', 'surge_l1npe_270_pos[]',
                'surge_l1n_0_pos[]', 'surge_l1n_90_pos[]', 'surge_l1n_180_pos[]', 'surge_l1n_270_pos[]'
            ]
            neg_field_names = [
                'surge_l1pe_0_neg[]', 'surge_l1pe_90_neg[]', 'surge_l1pe_180_neg[]', 'surge_l1pe_270_neg[]',
                'surge_npe_0_neg[]', 'surge_npe_90_neg[]', 'surge_npe_180_neg[]', 'surge_npe_270_neg[]',
                'surge_l1npe_0_neg[]', 'surge_l1npe_90_neg[]', 'surge_l1npe_180_neg[]', 'surge_l1npe_270_neg[]',
                'surge_l1n_0_neg[]', 'surge_l1n_90_neg[]', 'surge_l1n_180_neg[]', 'surge_l1n_270_neg[]'
            ]

            positive_values = {field: form_values(field) for field in pos_field_names}
            negative_values = {field: form_values(field) for field in neg_field_names}

            for row_idx in range(3, min(len(table.rows), 11)):
                for col_idx in range(1, len(table.rows[row_idx].cells)):
                    table.rows[row_idx].cells[col_idx].text = ''

            for row_idx, value_index in ((3, segment_start + 0), (5, segment_start + 1)):
                if row_idx < len(table.rows):
                    for col_offset, field in enumerate(pos_field_names, start=1):
                        values = positive_values.get(field, [])
                        table.rows[row_idx].cells[col_offset].text = (
                            _datasheet_text(values[value_index]) if value_index < len(values) else ''
                        )

            for row_idx, value_index in ((4, segment_start + 0), (6, segment_start + 1)):
                if row_idx < len(table.rows):
                    for col_offset, field in enumerate(neg_field_names, start=1):
                        values = negative_values.get(field, [])
                        table.rows[row_idx].cells[col_offset].text = (
                            _datasheet_text(values[value_index]) if value_index < len(values) else ''
                        )

        fill_power_table(doc.tables[3], 0)
        fill_power_table(doc.tables[4], 2)

        software_rows = []
        software_used = form_values('surge_software_used[]')
        software_versions = form_values('surge_software_version[]')
        software_count = max(len(software_used), len(software_versions), 0)
        for index in range(software_count):
            software_rows.append([
                software_used[index] if index < len(software_used) else '',
                software_versions[index] if index < len(software_versions) else ''
            ])
        if software_rows:
            _safe_repeating_row_values(doc.tables[6], 1, software_rows)

        _safe_table_cell_text(doc.tables[7], 0, 1, form_value('required_performance_criteria'))
        _safe_table_cell_text(doc.tables[7], 1, 1, form_value('met_performance_criteria'))

        tested_by = form_value('tested_by', assignment.test_person_name or '')
        test_date = form_value('test_date')
        _safe_table_cell_text(doc.tables[8], 1, 1, tested_by)
        _safe_table_cell_text(doc.tables[8], 2, 1, 'Generated electronically')
        _safe_table_cell_text(doc.tables[8], 3, 1, test_date)

        functional_text = (
            f"Functional check is conducted as per SOP reference number: "
            f"{form_value('functional_check_sop', 'IEC-SOP-505')}."
        )
        output_verification = form_value('output_verification')
        if output_verification:
            functional_text = f"{functional_text}\nOutput verification: {output_verification}"
        functional_check_notes = form_value('functional_check_notes')
        if functional_check_notes:
            functional_text = f"{functional_text}\nAdditional notes: {functional_check_notes}"
        _insert_images_after_paragraph_containing(
            doc,
            'Functional check is conducted as per SOP reference number:',
            form_data.get('surge_functional_check_images')
        )
        _set_paragraph_containing(
            doc,
            'Functional check is conducted as per SOP reference number:',
            functional_text
        )

        monitoring_bits = []
        if form_value('monitoring_parameters'):
            monitoring_bits.append(form_value('monitoring_parameters'))
        _set_paragraph_containing(
            doc,
            '<< Update this section with the monitoring parameters mentioned in the IEC-FRM-502 EMI EMC Questionnaire>>',
            '\n'.join(monitoring_bits)
        )

        setup_images = form_data.get('surge_setup_images')
        if not isinstance(setup_images, list):
            setup_images = [img for img in [
                form_data.get('surge_power_setup_image'),
                form_data.get('surge_signal_setup_image')
            ] if img]
            legacy_additional_images = form_data.get('surge_additional_setup_images')
            if isinstance(legacy_additional_images, list):
                setup_images.extend([img for img in legacy_additional_images if img])
            elif legacy_additional_images:
                setup_images.append(legacy_additional_images)

        if setup_images:
            if len(setup_images) >= 1:
                _insert_image_before_caption(
                    doc,
                    'Figure 1: Photograph of Surge test setup on Power Line',
                    setup_images[0]
                )
            if len(setup_images) >= 2:
                _insert_image_before_caption(
                    doc,
                    'Figure 2: Photograph of Surge test setup on Signal Line',
                    setup_images[1]
                )
            if len(setup_images) > 2:
                _insert_images_after_paragraph_containing(
                    doc,
                    'Figure 2: Photograph of Surge test setup on Signal Line',
                    setup_images[2:]
                )

        _set_paragraph_after_label(doc, 'RESULT:', form_value('result'))

        doc.save(output_path)

    @flask_app.route('/generate-surge-datasheet', methods=['POST'])
    @login_required
    def generate_surge_datasheet():
        """Persist the filled Surge form and send it to peer review."""
        try:
            payload = request.get_json(silent=True) or {}
            form_data = payload.get('form_data') or {}
            assignment_id = payload.get('assignment_id')
            test_request_id = payload.get('test_request_id') or payload.get('test_plan_id')
            tco_id = payload.get('tco_id')
            peer_reviewer_id = (
                payload.get('peer_reviewer_id')
                or form_data.get('peer_reviewer_id')
            )

            if not assignment_id:
                return jsonify({'success': False, 'message': 'Assignment ID is required'}), 400

            try:
                assignment_id = int(assignment_id)
            except (TypeError, ValueError):
                return jsonify({'success': False, 'message': 'Invalid assignment ID'}), 400

            assignment = db.session.get(PlannerEntry, assignment_id)
            if not assignment:
                return jsonify({'success': False, 'message': 'Assignment not found'}), 404

            if 'surge' not in (assignment.test_name or '').lower():
                return jsonify({'success': False, 'message': 'This generator is available only for Surge assignments'}), 400

            parent_request = db.session.get(
                EMCRequest,
                assignment.test_request_id or test_request_id
            )
            if not parent_request:
                return jsonify({'success': False, 'message': 'Parent test request not found'}), 404

            if not _can_access_iec_request(
                parent_request,
                allow_lab_engineer=True,
                require_assigned_lab_engineer=True
            ):
                return jsonify({'success': False, 'message': 'Access denied'}), 403

            if (
                current_user.role == 'lab_engineer'
                and assignment.engineer_user_id
                and assignment.engineer_user_id != current_user.id
            ):
                return jsonify({'success': False, 'message': 'You can only generate datasheets for your own assignments'}), 403

            effective_peer_reviewer_id = peer_reviewer_id
            if (
                not str(peer_reviewer_id or '').strip()
                and getattr(assignment, 'peer_reviewer_user_id', None)
            ):
                effective_peer_reviewer_id = str(assignment.peer_reviewer_user_id)

            peer_reviewer, reviewer_error = _resolve_peer_reviewer(
                effective_peer_reviewer_id,
                current_user.id
            )
            if reviewer_error:
                return jsonify({'success': False, 'message': reviewer_error}), 400

            from datasheet_gen import records as datasheet_records

            submitted_at = get_ist_now()
            assignment.datasheet_file_path = None
            assignment.datasheet_uploaded_at = submitted_at
            assignment.datasheet_uploaded_by = current_user.id
            assignment.peer_reviewer_user_id = peer_reviewer.id
            assignment.peer_review_assigned_at = submitted_at
            assignment.status = 'Peer Review'
            _append_datasheet_peer_review_comment(
                assignment,
                (
                    f"Surge datasheet form submitted to {peer_reviewer.username} for peer review. "
                    "Final Word datasheet will be generated after approval."
                ),
                current_user.username,
                'SENT FOR REVIEW'
            )

            test_date_raw = _datasheet_text(form_data.get('test_date'))
            if test_date_raw:
                try:
                    generated_test_date = datetime.strptime(test_date_raw, '%Y-%m-%d').date()
                    assignment.completion_date = generated_test_date
                    assignment.start_date = assignment.start_date or generated_test_date
                    assignment.end_date = generated_test_date
                except ValueError:
                    logger.warning('Unable to parse generated Surge test date: %s', test_date_raw)

            _update_parent_request_datasheet_status(assignment)
            datasheet_records.upsert_record(
                assignment,
                'SURGE',
                form_data,
                {},
                datasheet_records.SUBMITTED,
                generated_file_path='',
                user=current_user
            )

            return jsonify({
                'success': True,
                'message': f'Surge datasheet form sent to {peer_reviewer.username} for peer review',
                'peer_reviewer_name': peer_reviewer.username
            })

        except Exception as exc:
            db.session.rollback()
            logger.error('Error generating Surge datasheet: %s', exc)
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'message': 'An error occurred while sending the Surge datasheet for peer review'
            }), 500

    @flask_app.route('/cancel-test-assignment', methods=['POST'])
    @login_required
    def cancel_test_assignment():
        try:
            data = request.get_json()
            assignment_id = data.get('assignment_id')
            cancel_reason = data.get('cancel_reason')

            if not assignment_id or not cancel_reason:
                return jsonify({'success': False, 'message': 'Assignment ID and cancel reason are required'}), 400

            # Get the assignment
            assignment = db.session.get(PlannerEntry, assignment_id)
            if not assignment:
                return jsonify({'success': False, 'message': 'Assignment not found'}), 404

            # Update assignment status to cancelled
            assignment.status = 'cancelled'
            assignment.cancel_reason = cancel_reason
            assignment.cancelled_at = get_ist_now()
            assignment.cancelled_by = current_user.id

            db.session.commit()

            return jsonify({
                'success': True,
                'message': 'Test assignment cancelled successfully'
            })

        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error cancelling test assignment: {str(e)}")
            return jsonify({'success': False, 'message': f'Server error: {str(e)}'}), 500

    @flask_app.route('/admin-approval')
    @login_required
    def admin_approval():
        """Display admin approval page for test plans."""
        # Check if user has admin permission
        if current_user.role != 'admin':
            flash(
                'You do not have permission to access the admin approval page.', 'error')
            return redirect(url_for('index'))

        try:
            submitted_nulls_last = db.case(
                (EMCRequest.submitted_at.is_(None), 1),
                else_=0
            )

            approval_statuses = [
                'At Review',
                'Approved',
                'Test Plan Approved',
                'Assigned',
                'Assigned Lab Engineer',
                'In Progress',
                'Test Schedule In Progress',
                'Test Plan To Approve',
                'Admin Sign Off',
                'Draft Report',
                'Proceed Report',
                'Datasheet Uploaded',
                'Update plan',
                'Report Uploaded',
                'report_uploaded',
                'Need More Information',
                'Completed'
            ]

            at_review_requests = EMCRequest.query.options(
                joinedload(EMCRequest.service_types),
                joinedload(EMCRequest.assigned_engineer)
            ).filter(
                EMCRequest.status.in_(approval_statuses)
            ).order_by(
                submitted_nulls_last,
                EMCRequest.submitted_at.desc(),
                EMCRequest.created_at.desc()
            ).all()

            # Finished work goes to the end of the approval list.
            at_review_requests = _terminal_last(at_review_requests)

            test_plans = []
            for request in at_review_requests:
                assigned_display = request.assigned_engineer_name
                if not assigned_display and request.assigned_engineer:
                    assigned_display = request.assigned_engineer.username
                shared_review_thread = _get_combined_review_comment_thread(request)
                latest_shared_comment = shared_review_thread[-1] if shared_review_thread else None

                # Map DB status to UI keys/labels
                if request.status == 'At Review':
                    status_key = 'pending_approval'
                    status_label = request.status
                elif request.status in ('Test Schedule In Progress', 'Test Plan To Approve', 'In Progress', 'Admin Sign Off', 'Draft Report', 'Proceed Report', 'Datasheet Uploaded', 'Report Uploaded', 'report_uploaded', 'Update plan'):
                    status_key = 'in_progress'
                    status_label = request.status
                elif request.status in ('Approved', 'Test Plan Approved'):
                    status_key = 'approved'
                    status_label = request.status
                elif request.status in ('Assigned', 'Assigned Lab Engineer'):
                    status_key = 'assigned'
                    status_label = request.status
                elif request.status == 'Need More Information':
                    status_key = 'need_more_info'
                    status_label = request.status
                elif request.status == 'Completed':
                    status_key = 'completed'
                    status_label = request.status
                elif request.status == 'Cancelled':
                    status_key = 'cancelled'
                    status_label = request.status
                elif request.status == 'Rejected':
                    status_key = 'rejected'
                    status_label = request.status
                else:
                    status_key = 'other'
                    status_label = request.status or 'Unknown'

                test_plans.append({
                    'id': request.id,
                    'tco_id': request.tco_id or f'REQ-{request.id}',
                    'job_id': request.job_id,
                    'job_number': request.job_number or request.job_id or '',
                    'name': request.product_name or 'Unnamed Product',
                    'project': request.manufacturer or 'N/A',
                    'status_key': status_key,
                    'status_label': status_label,
                    'requester_name': request.requester_name or 'Unknown',
                    'service_types': _extract_service_types(request),
                    'submitted_date': request.submitted_at or request.created_at,
                    'assigned_engineer_id': request.assigned_engineer_id,
                    'assigned_engineer_name': request.assigned_engineer_name,
                    'assigned_engineer_display': assigned_display,
                    'has_review_messages': len(shared_review_thread) > 0,
                    'review_comments': latest_shared_comment.get('comment') if latest_shared_comment else None,
                    'reviewed_by': latest_shared_comment.get('username') if latest_shared_comment else getattr(request, 'reviewed_by', None),
                    'reviewed_at': latest_shared_comment.get('created_at') if latest_shared_comment else (
                        request.reviewed_at.isoformat() if getattr(request, 'reviewed_at', None) else None
                    )
                })

            assignable_users = User.query.filter(
                User.role.in_(['lab_engineer', 'admin']),
                User.is_active.is_(True)
            ).order_by(User.role.asc(), User.username.asc()).all()
            lab_engineers_data = [
                {
                    'id': user.id,
                    'name': user.username,
                    'email': user.email,
                    'role': user.role
                } for user in assignable_users if user.email
            ]

            assigned_filter = db.or_(
                EMCRequest.status == 'Assigned',
                EMCRequest.status == 'Assigned Lab Engineer'
            )
            approved_statuses = ['Approved', 'Test Plan Approved']
            status_counts = {
                'At Review': EMCRequest.query.filter_by(status='At Review').count(),
                'Approved': EMCRequest.query.filter(
                    EMCRequest.status.in_(approved_statuses)
                ).count(),
                'Assigned': EMCRequest.query.filter(assigned_filter).count(),
                # Treat admin review-ready statuses as part of in-progress work for stats
                'In Progress': EMCRequest.query.filter(
                    EMCRequest.status.in_(
                        ['In Progress', 'Test Schedule In Progress', 'Test Plan To Approve'])
                ).count()
            }

            statistics = {
                'pending_approval': status_counts.get('At Review', 0),
                'approved': status_counts.get('Approved', 0),
                'assigned': status_counts.get('Assigned', 0),
                'in_progress': status_counts.get('In Progress', 0)
            }

            return render_template('admin_approval.html',
                                   test_plans=test_plans,
                                   lab_engineers_data=lab_engineers_data,
                                   statistics=statistics,
                                   service_type_options=_get_service_type_filter_options(
                                       [plan.get('service_types', []) for plan in test_plans]
                                   ))
        except Exception as e:
            logger.error("Error loading admin approval page: %s", e)
            flash('Error loading admin approval page', 'error')
            return redirect(url_for('index'))

    @flask_app.route('/api/test-requests/<int:request_id>/admin-completed', methods=['POST'])
    @login_required
    def admin_completed(request_id):
        if current_user.role != 'admin':
            return jsonify({'success': False, 'error': 'Only admins can sign off'}), 403

        test_request = _get_request_or_404(request_id)
        data = request.get_json() or {}
        note = (data.get('note') or '').strip()

        try:
            # Update main request status to Completed
            test_request.status = 'Completed'
            test_request.reviewed_by = current_user.username
            test_request.reviewed_at = get_ist_now()

            # Update all planner entries for this request to completed (dual filter).
            planner_entries = PlannerEntry.query.filter(
                db.or_(*_planner_filters(test_request.id, test_request.tco_id))
            ).all()
            for entry in planner_entries:
                entry.status = 'completed'
                entry.updated_at = get_ist_now()

            # Append the completion note to the review comment thread if provided.
            if note:
                _append_review_comment_entry(
                    test_request=test_request,
                    comment=f'[Completion] {note}',
                    username=current_user.username,
                    role=current_user.role
                )

            db.session.commit()

            # The TCO is now complete - notify the requester and admins.
            send_completion_notification(
                test_request=test_request,
                completed_by=current_user.username
            )

            return jsonify({
                'success': True,
                'message': ('Report signed off successfully. Status updated to Completed. '
                            'Notification email sent to the requester and admins.')
            })

        except Exception as e:
            db.session.rollback()
            logger.error("Error in admin_completed: %s", e)
            return jsonify({'success': False, 'error': str(e)}), 500

    @flask_app.route('/api/admin/test-requests', methods=['GET'])
    @login_required
    def get_admin_test_requests():
        """Get test requests for admin approval page with optional show_all parameter."""
        # Check if user has admin permission
        if current_user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'You do not have permission to access this endpoint.'
            }), 403

        try:
            show_all = request.args.get('show_all', 'false').lower() == 'true'
            search_query = request.args.get('search', '').strip()
            status_filter = request.args.get('status', '').strip()
            service_type_filter = request.args.get('service_type', '').strip()
            sort_by, sort_dir = _normalize_request_sort_args(
                request.args.get('sort_by'),
                request.args.get('sort_dir')
            )

            submitted_nulls_last = db.case(
                (EMCRequest.submitted_at.is_(None), 1),
                else_=0
            )

            if show_all or status_filter:
                # Show all submitted test requests (excluding drafts)
                query = EMCRequest.query.filter(
                    EMCRequest.status != 'Draft'
                )
            else:
                # Show only approval statuses (default behavior)
                approval_statuses = [
                    'At Review',
                    'Approved',
                    'Test Plan Approved',
                    'Assigned',
                    'Assigned Lab Engineer',
                    'In Progress',
                    'Test Schedule In Progress',
                    'Test Plan To Approve',
                    'Admin Sign Off',
                    'Draft Report',
                    'Proceed Report',
                    'Datasheet Uploaded',
                    'Update plan',
                    'Report Uploaded',
                    'report_uploaded',
                    'Need More Information',
                    'Completed'
                ]
                query = EMCRequest.query.filter(
                    EMCRequest.status.in_(approval_statuses)
                )

            # Apply search filter
            if search_query:
                search_pattern = f'%{search_query}%'
                query = query.filter(
                    db.or_(
                        EMCRequest.product_name.like(search_pattern),
                        EMCRequest.tco_id.like(search_pattern),
                        EMCRequest.manufacturer.like(search_pattern),
                        EMCRequest.model_number.like(search_pattern),
                        EMCRequest.requester_name.like(search_pattern)
                    )
                )

            # Apply status filter
            if status_filter:
                status_filter_map = {
                    'pending_approval': ['At Review'],
                    'approved': ['Approved', 'Test Plan Approved'],
                    'assigned': ['Assigned', 'Assigned Lab Engineer'],
                    'in_progress': [
                        'In Progress',
                        'Test Schedule In Progress',
                        'Test Plan To Approve',
                        'Admin Sign Off',
                        'Draft Report',
                        'Proceed Report',
                        'Datasheet Uploaded',
                        'Update plan'
                    ],
                    'draft_report': ['Draft Report'],
                    'datasheet_uploaded': ['Datasheet Uploaded'],
                    'report_uploaded': ['Report Uploaded', 'report_uploaded'],
                    'admin_sign_off': ['Admin Sign Off'],
                    'need_more_info': ['Need More Information'],
                    'completed': ['Completed'],
                    'cancelled': ['Cancelled'],
                    'rejected': ['Rejected'],
                }

                filter_key = status_filter.lower()
                mapped_statuses = status_filter_map.get(filter_key)
                if mapped_statuses:
                    query = query.filter(EMCRequest.status.in_(mapped_statuses))
                else:
                    query = query.filter(
                        db.func.lower(EMCRequest.status) == filter_key
                    )

            # Order by submitted date
            query = query.order_by(
                submitted_nulls_last,
                EMCRequest.submitted_at.desc(),
                EMCRequest.created_at.desc()
            )

            # Eager load the assigned_engineer relationship to avoid lazy loading issues
            try:
                query = query.options(
                    joinedload(EMCRequest.assigned_engineer),
                    joinedload(EMCRequest.service_types)
                )
            except Exception as load_error:
                logger.warning(
                    "Could not eager load assigned_engineer relationship: %s", load_error)
                # Continue without eager loading - will use lazy loading instead

            requests = query.all()

            if service_type_filter:
                requests = [
                    test_request for test_request in requests
                    if _matches_service_type_filter(
                        _extract_service_types(test_request),
                        service_type_filter
                    )
                ]

            test_plans = []
            for test_request in requests:
                assigned_display = test_request.assigned_engineer_name
                if not assigned_display and test_request.assigned_engineer:
                    try:
                        assigned_display = test_request.assigned_engineer.username
                    except Exception:
                        assigned_display = None
                shared_review_thread = _get_combined_review_comment_thread(test_request)
                latest_shared_comment = shared_review_thread[-1] if shared_review_thread else None

                # Map DB status to UI keys/labels
                if test_request.status == 'At Review':
                    status_key = 'pending_approval'
                    status_label = test_request.status
                elif test_request.status in ('Test Schedule In Progress', 'Test Plan To Approve'):
                    status_key = 'in_progress'
                    status_label = test_request.status if test_request.status != 'Test Schedule In Progress' else 'Test Schedule In Progress'
                elif test_request.status == 'Admin Sign Off':
                    status_key = 'in_progress'
                    status_label = 'Admin Sign Off'
                elif test_request.status in ('Approved', 'Test Plan Approved'):
                    status_key = 'approved'
                    status_label = test_request.status
                elif test_request.status in ('Assigned', 'Assigned Lab Engineer'):
                    status_key = 'assigned'
                    status_label = test_request.status
                elif test_request.status == 'In Progress':
                    status_key = 'in_progress'
                    status_label = test_request.status
                elif test_request.status in ('Draft Report', 'Proceed Report', 'Datasheet Uploaded', 'Report Uploaded', 'report_uploaded', 'Update plan'):
                    status_key = 'in_progress'
                    status_label = test_request.status
                elif test_request.status == 'Need More Information':
                    status_key = 'need_more_info'
                    status_label = test_request.status
                elif test_request.status == 'Completed':
                    status_key = 'completed'
                    status_label = test_request.status
                elif test_request.status == 'Cancelled':
                    status_key = 'cancelled'
                    status_label = test_request.status
                elif test_request.status == 'Rejected':
                    status_key = 'rejected'
                    status_label = test_request.status
                else:
                    status_key = 'other'
                    status_label = test_request.status or 'Unknown'

                # Safely parse service_types JSON
                service_types = _extract_service_types(test_request)

                # Safely format dates
                submitted_date = None
                if test_request.submitted_at:
                    try:
                        submitted_date = test_request.submitted_at.isoformat()
                    except Exception:
                        submitted_date = None
                elif test_request.created_at:
                    try:
                        submitted_date = test_request.created_at.isoformat()
                    except Exception:
                        submitted_date = None

                test_plans.append({
                    'id': test_request.id,
                    'tco_id': test_request.tco_id or f'REQ-{test_request.id}',
                    'job_id': test_request.job_id,  # ADD THIS LINE
                    'job_number': test_request.job_number or test_request.job_id or '',
                    'name': test_request.product_name or 'Unnamed Product',
                    'project': test_request.manufacturer or 'N/A',
                    'status_key': status_key,
                    'status_label': status_label,
                    'requester_name': test_request.requester_name or 'Unknown',
                    'service_types': service_types,
                    'submitted_date': submitted_date,
                    'assigned_engineer_id': test_request.assigned_engineer_id,
                    'assigned_engineer_name': test_request.assigned_engineer_name,
                    'assigned_engineer_display': assigned_display,
                    'has_review_messages': len(shared_review_thread) > 0,
                    'review_comments': latest_shared_comment.get('comment') if latest_shared_comment else None,
                    'reviewed_by': latest_shared_comment.get('username') if latest_shared_comment else getattr(test_request, 'reviewed_by', None),
                    'reviewed_at': latest_shared_comment.get('created_at') if latest_shared_comment else (
                        test_request.reviewed_at.isoformat() if hasattr(test_request, 'reviewed_at') and test_request.reviewed_at else None
                    )
                })

            if sort_by:
                test_plans = _apply_request_identifier_sort(
                    test_plans,
                    sort_by,
                    sort_dir
                )

            return jsonify({
                'success': True,
                'data': test_plans
            })
        except Exception as e:
            import traceback
            logger.error("Error fetching admin test requests: %s", e)
            logger.error("Traceback: %s", traceback.format_exc())
            return jsonify({
                'success': False,
                'error': str(e),
                'traceback': traceback.format_exc() if logger.level <= logging.DEBUG else None
            }), 500

    @flask_app.errorhandler(413)
    def too_large(error):
        """Handle file size too large error."""
        return jsonify({
            'success': False,
            'error': 'File size exceeds the maximum limit of 50MB'
        }), 413

    @flask_app.route('/.well-known/appspecific/com.chrome.devtools.json')
    def chrome_devtools():
        """Handle Chrome DevTools configuration request."""
        # Return empty response for Chrome DevTools requests
        return '', 204

    @flask_app.errorhandler(404)
    def not_found(error):
        """Handle 404 errors."""
        # Ignore Chrome DevTools requests
        if request.path == '/.well-known/appspecific/com.chrome.devtools.json':
            return '', 204

        # Return JSON for API/AJAX requests
        if request.path.startswith('/api/') or request.is_json or request.headers.get('Accept') == 'application/json':
            return jsonify({
                'success': False,
                'error': 'Resource not found',
                'message': 'The requested resource was not found'
            }), 404

        # If user is not authenticated, redirect to login
        if not current_user.is_authenticated:
            flash('Please log in to access this page.', 'info')
            return redirect(url_for('auth.login'))

        # If user is authenticated, show 404 page
        return render_template('404.html'), 404

    @flask_app.errorhandler(500)
    def internal_error(error):
        """Handle 500 errors."""
        db.session.rollback()
        logger.error("Internal server error: %s", error)
        
        # Return JSON for API/AJAX requests
        if request.path.startswith('/api/') or request.is_json or request.headers.get('Accept') == 'application/json':
            return jsonify({
                'success': False,
                'error': 'Internal server error',
                'message': 'An error occurred while processing your request'
            }), 500
        
        return render_template('500.html'), 500

    return flask_app


if __name__ == '__main__':
    app = create_app()

    # Create database tables
    with app.app_context():
        db.create_all()

        # Initialize equipment database
        equipment_manager = EquipmentManager(app.config['EQUIPMENT_DATA_FILE'])
        equipment_manager.initialize_equipment_database()

    app.run(debug=True, host='0.0.0.0', port=3000)
